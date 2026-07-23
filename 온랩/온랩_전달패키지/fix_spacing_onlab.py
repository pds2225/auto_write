#!/usr/bin/env python3
"""
fix_spacing_onlab.py — 온랩 사업계획서 DOCX 줄간격 보정 (Word 렌더링용)

배경: 원본 양식(HWP)의 '줄간격 160%'는 글자크기 기준(12pt×1.6=19.2pt 고정)이지만,
Word 의 'Multiple 1.6'은 맑은 고딕의 큰 기본 행높이에 곱해져 실질 200%+ 로 부풀어
6페이지 문서가 9페이지로 늘어난다. 모든 단락의 배수 줄간격을 HWP 와 동일한
고정 행높이(글자크기 × 배수)로 변환해 원본 양식 외형을 되돌린다.

사용: python fix_spacing_onlab.py [대상.docx]   (기본: output\온랩_사업계획서_마켓게이트.docx)
"""
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

DEFAULT_TARGET = r"output\온랩_사업계획서_마켓게이트.docx"
DEFAULT_PT = 12.0        # 문서 기본 글자크기
DEFAULT_MULTIPLE = 1.6   # 스타일 기본 줄간격 배수


def _font_pt(paragraph):
    for run in paragraph.runs:
        if run.font.size:
            return run.font.size.pt
    try:
        sz = paragraph.style.font.size
        if sz:
            return sz.pt
    except Exception:
        pass
    return DEFAULT_PT


def fix_paragraph(paragraph):
    pf = paragraph.paragraph_format
    spacing = pf.line_spacing
    if isinstance(spacing, Pt) or (spacing is not None and not isinstance(spacing, float)):
        return False  # 이미 고정값
    multiple = spacing if isinstance(spacing, float) else DEFAULT_MULTIPLE
    exact = round(_font_pt(paragraph) * multiple, 1)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(exact)
    return True


def iter_all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def main():
    target = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_TARGET
    doc = Document(target)
    changed = sum(1 for p in iter_all_paragraphs(doc) if fix_paragraph(p))
    doc.save(target)
    print(f"줄간격 고정 변환: {changed}개 단락 → {target}")


if __name__ == "__main__":
    main()
