"""미래큐러스 AI인재실증형 작성본을 공고 5항목 기준으로 채점.
케이스 A 원칙: 채점만 한다(내용 생성/변형 없음)."""
import io, sys, json, datetime
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from auto_write.config import get_settings
from auto_write.services.openai_client import OpenAIService
from auto_write.services.evaluation_service import EvaluationService, EvalCriterion
from docx import Document

s = get_settings()
oai = OpenAIService(s)
print(f"[AI] available={oai.available} status={oai.status_text}")

eval_svc = EvaluationService(oai)

# 공고·양식에 명시 배점 없음 → 5항목 균등 20점(총 100). 세부기준은 양식 TABLE#2 원문 반영.
criteria = [
    EvalCriterion("문제인식", 20,
        "창업아이템의 필요성, 국내외 시장 현황 및 문제점, 문제 해결을 위한 개발 필요성",
        ["문제", "필요성", "시장", "현황", "고객"]),
    EvalCriterion("실현가능성", 20,
        "제품·서비스 개발/구체화 계획, 차별성 및 경쟁력 확보 전략, 사업비(정부지원+자부담) 집행 계획",
        ["개발", "차별성", "경쟁력", "사업비", "기술"]),
    EvalCriterion("성장전략", 20,
        "사업화 추진전략, 경쟁사 분석·목표시장 진입, 비즈니스모델(수익화), 투자유치(자금확보), 로드맵·사회적가치",
        ["사업화", "수익", "시장진입", "투자", "매출", "로드맵"]),
    EvalCriterion("팀구성", 20,
        "대표자 보유역량, 팀원 보유역량, 업무파트너 현황 및 활용 방안",
        ["대표자", "팀원", "역량", "파트너", "조직"]),
    EvalCriterion("AI인재활용계획", 20,
        "(핵심) AI 도입 필요성(기술적 근거·수치), AI를 제품·서비스에 적용하는 방법론과 협약기간(약7개월) 내 산출물·일정, AI 인재 채용계획(채용기간·인원 로드맵)",
        ["AI", "인재", "채용", "도입", "산출물", "정규직"]),
]

def docx_text(path):
    d = Document(path)
    parts = []
    for para in d.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)

DOC = r"D:\auto_write\app\tmp_quality_input\miraequrus_aijinjae_20260601.docx"
doc_text = docx_text(DOC)
print(f"[DOC] {DOC}")
print(f"[DOC] chars={len(doc_text)}")

if not oai.available:
    print("[STOP] AI 키 미연결 → score_document은 50% 폴백이라 실채점 불가. 키 설정 필요.")
    sys.exit(0)

# score_document은 내부적으로 doc_text[:6000]만 사용 → 긴 문서는 항목별 정확도 저하.
# 정확도 위해 항목별로 '전체 문서'를 직접 채점(complete_json)하여 6000자 한계를 우회.
def score_one(crit, full_text):
    sysp = (
        "당신은 정부지원사업(초기창업패키지 AI인재실증형) 심사위원입니다.\n"
        "제공된 사업계획서 전체를 아래 단일 평가항목 기준으로만 채점하세요.\n"
        "[채점 원칙]\n"
        f"1. 0~{crit.max_score}점 정수로 부여\n"
        "2. 구체적 수치·근거·실현가능성이 있으면 가점, 추상적·근거없는 서술은 감점\n"
        "3. 미작성(OO/00 등 미기재 플레이스홀더)·공란은 강하게 감점\n"
        "4. 이 항목과 무관한 내용은 점수에 반영하지 않음\n"
        "반환(JSON): {\"score\": 정수, \"strengths\": \"강점\", \"weaknesses\": \"미흡\", \"suggestion\": \"원문 기반 보완방향(없는 내용 생성 금지)\"}"
    )
    userp = json.dumps({
        "평가항목": crit.name,
        "세부기준": crit.description,
        "배점": crit.max_score,
        "사업계획서_전문": full_text[:14000],
    }, ensure_ascii=False)
    res = oai.complete_json(sysp, userp, max_tokens=1500)
    if not isinstance(res, dict):
        return None
    raw = int(res.get("score", 0))
    clamped = max(0, min(raw, crit.max_score))
    return {
        "name": crit.name, "max_score": crit.max_score, "score": clamped,
        "ratio": round(clamped / crit.max_score, 4),
        "strengths": str(res.get("strengths", "")),
        "weaknesses": str(res.get("weaknesses", "")),
        "suggestion": str(res.get("suggestion", "")),
    }

rows = []
for c in criteria:
    r = score_one(c, doc_text)
    if r is None:
        r = {"name": c.name, "max_score": c.max_score, "score": int(c.max_score*0.5),
             "ratio": 0.5, "strengths": "", "weaknesses": "채점 파싱 실패", "suggestion": "재시도"}
    rows.append(r)
    print(f"  [{r['score']:>2}/{r['max_score']}] {r['name']}  | 미흡: {r['weaknesses'][:60]}")

total = sum(r["score"] for r in rows)
maxt = sum(r["max_score"] for r in rows)
print(f"\n[TOTAL] {total}/{maxt} ({total/maxt*100:.1f}%)")
weak = [r["name"] for r in rows if r["ratio"] < 0.6]
print(f"[WEAK<60%] {weak}")

ts = "20260608"  # 고정(Date 사용 불가 환경)
out = {
    "document": DOC, "scored_at": ts, "total": total, "max_total": maxt,
    "pass_ratio": round(total/maxt, 4), "weak_criteria": weak,
    "note": "공고·양식에 명시 배점 없어 5항목 균등 20점 적용. 세부기준=양식 TABLE#2 원문. 케이스A: 채점만 수행.",
    "criteria": rows,
}
OUT_JSON = r"C:\Users\ekth3\.claude\jobs\db9a91b4\tmp\eval_result.json"
io.open(OUT_JSON, "w", encoding="utf-8").write(json.dumps(out, ensure_ascii=False, indent=2))
print(f"[SAVED] {OUT_JSON}")
