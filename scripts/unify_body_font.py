# -*- coding: utf-8 -*-
"""unify_body_font.py — DOCX '텍스트 본문'만 단일 글꼴·크기로 전역 통일.

배경: doc_quality_ops.unify_paragraph_formatting 은 '단락별 지배값' 통일이라
문서 전체를 하나의 글꼴/크기로 모으지 못한다. 이 스크립트는 본문(w:body 직계 단락)
전체를 지정 글꼴·크기(기본 맑은 고딕 11pt)로 강제 통일하되, 제목 위계는 보존한다.

통일 대상 = body 직계 단락 중:
  - 제목(□ 시작 / < ... > 표제목 / heading·title 스타일 / 지정 title 글꼴)이 아니고
  - 이미지·도형(drawing) 포함 단락이 아니며
  - 텍스트 run 이 1개 이상인 단락
표/헤더/푸터/텍스트박스는 구조상 자동 제외(doc.paragraphs = body 직계만).

안전: 색(w:color)·강조(w:b/w:u) 미변경(글꼴·크기만). 원본 덮어쓰기 금지(out==in 차단).
멱등(재실행 변경 0). 기존 doc_quality_ops 헬퍼 재사용.

사용법(PowerShell):
  $py = "C:\\...\\python.exe"
  & $py scripts\\unify_body_font.py 입력.docx 출력.docx            # 드라이런(미리보기)
  & $py scripts\\unify_body_font.py 입력.docx 출력.docx --apply     # 실제 적용
  옵션: --font "맑은 고딕" --pt 11 --title-fonts "HY헤드라인M,HY신명조"
"""
import argparse
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "app"))
from auto_write.services.doc_quality_ops import (  # noqa: E402
    _set_rpr_size,
    _set_rpr_fonts,
    _para_style_name,
    _element_has_drawing,
    _run_has_text,
)


def ea_font(run):
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        return None
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        return None
    return rf.get(qn("w:eastAsia"))


def is_title(para, title_fonts):
    txt = para.text.strip()
    if not txt:
        return False
    sname = _para_style_name(para) or ""
    if sname.startswith(("Heading", "Title")) or "제목" in sname:
        return True
    if txt.startswith("□"):
        return True
    if txt.startswith("<") and txt.endswith(">"):
        return True
    fonts = {ea_font(r) for r in para.runs if _run_has_text(r._element)}
    fonts.discard(None)
    if fonts and fonts <= title_fonts:
        return True
    return False


def unify_run(r, font, half_pt):
    rpr = r._element.find(qn("w:rPr"))
    if rpr is None:
        rpr = r._element.makeelement(qn("w:rPr"), {})
        r._element.insert(0, rpr)
    c1 = _set_rpr_size(rpr, half_pt)
    c2 = _set_rpr_fonts(rpr, font, font)
    return c1 or c2


def run(inp, outp, *, apply, font, pt, title_fonts):
    if Path(inp).resolve() == Path(outp).resolve():
        raise SystemExit("out==in 금지(원본 덮어쓰기 차단)")
    half_pt = str(int(round(pt * 2)))
    doc = Document(inp)
    targets, titles, skipped = [], [], []
    changed_runs = 0
    for para in doc.paragraphs:  # body 직계만(표/헤더/푸터/텍스트박스 제외)
        txt = para.text.strip()
        if not txt:
            continue
        if _element_has_drawing(para._p):
            skipped.append(txt[:40])
            continue
        if is_title(para, title_fonts):
            titles.append(txt[:45])
            continue
        targets.append(txt[:55])
        if apply:
            for r in para.runs:
                if _run_has_text(r._element) and unify_run(r, font, half_pt):
                    changed_runs += 1

    print(f"통일 목표: {font} {pt}pt")
    print("대상(본문) 단락:", len(targets))
    print("제외(제목):", len(titles))
    print("제외(이미지/도형):", len(skipped))
    print("--- 제외된 제목 ---")
    for t in titles:
        print("  [제목]", t)
    if apply:
        Path(outp).parent.mkdir(parents=True, exist_ok=True)
        doc.save(outp)
        print("변경 run:", changed_runs, "| 저장:", outp)
    else:
        print("(드라이런 — 변경 안 함. 실제 적용은 --apply)")
    return len(targets), changed_runs


def main(argv=None):
    ap = argparse.ArgumentParser(description="DOCX 본문 글꼴·크기 전역 통일(제목 위계 보존)")
    ap.add_argument("input")
    ap.add_argument("output")
    ap.add_argument("--apply", action="store_true", help="실제 적용(미지정 시 드라이런)")
    ap.add_argument("--font", default="맑은 고딕", help='본문 통일 글꼴(기본 "맑은 고딕")')
    ap.add_argument("--pt", type=float, default=11.0, help="본문 통일 크기 pt(기본 11)")
    ap.add_argument("--title-fonts", default="HY헤드라인M",
                    help="제목으로 간주할 글꼴 목록(콤마 구분, 기본 HY헤드라인M)")
    a = ap.parse_args(argv)
    title_fonts = {s.strip() for s in a.title_fonts.split(",") if s.strip()}
    run(a.input, a.output, apply=a.apply, font=a.font, pt=a.pt, title_fonts=title_fonts)


if __name__ == "__main__":
    main()
