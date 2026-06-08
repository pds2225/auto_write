"""미래큐러스 작성본에 실제 데이터 차트(간트·조직도)를 생성·삽입한 보완본 DOCX 생성 + 재채점.
케이스 A: 차트 데이터는 모두 문서 원문(표/본문)에 실재하는 값만 사용. 날조 없음."""
import io, sys, json, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from auto_write.config import get_settings
from auto_write.services.openai_client import OpenAIService
from auto_write.services.evaluation_service import EvalCriterion
from auto_write.services.chart_generator import gantt_chart, org_chart
from auto_write.services.chart_insert import insert_image_after_anchor

TMP = r"C:\Users\ekth3\.claude\jobs\db9a91b4\tmp"
SRC = r"D:\auto_write\app\tmp_quality_input\miraequrus_aijinjae_20260601.docx"
RESULTS = r"D:\auto_write\results"
os.makedirs(RESULTS, exist_ok=True)

# ── 1. 차트 PNG 생성 (T#6 추진일정 실데이터 / T#5 본문 팀 실데이터) ──
# 월 인덱스: 2026.10=0, 11=1, 12=2, 2027.01=3, 02=4 ; 막대 [start,end)
gantt_tasks = [
    {"name": "AI 인재 채용",            "start": 0, "end": 1},  # 26.10~26.10
    {"name": "위험도 분석 데이터 구축", "start": 0, "end": 2},  # 26.10~26.11
    {"name": "AI 위험도 분석 엔진 개발","start": 1, "end": 3},  # 26.11~26.12
    {"name": "Angel AI 플랫폼 개발",    "start": 2, "end": 4},  # 26.12~27.01
    {"name": "안전제어 칩 아키텍처 설계","start": 3, "end": 5}, # 27.01~27.02
]
gantt_png = gantt_chart(os.path.join(TMP, "chart_gantt.png"),
                        "AI 인재 활용 추진 일정 (2026.10~2027.02)", gantt_tasks)

org_nodes = [
    {"id": "ceo", "label": "대표 임상진", "parent": None},
    {"id": "vp",  "label": "부사장 김병진", "parent": "ceo"},
    {"id": "ai",  "label": "AI개발자(신규채용)", "parent": "ceo"},
]
org_png = org_chart(os.path.join(TMP, "chart_org.png"),
                    "팀 구성·역할 조직도", org_nodes)
print(f"[CHART] gantt={'OK' if gantt_png else 'FAIL'} org={'OK' if org_png else 'FAIL'}")

# ── 2. anchor 자동 탐색 + 삽입(체이닝, 원본 무변형) ──
def first_anchor(docpath, candidates):
    d = Document(docpath)
    texts = [p.text for p in d.paragraphs]
    for c in candidates:
        for t in texts:
            if c in t:
                return c
    return None

GANTT_ANCHORS = ["추진 일정", "추진일정", "사업 추진", "전체 사업단계", "로드맵", "사업추진 일정", "AI 인재 활용 계획"]
ORG_ANCHORS   = ["팀 구성", "팀구성", "대표자 및 팀원", "팀원 구성", "조직", "대표자"]

ga = first_anchor(SRC, GANTT_ANCHORS)
oa = first_anchor(SRC, ORG_ANCHORS)
print(f"[ANCHOR] gantt={ga!r} org={oa!r}")

step1 = os.path.join(RESULTS, "miraequrus_보완_step1.docx")
step2 = os.path.join(RESULTS, "miraequrus_보완_차트_20260608.docx")

ok1 = ok2 = None
if gantt_png:
    ok1 = insert_image_after_anchor(SRC, step1, ga or "사업", gantt_png,
                                    caption="[그림] AI 인재 활용 추진 일정 (출처: 본문 추진일정 표)")
    base_for_org = step1
else:
    base_for_org = SRC
if org_png:
    ok2 = insert_image_after_anchor(base_for_org, step2, oa or "팀", org_png,
                                    caption="[그림] 팀 구성·역할 조직도 (출처: 본문 팀구성)")
    FINAL = step2
else:
    # 간트만 있으면 step1이 최종
    FINAL = step1 if gantt_png else SRC
print(f"[INSERT] gantt_anchored={ok1} org_anchored={ok2}")
print(f"[FINAL] {FINAL}")

# ── 3. 재채점 (보완본 텍스트) ──
s = get_settings()
oai = OpenAIService(s)
if not oai.available:
    print("[STOP] AI 미연결 → 재채점 생략")
    sys.exit(0)

criteria = [
    EvalCriterion("문제인식", 20, "창업아이템 필요성, 국내외 시장현황·문제점, 개발 필요성", []),
    EvalCriterion("실현가능성", 20, "개발/구체화 계획, 차별성·경쟁력, 사업비 집행계획", []),
    EvalCriterion("성장전략", 20, "사업화 전략, 경쟁사분석, 수익모델, 투자유치, 로드맵·사회적가치", []),
    EvalCriterion("팀구성", 20, "대표자·팀원 역량, 업무파트너 활용", []),
    EvalCriterion("AI인재활용계획", 20, "(핵심)AI 도입 필요성·근거, 제품적용 방법론·7개월 산출물·일정, 채용 로드맵", []),
]

def docx_text(path):
    d = Document(path); parts = []
    for para in d.paragraphs:
        if para.text.strip(): parts.append(para.text)
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells): parts.append(" | ".join(cells))
    return "\n".join(parts)

def score_one(crit, full_text):
    sysp = ("당신은 정부지원사업(초기창업패키지 AI인재실증형) 심사위원입니다.\n"
            "사업계획서 전체를 아래 단일 평가항목 기준으로만 채점하세요.\n"
            f"1. 0~{crit.max_score}점 정수\n2. 구체 수치·근거·실현가능성 가점, 추상·근거없음 감점\n"
            "3. 미작성(OO/00 플레이스홀더)·공란 강한 감점\n"
            "반환(JSON): {\"score\":정수,\"strengths\":\"\",\"weaknesses\":\"\",\"suggestion\":\"\"}")
    userp = json.dumps({"평가항목": crit.name, "세부기준": crit.description,
                        "배점": crit.max_score, "사업계획서_전문": full_text[:14000]}, ensure_ascii=False)
    res = oai.complete_json(sysp, userp, max_tokens=1200)
    if not isinstance(res, dict): return None
    raw = int(res.get("score", 0)); cl = max(0, min(raw, crit.max_score))
    return {"name": crit.name, "max": crit.max_score, "score": cl,
            "weak": str(res.get("weaknesses", ""))[:80]}

txt = docx_text(FINAL)
print(f"[DOC] 보완본 chars={len(txt)}")
rows = []
for c in criteria:
    r = score_one(c, txt) or {"name": c.name, "max": c.max_score, "score": int(c.max_score*0.5), "weak": "파싱실패"}
    rows.append(r)
    print(f"  [{r['score']:>2}/{r['max']}] {r['name']} | {r['weak']}")
total = sum(r["score"] for r in rows)
print(f"\n[RESCORE] {total}/100  (이전 90/100)")
io.open(os.path.join(TMP, "rescore.json"), "w", encoding="utf-8").write(
    json.dumps({"total": total, "rows": rows, "final_doc": FINAL}, ensure_ascii=False, indent=2))
print(f"[SAVED] {os.path.join(TMP, 'rescore.json')}")
