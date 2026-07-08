# -*- coding: utf-8 -*-
"""submission_readiness.py — 제출 직전 검수 리포트 자동 생성.

읽기 전용 분석: 입력 DOCX 를 절대 수정하지 않는다.

CLI:
    python scripts\\submission_readiness.py <입력.docx> <출력리포트.md>

리포트 섹션:
    1. [확인필요]/[산출근거]/잔존빈칸 위치 목록 (가장 가까운 이전 제목 포함)
    2. 분량 실측 (문단수/표수/총문자수/페이지근사/NB 블록 제외 근사)
    3. 마스킹 위반 후보 (이메일·전화번호·대학교·성명 라벨 등, 보수적)
    4. NotebookLM 안내 블록 위치 목록
"""
from __future__ import annotations

import argparse
import io
import re
import sys
from pathlib import Path

from docx import Document
from docx.table import Table

# sys.path 에 app 추가
ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "app"))

# ─────────────────────────────────────────────
# 상수
# ─────────────────────────────────────────────

# 섹션 1: 확인 필요 패턴
CHECK_PATTERNS = [
    r"\[확인필요\]",
    r"\[산출근거\]",
    r"_{3,}",           # ___ 3개 이상
    r"\[\s*\]",         # [ ] 빈 괄호
    r"\(\s*작성\s*\)",  # (작성)
    r"〇{2,}",          # 〇〇 이상 (작성 예비란)
    r"○{2,}",
    r"\(\s*\)",         # () 빈 괄호
]
CHECK_RE = re.compile("|".join(CHECK_PATTERNS))

# NotebookLM 블록 감지: 구분선 또는 NotebookLM 키워드
NB_DIVIDER_RE = re.compile(r"─{10,}|━{10,}|-{10,}")
NB_KEYWORD_RE = re.compile(r"NotebookLM|슬라이드\s*프롬프트|슬라이드를\s*생성")

# 마스킹: 이메일, 전화번호
EMAIL_RE = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(r"0\d{1,2}[-.\s]?\d{3,4}[-.\s]?\d{4}")
UNIV_RE = re.compile(r"(대학교|대학원|전문대학)")
NAME_LABEL_RE = re.compile(r"(성명|대표자|책임자|연구책임|주관기관장)")

# 제목 패턴 (가장 가까운 이전 제목 탐지용)
HEADING_RE = re.compile(
    r"^(□|■|▶|▷|◆|◇|【|[<〈《]|[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮]|[1-9]\.|제\s*\d+\s*[장절항])"
)
CHARS_PER_PAGE = 900  # 페이지 근사 기준


# ─────────────────────────────────────────────
# 문서 순회 헬퍼 (image_apply._iter 패턴 참고)
# ─────────────────────────────────────────────

def _iter_table_cells(table: Table, top_table: Table, table_idx: int):
    """표 셀 안의 단락을 순회한다. (단락, top_table, table_idx, row_idx, col_idx)"""
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                yield para, top_table, table_idx, row_idx, col_idx
            for nested in cell.tables:
                yield from _iter_table_cells(nested, top_table, table_idx)


def iter_all_contexts(doc: Document):
    """문서 전체 단락을 순서대로 순회.

    Yields:
        (para, in_table: bool, table_idx: int | None, row: int | None, col: int | None)
    """
    for para in doc.paragraphs:
        yield para, False, None, None, None
    for t_idx, table in enumerate(doc.tables):
        for para, _, _, r, c in _iter_table_cells(table, table, t_idx):
            yield para, True, t_idx, r, c


# ─────────────────────────────────────────────
# 제목 추적
# ─────────────────────────────────────────────

def _is_heading(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    # 스타일 기반 탐지는 없으므로 텍스트 패턴으로만 판단
    return bool(HEADING_RE.match(t))


def _para_style_is_heading(para) -> bool:
    try:
        sname = (para.style.name or "").lower()
        return "heading" in sname or "title" in sname
    except Exception:
        return False


# ─────────────────────────────────────────────
# 섹션 1: 확인필요 / 잔존빈칸
# ─────────────────────────────────────────────

def scan_check_items(doc: Document) -> list[dict]:
    results = []
    last_heading = "(제목 없음)"

    for para, in_table, t_idx, r_idx, c_idx in iter_all_contexts(doc):
        text = para.text or ""
        stripped = text.strip()

        # 제목 갱신 (본문 단락일 때만 또는 표 셀이어도 제목처럼 보이면 갱신)
        if not in_table and (_is_heading(stripped) or _para_style_is_heading(para)):
            if stripped:
                last_heading = stripped[:80]
        elif not in_table and _para_style_is_heading(para) and stripped:
            last_heading = stripped[:80]

        matches = list(CHECK_RE.finditer(text))
        for m in matches:
            loc = f"표 #{t_idx + 1} ({r_idx + 1}행 {c_idx + 1}열)" if in_table else "본문"
            results.append({
                "heading": last_heading,
                "location": loc,
                "match": m.group(),
                "context": text.strip()[:100],
            })

    return results


# ─────────────────────────────────────────────
# 섹션 2: 분량 실측
# ─────────────────────────────────────────────

def measure_volume(doc: Document) -> dict:
    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)

    # 전체 문자 수 (공백 포함)
    total_chars = 0
    nb_chars = 0
    in_nb_block = False

    for para in doc.paragraphs:
        text = para.text or ""
        total_chars += len(text)
        # NB 블록 진입/탈출 탐지
        if NB_DIVIDER_RE.search(text) or NB_KEYWORD_RE.search(text):
            in_nb_block = not in_nb_block
            nb_chars += len(text)
        elif in_nb_block:
            nb_chars += len(text)

    # 표 셀 문자도 합산
    for t_idx, table in enumerate(doc.tables):
        for para, _, _, r, c in _iter_table_cells(table, table, t_idx):
            total_chars += len(para.text or "")

    page_approx = max(1, round(total_chars / CHARS_PER_PAGE, 1))
    page_excl_nb = max(1, round((total_chars - nb_chars) / CHARS_PER_PAGE, 1))

    return {
        "para_count": para_count,
        "table_count": table_count,
        "total_chars": total_chars,
        "nb_chars": nb_chars,
        "page_approx": page_approx,
        "page_excl_nb": page_excl_nb,
    }


# ─────────────────────────────────────────────
# 섹션 3: 마스킹 위반 후보
# ─────────────────────────────────────────────

def scan_masking_candidates(doc: Document) -> list[dict]:
    candidates = []

    def _add(category: str, text: str, loc: str, value: str):
        candidates.append({
            "category": category,
            "location": loc,
            "value": value[:80],
            "context": text.strip()[:100],
        })

    for para, in_table, t_idx, r_idx, c_idx in iter_all_contexts(doc):
        text = para.text or ""
        loc = f"표 #{t_idx + 1} ({r_idx + 1}행 {c_idx + 1}열)" if in_table else "본문"

        for m in EMAIL_RE.finditer(text):
            _add("이메일", text, loc, m.group())

        for m in PHONE_RE.finditer(text):
            _add("전화번호", text, loc, m.group())

        if UNIV_RE.search(text):
            _add("대학교/대학원 언급", text, loc, text.strip()[:60])

        if NAME_LABEL_RE.search(text):
            # 라벨 줄 자체 + 인접 텍스트만 후보로 올림
            _add("성명/대표자/책임자 라벨", text, loc, text.strip()[:60])

    return candidates


# ─────────────────────────────────────────────
# 섹션 4: NotebookLM 블록 탐지
# ─────────────────────────────────────────────

def scan_notebooklm_blocks(doc: Document) -> list[dict]:
    """NB 블록의 시작 위치와 직전 제목을 반환한다."""
    blocks = []
    last_heading = "(제목 없음)"
    in_block = False

    for para, in_table, t_idx, r_idx, c_idx in iter_all_contexts(doc):
        text = para.text or ""
        stripped = text.strip()

        # 제목 갱신
        if not in_table and (_is_heading(stripped) or _para_style_is_heading(para)):
            if stripped:
                last_heading = stripped[:80]

        is_divider = bool(NB_DIVIDER_RE.search(text))
        is_nb_kw = bool(NB_KEYWORD_RE.search(text))

        if not in_block and (is_divider or is_nb_kw):
            # 블록 시작
            in_block = True
            loc = f"표 #{t_idx + 1} ({r_idx + 1}행 {c_idx + 1}열)" if in_table else "본문"
            blocks.append({
                "block_no": len(blocks) + 1,
                "heading": last_heading,
                "location": loc,
                "trigger_text": stripped[:80],
            })
        elif in_block and is_divider:
            # 구분선으로 블록 종료
            in_block = False

    return blocks


# ─────────────────────────────────────────────
# 리포트 렌더링
# ─────────────────────────────────────────────

def render_report(
    input_path: str,
    check_items: list[dict],
    volume: dict,
    masking: list[dict],
    nb_blocks: list[dict],
) -> str:
    lines = []
    a = lines.append

    a(f"# 제출 준비 검수 리포트")
    a(f"")
    a(f"- **입력 문서**: `{input_path}`")
    a(f"- **생성 시각**: 자동 생성 (python submission_readiness.py)")
    a(f"- **주의**: 이 리포트는 읽기 전용 분석입니다. 원본 DOCX 는 수정되지 않았습니다.")
    a(f"")

    # ── 섹션 1 ──
    a(f"---")
    a(f"")
    a(f"## 1. [확인필요] / [산출근거] / 잔존 빈칸 위치 목록")
    a(f"")
    a(f"> 총 **{len(check_items)}건** 감지됨")
    a(f"")
    if check_items:
        a(f"| # | 직전 제목 | 위치 | 매치 패턴 | 문맥 (앞 100자) |")
        a(f"|---|-----------|------|-----------|-----------------|")
        for i, item in enumerate(check_items, 1):
            heading = item["heading"].replace("|", "｜")
            loc = item["location"].replace("|", "｜")
            match = item["match"].replace("|", "｜")
            ctx = item["context"].replace("|", "｜").replace("\n", " ")
            a(f"| {i} | {heading} | {loc} | `{match}` | {ctx} |")
    else:
        a(f"감지된 항목이 없습니다.")
    a(f"")

    # ── 섹션 2 ──
    a(f"---")
    a(f"")
    a(f"## 2. 분량 실측")
    a(f"")
    page_judgment = "✅ 15p 이내" if volume["page_approx"] <= 15 else "⚠️ 15p 초과 가능"
    excl_judgment = "✅ 15p 이내" if volume["page_excl_nb"] <= 15 else "⚠️ 15p 초과 가능"

    a(f"| 항목 | 값 |")
    a(f"|------|-----|")
    a(f"| 본문 단락 수 | {volume['para_count']:,} |")
    a(f"| 표 수 | {volume['table_count']:,} |")
    a(f"| 총 문자 수 (공백 포함) | {volume['total_chars']:,} |")
    a(f"| NotebookLM 블록 문자 수 (근사) | {volume['nb_chars']:,} |")
    a(f"| **페이지 근사** (÷{CHARS_PER_PAGE}) | **{volume['page_approx']}p** — {page_judgment} |")
    a(f"| NB 블록 제외 페이지 근사 | {volume['page_excl_nb']}p — {excl_judgment} |")
    a(f"")
    a(f"> ※ 페이지 근사는 문자수/{CHARS_PER_PAGE} 기준입니다. 실제 Word 페이지 수와 다를 수 있습니다.")
    a(f"")

    # ── 섹션 3 ──
    a(f"---")
    a(f"")
    a(f"## 3. 마스킹 위반 후보")
    a(f"")
    a(f"> 총 **{len(masking)}건** 후보 — **오탐 가능합니다. 판단은 사용자에게 있습니다.**")
    a(f">")
    a(f"> 보수적 탐지 기준: 이메일·전화번호 정규식 매치 / '대학교·대학원' 포함 줄 /")
    a(f"> '성명·대표자·책임자' 라벨 표 셀.")
    a(f"")
    if masking:
        a(f"| # | 유형 | 위치 | 후보 값 | 문맥 |")
        a(f"|---|------|------|---------|------|")
        for i, item in enumerate(masking, 1):
            cat = item["category"].replace("|", "｜")
            loc = item["location"].replace("|", "｜")
            val = item["value"].replace("|", "｜")
            ctx = item["context"].replace("|", "｜").replace("\n", " ")
            a(f"| {i} | {cat} | {loc} | `{val}` | {ctx} |")
    else:
        a(f"마스킹 위반 후보가 없습니다.")
    a(f"")

    # ── 섹션 4 ──
    a(f"---")
    a(f"")
    a(f"## 4. NotebookLM 안내 블록 위치 목록")
    a(f"")
    a(f"> 총 **{len(nb_blocks)}세트** 감지됨 (예상: 8세트)")
    a(f"")
    if nb_blocks:
        a(f"| # | 블록 번호 | 직전 제목 | 위치 | 시작 텍스트 |")
        a(f"|---|-----------|-----------|------|-------------|")
        for item in nb_blocks:
            heading = item["heading"].replace("|", "｜")
            loc = item["location"].replace("|", "｜")
            trigger = item["trigger_text"].replace("|", "｜")
            a(f"| {item['block_no']} | #{item['block_no']} | {heading} | {loc} | {trigger} |")
    else:
        a(f"NotebookLM 블록이 감지되지 않았습니다.")
    a(f"")
    a(f"---")
    a(f"")
    a(f"*리포트 끝 — 입력 DOCX 는 수정되지 않았습니다.*")

    return "\n".join(lines)


# ─────────────────────────────────────────────
# main
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="제출 직전 검수 리포트 자동 생성 (읽기 전용)"
    )
    parser.add_argument("input_docx", help="입력 DOCX 경로")
    parser.add_argument("output_md", help="출력 리포트 .md 경로")
    args = parser.parse_args()

    input_path = Path(args.input_docx)
    output_path = Path(args.output_md)

    if not input_path.exists():
        print(f"[오류] 입력 파일을 찾을 수 없습니다: {input_path}", file=sys.stderr)
        sys.exit(1)

    print(f"[submission_readiness] 입력: {input_path}")
    print(f"[submission_readiness] 출력: {output_path}")

    doc = Document(str(input_path))

    print("[1/4] [확인필요]/잔존빈칸 스캔 중...")
    check_items = scan_check_items(doc)
    print(f"      → {len(check_items)}건 발견")

    print("[2/4] 분량 실측 중...")
    volume = measure_volume(doc)
    print(f"      → 페이지 근사 {volume['page_approx']}p / NB 제외 {volume['page_excl_nb']}p")

    print("[3/4] 마스킹 위반 후보 스캔 중...")
    masking = scan_masking_candidates(doc)
    print(f"      → {len(masking)}건 후보")

    print("[4/4] NotebookLM 블록 탐지 중...")
    nb_blocks = scan_notebooklm_blocks(doc)
    print(f"      → {len(nb_blocks)}세트 감지")

    report = render_report(
        input_path=str(input_path),
        check_items=check_items,
        volume=volume,
        masking=masking,
        nb_blocks=nb_blocks,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with io.open(str(output_path), "w", encoding="utf-8") as f:
        f.write(report)

    print(f"\n[완료] 리포트 저장: {output_path}")
    print(f"  - [확인필요] 등: {len(check_items)}건")
    print(f"  - 페이지 근사: {volume['page_approx']}p")
    print(f"  - 마스킹 후보: {len(masking)}건")
    print(f"  - NotebookLM 블록: {len(nb_blocks)}세트")


if __name__ == "__main__":
    main()
