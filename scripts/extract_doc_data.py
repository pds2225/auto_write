"""미래큐러스 작성본에서 (1)미기재 플레이스홀더 위치, (2)차트화 가능 실제데이터 추출.
케이스A: 읽기만 한다(원문 무변형)."""
import io, sys, re, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

DOC = r"D:\auto_write\app\tmp_quality_input\miraequrus_aijinjae_20260601.docx"
doc = Document(DOC)

# (1) 미기재 플레이스홀더 패턴
PH = re.compile(r"00\s*(백만원|만원|원|명|개|건|년|월|일|%|시간|시)|○○|OO[가-힣A-Za-z]|[Xx][Xx]\s|△△|\(00\)|00,000|0,000|00\.00|00:00")
ph_para, ph_cell = [], []
for p in doc.paragraphs:
    t = p.text.strip()
    if t and PH.search(t):
        ph_para.append(t[:140])
for ti, tbl in enumerate(doc.tables):
    seen = set()
    for row in tbl.rows:
        for c in row.cells:
            ct = c.text.strip()
            if ct and ct not in seen and PH.search(ct):
                seen.add(ct)
                ph_cell.append({"table": ti, "text": ct[:100]})

# (2) 표 전체 덤프 (중복 병합셀 정리)
def dedup_row(cells):
    out = []
    for x in cells:
        x = x.strip().replace("\n", " ")
        if not out or out[-1] != x:
            out.append(x)
    return out

tables = []
for ti, tbl in enumerate(doc.tables):
    rows = []
    for row in tbl.rows:
        rows.append(dedup_row([c.text for c in row.cells]))
    head = " ".join(rows[0]) if rows else ""
    # 차트 후보 분류
    kind = None
    blob = head + " " + (" ".join(rows[1][:3]) if len(rows) > 1 else "")
    if re.search(r"추진|일정|기간|단계|로드맵|‘2[567]|26\.|27\.", " ".join(" ".join(r) for r in rows)):
        kind = "schedule"
    if re.search(r"팀|대표|직위|역량|담당", head):
        kind = "team" if kind is None else kind
    if re.search(r"사업비|비목|집행|정부지원|자부담|재료비|인건비", head):
        kind = "budget"
    tables.append({"idx": ti, "shape": [len(tbl.rows), len(tbl.columns)], "kind": kind, "head": head[:80], "rows": rows})

out = {
    "doc": DOC,
    "placeholder_paragraphs": ph_para,
    "placeholder_cells": ph_cell,
    "tables": tables,
}
OUT = r"C:\Users\ekth3\.claude\jobs\db9a91b4\tmp\doc_data.json"
io.open(OUT, "w", encoding="utf-8").write(json.dumps(out, ensure_ascii=False, indent=2))

print(f"[PH] 미기재 단락={len(ph_para)} 미기재 셀={len(ph_cell)}")
print("--- 미기재 셀 (사용자 작성 필요) ---")
for x in ph_cell[:25]:
    print(f"  T#{x['table']}: {x['text']}")
print("--- 차트 후보 표 ---")
for t in tables:
    if t["kind"]:
        print(f"  T#{t['idx']} [{t['kind']}] {t['shape']} :: {t['head']}")
print(f"[SAVED] {OUT}")
