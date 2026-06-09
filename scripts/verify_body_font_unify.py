# -*- coding: utf-8 -*-
"""verify_body_font_unify.py — unify_body_font 결과 검증(원본 vs 결과).

확인 항목:
  - 구조: body 단락 수 동일, 본문 텍스트 내용 불변
  - 통일: 본문 대상 run 이 100% 목표 글꼴·크기인지
  - 제목 보존: 제목 단락 폰트가 원본과 동일한지
  - 색·강조 보존: body 의 (bold, underline, color) 개수 원본=결과
  - 표 비훼손: 표 run 수·폰트 시그니처 동일

사용법:
  & $py scripts\\verify_body_font_unify.py 원본.docx 결과.docx [--font "맑은 고딕"] [--pt 11]
"""
import argparse
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "app"))
sys.path.insert(0, str(Path(__file__).resolve().parent))
from auto_write.services.doc_quality_ops import _element_has_drawing, _run_has_text  # noqa: E402
from unify_body_font import is_title  # noqa: E402


def run_font(r):
    rpr = r._element.find(qn("w:rPr"))
    ea = asc = sz = None
    if rpr is not None:
        rf = rpr.find(qn("w:rFonts"))
        if rf is not None:
            ea = rf.get(qn("w:eastAsia"))
            asc = rf.get(qn("w:ascii"))
        szel = rpr.find(qn("w:sz"))
        if szel is not None:
            sz = szel.get(qn("w:val"))
    return (ea, asc, sz)


def emph_counts(doc):
    nb = nu = nc = 0
    for p in doc.paragraphs:
        for r in p.runs:
            rpr = r._element.find(qn("w:rPr"))
            if rpr is None:
                continue
            nb += rpr.find(qn("w:b")) is not None
            nu += rpr.find(qn("w:u")) is not None
            nc += rpr.find(qn("w:color")) is not None
    return (nb, nu, nc)


def table_signature(doc):
    runs, text = 0, []
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        if _run_has_text(r._element):
                            runs += 1
                            text.append(run_font(r))
    return runs, text


def main(argv=None):
    ap = argparse.ArgumentParser()
    ap.add_argument("original")
    ap.add_argument("result")
    ap.add_argument("--font", default="맑은 고딕")
    ap.add_argument("--pt", type=float, default=11.0)
    a = ap.parse_args(argv)
    target_sz = str(int(round(a.pt * 2)))
    title_fonts = {"HY헤드라인M"}

    orig, res = Document(a.original), Document(a.result)
    op, rp = orig.paragraphs, res.paragraphs

    ok = True
    print("=== 구조 ===")
    same_len = len(op) == len(rp)
    ok &= same_len
    print(f"body 단락수 원본 {len(op)} / 결과 {len(rp)} => {'OK' if same_len else 'MISMATCH'}")
    text_mismatch = [i for i, (x, y) in enumerate(zip(op, rp)) if x.text != y.text]
    ok &= not text_mismatch
    print(f"본문 텍스트 불일치 {len(text_mismatch)} => {'OK(내용보존)' if not text_mismatch else 'CHANGED'}")

    body_runs = unified = 0
    title_total = title_preserved = 0
    bad = []
    for x, y in zip(op, rp):
        txt = y.text.strip()
        if not txt or _element_has_drawing(y._p):
            continue
        if is_title(y, title_fonts):
            title_total += 1
            fx = [run_font(r) for r in x.runs if _run_has_text(r._element)]
            fy = [run_font(r) for r in y.runs if _run_has_text(r._element)]
            if fx == fy:
                title_preserved += 1
            else:
                bad.append(("TITLE-CHANGED", txt[:30]))
            continue
        for r in y.runs:
            if not _run_has_text(r._element):
                continue
            body_runs += 1
            ea, asc, sz = run_font(r)
            if ea == a.font and asc == a.font and sz == target_sz:
                unified += 1
            else:
                bad.append((txt[:25], ea, asc, sz))

    print("=== 본문 통일 ===")
    u_ok = body_runs == unified
    ok &= u_ok
    print(f"대상 run {body_runs} / 통일 {unified} => {'OK(100%)' if u_ok else f'미통일 {body_runs-unified}'}")
    print("=== 제목 보존 ===")
    t_ok = title_total == title_preserved
    ok &= t_ok
    print(f"제목 {title_total} / 보존 {title_preserved} => {'OK' if t_ok else 'CHANGED'}")
    print("=== 색·강조 보존(body) ===")
    eo, er = emph_counts(orig), emph_counts(res)
    e_ok = eo == er
    ok &= e_ok
    print(f"(bold,underline,color) 원본 {eo} / 결과 {er} => {'OK' if e_ok else 'CHANGED'}")
    print("=== 표 비훼손 ===")
    to, tr = table_signature(orig), table_signature(res)
    tb_ok = to == tr
    ok &= tb_ok
    print(f"표 run수 원본 {to[0]} / 결과 {tr[0]} => {'OK' if tb_ok else 'CHANGED'}")
    if bad:
        print("--- 이상 항목 ---")
        for z in bad[:15]:
            print("  ", z)
    print("\n=== 종합:", "PASS" if ok else "FAIL", "===")
    return 0 if ok else 1


if __name__ == "__main__":
    raise SystemExit(main())
