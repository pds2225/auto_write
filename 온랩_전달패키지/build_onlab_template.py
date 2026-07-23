#!/usr/bin/env python3
"""
build_onlab_template.py — 2026 한난 온랩(On-Lab) 사업계획서 양식 DOCX 생성
공고 붙임1 양식(HWP) 구조를 그대로 DOCX로 옮긴다.
- 글자 12pt, 맑은 고딕, 줄간격 160% (공고 요구사항)
- 파란색(0070C0) 안내문구: inject.py 가 자동 제거
- 섹션 제목: 1x1 표(음영 바) → inject_after_keyword 의 섹션 경계로 동작
- 세부항목 제목: "N-M." 패턴 → _SECTION_HEADING_RE 경계와 일치
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0x00, 0x70, 0xC0)
FONT = "맑은 고딕"


def set_font(run, size=12, bold=False, color=None, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font)


def para(doc, text="", size=12, bold=False, color=None, align=None, spacing=1.6):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = spacing
    p.paragraph_format.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, color=color)
    return p


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def cell_write(cell, text, size=12, bold=False, align=None, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.3
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)


def heading_bar(doc, title):
    """섹션 제목을 1x1 음영 표로 생성 — 주입 시 섹션 경계 역할."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_write(t.rows[0].cells[0], title, size=14, bold=True)
    shade_cell(t.rows[0].cells[0], "DEEAF6")
    return t


def guide(doc, lines):
    """파란 안내문구(자동 제거 대상)."""
    for ln in lines:
        para(doc, ln, size=11, color=BLUE, spacing=1.3)


def subhead(doc, text):
    para(doc, text, size=12, bold=True)
    para(doc, "  ㅇ ", size=12)  # placeholder — 주입 시 교체됨


def main():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.line_spacing = 1.6
    pf.space_after = Pt(0)

    # 여백 2cm (5페이지 내외 분량 관리)
    from docx.shared import Cm
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2)
        sec.right_margin = Cm(2)

    # ── 제목 ──
    para(doc, "2026 한난 온랩(On-Lab) 사업계획서", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    guide(doc, [
        "☞ 사업계획서는 5페이지 내외로 작성",
        "☞ 그림, 표 활용 가능",
        "☞ 글자 크기 12pt, 폰트 '맑은 고딕' 또는 '굴림' 사용, 줄간격 160%",
        "☞ 항목별 세부 항목에 해당하는 내용을 작성하고 필요에 따라 생략/추가 가능",
    ])


    # ── 개요 표 (표0) ──
    t = doc.add_table(rows=5, cols=2)
    t.style = "Table Grid"
    labels = ["팀  명", "대표자", "참가 분야", "아이템명", "아이템 개요(50자)"]
    for i, lab in enumerate(labels):
        cell_write(t.rows[i].cells[0], lab, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(t.rows[i].cells[0], "F2F2F2")
        cell_write(t.rows[i].cells[1], "")


    # ── 1. 문제인식 (표1) ──
    heading_bar(doc, "1. 문제인식")
    guide(doc, [
        "❑ 고객·시장·사회에서 발생하고 있는 문제, 불편, 비효율 또는 해결 필요성",
        "❑ 해당 문제가 발생하게 된 배경과 기존 제품·서비스 또는 해결방식의 한계",
        "❑ 창업을 통해 해결하고자 하는 목표와 기대효과",
    ])
    subhead(doc, "1-1. 고객·시장의 문제와 해결 필요성")
    subhead(doc, "1-2. 문제의 배경과 기존 해결방식의 한계")
    subhead(doc, "1-3. 창업 목표와 기대효과")


    # ── 2. 실현가능성 (표2) ──
    heading_bar(doc, "2. 실현가능성")
    guide(doc, [
        "❑ 제품·서비스의 개요, 핵심 기능, 제공 방식 및 고객에게 전달되는 주요 가치",
        "❑ 제시한 문제를 해결하는 방식과 기존 제품·서비스 대비 차별성",
        "❑ 현재 준비 수준, 개발·구현 계획, 생산·운영 인프라 등 실제 실행 가능성",
    ])
    subhead(doc, "2-1. 제품·서비스 개요 및 핵심 기능")
    subhead(doc, "2-2. 문제 해결방식과 차별성")
    subhead(doc, "2-3. 준비 수준 및 개발·구현 계획")


    # ── 3. 성장전략 (표3) ──
    heading_bar(doc, "3. 성장전략")
    guide(doc, [
        "❑ 목표시장, 주요 고객, 시장 규모 및 수요 발생 근거",
        "❑ 경쟁사 또는 유사 제품·서비스 대비 경쟁우위와 시장 진입전략",
        "❑ 수익모델, 판매·마케팅 전략, 자금 활용계획 및 향후 확장 가능성",
    ])
    subhead(doc, "3-1. 목표시장 및 시장 규모")
    subhead(doc, "3-2. 경쟁우위 및 시장 진입전략")
    subhead(doc, "3-3. 수익모델 및 자금 활용계획")


    # ── 4. 팀 구성 (표4) + 팀 표(표5) ──
    heading_bar(doc, "4. 팀 구성")
    guide(doc, [
        "❑ 대표자 및 팀원의 주요 경력, 전문성, 역할 분담 등 사업 추진 역량",
        "❑ 보유 기술, 네트워크, 장비, 시설, 지식재산권 등 활용 가능한 자원",
        "❑ 부족한 역량 보완계획, 외부 협력 또는 인력 확보 계획",
    ])
    subhead(doc, "4-1. 대표자 및 팀원의 사업 추진 역량")
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    for j, h in enumerate(["구분", "역할", "주요 경력 및 보유 역량"]):
        cell_write(t.rows[0].cells[j], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(t.rows[0].cells[j], "F2F2F2")
    subhead(doc, "4-2. 활용 가능한 자원(기술·네트워크·지식재산권)")
    subhead(doc, "4-3. 부족 역량 보완 및 인력 확보 계획")


    # ── 5. 사업화 추진 계획 (표6) + 일정 표(표7) ──
    heading_bar(doc, "5. 사업화 추진 계획")
    guide(doc, [
        "❑ 시제품 개발, 고객 검증, 제품·서비스 고도화, 출시 등 단계별 사업화 추진계획",
        "❑ 생산·운영, 유통·판매, 제휴, 마케팅 등 실제 사업 실행을 위한 세부 추진방안",
        "❑ 향후 제품·서비스 개선, 고객군 확대, 시장 확장, 추가 수익모델 등 발전방안",
    ])
    subhead(doc, "5-1. 단계별 사업화 추진계획")
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    for j, h in enumerate(["단계", "추진 내용", "추진 기간", "세부 내용"]):
        cell_write(t.rows[0].cells[j], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(t.rows[0].cells[j], "F2F2F2")
    subhead(doc, "5-2. 사업 실행 세부 추진방안")
    subhead(doc, "5-3. 향후 발전방안")


    # ── 6. 자유주제 (표8) ──
    heading_bar(doc, "6. 자유주제")
    guide(doc, ["❑ (기타) 위 제시된 항목 외 추가하고자 하는 내용 기재"])
    subhead(doc, "6-1. 온랩(On-Lab) 참여 목표 및 활용 계획")

    out = "templates/온랩_사업계획서_양식.docx"
    doc.save(out)
    print(f"저장 완료: {out}")


if __name__ == "__main__":
    main()
