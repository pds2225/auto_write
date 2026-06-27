"""
사업계획서 자동 작성기 v2
- 작성완료본 → 빈 양식 자동 채우기
- 품질 자동 채점 (길이·데이터·구조·키워드)
- 미흡 섹션 AI 재작성 (Claude) + Before/After 수락/거절
- DOCX 미리보기 + 다운로드
"""

import streamlit as st
import io, re, os
from difflib import SequenceMatcher
from docx import Document
from core.criteria_mapper import map_heading, UNKNOWN
from core.ai_writer import build_non_empty_text_content

try:
    import mammoth; HAS_MAMMOTH = True
except ImportError:
    HAS_MAMMOTH = False

try:
    import anthropic; HAS_CLAUDE = True
except ImportError:
    HAS_CLAUDE = False

# ── 페이지 설정 ──────────────────────────────────────────
st.set_page_config(page_title="사업계획서 자동 작성기", page_icon="📝", layout="wide")
st.markdown("""
<style>
section[data-testid="stSidebar"] { background: #1b3d6e; }
section[data-testid="stSidebar"] * { color: #fff !important; }
.stProgress > div > div { background: #1b3d6e; }
.badge { display:inline-block; padding:2px 10px; border-radius:10px;
         font-size:12px; font-weight:700; margin-left:8px; }
.good  { background:#e6f4ec; color:#2a7a4b; }
.warn  { background:#fff3d6; color:#a06000; }
.bad   { background:#fde8e6; color:#c0392b; }
.card  { border:1px solid #dde4ee; border-radius:8px; padding:16px;
         margin-bottom:12px; background:#fbfcfd; }
.card-warn { border:1px solid #f0ad4e; background:#fffbf0; }
.card-bad  { border:1px solid #d9534f; background:#fef5f4; }
</style>
""", unsafe_allow_html=True)

# ── 세션 초기화 ──────────────────────────────────────────
for k, v in {
    "step": 1,
    "src_sections": {},
    "tmpl_bytes": None,
    "tmpl_sections": {},
    "matches": {},
    "filled": {},
    "scores": {},
    "rewrites": {},          # {heading: "AI 재작성 텍스트"}
    "accepted": set(),       # 수락된 heading set
}.items():
    if k not in st.session_state:
        st.session_state[k] = v

S = st.session_state  # 짧은 별칭

# ── 유틸 함수 ────────────────────────────────────────────
HEADING_RE = [
    r'^[0-9]+\.',  r'^[0-9]+-[0-9]+', r'^[가-힣]\.',
    r'^\([0-9]+\)', r'^[①②③④⑤⑥]',
    r'^[■◆▶]\s',   r'^제\s*[0-9]+\s*[장절항]',
]

# 템플릿의 안내문/유의사항 줄 — 매핑 대상 섹션에서 제외
ANNOTATION_RE = re.compile(
    r'^\s*[\*※▷◎]\s*|^\s*\(유의사항\)|재창업사업화\s*사업비|기본\s*제공한\s*질문|'
    r'진단결과\s*도출|필요\s*시\s*항목|필요\s*시\s*칸|필요분야만\s*작성',
    re.IGNORECASE,
)

def _is_heading(para):
    t = para.text.strip()
    if not t or len(t) > 80: return False
    if "heading" in para.style.name.lower(): return True
    for p in HEADING_RE:
        if re.match(p, t): return True
    runs_bold = [r.bold for r in para.runs if r.text.strip()]
    return bool(runs_bold) and all(runs_bold) and len(t) < 60


def extract_sections(docx_bytes):
    """DOCX → {헤딩: 본문} dict"""
    doc = Document(io.BytesIO(docx_bytes))
    secs, cur_h, cur_lines = {}, "서두", []

    def flush():
        txt = "\n".join(cur_lines).strip()
        if txt: secs[cur_h] = txt

    for p in doc.paragraphs:
        if _is_heading(p):
            flush(); cur_h = p.text.strip(); cur_lines = []
        else:
            t = p.text.strip()
            if t: cur_lines.append(t)
    flush()

    # 표 → "[표N]" 섹션으로 추가
    for i, tbl in enumerate(doc.tables):
        rows = []
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells): rows.append(" | ".join(cells))
        if rows: secs[f"[표{i+1}]"] = "\n".join(rows)

    return secs


def extract_tmpl_sections(docx_bytes):
    """템플릿 DOCX → {헤딩: ""} dict (채워야 할 빈칸 목록)"""
    doc = Document(io.BytesIO(docx_bytes))
    secs = {}
    for p in doc.paragraphs:
        if _is_heading(p):
            t = p.text.strip()
            if not ANNOTATION_RE.match(t):  # 안내문/유의사항 줄 제외
                secs[t] = ""
    return secs


def _sim(a, b): return SequenceMatcher(None, a, b).ratio()

def _kw(a, b):
    wa = set(re.findall(r'[가-힣a-zA-Z]{2,}', a))
    wb = set(re.findall(r'[가-힣a-zA-Z]{2,}', b))
    if not wa or not wb: return 0.0
    return len(wa & wb) / max(len(wa), len(wb))


def auto_match(src, tmpl):
    """소스 섹션 → 템플릿 섹션 자동 매핑 {tmpl_h: (src_h, score)}

    1단계: criteria_mapper 카테고리 일치 후보끼리 비교 (정밀)
    2단계: UNKNOWN이거나 후보 없으면 전체 소스 대상 폴백 (관대)
    """
    src_keys = [sh for sh in src if not sh.startswith("[표")]
    src_cat  = {sh: map_heading(sh).category for sh in src_keys}
    tmpl_cat = {th: map_heading(th).category for th in tmpl}

    used, out = set(), {}
    for th in tmpl:
        th_cat = tmpl_cat[th]

        # 1단계: 같은 카테고리 후보
        if th_cat != UNKNOWN:
            candidates = [sh for sh in src_keys
                          if sh not in used and src_cat[sh] == th_cat]
        else:
            candidates = []

        # 2단계 폴백: 카테고리 불명확하거나 후보 없음
        if not candidates:
            candidates = [sh for sh in src_keys if sh not in used]

        best, best_s = None, 0.0
        for sh in candidates:
            s = _sim(th, sh) * 0.55 + _kw(th, sh) * 0.45
            if s > best_s:
                best_s = s; best = sh

        # 카테고리 일치 시 역치 완화 (0.15), 폴백 시 기존 역치 유지 (0.25)
        threshold = 0.15 if (th_cat != UNKNOWN and best and
                             src_cat.get(best) == th_cat) else 0.25
        if best and best_s >= threshold:
            out[th] = (best, round(best_s, 2)); used.add(best)
        else:
            out[th] = (None, 0.0)
    return out


def score_section(text):
    """섹션 품질 점수 0~100"""
    if not text or not text.strip(): return 0
    s, t = 0, text.strip()
    n = len(t)
    # 길이 30점
    if n >= 500: s += 30
    elif n >= 300: s += 22
    elif n >= 150: s += 14
    elif n >= 80:  s += 8
    elif n >= 30:  s += 4
    # 숫자·데이터 25점
    if re.search(r'\d+[%억만원개사명년월]', t): s += 25
    elif re.search(r'\d{2,}', t): s += 12
    # 구조 20점
    b = len(re.findall(r'[\n◦•\-]\s*[가-힣a-zA-Z]', t))
    nl = t.count('\n')
    if b >= 4 or nl >= 5: s += 20
    elif b >= 2 or nl >= 3: s += 12
    elif b >= 1 or nl >= 1: s += 6
    # 키워드 25점
    kws = ['문제','해결','시장','고객','수익','경쟁','기술','계획','목표','전략','차별','성장']
    s += min(25, sum(1 for k in kws if k in t) * 4)
    return min(100, s)


def build_docx(tmpl_bytes, filled):
    """템플릿 DOCX에 filled 내용을 주입해 bytes 반환"""
    doc = Document(io.BytesIO(tmpl_bytes))
    cur_h, inserted = None, set()
    for para in doc.paragraphs:
        if _is_heading(para):
            cur_h = para.text.strip()
        elif cur_h and cur_h not in inserted and cur_h in filled and filled[cur_h]:
            # 첫 번째 빈/안내 단락에 삽입
            for run in para.runs: run.text = ""
            if para.runs:
                para.runs[0].text = filled[cur_h]
            else:
                para.add_run(filled[cur_h])
            inserted.add(cur_h)
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


def docx_to_html(docx_bytes):
    if not HAS_MAMMOTH:
        return "<p style='color:gray'>mammoth 미설치 — 텍스트 미리보기만 가능합니다.</p>"
    try:
        r = mammoth.convert_to_html(io.BytesIO(docx_bytes))
        return r.value
    except Exception as e:
        return f"<p style='color:red'>미리보기 오류: {e}</p>"


def ai_rewrite(heading, content, api_key):
    """Claude API로 섹션 재작성"""
    if not HAS_CLAUDE:
        return "anthropic 패키지가 설치되지 않았습니다."
    if not api_key:
        return "API 키를 사이드바에 입력해 주세요."
    try:
        client = anthropic.Anthropic(api_key=api_key)
        user_prompt = (
            f"섹션명: {heading}\n\n현재 내용:\n{content or '(없음)'}\n\n"
            "개선된 내용만 출력하세요 (섹션명 반복 불필요)."
        )
        content_blocks = build_non_empty_text_content(user_prompt)
        msg = client.messages.create(
            model="claude-opus-4-5",
            max_tokens=1500,
            system=(
                "당신은 정부지원사업 사업계획서 전문 작성자입니다. "
                "주어진 섹션을 평가 기준에 맞게 개선합니다:\n"
                "- 구체적 수치·데이터 포함\n"
                "- 문제→원인→해결 논리 구조\n"
                "- ◦ - 등 불릿으로 가독성 확보\n"
                "- 300자 이상, 과장 없이, 한국어로"
            ),
            messages=[{"role": "user", "content": content_blocks}],
        )
        return msg.content[0].text.strip()
    except Exception as e:
        return f"[오류] {e}"


# ── 사이드바 ─────────────────────────────────────────────
with st.sidebar:
    st.markdown("## 📝 사업계획서\n## 자동 작성기")
    st.markdown("---")
    steps = ["① 파일 업로드", "② 섹션 매핑", "③ 미리보기·편집", "④ 품질채점·재작성", "⑤ 다운로드"]
    chosen = st.radio("단계", steps, index=S.step - 1, key="sidebar_step")
    S.step = steps.index(chosen) + 1
    st.markdown("---")
    st.markdown("#### Claude API 키")
    st.caption("재작성 기능에 사용 (선택)")
    api_key_in = st.text_input("API Key", type="password", key="api_key_field",
                               placeholder="sk-ant-...")
    if api_key_in: st.caption("✅ 키 입력됨")

api_key = api_key_in or os.environ.get("ANTHROPIC_API_KEY", "")

# ── STEP 1: 파일 업로드 ──────────────────────────────────
if S.step == 1:
    st.title("① 파일 업로드")
    st.caption("두 파일을 모두 업로드하면 다음 단계로 진행할 수 있습니다.")
    c1, c2 = st.columns(2)

    with c1:
        st.subheader("📄 작성완료본")
        st.caption("내용을 추출할 완성된 사업계획서")
        src_file = st.file_uploader("완성본 DOCX", type=["docx"], key="f_src")
        if src_file:
            S.src_sections = extract_sections(src_file.read())
            st.success(f"{len(S.src_sections)}개 섹션 추출 완료")
            with st.expander("추출 섹션 목록"):
                for k, v in S.src_sections.items():
                    st.markdown(f"- **{k}**: {v[:60]}{'…' if len(v)>60 else ''}")

    with c2:
        st.subheader("📋 빈 양식 (템플릿)")
        st.caption("내용을 채워 넣을 빈 사업계획서 양식")
        tmpl_file = st.file_uploader("양식 DOCX", type=["docx"], key="f_tmpl")
        if tmpl_file:
            raw = tmpl_file.read()
            S.tmpl_bytes = raw
            S.tmpl_sections = extract_tmpl_sections(raw)
            st.success(f"{len(S.tmpl_sections)}개 섹션 확인")
            with st.expander("양식 섹션 목록"):
                for k in S.tmpl_sections:
                    st.markdown(f"- {k}")

    st.divider()
    if S.src_sections and S.tmpl_bytes:
        if st.button("➡ 섹션 매핑으로", type="primary", use_container_width=True):
            S.matches = auto_match(S.src_sections, S.tmpl_sections)
            S.filled = {
                th: (S.src_sections.get(sh, "") if sh else "")
                for th, (sh, _) in S.matches.items()
            }
            S.scores = {th: score_section(cnt) for th, cnt in S.filled.items()}
            S.step = 2; st.rerun()


# ── STEP 2: 섹션 매핑 ────────────────────────────────────
elif S.step == 2:
    st.title("② 섹션 매핑 확인")
    st.caption("자동 매핑이 틀렸으면 수동으로 변경하세요.")

    if not S.matches:
        st.warning("먼저 ① 파일 업로드를 완료하세요."); st.stop()

    src_opts = ["(없음)"] + [k for k in S.src_sections if not k.startswith("[표")]
    new_matches, new_filled = {}, {}

    for th, (sh, sc) in S.matches.items():
        icon = "🟢" if sc >= 0.6 else ("🟡" if sc >= 0.3 else "🔴")
        th_cat = map_heading(th).category
        cat_tag = f"[{th_cat}] " if th_cat != UNKNOWN else ""
        with st.expander(f"{icon} {cat_tag}**{th}** → {sh or '없음'} ({sc:.0%})"):
            cur_idx = src_opts.index(sh) if sh in src_opts else 0
            sel = st.selectbox("소스 섹션 선택", src_opts, index=cur_idx, key=f"m_{th}")
            if sel == "(없음)":
                new_matches[th] = (None, 0.0); new_filled[th] = ""
            else:
                new_matches[th] = (sel, sc)
                new_filled[th] = S.src_sections.get(sel, "")
                st.text_area("소스 내용 미리보기", new_filled[th][:400], height=90,
                             disabled=True, key=f"pv_{th}")

    mapped = sum(1 for v in new_matches.values() if v[0])
    st.info(f"매핑 현황: **{mapped} / {len(new_matches)}** 섹션")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("⬅ 이전", use_container_width=True):
            S.step = 1; st.rerun()
    with col2:
        if st.button("➡ 미리보기·편집으로", type="primary", use_container_width=True):
            S.matches = new_matches; S.filled = new_filled
            S.scores = {th: score_section(cnt) for th, cnt in new_filled.items()}
            S.step = 3; st.rerun()


# ── STEP 3: 미리보기·편집 ────────────────────────────────
elif S.step == 3:
    st.title("③ 미리보기 & 편집")

    tab_a, tab_b = st.tabs(["📄 DOCX 미리보기", "✏️ 섹션별 편집"])

    with tab_a:
        doc_bytes = build_docx(S.tmpl_bytes, S.filled)
        html = docx_to_html(doc_bytes)
        st.markdown(
            f'<div style="border:1px solid #ccc;border-radius:6px;padding:24px 32px;'
            f'background:#fff;line-height:1.85;font-size:14px">{html}</div>',
            unsafe_allow_html=True,
        )

    with tab_b:
        updated = {}
        for th, cnt in S.filled.items():
            q = S.scores.get(th, 0)
            badge = ("good" if q >= 70 else "warn" if q >= 40 else "bad")
            label = ("양호" if q >= 70 else "보통" if q >= 40 else "미흡")
            st.markdown(
                f'**{th}** <span class="badge {badge}">{q}점 {label}</span>',
                unsafe_allow_html=True,
            )
            updated[th] = st.text_area("", cnt, height=110, key=f"ed_{th}",
                                        label_visibility="collapsed")
        if st.button("✅ 편집 반영", type="primary", use_container_width=True):
            S.filled = updated
            S.scores = {th: score_section(c) for th, c in updated.items()}
            st.success("반영됨"); st.rerun()

    st.divider()
    col1, col2 = st.columns(2)
    with col1:
        if st.button("⬅ 이전", use_container_width=True): S.step = 2; st.rerun()
    with col2:
        if st.button("➡ 품질채점·재작성으로", type="primary", use_container_width=True):
            S.step = 4; st.rerun()


# ── STEP 4: 품질채점·재작성 ──────────────────────────────
elif S.step == 4:
    st.title("④ 품질 채점 & AI 재작성")

    # 요약 메트릭
    scores = list(S.scores.values())
    if scores:
        avg = sum(scores) / len(scores)
        g = sum(1 for s in scores if s >= 70)
        w = sum(1 for s in scores if 40 <= s < 70)
        b = sum(1 for s in scores if s < 40)
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("평균 점수", f"{avg:.0f}점")
        m2.metric("🟢 양호(70+)", f"{g}개")
        m3.metric("🟡 보통(40~)", f"{w}개")
        m4.metric("🔴 미흡(~40)", f"{b}개", delta=f"재작성 권고 {b}개" if b else None,
                  delta_color="inverse")

    if not api_key:
        st.warning("🔑 AI 재작성을 사용하려면 사이드바에 Claude API 키를 입력하세요. "
                   "수동 편집은 API 키 없이도 가능합니다.")

    st.divider()

    # 낮은 점수 먼저 정렬
    sorted_h = sorted(S.filled.keys(), key=lambda h: S.scores.get(h, 0))

    for h in sorted_h:
        content = S.filled.get(h, "")
        q = S.scores.get(h, 0)
        is_accepted = h in S.accepted
        has_rewrite = h in S.rewrites

        # 카드 CSS 클래스
        card_cls = "card" if q >= 70 else ("card-warn" if q >= 40 else "card-bad")
        badge_cls = "good" if q >= 70 else ("warn" if q >= 40 else "bad")
        badge_txt = ("✓ 양호" if q >= 70 else "△ 보통" if q >= 40 else "✗ 미흡 — 재작성 권고")

        st.markdown(
            f'<div class="{card_cls}">'
            f'<strong>{h}</strong>'
            f'<span class="badge {badge_cls}">{q}점 {badge_txt}</span>'
            f'</div>',
            unsafe_allow_html=True,
        )

        col_txt, col_btn = st.columns([4, 1])

        with col_txt:
            # 수락된 재작성본 vs 원본 표시
            show = S.rewrites[h] if is_accepted else content
            label = "✅ 재작성 적용됨" if is_accepted else "현재 내용"
            edited = st.text_area(label, show, height=110, key=f"qa_{h}")
            # 편집 내용 즉시 반영
            if edited != show:
                S.filled[h] = edited
                S.scores[h] = score_section(edited)

        with col_btn:
            st.write("")  # 세로 정렬용
            if is_accepted:
                # 수락 상태 → 원본 복원 버튼
                if st.button("↩ 원본", key=f"rst_{h}", use_container_width=True):
                    S.accepted.discard(h)
                    S.filled[h] = content  # 원본 복원
                    S.scores[h] = score_section(content)
                    st.rerun()
            else:
                # 재작성 버튼
                disabled = not api_key
                tooltip = "API 키 필요" if disabled else f"{h} 재작성"
                if st.button("🤖 AI 재작성", key=f"rw_{h}",
                             disabled=disabled, help=tooltip,
                             use_container_width=True):
                    with st.spinner("재작성 중..."):
                        result = ai_rewrite(h, content, api_key)
                    S.rewrites[h] = result
                    st.rerun()

        # ── Before / After 비교 패널 ──
        if has_rewrite and not is_accepted:
            rw_text = S.rewrites[h]
            with st.expander("📊 재작성 결과 비교 — 수락 또는 거절", expanded=True):
                bc1, bc2 = st.columns(2)
                with bc1:
                    st.markdown("**원본**")
                    st.text_area("", content, height=160, disabled=True, key=f"b_orig_{h}")
                with bc2:
                    st.markdown("**AI 재작성 (편집 가능)**")
                    edited_rw = st.text_area("", rw_text, height=160, key=f"b_rw_{h}")

                qa1, qa2 = st.columns(2)
                with qa1:
                    if st.button("✅ 수락 — 재작성본 적용", key=f"acc_{h}",
                                 type="primary", use_container_width=True):
                        final = edited_rw  # 편집된 버전 사용
                        S.rewrites[h] = final
                        S.filled[h] = final
                        S.scores[h] = score_section(final)
                        S.accepted.add(h)
                        st.rerun()
                with qa2:
                    if st.button("❌ 거절 — 원본 유지", key=f"rej_{h}",
                                 use_container_width=True):
                        del S.rewrites[h]
                        st.rerun()

        st.write("")  # 섹션 간격

    st.divider()
    col1, col2 = st.columns(2)
    with col1:
        if st.button("⬅ 이전", use_container_width=True): S.step = 3; st.rerun()
    with col2:
        if st.button("➡ 다운로드", type="primary", use_container_width=True):
            S.step = 5; st.rerun()


# ── STEP 5: 다운로드 ─────────────────────────────────────
elif S.step == 5:
    st.title("⑤ 최종 다운로드")

    final_bytes = build_docx(S.tmpl_bytes, S.filled)

    st.success("🎉 사업계획서 작성 완료!")

    # 최종 품질 요약
    scores = S.scores
    if scores:
        avg = sum(scores.values()) / len(scores)
        accepted_count = len(S.accepted)
        st.markdown(
            f"**최종 평균 품질: {avg:.0f}점** "
            f"| AI 재작성 적용: **{accepted_count}개** 섹션"
        )

    # 미리보기
    st.markdown("### 최종 미리보기")
    html = docx_to_html(final_bytes)
    st.markdown(
        f'<div style="border:1px solid #ccc;border-radius:6px;padding:24px 32px;'
        f'background:#fff;line-height:1.85;font-size:14px;max-height:600px;'
        f'overflow-y:auto">{html}</div>',
        unsafe_allow_html=True,
    )

    st.divider()

    # 다운로드
    st.download_button(
        label="📥 완성된 사업계획서 다운로드 (.docx)",
        data=final_bytes,
        file_name="사업계획서_완성본.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary",
        use_container_width=True,
    )

    st.divider()
    if st.button("🔄 새 파일로 다시 시작", use_container_width=True):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()
