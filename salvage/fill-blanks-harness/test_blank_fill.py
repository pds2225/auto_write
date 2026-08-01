"""빈칸 채우기 공통 골격(blank_fill_service) 회귀 테스트.

검증 항목
  1) [필수] 목차 보존: 채움 전/후 섹션 제목 문단 텍스트 집합이 완전히 동일.
  2) plan 채움 동작: plan 값이 빈 칸에 실제로 들어감.
  3) 원본 덮어쓰기 차단: output==input 이면 ValueError.
  4) 멱등성: 같은 plan 으로 두 번 실행해도 본문 중복·손상 없음(문단 수 안정).
  5) plan 없음: 에러 없이 안전 복사 + report.filled=False.

실행: (app 디렉토리 기준)  python -m pytest tests/test_blank_fill.py -q
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from auto_write.services.blank_fill_service import fill_blanks


# --------------------------------------------------------------------------- 목차(섹션 제목) 상수
SECTION_TITLES = [
    "1. 기업명 식별표",
    "2. 사업 개요",
    "3. AI 인재활용 계획",
    "4. 추진 일정",
]


def _make_sample(path: Path) -> None:
    """섹션 제목 + 그 아래 빈 문단/빈 셀 + 더미값 표를 가진 샘플 DOCX 생성."""
    d = Document()

    # 1) 기업명 식별표(제목 + 빈칸 표) — identity 채움 대상
    d.add_paragraph(SECTION_TITLES[0])
    t_id = d.add_table(rows=2, cols=2)
    t_id.cell(0, 0).text = "기업명"
    t_id.cell(0, 1).text = ""  # 빈칸
    t_id.cell(1, 0).text = "사업자등록번호"
    t_id.cell(1, 1).text = ""  # 빈칸

    # 2) 사업 개요(제목 + 빈 문단)
    d.add_paragraph(SECTION_TITLES[1])
    d.add_paragraph("")  # 빈 문단

    # 3) AI 인재활용 계획(제목 + 가이드 앵커 문단) — paragraph_fills 대상
    d.add_paragraph(SECTION_TITLES[2])
    d.add_paragraph("ㅇ AI 도입 필요성 작성")

    # 4) 추진 일정(제목 + 더미값 표) — residual 스캔/정리 대상
    d.add_paragraph(SECTION_TITLES[3])
    t_dummy = d.add_table(rows=2, cols=2)
    t_dummy.cell(0, 0).text = "구분"
    t_dummy.cell(0, 1).text = "내용"
    t_dummy.cell(1, 0).text = "1"
    t_dummy.cell(1, 1).text = "OOOOO"  # 더미값(RESIDUAL_RE 매칭)

    d.save(str(path))


def _section_titles_in(path: Path) -> set[str]:
    """문서의 섹션 제목(목차) 문단 텍스트 집합을 반환."""
    doc = Document(str(path))
    titles = {p.text.strip() for p in doc.paragraphs if p.text.strip()}
    return {t for t in titles if t in set(SECTION_TITLES)}


def _sample_plan() -> dict:
    """빈칸을 채울 최소 plan."""
    return {
        "identity": {
            "기업명": "테스트기업",
            "사업자등록번호": "123-45-67890",
        },
        "paragraph_fills": [
            {
                "anchor": "ㅇ AI 도입 필요성 작성",
                "lines": [
                    "ㅇ AI 도입 필요성",
                    "· 첫 번째 근거",
                    "· 두 번째 근거",
                ],
            }
        ],
    }


# --------------------------------------------------------------------------- 1
def test_section_titles_preserved(tmp_path: Path):
    """[필수] 채움 전/후 섹션 제목(목차) 집합이 완전히 동일해야 한다."""
    src = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    _make_sample(src)

    before = _section_titles_in(src)
    assert before == set(SECTION_TITLES)  # 사전 조건: 4개 제목 모두 존재

    fill_blanks(src, out, plan=_sample_plan(), backup_root=tmp_path / "bk")

    after = _section_titles_in(out)
    # 목차 제목 집합이 완전히 동일 — 삭제·변경 없음
    assert after == before


# --------------------------------------------------------------------------- 2
def test_plan_fills_blanks(tmp_path: Path):
    """plan 으로 준 값이 실제로 빈 칸/앵커에 들어가야 한다."""
    src = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    _make_sample(src)

    report = fill_blanks(src, out, plan=_sample_plan(), backup_root=tmp_path / "bk")

    assert report["plan_applied"] is True
    assert report["filled"] is True
    assert report["identity_filled"] >= 1

    doc = Document(str(out))
    all_cell_text = " ".join(c.text for tb in doc.tables for c in tb._cells)
    assert "테스트기업" in all_cell_text
    assert "123-45-67890" in all_cell_text

    para_text = "\n".join(p.text for p in doc.paragraphs)
    assert "첫 번째 근거" in para_text  # 앵커 뒤 본문 삽입 확인


# --------------------------------------------------------------------------- 3
def test_output_never_overwrites_input(tmp_path: Path):
    """출력==입력이면 ValueError 로 차단해야 한다."""
    src = tmp_path / "same.docx"
    _make_sample(src)
    with pytest.raises(ValueError):
        fill_blanks(src, src, plan=_sample_plan(), backup_root=tmp_path / "bk")


# --------------------------------------------------------------------------- 4
def test_idempotent_no_body_duplication(tmp_path: Path):
    """같은 plan 으로 원본을 두 번 채워도 본문 문단 수가 안정적이어야 한다.

    재실행은 항상 '원본'을 입력으로 1회 적용 → 결과 문단 수가 동일해야 한다.
    """
    src = tmp_path / "in.docx"
    out1 = tmp_path / "out1.docx"
    out2 = tmp_path / "out2.docx"
    _make_sample(src)

    fill_blanks(src, out1, plan=_sample_plan(), backup_root=tmp_path / "bk")
    fill_blanks(src, out2, plan=_sample_plan(), backup_root=tmp_path / "bk")

    n1 = len(Document(str(out1)).paragraphs)
    n2 = len(Document(str(out2)).paragraphs)
    assert n1 == n2  # 본문 중복 없음

    # 목차도 양쪽 모두 보존
    assert _section_titles_in(out1) == set(SECTION_TITLES)
    assert _section_titles_in(out2) == set(SECTION_TITLES)


# --------------------------------------------------------------------------- 5
def test_no_plan_safe_copy(tmp_path: Path):
    """plan 이 없으면 에러 없이 안전 복사하고 report.filled=False 여야 한다."""
    src = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    _make_sample(src)

    report = fill_blanks(src, out, plan=None, backup_root=tmp_path / "bk")

    assert out.exists()
    assert report["filled"] is False
    assert report["plan_applied"] is False
    assert report.get("reason") == "no plan provided"
    # 안전 복사이므로 목차 보존
    assert _section_titles_in(out) == set(SECTION_TITLES)
    # 백업이 실제로 만들어졌는지
    assert Path(report["backup_dir"]).exists()
