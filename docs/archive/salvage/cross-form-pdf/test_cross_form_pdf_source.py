"""test_cross_form_pdf_source.py — cross_form_autofill 의 PDF(플랫 텍스트) 소스 지원.

완성본 PDF 를 소스로 받아 '짧은 식별항목'(기업명·대표자·연락처 등)의 값을
타깃 빈양식에 전사하는 경로를 검증한다. 에세이 본문(긴 서술)은 _MAX_VALUE_LEN
폐기 가드로 차단한다.

안전 원칙: 오매칭<빈칸 · 날조0 · 원본미수정 · 회귀0.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from auto_write.services import cross_form_autofill as cfa
from auto_write.services.cross_form_autofill import (
    _cluster_members,
    _extract_source_from_text,
    _find_label_spans,
    autofill_from_source,
    match_fields,
)


# --- _find_label_spans / _extract_source_from_text 단위 --------------------------

def _anchor(*labels: str) -> dict[str, str]:
    """표시라벨 목록 → {정규화키: 표시라벨} 앵커."""
    return {cfa._key(lbl): lbl for lbl in labels}


def test_extract_multi_anchor_one_line():
    """한 줄 멀티앵커: '기업(예정)명 폴파이프 대표자 홍길동' → 각 값 정확 추출."""
    anchor = _anchor("기업명", "대표자")
    text = "기업(예정)명 폴파이프 대표자 홍길동"
    fields, originals = _extract_source_from_text(text, anchor)
    assert fields.get(cfa._key("기업명")) == "폴파이프"
    assert fields.get(cfa._key("대표자")) == "홍길동"
    # 원본 라벨맵이 채워진다(match_fields 의 H3 괄호 대조에 쓰임)
    assert cfa._key("기업명") in originals


def test_value_that_is_a_label_is_discarded():
    """값 자리에 또 다른 앵커 라벨이 오면 그 필드는 빈칸(오매칭<빈칸)."""
    anchor = _anchor("기업명", "대표자")
    # '기업명' 다음에 값 없이 바로 '대표자' 라벨 → 기업명 값='대표자'(라벨) → 폐기
    text = "기업명 대표자 홍길동"
    fields, _ = _extract_source_from_text(text, anchor)
    assert cfa._key("기업명") not in fields  # 값이 라벨이라 폐기
    assert fields.get(cfa._key("대표자")) == "홍길동"


def test_oversized_value_discarded():
    """_MAX_VALUE_LEN 초과 서술(에세이 본문)은 추출 안 됨."""
    anchor = _anchor("사업명")
    long_value = "가" * (cfa._MAX_VALUE_LEN + 5)
    text = f"사업명 {long_value}"
    fields, _ = _extract_source_from_text(text, anchor)
    assert cfa._key("사업명") not in fields


def test_blank_value_discarded():
    """빈칸/밑줄/점 값은 폐기(_is_fill_blank)."""
    anchor = _anchor("연락처")
    for blank in ("____", "......", "   ", "- - -"):
        fields, _ = _extract_source_from_text(f"연락처 {blank}", anchor)
        assert cfa._key("연락처") not in fields, f"blank={blank!r}"


def test_first_wins_for_same_key():
    """같은 정규화키는 처음 값 유지(first-wins)."""
    anchor = _anchor("대표자")
    text = "대표자 홍길동\n대표자 김철수"
    fields, _ = _extract_source_from_text(text, anchor)
    assert fields.get(cfa._key("대표자")) == "홍길동"


def test_value_extracted_from_multiline():
    """여러 줄에 흩어진 식별항목을 각 줄에서 추출."""
    anchor = _anchor("기업명", "대표자", "연락처")
    text = "기업명 폴파이프\n대표자 홍길동\n연락처 010-1234-5678"
    fields, _ = _extract_source_from_text(text, anchor)
    assert fields.get(cfa._key("기업명")) == "폴파이프"
    assert fields.get(cfa._key("대표자")) == "홍길동"
    assert fields.get(cfa._key("연락처")) == "010-1234-5678"


# --- _cluster_members ------------------------------------------------------------

def test_cluster_members_returns_synonyms():
    """동의어 멤버 반환 확인."""
    members = _cluster_members(cfa._key("대표자"))
    norm_members = {nk for nk, _disp in members}
    assert cfa._key("성명") in norm_members
    assert cfa._key("대표이사") in norm_members
    assert cfa._key("대표자") in norm_members


def test_cluster_members_empty_for_unknown():
    """클러스터 미등록 라벨은 빈 리스트."""
    assert _cluster_members(cfa._key("존재하지않는라벨XYZ")) == []


# --- 함수 단위 통합: 텍스트 추출 + 매칭 + 채움 (실 PDF 의존 회피) --------------------

def test_text_extract_match_and_fill(tmp_path: Path):
    """_extract_source_from_text → match_fields → 채움 을 함수 단위로 묶어 검증.

    실제 PDF 생성을 피하고, 추출된 소스 dict 가 기존 match/채움 경로에서
    high 전사되는지 확인한다(소스 dict 형태 동일성 증명).
    """
    # (a) 타깃 빈양식 .docx — 식별항목 표
    tgt = tmp_path / "target.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "기업명"
    table.cell(0, 1).text = ""
    table.cell(1, 0).text = "대표자"
    table.cell(1, 1).text = ""
    doc.save(str(tgt))

    # (b) 소스: PDF 텍스트를 직접 구성(앵커 방식 추출)
    targets = cfa.find_target_fields(tgt)
    anchor: dict[str, str] = {}
    for t in targets:
        anchor[t["normalized"]] = t["orig_label"]
        for nk, disp in _cluster_members(t["normalized"]):
            anchor.setdefault(nk, disp)

    src_text = "기업명 폴파이프 대표자 홍길동"
    src_fields, src_originals = _extract_source_from_text(src_text, anchor)

    matches = match_fields(src_fields, targets, src_originals)
    high = [m for m in matches if m.confidence == "high" and m.value]
    vals = {m.normalized: m.value for m in high}
    assert vals.get(cfa._key("기업명")) == "폴파이프"
    assert vals.get(cfa._key("대표자")) == "홍길동"


# --- 확장자 정책: 소스 .pdf 허용 / 타깃 .pdf 거부 -------------------------------

def test_target_pdf_rejected(tmp_path: Path):
    """타깃이 .pdf 면 ok=False + notes(기존처럼 거부 — 값을 써넣을 수 없음)."""
    # 소스(docx) 준비
    src = tmp_path / "source.docx"
    sdoc = Document()
    st = sdoc.add_table(rows=1, cols=2)
    st.cell(0, 0).text = "기업명"
    st.cell(0, 1).text = "폴파이프"
    sdoc.save(str(src))
    # 타깃을 .pdf 로 지정(파일 내용 무관 — 확장자에서 거부돼야 함)
    tgt_pdf = tmp_path / "target.pdf"
    tgt_pdf.write_bytes(b"%PDF-1.4 dummy")
    out = tmp_path / "out.docx"

    report = autofill_from_source(src, tgt_pdf, out)
    assert report.ok is False
    assert any("타깃" in n for n in report.notes)


def test_source_pdf_not_rejected_by_extension(tmp_path: Path, monkeypatch):
    """소스가 .pdf 면 확장자 정책에서 거부되지 않는다(_SOURCE_EXTS 허용).

    실제 PDF 텍스트 추출은 monkeypatch 로 식별항목 텍스트를 주입해 회피한다.
    """
    # 타깃 빈양식 .docx
    tgt = tmp_path / "target.docx"
    tdoc = Document()
    tbl = tdoc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "기업명"
    tbl.cell(0, 1).text = ""
    tbl.cell(1, 0).text = "대표자"
    tbl.cell(1, 1).text = ""
    tdoc.save(str(tgt))

    # 소스 .pdf (내용 무관 — extract_text 를 monkeypatch)
    src_pdf = tmp_path / "source.pdf"
    src_pdf.write_bytes(b"%PDF-1.4 dummy")
    out = tmp_path / "out.docx"

    # doc_text_extract.extract_text 가 식별항목 텍스트를 반환하도록 가짜 주입
    import auto_write.services.doc_text_extract as dte

    def _fake_extract_text(path):
        return "기업명 폴파이프 대표자 홍길동", []

    monkeypatch.setattr(dte, "extract_text", _fake_extract_text)

    report = autofill_from_source(src_pdf, tgt, out)
    # 확장자 거부 메시지가 '소스'에 대해 나오면 안 된다
    assert not any("소스 비지원 확장자" in n for n in report.notes), report.notes
    # 식별항목 전사 성공
    assert report.transcribed >= 2, report.as_dict()
    assert out.exists()
    # PDF 소스 사용 노트 1줄
    assert any("PDF 소스" in n for n in report.notes), report.notes


def test_empty_pdf_text_is_safe(tmp_path: Path, monkeypatch):
    """텍스트가 비면(스캔 PDF) 채움 0 + 안내 노트(안전, 크래시 없음)."""
    tgt = tmp_path / "target.docx"
    tdoc = Document()
    tbl = tdoc.add_table(rows=1, cols=2)
    tbl.cell(0, 0).text = "기업명"
    tbl.cell(0, 1).text = ""
    tdoc.save(str(tgt))

    src_pdf = tmp_path / "scan.pdf"
    src_pdf.write_bytes(b"%PDF-1.4 dummy")
    out = tmp_path / "out.docx"

    import auto_write.services.doc_text_extract as dte
    monkeypatch.setattr(dte, "extract_text", lambda path: ("", ["스캔본"]))

    report = autofill_from_source(src_pdf, tgt, out)
    assert report.transcribed == 0
    assert any("스캔" in n or "추출 결과 없음" in n for n in report.notes), report.notes


def test_source_not_modified(tmp_path: Path, monkeypatch):
    """소스·타깃 원본 미수정 확인(PDF 소스 경로)."""
    tgt = tmp_path / "target.docx"
    tdoc = Document()
    tbl = tdoc.add_table(rows=1, cols=2)
    tbl.cell(0, 0).text = "기업명"
    tbl.cell(0, 1).text = ""
    tdoc.save(str(tgt))
    tgt_bytes_before = tgt.read_bytes()

    src_pdf = tmp_path / "source.pdf"
    src_pdf.write_bytes(b"%PDF-1.4 dummy")
    src_bytes_before = src_pdf.read_bytes()
    out = tmp_path / "out.docx"

    import auto_write.services.doc_text_extract as dte
    monkeypatch.setattr(dte, "extract_text",
                        lambda path: ("기업명 폴파이프", []))

    autofill_from_source(src_pdf, tgt, out)
    assert tgt.read_bytes() == tgt_bytes_before
    assert src_pdf.read_bytes() == src_bytes_before
