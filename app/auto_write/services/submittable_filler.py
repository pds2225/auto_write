"""렌더된 사업계획서 DOCX를 '제출 가능' 수준으로 마감하는 후처리기.

배경:
- 템플릿 분석기는 일반현황 식별표(기업명/개업연월일 등)나 병합이 심한 표를
  profile.tables에 포함하지 못하는 경우가 있어, 렌더 단계에서 비어 있거나
  더미값(OOOOO, 00백만원, < 예시 > 등)이 그대로 남는다.
- 본 모듈은 *렌더가 끝난* output.docx 를 입력으로 받아, 실제 데이터(plan)를
  '행 라벨 + 논리적 셀' 기준으로 채워 넣고 남은 더미/가이드 마커를 정리한다.

설계 원칙(클린 아키텍처):
- 도메인 데이터(plan)는 외부에서 주입한다. 본 모듈에는 특정 기업 정보를
  하드코딩하지 않는다(제너릭 엔진).
- 병합 셀 안전성: python-docx 의 row.cells 는 병합된 셀을 중복 반환하므로,
  내부 w:tc 동일성으로 중복 제거한 '논리적 셀' 순서로만 기록한다.
- 기존 렌더 결과(PSST 본문 등)는 건드리지 않고, 비어있거나 더미인 곳만 보정한다.
"""
from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.text.paragraph import Paragraph

from .docx_ops import set_cell_text


class SubmittableFiller:
    # 잔여 더미 토큰(정리/검증용)
    RESIDUAL_RE = re.compile(
        r"(O{3,}|○{2,}|0{2,}\s*\(명\)|0{2,}\s*백만원|O{2,}\.O{2,}|"
        r"OOO-OO-OOOOO|OO학|OO기술|예정\(['’]?00|완료\(['’]?00|‘00|’00\.0|…|\.{3,})"
    )
    # 비어있는 가이드 표시(예: < 예시 : ... >) — 채워지지 않은 안내문구
    GUIDE_LINE_RE = re.compile(r"^<.*>$")

    def __init__(self, plan: dict[str, Any]):
        self.plan = plan

    # ----------------------------------------------------------------- utils
    @staticmethod
    def _norm(text: str) -> str:
        return re.sub(r"\s+", " ", str(text or "")).strip()

    @staticmethod
    def _key(text: str) -> str:
        """라벨 비교용: 공백/괄호 보조설명 제거한 핵심 키."""
        t = re.sub(r"\(.*?\)", "", str(text or ""))
        return re.sub(r"\s+", "", t).strip()

    @staticmethod
    def _logical_cells(row) -> list:
        """병합 중복을 제거한 논리적 셀(열) 목록을 순서대로 반환."""
        seen: set[int] = set()
        result = []
        for cell in row.cells:
            ident = id(cell._tc)
            if ident in seen:
                continue
            seen.add(ident)
            result.append(cell)
        return result

    @staticmethod
    def _set_paragraph_text(paragraph, text: str) -> None:
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(text)

    # ----------------------------------------------------------------- steps
    def _fill_identity(self, doc, report) -> None:
        identity = self.plan.get("identity") or {}
        if not identity:
            return
        key_map = {self._key(k): v for k, v in identity.items()}
        for table in doc.tables:
            cells_text = " ".join(c.text for c in table._cells)
            if "기업명" not in cells_text and "사업자등록번호" not in cells_text:
                continue
            for row in table.rows:
                logical = self._logical_cells(row)
                for i, cell in enumerate(logical):
                    ck = self._key(cell.text)
                    if not ck or i + 1 >= len(logical):
                        continue
                    for label_key, value in key_map.items():
                        if label_key and ck.startswith(label_key):
                            set_cell_text(logical[i + 1], str(value))
                            report["identity_filled"] += 1
                            break

    def _find_overview_table(self, doc):
        for table in doc.tables:
            joined = self._key(" ".join(c.text for c in table._cells))
            if "아이템개요" in joined and "문제인식" in joined and "팀구성" in joined:
                return table
        return None

    def _fill_overview(self, doc, report) -> None:
        overview = self.plan.get("overview") or {}
        if not overview:
            return
        table = self._find_overview_table(doc)
        if table is None:
            report["notes"].append("개요 요약표를 찾지 못함")
            return
        key_map = {self._key(k): v for k, v in overview.items()}
        for row in table.rows:
            logical = self._logical_cells(row)
            if not logical:
                continue
            head = self._key(logical[0].text)
            # r0: 명칭 | 값 | 범주 | 값
            if head.startswith("명칭") and len(logical) >= 4:
                if "명칭" in key_map:
                    set_cell_text(logical[1], str(key_map["명칭"]))
                    report["overview_filled"] += 1
                if "범주" in key_map:
                    set_cell_text(logical[3], str(key_map["범주"]))
                    report["overview_filled"] += 1
                continue
            for label_key, value in key_map.items():
                if label_key in ("명칭", "범주"):
                    continue
                if head.startswith(label_key) and len(logical) >= 2:
                    set_cell_text(logical[1], str(value))
                    report["overview_filled"] += 1
                    break

    @staticmethod
    def _iter_all_paragraphs(doc):
        """본문 직계 단락 + 모든 표 셀(중첩 표 포함) 단락을 순회한다.

        정부 사업계획서 양식은 표 기반이라 앵커가 표 셀 안에 있는 경우가 많다.
        본문(doc.paragraphs)만 보면 표 셀 앵커를 놓쳐 채움이 누락된다.
        """
        for paragraph in doc.paragraphs:
            yield paragraph

        def _walk(table):
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph
                    for nested in cell.tables:
                        yield from _walk(nested)

        for table in doc.tables:
            yield from _walk(table)

    def _apply_paragraph_fills(self, doc, report) -> None:
        """본문 가이드 문구(예: 'ㅇ ... 세부내용 작성')를 실제 내용 + 하위 불릿으로 교체.

        profile에 섹션으로 잡히지 않아 도구가 채우지 못한 본문(예: 5. AI 인재활용 계획)을
        앵커 문단 기준으로 채운다. 첫 줄은 앵커 문단에 기록하고, 나머지 줄은 동일 서식의
        새 문단으로 앵커 바로 뒤에 순서대로 삽입한다.
        """
        for fill in self.plan.get("paragraph_fills") or []:
            anchor = self._norm(fill.get("anchor", ""))
            lines = [str(line) for line in (fill.get("lines") or []) if str(line).strip()]
            if not anchor or not lines:
                continue
            target = None
            for paragraph in self._iter_all_paragraphs(doc):
                if self._norm(paragraph.text) == anchor:
                    target = paragraph
                    break
            if target is None:
                report["notes"].append(f"앵커 미발견(본문/표): {anchor[:30]}")
                continue
            self._set_paragraph_text(target, lines[0])
            cursor = target
            for line in lines[1:]:
                new_p = deepcopy(cursor._p)
                cursor._p.addnext(new_p)
                cursor = Paragraph(new_p, cursor._parent)
                self._set_paragraph_text(cursor, line)
            report["paragraphs_filled"] += 1

    def _apply_row_rewrites(self, doc, report) -> None:
        for rw in self.plan.get("row_rewrites") or []:
            ti = int(rw.get("table_index", -1))
            ri = int(rw.get("row", -1))
            cols = rw.get("cols") or []
            if ti < 0 or ti >= len(doc.tables):
                report["notes"].append(f"row_rewrite 표 범위초과 ti={ti}")
                continue
            table = doc.tables[ti]
            if ri < 0 or ri >= len(table.rows):
                report["notes"].append(f"row_rewrite 행 범위초과 ti={ti} ri={ri}")
                continue
            logical = self._logical_cells(table.rows[ri])
            for ci, value in enumerate(cols):
                if value is None or ci >= len(logical):
                    continue
                set_cell_text(logical[ci], str(value))
                report["rows_rewritten"] += 1

    def _apply_replacements(self, doc, report) -> None:
        repl = {self._norm(k): v for k, v in (self.plan.get("replacements") or {}).items()}
        prefix = self.plan.get("replacements_prefix") or {}
        prefix_norm = {self._norm(k): v for k, v in prefix.items()}

        def handle(text: str):
            norm = self._norm(text)
            if not norm:
                return None
            if norm in repl:
                return str(repl[norm])
            for pre, val in prefix_norm.items():
                if pre and norm.startswith(pre):
                    return str(val)
            return None

        # tables
        for table in doc.tables:
            for cell in table._cells:
                new = handle(cell.text)
                if new is not None:
                    set_cell_text(cell, new)
                    report["replacements"] += 1
        # paragraphs (본문 가이드 문구)
        for paragraph in doc.paragraphs:
            new = handle(paragraph.text)
            if new is not None:
                self._set_paragraph_text(paragraph, new)
                report["replacements"] += 1

    def _blank_guides(self, doc, report) -> None:
        """채워지지 않은 '< ... >' 안내문구 문단/셀을 제거(빈칸 처리)."""
        for paragraph in doc.paragraphs:
            if self.GUIDE_LINE_RE.match(self._norm(paragraph.text)):
                self._set_paragraph_text(paragraph, "")
                report["guides_blanked"] += 1
        for table in doc.tables:
            for cell in table._cells:
                if self.GUIDE_LINE_RE.match(self._norm(cell.text)):
                    set_cell_text(cell, "")
                    report["guides_blanked"] += 1

    def _blank_residual(self, doc, report) -> None:
        """남은 더미 토큰을 빈칸으로 정리(허위 수치 생성 금지). 표 셀만 대상."""
        for table in doc.tables:
            for cell in table._cells:
                txt = cell.text
                if not txt.strip():
                    continue
                if self.RESIDUAL_RE.search(txt):
                    cleaned = self.RESIDUAL_RE.sub("", txt)
                    cleaned = re.sub(r"[ \t]{2,}", " ", cleaned).strip(" /·,")
                    if cleaned != txt.strip():
                        set_cell_text(cell, cleaned)
                        report["residual_cleaned"] += 1

    def scan_residual(self, doc) -> list[str]:
        found: list[str] = []
        for i, p in enumerate(doc.paragraphs):
            if self.RESIDUAL_RE.search(p.text):
                found.append(f"para[{i}] {self._norm(p.text)[:60]}")
        for ti, table in enumerate(doc.tables):
            for cell in table._cells:
                if self.RESIDUAL_RE.search(cell.text):
                    found.append(f"table[{ti}] {self._norm(cell.text)[:50]}")
        # 중복 축약
        uniq: list[str] = []
        seen: set[str] = set()
        for f in found:
            if f in seen:
                continue
            seen.add(f)
            uniq.append(f)
        return uniq

    # ----------------------------------------------------------------- entry
    def finalize(self, input_docx: Path, output_docx: Path) -> dict[str, Any]:
        report: dict[str, Any] = {
            "identity_filled": 0,
            "overview_filled": 0,
            "rows_rewritten": 0,
            "replacements": 0,
            "residual_cleaned": 0,
            "guides_blanked": 0,
            "paragraphs_filled": 0,
            "notes": [],
        }
        doc = Document(str(input_docx))
        self._fill_identity(doc, report)
        self._fill_overview(doc, report)
        self._apply_paragraph_fills(doc, report)
        self._apply_row_rewrites(doc, report)
        self._apply_replacements(doc, report)
        self._blank_guides(doc, report)
        self._blank_residual(doc, report)
        output_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_docx))
        # 검증 스캔(저장본 기준)
        report["residual_remaining"] = self.scan_residual(Document(str(output_docx)))
        return report
