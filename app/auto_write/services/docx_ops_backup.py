from __future__ import annotations

import copy
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from lxml import etree

GUIDE_MARKER_RE = re.compile(r"(※|<[^>]+>|기재|작성요령|작성방법|예시|OOO|○○○)")


def _normalize_match_text(text: str) -> str:
    normalized = re.sub(r"\s+", " ", text or "").strip()
    normalized = normalized.replace("–", "-").replace("—", "-").replace("‑", "-")
    for _ in range(2):
        if len(normalized) < 16 or len(normalized) % 2 != 0:
            break
        half = len(normalized) // 2
        left = normalized[:half].strip()
        right = normalized[half:].strip()
        if left and left == right:
            normalized = left
            continue
        break
    return normalized


def _match_anchor(text: str, anchor_text: str) -> bool:
    anchor = _normalize_match_text(anchor_text)
    candidate = _normalize_match_text(text)
    if not anchor or not candidate:
        return False
    if anchor in candidate:
        return True
    anchor_compact = anchor.replace(" ", "")
    candidate_compact = candidate.replace(" ", "")
    if anchor_compact and anchor_compact in candidate_compact:
        return True
    return False


def _paragraph_text(paragraph: etree._Element) -> str:
    return "".join(node.text or "" for node in paragraph.iter(qn("w:t"))).strip()


def _iter_body_paragraphs(doc: Document) -> list[etree._Element]:
    body = doc.element.body
    return [element for element in body if element.tag == qn("w:p")]


def _write_para(para, text: str) -> None:
    runs = para.findall(qn("w:r"))
    if runs:
        for run in runs:
            for text_node in run.findall(qn("w:t")):
                text_node.text = ""
        text_nodes = runs[0].findall(qn("w:t"))
        if text_nodes:
            text_nodes[0].text = text
        else:
            element = etree.SubElement(runs[0], qn("w:t"))
            element.text = text
    else:
        run = etree.SubElement(para, qn("w:r"))
        element = etree.SubElement(run, qn("w:t"))
        element.text = text


def set_cell_text(cell, text: str) -> None:
    tc = cell._tc
    paragraphs = tc.findall(qn("w:p"))
    if not paragraphs:
        return
    lines = text.splitlines() or [text]
    _write_para(paragraphs[0], lines[0])
    for extra in paragraphs[1:]:
        tc.remove(extra)
    ref = paragraphs[0]
    for line in lines[1:]:
        new_paragraph = copy.deepcopy(ref)
        for run in new_paragraph.findall(qn("w:r")):
            for text_node in run.findall(qn("w:t")):
                text_node.text = ""
        _write_para(new_paragraph, line)
        tc.append(new_paragraph)


def find_paragraph(doc: Document, anchor_text: str):
    for paragraph in _iter_body_paragraphs(doc):
        text = _paragraph_text(paragraph)
        if _match_anchor(text, anchor_text):
            return paragraph
    return None


def find_cell(doc: Document, anchor_text: str):
    anchor = _normalize_match_text(anchor_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text or ""
                if _match_anchor(text, anchor):
                    return cell
    return None


def insert_text_after_paragraph(doc: Document, anchor_text: str, text: str, allow_create_paragraph: bool = False) -> bool:
    paragraph = find_paragraph(doc, anchor_text)
    if paragraph is None:
        cell = find_cell(doc, anchor_text)
        if cell is None:
            return False
        existing = cell.text.strip()
        if existing and not GUIDE_MARKER_RE.search(existing):
            return False
        set_cell_text(cell, text)
        return True

    body = doc.element.body
    body_items = list(body)
    try:
        para_index = body_items.index(paragraph)
    except ValueError:
        para_index = -1
    if para_index >= 0:
        lines = text.splitlines() or [text]
        inserted = 0
        for next_index in range(para_index + 1, min(para_index + 12, len(body_items))):
            candidate = body_items[next_index]
            if candidate.tag != qn("w:p"):
                continue
            if _paragraph_text(candidate):
                continue
            _write_para(candidate, lines[inserted])
            inserted += 1
            if inserted >= len(lines):
                return True
        if inserted > 0:
            return True
    if not allow_create_paragraph or para_index < 0:
        return False
    reference = paragraph
    for offset, line in enumerate(text.splitlines() or [text]):
        new_paragraph = copy.deepcopy(reference)
        for run in new_paragraph.findall(qn("w:r")):
            for text_node in run.findall(qn("w:t")):
                text_node.text = ""
        _write_para(new_paragraph, line)
        body.insert(para_index + 1 + offset, new_paragraph)
    return True


def _get_image_xml(doc: Document, image_path: Path, width_cm: float):
    temp_paragraph = OxmlElement("w:p")
    doc.element.body.append(temp_paragraph)
    from docx.text.paragraph import Paragraph

    paragraph = Paragraph(temp_paragraph, doc)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    drawings = temp_paragraph.findall(".//" + qn("w:drawing"))
    drawing_xml = copy.deepcopy(drawings[0]) if drawings else None
    doc.element.body.remove(temp_paragraph)
    return drawing_xml


def insert_image_in_cell(doc: Document, table_index: int, row: int, cell_index: int, image_path: Path, width_cm: float = 12.0) -> bool:
    try:
        table = doc.tables[table_index]
        cell = table.rows[row].cells[cell_index]
    except Exception:
        return False
    drawing_xml = _get_image_xml(doc, image_path, width_cm)
    if drawing_xml is None:
        return False
    tc = cell._tc
    paragraphs = tc.findall(qn("w:p"))
    if not paragraphs:
        return False
    target = paragraphs[0]
    for paragraph in paragraphs[1:]:
        if _paragraph_text(paragraph):
            continue
        tc.remove(paragraph)
    for run in list(target.findall(qn("w:r"))):
        target.remove(run)
    new_run = OxmlElement("w:r")
    new_run.append(drawing_xml)
    target.append(new_run)
    return True


def _set_paragraph_center(paragraph: etree._Element) -> None:
    ppr = paragraph.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph.insert(0, ppr)
    jc = ppr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        ppr.append(jc)
    jc.set(qn("w:val"), "center")


def insert_image_after_paragraph(
    doc: Document,
    anchor_text: str,
    image_path: Path,
    width_cm: float = 14.0,
    insert_offset: int = 1,
    allow_create_paragraph: bool = False,
) -> bool:
    paragraph = find_paragraph(doc, anchor_text)
    if paragraph is None:
        return False
    drawing_xml = _get_image_xml(doc, image_path, width_cm)
    if drawing_xml is None:
        return False
    body = doc.element.body
    try:
        base_index = list(body).index(paragraph)
    except ValueError:
        return False

    body_items = list(body)
    for next_index in range(base_index + 1, min(base_index + 8, len(body_items))):
        candidate = body_items[next_index]
        if candidate.tag != qn("w:p"):
            continue
        if _paragraph_text(candidate):
            continue
        for run in list(candidate.findall(qn("w:r"))):
            candidate.remove(run)
        image_run = OxmlElement("w:r")
        image_run.append(copy.deepcopy(drawing_xml))
        candidate.append(image_run)
        _set_paragraph_center(candidate)
        return True

    if not allow_create_paragraph:
        return False
    image_paragraph = OxmlElement("w:p")
    _set_paragraph_center(image_paragraph)
    image_run = OxmlElement("w:r")
    image_run.append(drawing_xml)
    image_paragraph.append(image_run)
    body.insert(base_index + max(1, insert_offset), image_paragraph)
    return True
