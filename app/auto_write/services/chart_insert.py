"""chart_insert.py

``chart_generator`` 가 만든 PNG 차트를 **python-docx 로 DOCX 본문에 삽입**하는
헬퍼. 본 모듈은 원본 DOCX 를 절대 덮어쓰지 않으며, 항상 별도 출력 경로(out_docx)
로 복사본을 저장한다.

핵심 함수:
  - ``insert_image_after_anchor`` : anchor_text 를 포함하는 첫 단락 바로 뒤에
    이미지(+선택 캡션)를 삽입한다. anchor 를 찾으면 그 자리 뒤에 넣고 True,
    못 찾으면 문서 끝에 추가하고 False 를 반환한다.
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.shared import Inches


def insert_image_after_anchor(
    in_docx: str,
    out_docx: str,
    anchor_text: str,
    png_path: str,
    caption: str = "",
    width_inches: float = 6.0,
) -> bool:
    """anchor_text 를 포함하는 첫 단락 바로 뒤에 이미지(+캡션)를 삽입한다.

    Args:
        in_docx: 원본 DOCX 경로(읽기 전용, 절대 덮어쓰지 않음).
        out_docx: 결과 DOCX 경로. **in_docx 와 같으면 ValueError**.
        anchor_text: 삽입 기준이 되는 본문 텍스트(부분 일치).
        png_path: 삽입할 PNG 이미지 경로.
        caption: 이미지 아래 캡션(빈 문자열이면 캡션 생략).
        width_inches: 이미지 폭(인치).

    Returns:
        anchor 를 찾아 그 뒤에 삽입했으면 True, 못 찾아 문서 끝에 추가했으면 False.

    Raises:
        ValueError: in_docx 와 out_docx 가 같은 경로일 때(원본 덮어쓰기 금지).
        FileNotFoundError: in_docx 또는 png_path 가 존재하지 않을 때.
    """
    in_path = Path(in_docx)
    out_path = Path(out_docx)
    if in_path.resolve() == out_path.resolve():
        raise ValueError("in_docx 와 out_docx 가 같습니다. 원본 덮어쓰기는 금지입니다.")
    if not in_path.exists():
        raise FileNotFoundError(f"입력 DOCX 가 없습니다: {in_docx}")
    if not Path(png_path).exists():
        raise FileNotFoundError(f"이미지 파일이 없습니다: {png_path}")

    # 원본은 손대지 않고 복사본을 만든 뒤 복사본만 수정한다.
    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(str(in_path), str(out_path))

    doc = Document(str(out_path))

    anchor_para = None
    for para in doc.paragraphs:
        if anchor_text and anchor_text in para.text:
            anchor_para = para
            break

    if anchor_para is not None:
        # anchor 바로 뒤에 (캡션 먼저 insert_before 하면 순서가 뒤집히므로)
        # 이미지 단락 → 캡션 단락 순으로 anchor 다음에 위치하도록 삽입한다.
        # python-docx 표준: 다음 단락 앞에 끼워넣는다. anchor 가 마지막이면 끝에 add.
        next_para = _next_paragraph(anchor_para)
        if next_para is not None:
            img_para = next_para.insert_paragraph_before()
            _add_picture(img_para, png_path, width_inches)
            if caption:
                cap_para = next_para.insert_paragraph_before()
                _add_caption(cap_para, caption)
        else:
            img_para = doc.add_paragraph()
            _add_picture(img_para, png_path, width_inches)
            if caption:
                _add_caption(doc.add_paragraph(), caption)
        doc.save(str(out_path))
        return True

    # anchor 미발견 → 문서 끝에 추가하고 False
    img_para = doc.add_paragraph()
    _add_picture(img_para, png_path, width_inches)
    if caption:
        _add_caption(doc.add_paragraph(), caption)
    doc.save(str(out_path))
    return False


def _next_paragraph(paragraph):
    """주어진 단락 바로 다음의 본문 단락을 반환한다(없으면 None)."""
    from docx.text.paragraph import Paragraph

    sib = paragraph._p.getnext()
    while sib is not None:
        if sib.tag == paragraph._p.tag:  # w:p
            return Paragraph(sib, paragraph._parent)
        sib = sib.getnext()
    return None


def _add_picture(paragraph, png_path: str, width_inches: float) -> None:
    run = paragraph.add_run()
    run.add_picture(png_path, width=Inches(width_inches))
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:  # pragma: no cover - 정렬 실패는 치명적이지 않음
        pass


def _add_caption(paragraph, caption: str) -> None:
    run = paragraph.add_run(caption)
    run.italic = True
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:  # pragma: no cover
        pass
