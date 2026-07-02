"""cross_form_autofill.py — 완성된 사업계획서 A(소스)의 라벨-값을
빈 양식 B(타깃)의 유사 항목 칸에 자동 전사하는 결정론 엔진(v1, AI無).

배경/목적
---------
같은 사업계획서를 여러 운영기관 양식에 다시 옮겨 적는 반복 작업을 줄인다.
A 문서에서 "라벨→값" 쌍을 뽑고, B 문서의 빈 항목 칸 라벨과 의미가 같은(동의어 포함)
칸을 보수적으로 매칭해 값을 전사한다.

핵심 안전 원칙
-------------
- **오매칭은 빈칸보다 나쁘다** → **high 만 자동 전사한다.**
  정확일치/동의어 단일후보(high)만 채운다. 퍼지(자카드·부분문자열)·약매칭·충돌은
  자동 채우지 않고 needs_confirm 으로 제안만 한다(접미사 다른 합성어 — 사업명/사업자명,
  주소/주소지정 — 가 자동 전사되는 것을 원천 차단).
- **날조 0** — 소스에 실제로 존재하는(비어있지 않은) 값만 옮긴다. 없으면 안 채운다.
- **원본 미수정** — A·B 는 절대 변형하지 않는다. out==source 또는 out==target 이면 ValueError.
  중간본은 tempfile 로 처리한다.
- **전사값 보존** — 전사 셀에는 ``docx_ops.set_cell_text`` 로 값을 직접 기입한다.
  잔여물/안내 청소 패스를 돌리지 않으므로 ``○○○`` 마스킹 이름·``OOO-OO-OOOOO`` 같은
  값이 훼손되지 않는다(거짓 보고 방지). transcribed 는 실제 저장 문서 기준으로 센다.
- AI 호출 없음 — 동일 입력, 동일 결과(결정론). use_ai 는 v1 에서 미사용(향후 확장 슬롯).

재사용
------
- ``SubmittableFiller._key`` / ``_logical_cells`` 의 정규화·병합셀 안전 순회를 재사용한다
  (여기서는 모듈 함수 _key/_logical_cells 로 노출해 일관 사용).
- 전사는 타깃 칸 좌표(table_index/row/value_cell)에 ``docx_ops.set_cell_text`` 로 직접
  기입한다(라벨 다음 칸이 아니라 정확한 값셀에, 잔여물 청소 없이).
- A/B 가 .hwp/.hwpx 면 ``hwp_docx_convert`` 로 DOCX 변환, out 이 .hwp 면 docx_to_hwp.
"""

from __future__ import annotations

import json
import re
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from .docx_ops import set_cell_text
from .submittable_filler import SubmittableFiller

_HWP_EXTS = {".hwp", ".hwpx"}
_SUPPORTED_EXTS = {".docx", ".hwp", ".hwpx"}  # H7: 입력 지원 확장자 화이트리스트


# --- 라벨 정규화 / 병합셀 순회(SubmittableFiller 규칙 재사용) -------------------

# 실제 양식 라벨에 거의 항상 붙는 선행 '장식'(글머리표/순번)을 벗겨 라벨 변형
# recall 을 높인다. 의미를 바꾸지 않는 접두만 제거하므로 오매칭을 새로 만들지 않는다
# (장식을 벗긴 뒤에도 '사업명'≠'사업자명'은 그대로 다른 키다).
_BULLET_PREFIX_RE = re.compile(r"^[○●◯◌▶▷◀◁◆◇■□▪▫▸▹◦‣⦁·∙•*※→⇒]+")
_NUM_PREFIX_RES = (
    re.compile(r"^\d{1,2}\s*[.)]\s*"),                       # 1. 1) 12.
    re.compile(r"^[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]\s*"),         # 원문자 숫자
    re.compile(r"^[가나다라마바사아자차카타파하]\s*[.)]\s*"),  # 가. 나)
    re.compile(r"^[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ]\s*[.)]?\s*"),               # 로마자 순번
)
_HAS_WORD_RE = re.compile(r"[가-힣A-Za-z0-9]")


def _strip_label_decoration(key: str) -> str:
    """정규화 라벨 키에서 선행 글머리표/순번 접두를 보수적으로 제거한다.

    글머리표 1패스 + 순번 1패스를 최대 2회 반복한다('1. ○ 라벨' 같은 중첩 대응).
    구분자(. 또는 ))가 없는 정상 라벨('1차년도', '가산점')은 깎지 않는다.
    제거 결과에 글자(한글/영숫자)가 하나도 없으면 원본을 그대로 둔다
    (마스킹 라벨 '○○○' 같은 전부-기호 토큰 보호).
    """
    s = key
    for _ in range(2):
        before = s
        s = _BULLET_PREFIX_RE.sub("", s, count=1)
        for rx in _NUM_PREFIX_RES:
            s2 = rx.sub("", s, count=1)
            if s2 != s:
                s = s2
                break
        if s == before:
            break
    s = s.strip()
    if not _HAS_WORD_RE.search(s):
        return key
    return s


def _key(text: str) -> str:
    """라벨 비교용 핵심 키: 괄호·공백 제거 + 선행 글머리표/순번 장식 제거.

    SubmittableFiller._key(괄호·공백 제거)에 더해, 실제 양식 라벨에 흔한
    글머리표(○ ▶ ·)·순번(1. ① 가. Ⅰ.) 접두를 벗겨 라벨 변형 recall 을 높인다.
    의미를 바꾸지 않는 장식만 제거하므로 보수성(오매칭 차단)은 그대로다.
    """
    return _strip_label_decoration(SubmittableFiller._key(text))


def _logical_cells(row) -> list:
    """병합 중복을 제거한 논리적 셀 목록(SubmittableFiller._logical_cells 와 동일 규칙)."""
    return SubmittableFiller._logical_cells(row)


def _iter_all_tables(container) -> list:
    """문서/셀의 모든 표를 문서 순서로(중첩표 포함) 평탄화해 반환한다.

    재발 클래스 D 차단용 **단일 순회 헬퍼**. ``python-docx`` 의 ``doc.tables`` 는
    **최상위 표만** 준다(표 셀 '안'에 들어간 중첩표를 통째로 놓친다). 이 헬퍼는 각 표의
    논리 셀 안 중첩표까지 재귀로 훑어, 소스추출·타깃탐지·체크박스·전사기입이 **동일한
    표 좌표계**(같은 평탄화 순서 = ``table_index``)를 공유하게 한다. 과거 앵커탐색·셀
    인라인 갭이 함수마다 제각각 순회하다 표 셀·중첩표를 반복해 놓친 재발을 구조적으로 막는다.

    - **순서**: 전위(부모 표 → 그 셀 안 중첩표 → 다음 부모). 탐지와 기입이 같은 함수를
      쓰므로 인덱스가 항상 정합한다.
    - **회귀 0**: 중첩표가 없으면 결과는 ``list(container.tables)`` 와 완전히 동일하다
      (기존 최상위-표 동작을 그대로 보존).
    - ``_logical_cells`` 로 병합 중복 셀을 건너뛰어 같은 중첩표를 두 번 담지 않는다.
      표는 아래로만 중첩되므로 재귀는 유한하다.
    """
    out: list = []
    for table in container.tables:
        out.append(table)
        for row in table.rows:
            for cell in _logical_cells(row):
                out.extend(_iter_all_tables(cell))
    return out


# --- 동의어 클러스터 ----------------------------------------------------------
# 각 클러스터는 의미가 같은 라벨들의 모음이다. 정규화(_key) 기준으로 비교한다.
# 확장 가능: 새 클러스터를 리스트에 추가하면 _CLUSTER_OF 가 자동 반영된다.
SYNONYMS: list[list[str]] = [
    ["기업명", "신청기관", "업체명", "회사명", "상호", "기관명", "법인명",
     "회사명칭", "기업체명", "상호명",
     "신청기업명", "신청업체명", "사업체명", "회사이름", "기업이름", "기업명칭",
     # 2026-07-01 2차 확장 (주관/수행기관명=운영주체, 수요/공급기업명=B2B상대 → 제외)
     "신청자상호", "사업자명", "단체명", "팀명", "창업기업명", "참여기업명",
     "지원기업명", "입주기업명", "신청법인명"],
    ["대표자", "성명", "대표", "대표이사", "대표자명", "대표자성명",
     "성명(대표자)", "신청인",
     # 신청자명↔성명 미매칭(실측) 해소. 이름값 가드가 오전사 차단.
     "신청자명", "신청자", "신청인명", "신청자성명", "신청인성명",
     "참가자명", "참가자", "성함", "대표성명", "대표자님", "멘티명",
     # 2026-07-01 2차 확장 (책임자류는 담당자명 신설로 분리 — 대표≠담당)
     "기업대표", "법인대표", "신청대표", "신청자이름", "대표자이름"],
    ["사업명", "과제명", "사업명칭", "과제명칭", "사업과제명", "과제명(사업명)",
     "사업아이템명", "아이템명", "사업아이템", "창업아이템", "창업아이템명",
     "프로젝트명", "사업제목", "아이템",
     # 2026-07-01 2차 확장 (제품명·서비스명·브랜드명은 제품서비스명 신설로 분리)
     "지원사업명", "신청사업명", "신청과제명", "개발과제명", "연구과제명",
     "과제제목", "사업제안명", "제안과제명", "창업아이디어명"],
    ["사업자등록번호", "사업자번호", "사업자등록No", "사업자등록번호No",
     "사업자등록증번호", "사업자등록증상번호",
     # 2026-07-01 2차 확장 ('고유번호'·'등록번호' 단독은 초일반 → 제외)
     "개인사업자번호", "법인사업자번호"],
    ["연락처", "전화", "전화번호", "휴대전화", "연락전화", "대표전화",
     "핸드폰", "휴대폰",
     "연락번호", "휴대폰번호", "핸드폰번호", "이동전화", "담당자연락처",
     "담당자전화", "휴대전화번호",
     # 2026-07-01 2차 확장
     "유선전화", "회사전화", "사무실전화", "사업장전화", "대표번호",
     "대표연락처", "개인연락처", "비상연락처", "신청자연락처", "연락가능번호",
     "모바일번호"],
    ["주소", "소재지", "사업장주소", "사업장소재지", "주된사무소소재지",
     "본사주소", "회사주소",
     # 주소지↔자택주소 미매칭(실측) 해소. 한 양식에 2칸이면 conflict 로 자동전사 보류.
     "주소지", "자택주소", "거주지", "거주지주소", "본점소재지",
     "회사소재지", "기업소재지", "신청인주소", "대표자주소", "본사소재지",
     "사업장주소지", "실거주지",
     # 2026-07-01 2차 확장
     "사업장", "본점주소", "주사무소", "주사무소소재지", "사무실주소",
     "공장주소", "공장소재지", "지점주소", "지사주소", "영업장주소",
     "실제사업장주소", "거소지"],
    ["업종", "산업분류", "업태", "주업종", "업종업태",
     # 2026-07-01 2차 확장 ('사업분야'는 지원분야 신설과 혼동 → 제외)
     "업종명", "업태명", "표준산업분류", "산업분류코드", "영위업종",
     "영업종목", "업종코드", "종목"],
    ["이메일", "전자우편", "email", "e-mail", "메일", "이메일주소", "전자메일",
     "e메일", "메일주소", "담당자이메일", "대표이메일", "이메일(email)",
     # 2026-07-01 2차 확장
     "이메일계정", "전자우편주소", "회사메일", "수신이메일", "연락이메일"],
    ["설립일", "설립연월일", "창업일", "설립일자", "개업연월일",
     "창업연월일", "설립년월일",
     "개업일", "개업일자", "회사설립일", "법인설립일", "설립연도", "창업연도",
     # 2026-07-01 2차 확장
     "사업개시일", "사업개시일자", "영업개시일", "창업일자", "최초창업일",
     "사업자등록일", "법인전환일", "법인등기일"],
    # --- 실무 라벨 변형(2026-06-23 확장): 정부지원사업 양식 빈출 항목 ---
    ["직원수", "종업원수", "상시근로자수", "고용인원", "임직원수",
     "근로자수", "상시종업원수", "총직원수", "재직인원",
     # 2026-07-01 2차 확장
     "직원현황", "인력현황", "고용현황", "상시고용인원", "현재직원수",
     "근무인원", "4대보험가입자수", "고용보험가입자수"],
    ["자본금", "납입자본금", "자본금규모", "자본금액",
     # 2026-07-01 2차 확장
     "출자금", "출자금액", "납입금액", "자기자본", "자본총계",
     "설립자본금", "등기자본금"],
    ["팩스", "팩스번호", "모사전송", "fax",
     # 2026-07-01 2차 확장
     "팩스연락처", "회사팩스"],
    ["홈페이지", "웹사이트", "홈페이지주소", "회사홈페이지", "website",
     # 2026-07-01 2차 확장
     "url", "사이트", "웹주소", "웹페이지", "공식홈페이지", "자사몰",
     "쇼핑몰주소"],
    ["직위", "직급", "직책",
     # 슬래시는 _key 에서 유지되므로 슬래시 포함 라벨을 그대로 추가
     "직위/직책", "직책/직위", "담당자직위", "대표직위",
     # 2026-07-01 2차 확장 ('역할'·'담당역할'은 초일반 → 제외)
     "직함", "담당자직책", "담당자직급", "신청자직위", "책임자직위", "소속직위"],
    ["부서", "소속", "소속부서", "부서명", "담당부서", "소속팀",
     # 2026-07-01 2차 확장
     "소속명", "소속기관", "소속회사", "소속기업", "담당팀", "근무부서",
     "신청부서", "관리부서"],
    ["생년월일", "출생년월일", "출생연월일", "생일", "출생일",
     # 2026-07-01 2차 확장
     "생년월일자", "출생일자", "만생년월일", "신청자생년월일", "대표자생년월일"],
    ["법인등록번호", "법인등기번호", "법인번호",
     # 2026-07-01 2차 확장
     "등기번호", "법인등기부번호", "법인고유번호"],
    ["성별", "성별구분",
     # 2026-07-01 2차 확장 ('남/여'·'남여'는 값이라 라벨 아님 → 제외)
     "남녀구분", "신청자성별", "대표자성별", "참가자성별"],
    # === 2026-07-01 신설 5클러스터 (사용자 요청) ===
    # 담당자명 — 대표자와 분리(실무담당자). 대표자 클러스터엔 넣지 않음(사용자 강조).
    ["담당자", "실무자", "담당자명", "실무담당자", "신청담당자", "과제담당자",
     "연락담당자", "책임자", "담당책임자"],
    # 지원분야 — 공고별 신청유형. '분야/유형/트랙/리그' 단독(초일반)은 제외.
    ["지원분야", "신청분야", "지원유형", "모집분야", "신청유형", "모집부문",
     "지원부문"],
    # 제품/서비스명 — 사업명과 분리(사업명≠제품명 가능).
    ["제품명", "서비스명", "솔루션명", "플랫폼명", "앱명", "브랜드명",
     "주력제품명", "제품서비스명"],
    # 매출액
    ["매출액", "매출", "연매출", "최근매출액", "전년도매출액", "당기매출액",
     "매출현황", "매출규모"],
    # 사업비 — 총사업비/소요예산 계열만(정부지원금·자부담금은 서로 다른 값이라 별도 항목 → 제외).
    ["총사업비", "사업비", "소요예산", "총소요예산"],
    # 사업자 형태(선택칸, master 병합) — 체크박스 그룹 라벨 ↔ 소스 값('개인사업자'/'법인') 매칭용.
    ["사업자형태", "사업자구분", "기업형태", "사업자유형", "기업구분",
     "법인형태", "창업기업형태", "사업형태"],
]

# 정규화 라벨 → 대표키(클러스터의 첫 원소). 같은 대표키면 동의어로 본다.
_CLUSTER_OF: dict[str, str] = {}
for _cluster in SYNONYMS:
    _rep = _key(_cluster[0])
    for _alias in _cluster:
        _CLUSTER_OF[_key(_alias)] = _rep


def _cluster_rep(norm_label: str) -> Optional[str]:
    """정규화 라벨이 속한 동의어 클러스터 대표키. 없으면 None."""
    return _CLUSTER_OF.get(norm_label)


# 이름(성명)류 동의어 클러스터 대표키 — 이 클러스터 타깃엔 '이름 모양' 값만 high 전사.
_NAME_FIELD_REP: Optional[str] = _CLUSTER_OF.get(_key("대표자"))

# 이름이 아님을 드러내는 표지: 나열 구분자(콤마·가운뎃점·세미콜론·파이프 등) / 역할·책임 서술어.
_NON_NAME_RE = re.compile(r"[,，;/·•∙ㆍ、|&]|및|총괄|담당|수행|자문")


def _is_name_field(norm_label: str) -> bool:
    """정규화 라벨이 이름(성명)류 동의어 클러스터에 속하면 True."""
    return _NAME_FIELD_REP is not None and _cluster_rep(norm_label) == _NAME_FIELD_REP


def _looks_like_name(value: str) -> bool:
    """값이 사람 이름(또는 마스킹 ○○○)으로 그럴듯하면 True.

    이름은 짧고 나열·역할서술이 없다. 역할분담 서술("대표자 : 기술개발, 특허전략 및
    사업화 총괄")이 이름칸에 high 전사되는 실측 오류를 차단한다(오매칭<빈칸). 가드는
    이름필드에만 적용하므로 비이름 필드(연락처 등)의 콤마 값은 영향받지 않는다.
    """
    v = (value or "").strip()
    if not v:
        return False
    if len(v) > 20:               # 외국식 이름을 포함해도 이름은 길지 않다
        return False
    return not _NON_NAME_RE.search(v)


_BRACKET_RE = re.compile(r"[\(（]([^\)）]*)[\)）]")


def _bracket_tokens(text: str) -> set[str]:
    """라벨의 괄호 안 토큰 집합(공백 제거). 예: '금액(국고)' → {'국고'}."""
    return {
        re.sub(r"\s+", "", m).strip()
        for m in _BRACKET_RE.findall(str(text or ""))
        if re.sub(r"\s+", "", m).strip()
    }


def _bracket_conflict(label_a: str, label_b: str) -> bool:
    """두 원본 라벨의 괄호 토큰이 양쪽에 존재하고 서로 다르면 True(H3).

    '금액(국고)' vs '금액(자부담)' → _key 는 둘 다 '금액'이라 정확일치하지만
    괄호 구별 토큰이 달라 같은 항목이 아니다 → high 금지.
    """
    ta = _bracket_tokens(label_a)
    tb = _bracket_tokens(label_b)
    if not ta or not tb:
        return False
    return ta != tb


# --- 데이터 모델 --------------------------------------------------------------

@dataclass
class Match:
    target_label: str          # 타깃 원본 라벨 텍스트
    normalized: str            # 타깃 정규화 라벨
    source_label: str          # 매칭된 소스 원본 라벨(자동전사 대상이 아니면 "")
    value: str                 # 전사할 값(소스 실값; 전사 대상이 아니면 "")
    confidence: str            # "high"(자동전사) | "fuzzy" | "conflict" | "low"
    table_index: int           # 타깃 표 인덱스("paragraph" 칸이면 -1)
    row: int                   # 타깃 행("paragraph" 칸이면 -1)
    value_cell: int            # 타깃 값 셀(논리 셀) 인덱스("paragraph" 칸이면 -1)
    kind: str = "table"        # 채울 칸 종류: "table" | "paragraph" | "cell_paragraph"
    para_index: int = -1       # 본문 단락 인덱스(kind=="paragraph" 일 때만 유효)
    cell_para_index: int = -1  # 셀 내부 단락 인덱스(kind=="cell_paragraph" 일 때만; value_cell=논리셀)
    fill_start: int = -1       # 단락 기입 구간 시작(kind=="paragraph"/"cell_paragraph", 멀티필드)
    fill_end: int = -1         # 단락 기입 구간 끝(이 구간을 ' '+값 으로 교체)
    candidates: list[str] = field(default_factory=list)  # 보고용 후보(자동전사 아님)

    def as_dict(self) -> dict[str, Any]:
        return {
            "target_label": self.target_label,
            "normalized": self.normalized,
            "source_label": self.source_label,
            "value": self.value,
            "confidence": self.confidence,
            "table_index": self.table_index,
            "row": self.row,
            "value_cell": self.value_cell,
            "kind": self.kind,
            "para_index": self.para_index,
        }


@dataclass
class AutofillReport:
    source: str
    target: str
    output: str
    transcribed: int = 0
    confirmed: int = 0         # 그중 사용자 확정(confirmations)으로 채운 칸 수
    checkbox_checked: int = 0  # 자동 체크한 선택칸 수(□→■)
    src_fields: int = 0        # 소스에서 추출한 라벨-값 수(진단용, H6)
    tgt_fields: int = 0        # 타깃에서 찾은 빈칸 수(진단용, H6)
    matches: list[Match] = field(default_factory=list)        # 실제 전사한 매칭
    needs_confirm: list[dict[str, Any]] = field(default_factory=list)  # 애매/충돌(보류)
    unmatched_targets: list[dict[str, Any]] = field(default_factory=list)
    checkbox_groups: list[dict[str, Any]] = field(default_factory=list)  # 선택칸 판정 결과
    ok: bool = False
    notes: list[str] = field(default_factory=list)

    def as_dict(self) -> dict[str, Any]:
        return {
            "source": self.source,
            "target": self.target,
            "output": self.output,
            "transcribed": self.transcribed,
            "confirmed": self.confirmed,
            "checkbox_checked": self.checkbox_checked,
            "src_fields": self.src_fields,
            "tgt_fields": self.tgt_fields,
            "matches": [m.as_dict() for m in self.matches],
            "needs_confirm": self.needs_confirm,
            "unmatched_targets": self.unmatched_targets,
            "checkbox_groups": self.checkbox_groups,
            "ok": self.ok,
            "notes": self.notes,
        }


# --- 추출: 소스에서 라벨→값 ----------------------------------------------------

_PARA_LABEL_VALUE_RE = re.compile(r"^\s*([^:：]{1,40})\s*[:：]\s*(.+?)\s*$")

# 순수 숫자/금액/날짜 패턴(값이 라벨로 둔갑하는 행렬표에서 라벨다움 판정용).
# 예: "30,000,000", "60,000원", "2026-04-10", "2026.04.10", "10%", "1,234.5"
_NUMERIC_VALUE_RE = re.compile(
    r"^[\s\d.,\-/원만억천백십개%~()–—:]+$"
)


def _looks_numeric(text: str) -> bool:
    """값이 순수 숫자/금액/날짜/수량 패턴이면 True(라벨로 등록 부적합)."""
    t = (text or "").strip()
    if not t:
        return False
    if not any(ch.isdigit() for ch in t):
        return False
    return bool(_NUMERIC_VALUE_RE.match(t))


def _is_numeric_data_row(logical) -> bool:
    """행의 비어있지 않은 셀이 대부분(과반 이상) 순수 숫자/금액이면 데이터 행."""
    vals = [c.text.strip() for c in logical if c.text and c.text.strip()]
    if len(vals) < 2:
        return False
    numeric = sum(1 for v in vals if _looks_numeric(v))
    return numeric >= len(vals) / 2


def _is_matrix_header_row(table_rows, ri: int) -> bool:
    """행렬형 헤더 행(라벨만 가득 + 다음 행이 숫자 데이터)이면 True.

    예: 행0=[국고보조금, 자기부담금, 총사업비] (라벨들), 행1=[30M, 10M, 40M] (숫자).
    이런 헤더 행은 짝수=라벨 위치단정이 라벨→라벨 페어링을 만들므로 건너뛴다.
    """
    if ri + 1 >= len(table_rows):
        return False
    logical = _logical_cells(table_rows[ri])
    non_empty = [c.text.strip() for c in logical if c.text and c.text.strip()]
    if len(non_empty) < 3:
        return False
    # 헤더 행 자신은 비숫자 라벨들로 가득해야 한다.
    if any(_looks_numeric(v) for v in non_empty):
        return False
    # 바로 다음 행이 숫자 데이터 행이어야 한다.
    return _is_numeric_data_row(_logical_cells(table_rows[ri + 1]))


def _collect_label_keys(doc) -> set[str]:
    """문서에서 '라벨로도 등장하는' 셀들의 정규화 키 집합(값 둔갑 차단용).

    두 종류를 수집한다:
    1. 짝수 인덱스(0,2,…) 셀 — 정상 라벨-값 쌍의 라벨 위치.
    2. **반복 셀** — 같은 열 인덱스에서 2개 이상 행에 같은 텍스트로 나타나는 셀.
       행렬/세로병합 예산표의 행머리(예: 국고보조금)·범주 라벨(인증평가·현금)은
       값이 아니라 분류 라벨인데 홀수 인덱스에 와도 반복되므로 이렇게 잡는다.

    행렬/세로배치 예산표에서 한 셀이 '값'으로 추출돼도 같은 문서에 '라벨'로도
    등장하면 그 (라벨,값) 쌍을 폐기하기 위한 라벨 사전이다.
    """
    keys: set[str] = set()
    # (col_index, key) → 출현 횟수
    col_counts: dict[tuple[int, str], int] = {}
    for table in _iter_all_tables(doc):
        rows = table.rows
        for ri, row in enumerate(rows):
            logical = _logical_cells(row)
            header_row = _is_matrix_header_row(rows, ri)
            for i, cell in enumerate(logical):
                k = _key(cell.text)
                if not k:
                    continue
                if i % 2 == 0 or header_row:  # 짝수 인덱스 또는 행렬 헤더 = 라벨 위치
                    keys.add(k)
                if not _looks_numeric(cell.text):  # 숫자는 반복돼도 라벨 아님
                    col_counts[(i, k)] = col_counts.get((i, k), 0) + 1
    # 같은 열에서 2회 이상 반복된 비숫자 셀 = 분류/행머리 라벨
    for (_, k), cnt in col_counts.items():
        if cnt >= 2:
            keys.add(k)
    return keys


def _is_bad_value(value: str, label_keys: set[str]) -> bool:
    """추출된 '값'이 실제로는 또 다른 라벨이라서 폐기해야 하면 True(C1/H1/M3 가드).

    (a) 값의 정규화 키가 동의어 클러스터 라벨이거나,
    (b) 같은 문서에서 라벨로도 등장하면(짝수 인덱스 또는 반복 셀) → 라벨→라벨
        오염이므로 폐기.

    숫자/금액/날짜 값(예: 123-45-67890, 30,000,000)은 그 자체로 라벨이 될 수 없어
    라벨→라벨 전사를 일으키지 않으므로 폐기하지 않는다(사업자등록번호 등 실값 보존).
    동의어 클러스터(a)는 label_keys 에 이미 그 라벨이 함께 등장할 때만 의미가 있으므로
    (b) 와 결합해서만 적용한다(단발 '연락처' 값 보존 — 기존 회귀 보호).
    """
    vkey = _key(value)
    if not vkey:
        return True
    # (b) 같은 문서에서 라벨로도 등장 → 라벨→라벨 오염
    if vkey in label_keys:
        return True
    # (a) 동의어 클러스터 라벨이면서, 그 클러스터의 다른 라벨이 문서 라벨로도 존재 →
    #     예산표의 분류 라벨이 값으로 둔갑한 경우. 단발 값(연락처 1회)은 보존.
    rep = _cluster_rep(vkey)
    if rep is not None:
        for lk in label_keys:
            if _cluster_rep(lk) == rep and lk != vkey:
                return True
    return False


def _vertical_header_cols(header: list) -> list[int]:
    """헤더 행이 '세로형 카드 헤더'면 인식 라벨 열 인덱스 목록, 아니면 [].

    세로형 정보 카드(표지·개요 박스)는 라벨을 한 줄에 나열하고(성명·생년월일·연락처)
    **아래 줄**에 값을 적는다. 소스 추출을 위해 이런 헤더 행을 판별한다.

    판별 규칙(가로 라벨|값 행과의 확실한 구분 = 오추출<빈칸):
      - 비어있지 않은 논리셀이 **2개 이상**이고 **전부 동의어 클러스터 라벨**일 때만.
        한 칸이라도 라벨 아닌 값(예: '밸류업(주)')이 섞이면 가로 라벨|값 행이므로
        세로 헤더로 보지 않는다(그 행은 짝수-인덱스 가로 추출이 처리).
      - 예산 행렬 헤더([국고보조금, 자기부담금, 총사업비])는 앞 두 라벨이 클러스터
        미등록이라 '전부 라벨' 조건에서 탈락 → 세로 값 추출 대상이 아니다.
      - **동의어 중복 배제**: 같은 동의어 클러스터 라벨이 두 칸 이상이면(예 부서·소속)
        정상 정보 카드가 아니다. 진짜 카드는 서로 다른 필드를 나열한다. 이 중복은
        가로 라벨|값 더블페어 행([직위|대표|부서|소속], 대표·소속이 값이지만 라벨
        어휘)을 세로로 오인한 신호 → 폐기(오추출<빈칸, 리뷰 MEDIUM 잔여 차단).
    반환된 열 인덱스만 아래 행에서 값을 읽는다.
    """
    non_empty = [i for i, c in enumerate(header) if (c.text or "").strip()]
    if len(non_empty) < 2:
        return []
    reps: list[str] = []
    for i in non_empty:
        rep = _cluster_rep(_key(header[i].text))
        if rep is None:
            return []  # 라벨 아닌 셀 존재 → 세로 헤더 아님
        reps.append(rep)
    if len(set(reps)) != len(reps):
        return []  # 같은 클러스터 라벨 중복 → 정상 카드 아님(가로 더블페어 오인 차단)
    return non_empty


def _extract_vertical_cards(doc, put) -> None:
    """세로형(라벨 위 / 값 아래) 정보 카드에서 (라벨,값)을 추출해 put 한다.

    가로 짝수-인덱스 추출이 못 읽는 구조(헤더=라벨 나열, 아래=값)를 보강한다.
    보수 가드(오추출<빈칸):
      - 헤더 행이 ``_vertical_header_cols`` 로 '전부 인식 라벨(≥2)'일 때만.
      - ``len(header)==len(below)`` (병합 정렬 보장).
      - **아래 행 라벨 가드(가로 오인 차단)**: 아래 행의 인식 열 셀 중 하나라도 그
        자체가 필드 라벨(``_cluster_rep`` 존재)이면, 이 행은 세로 헤더가 아니라
        '가로 라벨|값 행'을 오인한 것으로 보고 **카드 전체를 폐기**한다. 예:
        ``[직위|대표] / [성명|홍길동]`` — 아래 [성명,홍길동]의 '성명'이 라벨이라
        세로로 읽으면 ``대표→홍길동`` 오추출이 되지만, 가로 추출이 ``성명→홍길동``
        정답을 낸다. (값 셀 단어가 라벨 어휘와 겹치는 정상 가로 표 보호.)
      - **로스터 가드**: 아래-아래 행(ri+2)의 인식 열에 값(라벨 아님)이 **하나라도**
        있으면 다중 레코드 명부로 보고 카드 전체를 건너뛴다(첫 레코드를 대표값으로
        오인 기입 방지). 2번째 레코드의 일부 칸이 비어 있어도 로스터로 판정한다.
    값이 없는 칸은 건너뛴다(날조 0). ``put`` 은 ``_extract_source`` 의 ``_put``
    (이미 있으면 유지). 세로가 **가로보다 먼저** 실행되므로(순수 세로 카드의 헤더 행이
    가로에 라벨→라벨로 오염되는 것을 막기 위해), 위 '아래 행 라벨 가드'가 진짜 가로
    행을 세로가 잘못 선점하는 유일한 경로를 차단한다.
    """
    for table in _iter_all_tables(doc):
        rows = table.rows
        for ri in range(len(rows) - 1):
            header = _logical_cells(rows[ri])
            below = _logical_cells(rows[ri + 1])
            if not header or len(header) != len(below):
                continue
            cols = _vertical_header_cols(header)
            if not cols:
                continue
            # 아래 행 라벨 가드: 아래 인식 열 셀 중 하나라도 필드 라벨이면 세로 헤더가
            # 아니라 가로 라벨|값 행 오인 → 카드 전체 폐기(가로 추출이 정답을 낸다).
            if any(
                (below[ci].text or "").strip()
                and _cluster_rep(_key(below[ci].text)) is not None
                for ci in cols
            ):
                continue
            # 로스터(다중 레코드) 가드: 아래-아래 행의 인식 열에 값(라벨 아님)이 하나라도
            # 있으면 명부 → 카드 전체 건너뜀. 2번째 레코드가 일부 비어도 로스터로 본다.
            if ri + 2 < len(rows):
                below2 = _logical_cells(rows[ri + 2])
                if len(below2) == len(header) and any(
                    (below2[ci].text or "").strip()
                    and _cluster_rep(_key(below2[ci].text)) is None
                    for ci in cols
                ):
                    continue
            for ci in cols:
                value = (below[ci].text or "").strip()
                if not value:
                    continue                          # 값 없음 → 건너뜀(날조 0)
                label_text = header[ci].text
                label = _key(label_text)
                if not label:
                    continue
                put(label, label_text, value)


def extract_source_fields(docx_path: str | Path) -> dict[str, str]:
    """소스 DOCX 의 표에서 (라벨,값) 쌍을 추출한다.

    - **짝수 인덱스(0,2,4…)만 라벨로** 본다: col0=라벨,col1=값,col2=라벨,col3=값.
      값 셀(홀수 인덱스)을 라벨로 쓰지 않으므로 ``[항목][연락처][010-…]`` 같은 행에서
      값 셀(연락처/010-…)이 라벨로 오염되지 않는다.
    - **방어 가드(C1/H1/M3)**: 추출된 '값'이 (a)동의어 클러스터 라벨이거나
      (b)같은 문서에서 라벨로도 등장하거나 (c)순수 숫자/금액/날짜 패턴이면 그 쌍을
      폐기한다. 행렬/세로배치/병합 예산표에서 라벨이 값으로 둔갑해 high 오전사되는
      것을 차단한다(오매칭은 빈칸보다 나쁘다).
    - 정규화 라벨(_key)을 키로, 라벨·값 둘 다 비어있지 않을 때만 담는다(날조 0의 출발점).
    - 같은 정규화 라벨이 여러 번 나오면 처음 채워진 값을 유지한다.
    - **세로형 카드**(헤더 행=라벨 나열 / 아래 행=값)도 추출한다: 헤더가 전부 인식
      라벨(≥2)일 때만 아래 행 값을 읽는다(_extract_vertical_cards). 표지·개요 박스의
      세로형 정보 카드가 통째로 유실되던 갭을 메운다(가로 짝수-인덱스가 못 읽는 구조).
    - 보조: 본문 단락의 "라벨: 값" 패턴도 추출(표에 없을 때만 보강).
    """
    fields, _ = _extract_source(docx_path)
    return fields


def _extract_source(docx_path: str | Path) -> tuple[dict[str, str], dict[str, str]]:
    """extract_source_fields 의 내부 구현. (정규화필드, 원본라벨맵)을 함께 반환.

    원본라벨맵은 {정규화키: 원본 라벨 텍스트} 로, H3(괄호 토큰 대조)에 쓰인다.
    """
    doc = Document(str(docx_path))
    fields: dict[str, str] = {}
    originals: dict[str, str] = {}
    label_keys = _collect_label_keys(doc)

    def _put(label: str, label_text: str, value: str) -> None:
        if label not in fields:
            fields[label] = value
            originals[label] = SubmittableFiller._norm(label_text)

    # 세로형(라벨 위 / 값 아래) 정보 카드에서 값 추출 — 가로 짝수-인덱스 추출이 못
    # 읽는 표지형 카드(성명·생년월일·연락처를 위에 나열, 아래 행에 값). **가로보다
    # 먼저** 실행한다: 순수 세로 카드의 헤더 행(라벨만 나열)은 가로 짝수-인덱스가
    # 라벨→라벨(기업명→대표자)로 오염시켜 _put 이 그 오염값을 먼저 고정하기 때문.
    # 둘은 상호배타(진짜 가로 라벨|값 행은 col1 이 값이라 세로 헤더 조건 미충족)이므로
    # 세로-우선은 올바른 가로 추출을 절대 덮지 않는다(오추출<빈칸).
    _extract_vertical_cards(doc, _put)

    for table in _iter_all_tables(doc):
        rows = table.rows
        for ri, row in enumerate(rows):
            logical = _logical_cells(row)
            # 행렬형 헤더 행(라벨만 + 다음 행 숫자)은 짝수 페어링 비적용(C1)
            if _is_matrix_header_row(rows, ri):
                continue
            # 짝수 인덱스만 라벨: (0,1)(2,3)(4,5)… 쌍으로 처리
            for i in range(0, len(logical) - 1, 2):
                label_text = logical[i].text
                label = _key(label_text)
                value = (logical[i + 1].text or "").strip()
                if not label or not value:
                    continue
                # 라벨 자리에 순수 숫자/금액/날짜가 오면(행렬표 데이터행) 라벨 등록 제외(M3)
                if _looks_numeric(label_text):
                    continue
                # 값이 라벨/동의어로 둔갑한 행렬표 오염 차단(C1/H1)
                if _is_bad_value(value, label_keys):
                    continue
                _put(label, label_text, value)

    # 보조: 본문 "라벨: 값" — 표에서 못 얻은 라벨만 보강(한 줄 여러 칸 대응).
    # 멀티필드 줄("기업명 : A    대표자 : B")을 칸별로 추출해 첫 칸 값에 둘째 칸이
    # 묶이는 오염을 막는다. 빈칸(밑줄/공백)인 소스 칸은 값이 아니므로 등록 제외(날조 0).
    for para in doc.paragraphs:
        for label_raw, value_raw, _s, _e in _iter_line_fields(para.text or ""):
            label = _key(label_raw)
            value = (value_raw or "").strip()
            if not label or not value:
                continue
            if _is_fill_blank(value_raw):
                continue  # 소스 칸이 비어 있음(밑줄/공백) → 값 아님
            if label not in fields:
                _put(label, label_raw, value)

    return fields, originals


# --- 탐지: 타깃에서 빈 값칸 ----------------------------------------------------

# 한 줄 안의 "라벨[:：]값" 칸을 **모두** 찾는 패턴(멀티필드 단락 대응).
# 실제 정부양식 신청서 표지는 한 줄에 여러 칸을 나란히 둔다(기업명 : __  대표자 : __).
# 첫 콜론만 보던 기존 방식은 둘째 칸을 값으로 오인해 그 줄 전체를 놓쳤다 → 각 칸을 개별 인식.
#  - label: 콜론/탭 없는 1~39자(앞 공백 제외) — 긴 문장이 라벨로 둔갑하는 것 차단.
#  - value: 비탐욕. 다음 칸 시작(2+공백 또는 탭 뒤 비공백) 또는 줄 끝까지.
# 콜론 뒤 정렬 공백은 value 그룹 앞쪽에 포함된다(빈칸 판정·기입 시 한 칸 공백으로 정규화).
_LINE_FIELD_RE = re.compile(
    r"(?P<label>[^\s:：][^:：\t]{0,38}?)"
    r"[ \t　]*(?P<colon>[:：])"
    r"(?P<value>.*?)"
    r"(?=(?:[ \t　]{2,}|\t)\S|[ \t　]*$)"
)

# 콜론 뒤 '값 자리'가 빈칸으로 인정되는 채움기호 집합(전부-이 문자거나 공백/빈문자).
# 밑줄·점·대시·전각공백 등 명백한 빈칸 표기만 인정한다. ○●□▢ 같은 동그라미/네모는
# 블라인드 마스킹값(○○○)·체크박스와 혼동되므로 **빈칸으로 보지 않는다**(보존 우선).
_FILL_BLANK_RE = re.compile(
    r"^[\s_.·․‥…　─━―—–\-]*$"
)


def _iter_line_fields(text: str):
    """한 줄 텍스트에서 "라벨[:：]값" 칸들을 왼쪽부터 순회한다(멀티필드 대응).

    각 칸을 ``(label_raw, value_raw, fill_start, fill_end)`` 로 내준다.
      - label_raw: 콜론 앞 라벨 원문(선행 장식 포함; 정규화는 _key 담당).
      - value_raw: 콜론 뒤 값 원문(빈칸이면 공백/밑줄/점 등; 콜론 뒤 정렬공백 포함).
      - fill_start/fill_end: value_raw 가 차지하는 text 내 [시작,끝) 구간. 이 구간을
        ``" " + 값`` 으로 교체하면 라벨·콜론·칸 간격을 보존하며 값을 기입할 수 있다.
    칸과 칸 사이는 **2칸 이상 공백 또는 탭**으로 구분된 것으로 본다(정부양식 표지 정렬 관습).
    """
    for m in _LINE_FIELD_RE.finditer(text or ""):
        yield (m.group("label"), m.group("value"),
               m.start("value"), m.end("value"))


def _is_fill_blank(rest: str) -> bool:
    """콜론 뒤 텍스트가 '빈칸'(공백/밑줄/점/대시만 또는 빈문자)이면 True.

    한 글자라도 실제 글자(한글/영숫자)나 ○·□ 등 비-채움기호가 있으면 False
    → 이미 채워졌거나 마스킹값/문장이므로 덮어쓰지 않는다(오기입·훼손 차단).
    """
    return bool(_FILL_BLANK_RE.match(rest or ""))


def _is_visible_blank(rest: str) -> bool:
    """콜론 뒤 값이 '보이는 빈칸'(밑줄·점·대시 등 실제 채움선이 있음)이면 True.

    **표 셀 인라인 필드 전용** 판정. ``_is_fill_blank`` 는 콜론 뒤가 완전히 비어도
    (공백/빈문자) 빈칸으로 보지만, 표 셀에서 '라벨:'(콜론 뒤 공백)은 **옆 값칸**을
    가리키는 경우와 구별이 안 되므로 인라인으로 오기입하면 안 된다. 그래서 셀 인라인은
    스트립 후에도 비어있지 않은(밑줄 ______ 같은 가시적 채움선) 경우만 인정한다.
    본문 단락 경로(_is_fill_blank)는 '옆 칸' 개념이 없어 기존대로 둔다.
    """
    return bool((rest or "").strip()) and _is_fill_blank(rest)


# 예시 플레이스홀더(=실제로는 채울 빈칸) 판별 — 적대검증 확정 보수판(O마스크 제외).
_PH_DATE_RE = re.compile(r"(?<!\d)\d{4}\s*[.\-/]\s*(\d{1,2})\s*[.\-/]\s*(\d{1,2})\s*\.?(?!\d)")
_PH_DATE_KR_RE = re.compile(r"(?<!\d)\d{4}\s*년\s*(\d{1,2})\s*월\s*(\d{1,2})\s*일?")
_PH_ZERO_QTY_RE = re.compile(
    r"(?<![0-9,])0{2,}\s*(?:억원|천만원|백만원|만원|원|명|건|개|회|년|개월|%|퍼센트)(?![0-9])")
_PH_DUMMY_REG_RE = re.compile(r"(?<!\d)0{2,3}\s*-\s*0{2}\s*-\s*0{4,5}(?!\d)")
# 빈 괄호 예시(우편번호·전화 자리 '( - )'·'( )'·'()') — 괄호 안에 글자·숫자 없이
# 구분기호(하이픈·점·가운뎃점)·공백만. '(주)'·'(041-1234)'처럼 글자·숫자가 있으면
# 매치 안 됨 → 실값 보존. 전각 괄호（）도 포함.
_PH_EMPTY_PAREN_RE = re.compile(r"^[(（]\s*[-–—.·ㅡ~\s]*[)）]$")


def _is_obvious_placeholder(value: str) -> bool:
    """값이 '명백한 예시 플레이스홀더'(=실제로는 채울 빈칸)면 True (보수적).

    3종만 인정: ①불가능 날짜(월 또는 일이 00 — 2025.03.17 같은 정상 날짜 보존)
    ②전부-0 수량(000억원/00건 — '100억원'·'2,000명'은 부정후방탐색으로 배제)
    ③더미 등록번호(000-00-00000 — 실번호 327-29-01754 보존).
    **O마스크(OOO/○○)는 의도적으로 제외** — 영문 실단어(GOOGLE/SOHO/O2O)를 오판해
    실값을 덮어쓰는 경로를 원천 차단(적대검증 치명결함 반영). 0 없는 그럴듯한 실값
    (경기도 성남시 등)은 False → 보존(덮어쓰기 금지).
    """
    v = (value or "").strip()
    if not v:
        return False
    for rx in (_PH_DATE_RE, _PH_DATE_KR_RE):
        m = rx.search(v)
        # 월·일이 '둘 다' 0 인 더미 날짜만(2000.00.00.) — 한쪽만 0(2020.0.5 버전문자열)은 보존
        if m and int(m.group(1)) == 0 and int(m.group(2)) == 0:
            return True
    if _PH_ZERO_QTY_RE.search(v):
        return True
    if _PH_DUMMY_REG_RE.search(v):
        return True
    # 빈 괄호 예시 '( - )'·'( )'·'()' (우편번호·전화 자리) → 채울 빈칸으로 승격.
    if _PH_EMPTY_PAREN_RE.match(v):
        return True
    return False


# 가짜 타깃(채울 라벨로 부적합한 잡음) 판별 — R7 안전핀(클러스터 라벨 절대 보호).
_PURE_INDEX_RE = re.compile(r"^\s*(?:\d{1,3}|[①-⑳]|[Ⅰ-Ⅻ]|\d{1,3}\s*[.)])\s*$")
_ELLIPSIS_ONLY_RE = re.compile(r"^[.．…‥⋯\s]+$")
_EXAMPLE_OZERO_RE = re.compile(
    r"^[Oo0○Ｏ]{3,}(?:\s*(?:대표|과장|부장|차장|대리|사원|팀장|이사|주임))?$")
_NOTICE_PREFIX = ("※", "☞", "주)")


def _is_noise_label(label_text: str, norm: str) -> bool:
    """채울 타깃의 '라벨'로 부적합한 잡음이면 True(표 번호·예시토큰·생략기호·안내문).

    **R7 안전핀**: 동의어 클러스터(_CLUSTER_OF) 등록 라벨은 절대 잡지 않는다
    (진짜 필드 라벨 오탈락 0). 클러스터 미등록 라벨에만 잡음 패턴을 적용한다.
    """
    t = (label_text or "").strip()
    if not t:
        return True
    if _cluster_rep(norm) is not None:        # R7: 진짜 라벨 보호
        return False
    if _PURE_INDEX_RE.match(t):               # 표 번호 '2'/'3'/'①'
        return True
    if _ELLIPSIS_ONLY_RE.match(t):            # '...' 단독
        return True
    if _EXAMPLE_OZERO_RE.match(re.sub(r"\s+", "", t)):   # 'OOO'/'000 대표' 예시
        return True
    if t.startswith(_NOTICE_PREFIX):          # 안내문(※/☞/주))
        return True
    # 주의: '고용 계획(00명)'·'매출 목표(000억원)' 같은 진짜 수량필드 라벨을 드롭하지
    # 않도록, 라벨 텍스트에 _is_obvious_placeholder 를 적용하지 않는다(recall 보호).
    return False


def _cell_blank(cell) -> bool:
    """논리 셀이 '빈칸'(빈 문자열 또는 명백한 예시 플레이스홀더)이면 True(덮어쓰기 대상)."""
    t = (cell.text or "").strip()
    return (not t) or _is_obvious_placeholder(t)


def _row_all_blank(logical) -> bool:
    """행의 모든 논리 셀이 빈칸이면 True('값 대기' 행). 빈 행은 False(보수)."""
    return bool(logical) and all(_cell_blank(c) for c in logical)


def find_target_fields(docx_path: str | Path) -> list[dict[str, Any]]:
    """타깃 DOCX 에서 채울 빈칸을 식별한다(표 칸 + 본문 단락형 빈칸).

    반환 각 항목: {orig_label, normalized, kind, table_index, row, value_cell, para_index}
    - **표 칸**(kind="table"): 라벨 셀 바로 다음 논리 셀이 비어 있으면 후보.
      value_cell 은 라벨 셀 다음 논리 셀 인덱스. para_index 는 -1.
    - **세로형 표 칸**(kind="table"): 헤더 행에 인식 라벨(동의어 클러스터)이 있고
      바로 아래 행이 전부 빈칸이면 아래 칸을 값칸으로 본다(표지형 정보 카드). 이때
      row 는 '아래' 행(ri+1), value_cell 은 열 인덱스 → 표 기입 로직 그대로 재사용.
      로스터(다중 레코드)·병합 어긋남은 제외(보수).
    - **본문 단락형 빈칸**(kind="paragraph"): ``라벨 : ____`` 처럼 콜론 뒤가 빈칸인
      단락. para_index 는 ``doc.paragraphs`` 기준 인덱스, 표 좌표(table_index/row/
      value_cell)는 -1. 콜론 뒤에 실제 값/마스킹(○○○)/문장이 있으면 후보 제외
      (덮어쓰기·훼손 금지). 표 칸 다음에 덧붙이므로 같은 라벨이 표·단락에 모두
      있으면 H2 dedup 에서 표 칸이 우선된다(표가 더 신뢰도 높음).
    - 인접 값칸이 이미 채워져 있으면(빈칸 아님) 후보에서 제외한다(덮어쓰기 금지).
    """
    doc = Document(str(docx_path))
    targets: list[dict[str, Any]] = []
    for ti, table in enumerate(_iter_all_tables(doc)):
        for ri, row in enumerate(table.rows):
            logical = _logical_cells(row)
            for i in range(len(logical) - 1):
                label_text = logical[i].text or ""
                label = _key(label_text)
                if not label:
                    continue
                if _is_noise_label(label_text, label):
                    continue  # 표 번호·예시토큰·안내문 = 가짜 타깃(R7 클러스터 라벨은 보호)
                value_text = (logical[i + 1].text or "").strip()
                if value_text and not _is_obvious_placeholder(value_text):
                    continue  # 실제 값이 있음 → 후보 아님(덮어쓰기 금지)
                # 빈칸 또는 명백한 예시 플레이스홀더(2000.00.00.·000억원) → 채울 후보 승격
                targets.append({
                    "orig_label": SubmittableFiller._norm(label_text),
                    "normalized": label,
                    "kind": "table",
                    "table_index": ti,
                    "row": ri,
                    "value_cell": i + 1,
                    "para_index": -1,
                })

            # 표 셀 '안'의 인라인 "라벨 : ______" 빈칸(개요/표지 박스에 흔함).
            # doc.paragraphs 는 셀 단락을 포함하지 않아 아래 본문 스캔이 못 잡는다 →
            # 각 논리 셀의 단락을 직접 훑어 '보이는 빈칸' 인라인 필드만 채운다(멀티필드).
            for ci, cell in enumerate(logical):
                for cpi, cpara in enumerate(cell.paragraphs):
                    for label_raw, value_raw, fstart, fend in _iter_line_fields(cpara.text or ""):
                        label = _key(label_raw)
                        if not label:
                            continue
                        # 셀 인라인은 '보이는 빈칸'(밑줄/점/대시)만 — 콜론 뒤 공백/빈값은
                        # 옆 값칸 패턴과 모호하므로 제외(오기입 방지).
                        if not _is_visible_blank(value_raw):
                            continue
                        targets.append({
                            "orig_label": SubmittableFiller._norm(label_raw),
                            "normalized": label,
                            "kind": "cell_paragraph",
                            "table_index": ti,
                            "row": ri,
                            "value_cell": ci,
                            "para_index": -1,
                            "cell_para_index": cpi,
                            "fill_start": fstart,
                            "fill_end": fend,
                        })

        # 세로형(라벨 위 / 값 아래) 정보 카드: 헤더 행에 인식 라벨이 있고 바로 아래
        # 행이 '전부 빈칸'이면, 각 열의 아래 칸을 그 라벨의 값칸으로 본다. 가로 스캔
        # (라벨|값)이 못 잡는 표지형 카드(성명·생년월일·연락처를 위에 나열하고 그 아래
        # 줄에 값을 적는 구조)를 채운다. value_cell 은 '아래' 행(row=ri+1)의 열 인덱스라
        # 기존 표 기입 로직이 그대로 재사용된다(kind="table").
        # 안전 가드(오매칭<빈칸): ①동의어 클러스터에 등록된 인식 라벨만 ②아래 행이
        # 전부 빈칸일 때만(값이 하나라도 있으면 로스터/데이터 행) ③아래에 빈 행이 또
        # 이어지면 다중 레코드 로스터로 보고 제외(성명 헤더로 여러 명 명부에 일괄기입 방지)
        # ④병합으로 위·아래 논리셀 수가 어긋나면 세로 정렬을 보장할 수 없어 제외.
        # 가로 타깃 뒤에 추가 → 같은 라벨은 H2 dedup 에서 가로가 우선(더 신뢰도 높음).
        rows = table.rows
        for ri in range(len(rows) - 1):
            header = _logical_cells(rows[ri])
            below = _logical_cells(rows[ri + 1])
            if not header or len(header) != len(below):
                continue
            recognized = [
                ci for ci in range(len(header))
                if (header[ci].text or "").strip()
                and _cluster_rep(_key(header[ci].text)) is not None
            ]
            if not recognized:
                continue
            if not _row_all_blank(below):
                continue  # 아래 행에 값 존재 → 세로 카드 아님(로스터/데이터 행 보호)
            if ri + 2 < len(rows) and _row_all_blank(_logical_cells(rows[ri + 2])):
                continue  # 아래 빈 행이 연속 → 다중 레코드 로스터 → 세로 채움 금지
            for ci in recognized:
                if not _cell_blank(below[ci]):
                    continue
                label_text = header[ci].text or ""
                targets.append({
                    "orig_label": SubmittableFiller._norm(label_text),
                    "normalized": _key(label_text),
                    "kind": "table",
                    "table_index": ti,
                    "row": ri + 1,       # 값칸은 헤더 '아래' 행
                    "value_cell": ci,
                    "para_index": -1,
                })

    # 본문 단락형 빈칸: 한 줄의 각 "라벨 : <빈칸>" 칸(멀티필드 대응).
    # 표 칸 다음에 추가 → H2 dedup 에서 표 우선. 같은 줄 여러 칸은 각각 타깃이 된다.
    for pi, para in enumerate(doc.paragraphs):
        for label_raw, value_raw, fstart, fend in _iter_line_fields(para.text or ""):
            label = _key(label_raw)
            if not label:
                continue
            if not _is_fill_blank(value_raw):
                continue  # 이미 채워짐 / 마스킹 / 문장 → 후보 아님
            targets.append({
                "orig_label": SubmittableFiller._norm(label_raw),
                "normalized": label,
                "kind": "paragraph",
                "table_index": -1,
                "row": -1,
                "value_cell": -1,
                "para_index": pi,
                "fill_start": fstart,
                "fill_end": fend,
            })
    return targets


# --- 매칭: 타깃 라벨 ↔ 소스 라벨 ----------------------------------------------

def _tokens(norm_label: str) -> set[str]:
    """정규화 라벨을 문자 단위 토큰 집합으로(한국어 라벨은 짧아 문자 자카드가 안정적)."""
    return set(norm_label)


def _jaccard(a: set[str], b: set[str]) -> float:
    if not a or not b:
        return 0.0
    inter = len(a & b)
    union = len(a | b)
    return inter / union if union else 0.0


_FUZZY_THRESHOLD = 0.6


def _best_source_for_target(
    norm_target: str, source: dict[str, str]
) -> tuple[Optional[str], str, list[str]]:
    """타깃 정규화 라벨에 대한 best 소스 라벨을 보수적으로 고른다.

    반환: (선택 소스라벨 or None, confidence, 후보 리스트)
    - **자동 전사(high)는 ①정규화 정확일치 ②동의어 같은 클러스터 단일후보 만.**
    - 동의어 클러스터에 소스 후보가 2개 이상이면 ``"conflict"`` (자동전사 금지).
    - 정확일치·동의어가 없으면 퍼지(자카드·부분문자열)로 **제안만** → ``"fuzzy"``
      (자동전사 금지, needs_confirm). 접미사 다른 합성어(사업명/사업자명, 주소/주소지정)는
      클러스터가 다르므로 high 가 될 수 없고, 퍼지로만 제안된다.
    - 아무 후보도 없으면 ``"low"`` (미매칭).
    """
    # ① 정확일치 → high (단일)
    if norm_target in source:
        return norm_target, "high", []

    # ② 동의어 클러스터
    rep = _cluster_rep(norm_target)
    if rep is not None:
        syn_hits = [s for s in source if _cluster_rep(s) == rep]
        if len(syn_hits) == 1:
            return syn_hits[0], "high", []
        if len(syn_hits) > 1:
            return None, "conflict", syn_hits  # 충돌 → 자동전사 금지

    # ③ 퍼지: 자카드 ≥ 임계 또는 포함+길이차 작음 → 제안만(fuzzy), 자동전사 금지
    t_tokens = _tokens(norm_target)
    scored: list[tuple[float, str]] = []
    for s in source:
        s_tokens = _tokens(s)
        jac = _jaccard(t_tokens, s_tokens)
        contains = (norm_target in s or s in norm_target)
        len_close = abs(len(norm_target) - len(s)) <= 2
        if jac >= _FUZZY_THRESHOLD or (contains and len_close):
            scored.append((jac, s))
    if not scored:
        return None, "low", []
    scored.sort(key=lambda x: x[0], reverse=True)
    top_score = scored[0][0]
    top = [s for sc, s in scored if abs(sc - top_score) < 1e-9]
    return None, "fuzzy", top  # 단일이든 동점이든 자동전사 금지(제안만)


def match_fields(
    source: dict[str, str],
    targets: list[dict[str, Any]],
    source_originals: Optional[dict[str, str]] = None,
) -> list[Match]:
    """각 타깃에 best 소스 매칭을 산출한다.

    **high 단일 후보만** Match(전사 대상: source_label/value 채움)로 만든다.
    그 외(conflict·fuzzy·low)는 confidence 를 명확히 구분하고 ``value=""`` 로 표시 →
    autofill 단계가 needs_confirm/unmatched 로 분리한다. (충돌·퍼지가 confident 로
    새지 않도록 source_label 도 비운다.)

    추가 강등(오매칭은 빈칸보다 나쁘다):
    - **H2 중복 타깃**: 동일 정규화 타깃 라벨이 2개 이상이면 첫 1회만 high,
      나머지는 ``"duplicate"`` 로 강등(같은 값을 N칸에 복제 전사 방지).
    - **H3 괄호 토큰 충돌**: 정확일치라도 소스/타깃 원본 라벨의 괄호 토큰이
      서로 다르면(금액(국고) vs 금액(자부담)) ``"bracket"`` 로 강등(false high 방지).
    """
    source_originals = source_originals or {}
    matches: list[Match] = []
    high_seen: set[str] = set()  # 이미 high 전사된 정규화 타깃(H2 dedup)

    for tgt in targets:
        norm = tgt["normalized"]
        src_label, conf, candidates = _best_source_for_target(norm, source)

        demote_to: Optional[str] = None
        if conf == "high" and src_label is not None:
            # H3: 괄호 토큰이 다르면 high 금지(정확일치라도 다른 항목)
            if _bracket_conflict(source_originals.get(src_label, src_label),
                                 tgt["orig_label"]):
                demote_to = "bracket"
            # H2: 동일 정규화 타깃이 이미 high 전사됐으면 둘째부터 강등
            elif norm in high_seen:
                demote_to = "duplicate"
            # 값-타입: 이름(성명)류 타깃에 이름 모양 아닌 값(역할서술 등)이면 high 금지
            #          → needs_confirm 으로 노출(역할분담 서술 오전사 차단, 오매칭<빈칸)
            elif _is_name_field(norm) and not _looks_like_name(source[src_label]):
                demote_to = "value_type"

        if conf == "high" and src_label is not None and demote_to is None:
            high_seen.add(norm)
            matches.append(Match(
                target_label=tgt["orig_label"],
                normalized=norm,
                source_label=src_label,
                value=source[src_label],
                confidence="high",
                table_index=tgt["table_index"],
                row=tgt["row"],
                value_cell=tgt["value_cell"],
                kind=tgt.get("kind", "table"),
                para_index=tgt.get("para_index", -1),
                cell_para_index=tgt.get("cell_para_index", -1),
                fill_start=tgt.get("fill_start", -1),
                fill_end=tgt.get("fill_end", -1),
            ))
        else:
            # conflict / fuzzy / low / duplicate / bracket → 전사 보류.
            # source_label/value 모두 비움. 후보만 candidates 에 남긴다.
            final_conf = demote_to if demote_to is not None else conf
            cand = list(candidates)
            if demote_to is not None and src_label is not None and not cand:
                cand = [src_label]
            matches.append(Match(
                target_label=tgt["orig_label"],
                normalized=norm,
                source_label="",
                value="",
                confidence=final_conf,
                table_index=tgt["table_index"],
                row=tgt["row"],
                value_cell=tgt["value_cell"],
                kind=tgt.get("kind", "table"),
                para_index=tgt.get("para_index", -1),
                cell_para_index=tgt.get("cell_para_index", -1),
                fill_start=tgt.get("fill_start", -1),
                fill_end=tgt.get("fill_end", -1),
                candidates=cand,
            ))
    return matches


# --- 체크박스(선택칸) 자동 체크 -----------------------------------------------
# 정부양식의 "사업자 형태 | □ 개인 | □ 법인" 같은 선택칸을 소스 값으로 채운다.
# 안전 원칙(텍스트 전사와 동일): 오매칭은 빈칸보다 나쁘다 → **정확히 한 옵션에만
# 매칭될 때만** 체크. 모호(0개·2개+)·소스 무값이면 보류(자동 체크 금지, 날조 0).

# 빈(미선택) 체크박스 / 이미 체크된 박스 기호 집합. 사용자 선택: □ → ■.
_EMPTY_BOX_CHARS = "□☐▢◻◽⬚〼"
_CHECKED_BOX_CHARS = "■☑☒✓✔◼◾🗹"
_EMPTY_BOX_RE = re.compile("[" + _EMPTY_BOX_CHARS + "]")
_ANY_BOX_RE = re.compile("[" + _EMPTY_BOX_CHARS + _CHECKED_BOX_CHARS + "]")
_CHECK_MARK = "■"  # □ 를 이 기호로 치환(사용자 확정)

# 선택칸 값↔옵션 정규화 사전(알려진 동의어만 표준형으로 환원; 모르면 그대로 둔다).
# **부분문자열 매칭을 쓰지 않는 이유**: '개인'(옵션) in '개인정보보호'(값) / '법인' in
# '법인영업…' 처럼 옵션명이 무관한 값의 일부면 오체크된다(적대검증 HIGH 3건). 그래서
# 값·옵션을 이 사전으로 환원한 뒤 **정확일치만** 인정한다 — 서술형·짧은 절단값·예시값은
# 사전에 없어 자동으로 보류된다(오매칭 < 빈칸·날조 0 불변).
_CHOICE_VALUE_MAP: dict[str, str] = {
    "개인": "개인", "개인사업자": "개인", "개인기업": "개인",
    "법인": "법인", "법인사업자": "법인", "법인기업": "법인",
    "주식회사": "법인", "㈜": "법인", "유한회사": "법인",
    "유한책임회사": "법인", "합자회사": "법인", "합명회사": "법인",
}


def _normalize_choice(norm_key: str) -> str:
    """선택칸 값/옵션 정규화 키를 비교용 표준형으로 환원(알려진 동의어만)."""
    return _CHOICE_VALUE_MAP.get(norm_key, norm_key)


def _option_text(cell_text: str) -> str:
    """체크박스 기호를 제거한 옵션 라벨 텍스트('□ 개인' → '개인')."""
    return _ANY_BOX_RE.sub("", (cell_text or "").strip()).strip()


def _is_option_cell(cell_text: str) -> bool:
    """셀이 '체크박스 1개 + 옵션명' 형태면 True.

    - 박스가 **정확히 1개**여야 한다(0개=일반 라벨/값, 2개+=한 셀 복수박스는
      위치 모호로 보류·미지원). 박스 외에 실제 글자(옵션명)가 있어야 한다
      ('□' 단독 = 매트릭스 체크칸은 옵션 아님).
    """
    s = (cell_text or "").strip()
    if len(_ANY_BOX_RE.findall(s)) != 1:
        return False
    return bool(_HAS_WORD_RE.search(_ANY_BOX_RE.sub("", s)))


def find_checkbox_targets(docx_path: str | Path) -> list[dict[str, Any]]:
    """타깃 DOCX 표에서 '라벨 + 연속 □옵션' 선택칸 그룹을 탐지한다.

    반환 각 항목: {label, normalized, table_index, row, options:[{cell_index,
    text, normalized}]}. 라벨은 옵션 그룹 바로 앞의 비-옵션 논리셀이다. 옵션이
    행 맨 앞이라 라벨이 없거나 옵션이 2개 미만이면 그룹으로 보지 않는다(보수적).
    """
    doc = Document(str(docx_path))
    groups: list[dict[str, Any]] = []
    for ti, table in enumerate(_iter_all_tables(doc)):
        for ri, row in enumerate(table.rows):
            logical = _logical_cells(row)
            n = len(logical)
            i = 0
            while i < n:
                if not _is_option_cell(logical[i].text):
                    i += 1
                    continue
                # 옵션 그룹 시작 — 직전 비옵션 셀이 라벨.
                label_text = logical[i - 1].text if i - 1 >= 0 else ""
                options: list[dict[str, Any]] = []
                j = i
                while j < n and _is_option_cell(logical[j].text):
                    opt = _option_text(logical[j].text)
                    options.append({
                        "cell_index": j,
                        "text": opt,
                        "normalized": _key(opt),
                    })
                    j += 1
                label_norm = _key(label_text)
                if label_norm and len(options) >= 2:
                    groups.append({
                        "label": SubmittableFiller._norm(label_text),
                        "normalized": label_norm,
                        "table_index": ti,
                        "row": ri,
                        "options": options,
                    })
                i = j
    return groups


def match_checkbox_groups(
    source: dict[str, str],
    groups: list[dict[str, Any]],
    source_originals: Optional[dict[str, str]] = None,
) -> list[dict[str, Any]]:
    """각 선택칸 그룹에 대해 체크할 옵션을 보수적으로 결정한다.

    - 그룹 라벨을 ``_best_source_for_target`` 로 소스 매칭(**high 단일후보만** 값 사용).
    - 값·옵션을 ``_normalize_choice`` 로 환원한 뒤 **정확일치**가 정확히 하나면 체크
      (checked_option_index). 부분문자열 매칭은 쓰지 않는다 — '개인'이 '개인정보보호'의
      일부거나 짧은 절단값('소')이 '중소기업'의 일부면 오체크되기 때문(오매칭 < 빈칸).
    - 소스 값이 예시 플레이스홀더(``_is_obvious_placeholder``)면 체크 금지(날조 0).
    - 0개/2개+/무값/플레이스홀더면 -1(보류).
    반환 각 항목: 그룹 정보 + {source_label, source_value, checked_option_index,
    confidence("high"|"ambiguous"|"no_match"|"no_source"|"placeholder")}.
    """
    results: list[dict[str, Any]] = []
    for g in groups:
        label_norm = g["normalized"]
        src_label, conf, _cands = _best_source_for_target(label_norm, source)
        value = source.get(src_label, "") if (conf == "high" and src_label) else ""

        checked_idx = -1
        if value and _is_obvious_placeholder(value):
            decision = "placeholder"         # 예시값(00법인 등) → 체크 금지(날조 0)
        elif value:
            vnorm = _normalize_choice(_key(value))
            hits: list[int] = []
            for oi, opt in enumerate(g["options"]):
                okey = opt["normalized"]
                if okey and _normalize_choice(okey) == vnorm:  # 정확일치만(환원 후)
                    hits.append(oi)
            if len(hits) == 1:
                checked_idx = hits[0]
                decision = "high"
            elif hits:
                decision = "ambiguous"       # 2개+ 매칭 → 보류
            else:
                decision = "no_match"        # 소스 값이 어느 옵션과도 정확일치 안 함
        else:
            decision = "no_source"           # 소스에 그 라벨 값이 없음(날조 0)

        results.append({
            "label": g["label"],
            "normalized": label_norm,
            "table_index": g["table_index"],
            "row": g["row"],
            "options": g["options"],
            "source_label": src_label if value else "",
            "source_value": value,
            "checked_option_index": checked_idx,
            "confidence": "high" if checked_idx >= 0 else decision,
        })
    return results


def _check_option_cell(cell) -> bool:
    """셀의 빈 체크박스(첫 1개)를 ■ 로 치환한다. 옵션 글자·서식은 보존.

    이미 체크된 박스(■/☑ 등)거나 빈 박스가 없으면 무변경(멱등). 반환=변경 여부.
    run 단위로 치환해 옵션 텍스트("개인")와 다른 run 서식을 건드리지 않는다.
    ``para.runs`` 가 아니라 ``.//w:r`` 전체를 도므로 <w:hyperlink>/필드로 감싸인
    run 안의 □ 도 잡는다(R11 선례 — 검출 cell.text 와 기입 범위 일치, 조용한 no-op 방지).
    """
    from docx.oxml.ns import qn
    from docx.text.run import Run

    for para in cell.paragraphs:
        for r_el in para._p.findall(".//" + qn("w:r")):
            run = Run(r_el, para)
            if _EMPTY_BOX_RE.search(run.text):
                run.text = _EMPTY_BOX_RE.sub(_CHECK_MARK, run.text, count=1)
                return True
    return False


# --- 사용자 확정(needs_confirm 적용) -----------------------------------------

def _normalize_confirmations(
    confirmations: Optional[dict[str, str]],
) -> dict[str, str]:
    """사용자 확정 ``{타깃라벨: 소스라벨}`` 을 정규화 키 맵으로 변환한다.

    양쪽 라벨을 ``_key`` 로 정규화하므로 사용자가 보고서에 보이는 라벨을 그대로
    (장식·공백 포함) 복사해 넣어도 매칭된다. 키/값이 비면 버린다(잡음 차단).
    """
    norm: dict[str, str] = {}
    for raw_t, raw_s in (confirmations or {}).items():
        tk = _key(str(raw_t))
        sk = _key(str(raw_s))
        if tk and sk:
            norm[tk] = sk
    return norm


def _apply_confirmations(
    all_matches: list[Match],
    confident: list[Match],
    src_fields: dict[str, str],
    conf_norm: dict[str, str],
) -> tuple[list[Match], list[str]]:
    """비-high 매칭 중 사용자 확정된 타깃을 소스 실값으로 승격한다.

    - 확정 ``{타깃norm: 소스norm}`` 에서 소스 라벨이 소스에 실제 값으로 존재할 때만
      승격(날조 0). 없으면 채우지 않고 notes 로 알린다.
    - 승격된 Match 는 confidence="confirmed", source_label/value 가 채워진다(부수효과로
      해당 Match 객체를 직접 갱신 → 호출부의 needs_confirm 집계에서 자동 제외됨).
    - **동일 정규화 타깃 빈칸이 복수면 모두 채운다**(high 경로의 H2 dedup과 다름):
      사용자가 그 라벨을 명시적으로 확정했으므로 같은 소스 실값을 N칸에 채운다.
    - 반환: (승격된 Match 리스트, 진단 notes).
    """
    confirmed_matches: list[Match] = []
    notes: list[str] = []
    confident_ids = {id(m) for m in confident}
    used: set[str] = set()

    for m in all_matches:
        if id(m) in confident_ids:
            continue
        src_key = conf_norm.get(m.normalized)
        if src_key is None:
            continue
        value = src_fields.get(src_key, "")
        if not value:  # 소스에 그 라벨/값이 없음 → 날조 금지
            if m.normalized not in used:
                notes.append(
                    f"확정 무시('{m.normalized}'→'{src_key}'): 소스에 해당 값이 없음")
            used.add(m.normalized)
            continue
        m.source_label = src_key
        m.value = value
        m.confidence = "confirmed"
        confirmed_matches.append(m)
        used.add(m.normalized)

    # 어떤 빈칸 타깃에도 적용되지 않은 확정은 사용자에게 알린다.
    for tk in conf_norm:
        if tk not in used:
            notes.append(f"확정 미적용('{tk}'): 일치하는 빈칸 타깃이 없음")
    return confirmed_matches, notes


# --- 단락형 빈칸 기입 ---------------------------------------------------------

def _fill_paragraph_fields(para, matches: list[Match]) -> list[Match]:
    """한 본문 단락의 여러 "라벨 : <빈칸>" 칸에 값을 **한 번에** 기입한다.

    각 Match 의 빈칸 구간 ``[fill_start, fill_end)`` 를 ``" " + 값`` 으로 교체한다
    (라벨·콜론은 보존, 콜론 뒤 정렬공백/빈칸 표기는 한 칸 공백으로 정규화, 칸 사이
    간격은 그대로 보존). 여러 칸을 채울 때 인덱스가 밀리지 않도록 **오른쪽 칸부터**
    (fill_start 내림차순) 교체한다. 교체된 전체 줄 텍스트를 첫 run 에 쓰고 나머지 run
    텍스트는 비운다(set_cell_text 와 동일한 보수적 방식 — run 서식 유지).
    값이 비었거나 구간이 현재 단락 텍스트 범위를 벗어나면 그 칸은 건너뛴다(안전).
    반환: 실제 기입한 Match 리스트(전사/확정 카운트용).
    """
    text = para.text or ""
    n = len(text)
    valid = [
        m for m in matches
        if 0 <= m.fill_start <= m.fill_end <= n and str(m.value)
    ]
    if not valid:
        return []
    written = sorted(valid, key=lambda m: m.fill_start, reverse=True)
    new_text = text
    for m in written:
        new_text = (new_text[: m.fill_start]
                    + " " + str(m.value)
                    + new_text[m.fill_end:])
    runs = para.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""
    else:
        para.add_run(new_text)
    return written


# --- 변환 헬퍼(HWP) -----------------------------------------------------------

def _to_docx_if_needed(path: Path, tmpdir: Path, report: AutofillReport) -> Optional[Path]:
    """입력이 .hwp/.hwpx 면 DOCX 로 변환해 그 경로를 반환. 실패 시 None."""
    if path.suffix.lower() not in _HWP_EXTS:
        return path
    from .hwp_docx_convert import hwp_to_docx

    out = tmpdir / (path.stem + "_in.docx")
    rep = hwp_to_docx(path, out)
    if rep.ok:
        report.notes.append(f"HWP 변환({path.name} → DOCX, {rep.method})")
        return out
    report.notes.append(f"HWP 변환 실패({path.name}): {'; '.join(rep.notes)}")
    return None


# --- 진입점 -------------------------------------------------------------------

def autofill_from_source(
    source: str | Path,
    target: str | Path,
    out: str | Path,
    *,
    use_ai: bool = False,
    confirmations: Optional[dict[str, str]] = None,
    enable_checkbox: bool = True,
) -> AutofillReport:
    """소스 A 의 값을 타깃 B 의 빈 칸에 전사해 out 으로 저장한다(원본 미수정).

    절차: (HWP→DOCX 변환) → extract(소스) → find(타깃 빈칸) → match(보수)
          → **high 매칭** + **사용자 확정(confirmations) 매칭** 을 해당 값셀에
            set_cell_text 로 직접 기입(잔여물/안내 청소 없음)
          → (out 이 .hwp 면 docx_to_hwp).
    잔여물/안내 청소 패스를 돌리지 않으므로 ``○○○``·``OOO-OO-OOOOO`` 같은 전사값이 보존된다.
    transcribed 는 실제 저장 문서에서 비어있지 않게 기입된 셀 수다.

    confirmations
    -------------
    ``{타깃 라벨: 소스 라벨}`` 형태의 사용자 확정 맵. high 자동전사로는 보류된
    퍼지/충돌 후보(needs_confirm)를, 사용자가 고른 소스 라벨의 **실값**으로 채운다.
    - 양쪽 라벨은 정규화(_key)되므로 보고서에 보이는 라벨을 그대로(장식·공백 포함)
      복사해 넣어도 된다.
    - **날조 0**: 확정이 가리키는 소스 라벨이 실제 값으로 존재할 때만 채운다.
      없으면 채우지 않고 notes 로 알린다(없는 값을 지어내지 않는다).
    - 한계: 라벨이 정규화 기준으로 같으면(예: 괄호 토큰만 다른 '금액(국고)'/'금액(자부담)')
      소스에서 먼저 추출된 값이 쓰인다(사용자가 확정 시 책임).
    """
    source = Path(source)
    target = Path(target)
    out = Path(out)

    # 원본 보호: 출력이 소스/타깃과 같으면 거부(중간본은 tempfile)
    if source.exists() and out.resolve() == source.resolve():
        raise ValueError("출력 경로가 소스와 같습니다. 원본 덮어쓰기는 금지입니다.")
    if target.exists() and out.resolve() == target.resolve():
        raise ValueError("출력 경로가 타깃과 같습니다. 원본 덮어쓰기는 금지입니다.")

    report = AutofillReport(source=str(source), target=str(target), output=str(out))

    if not source.exists():
        raise FileNotFoundError(f"소스 파일이 없습니다: {source}")
    if not target.exists():
        raise FileNotFoundError(f"타깃 파일이 없습니다: {target}")
    if use_ai:
        report.notes.append("use_ai=True 는 v1 에서 미지원 — 결정론 매칭으로 진행")

    # H7: 지원 확장자 화이트리스트(.docx/.hwp/.hwpx). PDF/xlsx/위장파일은 보수적 거부.
    for role, p in (("소스", source), ("타깃", target)):
        if p.suffix.lower() not in _SUPPORTED_EXTS:
            report.ok = False
            report.notes.append(
                f"{role} 비지원 확장자({p.suffix or '없음'}) — .docx/.hwp/.hwpx 만 지원")
            return report

    # M6: 출력 상위 폴더가 없으면 자동 생성(중첩 경로 저장 시 크래시 방지)
    if out.parent and not out.parent.exists():
        out.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix="xform_") as td:
        tmpdir = Path(td)

        src_docx = _to_docx_if_needed(source, tmpdir, report)
        tgt_docx = _to_docx_if_needed(target, tmpdir, report)
        if src_docx is None or tgt_docx is None:
            report.ok = False
            return report

        # H7: 손상/위장 DOCX(PackageNotFoundError)는 raw traceback 대신 ok=False+notes.
        try:
            src_fields, src_originals = _extract_source(src_docx)
            tgt_fields = find_target_fields(tgt_docx)
        except PackageNotFoundError as exc:
            report.ok = False
            report.notes.append(f"DOCX 로딩 실패(손상/위장 파일 가능): {exc}")
            return report

        report.src_fields = len(src_fields)
        report.tgt_fields = len(tgt_fields)
        all_matches = match_fields(src_fields, tgt_fields, src_originals)

        # confident: confidence=="high" AND value 비어있지 않음 (충돌·퍼지·low 제외)
        confident = [
            m for m in all_matches
            if m.confidence == "high" and m.source_label and m.value
        ]

        # 사용자 확정(needs_confirm 적용): 비-high 매칭을 소스 실값으로 승격(날조 0).
        conf_norm = _normalize_confirmations(confirmations)
        confirmed_matches, confirm_notes = _apply_confirmations(
            all_matches, confident, src_fields, conf_norm)
        report.notes.extend(confirm_notes)

        # 전사 대상 = high 자동 + 사용자 확정. 그 외만 needs_confirm/unmatched 로 분리.
        to_write = confident + confirmed_matches
        write_ids = {id(m) for m in to_write}
        for m in all_matches:
            if id(m) in write_ids:
                continue
            if m.confidence == "low":
                report.unmatched_targets.append({
                    "target_label": m.target_label,
                    "normalized": m.normalized,
                })
            else:
                # conflict / fuzzy → 사람 확인 필요(자동전사 금지)
                report.needs_confirm.append({
                    "target_label": m.target_label,
                    "normalized": m.normalized,
                    "candidates": list(m.candidates),
                    "confidence": m.confidence,
                })

        # 전사 출력은 DOCX 중간본(tmp). out 이 .docx 면 바로, .hwp 면 변환.
        out_is_hwp = out.suffix.lower() in _HWP_EXTS
        docx_out = (tmpdir / (out.stem + "_filled.docx")) if out_is_hwp else out

        # 직접 쓰기: 타깃을 열어 (table_index,row,value_cell) 좌표에 set_cell_text 로
        # 값을 직접 기입한다. 잔여물/안내 청소 패스를 돌리지 않으므로 ○○○·OOO-… 보존.
        doc = Document(str(tgt_docx))
        # 탐지(find_target_fields/find_checkbox_targets)와 **동일한 평탄화 순서**로
        # 표를 나열해 table_index 정합을 보장한다(중첩표 포함). 루프마다 재계산하지 않게
        # 한 번만 만든다.
        all_tables = _iter_all_tables(doc)
        transcribed = 0
        confirmed_written = 0
        para_groups: dict[int, list[Match]] = {}
        # 표 셀 '안' 인라인 빈칸: (표, 행, 논리셀, 셀단락) 별로 모아 한 번에 기입.
        cell_para_groups: dict[tuple[int, int, int, int], list[Match]] = {}
        for m in to_write:
            # 본문 단락형 빈칸: 같은 단락의 여러 칸을 모아 한 번에 기입(인덱스 밀림 방지)
            if m.kind == "paragraph":
                para_groups.setdefault(m.para_index, []).append(m)
                continue
            if m.kind == "cell_paragraph":
                cell_para_groups.setdefault(
                    (m.table_index, m.row, m.value_cell, m.cell_para_index), []
                ).append(m)
                continue
            if m.table_index >= len(all_tables):
                report.notes.append(f"전사 표 범위초과 ti={m.table_index}")
                continue
            table = all_tables[m.table_index]
            if m.row >= len(table.rows):
                report.notes.append(f"전사 행 범위초과 ti={m.table_index} ri={m.row}")
                continue
            logical = _logical_cells(table.rows[m.row])
            if m.value_cell >= len(logical):
                report.notes.append(
                    f"전사 셀 범위초과 ti={m.table_index} ri={m.row} ci={m.value_cell}")
                continue
            cur = (logical[m.value_cell].text or "").strip()
            if cur and not _is_obvious_placeholder(cur):
                # 2중 게이트: 빈칸도 예시 플레이스홀더도 아닌 실값이면 덮어쓰기 금지(불변원칙)
                report.notes.append(
                    f"전사 보류(실값 보존) ti={m.table_index} ri={m.row} ci={m.value_cell}")
                continue
            set_cell_text(logical[m.value_cell], str(m.value))
            transcribed += 1
            if m.confidence == "confirmed":
                confirmed_written += 1

        # 본문 단락형 빈칸: 같은 단락의 여러 칸을 한 번에 기입(라벨·콜론·칸 간격 보존)
        for pi, group in para_groups.items():
            if pi < 0 or pi >= len(doc.paragraphs):
                report.notes.append(f"전사 단락 범위초과 pi={pi}")
                continue
            for m in _fill_paragraph_fields(doc.paragraphs[pi], group):
                transcribed += 1
                if m.confidence == "confirmed":
                    confirmed_written += 1

        # 표 셀 '안' 인라인 빈칸: 셀 단락을 찾아 본문 단락과 동일하게 한 번에 기입.
        # (표-라벨 경로는 '빈/예시' 값셀만 덮으므로, 글자가 든 인라인 셀과 겹치지 않는다.)
        for (ti, ri, ci, cpi), group in cell_para_groups.items():
            if ti < 0 or ti >= len(all_tables):
                report.notes.append(f"전사 셀단락 표범위초과 ti={ti}")
                continue
            rows = all_tables[ti].rows
            if ri < 0 or ri >= len(rows):
                report.notes.append(f"전사 셀단락 행범위초과 ti={ti} ri={ri}")
                continue
            logical = _logical_cells(rows[ri])
            if ci < 0 or ci >= len(logical):
                report.notes.append(f"전사 셀단락 셀범위초과 ti={ti} ri={ri} ci={ci}")
                continue
            cell_paras = logical[ci].paragraphs
            if cpi < 0 or cpi >= len(cell_paras):
                report.notes.append(f"전사 셀단락 단락범위초과 ti={ti} ri={ri} ci={ci} cpi={cpi}")
                continue
            for m in _fill_paragraph_fields(cell_paras[cpi], group):
                transcribed += 1
                if m.confidence == "confirmed":
                    confirmed_written += 1

        # 체크박스(선택칸) 자동 체크: 사업자형태 등 '□옵션' → ■ (보수적, high 단일매칭만).
        # 좌표(table_index/row/cell_index)는 tgt_docx 와 doc 가 같은 파일이라 일치한다.
        # 체크박스 처리 실패가 텍스트 전사 결과를 무효화하지 않도록 예외를 가둔다.
        checkbox_checked = 0
        if enable_checkbox:
            try:
                groups = find_checkbox_targets(tgt_docx)
                decided = match_checkbox_groups(groups and src_fields or {},
                                                groups, src_originals)
                report.checkbox_groups = [
                    {k: d[k] for k in ("label", "normalized", "source_value",
                                       "checked_option_index", "confidence")}
                    for d in decided
                ]
                for d in decided:
                    ci = d["checked_option_index"]
                    if ci < 0:
                        continue
                    if d["table_index"] >= len(all_tables):
                        continue
                    rows = all_tables[d["table_index"]].rows
                    if d["row"] >= len(rows):
                        continue
                    logical = _logical_cells(rows[d["row"]])
                    cell_idx = d["options"][ci]["cell_index"]
                    if cell_idx >= len(logical):
                        continue
                    if _check_option_cell(logical[cell_idx]):
                        checkbox_checked += 1
            except Exception as exc:  # noqa: BLE001 — 전사 결과 보호
                report.notes.append(f"체크박스 처리 건너뜀: {exc}")

        doc.save(str(docx_out))

        report.matches = to_write
        report.transcribed = transcribed
        report.confirmed = confirmed_written
        report.checkbox_checked = checkbox_checked

        # H6: 전사 0건/타깃 빈칸 0건은 "성공"으로 오인하지 않도록 보수적 처리.
        #     파일 저장은 됐어도 전사 성과가 없으므로 ok=False + 진단 경고.
        #     단 체크박스 자동 체크가 있었으면 성과로 인정한다(checkbox_checked>0).
        nothing_done = (
            (transcribed == 0 and checkbox_checked == 0)
            or (report.tgt_fields == 0 and checkbox_checked == 0)
        )
        if nothing_done:
            if report.tgt_fields == 0:
                report.notes.append(
                    "타깃에서 '라벨|빈칸' 구조를 찾지 못함(세로형/단일열/마스킹 등 구조 불일치 가능)")
            else:
                report.notes.append(
                    f"전사 0건 — 소스 {report.src_fields}필드/타깃 {report.tgt_fields}빈칸이나 "
                    f"보수적 매칭(high)에서 자동전사 대상 없음")

        if out_is_hwp:
            from .hwp_docx_convert import docx_to_hwp

            conv = docx_to_hwp(docx_out, out)
            if conv.ok:
                report.ok = not nothing_done
            else:
                # COM 미가용 등 — DOCX 보존(예외 전파 금지)
                fallback = out.with_suffix(".docx")
                Document(str(docx_out)).save(str(fallback))
                report.output = str(fallback)
                report.ok = False
                report.notes.append(
                    f"HWP 출력 실패 — DOCX 보존: {fallback.name} ({'; '.join(conv.notes)})")
        else:
            report.ok = not nothing_done

    return report


# --- 배치: 공고 폴더 양식 일괄 채우기 + HWP 변환 --------------------------------

_FORM_GLOB_PATTERNS = ("*.docx", "*.hwp", "*.hwpx")

_DEFAULT_SOURCE_KEYWORDS: tuple[str, ...] = (
    "사업계획서", "사업계획", "신청서", "완성", "k-global", "kglobal", "star",
    "제출", "작성", "이력서", "경영지도", "박다솜",
)

_RESUME_BONUS_KEYWORD = "이력서"
_RESUME_PENALTY_KEYWORDS: tuple[str, ...] = ("신청서", "동의서", "추천서")

_NON_FORM_NAME_KEYWORDS: tuple[str, ...] = (
    "공고", "공고문", "모집", "안내", "포스터", "poster", "붙임", "첨부",
    "download", "manifest",
)

_BATCH_OUTPUT_SKIP_DIRS = frozenset({"filled", "filled_out", "__pycache__"})


@dataclass
class BatchAutofillItem:
    target: str
    source: str
    output: str
    hwp_output: str = ""
    ok: bool = False
    transcribed: int = 0
    needs_confirm_count: int = 0
    hwp_ok: bool = False
    notes: list[str] = field(default_factory=list)
    # 양식별 상세(사람 요약·기계 인수인계용): 어느 칸이 애매한지(확정 명령)와
    # 완성본에 값이 없어 비운 칸. 집계(needs_confirm_count)만으론 "어느 칸? 무슨 명령?"을
    # 알 수 없어 배치 사용자가 각 양식을 단일 모드로 다시 돌려야 했던 갭을 메운다.
    needs_confirm: list[dict] = field(default_factory=list)
    unmatched_targets: list[dict] = field(default_factory=list)


@dataclass
class BatchAutofillReport:
    notice_folder: str
    source_pool: str
    output_dir: str
    items: list[BatchAutofillItem] = field(default_factory=list)
    skipped_targets: list[str] = field(default_factory=list)

    @property
    def ok_count(self) -> int:
        return sum(1 for i in self.items if i.ok)

    @property
    def hwp_count(self) -> int:
        return sum(1 for i in self.items if i.hwp_ok)


def _name_has_any_keyword(name: str, keywords: tuple[str, ...]) -> bool:
    lowered = name.lower()
    return any(kw.lower() in lowered for kw in keywords)


def is_skipped_non_form(path: Path) -> bool:
    """공고문·안내·포스터 등 양식이 아닌 첨부는 배치 타깃에서 제외한다."""
    return _name_has_any_keyword(path.stem, _NON_FORM_NAME_KEYWORDS)


def discover_form_targets(notice_folder: str | Path) -> list[Path]:
    """공고 폴더 최상위에서 양식 후보(.docx/.hwp/.hwpx)를 찾는다."""
    folder = Path(notice_folder)
    if not folder.is_dir():
        raise FileNotFoundError(f"공고 폴더가 없습니다: {folder}")

    found: list[Path] = []
    seen: set[str] = set()
    for pattern in _FORM_GLOB_PATTERNS:
        for path in sorted(folder.glob(pattern)):
            key = str(path.resolve())
            if key in seen:
                continue
            seen.add(key)
            if is_skipped_non_form(path):
                continue
            if path.suffix.lower() not in _SUPPORTED_EXTS:
                continue
            found.append(path)
    return found


def list_source_pool(pool_dir: str | Path, *, recursive: bool = False) -> list[Path]:
    """완성본 A 후보 폴더에서 지원 확장자 파일을 나열한다."""
    folder = Path(pool_dir)
    if not folder.is_dir():
        raise FileNotFoundError(f"소스 풀 폴더가 없습니다: {folder}")

    files: list[Path] = []
    seen: set[str] = set()
    glob_fn = folder.rglob if recursive else folder.glob
    for pattern in _FORM_GLOB_PATTERNS:
        for path in sorted(glob_fn(pattern)):
            if path.suffix.lower() not in _SUPPORTED_EXTS:
                continue
            key = str(path.resolve())
            if key in seen:
                continue
            seen.add(key)
            files.append(path)
    return files


def _parse_yyyymmdd_from_name(name: str) -> int | None:
    """파일명/stem 에서 YYYYMMDD 를 추출한다(없으면 None)."""
    import re

    for m in re.finditer(r"(\d{4})[.\-_]?(\d{2})[.\-_]?(\d{2})", name):
        y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if 1900 <= y <= 2100 and 1 <= mo <= 12 and 1 <= d <= 31:
            return y * 10000 + mo * 100 + d
    for m in re.finditer(r"(?<!\d)(\d{8})(?!\d)", name):
        raw = m.group(1)
        y, mo, d = int(raw[:4]), int(raw[4:6]), int(raw[6:8])
        if 1900 <= y <= 2100 and 1 <= mo <= 12 and 1 <= d <= 31:
            return y * 10000 + mo * 100 + d
    return None


def _score_resume_filename(path: Path, *, prefer_resume: bool) -> tuple[int, int, int]:
    """이력서 소스 선택용 파일명 점수: (이력서보너스, YYYYMMDD, 감점)."""
    if not prefer_resume:
        return 0, 0, 0
    stem = path.stem
    bonus = 3 if _RESUME_BONUS_KEYWORD in stem else 0
    date_val = _parse_yyyymmdd_from_name(stem) or 0
    penalty = sum(1 for kw in _RESUME_PENALTY_KEYWORDS if kw in stem)
    return bonus, date_val, penalty


def _score_source_candidate(
    path: Path,
    keywords: tuple[str, ...],
    *,
    prefer_resume: bool = False,
) -> tuple[int, int, int, int, float]:
    """소스 후보 점수: (키워드적중, 이력서보너스, YYYYMMDD, 감점, mtime)."""
    hits = sum(1 for kw in keywords if kw.lower() in path.stem.lower())
    bonus, date_val, penalty = _score_resume_filename(path, prefer_resume=prefer_resume)
    return hits, bonus, date_val, penalty, path.stat().st_mtime


def _source_sort_key(
    dry_run: int,
    kw_hits: int,
    resume_bonus: int,
    date_val: int,
    penalty: int,
    mtime: float,
) -> tuple[int, int, int, int, int, float]:
    """dry-run → 이력서보너스 → 파일명날짜 → 키워드 → 감점(낮을수록 좋음) → mtime."""
    return dry_run, resume_bonus, date_val, kw_hits, -penalty, mtime


@dataclass
class SourcePickScore:
    path: str
    dry_run: int = 0
    keyword_hits: int = 0
    resume_bonus: int = 0
    filename_date: int = 0
    penalty: int = 0
    mtime: float = 0.0

    @property
    def sort_key(self) -> tuple[int, int, int, int, int, float]:
        return _source_sort_key(
            self.dry_run, self.keyword_hits, self.resume_bonus,
            self.filename_date, self.penalty, self.mtime,
        )


@dataclass
class SourcePickReport:
    pool_dir: str
    recommended: str
    recursive: bool
    prefer_resume: bool
    target: str
    scores: list[SourcePickScore] = field(default_factory=list)

    def as_dict(self) -> dict:
        return {
            "pool_dir": self.pool_dir,
            "recommended": self.recommended,
            "recursive": self.recursive,
            "prefer_resume": self.prefer_resume,
            "target": self.target,
            "scores": [
                {
                    "path": s.path,
                    "dry_run": s.dry_run,
                    "keyword_hits": s.keyword_hits,
                    "resume_bonus": s.resume_bonus,
                    "filename_date": s.filename_date,
                    "penalty": s.penalty,
                    "mtime": s.mtime,
                }
                for s in self.scores
            ],
        }


def rank_source_pool(
    pool_dir: str | Path,
    target: Path | None = None,
    keywords: tuple[str, ...] | None = None,
    *,
    recursive: bool = False,
    prefer_resume: bool = False,
    use_dry_run: bool = True,
) -> SourcePickReport:
    """소스 풀 후보를 점수순으로 정렬해 추천 1개와 breakdown 을 반환한다."""
    keywords = keywords or _DEFAULT_SOURCE_KEYWORDS
    pool = Path(pool_dir)
    candidates = list_source_pool(pool, recursive=recursive)
    scores: list[SourcePickScore] = []
    for cand in candidates:
        dr = 0
        if use_dry_run and target is not None:
            dr = dry_run_transcribe_score(cand, target)
        kw_hits, bonus, date_val, penalty, mtime = _score_source_candidate(
            cand, keywords, prefer_resume=prefer_resume,
        )
        scores.append(SourcePickScore(
            path=str(cand),
            dry_run=dr,
            keyword_hits=kw_hits,
            resume_bonus=bonus,
            filename_date=date_val,
            penalty=penalty,
            mtime=mtime,
        ))
    scores.sort(key=lambda s: s.sort_key, reverse=True)
    recommended = scores[0].path if scores else ""
    return SourcePickReport(
        pool_dir=str(pool),
        recommended=recommended,
        recursive=recursive,
        prefer_resume=prefer_resume,
        target=str(target) if target else "",
        scores=scores,
    )


def format_source_pick_korean(report: SourcePickReport) -> str:
    """pick 서브커맨드용 한국어 요약."""
    lines = [
        f"소스 풀: {report.pool_dir}",
        f"재귀 스캔: {'예' if report.recursive else '아니오'}",
        f"이력서 우선: {'예' if report.prefer_resume else '아니오'}",
    ]
    if report.target:
        lines.append(f"타깃(dry-run): {report.target}")
    if not report.recommended:
        lines.append("추천 소스: (후보 없음)")
        return "\n".join(lines)
    lines.append(f"추천 소스: {report.recommended}")
    lines.append("")
    lines.append("상위 후보 점수:")
    for i, s in enumerate(report.scores[:5], 1):
        name = Path(s.path).name
        parts = [
            f"dry-run={s.dry_run}",
            f"키워드={s.keyword_hits}",
        ]
        if report.prefer_resume:
            parts.extend([
                f"이력서보너스={s.resume_bonus}",
                f"파일명날짜={s.filename_date or '-'}",
                f"감점={s.penalty}",
            ])
        lines.append(f"  {i}. {name} ({', '.join(parts)})")
    return "\n".join(lines)


def pick_source_from_pool(
    pool_dir: str | Path,
    target: Path,
    keywords: tuple[str, ...] | None = None,
    *,
    recursive: bool = False,
    prefer_resume: bool = False,
    use_dry_run: bool = True,
) -> Path | None:
    """타깃마다 소스 A 1개를 고른다(dry-run 매칭 점수 → 키워드 → 최신 mtime)."""
    report = rank_source_pool(
        pool_dir, target, keywords,
        recursive=recursive, prefer_resume=prefer_resume, use_dry_run=use_dry_run,
    )
    if not report.recommended:
        return None
    if len(report.scores) == 1:
        return Path(report.recommended)

    if use_dry_run and report.scores[0].dry_run > 0:
        return Path(report.recommended)

    # dry-run 전부 0이면 파일명 키워드·이력서 점수 경로로 폴백
    best = report.scores[0]
    if best.keyword_hits > 0 or best.resume_bonus > 0:
        top_kw = best.keyword_hits
        top_bonus = best.resume_bonus
        top_date = best.filename_date
        top_penalty = best.penalty
        tied = [
            s for s in report.scores
            if s.keyword_hits == top_kw
            and s.resume_bonus == top_bonus
            and s.filename_date == top_date
            and s.penalty == top_penalty
        ]
        return Path(max(tied, key=lambda s: s.mtime).path)
    return Path(max(report.scores, key=lambda s: s.mtime).path)


def dry_run_transcribe_score(source: Path, target: Path) -> int:
    """소스-타깃 쌍의 예상 high 자동전사 건수(출력 파일 없음, 가벼운 매칭)."""
    source = Path(source)
    target = Path(target)
    if source.suffix.lower() not in _SUPPORTED_EXTS:
        return 0
    if target.suffix.lower() not in _SUPPORTED_EXTS:
        return 0

    with tempfile.TemporaryDirectory(prefix="xform_dry_") as td:
        tmpdir = Path(td)
        probe = AutofillReport(source=str(source), target=str(target), output="")
        src_docx = _to_docx_if_needed(source, tmpdir, probe)
        tgt_docx = _to_docx_if_needed(target, tmpdir, probe)
        if src_docx is None or tgt_docx is None:
            return 0
        try:
            src_fields, src_originals = _extract_source(src_docx)
            tgt_fields = find_target_fields(tgt_docx)
        except PackageNotFoundError:
            return 0
        if not src_fields or not tgt_fields:
            return 0
        matches = match_fields(src_fields, tgt_fields, src_originals)
        return sum(
            1 for m in matches
            if m.confidence == "high" and m.source_label and m.value
        )


def try_convert_filled_docx_to_hwp(docx_path: Path) -> tuple[Optional[Path], list[str]]:
    """채운 DOCX 를 같은 폴더에 .hwp 로 변환 시도한다(COM 불가 시 예외 없이 실패 보고)."""
    from .hwp_docx_convert import docx_to_hwp

    docx_path = Path(docx_path)
    if docx_path.suffix.lower() != ".docx" or not docx_path.exists():
        return None, ["DOCX 출력이 없어 HWP 변환을 건너뜁니다"]
    hwp_path = docx_path.with_suffix(".hwp")
    if hwp_path.resolve() == docx_path.resolve():
        return None, ["HWP 출력 경로가 DOCX 와 같습니다"]
    conv = docx_to_hwp(docx_path, hwp_path)
    if conv.ok:
        return hwp_path, []
    return None, list(conv.notes)


def format_batch_summary_korean(report: BatchAutofillReport) -> str:
    """사용자용 한국어 요약(채팅·팝업에 붙여넣기 좋은 문장)."""
    lines: list[str] = []
    n_ok = report.ok_count
    n_total = len(report.items)
    n_confirm = sum(i.needs_confirm_count for i in report.items)
    n_hwp = report.hwp_count
    n_fail = n_total - n_ok

    if n_ok:
        lines.append(f"양식 {n_ok}개 채움")
    if n_confirm:
        lines.append(f"{n_confirm}칸 확인필요")
    if n_hwp:
        lines.append(f"HWP {n_hwp}개 생성")
    elif n_ok:
        lines.append("HWP 변환 실패 — DOCX만 저장됨(한글 프로그램에서 열어 주세요)")
    if n_fail:
        lines.append(f"실패 {n_fail}건(세부는 아래 참고)")
    if report.skipped_targets:
        lines.append(f"공고문 등 제외 {len(report.skipped_targets)}개")
    if report.output_dir:
        lines.append(f"저장: {report.output_dir}")

    for item in report.items:
        if not item.ok:
            name = Path(item.target).name
            reason = item.notes[0] if item.notes else "채움 실패"
            lines.append(f"  · {name}: {reason}")
        elif item.needs_confirm_count and not n_confirm:
            pass
        elif item.needs_confirm_count:
            name = Path(item.target).name
            lines.append(f"  · {name}: 확인필요 {item.needs_confirm_count}칸")
        elif item.ok and not item.hwp_ok and item.output.endswith(".docx"):
            name = Path(item.target).name
            if any("HWP" in n or "한글" in n for n in item.notes):
                lines.append(f"  · {name}: DOCX만 저장(HWP 변환 실패)")

    if not lines:
        return "채울 양식을 찾지 못했습니다. 폴더에 신청서·참가서류 파일이 있는지 확인해 주세요."
    return "\n".join(lines)


def format_batch_detail_korean(
    report: BatchAutofillReport, per_form_limit: int = 6
) -> str:
    """배치 채움의 **양식별 '확인 필요(확정 명령)·빈칸' 상세**를 사람이 읽게 만든다.

    ``format_batch_summary_korean`` 은 '확인필요 N칸' 같은 집계만 주므로, 비개발자는
    어느 칸이 애매하고 무슨 ``--confirm`` 명령을 붙여야 하는지 알 수 없어 각 양식을
    단일 모드로 다시 돌려야 했다. 이 함수는 채운 양식마다 needs_confirm 후보 + 그대로
    붙여 다시 실행할 수 있는 ``--confirm "타깃=소스"`` 명령과 빈칸(직접 채울 칸) 목록을
    보여줘, 배치에서도 사람 확인 루프를 닫을 수 있게 한다.

    읽기 전용 — report 만 읽고 값을 지어내지 않는다(빈칸은 완성본에 값이 없어 비운 칸).
    확인 필요·빈칸이 하나도 없으면 빈 문자열을 돌려 준다(호출부에서 출력 생략).
    """
    blocks: list[str] = []
    for item in report.items:
        if not item.ok:
            continue
        nc = list(item.needs_confirm)
        um = list(item.unmatched_targets)
        if not nc and not um:
            continue

        head = f"■ {Path(item.target).name}  (채움 {item.transcribed}"
        if nc:
            head += f" · 확인필요 {len(nc)}"
        if um:
            head += f" · 빈칸 {len(um)}"
        head += "칸)"
        block: list[str] = [head]

        if nc:
            block.append("  확인 필요(붙여 다시 실행하면 채워집니다):")
            for entry in nc[:per_form_limit]:
                tgt = entry.get("target_label") or entry.get("normalized") or ""
                cands = [c for c in (entry.get("candidates") or []) if c]
                if cands:
                    cand_str = ", ".join(_summary_shorten(c, 24) for c in cands[:4])
                    block.append(f"    · {_summary_shorten(tgt)} → 후보: {cand_str}")
                    block.append(f"        확정: {_confirm_hint(tgt, cands[0])}")
                else:
                    block.append(
                        f"    · {_summary_shorten(tgt)} → 알맞은 후보 없음(직접 채워 주세요)")
            extra = len(nc) - per_form_limit
            if extra > 0:
                block.append(f"    …외 {extra}칸")

        if um:
            block.append("  빈칸(완성본에 값이 없어 비움, 직접 채워 주세요):")
            for entry in um[:per_form_limit]:
                tgt = entry.get("target_label") or entry.get("normalized") or ""
                block.append(f"    · {_summary_shorten(tgt)}")
            extra = len(um) - per_form_limit
            if extra > 0:
                block.append(f"    …외 {extra}칸")

        blocks.append("\n".join(block))

    if not blocks:
        return ""
    header = "[양식별 확인 필요·빈칸] (없는 값은 지어내지 않습니다)"
    return header + "\n" + "\n".join(blocks)


# --- 단일 전사 결과 사람용 요약 -----------------------------------------------

def _summary_shorten(text: str, limit: int = 40) -> str:
    """요약 한 줄에 넣기 좋게 공백 정리 + 과길이 말줄임(값을 바꾸지 않고 표시만)."""
    t = " ".join(str(text or "").split())
    if len(t) <= limit:
        return t
    return t[: max(1, limit - 1)] + "…"


def _confirm_hint(tgt: str, src: str) -> str:
    """복붙 가능한 확정 안내(단일·배치 공용). 라벨/소스에 ``"`` 나 ``=`` 가 있으면
    인라인 ``--confirm "타깃=소스"`` 파싱이 깨지므로(파서가 첫 ``=`` 에서 분리·셸 따옴표
    붕괴 → 엉뚱한 칸 기입) 값 손상 없이 특수문자를 담는 ``--confirm-file``(JSON) 로
    안내한다. 특수문자가 없으면 종전 인라인 명령을 그대로 방출(하위호환).

    읽기 전용 표시용 — 값을 바꾸지 않는다.
    """
    if '"' in tgt or '"' in src or '=' in tgt or '=' in src:
        pair = json.dumps({tgt: src}, ensure_ascii=False)  # 특수문자를 안전히 담은 유효 JSON
        return (f'라벨에 특수문자(\" 또는 =)가 있어 --confirm-file 사용 → '
                f'{pair} 를 JSON 파일로 저장 후 --confirm-file 파일.json')
    return f'--confirm "{tgt}={src}"'


def format_single_summary_korean(report: AutofillReport) -> str:
    """단일 A→B 전사 결과를 비개발자용 한국어 요약으로 만든다.

    원본 JSON 덤프 대신 **무엇이 채워졌고 / 무엇을 확인해야 하고 / 무엇을 직접
    채워야 하는지**를 한눈에 보여준다. 확인 필요(needs_confirm) 칸은 그대로 붙여
    다시 실행할 수 있는 ``--confirm "타깃=소스"`` 명령을 함께 제시해, 비개발자가
    raw JSON 을 뜯어보지 않고도 사람 확인 루프를 닫을 수 있게 한다.

    읽기 전용 — ``report`` 만 읽고 값을 지어내지 않는다(전사 개수는 실제 저장 기준
    ``transcribed``, 확인·빈칸 목록은 report 의 실제 항목).
    """
    lines: list[str] = ["완성본 → 빈 양식 전사 결과"]
    if report.source:
        lines.append(f"  소스: {Path(report.source).name}")
    if report.target:
        lines.append(f"  양식: {Path(report.target).name}")
    if report.output:
        lines.append(f"  결과: {Path(report.output).name}")

    if not report.ok:
        reason = report.notes[0] if report.notes else "전사할 항목을 찾지 못했습니다."
        lines.append("")
        lines.append(f"[아직 제출본 아님] {_summary_shorten(reason, 80)}")

    # [채운 칸] 자동으로 채운 칸 (실제 저장 기준 transcribed, 목록은 매칭)
    matches = list(report.matches)
    if report.transcribed or matches:
        lines.append("")
        lines.append(f"[자동으로 채운 칸] {report.transcribed}개")
        for m in matches:
            tgt = _summary_shorten(m.target_label or m.normalized)
            val = _summary_shorten(m.value, 50)
            tag = " (사용자 확정)" if m.confidence == "confirmed" else ""
            src = m.source_label
            if src and _key(src) != m.normalized:
                lines.append(f"  · {tgt} ← {val}  (소스 라벨: {_summary_shorten(src)}){tag}")
            else:
                lines.append(f"  · {tgt} ← {val}{tag}")
        # 미기입 사유는 여러 가지(표 셀에 이미 실값·단락 필드 위치 미확정·빈값)라
        # 원인을 단정하지 않는다(비개발자 오해 방지 — 코드리뷰 MEDIUM#1).
        skipped = len(matches) - report.transcribed
        if skipped > 0:
            lines.append(
                f"  ※ 위 {skipped}칸은 채우지 않았습니다"
                "(양식에 이미 값이 있거나 채울 위치를 확정하지 못함).")

    # [선택칸] 자동 체크한 선택칸(box -> checked)
    checked_groups = [
        g for g in report.checkbox_groups
        if isinstance(g.get("checked_option_index"), int)
        and g["checked_option_index"] >= 0
    ]
    if report.checkbox_checked or checked_groups:
        lines.append("")
        lines.append(f"[자동 체크한 선택칸] {report.checkbox_checked}개")
        for g in checked_groups:
            label = _summary_shorten(g.get("label") or g.get("normalized") or "")
            sv = _summary_shorten(g.get("source_value") or "", 30)
            lines.append(f"  · {label}: '{sv}' 로 체크" if sv else f"  · {label}: 체크")

    # [확인 필요] 비슷하지만 애매 — 후보 + 복붙 가능한 확정 명령
    if report.needs_confirm:
        lines.append("")
        lines.append(
            f"[확인 필요] 비슷하지만 확실치 않아 비워둠 {len(report.needs_confirm)}개")
        for nc in report.needs_confirm[:12]:
            tgt = nc.get("target_label") or nc.get("normalized") or ""
            cands = [c for c in (nc.get("candidates") or []) if c]
            if cands:
                cand_str = ", ".join(_summary_shorten(c, 24) for c in cands[:4])
                lines.append(f"  · {_summary_shorten(tgt)} → 후보: {cand_str}")
                lines.append(f"      확정하려면: {_confirm_hint(tgt, cands[0])}")
            else:
                lines.append(
                    f"  · {_summary_shorten(tgt)} → 알맞은 후보 없음(직접 채워 주세요)")
        extra = len(report.needs_confirm) - 12
        if extra > 0:
            lines.append(f"  …외 {extra}개")

    # [빈칸] 완성본에 값이 없어 비워둔 칸(날조 0 — 직접 채움)
    if report.unmatched_targets:
        lines.append("")
        lines.append(
            f"[빈칸] 완성본에 값이 없어 비워둔 칸 {len(report.unmatched_targets)}개 (직접 채워 주세요)")
        for um in report.unmatched_targets[:12]:
            tgt = um.get("target_label") or um.get("normalized") or ""
            lines.append(f"  · {_summary_shorten(tgt)}")
        extra = len(report.unmatched_targets) - 12
        if extra > 0:
            lines.append(f"  …외 {extra}개")

    body = (report.transcribed or matches or checked_groups or report.checkbox_checked
            or report.needs_confirm or report.unmatched_targets)
    if report.ok and not body:
        lines.append("")
        lines.append("  (채울 빈칸이 없거나 이미 모두 채워져 있습니다.)")

    tail: list[str] = []
    if report.needs_confirm:
        tail.append("확인 필요 칸은 위 --confirm 명령을 붙여 다시 실행하면 채워집니다.")
    if report.unmatched_targets:
        tail.append("빈칸은 완성본에 값이 없어 비운 것이라 직접 입력하세요"
                    "(없는 값은 지어내지 않습니다).")
    if tail:
        lines.append("")
        lines.append("다음: " + " ".join(tail))

    return "\n".join(lines)


def batch_autofill_from_pool(
    notice_folder: str | Path,
    source_pool: str | Path,
    output_subdir: str = "filled",
    *,
    source_keywords: tuple[str, ...] | None = None,
    recursive: bool = False,
    prefer_resume: bool = False,
    use_ai: bool = False,
    confirmations: Optional[dict[str, str]] = None,
    enable_checkbox: bool = True,
    convert_hwp: bool = True,
) -> BatchAutofillReport:
    """공고 폴더의 양식들을 소스 풀에서 고른 A 로 일괄 채운다."""
    notice = Path(notice_folder)
    pool = Path(source_pool)
    out_dir = notice / output_subdir
    out_dir.mkdir(parents=True, exist_ok=True)

    batch = BatchAutofillReport(
        notice_folder=str(notice),
        source_pool=str(pool),
        output_dir=str(out_dir),
    )

    # 스킵된 공고문 파일명 기록(진단용)
    for pattern in _FORM_GLOB_PATTERNS:
        for path in notice.glob(pattern):
            if is_skipped_non_form(path):
                batch.skipped_targets.append(str(path))

    targets = discover_form_targets(notice)
    if not targets:
        return batch

    for target in targets:
        item = BatchAutofillItem(target=str(target), source="", output="")
        source = pick_source_from_pool(
            pool, target, source_keywords,
            recursive=recursive, prefer_resume=prefer_resume,
        )
        if source is None:
            item.notes.append("완성본 폴더에 소스 파일이 없습니다")
            batch.items.append(item)
            continue

        item.source = str(source)
        out_path = out_dir / f"{target.stem}_filled.docx"
        item.output = str(out_path)

        try:
            rep = autofill_from_source(
                source, target, out_path,
                use_ai=use_ai,
                confirmations=confirmations,
                enable_checkbox=enable_checkbox,
            )
        except (FileNotFoundError, ValueError, OSError) as exc:
            item.notes.append(str(exc))
            batch.items.append(item)
            continue

        item.transcribed = rep.transcribed
        item.needs_confirm_count = len(rep.needs_confirm)
        # 양식별 상세를 보존해 배치에서도 "어느 칸이 확인 필요/빈칸인지 + 확정 명령"을
        # 단일 모드처럼 안내한다(값을 지어내지 않고 rep 의 실제 항목만 복사).
        item.needs_confirm = list(rep.needs_confirm)
        item.unmatched_targets = list(rep.unmatched_targets)
        item.ok = rep.ok
        item.notes.extend(rep.notes[:3])
        item.output = rep.output

        out_file = Path(rep.output)
        if convert_hwp and item.ok and out_file.suffix.lower() == ".docx" and out_file.exists():
            hwp_path, hwp_notes = try_convert_filled_docx_to_hwp(out_file)
            if hwp_path:
                item.hwp_output = str(hwp_path)
                item.hwp_ok = True
            else:
                item.notes.append(
                    "HWP 변환 실패 — DOCX만 저장됨"
                    + (f" ({hwp_notes[0]})" if hwp_notes else ""))

        batch.items.append(item)

    return batch
