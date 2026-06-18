"""usage_acceptance.py — 실사용 기준 수용 검사(acceptance) 엔진.

doc_quality_score(서식 청소 점수)와 측정 대상이 다르다.
여기서는 '심사위원·접수처 관점'에서 제출을 막는 하드페일 결함을 찾는다.
fail 등급 결함이 1개라도 있으면 제출불가(DRAFT) 판정이다.

- AI 호출 없음. 동일 입력 → 동일 결과(결정론).
- 텍스트 검사는 본문 직계 단락 + 표 셀(중첩 포함) + 머리글·바닥글·텍스트박스의
  표시 텍스트를 모두 본다(숨은 영역의 마커·자리표시 미탐 방지).
- 폰트 검사(_iter_all_runs)는 본문·표만 본다 — 머리글/바닥글은 양식 고정 영역이라
  포함하면 정상 문서가 혼용 오탐(fail)되는 위험이 더 크다(오탐 0 우선).
- 공고 조건부 요건(블라인드 마스킹·요구 산출형식·분량 등)은 AcceptanceConfig 로
  묶어 전달한다. config 미지정이면 전부 기본값 = 현행 동작.

검사 항목 (check_id / 등급)
---------------------------------
unresolved_markers     [확인필요] 등 미해결 작성 마커 잔존            fail
self_inserted_blocks   파이프라인 자기삽입 안내블록(NotebookLM 등) 잔존 fail
template_placeholders  양식 자리표시(<사진(이미지)>, OOO 등) 잔존      fail
unchecked_choices      '택 1'/'해당여부' 행에 선택 표시 없음           fail
empty_label_fields     명칭 등 필수 라벨 옆 칸 공란                    fail
font_name_mixing       폰트 이름 혼용(표 포함, 허용 4종 초과)          fail
font_size_spread       글자크기 과다 분산(7종 초과)·이상치(<8/>18pt)   warn
empty_table_rows       완전히 빈 표 행                                 warn
recruit_date_conflict  채용시기 표기 상호 모순                         warn
masking_violation      블라인드 마스킹 위반(실명 잔존) — blind_review 시만 fail(조건부)
residual_colored_runs  검정 아닌 유색 텍스트 잔존(파란 안내문구 등)      fail
paren_choices          괄호형 ( ) 선택란 미선택 의심                    warn(선도입)
empty_label_fields_ext 라벨 변형·확장 목록 공란 의심                    warn(선도입)
empty_image_slots      빈 그림/사진 칸 의심                             warn(선도입)
page_overflow          분량 제한 초과 의심(근사) — config 지정 시만     warn(조건부)
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Iterator

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

SEV_FAIL = "fail"
SEV_WARN = "warn"

_MAX_SAMPLES = 5

# --- 패턴 ------------------------------------------------------------------
_MARKER_RE = re.compile(r"\[확인필요[^\]]*\]|\[산출근거[^\]]*\]\s*\?|\[TODO[^\]]*\]")
_SELF_BLOCK_RE = re.compile(
    r"NotebookLM\s*슬라이드\s*생성용\s*프롬프트|이\s*블록은\s*삭제하세요|슬라이드\s*생성에\s*붙여넣으세요"
    r"|\[작성\s*보강\s*가이드\]|\[AI\s*작성\s*보강\]|섹션은\s*삭제하세요|\(작성\s*필요\s*—"
)
# 제거기(image_apply.strip_notebooklm_blocks)가 같은 정의를 쓰도록 공개한다 —
# 검출과 제거의 패턴이 어긋나면 '지웠는데 검출됨' 류의 재발이 생긴다.
SELF_BLOCK_RE = _SELF_BLOCK_RE

# 쓰기 모듈(psst_fill·bizplan_ai_writer)이 import 해 쓰는 표준 문구 —
# 검출(_SELF_BLOCK_RE)과 동일한 곳에서 정의해 '문구가 갈라져 검출 실패'(ACC-6)
# 재발을 막는다(R5 의 검출-제거 패턴 공유 원칙).
SCAFFOLD_HEADING = "■ [작성 보강 가이드] PSST 미흡·누락 영역"
SCAFFOLD_DELETE_NOTICE = (
    "아래는 자동 점검 결과 보강이 필요한 영역입니다. 각 항목을 본문 해당 절에 "
    "구체적으로 작성한 뒤 이 가이드 섹션은 삭제하세요. (내용은 자동 생성되지 않습니다)"
)
SCAFFOLD_ITEM_SUFFIX = "(작성 필요 — 구체 내용·근거 수치 기입)"
AI_SECTION_HEADING = "■ [AI 작성 보강] PSST 영역 초안 (근거 명시 · 검토 필수)"
AI_SECTION_DELETE_NOTICE = (
    "아래는 AI 가 입력 정보 기반으로 작성한 초안입니다. 수치의 [산출근거]·[확인필요] 를 "
    "반드시 검토하고, 본문 해당 절로 옮긴 뒤 이 섹션은 삭제하세요."
)
# ○○○/OOO 는 '빈 자리표시'이기도 하지만 블라인드 공고에서는 '올바른 마스킹'이다.
# blind_review=True 면 placeholder 검사에서 제외된다(ACC-1 오탐 방지).
_MASK_PLACEHOLDER_RE = re.compile(r"(?<!\w)(OOO|○○○)(?!\w)")
_PLACEHOLDER_RES = (
    re.compile(r"<\s*사진\s*\(이미지\)[^>]*>"),
    re.compile(r"<[^<>\n]{0,30}제목\s*>"),
    _MASK_PLACEHOLDER_RE,
)
# 블라인드 평가에서 마스킹해야 하는 라벨(공백·콜론 정규화 후 비교)
_BLIND_LABEL_NORM = {"성명", "대표자성명", "대표자명", "직장명", "대학명"}
_BLIND_INLINE_RE = re.compile(
    r"(대표자\s*성명|대표자명|성명|직장명|대학명)\s*[:：]\s*([^\s,/|·]{1,20})"
)
_HANGUL_NAME_RE = re.compile(r"[가-힣]{2,10}")
_CHECKBOX_ROW_RE = re.compile(r"택\s*1|해당\s*여부")
_CHECKED_MARKS = ("■", "☑", "✔", "▣", "☒", "√")
_LABEL_FIELDS = ("명 칭", "명칭", "기업명", "대표자명", "창업아이템명")
_DATE_TOKEN_RE = re.compile(r"[’'‘]\s?(\d{2})\s*\.\s*(\d{1,2})")
_DATE_TOKEN_KR_RE = re.compile(r"(\d{4})\s*년\s*(\d{1,2})\s*월")  # '2026년 3월' 한국식 표기
# docx_ops._PRESERVE_COLORS 와 동일(흰색 계열 보존색) — 검정/무지정 외 허용 색
_COLOR_PRESERVE = {"ffffff", "fffffe", "f2f2f2"}
# --- US-3c warn 선도입 검사용(오탐 표면적 큼 — 음성 코퍼스 검증 후 차기 fail 승격) ---
_LABEL_FIELDS_EXT = ("사업자등록번호", "생년월일", "연락처", "신청분야")
_norm_label = lambda s: re.sub(r"[\s:：]", "", s)  # noqa: E731 — 라벨 비교 정규화
_LABEL_NORM_ALL = ({_norm_label(x) for x in _LABEL_FIELDS}
                   | {_norm_label(x) for x in _LABEL_FIELDS_EXT})
_IMAGE_LABEL_RE = re.compile(r"사진|이미지|그림|로고")
_EMPTY_PAREN_RE = re.compile(r"\(\s*\)")
_FILLED_PAREN_RE = re.compile(r"\(\s*[VvOX○●√■]\s*\)")
_STANDALONE_V_RE = re.compile(r"(?<![A-Za-z0-9])[V√●](?![A-Za-z0-9])")  # 'TV' 의 V 는 제외
# 템플릿이 정상적으로 쓰는 폰트 조합(KISED 양식 기준) — 이 수를 넘으면 혼용으로 본다.
_ALLOWED_FONT_KINDS = 4
_ALLOWED_SIZE_KINDS = 7


@dataclass(frozen=True)
class AcceptanceConfig:
    """공고 조건부 수용 요건 묶음 — 기본값은 전부 현행 동작(보수적, 오탐 0).

    required_format 은 여기 '보관'만 한다. 판정은 run_acceptance(DOCX 내부 검사)가
    아니라 파이프라인 레벨의 최종 산출 파일 확장자 게이트에서 한다 — 그렇지 않으면
    변환 전 DOCX 가 HWP 요구 공고에서 영구 fail 이 된다.
    """
    blind_review: bool = False          # True: ○○○ 마스킹 허용 + 실명 잔존 검출(fail)
    required_format: str | None = None  # 예: "hwp" — 파이프라인 레벨에서 판정
    max_pages: int | None = None        # 본문 분량 제한 — None 이면 검사 안 함
    ai_section_max: int | None = None   # AI활용계획 등 섹션 분량 제한
    allowed_fonts: int = _ALLOWED_FONT_KINDS


@dataclass
class CheckResult:
    check_id: str
    label: str
    severity: str
    defects: int
    samples: list[str] = field(default_factory=list)
    detail: str = ""

    @property
    def passed(self) -> bool:
        return self.defects == 0

    def as_dict(self) -> dict[str, Any]:
        return {
            "check_id": self.check_id, "label": self.label,
            "severity": self.severity, "defects": self.defects,
            "passed": self.passed, "samples": self.samples, "detail": self.detail,
        }


@dataclass
class AcceptanceReport:
    source: str
    results: list[CheckResult] = field(default_factory=list)

    @property
    def submittable(self) -> bool:
        return all(r.passed for r in self.results if r.severity == SEV_FAIL)

    @property
    def fail_defects(self) -> int:
        return sum(r.defects for r in self.results if r.severity == SEV_FAIL)

    @property
    def warn_defects(self) -> int:
        return sum(r.defects for r in self.results if r.severity == SEV_WARN)

    def as_dict(self) -> dict[str, Any]:
        return {
            "source": self.source,
            "submittable": self.submittable,
            "verdict": "제출가능" if self.submittable else "제출불가(DRAFT)",
            "fail_defects": self.fail_defects,
            "warn_defects": self.warn_defects,
            "checks": [r.as_dict() for r in self.results],
        }


# --- 순회 도우미 ------------------------------------------------------------

def _iter_cells(tables) -> Iterator:
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                yield cell
                if cell.tables:
                    yield from _iter_cells(cell.tables)


def _iter_tables_recursive(tables) -> Iterator:
    """표를 중첩(셀 안의 표) 포함해 전부 낸다 — 행 단위 검사들의 공용 순회.

    (_iter_cells 는 lxml proxy 유지 주석이 걸린 민감 경로라 의도적으로 따로 둔다.)
    """
    for table in tables:
        yield table
        for row in table.rows:
            for cell in row.cells:
                if cell.tables:
                    yield from _iter_tables_recursive(cell.tables)


def _dedup_cells(doc: Document):
    """병합 중복을 제거한 셀 목록.

    주의: lxml 요소 proxy 는 참조가 사라지면 메모리(id)가 재사용될 수 있어
    id 만 저장하면 서로 다른 셀이 '본 것'으로 오판된다(과소 집계 버그).
    refs 리스트로 proxy 참조를 유지해 id 안정성을 보장한다.
    """
    seen_ids: set[int] = set()
    refs: list = []          # proxy 생존 유지용 — 삭제 금지
    out = []
    for cell in _iter_cells(doc.tables):
        tc = cell._tc
        if id(tc) in seen_ids:
            continue
        seen_ids.add(id(tc))
        refs.append(tc)
        out.append(cell)
    return out


# 제거기(image_apply.strip_notebooklm_blocks)가 검출과 같은 셀 순회(병합 중복
# 제거 포함)를 쓰도록 공개한다 — 순회가 어긋나면 '지웠는데 검출됨'이 재발한다.
dedup_cells = _dedup_cells


def _iter_extra_paragraphs(doc: Document) -> Iterator[tuple[str, Paragraph]]:
    """머리글·바닥글 단락(표 셀 포함) — 본문 밖에 숨은 표시 텍스트(ACC-9).

    명시 정의가 있을 때만 본다 — is_linked_to_previous 인 빈 머리글에 접근하면
    python-docx 가 part 를 새로 만드는 부작용이 있어(읽기 전용 위반) 건너뛴다.
    """
    for section in doc.sections:
        for label, hf in (("머리글", section.header), ("바닥글", section.footer)):
            if hf is None or hf.is_linked_to_previous:
                continue
            for p in hf.paragraphs:
                yield (label, p)
            for cell in _iter_cells(hf.tables):
                for p in cell.paragraphs:
                    yield (label, p)


def _iter_textbox_paragraphs(doc: Document) -> Iterator[Paragraph]:
    """본문 내 텍스트박스(w:txbxContent) 단락 — 그림틀 안 텍스트도 표시 영역이다."""
    for txbx in doc.element.body.iter(qn("w:txbxContent")):
        for p_el in txbx.iter(qn("w:p")):
            yield Paragraph(p_el, doc)


def _iter_all_texts(doc: Document) -> Iterator[tuple[str, str]]:
    """(위치, 텍스트) — 본문 단락·표 셀(중첩)·머리글·바닥글·텍스트박스를 모두 낸다."""
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            yield ("본문", t)
    for cell in _dedup_cells(doc):
        t = cell.text.strip()
        if t:
            yield ("표", t)
    for where, p in _iter_extra_paragraphs(doc):
        t = p.text.strip()
        if t:
            yield (where, t)
    for p in _iter_textbox_paragraphs(doc):
        t = p.text.strip()
        if t:
            yield ("텍스트박스", t)


def _iter_all_runs(doc: Document):
    # 의도적으로 본문·표만 — 머리글/바닥글 폰트는 양식 고정 영역이라 포함 시
    # 정상 문서가 font_name_mixing(fail) 오탐될 위험이 크다(모듈 docstring 참조).
    for p in doc.paragraphs:
        yield from p.runs
    for cell in _dedup_cells(doc):
        for p in cell.paragraphs:
            yield from p.runs


def _row_cells_dedup(row) -> list[str]:
    """병합 셀 중복을 제거한 행 셀 텍스트 목록."""
    out: list[str] = []
    seen: list = []
    for cell in row.cells:
        tc = cell._tc
        if any(tc is s for s in seen):
            continue
        seen.append(tc)
        out.append(cell.text.strip())
    return out


def _scan_regex(doc: Document, regexes) -> tuple[int, list[str]]:
    n, samples = 0, []
    for where, text in _iter_all_texts(doc):
        hits = 0
        for rx in regexes:
            hits += len(rx.findall(text))
        if hits:
            n += hits
            if len(samples) < _MAX_SAMPLES:
                samples.append(f"[{where}] {text[:40]}")
    return n, samples


# --- 개별 검사 ---------------------------------------------------------------

def check_unresolved_markers(doc: Document, config: AcceptanceConfig | None = None) -> CheckResult:
    n, s = _scan_regex(doc, (_MARKER_RE,))
    return CheckResult("unresolved_markers", "[확인필요] 등 미해결 마커", SEV_FAIL, n, s,
                       f"마커 {n}개 — 전부 실제 값으로 치환해야 제출 가능")


def check_self_inserted_blocks(doc: Document, config: AcceptanceConfig | None = None) -> CheckResult:
    n = 0
    samples: list[str] = []
    for where, text in _iter_all_texts(doc):
        if _SELF_BLOCK_RE.search(text):
            n += 1
            if len(samples) < _MAX_SAMPLES:
                samples.append(f"[{where}] {text[:40]}")
    return CheckResult("self_inserted_blocks", "자기삽입 안내블록 잔존", SEV_FAIL, n, samples,
                       f"파이프라인이 삽입한 작업용 블록 {n}개 잔존 — 제출본에서는 0이어야 함")


def check_template_placeholders(doc: Document, config: AcceptanceConfig | None = None) -> CheckResult:
    """자리표시 검출 — 패턴이 겹쳐도(동일 구간 이중매치) 1건으로 센다.

    blind_review=True 면 ○○○/OOO 패턴을 제외한다 — 블라인드 공고에서 그것은
    빈 자리표시가 아니라 '올바른 마스킹'이다(실명 잔존은 masking_violation 이 잡음).
    """
    blind = config is not None and config.blind_review
    regexes = tuple(rx for rx in _PLACEHOLDER_RES
                    if not (blind and rx is _MASK_PLACEHOLDER_RE))
    n = 0
    samples: list[str] = []
    for where, text in _iter_all_texts(doc):
        spans: list[tuple[int, int]] = []
        for rx in regexes:
            spans.extend(m.span() for m in rx.finditer(text))
        if not spans:
            continue
        spans.sort()
        merged = 0
        cur_end = -1
        for a, b in spans:
            if a >= cur_end:
                merged += 1
                cur_end = b
            else:
                cur_end = max(cur_end, b)
        n += merged
        if len(samples) < _MAX_SAMPLES:
            samples.append(f"[{where}] {text[:40]}")
    return CheckResult("template_placeholders", "양식 자리표시 잔존", SEV_FAIL, n, samples,
                       f"<사진(이미지)>·OOO 등 자리표시 {n}개")


def check_unchecked_choices(doc: Document, config: AcceptanceConfig | None = None) -> CheckResult:
    defects = 0
    samples: list[str] = []
    seen_labels: set[str] = set()
    for table in _iter_tables_recursive(doc.tables):
        for row in table.rows:
            cells = _row_cells_dedup(row)
            row_text = " ".join(cells)
            if not _CHECKBOX_ROW_RE.search(row_text):
                continue
            label = cells[0][:20] if cells else row_text[:20]
            if label in seen_labels:
                continue
            seen_labels.add(label)
            checked = (any(m in row_text for m in _CHECKED_MARKS)
                       or _STANDALONE_V_RE.search(row_text))  # 'V' 표기 체크 인정(오탐 방지)
            if "□" in row_text and not checked:
                defects += 1
                if len(samples) < _MAX_SAMPLES:
                    samples.append(f"[표] {row_text[:40]}")
    return CheckResult("unchecked_choices", "선택란(택1·해당여부) 미체크", SEV_FAIL, defects, samples,
                       f"체크 표시(■ 등) 없는 선택 행 {defects}개")


def check_empty_label_fields(doc: Document, config: AcceptanceConfig | None = None) -> CheckResult:
    defects = 0
    samples: list[str] = []
    for table in _iter_tables_recursive(doc.tables):
        for row in table.rows:
            cells = _row_cells_dedup(row)
            for i, c in enumerate(cells[:-1]):
                if c in _LABEL_FIELDS and not cells[i + 1]:
                    defects += 1
                    if len(samples) < _MAX_SAMPLES:
                        samples.append(f"[표] '{c}' 옆 칸 공란")
    return CheckResult("empty_label_fields", "필수 라벨 옆 칸 공란", SEV_FAIL, defects, samples,
                       f"공란 필수칸 {defects}개")


def _is_masked_value(v: str) -> bool:
    return ("○" in v) or bool(re.fullmatch(r"O+", v))


def check_residual_colored_runs(doc: Document, config: AcceptanceConfig | None = None) -> CheckResult:
    """검정 아닌 유색 본문 텍스트 잔존(파란 안내문구·회색 가이드 등) — fail (ACC-3).

    공고 규칙: 양식의 색 있는 안내문구는 삭제하고 본문은 검정 글씨.
    - 색 미지정(상속)·검정(000000)·흰색 계열(_COLOR_PRESERVE)은 통과.
    - 자기삽입 블록(_SELF_BLOCK_RE 단락)은 self_inserted_blocks 가 잡으므로 제외
      (이중 집계 방지). 단락 단위로 센다.
    - 순회 범위: 본문·표 셀 + 머리글/바닥글 + 텍스트박스(ACC-9, 다른 fail 검사와 동일).
    """
    defects = 0
    samples: list[str] = []

    def _scan_para(where: str, p) -> None:
        nonlocal defects
        text = (p.text or "").strip()
        if not text or _SELF_BLOCK_RE.search(text):
            return
        for run in p.runs:
            if not (run.text or "").strip():
                continue
            try:
                color = run.font.color
                rgb = color.rgb if (color is not None and color.type is not None) else None
            except Exception:
                rgb = None
            if rgb is None:
                continue
            hexv = str(rgb).lower()
            if hexv == "000000" or hexv in _COLOR_PRESERVE:
                continue
            defects += 1
            if len(samples) < _MAX_SAMPLES:
                samples.append(f"[{where}] #{hexv} '{text[:30]}'")
            return  # 단락당 1건

    for p in doc.paragraphs:
        _scan_para("본문", p)
    for cell in _dedup_cells(doc):
        for p in cell.paragraphs:
            _scan_para("표", p)
    # 머리글·바닥글·텍스트박스의 유색 안내문구도 본다(다른 fail 검사와 동일 범위 — ACC-9).
    for where, p in _iter_extra_paragraphs(doc):
        _scan_para(where, p)
    for p in _iter_textbox_paragraphs(doc):
        _scan_para("텍스트박스", p)
    return CheckResult("residual_colored_runs", "검정 아닌 유색 텍스트 잔존", SEV_FAIL, defects, samples,
                       f"유색 텍스트 단락 {defects}개 — 안내문구 삭제·본문 검정 규칙 위반 의심")


def check_masking_violation(doc: Document, config: AcceptanceConfig | None = None) -> CheckResult:
    """블라인드 평가 위반(실명·기관명 잔존) — config.blind_review=True 일 때만 활성.

    블라인드 공고는 성명·직장명·대학명을 ○ 마스킹해야 하며 실명 기재는 공고
    위반(탈락 사유)이다. 비블라인드(기본)에서는 defects=0 — 오탐 0 원칙.
    값이 통째로 한글 2~10자일 때만 잡는 보수적 패턴(라벨 옆 칸 / '라벨: 값' 한정).
    """
    label = "블라인드 마스킹 위반(실명 잔존)"
    if config is None or not config.blind_review:
        return CheckResult("masking_violation", label, SEV_FAIL, 0, [],
                           "비블라인드 공고 — 검사 비활성(blind_review 시 활성)")
    defects = 0
    samples: list[str] = []

    def _flag(where: str, lbl: str, value: str) -> None:
        nonlocal defects
        defects += 1
        if len(samples) < _MAX_SAMPLES:
            samples.append(f"[{where}] '{lbl}' 칸 실명 의심({value[0]}…) — ○ 마스킹 필요")

    for where, text in _iter_all_texts(doc):
        for lbl, value in _BLIND_INLINE_RE.findall(text):
            v = value.strip()
            if _is_masked_value(v):
                continue
            if _HANGUL_NAME_RE.fullmatch(v):
                _flag(where, lbl, v)

    for table in _iter_tables_recursive(doc.tables):
        for row in table.rows:
            cells = _row_cells_dedup(row)
            for i, c in enumerate(cells[:-1]):
                if re.sub(r"[\s:：]", "", c) not in _BLIND_LABEL_NORM:
                    continue
                v = cells[i + 1].strip()
                if not v or _is_masked_value(v):
                    continue
                if _HANGUL_NAME_RE.fullmatch(v):
                    _flag("표", c, v)
    return CheckResult("masking_violation", label, SEV_FAIL, defects, samples,
                       f"마스킹 안 된 실명 의심 {defects}건 — 블라인드 공고 위반(탈락 사유)")


def check_font_name_mixing(doc: Document, config: AcceptanceConfig | None = None) -> CheckResult:
    """폰트 혼용 검사 — ascii(영문)와 eastAsia(한글) 슬롯을 분리 집계한다.

    한글 문서는 run 하나가 정상적으로 (ascii=Arial, eastAsia=맑은 고딕) 페어를
    갖는다. 두 슬롯을 한 집합에 합산하면 페어가 2종으로 부풀어 정상 문서가
    fail 오탐된다(ACC-8). 슬롯별로 허용 종수를 따로 적용해 페어 오탐을 없애되,
    슬롯 안의 진짜 혼용(예: 한글 폰트 6종)은 그대로 잡는다.
    """
    allowed = config.allowed_fonts if config is not None else _ALLOWED_FONT_KINDS
    ascii_names: set[str] = set()
    ea_names: set[str] = set()

    def _style_fonts(style) -> tuple[str | None, str | None]:
        """단락 스타일 체인(base_style 포함)에서 명시된 ascii/eastAsia 폰트를 찾는다."""
        asc = ea = None
        seen: set[str] = set()
        st = style
        while st is not None and getattr(st, "style_id", None) not in seen:
            seen.add(st.style_id)
            try:
                if asc is None:
                    asc = st.font.name
                rpr = st.element.rPr
                if ea is None and rpr is not None and rpr.rFonts is not None:
                    ea = rpr.rFonts.get(qn("w:eastAsia"))
            except Exception:
                pass
            st = getattr(st, "base_style", None)
        return asc, ea

    def _scan_para(p) -> None:
        style_asc = style_ea = None
        style_resolved = False
        for run in p.runs:
            if not (run.text or "").strip():
                continue
            r_asc = run.font.name
            r_ea = None
            rpr = run._element.rPr
            if rpr is not None and rpr.rFonts is not None:
                r_ea = rpr.rFonts.get(qn("w:eastAsia"))
            # run 에 직접 지정이 없으면 단락 스타일 체인의 유효 폰트로 해석한다(ACC-7)
            if (r_asc is None or r_ea is None) and not style_resolved:
                style_asc, style_ea = _style_fonts(getattr(p, "style", None))
                style_resolved = True
            if r_asc is None:
                r_asc = style_asc
            if r_ea is None:
                r_ea = style_ea
            if r_asc:
                ascii_names.add(r_asc)
            if r_ea:
                ea_names.add(r_ea)

    for p in doc.paragraphs:
        _scan_para(p)
    for cell in _dedup_cells(doc):
        for p in cell.paragraphs:
            _scan_para(p)

    over = max(0, len(ascii_names) - allowed) + max(0, len(ea_names) - allowed)
    samples = [f"ascii:{n}" for n in sorted(ascii_names)[:_MAX_SAMPLES]] + \
              [f"eastAsia:{n}" for n in sorted(ea_names)[:_MAX_SAMPLES]]
    return CheckResult("font_name_mixing", "폰트 이름 혼용(표 포함)", SEV_FAIL, over,
                       samples,
                       f"ascii {len(ascii_names)}종 / eastAsia {len(ea_names)}종 (허용 슬롯별 {allowed}종)")


def check_paren_choices(doc: Document, config: AcceptanceConfig | None = None) -> CheckResult:
    """괄호형 선택란 미선택 의심 — warn (ACC-10, 선도입).

    '택 1/해당여부' 행에 □ 가 없고 '( )' 빈 괄호만 있으면 미선택 의심.
    채움 괄호 '(V)/(O)/(○)' 가 하나라도 있으면 통과. □ 형은 기존 fail 검사 담당.
    """
    defects = 0
    samples: list[str] = []
    for table in _iter_tables_recursive(doc.tables):
        for row in table.rows:
            row_text = " ".join(_row_cells_dedup(row))
            if not _CHECKBOX_ROW_RE.search(row_text) or "□" in row_text:
                continue
            if _EMPTY_PAREN_RE.search(row_text) and not _FILLED_PAREN_RE.search(row_text):
                defects += 1
                if len(samples) < _MAX_SAMPLES:
                    samples.append(f"[표] {row_text[:40]}")
    return CheckResult("paren_choices", "괄호형 선택란 미선택 의심", SEV_WARN, defects, samples,
                       f"빈 괄호 선택 행 {defects}개 — ( ) 안에 V/○ 표기 필요")


def check_empty_label_fields_ext(doc: Document, config: AcceptanceConfig | None = None) -> CheckResult:
    """필수 라벨 옆 칸 공란(확장) — warn (ACC-11, 선도입).

    기존 fail 검사(정확일치 5종)가 놓치는 라벨 변형('기 업 명', '명칭 :')과
    확장 라벨(사업자등록번호·생년월일·연락처·신청분야)을 정규화 비교로 잡는다.
    기존 fail 검사가 이미 센 항목(정확일치)은 제외해 이중 집계를 막는다.
    """
    defects = 0
    samples: list[str] = []
    for table in _iter_tables_recursive(doc.tables):
        for row in table.rows:
            cells = _row_cells_dedup(row)
            for i, c in enumerate(cells[:-1]):
                if c in _LABEL_FIELDS:        # 기존 fail 검사 담당분 제외
                    continue
                if _norm_label(c) in _LABEL_NORM_ALL and not cells[i + 1]:
                    defects += 1
                    if len(samples) < _MAX_SAMPLES:
                        samples.append(f"[표] '{c}' 옆 칸 공란(확장 검사)")
    return CheckResult("empty_label_fields_ext", "필수 라벨 옆 칸 공란(확장)", SEV_WARN, defects, samples,
                       f"공란 의심 {defects}개 — 라벨 변형·확장 목록 기준(warn 선도입)")


def check_empty_image_slots(doc: Document, config: AcceptanceConfig | None = None) -> CheckResult:
    """빈 그림/사진 칸 의심 — warn (ACC-12, 선도입).

    '사진/이미지/그림/로고' 짧은 라벨 셀의 옆 칸에 텍스트도 그림(w:drawing/w:pict)도
    없으면 미첨부 의심. 자리표시 텍스트(<사진(이미지)>)는 placeholder 검사 담당.
    """
    defects = 0
    samples: list[str] = []

    def _cell_has_image(cell) -> bool:
        tc = cell._tc
        return bool(tc.findall(".//" + qn("w:drawing")) or tc.findall(".//" + qn("w:pict")))

    for table in _iter_tables_recursive(doc.tables):
        for row in table.rows:
            # 병합 중복 제거하되 그림 검사를 위해 셀 객체가 필요하다
            seen: list = []
            cells = []
            for cell in row.cells:
                if any(cell._tc is s for s in seen):
                    continue
                seen.append(cell._tc)
                cells.append(cell)
            for i, cell in enumerate(cells[:-1]):
                t = cell.text.strip()
                if not t or len(t) > 12 or not _IMAGE_LABEL_RE.search(t):
                    continue
                nxt = cells[i + 1]
                if not nxt.text.strip() and not _cell_has_image(nxt):
                    defects += 1
                    if len(samples) < _MAX_SAMPLES:
                        samples.append(f"[표] '{t}' 옆 칸에 텍스트·그림 없음")
    return CheckResult("empty_image_slots", "빈 그림/사진 칸 의심", SEV_WARN, defects, samples,
                       f"그림 미첨부 의심 칸 {defects}개")


_HEADING_LIKE_RE = re.compile(r"^\s*(?:[ⅠⅡⅢⅣⅤ]|\d+\s*[.)]|■|□\s*[가-힣]+\s*:)")
_AI_SECTION_RE = re.compile(r"AI.{0,10}(활용|인재).{0,8}계획")


def check_page_overflow(doc: Document, config: AcceptanceConfig | None = None) -> CheckResult:
    """분량 제한 초과 의심(근사) — warn (ACC-4, config 지정 시에만 활성).

    python-docx 는 페이지를 정확히 못 세므로 본문 글자수 기반 근사(1p ≈ 1,500자)로
    경고만 한다. 정확 판정은 파이프라인 레벨의 미리보기 렌더러(page PNG) 몫.
    config.max_pages / ai_section_max 미지정(기본)이면 검사 비활성 — 오탐 0.
    """
    label = "분량 제한 초과 의심(근사)"
    if config is None or (config.max_pages is None and config.ai_section_max is None):
        return CheckResult("page_overflow", label, SEV_WARN, 0, [],
                           "분량 제한 미지정 — 검사 비활성")
    defects = 0
    samples: list[str] = []
    total_chars = sum(len(t) for _, t in _iter_all_texts(doc))
    est = max(1, -(-total_chars // 1500))
    if config.max_pages is not None and est > config.max_pages:
        defects += 1
        samples.append(f"본문 약 {total_chars:,}자 ≈ {est}p > 제한 {config.max_pages}p")
    if config.ai_section_max is not None:
        chars = 0
        in_section = False
        for p in doc.paragraphs:
            t = p.text.strip()
            if not in_section:
                if t and _AI_SECTION_RE.search(t):
                    in_section = True
                continue
            if t and _HEADING_LIKE_RE.match(t) and not _AI_SECTION_RE.search(t):
                break
            chars += len(t)
        if in_section:
            est_ai = max(1, -(-chars // 1500))
            if est_ai > config.ai_section_max:
                defects += 1
                samples.append(f"AI 계획 섹션 약 {chars:,}자 ≈ {est_ai}p > 제한 {config.ai_section_max}p")
    return CheckResult("page_overflow", label, SEV_WARN, defects, samples,
                       f"분량 초과 의심 {defects}건 — 근사 추정(정확 판정은 렌더러 필요)")


def check_font_size_spread(doc: Document, config: AcceptanceConfig | None = None) -> CheckResult:
    sizes: set[float] = set()
    outliers = 0
    for run in _iter_all_runs(doc):
        if not (run.text or "").strip() or run.font.size is None:
            continue
        pt = run.font.size.pt
        sizes.add(pt)
        if pt < 8 or pt > 18:
            outliers += 1
    over = max(0, len(sizes) - _ALLOWED_SIZE_KINDS) + (1 if outliers else 0)
    return CheckResult("font_size_spread", "글자크기 분산·이상치", SEV_WARN, over,
                       [f"{s}pt" for s in sorted(sizes)],
                       f"크기 {len(sizes)}종, 이상치(8pt 미만/18pt 초과) {outliers}개")


def check_empty_table_rows(doc: Document, config: AcceptanceConfig | None = None) -> CheckResult:
    defects = 0
    for table in doc.tables:
        for row in table.rows:
            cells = _row_cells_dedup(row)
            if cells and all(not c for c in cells):
                defects += 1
    return CheckResult("empty_table_rows", "완전히 빈 표 행", SEV_WARN, defects, [],
                       f"빈 행 {defects}개")


def check_recruit_date_conflict(doc: Document, config: AcceptanceConfig | None = None) -> CheckResult:
    tokens: set[str] = set()
    samples: list[str] = []
    for where, text in _iter_all_texts(doc):
        if "채용" not in text:
            continue
        found = [(y, m) for y, m in _DATE_TOKEN_RE.findall(text)]
        found += [(y[-2:], m) for y, m in _DATE_TOKEN_KR_RE.findall(text)]  # '2026년 3월' → '26.3
        for y, m in found:
            tok = f"'{y}.{int(m)}"
            if tok not in tokens and len(samples) < _MAX_SAMPLES:
                samples.append(f"[{where}] {tok} ← {text[:30]}")
            tokens.add(tok)
    defects = max(0, len(tokens) - 1)
    return CheckResult("recruit_date_conflict", "채용시기 표기 모순", SEV_WARN, defects, samples,
                       f"서로 다른 채용시기 표기 {len(tokens)}종: {sorted(tokens)}")


_ALL_CHECKS = (
    check_unresolved_markers,
    check_self_inserted_blocks,
    check_template_placeholders,
    check_unchecked_choices,
    check_empty_label_fields,
    check_masking_violation,
    check_residual_colored_runs,
    check_font_name_mixing,
    check_font_size_spread,
    check_empty_table_rows,
    check_recruit_date_conflict,
    check_paren_choices,
    check_empty_label_fields_ext,
    check_empty_image_slots,
    check_page_overflow,
)


def run_acceptance(path: str | Path, config: AcceptanceConfig | None = None) -> AcceptanceReport:
    """DOCX 1개에 대해 전체 수용 검사를 실행한다 (읽기 전용).

    config 미지정(None)이면 AcceptanceConfig 기본값으로 동작한다 — 현행과 동일.
    """
    cfg = config if config is not None else AcceptanceConfig()
    path = Path(path)
    doc = Document(str(path))
    report = AcceptanceReport(source=str(path))
    for check in _ALL_CHECKS:
        report.results.append(check(doc, cfg))
    return report


def backup_existing_output(target: str | Path) -> str:
    """고정 산출명이 기존 파일을 덮어쓰기 전에 타임스탬프 백업으로 보존한다(PIPE-2).

    재실행이 이전 산출물(사용자가 수정했을 수 있는 _DRAFT 포함)을 무경고로
    파괴하지 않게 한다. 백업했으면 백업 경로, 대상이 없으면 빈 문자열.
    """
    target = Path(target)
    if not target.exists():
        return ""
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    bak = target.with_name(f"{target.stem}_prev{stamp}{target.suffix}")
    i = 0
    while bak.exists():
        i += 1
        bak = target.with_name(f"{target.stem}_prev{stamp}_{i}{target.suffix}")
    target.replace(bak)
    return str(bak)


def force_draft_name(path: Path, *, avoid: Path | None = None) -> tuple[Path, str]:
    """게이트 정책(R7/R8) 파일명 유틸 — 제출불가 판정 파일에 ``_DRAFT`` 를 강제한다.

    검사(run_acceptance)는 읽기 전용이고, 이 헬퍼는 게이트(파이프라인)만 호출한다.
    이미 ``_DRAFT`` 이름이면 그대로 둔다. 목표 이름이 ``avoid``(예: 입력 원본)와
    겹치면 덮어쓰지 않고 ``_DRAFT2`` 로 마킹한다(원본 보존).

    Returns:
        (최종 경로, 오류 문자열). rename 실패(파일 잠금 등) 시 경로는 원래대로
        남고 오류 문자열이 채워진다 — 호출자는 이를 사용자에게 반드시 알려야 한다.
    """
    if path.stem.endswith("_DRAFT") or path.stem.endswith("_DRAFT2"):
        return path, ""
    draft = path.with_name(f"{path.stem}_DRAFT{path.suffix}")
    if avoid is not None and draft.resolve() == avoid.resolve():
        draft = path.with_name(f"{path.stem}_DRAFT2{path.suffix}")
    if draft.exists():
        # 기존 _DRAFT(사용자가 수정 중일 수 있는 파일)를 무경고로 파괴하지 않는다(PIPE-2)
        backup_existing_output(draft)
    try:
        path.replace(draft)
    except OSError as exc:
        return path, f"{type(exc).__name__}: {exc}"
    return draft, ""
