"""psst_fill.py — PSST 검사 결과를 바탕으로 '작성 뼈대 + 가이드'를 DOCX 에 삽입한다.

``psst_check.check_psst`` 가 '누락/미흡'으로 판정한 영역에 대해, 문서 끝에
**작성 보강 가이드 섹션**을 추가한다. 각 영역의 빠진 하위항목을 체크리스트로
나열해 사용자가 직접 무엇을 써야 하는지 알 수 있게 한다.

안전 원칙
---------
- 원본 DOCX 는 절대 덮어쓰지 않는다(``out_docx == in_docx`` 면 ``ValueError``).
- **내용을 지어내지 않는다.** 누락 항목에 '(작성 필요)' 자리표시만 넣는다.
  실제 알맹이는 사용자(또는 AI 초안)가 채워야 한다.
- 가이드는 문서 **끝**에 모아 추가한다(본문 위치 오판으로 인한 훼손 방지).
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any
import shutil

from docx import Document
from docx.shared import Pt, RGBColor

from .psst_check import check_psst

_DEFAULT_TARGET_GRADES = ("누락", "미흡")

# 영역별 작성 힌트(무엇을 어떻게 쓰면 좋은지 — 결정론 안내, 내용 날조 아님)
_AREA_HINTS = {
    "problem": "고객/시장의 문제를 수치 근거와 함께 구체적으로 기술하세요.",
    "solution": "해결방안·핵심기능·차별성·구현 가능성(TRL/시제품)을 적으세요.",
    "scale": "시장규모(TAM/SAM/SOM)·수익모델·판로·KPI/매출계획을 적으세요.",
    "team": "대표자 역량·팀 구성·외부 협력·수행 경험을 적으세요.",
}


@dataclass
class PSSTFillReport:
    areas_scaffolded: int = 0
    items_added: int = 0
    target_grades: tuple[str, ...] = _DEFAULT_TARGET_GRADES
    output_docx: str = ""
    scaffolded_areas: list[str] = field(default_factory=list)
    overall_ratio: float = 0.0

    def as_dict(self) -> dict[str, Any]:
        return {
            "areas_scaffolded": self.areas_scaffolded,
            "items_added": self.items_added,
            "target_grades": list(self.target_grades),
            "output_docx": self.output_docx,
            "scaffolded_areas": self.scaffolded_areas,
            "overall_ratio": round(self.overall_ratio, 3),
        }


def _add_heading(doc: Document, text: str, *, size: int = 13) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def _add_area_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)


def _add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    try:
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    except Exception:  # pragma: no cover
        pass


def _add_checkitem(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.add_run(text)


def apply_psst_scaffold(
    in_docx: str,
    out_docx: str,
    *,
    target_grades: tuple[str, ...] = _DEFAULT_TARGET_GRADES,
) -> PSSTFillReport:
    """PSST '누락/미흡' 영역에 작성 뼈대+가이드를 문서 끝에 삽입한다.

    Args:
        in_docx: 원본 DOCX(읽기 전용).
        out_docx: 결과 DOCX 경로. **in_docx 와 같으면 ValueError**.
        target_grades: 보강 대상 등급(기본: 누락·미흡).

    Returns:
        PSSTFillReport — 보강 결과 집계. 보강 대상이 없으면 변경 없이 복사본만 만든다.
    """
    in_path = Path(in_docx)
    out_path = Path(out_docx)
    if in_path.resolve() == out_path.resolve():
        raise ValueError("in_docx 와 out_docx 가 같습니다. 원본 덮어쓰기는 금지입니다.")
    if not in_path.exists():
        raise FileNotFoundError(f"입력 DOCX 가 없습니다: {in_docx}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(str(in_path), str(out_path))

    doc = Document(str(out_path))
    psst = check_psst(doc)
    report = PSSTFillReport(
        target_grades=tuple(target_grades),
        output_docx=str(out_path),
        overall_ratio=psst.overall_ratio,
    )

    weak = [a for a in psst.areas if a.grade in target_grades]
    if not weak:
        doc.save(str(out_path))
        return report

    _add_heading(doc, "■ [작성 보강 가이드] PSST 미흡·누락 영역")
    _add_note(
        doc,
        "아래는 자동 점검 결과 보강이 필요한 영역입니다. 각 항목을 본문 해당 절에 "
        "구체적으로 작성한 뒤 이 가이드 섹션은 삭제하세요. (내용은 자동 생성되지 않습니다)",
    )

    for area in weak:
        section_state = "섹션 없음" if not area.section_present else "섹션 있음"
        _add_area_title(
            doc, f"▶ {area.label} — 등급: {area.grade} ({section_state})"
        )
        hint = _AREA_HINTS.get(area.area)
        if hint:
            _add_note(doc, f"작성 방향: {hint}")
        if not area.section_present:
            _add_note(doc, "이 영역의 섹션 헤더가 문서에 없습니다 — 헤더부터 추가하세요.")
        targets = area.missing_items or [it for it in ()]
        if not targets:
            _add_checkitem(doc, "  □ (세부 항목 보강: 내용을 더 구체화하세요)")
            report.items_added += 1
        for item in targets:
            _add_checkitem(doc, f"  □ {item}: (작성 필요 — 구체 내용·근거 수치 기입)")
            report.items_added += 1
        report.areas_scaffolded += 1
        report.scaffolded_areas.append(area.label)

    doc.save(str(out_path))
    return report


def apply_psst_scaffold_docx(
    path: str,
    out_docx: str,
    *,
    target_grades: tuple[str, ...] = _DEFAULT_TARGET_GRADES,
) -> PSSTFillReport:
    return apply_psst_scaffold(path, out_docx, target_grades=target_grades)
