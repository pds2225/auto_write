from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document

from ..models import EvidenceSource, GeneratedImage, ProjectInput, TemplateProfile


class QAService:
    GUIDE_MARKER_RE = re.compile(r"(※|<[^>]+>|기재|작성요령|작성방법|예시|OOO|○○○)")
    CRITICAL_GUIDE_MARKER_RE = re.compile(
        r"(<[^>]*(기재|작성|입력|예시|제목|OOO|○○○|000)[^>]*>|OOO|○○○|000|작성요령|작성방법)"
    )

    @staticmethod
    def _msg(category: str, message: str) -> str:
        return f"[{category}] {message}"

    def _normalize_match_text(self, text: str) -> str:
        normalized = re.sub(r"\s+", " ", text or "").strip()
        normalized = normalized.replace("–", "-").replace("—", "-").replace("‑", "-")
        for _ in range(2):
            if len(normalized) < 16 or len(normalized) % 2 != 0:
                break
            half = len(normalized) // 2
            left = normalized[:half].strip()
            right = normalized[half:].strip()
            if left and left == right:
                normalized = left
                continue
            break
        return normalized

    def _match_anchor(self, candidate: str, anchor_text: str) -> bool:
        anchor = self._normalize_match_text(anchor_text)
        text = self._normalize_match_text(candidate)
        if not anchor or not text:
            return False
        if anchor in text:
            return True
        compact_anchor = anchor.replace(" ", "")
        compact_text = text.replace(" ", "")
        return bool(compact_anchor and compact_anchor in compact_text)

    @staticmethod
    def _paragraph_text(paragraph) -> str:
        parts = [
            str(node.text or "").strip()
            for node in paragraph.iter()
            if str(getattr(node, "tag", "")).endswith("}t")
        ]
        return " ".join(part for part in parts if part).strip()

    def _is_meaningful_text(self, text: str, anchor_text: str = "") -> bool:
        candidate = str(text or "").strip()
        if not candidate:
            return False
        if self.GUIDE_MARKER_RE.search(candidate):
            return False
        if anchor_text and self._match_anchor(candidate, anchor_text):
            return False
        return True

    def _table_has_meaningful_text(self, table_element, anchor_text: str) -> bool:
        for cell in table_element.iter():
            if not str(getattr(cell, "tag", "")).endswith("}tc"):
                continue
            texts = [str(node.text or "").strip() for node in cell.iter() if str(getattr(node, "tag", "")).endswith("}t")]
            joined = re.sub(r"\s+", " ", " ".join(item for item in texts if item)).strip()
            if self._is_meaningful_text(joined, anchor_text):
                return True
        return False

    def _has_content_after_anchor(self, document: Document, anchor_text: str) -> bool:
        body_items = list(document.element.body)
        for index, element in enumerate(body_items):
            if not str(getattr(element, "tag", "")).endswith("}p"):
                continue
            current = self._paragraph_text(element)
            if not self._match_anchor(current, anchor_text):
                continue
            for next_index in range(index + 1, min(index + 10, len(body_items))):
                next_item = body_items[next_index]
                tag = str(getattr(next_item, "tag", ""))
                if tag.endswith("}p"):
                    if self._is_meaningful_text(self._paragraph_text(next_item), anchor_text):
                        return True
                    continue
                if tag.endswith("}tbl") and self._table_has_meaningful_text(next_item, anchor_text):
                    return True
        return False

    def _cell_has_meaningful_text(self, document: Document, table_index: int, row: int, cell: int) -> bool:
        if table_index < 0 or table_index >= len(document.tables):
            return False
        table = document.tables[table_index]
        if row < 0 or row >= len(table.rows):
            return False
        if cell < 0 or cell >= len(table.rows[row].cells):
            return False
        text = table.rows[row].cells[cell].text.strip()
        return self._is_meaningful_text(text)

    def _collect_guide_markers(self, document: Document, limit: int = 12) -> list[str]:
        found: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if self.GUIDE_MARKER_RE.search(text):
                found.append(text[:120])
            if len(found) >= limit:
                return found
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if not text:
                        continue
                    if self.GUIDE_MARKER_RE.search(text):
                        found.append(text[:120])
                    if len(found) >= limit:
                        return found
        return found

    def build_report(
        self,
        profile: TemplateProfile,
        project_input: ProjectInput,
        render_result: dict[str, Any],
        images: list[GeneratedImage],
        evidence: list[EvidenceSource],
    ) -> dict[str, Any]:
        errors: list[str] = []
        warnings: list[str] = []
        answers = project_input.answers
        output_path = Path(render_result["output_path"])
        output_doc = Document(str(output_path)) if output_path.exists() else None
        section_by_field_id = {section.field_id: section for section in profile.sections}
        table_by_id = {table.table_id: table for table in profile.tables}

        for question in profile.questions:
            if not question.required:
                continue
            target = question.target.get("kind")
            if target == "project_meta":
                value = project_input.project_meta.get(question.target.get("key", ""))
                if not str(value or "").strip():
                    errors.append(self._msg("필수입력", f"'{question.label}' 항목이 비어 있습니다."))
            elif target == "organization_profile":
                value = project_input.organization_profile.get(question.target.get("key", ""))
                if not str(value or "").strip():
                    errors.append(self._msg("필수입력", f"'{question.label}' 항목이 비어 있습니다."))
            elif target == "section":
                section = section_by_field_id.get(question.question_id)
                if output_doc is not None and section is not None:
                    if not self._has_content_after_anchor(output_doc, section.anchor_text):
                        errors.append(self._msg("필수입력", f"'{question.label}' 항목이 결과 문서에서 비어 있습니다."))
                else:
                    value = str(answers.get(question.question_id, "") or "").strip()
                    if not value:
                        errors.append(self._msg("필수입력", f"'{question.label}' 항목이 비어 있습니다."))
            elif target == "table_cell":
                table = table_by_id.get(str(question.target.get("table_id", "")))
                cell_id = str(question.target.get("cell_id", ""))
                table_cell = next((cell for cell in table.cells if cell.cell_id == cell_id), None) if table else None
                if output_doc is not None and table is not None and table_cell is not None:
                    if not self._cell_has_meaningful_text(output_doc, table.table_index, table_cell.row, table_cell.cell):
                        errors.append(self._msg("필수입력", f"'{question.label}' 항목이 결과 문서에서 비어 있습니다."))
                else:
                    value = str(answers.get(question.question_id, "") or "").strip()
                    if not value:
                        errors.append(self._msg("필수입력", f"'{question.label}' 항목이 비어 있습니다."))
            else:
                value = answers.get(question.question_id)
                if not str(value or "").strip():
                    errors.append(self._msg("필수입력", f"'{question.label}' 항목이 비어 있습니다."))

        image_by_slot = {image.slot_id for image in images}
        for slot in profile.image_slots:
            if slot.required and slot.slot_id not in image_by_slot:
                errors.append(self._msg("필수이미지", f"'{slot.label}' 이미지가 생성되지 않았습니다."))

        if render_result.get("errors"):
            for issue in render_result["errors"]:
                errors.append(self._msg("렌더링", str(issue)))
        if render_result.get("warnings"):
            for issue in render_result["warnings"]:
                warnings.append(self._msg("렌더링", str(issue)))

        for table in profile.tables:
            for cell in table.cells:
                value = str(answers.get(cell.cell_id, "") or "")
                if len(value) > 180:
                    warnings.append(self._msg("표길이", f"'{cell.label}' 내용이 길어 줄바꿈이 많을 수 있습니다."))

        if output_doc is not None:
            guide_markers = self._collect_guide_markers(output_doc)
            for marker in guide_markers[:6]:
                if self.CRITICAL_GUIDE_MARKER_RE.search(marker):
                    errors.append(self._msg("가이드문구", f"결과 문서에 가이드 문구가 남아 있습니다: {marker}"))
                else:
                    warnings.append(self._msg("가이드문구", f"결과 문서에 가이드 문구가 남아 있습니다: {marker}"))

        if project_input.evidence_requests and not evidence:
            warnings.append(self._msg("출처", "검색 요청은 있었지만 저장된 출처가 없습니다."))

        for request in project_input.evidence_requests:
            if request.topic and not any(source.topic == request.topic for source in evidence):
                warnings.append(self._msg("출처", f"'{request.topic}' 주제의 검색 결과를 찾지 못했습니다."))

        if not output_path.exists():
            errors.append(self._msg("산출물", "최종 DOCX 파일이 생성되지 않았습니다."))

        return {
            "passed": not errors,
            "error_count": len(errors),
            "warning_count": len(warnings),
            "errors": errors,
            "warnings": warnings,
            "summary": {
                "sections_written": render_result.get("sections_written", 0),
                "cells_written": render_result.get("cells_written", 0),
                "images_written": render_result.get("images_written", 0),
                "sources_collected": len(evidence),
            },
        }
