"""doc_quality_ops.py

완성된 DOCX(사업계획서·보고서·제출문서)에 대한 **결정론적 후처리 연산** 모음.

설계 원칙
---------
1. 모든 함수는 python-docx ``Document`` 를 받아 in-place 로 수정하고, 변경 횟수를 반환한다.
2. AI 를 호출하지 않는다(규칙 기반). 키가 없어도 동작한다.
3. 원본 구조를 최소 변경한다 — 텍스트 노드(``w:t``) 단위로만 손대고, 단락/런 구조는
   가능한 보존한다. 파괴적 삭제는 "명백한 빈 단락" 과 "명백한 양식 안내 단락" 으로 제한.
4. 검증된 헬퍼는 ``docx_ops`` 에서 재사용한다(중복 구현 금지).

각 함수의 반환값은 정수(변경 건수)이며, 오케스트레이터가 품질 리포트에 집계한다.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

# 검증된 헬퍼 재사용 (docx_ops.py)
from .docx_ops import (
    _iter_body_paragraphs,
    _paragraph_text,
    GUIDE_MARKER_RE,
    _PRESERVE_COLORS,
    _normalize_color_value,
    _set_run_color_black_unless_preserved,
)

# ---------------------------------------------------------------------------
# 상수 / 정규식
# ---------------------------------------------------------------------------

# 글머리표로 쓰이는 대표 기호들
_BULLET_SYMBOLS = "○●◦◌∙·•‣▪▫■□◇◆▶▷►▸-–—*ㅇㅁ"
# 줄 맨 앞 "공백* 글머리표+ 공백+" 패턴 → 기호 뒤 공백을 1칸으로
_BULLET_PREFIX_RE = re.compile(
    rf"^([ \t 　]*)([{re.escape(_BULLET_SYMBOLS)}]+)([ \t 　]{{2,}})"
)
# 내부 다중 공백(2칸 이상) → 1칸 (줄바꿈/탭 제외, 일반/전각/nbsp 공백 대상)
_MULTI_SPACE_RE = re.compile(r"[  　]{2,}")
# 강조 대상 핵심 키워드 (정량 성과 지표)
_EMPHASIS_KEYWORDS = (
    "매출", "영업이익", "순이익", "고용", "채용", "수출", "특허", "인증",
    "R&D", "연구개발", "KPI", "목표", "기대효과", "투자유치", "점유율",
    "성장률", "ROI", "매출액", "거래액", "MAU", "전환율", "원가절감",
)
# 숫자/단위 패턴 (강조는 "키워드 + 실제 수치" 동반 시에만 → 과잉 강조 방지)
# 반드시 '아라비아/전각 숫자' 가 있어야 정량 성과로 인정한다.
# 단독 한글 단위(개/회/점/위/차/명/건/배)는 '개발·관점·위해·차별·설명·요건' 등에서
# 오탐을 유발하므로 수치로 인정하지 않는다(require_numeric 강화).
_NUMERIC_RE = re.compile(r"[0-9０-９]")
# "명백한 양식 안내" 로 판단할 보수적 패턴 (단락 통째 삭제 대상)
_PURE_GUIDE_RE = re.compile(
    r"^\s*(?:[<\(［【]\s*)?(?:작성\s*요령|작성\s*방법|작성\s*예시|기재\s*요령|기재\s*방법|"
    r"유의\s*사항|예\s*시|참고\s*용|※[^。\n]{0,40}(?:작성|기재|예시))"
)
# 보강 패턴 — _PURE_GUIDE_RE(줄 시작 고정사전)의 사각지대를 보완한다.
# (a) 글머리표(○ ㅇ - * 등) 접두가 붙은 안내, (b) 문장 중간/끝의 '명령형 어미'
# (기재하시오·작성하세요·입력 바랍니다·작성할 것 등) 로 끝나는 안내만 대상으로 한다.
# 선언형(작성한다/작성합니다/기재하였다)은 어미가 '하시오/하세요/바랍니다'가 아니므로
# 매칭되지 않아 일반 서술문의 과삭제를 막는다. 동사 앞 텍스트는 {0,80} 으로 제한.
_GUIDE_EXTRA_RE = re.compile(
    r"^[\s○●▪◦·•‣⁃\-–—－*ㅇㅁ<\(（［【]*"
    r"(?:"
    r"(?:작성|기재|기입|입력)\s*(?:요령|방법|예시|방식|요망)"
    r"|[^\n]{0,80}?(?:기재|작성|기입|입력|서술|기술)\s*"
    r"(?:하(?:시오|십시오|세요|시기\s*바랍니다)|해\s*주(?:십시오|시기\s*바랍니다|세요)"
    r"|하여\s*주(?:십시오|시기\s*바랍니다|세요)|하기\s*바랍니다|할\s*것|바랍니다|바람|요망|요함)"
    r")"
)


def _is_guide_text(text: str) -> bool:
    """단락/표행 텍스트가 '명백한 양식 안내문구'인지 단일 판정(삭제·채점 공통 기준).

    삭제(remove_guide_paragraphs / remove_table_guide_rows)와 채점(doc_quality_score._scan_guide)
    이 동일 기준을 쓰도록 단일 진실원천으로 둔다. ``_PURE_GUIDE_RE``(줄 시작 고정사전) 또는
    ``_GUIDE_EXTRA_RE``(글머리표 접두·명령형 어미)에 매칭되면 안내문구로 본다.

    주의: ``docx_ops.GUIDE_MARKER_RE`` 는 '셀에 든 기존 텍스트를 덮어써도 되는가'(빈 자리 표시)
    판단용으로 ``기재``·``<...>`` 등을 모두 잡는 훨씬 광범위한 패턴이다. 삭제 기준에 합치면
    실제 본문까지 지워질 수 있어 의도적으로 ``_is_guide_text`` 에 포함하지 않는다(과삭제 방지).
    """
    if not text:
        return False
    return bool(_PURE_GUIDE_RE.search(text) or _GUIDE_EXTRA_RE.search(text))


@dataclass
class QualityOpsReport:
    """후처리 연산 결과 집계."""
    guide_paragraphs_removed: int = 0
    table_guide_rows_removed: int = 0
    bullet_spacing_fixed: int = 0
    table_cells_cleaned: int = 0
    empty_paragraphs_removed: int = 0
    paragraphs_emphasized: int = 0
    font_sizes_normalized: int = 0
    paragraphs_unified: int = 0
    colored_runs_normalized: int = 0
    notes: list[str] = field(default_factory=list)

    def as_dict(self) -> dict:
        return {
            "guide_paragraphs_removed": self.guide_paragraphs_removed,
            "table_guide_rows_removed": self.table_guide_rows_removed,
            "bullet_spacing_fixed": self.bullet_spacing_fixed,
            "table_cells_cleaned": self.table_cells_cleaned,
            "empty_paragraphs_removed": self.empty_paragraphs_removed,
            "paragraphs_emphasized": self.paragraphs_emphasized,
            "font_sizes_normalized": self.font_sizes_normalized,
            "paragraphs_unified": self.paragraphs_unified,
            "colored_runs_normalized": self.colored_runs_normalized,
            "notes": list(self.notes),
        }


# ---------------------------------------------------------------------------
# 내부 유틸
# ---------------------------------------------------------------------------

def _text_nodes(element):
    """주어진 oxml 요소 하위의 모든 w:t 노드를 반환."""
    return list(element.iter(qn("w:t")))


def _element_has_drawing(element) -> bool:
    """이미지/도형(w:drawing, w:pict)을 포함하면 True (삭제 금지 판단용)."""
    return bool(element.find(".//" + qn("w:drawing")) is not None
                or element.find(".//" + qn("w:pict")) is not None)


def _run_has_text(run) -> bool:
    """런에 비어있지 않은 텍스트 노드가 있으면 True."""
    return any((n.text or "") for n in run.findall(qn("w:t")))


def _bold_is_on(rpr) -> bool:
    """rPr 에 켜진 bold(<w:b>, val 이 false/0/off 가 아님)가 있으면 True."""
    if rpr is None:
        return False
    b = rpr.find(qn("w:b"))
    if b is None:
        return False
    val = b.get(qn("w:val"))
    return val is None or str(val).lower() not in ("0", "false", "off")


def _para_is_bold(para) -> bool:
    """텍스트가 있는 런 중 하나라도 bold 가 켜진 단락이면 True."""
    for run in para.findall(qn("w:r")):
        if _run_has_text(run) and _bold_is_on(run.find(qn("w:rPr"))):
            return True
    return False


# ---------------------------------------------------------------------------
# 1. 글머리표 공백 정리
# ---------------------------------------------------------------------------

def normalize_bullet_spacing(doc: Document) -> int:
    """글머리표(○ ㅇ · • - 등) 뒤 과다 공백과 단락 내부 다중 공백을 1칸으로 정리.

    런(run) 단위로 텍스트 노드만 수정하므로 서식은 보존된다. 보수적으로:
    - 단락의 첫 텍스트 노드에 대해서만 '글머리표+공백' 접두 패턴을 정리
    - 모든 텍스트 노드의 내부 2칸 이상 공백을 1칸으로 축소
    표 안/밖 단락 모두 대상으로 한다.
    """
    fixed = 0
    # 본문 + 표를 모두 포함하는 전체 단락 순회
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    for para in all_paragraphs:
        nodes = _text_nodes(para._p)
        if not nodes:
            continue
        changed = False
        # (a) 접두 글머리표 공백 — 첫 비어있지 않은 노드
        for node in nodes:
            if node.text is None or node.text == "":
                continue
            new = _BULLET_PREFIX_RE.sub(lambda m: f"{m.group(1)}{m.group(2)} ", node.text, count=1)
            if new != node.text:
                node.text = new
                changed = True
            break
        # (b) 내부 다중 공백 축소 (모든 노드)
        for node in nodes:
            if not node.text:
                continue
            new = _MULTI_SPACE_RE.sub(" ", node.text)
            if new != node.text:
                node.text = new
                changed = True
        if changed:
            fixed += 1
    return fixed


# ---------------------------------------------------------------------------
# 2. 표 내부 공백 정리
# ---------------------------------------------------------------------------

def cleanup_table_whitespace(doc: Document) -> int:
    """표 셀 텍스트의 앞뒤 공백·내부 다중 공백을 정리한다.

    셀 단위로 처리하되 런 구조는 보존:
    - 셀의 첫 텍스트 노드 왼쪽 공백 제거, 마지막 텍스트 노드 오른쪽 공백 제거
    - 각 노드 내부 다중 공백 1칸 축소
    이미지가 든 셀은 건드리지 않는다.
    """
    cleaned = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                if _element_has_drawing(tc):
                    continue
                nodes = [n for n in _text_nodes(tc) if n.text]
                if not nodes:
                    continue
                before = [n.text for n in nodes]
                # 내부 다중 공백 축소
                for n in nodes:
                    n.text = _MULTI_SPACE_RE.sub(" ", n.text)
                # 셀 경계 trim
                nodes[0].text = nodes[0].text.lstrip("  　\t")
                nodes[-1].text = nodes[-1].text.rstrip("  　\t")
                if [n.text for n in nodes] != before:
                    cleaned += 1
    return cleaned


# ---------------------------------------------------------------------------
# 3. 불필요한 빈 단락 삭제
# ---------------------------------------------------------------------------

def remove_empty_paragraphs(doc: Document, *, keep_single: bool = True) -> int:
    """연속된 완전 빈 단락을 축소한다.

    - 본문(body) 직계 단락만 대상으로 한다(표 셀 내부는 건드리지 않음 — 셀 구조 보존).
    - '완전 빈 단락' = 텍스트가 비어있고 이미지/도형이 없는 w:p.
    - 연속 빈 단락 그룹에서 ``keep_single`` 이면 1개만 남기고 제거, 아니면 모두 제거.
    - 문서 맨 끝의 마지막 빈 단락 1개는 Word 호환을 위해 보존한다.
    """
    body = doc.element.body
    children = list(body)
    # w:p 직계만 인덱싱
    para_flags = []  # (element, is_empty)
    for el in children:
        if el.tag != qn("w:p"):
            para_flags.append((el, None))
            continue
        is_empty = (not _paragraph_text(el)) and (not _element_has_drawing(el))
        para_flags.append((el, is_empty))

    to_remove = []
    run_start = None
    sequence = []

    def flush(seq):
        if not seq:
            return
        # keep_single: 첫 단락 보존, 나머지 제거
        start = 1 if keep_single else 0
        to_remove.extend(seq[start:])

    prev_empty = False
    for el, is_empty in para_flags:
        if is_empty is True:
            sequence.append(el)
            prev_empty = True
        else:
            if prev_empty:
                flush(sequence)
                sequence = []
            prev_empty = False
    if prev_empty:
        flush(sequence)

    # 문서 맨 끝 빈 단락 1개는 보존 (sectPr 직전 호환)
    removed = 0
    for el in to_remove:
        # 마지막 body 자식이면 보존
        if el is body[-1]:
            continue
        body.remove(el)
        removed += 1
    return removed


# ---------------------------------------------------------------------------
# 4. 핵심 문장 강조 (Bold)
# ---------------------------------------------------------------------------

def emphasize_key_sentences(
    doc: Document,
    *,
    keywords: tuple[str, ...] = _EMPHASIS_KEYWORDS,
    underline: bool = False,
    require_numeric: bool = True,
    max_emphasis_ratio: float = 0.15,
    hard_emphasis_ratio: float = 0.30,
    min_emphasis: int = 1,
    max_emphasis: int | None = None,
) -> int:
    """정량 성과(매출·고용·수출·특허·KPI 등) 핵심 키워드가 든 단락을 Bold 처리.

    과잉 강조(문서 전체가 굵어지는 현상)를 막기 위해 **비율 기반 예산**으로 강조량을
    제한한다:
    - 강조 허용 총량 = 본문 비어있지 않은 단락 수 × ``max_emphasis_ratio``(기본 15%).
      단 ``hard_emphasis_ratio``(기본 30%)를 절대 넘지 않는다 — 채점의 과잉강조
      게이트(0.35) 직전에서 차단된다.
    - **이미 강조된 단락(원본 포함)도 예산에 합산**한다. 원본이 이미 많이 굵으면
      추가 강조를 줄이거나 하지 않는다(재실행해도 누적 폭주 방지 = 멱등성).
    - '키워드 + 실제 숫자' 가 함께 있는 단락만 강조(require_numeric, 단독 한글 단위 제외).
    - 표 안 단락 제외(``_iter_body_paragraphs``), 길이 8자 미만(제목/캡션) 제외,
      이미 강조된 단락 제외.
    - ``max_emphasis`` 가 주어지면 추가 건수 절대 상한으로도 적용(하위호환, 기본 None).
    underline=True 면 밑줄도 함께 추가(기본 off).
    """
    paras = list(_iter_body_paragraphs(doc))
    total_nonempty = sum(1 for p in paras if _paragraph_text(p).strip())
    if total_nonempty == 0:
        return 0

    # 이미 강조된 단락(원본 포함) — 예산에서 차감
    existing_bold = sum(1 for p in paras if _para_is_bold(p))
    # 비율 기반 허용 총량(원본 포함). 작은 문서도 최소 강조는 가능하되 hard cap 으로 제한.
    hard_cap = int(total_nonempty * hard_emphasis_ratio)
    allowed_total = max(min_emphasis, round(total_nonempty * max_emphasis_ratio))
    allowed_total = min(allowed_total, hard_cap) if hard_cap > 0 else min(allowed_total, 1)
    budget = allowed_total - existing_bold
    if max_emphasis is not None:
        budget = min(budget, max_emphasis)
    if budget <= 0:
        return 0

    emphasized = 0
    for para in paras:
        if emphasized >= budget:
            break
        text = _paragraph_text(para)
        if len(text.strip()) < 8:
            continue
        if _para_is_bold(para):
            continue
        if not any(kw in text for kw in keywords):
            continue
        if require_numeric and not _NUMERIC_RE.search(text):
            continue
        runs = para.findall(qn("w:r"))
        if not runs:
            continue
        applied = False
        for run in runs:
            # 텍스트가 있는 런만
            if not _run_has_text(run):
                continue
            rpr = run.find(qn("w:rPr"))
            if rpr is None:
                rpr = run.makeelement(qn("w:rPr"), {})
                run.insert(0, rpr)
            if rpr.find(qn("w:b")) is None:
                rpr.append(run.makeelement(qn("w:b"), {}))
                applied = True
            if underline and rpr.find(qn("w:u")) is None:
                u = run.makeelement(qn("w:u"), {qn("w:val"): "single"})
                rpr.append(u)
                applied = True
        if applied:
            emphasized += 1
    return emphasized


# ---------------------------------------------------------------------------
# 5. 문단 수준별 글자 크기 표준화 (보수적/옵션)
# ---------------------------------------------------------------------------

def normalize_font_sizes(
    doc: Document,
    *,
    body_pt: float = 11.0,
    min_pt: float = 9.0,
    max_pt: float = 16.0,
    enable: bool = False,
) -> int:
    """본문 런의 글자 크기 이상치를 표준 범위로 보정(기본 비활성).

    문서 서식을 깨뜨릴 위험이 있어 ``enable=True`` 일 때만 동작한다. 동작 시:
    - 본문(표 밖) 단락의 런 중 크기가 명시되어 있고 [min_pt, max_pt] 범위를 벗어난 것만
      body_pt 로 보정한다(범위 내 값과 제목/표는 건드리지 않음).
    """
    if not enable:
        return 0
    normalized = 0
    for para in _iter_body_paragraphs(doc):
        for run in para.findall(qn("w:r")):
            rpr = run.find(qn("w:rPr"))
            if rpr is None:
                continue
            sz = rpr.find(qn("w:sz"))
            if sz is None:
                continue
            try:
                half_pt = float(sz.get(qn("w:val")))
            except (TypeError, ValueError):
                continue
            pt = half_pt / 2.0
            if pt < min_pt or pt > max_pt:
                sz.set(qn("w:val"), str(int(body_pt * 2)))
                szcs = rpr.find(qn("w:szCs"))
                if szcs is not None:
                    szcs.set(qn("w:val"), str(int(body_pt * 2)))
                normalized += 1
    return normalized


# ---------------------------------------------------------------------------
# 5-b. 단락별 텍스트 서식 통일 (크기·글꼴) — 지배값 기반·날조 없음
# ---------------------------------------------------------------------------

def _dominant(values: list[str | None]) -> str | None:
    """문서 순서를 유지하며 '가장 빈도 높은 실제 값'을 반환(동률은 첫 등장 우선).

    None(명시값 없음=테마 상속)은 후보에서 제외한다. 모든 값이 None 이면 None.
    """
    counts: dict[str, int] = {}
    order: list[str] = []
    for v in values:
        if v is None:
            continue
        if v not in counts:
            counts[v] = 0
            order.append(v)
        counts[v] += 1
    if not order:
        return None
    best = order[0]
    for v in order:
        if counts[v] > counts[best]:  # strictly greater → 동률 시 첫 등장 유지
            best = v
    return best


def _set_rpr_size(rpr, half_pt_val: str) -> bool:
    """rPr 의 w:sz / w:szCs 를 half_pt_val 로 맞춘다. 변경이 있으면 True."""
    changed = False
    for tag in ("w:sz", "w:szCs"):
        el = rpr.find(qn(tag))
        if el is None:
            el = rpr.makeelement(qn(tag), {})
            rpr.append(el)
            changed = True
        if el.get(qn("w:val")) != half_pt_val:
            el.set(qn("w:val"), half_pt_val)
            changed = True
    return changed


def _set_rpr_fonts(rpr, ascii_val: str | None, ea_val: str | None) -> bool:
    """rPr 의 w:rFonts 글꼴(ascii/hAnsi/eastAsia/cs)을 지배값으로 맞춘다. 변경 시 True."""
    changed = False
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)  # w:rFonts 는 rPr 의 첫 자식이어야 한다
        changed = True

    def _apply(attr: str, val: str | None) -> None:
        nonlocal changed
        if val is None:
            return
        if rfonts.get(qn(attr)) != val:
            rfonts.set(qn(attr), val)
            changed = True

    if ascii_val is not None:
        _apply("w:ascii", ascii_val)
        _apply("w:hAnsi", ascii_val)
    if ea_val is not None:
        _apply("w:eastAsia", ea_val)
        _apply("w:cs", ea_val)
    return changed


def _para_style_name(para) -> str:
    try:
        return para.style.name or ""
    except Exception:
        return ""


def unify_paragraph_formatting(
    doc: Document,
    *,
    scope: str = "body+tables",
    preserve_emphasis: bool = True,
    enable: bool = True,
) -> int:
    """단락별로 텍스트 런의 글자 크기·글꼴(font family)을 '그 단락의 지배값'으로 통일.

    "단락별 서식 통일" — 한 단락 안에서 런마다 크기/글꼴이 들쭉날쭉한 것을 그 단락에서
    가장 많이 쓰인 '이미 존재하는 값(지배값)'으로 맞춘다.

    안전 원칙:
    - 날조 금지: 어떤 런에도 명시 크기/글꼴이 없으면(전부 테마·스타일 상속) 그 단락은
      건드리지 않는다(테마 상속 보존). 지배값은 문서에 실재하는 값 중에서만 고른다.
    - 강조 보존(preserve_emphasis): w:b / w:u 는 손대지 않는다(크기·글꼴만 통일).
    - 제목/소제목(스타일명이 Heading/Title/제목)·이미지/도형 포함 단락은 건너뛴다.
    - 멱등성: 한 번 통일하면 재실행 시 0(이미 단일 값).
    - ``enable=False`` 면 0(비활성). 단, run_all 기본은 활성(normalize_font_sizes 처럼
      죽은 코드가 되지 않도록 live 기본값).

    scope 에 'tables' 가 포함되면 표 셀 단락까지 대상(병합셀 중복 제외).
    """
    if not enable:
        return 0

    paragraphs = list(doc.paragraphs)
    if "tables" in scope:
        seen: set[int] = set()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cid = id(cell._tc)
                    if cid in seen:
                        continue
                    seen.add(cid)
                    paragraphs.extend(cell.paragraphs)

    changed = 0
    for para in paragraphs:
        if _element_has_drawing(para._p):
            continue
        style_name = _para_style_name(para)
        if style_name.startswith(("Heading", "Title")) or "제목" in style_name:
            continue
        text_runs = [r for r in para.runs if _run_has_text(r._element)]
        if not text_runs:
            continue

        sizes: list[str | None] = []
        ascii_fonts: list[str | None] = []
        ea_fonts: list[str | None] = []
        for r in text_runs:
            rpr = r._element.find(qn("w:rPr"))
            sz = asc = ea = None
            if rpr is not None:
                sz_el = rpr.find(qn("w:sz"))
                if sz_el is not None:
                    sz = sz_el.get(qn("w:val"))
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is not None:
                    asc = rfonts.get(qn("w:ascii"))
                    ea = rfonts.get(qn("w:eastAsia"))
            sizes.append(sz)
            ascii_fonts.append(asc)
            ea_fonts.append(ea)

        dom_sz = _dominant(sizes)
        dom_asc = _dominant(ascii_fonts)
        dom_ea = _dominant(ea_fonts)
        if dom_sz is None and dom_asc is None and dom_ea is None:
            continue  # 전부 테마 상속 → 보존(날조 금지)

        para_changed = False
        for r in text_runs:
            rpr = r._element.find(qn("w:rPr"))
            if rpr is None:
                rpr = r._element.makeelement(qn("w:rPr"), {})
                r._element.insert(0, rpr)
            if dom_sz is not None and _set_rpr_size(rpr, dom_sz):
                para_changed = True
            if (dom_asc is not None or dom_ea is not None) and _set_rpr_fonts(rpr, dom_asc, dom_ea):
                para_changed = True
        if para_changed:
            changed += 1
    return changed


# ---------------------------------------------------------------------------
# 6. 명백한 양식 안내 단락 삭제 (보수적)
# ---------------------------------------------------------------------------

def remove_guide_paragraphs(doc: Document, *, max_len: int = 120) -> int:
    """'작성요령/예시/유의사항' 등 명백한 양식 안내 단락을 삭제(보수적).

    오삭제를 막기 위해:
    - body 직계 단락만 대상(표 셀은 보존 — qa_service/docx_ops 가 음영·색상 처리 담당)
    - ``_PURE_GUIDE_RE`` 로 시작하는 단락만 삭제(부분 포함은 제외)
    - 길이가 ``max_len`` 초과면 실제 내용일 수 있어 제외
    - 이미지/도형 포함 단락 제외
    - 문서 맨 끝 단락은 보존
    """
    body = doc.element.body
    removed = 0
    for para in list(_iter_body_paragraphs(doc)):
        if para is body[-1]:
            continue
        if _element_has_drawing(para):
            continue
        text = _paragraph_text(para)
        if not text or len(text) > max_len:
            continue
        if _is_guide_text(text):
            body.remove(para)
            removed += 1
    return removed


# ---------------------------------------------------------------------------
# 6-b. 표 셀에 박힌 양식 안내문구 삭제 (보수적) — body 직계 사각지대 보완
# ---------------------------------------------------------------------------

def _cell_row_text(row) -> str:
    """행의 모든 셀 텍스트를 합쳐 공백 정규화한 문자열. 병합셀 중복은 1회만."""
    parts: list[str] = []
    seen: set[int] = set()
    for cell in row.cells:
        tc_id = id(cell._tc)
        if tc_id in seen:
            continue
        seen.add(tc_id)
        parts.append(cell.text)
    return re.sub(r"\s+", " ", " ".join(parts)).strip()


def remove_table_guide_rows(doc: Document, *, max_len: int = 500) -> int:
    """표 셀에 박힌 '명백한 양식 안내문구'를 삭제한다(보수적).

    ``remove_guide_paragraphs`` 가 body 직계 단락만 보던 사각지대(표 셀)를 보완하되,
    동일한 ``_PURE_GUIDE_RE`` 기준만 사용해 일관성을 유지한다.

    - 표의 '비어있지 않은 모든 행' 이 ``_PURE_GUIDE_RE`` 로 시작 매칭되면
      (= 안내 전용 표) 표를 통째 제거한다.
    - 데이터 행과 안내 행이 섞인 표에서는 ``_PURE_GUIDE_RE`` 매칭 안내 행만 제거하되,
      데이터 행이 1개 이상 남는 경우에만 수행한다(표 골격 보존).

    오삭제 방지:
    - ``※``/``작성요령``/``예시`` 등으로 '시작' 하는 행만 대상(본문 데이터 행은 보존)
    - 이미지/도형이 든 행은 보존
    - 행 텍스트가 ``max_len`` 초과면 실제 내용일 수 있어 보존
    """
    removed = 0
    for table in list(doc.tables):
        flags: list[tuple] = []  # (row, is_guide, has_text)
        for row in table.rows:
            has_img = any(_element_has_drawing(cell._tc) for cell in row.cells)
            text = _cell_row_text(row)
            is_guide = (
                bool(text)
                and not has_img
                and len(text) <= max_len
                and _is_guide_text(text)
            )
            flags.append((row, is_guide, bool(text)))

        nonempty = [f for f in flags if f[2]]
        if not nonempty:
            continue

        if all(f[1] for f in nonempty):
            # 안내 전용 표 → 통째 제거
            tbl = table._tbl
            parent = tbl.getparent()
            if parent is not None:
                parent.remove(tbl)
                removed += len(nonempty)
            continue

        # 혼합 표: 데이터 행이 남을 때만 안내 행 제거
        if not any((not f[1]) for f in nonempty):
            continue
        for row, is_guide, _has_text in flags:
            if is_guide:
                tr = row._tr
                tp = tr.getparent()
                if tp is not None:
                    tp.remove(tr)
                    removed += 1
    return removed


# ---------------------------------------------------------------------------
# 통합 실행기
# ---------------------------------------------------------------------------

def _iter_all_paragraph_elements(doc: Document):
    """본문 직계 + 표 셀(중첩 포함) + 텍스트박스 + 머리글/바닥글 단락(w:p lxml 요소)을
    중복 없이 순회한다.

    유색 안내문구는 머리글·바닥글·텍스트박스에도 들어가므로(검출 ACC-3 와 동일 범위로
    정렬), 본문/표만 보던 사각지대를 보완한다.
    """
    for p in _iter_body_paragraphs(doc):
        yield p
    seen: set[int] = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cid = id(cell._tc)
                if cid in seen:
                    continue
                seen.add(cid)
                for p in cell._tc.iter(qn("w:p")):
                    yield p
    # 텍스트박스(w:txbxContent) 안 단락
    for txbx in doc.element.body.iter(qn("w:txbxContent")):
        for p in txbx.iter(qn("w:p")):
            yield p
    # 머리글·바닥글(명시 정의가 있을 때만 — 빈 linked 머리글 접근 시 part 신규생성 부작용 회피)
    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf is None or hf.is_linked_to_previous:
                continue
            for p in hf.paragraphs:
                yield p._p
            for tbl in hf.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p._p


def normalize_colored_text_to_black(doc: Document, *, enable: bool = True) -> int:
    """검정 아닌 유색 텍스트(파란 안내문구·회색 가이드 등)를 검정으로 정규화한다.

    ``usage_acceptance.check_residual_colored_runs`` (ACC-3) 가 fail 로 잡는 결함을
    교정한다 — '명시적으로 6자리 hex 색이 지정됐고 검정(000000)·보존색
    (_PRESERVE_COLORS, 흰색 계열)이 아닌' 런의 색만 검정으로 바꾼다.

    안전 원칙(최소 변경, 검출과 정합):
    - 색 미지정(테마 상속)·검정·보존색·테마색(w:themeColor)·``w:val="auto"`` 등
      **정규 6자리 hex 가 아닌 색은 건드리지 않는다**(검출도 비결함으로 보므로 — 역연산 정합).
    - 텍스트는 변경하지 않는다(날조 0). 굵게/밑줄 등 강조는 보존(색만 변경).
    - 검증된 ``docx_ops._set_run_color_black_unless_preserved`` 를 재사용(단일 출처).
    - 순회 범위는 본문+표+텍스트박스+머리글/바닥글(검출과 동일). 다만 검출이 제외하는
      자기삽입 블록(NotebookLM/스캐폴드) 런도 여기선 검정화될 수 있다 — 그 블록은
      strip/submit-clean 으로 별도 제거되므로 무해(게이트는 self_inserted_blocks 로 DRAFT 유지).

    Returns:
        색을 검정으로 바꾼 런 수(멱등 — 두 번째 실행은 0).
    """
    if not enable:
        return 0
    changed = 0
    for para in _iter_all_paragraph_elements(doc):
        for run in para.findall(qn("w:r")):
            if not "".join(t.text or "" for t in run.findall(qn("w:t"))).strip():
                continue  # 보이는 텍스트가 있는 런만
            rpr = run.find(qn("w:rPr"))
            if rpr is None:
                continue
            color = rpr.find(qn("w:color"))
            if color is None:
                continue  # 색 미지정(상속) — 검출도 통과시키므로 보존
            hexv = _normalize_color_value(
                color.get(qn("w:val")) or color.get("w:val") or color.get("val")
            )
            # 정규 6자리 hex 만 대상 — 'auto'·테마·단축형 등은 검출(color.rgb=None)이
            # 비결함으로 보므로 교정도 보존한다(역연산 계약·멱등성 유지).
            if not re.fullmatch(r"[0-9a-f]{6}", hexv) or hexv == "000000" or hexv in _PRESERVE_COLORS:
                continue
            _set_run_color_black_unless_preserved(run)
            changed += 1
    return changed


def run_all(
    doc: Document,
    *,
    remove_guides: bool = True,
    emphasize: bool = True,
    underline: bool = False,
    normalize_fonts: bool = False,
    unify_formatting: bool = True,
    normalize_colors: bool = True,
) -> QualityOpsReport:
    """모든 결정론적 후처리를 안전한 순서로 1회 적용하고 집계 리포트를 반환한다.

    순서: 안내삭제(body+표) → 글머리표공백 → 표공백 → 빈단락 → 서식통일 → 유색→검정 → 강조 → (옵션)폰트

    ``unify_formatting`` 은 기본 활성(live) — 단락별 크기/글꼴을 지배값으로 통일한다.
    강조(emphasize)보다 먼저 실행해 굵게 처리한 런의 서식이 덮이지 않게 한다.
    """
    report = QualityOpsReport()
    if remove_guides:
        report.guide_paragraphs_removed = remove_guide_paragraphs(doc)
        # 표 셀에 박힌 양식 안내문구도 제거(멱등 결함예방). 기존 함수가 run_all 에
        # 연결돼 있지 않던 사각지대 보완 — 데이터 행이 있는 표는 안내 행만, 안내 전용
        # 표는 통째 제거되며 데이터 행은 보존된다.
        report.table_guide_rows_removed = remove_table_guide_rows(doc)
    report.bullet_spacing_fixed = normalize_bullet_spacing(doc)
    report.table_cells_cleaned = cleanup_table_whitespace(doc)
    report.empty_paragraphs_removed = remove_empty_paragraphs(doc)
    report.paragraphs_unified = unify_paragraph_formatting(doc, enable=unify_formatting)
    report.colored_runs_normalized = normalize_colored_text_to_black(doc, enable=normalize_colors)
    if emphasize:
        report.paragraphs_emphasized = emphasize_key_sentences(doc, underline=underline)
    report.font_sizes_normalized = normalize_font_sizes(doc, enable=normalize_fonts)
    return report
