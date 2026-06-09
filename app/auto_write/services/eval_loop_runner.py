"""eval_loop_runner.py

공고 평가기준 채점 → 취약 섹션 식별 → 부분 재생성 → 재채점을, 목표점수 도달 또는
수렴까지 반복하는 Evaluation Loop 러너.

조합:
- EvaluationService : 공고 파싱·AI 채점·취약 섹션 식별·보완 컨텍스트 생성(채점 엔진)
- ProjectService.regenerate_sections : 취약 섹션만 재작성 후 재렌더(생성 엔진)

도메인 원칙:
- 원문에 없는 내용 생성 금지. 근거 부족으로 점수가 낮아 더 못 올리는 항목은
  창작하지 않고 needs_input(사용자 보완 입력 필요)로 보고한다.
- AI 채점 변동 흡수: 회당 scoring_passes 회 채점 후 보수적으로 '하한(min) 총점'을
  대표값으로 사용한다.
- 게이트(제출가능 판정)는 SubmissionPipeline에서 pass_ratio>=target & needs_input==0 &
  서식점수>=85 로 종합한다. 본 러너는 평가 루프 결과(EvalLoopReport)만 만든다.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

from ..utils import log_line
from .evaluation_service import EvalCriterion, EvalLoopReport, EvalResult


class EvalLoopRunner:
    def __init__(self, evaluation_service: Any, project_service: Any, storage: Any) -> None:
        self.evaluation_service = evaluation_service
        self.project_service = project_service
        self.storage = storage

    @staticmethod
    def _extract_doc_text(output_docx: Path) -> str:
        doc = Document(str(output_docx))
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(dict.fromkeys(cells)))
        return "\n".join(lines)

    @staticmethod
    def _total(scores: list[Any]) -> int:
        return sum(s.score for s in scores)

    def _score_conservative(
        self, doc_text: str, criteria: list[EvalCriterion], profile_questions: list[dict], scoring_passes: int
    ) -> list[Any]:
        """scoring_passes 회 채점 후 총점이 가장 낮은(보수적) 패스를 대표로 반환."""
        best: list[Any] | None = None
        for _ in range(max(1, scoring_passes)):
            scores = self.evaluation_service.score_document(doc_text, criteria, profile_questions)
            if best is None or self._total(scores) < self._total(best):
                best = scores
        return best or []

    def run(
        self,
        project_id: str,
        announcement_text: str = "",
        *,
        criteria: list[EvalCriterion] | None = None,
        target_score: int = 92,
        max_iterations: int = 3,
        scoring_passes: int = 2,
    ) -> EvalLoopReport:
        profile = self.project_service.load_profile_for_project(project_id)
        profile_questions = [q.model_dump() for q in profile.questions]
        valid_qids = {str(q.get("question_id", "")) for q in profile_questions if q.get("question_id")}

        if criteria is None:
            criteria = (
                self.evaluation_service.parse_announcement(announcement_text)
                if announcement_text.strip()
                else []
            )

        output_docx = self.storage.project_dir(project_id) / "output" / "output.docx"
        target_ratio = max(0.0, min(target_score / 100.0, 1.0))
        min_ratio = getattr(self.evaluation_service, "MIN_REWRITE_SCORE_RATIO", 0.6)

        iterations: list[EvalResult] = []
        needs_input: list[str] = []
        converged = False
        prev_total = -1
        max_iterations = max(1, int(max_iterations))

        for iteration in range(1, max_iterations + 1):
            if not criteria or not output_docx.exists():
                break
            doc_text = self._extract_doc_text(output_docx)
            rep_scores = self._score_conservative(doc_text, criteria, profile_questions, scoring_passes)
            result = self.evaluation_service.build_eval_result(iteration, rep_scores, profile_questions)
            iterations.append(result)
            log_line(
                f"[EvalLoop] iter={iteration} total={result.total_score}/{result.max_total} "
                f"ratio={result.pass_ratio:.3f}"
            )

            if result.pass_ratio >= target_ratio:
                converged = True
                break

            plateau = prev_total >= 0 and (result.total_score - prev_total) < 2
            prev_total = result.total_score

            weak_q = [q for q in result.weak_sections if q in valid_qids]
            weak_q = list(dict.fromkeys(weak_q))

            if plateau or not weak_q or iteration >= max_iterations:
                converged = converged or plateau
                break

            refinement = self.evaluation_service.build_refinement_context(
                weak_q, result.criteria_scores, profile_questions, ""
            )
            self.project_service.regenerate_sections(
                project_id, weak_q, refinement_context=refinement
            )

        final = iterations[-1] if iterations else None
        if final is not None and final.pass_ratio < target_ratio:
            for s in final.criteria_scores:
                if s.ratio < min_ratio and s.name not in needs_input:
                    needs_input.append(s.name)

        return EvalLoopReport(
            project_id=project_id,
            iterations=iterations,
            final_score=final.total_score if final else 0,
            final_max=final.max_total if final else 0,
            final_pass_ratio=final.pass_ratio if final else 0.0,
            converged=converged,
            announcement_criteria=[
                {"name": c.name, "max_score": c.max_score, "description": c.description}
                for c in criteria
            ],
            needs_input=needs_input,
        )
