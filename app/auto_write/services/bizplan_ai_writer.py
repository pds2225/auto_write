"""bizplan_ai_writer.py — PSST 약점 영역 본문을 'AI가 근거 명시하며' 작성/보강한다.

사용자 정책(승인): AI 가 빈 영역의 초안까지 채우되,
  - 입력(초안/브리프)에 있는 사실은 단정 기술,
  - 추정·시장수치·전망에는 **[산출근거: 출처(기관·연도), 계산식]** 을 반드시 병기,
  - 근거는 통계청/KOSIS/정부부처 등 **공식 자료 우선**,
  - 출처를 댈 수 없는 핵심 수치는 지어내지 말고 **[확인필요: …]** 로 표시한다.

정부지원사업 허위기재는 형사처벌 대상이므로, '무출처 날조' 는 프롬프트·후처리로 차단한다.

AI 키가 없으면 아무 것도 쓰지 않고(skipped=True) 안전하게 빠진다 — 이때는
``psst_fill`` 의 결정론적 작성 가이드가 대신 채운다(autopilot 단계에서).
"""

from __future__ import annotations

import json
import shutil
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.shared import Pt, RGBColor

from .psst_check import check_psst
# 검출기(usage_acceptance._SELF_BLOCK_RE)와 동일 정의의 표준 문구(패턴 공유 원칙)
from .usage_acceptance import AI_SECTION_DELETE_NOTICE, AI_SECTION_HEADING

_TARGET_GRADES = ("누락", "미흡")

_SYSTEM_PROMPT = (
    "당신은 한국 정부지원사업 사업계획서 전문 작성자입니다. 주어진 PSST 영역의 본문을 "
    "심사위원이 읽기 좋게 작성/보강하세요.\n\n"
    "[필수 규칙 — 위반 시 결과 폐기]\n"
    "1. 입력(초안/브리프)에 실제로 있는 사실만 단정적으로 기술한다.\n"
    "2. 추정·시장규모·성장률·전망 등 수치를 쓸 때는 반드시 문장 끝에 "
    "[산출근거: 출처(기관·연도), 계산식] 을 병기한다.\n"
    "3. 근거 출처는 통계청·KOSIS·정부부처·공신력 있는 기관 등 공식 자료를 우선한다.\n"
    "4. 출처를 댈 수 없는 핵심 수치는 절대 지어내지 말고 '[확인필요: 무엇]' 으로 표시한다.\n"
    "5. 허위·과장 금지(정부지원 허위기재는 형사처벌·환수·참여제한 대상).\n\n"
    "반환 형식(JSON object):\n"
    '{"paragraphs": ["문단1", "문단2", ...], '
    '"needs_confirm": ["사용자가 직접 확인/입력해야 할 항목"], '
    '"evidence_used": ["인용한 공식 출처"]}'
)


@dataclass
class AiWriteReport:
    areas_written: int = 0
    paragraphs_added: int = 0
    needs_confirm: list[str] = field(default_factory=list)
    evidence_used: list[str] = field(default_factory=list)
    written_areas: list[str] = field(default_factory=list)
    skipped: bool = False
    skip_reason: str = ""
    output_docx: str = ""

    def as_dict(self) -> dict[str, Any]:
        return {
            "areas_written": self.areas_written,
            "paragraphs_added": self.paragraphs_added,
            "needs_confirm": self.needs_confirm,
            "evidence_used": self.evidence_used,
            "written_areas": self.written_areas,
            "skipped": self.skipped,
            "skip_reason": self.skip_reason,
            "output_docx": self.output_docx,
        }


def _doc_text(doc: Document, *, limit: int = 8000) -> str:
    parts: list[str] = []
    total = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
            total += len(t)
            if total >= limit:
                break
    return "\n".join(parts)


def _add_heading(doc: Document, text: str, *, size: int = 13) -> None:
    run = doc.add_paragraph().add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def _add_area_title(doc: Document, text: str) -> None:
    run = doc.add_paragraph().add_run(text)
    run.bold = True
    run.font.size = Pt(11)


def _add_body(doc: Document, text: str) -> None:
    doc.add_paragraph().add_run(text)


def _add_note(doc: Document, text: str) -> None:
    run = doc.add_paragraph().add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    try:
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    except Exception:  # pragma: no cover
        pass


def ai_write_areas(
    in_docx: str,
    out_docx: str,
    *,
    brief: str = "",
    openai_service: Optional[Any] = None,
    target_grades: tuple[str, ...] = _TARGET_GRADES,
    max_paras_per_area: int = 4,
) -> AiWriteReport:
    """PSST 약점 영역 본문을 AI 로 작성/보강해 문서 끝 보강 섹션에 삽입한다.

    Args:
        in_docx: 원본 DOCX(읽기 전용).
        out_docx: 결과 DOCX. **in_docx 와 같으면 ValueError**.
        brief: 사용자 사업 브리프(아이디어·팀·수치 등 자유 텍스트).
        openai_service: ``OpenAIService`` 인스턴스. None 또는 미가용이면 skipped.
        target_grades: 작성 대상 등급(기본: 누락·미흡).

    Returns:
        AiWriteReport. AI 미가용 시 skipped=True 로 반환(문서는 복사만).
    """
    in_path = Path(in_docx)
    out_path = Path(out_docx)
    if in_path.resolve() == out_path.resolve():
        raise ValueError("in_docx 와 out_docx 가 같습니다. 원본 덮어쓰기는 금지입니다.")
    if not in_path.exists():
        raise FileNotFoundError(f"입력 DOCX 가 없습니다: {in_docx}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(str(in_path), str(out_path))

    doc = Document(str(out_path))
    report = AiWriteReport(output_docx=str(out_path))

    available = bool(openai_service is not None and getattr(openai_service, "available", False))
    if not available:
        report.skipped = True
        report.skip_reason = "AI 키 미연결 — psst_fill 가이드로 폴백"
        doc.save(str(out_path))
        return report

    psst = check_psst(doc)
    weak = [a for a in psst.areas if a.grade in target_grades]
    if not weak:
        doc.save(str(out_path))
        return report

    context_text = _doc_text(doc)
    wrote_any = False
    for area in weak:
        user_prompt = json.dumps({
            "area": area.label,
            "missing_items": area.missing_items,
            "max_paragraphs": max_paras_per_area,
            "user_brief": brief[:4000],
            "current_document_excerpt": context_text[:5000],
        }, ensure_ascii=False)
        result = openai_service.complete_json(_SYSTEM_PROMPT, user_prompt, max_tokens=2048)
        if not isinstance(result, dict):
            continue
        paragraphs = [str(p).strip() for p in result.get("paragraphs", []) if str(p).strip()]
        if not paragraphs:
            continue
        if not wrote_any:
            _add_heading(doc, AI_SECTION_HEADING)
            _add_note(doc, AI_SECTION_DELETE_NOTICE)
            wrote_any = True
        _add_area_title(doc, f"▶ {area.label}")
        for para in paragraphs[:max_paras_per_area]:
            _add_body(doc, para)
            report.paragraphs_added += 1
        nc = [str(x).strip() for x in result.get("needs_confirm", []) if str(x).strip()]
        ev = [str(x).strip() for x in result.get("evidence_used", []) if str(x).strip()]
        if nc:
            _add_note(doc, "확인필요: " + " / ".join(nc))
            report.needs_confirm.extend(nc)
        report.evidence_used.extend(ev)
        report.areas_written += 1
        report.written_areas.append(area.label)

    doc.save(str(out_path))
    return report
