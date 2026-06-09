from __future__ import annotations

import copy
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from lxml import etree

_PRESERVE_COLORS = {"ffffff", "fffffe", "f2f2f2"}
_PLACEHOLDER_SHADE_FILLS = {"ffff00", "fff2cc", "ffeb9c", "ffe699", "ffd966"}
GUIDE_MARKER_RE = re.compile(r"(※|<[^>]+>|기재|작성요령|작성방법|예시|OOO|○○○)")


def _normalize_match_text(text: str) -> str:
    normalized = re.sub(r"\s+", " ", text or "").strip()
    normalized = normalized.replace("–", "-").replace("—", "-").replace("‑", "-")
    for _ in range(2):
        if len(normalized) < 16 or len(normalized) % 2 != 0:
            break
        half = len(normalized) // 2
        left = normalized[:half].strip()
        right = normalized[half:].strip()
        if left and left == right:
            normalized = left
            continue
        break
    return normalized


def _match_anchor(text: str, anchor_text: str) -> bool:
    anchor = _normalize_match_text(anchor_text)
    candidate = _normalize_match_text(text)
    if not anchor or not candidate:
        return False
    if anchor in candidate:
        return True
    anchor_compact = anchor.replace(" ", "")
    candidate_compact = candidate.replace(" ", "")
    if anchor_compact and anchor_compact in candidate_compact:
        return True
    return False


def _paragraph_text(paragraph: etree._Element) -> str:
    return "".join(node.text or "" for node in paragraph.iter(qn("w:t"))).strip()


def _iter_body_paragraphs(doc: Document) -> list[etree._Element]:
    body = doc.element.body
    return [element for element in body if element.tag == qn("w:p")]


def _normalize_color_value(value: str | None) -> str:
    return re.sub(r"[^0-9a-fA-F]", "", str(value or "")).lower()


def _set_run_color_black_unless_preserved(run: etree._Element) -> None:
    rpr = run.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run.insert(0, rpr)
    highlight = rpr.find(qn("w:highlight"))
    if highlight is not None:
        rpr.remove(highlight)
    shd = rpr.find(qn("w:shd"))
    if shd is not None:
        rpr.remove(shd)
    color = rpr.find(qn("w:color"))
    current_color = ""
    if color is not None:
        current_color = _normalize_color_value(
            color.get(qn("w:val")) or color.get("w:val") or color.get("val")
        )
    if current_color in _PRESERVE_COLORS:
        return
    if color is None:
        color = OxmlElement("w:color")
        rpr.append(color)
    color.set(qn("w:val"), "000000")


def _write_para(para, text: str) -> None:
    runs = para.findall(qn("w:r"))
    if runs:
        for run in runs:
            for text_node in run.findall(qn("w:t")):
                text_node.text = ""
        target_run = runs[0]
        _set_run_color_black_unless_preserved(target_run)
        text_nodes = target_run.findall(qn("w:t"))
        if text_nodes:
            text_nodes[0].text = text
        else:
            element = etree.SubElement(target_run, qn("w:t"))
            element.text = text
    else:
        run = etree.SubElement(para, qn("w:r"))
        _set_run_color_black_unless_preserved(run)
        element = etree.SubElement(run, qn("w:t"))
        element.text = text


def logical_cells(row):
    """행의 '논리 셀'(서로 다른 w:tc)을 등장 순서대로 반환(가로 병합 중복 제거).

    python-docx 의 ``row.cells`` 는 가로 병합된 셀을 grid 칸 수만큼 '중복'해서 돌려준다.
    반면 템플릿 분석(analysis/docx_template)은 ``<w:tr>`` 안의 ``<w:tc>`` 를 직접 세어
    (= 논리 인덱스) 셀 좌표(TableCellProfile.cell)를 기록한다. 따라서 **텍스트 채움**은
    grid 인덱스가 아니라 이 논리 셀 목록으로 해석해야 병합 표에서도 같은 자리에 들어간다.
    (이미지 슬롯은 enumerate(row.cells) 기반의 grid 인덱스를 쓰므로 이 함수를 쓰지 않는다.)
    """
    result = []
    seen: set[int] = set()
    for cell in row.cells:
        cid = id(cell._tc)
        if cid in seen:
            continue
        seen.add(cid)
        result.append(cell)
    return result


def set_cell_text(cell, text: str) -> None:
    tc = cell._tc
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is not None:
        tc_shading = tc_pr.find(qn("w:shd"))
        if tc_shading is not None:
            fill = _normalize_color_value(
                tc_shading.get(qn("w:fill")) or tc_shading.get("w:fill") or tc_shading.get("fill")
            )
            if fill in _PLACEHOLDER_SHADE_FILLS:
                tc_pr.remove(tc_shading)
    paragraphs = tc.findall(qn("w:p"))
    if not paragraphs:
        return
    lines = text.splitlines() or [text]
    _write_para(paragraphs[0], lines[0])
    for extra in paragraphs[1:]:
        tc.remove(extra)
    ref = paragraphs[0]
    for line in lines[1:]:
        new_paragraph = copy.deepcopy(ref)
        for run in new_paragraph.findall(qn("w:r")):
            for text_node in run.findall(qn("w:t")):
                text_node.text = ""
        _write_para(new_paragraph, line)
        tc.append(new_paragraph)


def find_paragraph(doc: Document, anchor_text: str):
    for paragraph in _iter_body_paragraphs(doc):
        text = _paragraph_text(paragraph)
        if _match_anchor(text, anchor_text):
            return paragraph
    return None


def find_cell(doc: Document, anchor_text: str):
    anchor = _normalize_match_text(anchor_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text or ""
                if _match_anchor(text, anchor):
                    return cell
    return None


def insert_text_after_paragraph(doc: Document, anchor_text: str, text: str, allow_create_paragraph: bool = False) -> bool:
    paragraph = find_paragraph(doc, anchor_text)
    if paragraph is None:
        cell = find_cell(doc, anchor_text)
        if cell is None:
            return False
        existing = cell.text.strip()
        if existing and not GUIDE_MARKER_RE.search(existing):
            return False
        set_cell_text(cell, text)
        return True

    body = doc.element.body
    body_items = list(body)
    try:
        para_index = body_items.index(paragraph)
    except ValueError:
        para_index = -1
    if para_index >= 0:
        lines = text.splitlines() or [text]
        inserted = 0
        for next_index in range(para_index + 1, min(para_index + 12, len(body_items))):
            candidate = body_items[next_index]
            if candidate.tag != qn("w:p"):
                continue
            if _paragraph_text(candidate):
                continue
            _write_para(candidate, lines[inserted])
            inserted += 1
            if inserted >= len(lines):
                return True
        if inserted > 0:
            return True
    if not allow_create_paragraph or para_index < 0:
        return False
    reference = paragraph
    for offset, line in enumerate(text.splitlines() or [text]):
        new_paragraph = copy.deepcopy(reference)
        for run in new_paragraph.findall(qn("w:r")):
            for text_node in run.findall(qn("w:t")):
                text_node.text = ""
        _write_para(new_paragraph, line)
        body.insert(para_index + 1 + offset, new_paragraph)
    return True


def _get_image_xml(doc: Document, image_path: Path, width_cm: float):
    temp_paragraph = OxmlElement("w:p")
    doc.element.body.append(temp_paragraph)
    from docx.text.paragraph import Paragraph

    paragraph = Paragraph(temp_paragraph, doc)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    drawings = temp_paragraph.findall(".//" + qn("w:drawing"))
    drawing_xml = copy.deepcopy(drawings[0]) if drawings else None
    doc.element.body.remove(temp_paragraph)
    return drawing_xml


def insert_image_in_cell(doc: Document, table_index: int, row: int, cell_index: int, image_path: Path, width_cm: float = 12.0) -> bool:
    try:
        table = doc.tables[table_index]
        cell = table.rows[row].cells[cell_index]
    except Exception:
        return False
    drawing_xml = _get_image_xml(doc, image_path, width_cm)
    if drawing_xml is None:
        return False
    tc = cell._tc
    paragraphs = tc.findall(qn("w:p"))
    if not paragraphs:
        return False
    target = paragraphs[0]
    for paragraph in paragraphs[1:]:
        if _paragraph_text(paragraph):
            continue
        tc.remove(paragraph)
    for run in list(target.findall(qn("w:r"))):
        target.remove(run)
    new_run = OxmlElement("w:r")
    new_run.append(drawing_xml)
    target.append(new_run)
    return True


def _set_paragraph_center(paragraph: etree._Element) -> None:
    ppr = paragraph.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        paragraph.insert(0, ppr)
    jc = ppr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        ppr.append(jc)
    jc.set(qn("w:val"), "center")


def insert_image_after_paragraph(
    doc: Document,
    anchor_text: str,
    image_path: Path,
    width_cm: float = 14.0,
    insert_offset: int = 1,
    allow_create_paragraph: bool = False,
) -> bool:
    paragraph = find_paragraph(doc, anchor_text)
    if paragraph is None:
        return False
    drawing_xml = _get_image_xml(doc, image_path, width_cm)
    if drawing_xml is None:
        return False
    body = doc.element.body
    try:
        base_index = list(body).index(paragraph)
    except ValueError:
        return False

    body_items = list(body)
    for next_index in range(base_index + 1, min(base_index + 8, len(body_items))):
        candidate = body_items[next_index]
        if candidate.tag != qn("w:p"):
            continue
        if _paragraph_text(candidate):
            continue
        for run in list(candidate.findall(qn("w:r"))):
            candidate.remove(run)
        image_run = OxmlElement("w:r")
        image_run.append(copy.deepcopy(drawing_xml))
        candidate.append(image_run)
        _set_paragraph_center(candidate)
        return True

    if not allow_create_paragraph:
        return False
    image_paragraph = OxmlElement("w:p")
    _set_paragraph_center(image_paragraph)
    image_run = OxmlElement("w:r")
    image_run.append(drawing_xml)
    image_paragraph.append(image_run)
    body.insert(base_index + max(1, insert_offset), image_paragraph)
    return True
