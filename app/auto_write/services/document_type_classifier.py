"""document_type_classifier.py

생성/제출 문서의 **유형을 자동 분류**한다. 유형에 따라 적용할 품질 규칙
(구조 적합성, PSST 검사 여부 등)이 달라진다.

분류는 기본적으로 **규칙 기반(키워드 시그니처 + 가중 점수)** 이며, AI 키가 있고
``openai_service`` 가 주어지면 모호한 경우(상위 두 후보 점수차가 작을 때)에 한해
AI 보조 판정을 사용할 수 있다(선택).

유형 코드
---------
- business_plan        : 사업계획서
- rnd_plan             : R&D 연구개발계획서
- pitch_deck           : 발표평가 자료
- consulting_report    : 컨설팅 보고서
- policy_fund_report   : 정책자금 검토보고서
- certification_report : 인증 검토보고서
- export_report        : 수출컨설팅 보고서
- field_clinic_report  : 현장클리닉 보고서
- result_report        : 사업 결과·성과·정산 보고서
- generic_submission   : 기타 제출문서 (fallback)
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document


@dataclass
class DocTypeResult:
    type_code: str
    type_label: str
    confidence: float           # 0.0 ~ 1.0
    scores: dict[str, int] = field(default_factory=dict)
    matched_keywords: list[str] = field(default_factory=list)
    method: str = "rule"        # "rule" | "ai"

    def as_dict(self) -> dict[str, Any]:
        return {
            "type_code": self.type_code,
            "type_label": self.type_label,
            "confidence": round(self.confidence, 3),
            "method": self.method,
            "scores": self.scores,
            "matched_keywords": self.matched_keywords[:20],
        }


# 유형별 (라벨, 키워드 가중치) — 키워드는 부분일치, 가중치는 변별력에 비례
_TYPE_LABELS = {
    "business_plan": "사업계획서",
    "rnd_plan": "R&D 연구개발계획서",
    "pitch_deck": "발표평가 자료",
    "consulting_report": "컨설팅 보고서",
    "policy_fund_report": "정책자금 검토보고서",
    "certification_report": "인증 검토보고서",
    "export_report": "수출컨설팅 보고서",
    "field_clinic_report": "현장클리닉 보고서",
    "result_report": "사업 결과·성과 보고서",
    "generic_submission": "기타 제출문서",
}

_SIGNATURES: dict[str, list[tuple[str, int]]] = {
    "business_plan": [
        ("사업계획서", 5), ("문제인식", 3), ("실현가능성", 3), ("성장전략", 3),
        ("팀구성", 2), ("Problem", 2), ("Solution", 2), ("Scale", 2), ("Team", 2),
        ("창업아이템", 3), ("사업화", 2), ("PSST", 5), ("BM", 1), ("시장진입", 1),
    ],
    "rnd_plan": [
        ("연구개발", 4), ("기술개발목표", 4), ("R&D", 3), ("TRL", 4), ("성능지표", 3),
        ("실험방법", 3), ("연구개발계획서", 5), ("선행기술", 2), ("특허분석", 2),
        ("기술성", 2), ("개발내용", 2),
    ],
    "pitch_deck": [
        ("발표", 3), ("피칭", 4), ("IR", 3), ("데모데이", 4), ("발표평가", 5),
        ("Pitch", 3), ("투자유치", 2), ("Q&A", 2), ("슬라이드", 2),
    ],
    "consulting_report": [
        ("컨설팅", 4), ("기업현황", 3), ("진단결과", 4), ("개선과제", 4),
        ("실행계획", 2), ("기대효과", 2), ("컨설팅보고서", 5), ("경영진단", 3),
        ("SWOT", 1), ("As-Is", 2), ("To-Be", 2),
    ],
    "policy_fund_report": [
        ("정책자금", 5), ("자금용도", 4), ("상환재원", 5), ("매출추이", 3),
        ("담보", 2), ("보증", 2), ("신용평가", 3), ("융자", 3), ("운전자금", 2),
        ("시설자금", 2), ("리스크", 1),
    ],
    "certification_report": [
        ("인증", 4), ("인증요건", 5), ("충족여부", 4), ("미비서류", 4),
        ("보완과제", 3), ("벤처기업", 2), ("이노비즈", 3), ("메인비즈", 3),
        ("ISO", 2), ("인증검토", 5),
    ],
    "export_report": [
        ("수출", 4), ("수출컨설팅", 5), ("해외시장", 3), ("바이어", 4),
        ("FTA", 3), ("통관", 3), ("관세", 2), ("HS코드", 4), ("무역", 2),
        ("수출입", 3), ("글로벌진출", 2),
    ],
    "field_clinic_report": [
        ("현장클리닉", 5), ("현장진단", 4), ("클리닉", 4), ("현장방문", 3),
        ("애로사항", 3), ("처방", 3), ("현장지도", 4), ("개선처방", 4),
    ],
    # 사업 결과·성과·정산 보고서 — 고유 복합어 위주(bare "결과/성과/실적" 미사용)로
    # 기존 유형(특히 consulting_report "진단결과", rnd_plan "성능지표")과 충돌 방지.
    "result_report": [
        ("결과보고서", 5), ("성과보고서", 5), ("최종보고서", 4), ("완료보고서", 4),
        ("정산보고서", 4), ("사업비정산", 4), ("추진실적", 3), ("수행실적", 3),
        ("집행실적", 3), ("성과지표", 2), ("달성도", 2), ("정산", 2),
    ],
}

_MIN_SCORE = 4          # 이 점수 미만이면 generic 으로 분류
_AMBIGUITY_GAP = 3      # 1·2위 점수차가 이 값 이하면 모호 → (옵션)AI 보조


def _extract_text(doc: Document, *, limit: int = 12000) -> str:
    parts: list[str] = []
    total = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
            total += len(t)
            if total >= limit:
                return "\n".join(parts)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
                    total += len(t)
                    if total >= limit:
                        return "\n".join(parts)
    return "\n".join(parts)


def classify_text(text: str, *, filename: str = "") -> DocTypeResult:
    """문서 텍스트(+ 파일명)로 유형을 분류한다(규칙 기반)."""
    haystack = f"{filename}\n{text}"
    scores: dict[str, int] = {}
    matched: dict[str, list[str]] = {}
    for type_code, sigs in _SIGNATURES.items():
        s = 0
        hits: list[str] = []
        for kw, w in sigs:
            # 대소문자 무시(영문 키워드 대응)
            if kw.lower() in haystack.lower():
                s += w
                hits.append(kw)
        scores[type_code] = s
        matched[type_code] = hits

    ranked = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
    top_code, top_score = ranked[0]
    second_score = ranked[1][1] if len(ranked) > 1 else 0

    if top_score < _MIN_SCORE:
        return DocTypeResult(
            type_code="generic_submission",
            type_label=_TYPE_LABELS["generic_submission"],
            confidence=0.3,
            scores=scores,
            matched_keywords=[],
            method="rule",
        )

    # confidence: 1위 점수 비중 + 2위와의 격차 반영
    total = sum(max(0, v) for v in scores.values()) or 1
    gap_factor = min(1.0, (top_score - second_score) / max(1, top_score))
    confidence = min(0.99, 0.4 + 0.4 * (top_score / total) + 0.2 * gap_factor)

    return DocTypeResult(
        type_code=top_code,
        type_label=_TYPE_LABELS[top_code],
        confidence=confidence,
        scores=scores,
        matched_keywords=matched[top_code],
        method="rule",
    )


def classify_docx(
    path: str | Path,
    *,
    openai_service: Any | None = None,
) -> DocTypeResult:
    """DOCX 파일 경로로 유형을 분류한다.

    규칙 기반 결과가 모호(1·2위 격차 <= _AMBIGUITY_GAP)하고 ``openai_service`` 가
    사용 가능하면 AI 보조 판정을 시도한다(실패 시 규칙 결과 유지).
    """
    path = Path(path)
    doc = Document(str(path))
    text = _extract_text(doc)
    result = classify_text(text, filename=path.name)

    if openai_service is None or not getattr(openai_service, "available", False):
        return result

    ranked = sorted(result.scores.items(), key=lambda kv: kv[1], reverse=True)
    top = ranked[0][1]
    second = ranked[1][1] if len(ranked) > 1 else 0
    if (top - second) > _AMBIGUITY_GAP:
        return result  # 충분히 명확 → AI 불필요

    try:
        system = (
            "당신은 한국 정부지원사업 문서 분류 전문가입니다. 아래 후보 중 하나의 코드만 반환하세요.\n"
            + "\n".join(f"- {code}: {label}" for code, label in _TYPE_LABELS.items())
            + '\n반환 형식(JSON): {"type_code": "코드"}'
        )
        user = f"파일명: {path.name}\n본문 일부:\n{text[:4000]}"
        ai = openai_service.complete_json(system, user)
        code = str((ai or {}).get("type_code", "")).strip()
        if code in _TYPE_LABELS:
            result.type_code = code
            result.type_label = _TYPE_LABELS[code]
            result.method = "ai"
            result.confidence = max(result.confidence, 0.85)
    except Exception:
        pass
    return result
