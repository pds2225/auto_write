"""conversion_fidelity.py — DOCX↔HWP 변환 일치도 측정 하네스.

**구조(structural) 일치도만 측정한다** — 단락수·표수·셀수·이미지수·텍스트.
폰트·스타일·색상·레이아웃 등 시각 서식은 metric 범위 밖이다.
따라서 **구조 100% ≠ 시각 100%** 다(구조가 완전히 일치해도 글꼴/색/줄간격 등은
다를 수 있다). 이 한계는 리포트의 notes 에도 항상 명시된다.

설계 원칙
---------
- **결정론·AI無** — 동일 입력이면 동일 결과(난수·외부호출 없음).
- **입력 미수정** — 측정은 읽기 전용. roundtrip 중간 파일은 임시 디렉터리에.
- **COM 종속 정직 보고** — DOCX→HWP 는 한글 COM 대화형 전용이라
  COM 미가용/실패 시 측정을 생략하고 ``ok=False`` + 안내만 담는다(예외 전파 X).
"""

from __future__ import annotations

import re
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from docx import Document

# 텍스트 정규화: 공백 류를 단일 공백으로 접고 양끝을 다듬어 토큰/문자 일치율을 안정화.
_WS_RE = re.compile(r"\s+")

# 이미지 카운트용 XML 네임스페이스(drawing / blip).
_NS_DRAWING = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"

# overall 가중치 — 텍스트·표의 비중을 높인다(정부양식은 표·본문 중심).
_WEIGHTS = {
    "text": 0.35,
    "tables": 0.25,
    "cells": 0.20,
    "paragraphs": 0.10,
    "images": 0.10,
}


@dataclass
class FidelityReport:
    ok: bool = False
    overall_score: float = 0.0           # 0~100 종합 일치도
    metrics: dict[str, float] = field(default_factory=dict)   # metric별 일치율%
    counts: dict[str, Any] = field(default_factory=dict)      # a/b 측정치
    lost_items: list[str] = field(default_factory=list)       # 손실 항목 설명
    method: str = ""                     # "structural_compare" | "roundtrip" | ""
    notes: list[str] = field(default_factory=list)

    def as_dict(self) -> dict[str, Any]:
        return {
            "ok": self.ok,
            "overall_score": self.overall_score,
            "metrics": self.metrics,
            "counts": self.counts,
            "lost_items": self.lost_items,
            "method": self.method,
            "notes": self.notes,
        }


# --- 구조 측정 도우미 ---------------------------------------------------------

def _norm_text(text: str) -> str:
    return _WS_RE.sub(" ", text or "").strip()


def _count_images(doc: Document) -> int:
    """본문 XML 에서 그림(drawing) 요소 수를 센다.

    DOCX 의 이미지는 ``w:drawing`` 안의 ``a:blip``(실제 그림 참조)으로 표현된다.
    blip 이 있으면 그것을, 없으면 drawing 요소 수를 폴백으로 센다.
    """
    body = doc.element.body
    blips = body.findall(f".//{{{_NS_A}}}blip")
    if blips:
        return len(blips)
    return len(body.findall(f".//{{{_NS_DRAWING}}}drawing"))


def _measure(doc: Document) -> dict[str, Any]:
    """DOCX 의 구조 측정치를 모은다(단락·표·셀·이미지·텍스트)."""
    paragraphs = doc.paragraphs
    tables = doc.tables
    # 셀 수: 병합 안전 — 표별 row×cell 합(python-docx 의 row.cells 는 병합 셀을
    # 중복 반환할 수 있으나, a/b 동일 기준으로 세므로 일치율 측정에는 일관적이다).
    cell_count = 0
    for t in tables:
        for row in t.rows:
            cell_count += len(row.cells)

    text_parts = [p.text for p in paragraphs]
    for t in tables:
        for row in t.rows:
            text_parts.extend(c.text for c in row.cells)
    full_text = _norm_text("\n".join(text_parts))

    return {
        "paragraphs": len([p for p in paragraphs if p.text.strip()]),
        "tables": len(tables),
        "cells": cell_count,
        "images": _count_images(doc),
        "chars": len(full_text.replace(" ", "")),
        "tokens": len(full_text.split()),
        "_text": full_text,
    }


def _ratio(a: float, b: float) -> float:
    """min/max 일치율%(둘 다 0 이면 100 — '둘 다 없음'은 완전 일치)."""
    if a == 0 and b == 0:
        return 100.0
    hi = max(a, b)
    if hi == 0:
        return 100.0
    return round(min(a, b) / hi * 100.0, 2)


# --- 두 DOCX 구조 비교 --------------------------------------------------------

def compare_docx_structure(docx_a: str | Path, docx_b: str | Path) -> FidelityReport:
    """두 DOCX 의 **구조 일치도**를 비교한다(python-docx, 결정론).

    metric별 일치율% = min(a,b)/max(a,b)*100 (둘 다 0 이면 100).
    텍스트는 정규화 후 문자/토큰 일치율의 평균. overall 은 가중 평균
    (텍스트·표 비중↑). 동일 문서면 100.0.
    """
    report = FidelityReport(method="structural_compare")
    a = _measure(Document(str(docx_a)))
    b = _measure(Document(str(docx_b)))

    metrics = {
        "paragraphs": _ratio(a["paragraphs"], b["paragraphs"]),
        "tables": _ratio(a["tables"], b["tables"]),
        "cells": _ratio(a["cells"], b["cells"]),
        "images": _ratio(a["images"], b["images"]),
    }
    # 텍스트 일치율: 문자·토큰 두 척도의 평균.
    text_metric = round(
        (_ratio(a["chars"], b["chars"]) + _ratio(a["tokens"], b["tokens"])) / 2.0, 2)
    metrics["text"] = text_metric

    overall = round(sum(metrics[k] * w for k, w in _WEIGHTS.items()), 2)

    lost: list[str] = []
    _note_loss(lost, "단락", a["paragraphs"], b["paragraphs"])
    _note_loss(lost, "표", a["tables"], b["tables"])
    _note_loss(lost, "표 셀", a["cells"], b["cells"])
    _note_loss(lost, "이미지", a["images"], b["images"])
    if text_metric < 100.0:
        lost.append(
            f"텍스트 일치율 {text_metric}% — 문자 {a['chars']}→{b['chars']}, "
            f"토큰 {a['tokens']}→{b['tokens']}")

    report.metrics = metrics
    report.overall_score = overall
    report.counts = {
        "a": {k: a[k] for k in ("paragraphs", "tables", "cells", "images", "chars", "tokens")},
        "b": {k: b[k] for k in ("paragraphs", "tables", "cells", "images", "chars", "tokens")},
    }
    report.lost_items = lost
    report.ok = True
    report.notes.append(
        "구조(structural) 일치도만 측정 — 폰트·스타일·색상·레이아웃 등 "
        "시각 서식은 범위 밖. 구조 100% ≠ 시각 100%.")
    if overall < 100.0:
        report.notes.append(f"100% 미달 갭: {round(100.0 - overall, 2)}% (위 lost_items 참조).")
    return report


def _note_loss(lost: list[str], label: str, a: int, b: int) -> None:
    if a != b:
        lost.append(f"{label} 수 불일치: {a} → {b} ({'손실' if b < a else '증가'} {abs(a - b)})")


# --- roundtrip 측정 (DOCX → HWP → DOCX') --------------------------------------

def measure_roundtrip_fidelity(
    docx_path: str | Path, *, use_com: bool = True,
    work_dir: Optional[str | Path] = None,
) -> FidelityReport:
    """docx → (DOCX→HWP) → (HWP→DOCX) docx' 라운드트립 후 구조 일치도를 측정한다.

    **COM 미가용/실패 시** ``ok=False`` + 안내 notes 를 담고 측정을 생략한다
    (예외 전파 X). DOCX→HWP 는 한글 COM 대화형 전용이라 이 PC 인터랙티브
    세션에서만 실측 가능하다.

    중간 파일은 ``work_dir``(미지정 시 임시 디렉터리)에 생성하며 입력은 수정하지 않는다.
    """
    from .hwp_docx_convert import convert, hancom_com_available

    src = Path(docx_path)
    if not src.exists():
        raise FileNotFoundError(f"입력 파일이 없습니다: {src}")

    report = FidelityReport(method="roundtrip")
    if not (use_com and hancom_com_available()):
        report.notes.append(
            "DOCX→HWP COM 대화형 전용 — roundtrip 측정 불가, "
            "이 PC 인터랙티브에서만. (CI/백그라운드 세션에서는 한글 COM 이 안 뜸)")
        return report

    if work_dir is not None:
        return _run_roundtrip(src, Path(work_dir), report, use_com)
    with tempfile.TemporaryDirectory(prefix="fidelity_") as td:
        return _run_roundtrip(src, Path(td), report, use_com)


def _run_roundtrip(src: Path, work: Path, report: FidelityReport, use_com: bool) -> FidelityReport:
    work.mkdir(parents=True, exist_ok=True)
    hwp_mid = work / (src.stem + "_mid.hwp")
    docx_back = work / (src.stem + "_back.docx")

    r1 = convert(src, hwp_mid, use_com=use_com)
    if not r1.ok:
        report.notes.append("DOCX→HWP 변환 실패 — roundtrip 측정 불가.")
        report.notes.extend(r1.notes)
        return report
    r2 = convert(hwp_mid, docx_back, use_com=use_com)
    if not r2.ok:
        report.notes.append("HWP→DOCX 역변환 실패 — roundtrip 측정 불가.")
        report.notes.extend(r2.notes)
        return report

    cmp = compare_docx_structure(src, docx_back)
    cmp.method = "roundtrip"
    cmp.notes.insert(0, f"roundtrip 경로: DOCX→HWP({r1.method}) → HWP→DOCX({r2.method}).")
    return cmp
