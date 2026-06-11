"""image_apply.py — 인포그래픽 '제안'을 받아 그림 위치에 NotebookLM 슬라이드 프롬프트를 삽입한다.

``infographic_suggest.suggest_images_ai`` 가 만든 제안을 받아, 각 제안 위치(anchor) 뒤에
**NotebookLM 의 슬라이드 생성 기능에 그대로 붙여넣을 '프롬프트 텍스트 블록'**을 삽입한다.
사용자는 그 프롬프트를 NotebookLM 에 붙여넣어 슬라이드를 생성한 뒤, 안내 블록은 삭제하면 된다.

(이전 버전은 matplotlib(``chart_generator``)로 차트 PNG 를 직접 만들어 삽입했으나,
지금은 차트 직접 생성 대신 NotebookLM 슬라이드 프롬프트 삽입으로 동작한다.
``chart_generator``/``chart_insert`` 모듈 자체는 보존되며 여기서는 사용하지 않는다.)

안전 원칙
---------
- 원본 DOCX 는 절대 덮어쓰지 않는다(``out_docx == in_docx`` 면 ``ValueError``).
- 정부지원사업 문서 특성상 **없는 숫자는 만들지 않는다(날조 0)** — 슬라이드 프롬프트는
  "문서에 실제로 있는 수치만 쓰라" 는 규칙으로 생성된다(infographic_suggest 참고).
- 제안은 Claude(``openai_service`` 가용 시) 또는 키워드 규칙(폴백)으로 만든다.
"""

from __future__ import annotations

import re
import shutil
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

from .infographic_suggest import ImageSuggestion, suggest_images_ai
from .usage_acceptance import SELF_BLOCK_RE as _SELF_BLOCK_RE

_DIVIDER = "─" * 30


@dataclass
class ImageApplyItem:
    anchor_text: str
    visual_type: str
    action: str            # "notebooklm_prompt"
    anchor_found: bool
    note: str = ""

    def as_dict(self) -> dict[str, Any]:
        return {
            "anchor_text": self.anchor_text[:80],
            "visual_type": self.visual_type,
            "action": self.action,
            "anchor_found": self.anchor_found,
            "note": self.note,
        }


@dataclass
class ImageApplyReport:
    items: list[ImageApplyItem] = field(default_factory=list)
    prompts_inserted: int = 0
    anchors_missing: int = 0
    output_docx: str = ""

    def as_dict(self) -> dict[str, Any]:
        return {
            "prompts_inserted": self.prompts_inserted,
            "anchors_missing": self.anchors_missing,
            "output_docx": self.output_docx,
            "items": [i.as_dict() for i in self.items],
        }


def _iter_anchor_contexts(doc: Document):
    """(paragraph, top_table) 쌍을 순회한다.

    - 본문 직계 단락: top_table=None
    - 표 셀 안 단락(중첩 표 포함): 그 셀이 속한 **최상위 표**를 함께 돌려준다.
      (표 안 앵커는 셀 안이 아니라 표 '뒤'에 블록을 넣기 위함.)
    """
    for para in doc.paragraphs:
        yield para, None
    for table in doc.tables:
        yield from _iter_table_contexts(table, table)


def _iter_table_contexts(table, top_table):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                yield para, top_table
            for nested in cell.tables:
                yield from _iter_table_contexts(nested, top_table)


def _anchor_matches(anchor_text: str, para_text: str) -> bool:
    a = (anchor_text or "").strip()
    t = (para_text or "").strip()
    if not a or not t:
        return False
    key = a[:40]
    if key and key in t:
        return True
    # 표 헤더처럼 여러 셀을 합쳐 만든 앵커: 셀 텍스트가 앵커의 일부일 때도 매칭.
    if len(t) >= 4 and t in a:
        return True
    return False


def _find_anchor(doc: Document, anchor_text: str):
    """anchor_text 와 일치/부분일치하는 단락을 본문 + 표 셀에서 찾는다.

    Returns:
        (paragraph, top_table). 못 찾으면 (None, None).
        top_table 가 None 이 아니면 그 단락은 표 안에 있다(삽입은 표 뒤에 한다).
    """
    if not (anchor_text or "").strip():
        return None, None
    for para, table in _iter_anchor_contexts(doc):
        if _anchor_matches(anchor_text, para.text):
            return para, table
    return None, None


def _insert_paras_after(ref_elem, parent, callbacks: list) -> None:
    """ref_elem(<w:p> 또는 <w:tbl>) 바로 뒤에 새 단락들을 순서대로 끼워넣는다.

    본문/표 셀 어디에 있든 문서 순서상 정확히 ref 다음 위치에 삽입된다
    (python-docx 의 ``insert_paragraph_before`` 가 못 하는 '마지막 단락 뒤' 도 처리).
    """
    ref = ref_elem
    for cb in callbacks:
        new_p = OxmlElement("w:p")
        ref.addnext(new_p)
        cb(Paragraph(new_p, parent))
        ref = new_p


def _add_divider(paragraph) -> None:
    run = paragraph.add_run(_DIVIDER)
    run.font.size = Pt(8)
    try:
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    except Exception:  # pragma: no cover - 색상 실패는 치명적이지 않음
        pass


def _add_prompt_header(paragraph, sug: ImageSuggestion) -> None:
    run = paragraph.add_run(
        f"📊 [NotebookLM 슬라이드 생성용 프롬프트] · 유형: {sug.visual_type}  "
        f"(슬라이드 생성 후 이 블록은 삭제하세요)"
    )
    run.bold = True
    run.font.size = Pt(9)
    try:
        run.font.color.rgb = RGBColor(0x33, 0x66, 0xCC)
    except Exception:  # pragma: no cover
        pass


def _add_prompt_intro(paragraph) -> None:
    run = paragraph.add_run("↓ 아래 문장을 NotebookLM 슬라이드 생성에 붙여넣으세요")
    run.italic = True
    run.font.size = Pt(9)
    try:
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    except Exception:  # pragma: no cover
        pass


def _add_prompt_text(paragraph, sug: ImageSuggestion) -> None:
    text = sug.slide_prompt or sug.prompt or f"'{sug.visual_type}' 슬라이드를 만들어줘."
    run = paragraph.add_run(text)
    run.font.size = Pt(10)


def _apply_prompt_block(doc: Document, anchor_para, table, sug: ImageSuggestion) -> None:
    """anchor 뒤에 NotebookLM 프롬프트 블록을 삽입한다.

    - 본문 앵커: 그 단락 바로 뒤.
    - 표 안 앵커: 셀 안이 아니라 **표 전체 뒤**(큰 블록이 셀을 부풀리지 않도록).
    - 앵커 없음: 문서 끝(최후 폴백).
    """
    callbacks = [
        _add_divider,
        lambda p: _add_prompt_header(p, sug),
        _add_prompt_intro,
        lambda p: _add_prompt_text(p, sug),
        _add_divider,
    ]
    if anchor_para is None:
        for cb in callbacks:
            cb(doc.add_paragraph())
    elif table is not None:
        _insert_paras_after(table._tbl, doc, callbacks)
    else:
        _insert_paras_after(anchor_para._p, anchor_para._parent, callbacks)


def apply_images(
    in_docx: str,
    out_docx: str,
    *,
    max_items: int = 8,
    placeholder_only: bool = False,
    openai_service: Optional[Any] = None,
) -> ImageApplyReport:
    """인포그래픽 제안 위치에 NotebookLM 슬라이드 생성 프롬프트를 삽입한다.

    Args:
        in_docx: 원본 DOCX(읽기 전용, 덮어쓰지 않음).
        out_docx: 결과 DOCX 경로. **in_docx 와 같으면 ValueError**.
        max_items: 적용할 최대 제안 수(suggest_images_ai 의 max_suggestions).
        placeholder_only: (하위호환용, 동작에는 영향 없음 — 항상 프롬프트 블록을 삽입).
        openai_service: Claude 제안에 사용할 서비스. None/미가용이면 키워드 규칙으로 폴백.

    Returns:
        ImageApplyReport — 삽입 결과 집계.
    """
    in_path = Path(in_docx)
    out_path = Path(out_docx)
    if in_path.resolve() == out_path.resolve():
        raise ValueError("in_docx 와 out_docx 가 같습니다. 원본 덮어쓰기는 금지입니다.")
    if not in_path.exists():
        raise FileNotFoundError(f"입력 DOCX 가 없습니다: {in_docx}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(str(in_path), str(out_path))

    doc = Document(str(out_path))
    report = ImageApplyReport(output_docx=str(out_path))

    suggestions = suggest_images_ai(
        doc, openai_service=openai_service, max_suggestions=max_items
    ).suggestions

    for sug in suggestions:
        anchor_para, table = _find_anchor(doc, sug.anchor_text)
        anchor_found = anchor_para is not None

        _apply_prompt_block(doc, anchor_para, table, sug)
        report.prompts_inserted += 1
        note = "NotebookLM 슬라이드 프롬프트 삽입"
        if anchor_found and table is not None:
            note += " / 표 뒤 위치"
        if not anchor_found:
            report.anchors_missing += 1
            note += " / anchor 미발견(문서 끝에 추가)"

        report.items.append(ImageApplyItem(
            anchor_text=sug.anchor_text,
            visual_type=sug.visual_type,
            action="notebooklm_prompt",
            anchor_found=anchor_found,
            note=note,
        ))

    doc.save(str(out_path))
    return report


def apply_images_docx(
    path: str,
    out_docx: str,
    *,
    max_items: int = 8,
    placeholder_only: bool = False,
    openai_service: Optional[Any] = None,
) -> ImageApplyReport:
    return apply_images(
        path, out_docx,
        max_items=max_items,
        placeholder_only=placeholder_only,
        openai_service=openai_service,
    )


# --- NotebookLM 블록 제거 (삽입의 역연산) -------------------------------------

_PROMPT_INTRO_RE = re.compile(r"슬라이드\s*생성에\s*붙여넣으세요")


@dataclass
class StripReport:
    markers_removed: int = 0      # 헤더·안내 등 검출 마커 단락 수(usage_acceptance 결함 수와 대응)
    paragraphs_removed: int = 0   # 구분선·프롬프트 본문 포함 총 삭제 단락 수
    output_docx: str = ""

    def as_dict(self) -> dict[str, Any]:
        return {
            "markers_removed": self.markers_removed,
            "paragraphs_removed": self.paragraphs_removed,
            "output_docx": self.output_docx,
        }


def _p_text(elem) -> str:
    return "".join(elem.itertext())


def _is_p(elem) -> bool:
    return elem is not None and elem.tag == qn("w:p")


def _is_divider_p(elem) -> bool:
    t = _p_text(elem).strip()
    return len(t) >= 10 and set(t) <= {"─"}


def strip_notebooklm_blocks(in_docx: str, out_docx: str) -> StripReport:
    """잔존한 NotebookLM 작업용 블록을 제거한 '제출용 사본'을 만든다.

    ``apply_images`` 가 삽입한 블록(구분선·헤더·안내·프롬프트·구분선 5단락)을
    ``usage_acceptance`` 의 self_inserted_blocks 검출과 **같은 패턴**으로 찾아 지운다.
    실본문 오삭제 방지를 위해 프롬프트 본문 단락은 '안내문 바로 다음 + 구분선 직전'
    구조가 확인될 때만 함께 지운다(구조가 깨졌으면 마커·구분선만 지운다).

    원본은 절대 덮어쓰지 않는다(``out_docx == in_docx`` 면 ``ValueError``).
    """
    in_path = Path(in_docx)
    out_path = Path(out_docx)
    if in_path.resolve() == out_path.resolve():
        raise ValueError("in_docx 와 out_docx 가 같습니다. 원본 덮어쓰기는 금지입니다.")
    if not in_path.exists():
        raise FileNotFoundError(f"입력 DOCX 가 없습니다: {in_docx}")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(str(in_path), str(out_path))

    doc = Document(str(out_path))
    report = StripReport(output_docx=str(out_path))

    contexts = list(_iter_anchor_contexts(doc))   # 단락 proxy 보관(lxml id 재사용 방지)
    marked: list = []
    seen: set[int] = set()
    seen_markers: set[int] = set()

    def _mark(elem) -> None:
        if elem is not None and id(elem) not in seen:
            seen.add(id(elem))
            marked.append(elem)

    for para, _tbl in contexts:
        p = para._p
        text = para.text or ""
        if not _SELF_BLOCK_RE.search(text):
            continue
        if id(p) in seen_markers:                 # 병합 셀 중복 방문 방지
            continue
        seen_markers.add(id(p))
        _mark(p)
        prev = p.getprevious()
        if _is_p(prev) and _is_divider_p(prev):
            _mark(prev)
        nxt = p.getnext()
        if _is_p(nxt) and _is_divider_p(nxt):
            _mark(nxt)
        elif (_PROMPT_INTRO_RE.search(text) and _is_p(nxt)
              and not _SELF_BLOCK_RE.search(_p_text(nxt))):
            # 안내문 다음 단락 = 프롬프트 본문. 바로 뒤가 구분선일 때만 본문으로 확정.
            nxt2 = nxt.getnext()
            if _is_p(nxt2) and _is_divider_p(nxt2):
                _mark(nxt)
                _mark(nxt2)

    for elem in marked:
        parent = elem.getparent()
        if parent is None:
            continue
        parent.remove(elem)
        # 표 셀은 단락이 최소 1개 있어야 유효한 DOCX — 다 비면 빈 단락을 채운다.
        if parent.tag == qn("w:tc") and parent.find(qn("w:p")) is None:
            parent.append(OxmlElement("w:p"))

    report.markers_removed = len(seen_markers)
    report.paragraphs_removed = len(marked)
    doc.save(str(out_path))
    return report
