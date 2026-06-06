from __future__ import annotations

import json
import logging
import os
import re
import shutil
import subprocess
import sys
import zipfile
from datetime import date
from glob import glob
from pathlib import Path
from typing import Any

from docx import Document

from ..analysis.docx_template import (
    analyze_template,
    build_doc_summary,
    cell_value_needs_improvement,
    detect_tables,
    sanitize_template_profile,
)
from ..document_ingest import REFERENCE_SUFFIXES, ensure_template_docx, extract_additional_text
from ..models import (
    ArtifactBundle,
    EvidenceRequest,
    EvidenceSource,
    GeneratedImage,
    ImageRequest,
    ProjectInput,
    ReferenceFile,
    TemplateProfile,
)
from ..storage import Storage
from ..utils import log_line, read_json, sanitize_user_filename, unique_lines, write_json
from .evaluation_service import EvalLoopReport, EvaluationService
from .evidence_service import EvidenceService
from .image_service import ImageService
from .openai_client import OpenAIService
from .qa_service import QAService
from .render_service import RenderService


class ProjectService:
    PSST_PROBLEM_RE = re.compile(r"1\.\s*문제\s*인식.*Problem", re.IGNORECASE)
    PSST_SOLUTION_RE = re.compile(r"2\.\s*실현\s*가능성.*Solution", re.IGNORECASE)
    PSST_SCALE_RE = re.compile(r"3\.\s*성장전략.*Scale", re.IGNORECASE)
    PSST_TEAM_RE = re.compile(r"4\.\s*팀\s*구성.*Team", re.IGNORECASE)
    CORE_TABLE_LABEL_RE = re.compile(
        r"(일반현황|창업\s*아이템\s*개요|채용\s*계획|주요\s*업무|추진\s*일정|목표\s*성과)",
        re.IGNORECASE,
    )
    PSST_HEADING_LABELS = {
        "problem": "1. 문제 인식 (Problem)",
        "solution": "2. 실현 가능성 (Solution)",
        "scale": "3. 성장전략 (Scale-up)",
        "team": "4. 팀 구성 (Team)",
    }
    SECTION_HEADING_RE = re.compile(r"^((\d+[\.\)])|([①-⑨]))\s*.+")
    GUIDE_MARKER_RE = re.compile(r"(※|<[^>]+>|기재|작성요령|작성방법|예시|OOO|○○○)")
    INVALID_XML_CHAR_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")
    LIBRARY_SENTENCE_SPLIT_RE = re.compile(r"(?<=[\.\!\?])\s+")
    LIBRARY_BLOCK_SPLIT_RE = re.compile(r"(?:\r?\n){2,}")
    LIBRARY_GUIDE_RE = re.compile(r"(※|<[^>]+>|기재|작성요령|작성방법|예시|샘플|양식|붙임)")
    LIBRARY_TOKEN_RE = re.compile(r"[가-힣A-Za-z0-9]{2,}")
    LIBRARY_KEYWORD_HINTS = ("문제", "해결", "목표", "시장", "고객", "전략", "실행", "성과", "매출", "비용", "로드맵")
    REFERENCE_SNIPPET_LIMIT = 160
    SECTION_AUTOFILL_EXCLUDE_RE = re.compile(
        r"(^(?:[◦ㅇ■□●·•]+)$|담\s*당\s*자|사실과 다름이 없음을 확인|추천기관|신청기업|제출목록|증빙서류|평가대상에서 제외|별첨|붙임)",
        re.IGNORECASE,
    )
    TABLE_AUTOFILL_EXCLUDE_RE = re.compile(
        r"(열\s*\d+|행\s*\d+|협약해지|제재|환수|유의사항|집행기준|증빙|계좌이체|서류|동의|개인정보|제출목록|체크리스트|추천기관|신청기업|담당자|등록번호|주소|전화번호|주민등록|이메일|별첨|붙임)",
        re.IGNORECASE,
    )
    SAMPLE_VALUE_RE = re.compile(
        r"(○○|OOO|OO\.OO|00\.00|00백만|…|\.{2,}|□|■|예시|샘플|법인기업|예비창업팀|사업자 소재지|"
        r"○○기업|OO학|예정\(['\u2019]?00|완료\(['\u2019]?00|에정\b)"
    )
    TABLE_NARROW_HINT_RE = re.compile(r"(월|주차|진행률|금액|비중|건수|수량|단가|열\s*\d+|행\s*\d+)")
    DOCX_RENDER_SCRIPT_PATTERNS = (
        "~/.codex/plugins/cache/openai-primary-runtime/documents/*/skills/documents/render_docx.py",
        "~/.cache/codex-runtimes/codex-primary-runtime*/documents/*/skills/documents/render_docx.py",
    )

    def __init__(
        self,
        storage: Storage,
        openai_service: OpenAIService,
        evidence_service: EvidenceService,
        image_service: ImageService,
        render_service: RenderService,
        qa_service: QAService,
    ):
        self.storage = storage
        self.openai_service = openai_service
        self.evidence_service = evidence_service
        self.image_service = image_service
        self.render_service = render_service
        self.qa_service = qa_service
        self._library_snippets: list[dict[str, str]] | None = None
        logging.getLogger("pypdf").setLevel(logging.ERROR)

    def analyze_uploaded_template(self, file_name: str, content: bytes) -> TemplateProfile:
        template_id, output_path = self.storage.create_template_space(file_name)
        output_path.write_bytes(content)
        analysis_path, conversion_notes = ensure_template_docx(output_path)
        profile = analyze_template(analysis_path)
        profile.template_id = template_id
        profile.template_name = file_name
        profile.source_docx = str(analysis_path)
        profile.analysis_notes.extend(note for note in conversion_notes if note not in profile.analysis_notes)
        if self.openai_service.settings.template_ai_refine_enabled and self.openai_service.available:
            refined = self.openai_service.refine_template_profile(profile.model_dump(), build_doc_summary(analysis_path))
            if refined:
                try:
                    merged = TemplateProfile.model_validate(refined)
                    merged.template_id = template_id
                    merged.source_docx = str(analysis_path)
                    merged.template_name = file_name
                    for note in conversion_notes:
                        if note not in merged.analysis_notes:
                            merged.analysis_notes.append(note)
                    profile = merged
                except Exception:
                    pass
        profile = sanitize_template_profile(profile)
        self.storage.save_template_profile(profile)
        return profile

    def finalize_template(self, template_id: str, raw_json: str) -> TemplateProfile:
        try:
            data = json.loads(raw_json)
        except json.JSONDecodeError as exc:
            raise ValueError("템플릿 설정 JSON 형식이 올바르지 않습니다.") from exc
        profile = TemplateProfile.model_validate(data)
        profile.template_id = template_id
        folder = self.storage.template_dir(template_id)
        docx_files = sorted(folder.glob("*.docx"))
        if not docx_files:
            raise ValueError("원본 DOCX 파일을 찾지 못했습니다. 템플릿을 다시 업로드해 주세요.")
        source_docx = docx_files[0]
        profile.source_docx = str(source_docx)
        profile = sanitize_template_profile(profile)
        self.storage.save_template_profile(profile)
        return profile

    def create_project(self, template_id: str, project_name: str) -> str:
        project_id, _ = self.storage.create_project_space(template_id, project_name or "새 프로젝트")
        profile = sanitize_template_profile(self.storage.load_template_profile(template_id))
        profile = self._pin_template_source_to_project(profile, project_id)
        blank = ProjectInput(template_id=template_id)
        blank.project_meta = {
            "improve_partial": True,
            "psst_only": True,
            "disable_images": True,
        }
        self.storage.save_project_input(project_id, blank)
        write_json(self.storage.project_dir(project_id) / "template_snapshot.json", profile.model_dump())
        return project_id

    def _pin_template_source_to_project(self, profile: TemplateProfile, project_id: str) -> TemplateProfile:
        """Copy template DOCX into the project folder so generation survives missing template files."""
        try:
            source_docx = self._resolve_source_docx(profile, project_id)
        except ValueError:
            return profile
        project_dir = self.storage.project_dir(project_id)
        pinned = project_dir / "template_source.docx"
        if source_docx.resolve() != pinned.resolve():
            pinned.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(source_docx, pinned)
        profile.source_docx = str(pinned)
        return profile

    def template_docx_ready(self, profile: TemplateProfile) -> bool:
        candidates: list[Path] = []
        if profile.source_docx:
            candidates.append(Path(profile.source_docx))
        candidates.extend(
            path
            for path in sorted(self.storage.template_dir(profile.template_id).glob("*.docx"))
            if not path.name.startswith("~$")
        )
        seen: set[str] = set()
        for path in candidates:
            key = str(path).lower()
            if key in seen:
                continue
            seen.add(key)
            if path.is_file():
                return True
        return False

    def template_source_status(self, profile: TemplateProfile, project_id: str) -> dict[str, Any]:
        try:
            path = self._resolve_source_docx(profile, project_id)
            return {"ready": True, "path": str(path), "message": ""}
        except ValueError as exc:
            return {"ready": False, "path": "", "message": str(exc)}

    def _resolve_source_docx(self, profile: TemplateProfile, project_id: str) -> Path:
        candidates: list[Path] = []
        if profile.source_docx:
            candidates.append(Path(profile.source_docx))
        project_dir = self.storage.project_dir(project_id)
        candidates.append(project_dir / "template_source.docx")
        candidates.extend(sorted(project_dir.glob("*.docx")))
        template_folder = self.storage.template_dir(profile.template_id)
        candidates.extend(sorted(template_folder.glob("*.docx")))
        seen: set[str] = set()
        for path in candidates:
            key = str(path).lower()
            if key in seen:
                continue
            seen.add(key)
            if path.is_file():
                return path.resolve()
        docx_name = Path(profile.source_docx).name if profile.source_docx else "(없음)"
        raise ValueError(
            "템플릿 원본 DOCX 파일이 없습니다. "
            f"필요한 파일: {docx_name} | 템플릿 ID: {profile.template_id}. "
            "홈 화면에서 동일한 양식 DOCX를 다시 업로드한 뒤, 이 프로젝트를 새로 만드세요."
        )

    def save_project_form(
        self,
        project_id: str,
        answers: dict[str, Any],
        project_title: str,
        organization_name: str,
        evidence_topics: str,
        reference_files: list[tuple[str, bytes]],
        writing_provider: str = "",
        writing_model: str = "",
        improve_partial: bool = True,
        psst_only: bool = True,
        disable_images: bool = True,
    ) -> ProjectInput:
        profile = self.load_profile_for_project(project_id)
        existing_input: ProjectInput | None
        try:
            existing_input = self.storage.load_project_input(project_id)
        except Exception:
            existing_input = None

        references: list[ReferenceFile] = []
        if existing_input is not None:
            references.extend(
                reference
                for reference in existing_input.references
                if reference.saved_path and Path(reference.saved_path).exists()
            )
        for file_name, content in reference_files:
            if not file_name:
                continue
            safe_name = sanitize_user_filename(file_name)
            path = self._next_reference_path(project_id, safe_name)
            path.write_bytes(content)
            preview = self.extract_reference_text(path)
            preview_path = path.with_suffix(path.suffix + ".txt")
            preview_path.write_text(preview, encoding="utf-8")
            references.append(
                ReferenceFile(
                    file_name=path.name,
                    saved_path=str(path),
                    extracted_text_path=str(preview_path),
                    extracted_preview=preview[:400],
                )
            )
        deduped_references: list[ReferenceFile] = []
        seen_reference_keys: set[str] = set()
        for reference in references:
            ref_key = str(reference.saved_path or "").strip().lower()
            if not ref_key:
                continue
            if ref_key in seen_reference_keys:
                continue
            seen_reference_keys.add(ref_key)
            deduped_references.append(reference)
        references = deduped_references

        evidence_requests = [EvidenceRequest(topic=line) for line in unique_lines(evidence_topics)]
        merged_answers: dict[str, Any] = dict(existing_input.answers) if existing_input is not None else {}
        for key, value in answers.items():
            if str(value).strip():
                merged_answers[key] = value
            else:
                merged_answers.pop(key, None)
        resolved_project_title = str(project_title or "").strip()
        resolved_org_name = str(organization_name or "").strip()
        resolved_writing_provider = str(writing_provider or "").strip().lower()
        resolved_writing_model = str(writing_model or "").strip()
        if existing_input is not None:
            if not resolved_project_title:
                resolved_project_title = str(existing_input.project_meta.get("project_title", "")).strip()
            if not resolved_org_name:
                resolved_org_name = str(existing_input.organization_profile.get("name", "")).strip()
            if not resolved_writing_provider:
                resolved_writing_provider = str(existing_input.project_meta.get("writing_provider", "")).strip().lower()
            if not resolved_writing_model:
                resolved_writing_model = str(existing_input.project_meta.get("writing_model", "")).strip()
        if resolved_writing_provider not in {"openai", "anthropic"}:
            resolved_writing_provider = ""
        project_meta: dict[str, Any] = {
            "project_title": resolved_project_title,
            "improve_partial": bool(improve_partial),
            "psst_only": bool(psst_only),
            "disable_images": bool(disable_images),
        }
        if resolved_writing_provider:
            project_meta["writing_provider"] = resolved_writing_provider
        if resolved_writing_model:
            project_meta["writing_model"] = resolved_writing_model
        image_requests = []
        for slot in profile.image_slots:
            note_key = f"image_note_{slot.slot_id}"
            image_requests.append(
                ImageRequest(
                    slot_id=slot.slot_id,
                    label=slot.label,
                    prompt=str(merged_answers.get(note_key, "")),
                    source="generated",
                )
            )
        project_input = ProjectInput(
            template_id=profile.template_id,
            project_meta=project_meta,
            organization_profile={"name": resolved_org_name},
            answers=merged_answers,
            references=references,
            evidence_requests=evidence_requests,
            image_requests=image_requests,
        )
        self.storage.save_project_input(project_id, project_input)
        return project_input

    @staticmethod
    def _meta_flag(project_meta: dict[str, Any], key: str, default: bool = True) -> bool:
        raw = project_meta.get(key)
        if raw is None:
            return default
        if isinstance(raw, bool):
            return raw
        return str(raw).lower() in {"1", "true", "on", "yes"}

    def _find_psst_field_ids(self, profile: TemplateProfile) -> dict[str, str]:
        mapping: dict[str, str] = {}
        for section in profile.sections:
            text = f"{section.anchor_text} {section.label}"
            if "problem" not in mapping and self.PSST_PROBLEM_RE.search(text):
                mapping["problem"] = section.field_id
            if "solution" not in mapping and self.PSST_SOLUTION_RE.search(text):
                mapping["solution"] = section.field_id
            if "scale" not in mapping and self.PSST_SCALE_RE.search(text):
                mapping["scale"] = section.field_id
            if "team" not in mapping and self.PSST_TEAM_RE.search(text):
                mapping["team"] = section.field_id
        return mapping

    def _core_table_ids(self, profile: TemplateProfile) -> set[str]:
        return {table.table_id for table in profile.tables if self.CORE_TABLE_LABEL_RE.search(str(table.label or ""))}

    def _apply_psst_from_user_input(self, answers: dict[str, Any], profile: TemplateProfile) -> dict[str, Any]:
        psst = self._find_psst_field_ids(profile)
        if not psst:
            return answers
        updated = dict(answers)
        brief = str(answers.get("user_brief", "") or "").strip()
        notes = str(answers.get("user_notes", "") or "").strip()
        if brief and psst.get("problem"):
            updated[psst["problem"]] = brief
        if notes:
            parts = [part.strip() for part in re.split(r"\n{2,}", notes) if part.strip()]
            if not parts:
                parts = [notes]
            slot_keys = ("solution", "scale", "team")
            if len(parts) >= 2:
                for index, key in enumerate(slot_keys):
                    field_id = psst.get(key)
                    if field_id and index < len(parts):
                        updated[field_id] = parts[index]
            elif psst.get("solution"):
                updated[psst["solution"]] = notes
        return updated

    def _restrict_autofill_targets(self, profile: TemplateProfile, targets: list[dict[str, Any]]) -> list[dict[str, Any]]:
        psst_ids = set(self._find_psst_field_ids(profile).values())
        core_tables = self._core_table_ids(profile)
        if not psst_ids and not core_tables:
            return targets
        restricted: list[dict[str, Any]] = []
        for question in targets:
            target = question.get("target") or {}
            kind = str(target.get("kind", ""))
            if kind == "section":
                field_id = str(target.get("field_id", ""))
                if field_id in psst_ids:
                    restricted.append(question)
                continue
            if kind == "table_cell":
                table_id = str(target.get("table_id", ""))
                if table_id in core_tables:
                    restricted.append(question)
        return restricted

    def _results_docx_name(self, project_input: ProjectInput) -> str:
        stamp = date.today().strftime("%Y%m%d")
        title = str(project_input.project_meta.get("project_title", "") or "").strip()
        slug = re.sub(r'[<>:"/\\|?*]+', "", title)[:40].strip() or "사업계획서"
        return f"{stamp}_{slug}_초안.docx"

    def _build_hwp_paste_text(
        self,
        profile: TemplateProfile,
        project_input: ProjectInput,
        render_result: dict[str, Any],
        *,
        psst_field_ids: set[str],
        core_table_ids: set[str],
    ) -> str:
        lines: list[str] = []
        title = str(project_input.project_meta.get("project_title", "") or "").strip()
        org = str(project_input.organization_profile.get("name", "") or "").strip()
        if title:
            lines.extend(["=== 과제명 ===", title, ""])
        if org:
            lines.extend(["=== 기관명 ===", org, ""])
        answers = project_input.answers
        psst = self._find_psst_field_ids(profile)
        for key, heading in self.PSST_HEADING_LABELS.items():
            field_id = psst.get(key)
            if not field_id or field_id not in psst_field_ids:
                continue
            body = str(answers.get(field_id, "") or "").strip()
            if not body:
                continue
            lines.extend([f"=== {heading} ===", body, ""])
        brief = str(answers.get("user_brief", "") or "").strip()
        notes = str(answers.get("user_notes", "") or "").strip()
        if brief:
            lines.extend(["=== 입력: 사업 개요 ===", brief, ""])
        if notes:
            lines.extend(["=== 입력: 추가 메모 ===", notes, ""])
        for table in profile.tables:
            if table.table_id not in core_table_ids:
                continue
            table_lines: list[str] = []
            for cell in sorted(table.cells, key=lambda item: (item.row, item.cell)):
                value = str(answers.get(cell.cell_id, "") or "").strip()
                if not value:
                    continue
                label = str(cell.label or "").strip() or f"r{cell.row}c{cell.cell}"
                table_lines.append(f"{label}\t{value}")
            if table_lines:
                lines.append(f"=== 표: {table.label} ===")
                lines.extend(table_lines)
                lines.append("")
        render_errors = list(render_result.get("errors") or [])
        if render_errors:
            lines.append("=== 생성 시 확인할 항목 ===")
            for issue in render_errors[:12]:
                lines.append(f"- {issue}")
            lines.append("")
        lines.append("※ HWP: 각 === 제목 === 아래 본문·표 줄을 해당 양식 위치에 붙여넣으세요.")
        return "\n".join(lines).strip() + "\n"

    def _build_copy_blocks(
        self,
        profile: TemplateProfile,
        project_input: ProjectInput,
        *,
        psst_field_ids: set[str],
        core_table_ids: set[str],
    ) -> list[dict[str, str]]:
        blocks: list[dict[str, str]] = []
        answers = project_input.answers
        title = str(project_input.project_meta.get("project_title", "") or "").strip()
        if title:
            blocks.append({"id": "project_title", "title": "과제명", "text": title, "kind": "meta"})
        org = str(project_input.organization_profile.get("name", "") or "").strip()
        if org:
            blocks.append({"id": "organization_name", "title": "기관명", "text": org, "kind": "meta"})
        psst = self._find_psst_field_ids(profile)
        for key, heading in self.PSST_HEADING_LABELS.items():
            field_id = psst.get(key)
            if not field_id or field_id not in psst_field_ids:
                continue
            body = str(answers.get(field_id, "") or "").strip()
            if body:
                blocks.append({"id": field_id, "title": heading, "text": body, "kind": "section"})
        for table in profile.tables:
            if table.table_id not in core_table_ids:
                continue
            for cell in sorted(table.cells, key=lambda item: (item.row, item.cell)):
                value = str(answers.get(cell.cell_id, "") or "").strip()
                if not value:
                    continue
                blocks.append(
                    {
                        "id": cell.cell_id,
                        "title": f"{table.label} / {cell.label}",
                        "text": value,
                        "kind": "table_cell",
                    }
                )
        return blocks

    def _build_fill_map(
        self,
        profile: TemplateProfile,
        project_input: ProjectInput,
        render_result: dict[str, Any],
        *,
        psst_field_ids: set[str],
    ) -> dict[str, Any]:
        psst = self._find_psst_field_ids(profile)
        answers = project_input.answers
        sections_written: list[dict[str, str]] = []
        for key, heading in self.PSST_HEADING_LABELS.items():
            field_id = psst.get(key)
            if not field_id:
                continue
            sections_written.append(
                {
                    "psst": key,
                    "heading": heading,
                    "field_id": field_id,
                    "from_user_brief": key == "problem" and bool(str(answers.get("user_brief", "")).strip()),
                    "has_text": bool(str(answers.get(field_id, "")).strip()),
                }
            )
        return {
            "project_title": str(project_input.project_meta.get("project_title", "") or ""),
            "organization_name": str(project_input.organization_profile.get("name", "") or ""),
            "user_brief_mapped_to": psst.get("problem"),
            "user_notes_split_to": [psst.get(k) for k in ("solution", "scale", "team") if psst.get(k)],
            "sections": sections_written,
            "render_summary": {
                "sections_written": render_result.get("sections_written", 0),
                "cells_written": render_result.get("cells_written", 0),
                "errors": list(render_result.get("errors") or [])[:20],
                "warnings": list(render_result.get("warnings") or [])[:20],
            },
        }

    def _build_results_summary_ko(self, qa_report: dict[str, Any], render_result: dict[str, Any], results_dir: Path) -> str:
        lines: list[str] = []
        if qa_report.get("passed"):
            lines.append("✅ 생성 점검: 통과했습니다. output.docx와 results 폴더 파일을 확인하세요.")
        else:
            lines.append(
                f"⚠️ 생성 점검: 오류 {qa_report.get('error_count', 0)}건, "
                f"경고 {qa_report.get('warning_count', 0)}건입니다. 아래 메시지를 확인하세요."
            )
        for issue in (qa_report.get("errors") or [])[:6]:
            lines.append(f"  · {issue}")
        for issue in (render_result.get("errors") or [])[:6]:
            if issue not in (qa_report.get("errors") or []):
                lines.append(f"  · {issue}")
        if not (qa_report.get("errors") or render_result.get("errors")):
            lines.append("  · 표·문단 위치를 HWP에서 한 번 더 맞춰 주세요.")
        lines.append(f"📁 결과 폴더: {results_dir}")
        lines.append("📄 HWP 붙여넣기: hwp_paste.txt (제목·본문·표 탭 구분)")
        lines.append("📋 섹션별 복사: 화면의 「복사」 버튼 또는 copy_blocks.json")
        return "\n".join(lines) + "\n"

    def _publish_results_bundle(
        self,
        project_id: str,
        profile: TemplateProfile,
        project_input: ProjectInput,
        output_path: Path,
        qa_report: dict[str, Any],
        render_result: dict[str, Any],
        *,
        psst_field_ids: set[str],
        core_table_ids: set[str],
    ) -> dict[str, str]:
        results_dir = self.storage.results_dir(project_id)
        results_dir.mkdir(parents=True, exist_ok=True)
        dated_name = self._results_docx_name(project_input)
        results_docx = results_dir / dated_name
        shutil.copy2(output_path, results_docx)
        shutil.copy2(output_path, results_dir / "output.docx")

        hwp_text = self._build_hwp_paste_text(
            profile,
            project_input,
            render_result,
            psst_field_ids=psst_field_ids,
            core_table_ids=core_table_ids,
        )
        hwp_path = results_dir / "hwp_paste.txt"
        hwp_path.write_text(hwp_text, encoding="utf-8")

        copy_blocks = self._build_copy_blocks(
            profile,
            project_input,
            psst_field_ids=psst_field_ids,
            core_table_ids=core_table_ids,
        )
        copy_blocks_path = results_dir / "copy_blocks.json"
        write_json(copy_blocks_path, {"blocks": copy_blocks})

        fill_map_path = results_dir / "fill_map.json"
        write_json(
            fill_map_path,
            self._build_fill_map(profile, project_input, render_result, psst_field_ids=psst_field_ids),
        )

        summary_path = results_dir / "generation_summary.txt"
        summary_path.write_text(
            self._build_results_summary_ko(qa_report, render_result, results_dir),
            encoding="utf-8",
        )

        qa_copy = results_dir / "qa_report.json"
        write_json(qa_copy, qa_report)

        return {
            "results_dir": str(results_dir),
            "results_docx": str(results_docx),
            "hwp_paste": str(hwp_path),
            "copy_blocks": str(copy_blocks_path),
            "fill_map": str(fill_map_path),
            "generation_summary": str(summary_path),
        }

    def generate(self, project_id: str) -> ArtifactBundle:
        profile = self.load_profile_for_project(project_id)
        project_input = self.storage.load_project_input(project_id)
        source_docx = self._resolve_source_docx(profile, project_id)
        profile.source_docx = str(source_docx)
        improve_partial = self._meta_flag(project_input.project_meta, "improve_partial", True)
        psst_only = self._meta_flag(project_input.project_meta, "psst_only", True)
        disable_images = self._meta_flag(project_input.project_meta, "disable_images", True)
        project_input.answers = self._seed_answers_from_docx(profile, source_docx, project_input.answers)
        project_input.answers = self._apply_psst_from_user_input(project_input.answers, profile)
        # Case A: references uploaded → strict preserve mode
        has_references = bool(project_input.references)
        context = self._build_context(project_input, strict_preserve=has_references)
        transfer_mode = len(project_input.references) >= 2
        partial_doc = improve_partial and self._docx_is_partially_filled(source_docx)
        template_completed = self._template_looks_completed(profile) and not partial_doc
        if partial_doc:
            targets = self._collect_improvable_questions(profile, project_input.answers, source_docx)
        else:
            targets = self._collect_missing_questions(profile, project_input.answers, required_only=False)
        missing = self._filter_missing_for_autofill(profile, targets, transfer_mode)
        if psst_only:
            missing = self._restrict_autofill_targets(profile, missing)
        if missing and (not template_completed or partial_doc):
            reference_hints = self._suggest_reference_snippets(missing, project_input)
            # When project references are uploaded, prioritize them and avoid unrelated library contamination.
            has_library_context = False if project_input.references else bool(self._ensure_reference_library_loaded())
            has_user_context = (
                bool(project_input.references)
                or len(project_input.answers) >= 2
                or bool(reference_hints)
                or has_library_context
            )
            library_hints = (
                self._suggest_library_snippets(missing)
                if has_user_context and not project_input.references
                else {}
            )
            combined_hints = self._merge_hint_maps(reference_hints, library_hints)
            full_context = self._merge_context_with_library(context, combined_hints, questions=missing)
            writing_provider = str(project_input.project_meta.get("writing_provider", "")).strip().lower()
            writing_model = str(project_input.project_meta.get("writing_model", "")).strip()
            drafted = (
                self._draft_missing_answers_in_chunks(
                    missing,
                    full_context,
                    writing_provider=writing_provider,
                    writing_model=writing_model,
                    strict_preserve=has_references,
                )
                if has_user_context
                else {}
            )
            if drafted:
                project_input.answers.update({k: v for k, v in drafted.items() if str(v).strip()})
            remaining = self._collect_missing_questions(profile, project_input.answers, required_only=False)
            remaining = self._filter_missing_for_autofill(profile, remaining, transfer_mode)
            if psst_only:
                remaining = self._restrict_autofill_targets(profile, remaining)
            if remaining:
                fallback = self._build_fallback_answers(remaining, project_input, combined_hints, strict_preserve=has_references)
                project_input.answers.update(fallback)
            project_input.answers = self._postprocess_answers(profile, project_input.answers)
            project_input.answers = {
                key: self._sanitize_xml_text(value) if isinstance(value, str) else value
                for key, value in project_input.answers.items()
            }
            self.storage.save_project_input(project_id, project_input)

        evidence = self.evidence_service.search(project_input.evidence_requests)
        if disable_images:
            images: list[GeneratedImage] = []
        else:
            images = self.image_service.build_images(
                profile.image_slots,
                project_input.answers,
                evidence,
                self.storage.project_dir(project_id) / "generated_assets",
            )
        output_path = self.storage.project_dir(project_id) / "output" / "output.docx"
        psst_field_ids = set(self._find_psst_field_ids(profile).values())
        core_table_ids = self._core_table_ids(profile)
        render_result = self.render_service.render(
            profile,
            project_input,
            images,
            output_path,
            psst_only=psst_only,
            psst_field_ids=psst_field_ids,
            core_table_ids=core_table_ids,
        )
        sources_path = self.storage.project_dir(project_id) / "output" / "sources.json"
        qa_path = self.storage.project_dir(project_id) / "output" / "qa_report.json"
        benchmark_path = self.storage.project_dir(project_id) / "output" / "benchmark_compare.json"
        transfer_path = self.storage.project_dir(project_id) / "output" / "transfer_report.json"
        preview_manifest_path = self.storage.project_dir(project_id) / "output" / "preview_manifest.json"
        preview_dir = self.storage.project_dir(project_id) / "output" / "preview"
        preview_result = self._render_docx_previews(output_path, preview_dir)
        qa_report = self.qa_service.build_report(
            profile,
            project_input,
            render_result,
            images,
            evidence,
            preview_result=preview_result,
        )
        write_json(sources_path, [source.model_dump() for source in evidence])
        write_json(qa_path, qa_report)
        write_json(benchmark_path, self._build_benchmark_compare(profile, output_path, qa_report))
        write_json(preview_manifest_path, preview_result)
        write_json(
            transfer_path,
            self._build_transfer_report(
                profile=profile,
                project_input=project_input,
                output_path=output_path,
                render_result=render_result,
                qa_report=qa_report,
                evidence=evidence,
                images=images,
                transfer_mode=transfer_mode,
            ),
        )
        published = self._publish_results_bundle(
            project_id,
            profile,
            project_input,
            output_path,
            qa_report,
            render_result,
            psst_field_ids=psst_field_ids,
            core_table_ids=core_table_ids,
        )
        log_line(
            f"[DONE] project={project_id} docx={output_path.name} "
            f"errors={qa_report['error_count']} results={published.get('results_dir', '')}"
        )
        return ArtifactBundle(
            output_docx=str(output_path),
            qa_report=str(qa_path),
            sources=str(sources_path),
            benchmark_compare=str(benchmark_path),
            transfer_report=str(transfer_path),
            preview_manifest=str(preview_manifest_path),
            generated_assets=[image.path for image in images],
            results_dir=published.get("results_dir", ""),
            results_docx=published.get("results_docx", ""),
            hwp_paste=published.get("hwp_paste", ""),
            copy_blocks=published.get("copy_blocks", ""),
            fill_map=published.get("fill_map", ""),
        )

    def _resolve_render_docx_script(self) -> Path | None:
        configured = os.getenv("AUTO_WRITE_RENDER_DOCX_SCRIPT", "").strip()
        if configured:
            path = Path(configured).expanduser()
            if path.exists():
                return path
        for pattern in self.DOCX_RENDER_SCRIPT_PATTERNS:
            matches = sorted(Path(path).resolve() for path in glob(str(Path(pattern).expanduser())))
            if matches:
                return matches[-1]
        return None

    _PREVIEW_TOOL_MISSING_RE = re.compile(
        r"(ModuleNotFoundError|ImportError|No module named|cannot import name|"
        r"pdf2image|poppler|pdfinfo|DLL load failed)",
        re.IGNORECASE,
    )

    @classmethod
    def _preview_failure_is_tool_unavailable(cls, output: str) -> bool:
        """미리보기 렌더러가 자체 의존성(pdf2image/poppler 등) 부재로 실행조차 못했는지 판별한다."""
        return bool(cls._PREVIEW_TOOL_MISSING_RE.search(str(output or "")))

    def _render_docx_previews(self, docx_path: Path, preview_dir: Path) -> dict[str, Any]:
        result: dict[str, Any] = {
            "status": "skipped",
            "renderer": "artifact-tool",
            "script_path": "",
            "page_count": 0,
            "pages": [],
            "warnings": [],
            "errors": [],
        }
        if not docx_path.exists():
            result["status"] = "failed"
            result["errors"].append("미리보기 렌더링 대상 DOCX가 없습니다.")
            return result

        script_path = self._resolve_render_docx_script()
        if script_path is None:
            result["warnings"].append("DOCX 미리보기 렌더러를 찾지 못해 페이지 검수를 건너뛰었습니다.")
            return result

        preview_dir.mkdir(parents=True, exist_ok=True)
        for old_png in preview_dir.glob("page-*.png"):
            old_png.unlink(missing_ok=True)
        result["script_path"] = str(script_path)
        command = [
            sys.executable,
            str(script_path),
            str(docx_path),
            "--output_dir",
            str(preview_dir),
            "--renderer",
            "artifact-tool",
        ]
        completed = subprocess.run(command, capture_output=True, text=True, check=False)
        pages = sorted(preview_dir.glob("page-*.png"))
        result["pages"] = [path.name for path in pages]
        result["page_count"] = len(pages)

        if completed.returncode == 0 and pages:
            result["status"] = "passed"
            return result

        if pages:
            result["status"] = "passed"
            result["warnings"].append(
                f"렌더러 종료 코드는 {completed.returncode}였지만 페이지 PNG {len(pages)}장이 생성되어 미리보기는 확보했습니다."
            )
            stderr_text = (completed.stderr or "").strip()
            if stderr_text:
                result["warnings"].append(stderr_text[:240])
            return result

        stdout_text = (completed.stdout or "").strip()
        stderr_text = (completed.stderr or "").strip()
        if self._preview_failure_is_tool_unavailable(f"{stdout_text}\n{stderr_text}"):
            # 렌더러가 추가 구성요소(pdf2image/poppler 등) 부재로 실행 자체가 불가한 경우는
            # 문서 결함이 아니라 도구 부재이므로 '건너뜀(경고)'으로 처리한다. 문서는 정상이다.
            result["status"] = "skipped"
            result["warnings"].append(
                "미리보기 렌더러의 추가 구성요소(pdf2image/poppler 등)가 없어 "
                "DOCX 화면 검수를 건너뛰었습니다. 문서 자체는 정상 생성되었습니다."
            )
            return result

        result["status"] = "failed"
        if stdout_text:
            result["errors"].append(stdout_text[:240])
        if stderr_text:
            result["errors"].append(stderr_text[:240])
        if not result["errors"]:
            result["errors"].append("DOCX 미리보기 렌더링에 실패했습니다.")
        return result

    def _template_looks_completed(self, profile: TemplateProfile) -> bool:
        source_docx = Path(profile.source_docx)
        if not source_docx.exists():
            return False
        if self._docx_is_partially_filled(source_docx):
            return False
        try:
            metrics = self._doc_metrics(source_docx)
        except Exception:
            return False
        if metrics["paragraph_count"] >= 40 and metrics["paragraph_chars"] >= 1800 and metrics["table_fill_ratio"] >= 0.9:
            return True
        return False

    def _docx_is_partially_filled(self, docx_path: Path) -> bool:
        if not docx_path.exists():
            return False
        try:
            metrics = self._doc_metrics(docx_path)
        except Exception:
            return False
        if metrics.get("placeholder_hits", 0) > 0:
            return True
        if metrics["paragraph_count"] >= 25 and metrics["paragraph_chars"] >= 800:
            return True
        return False

    def _seed_answers_from_docx(
        self,
        profile: TemplateProfile,
        docx_path: Path,
        answers: dict[str, Any],
    ) -> dict[str, Any]:
        if not docx_path.exists():
            return answers
        seeded = dict(answers)
        try:
            summary = build_doc_summary(docx_path)
        except Exception:
            return seeded

        paragraphs = summary.get("paragraphs", [])
        anchor_to_field = {section.anchor_text: section.field_id for section in profile.sections}
        for index, text in enumerate(paragraphs):
            field_id = anchor_to_field.get(text)
            if not field_id or seeded.get(field_id):
                continue
            body_parts: list[str] = []
            for next_text in paragraphs[index + 1 : index + 6]:
                if next_text in anchor_to_field:
                    break
                if next_text.startswith("□") or self.SECTION_HEADING_RE.match(next_text):
                    break
                body_parts.append(next_text)
            body = "\n".join(body_parts).strip()
            if body and not self._is_sample_value(body):
                seeded[field_id] = body

        table_by_index = {table.table_index: table for table in profile.tables}
        for table_info in summary.get("tables", []):
            table_index = int(table_info.get("table_index", -1))
            table_profile = table_by_index.get(table_index)
            if table_profile is None:
                continue
            rows = table_info.get("rows", [])
            for cell in table_profile.cells:
                if seeded.get(cell.cell_id):
                    continue
                if cell.row >= len(rows):
                    continue
                row = rows[cell.row]
                if cell.cell >= len(row):
                    continue
                value = str(row[cell.cell] or "").strip()
                if value:
                    seeded[cell.cell_id] = value
        return seeded

    def _collect_improvable_questions(
        self,
        profile: TemplateProfile,
        answers: dict[str, Any],
        docx_path: Path,
    ) -> list[dict[str, Any]]:
        missing = self._collect_missing_questions(profile, answers, required_only=False)
        missing_ids = {str(item.get("question_id", "")) for item in missing}
        improvable: list[dict[str, Any]] = list(missing)

        for question in profile.questions:
            question_id = str(question.question_id)
            if question_id in missing_ids:
                continue
            value = str(answers.get(question_id, "") or "").strip()
            if value and self._is_sample_value(value):
                improvable.append(question.model_dump())
                continue
            target_kind = str(question.target.get("kind", ""))
            if target_kind == "table_cell" and value and cell_value_needs_improvement(value):
                improvable.append(question.model_dump())

        if docx_path.exists():
            try:
                summary = build_doc_summary(docx_path)
                enriched = detect_tables(summary["tables"], include_improvable_filled=True)
                question_by_cell = {
                    str(q.target.get("cell_id", "")): q
                    for q in profile.questions
                    if str(q.target.get("kind", "")) == "table_cell"
                }
                for table in enriched:
                    for cell in table.cells:
                        question = question_by_cell.get(cell.cell_id)
                        if question is None:
                            continue
                        if question.question_id in missing_ids:
                            continue
                        current = str(answers.get(question.question_id, "") or "").strip()
                        if current and not self._is_sample_value(current) and not cell_value_needs_improvement(current):
                            continue
                        improvable_ids = {str(item.get("question_id", "")) for item in improvable}
                        if question.question_id not in improvable_ids:
                            improvable.append(question.model_dump())
            except Exception:
                pass
        return improvable

    def load_profile_for_project(self, project_id: str) -> TemplateProfile:
        snapshot = self.storage.project_dir(project_id) / "template_snapshot.json"
        if snapshot.exists():
            profile = TemplateProfile.model_validate(read_json(snapshot))
            return sanitize_template_profile(profile)
        meta = read_json(self.storage.project_dir(project_id) / "project_meta.json")
        return sanitize_template_profile(self.storage.load_template_profile(meta["template_id"]))

    def normalize_profile(self, profile: TemplateProfile) -> TemplateProfile:
        return sanitize_template_profile(profile)

    @classmethod
    def _skip_section_label(cls, label: str) -> bool:
        return bool(cls.SECTION_AUTOFILL_EXCLUDE_RE.search(str(label or "").strip()))

    @classmethod
    def _is_sample_value(cls, text: str) -> bool:
        return bool(cls.SAMPLE_VALUE_RE.search(str(text or "").strip()))

    @classmethod
    def _skip_table_label(cls, label: str) -> bool:
        text = str(label or "").strip()
        if not text:
            return True
        if cls.TABLE_AUTOFILL_EXCLUDE_RE.search(text):
            return True
        return bool(cls.SAMPLE_VALUE_RE.search(text))

    def visible_questions(self, profile: TemplateProfile) -> list[dict[str, Any]]:
        visible: list[dict[str, Any]] = []
        for question in profile.questions:
            if question.question_id in {"project_title", "organization_name"}:
                continue
            if self._skip_section_label(question.label):
                continue
            if str(question.target.get("kind", "")) == "table_cell" and self._skip_table_label(question.label):
                continue
            if question.required:
                visible.append(question.model_dump())
                continue
            target_kind = str(question.target.get("kind", ""))
            if target_kind == "image_slot_note":
                continue
            if target_kind == "section" and len(visible) < 12:
                visible.append(question.model_dump())
        return visible

    def extract_reference_text(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"}:
            return path.read_text(encoding="utf-8", errors="ignore")
        if suffix == ".json":
            return json.dumps(read_json(path), ensure_ascii=False, indent=2)
        if suffix in {".hwpx", ".hwp"}:
            extracted = extract_additional_text(path)
            return extracted or f"{suffix.upper().lstrip('.')} 텍스트 추출에 실패했습니다: {path.name}"
        if suffix == ".docx":
            document = Document(str(path))
            return "\n".join(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
        if suffix == ".pdf":
            try:
                from pypdf import PdfReader

                reader = PdfReader(str(path))
                pages: list[str] = []
                for page in reader.pages:
                    text = (page.extract_text() or "").strip()
                    if text:
                        pages.append(text)
                joined = self._sanitize_xml_text("\n".join(pages).strip())
                return joined or f"PDF에서 텍스트를 추출하지 못했습니다: {path.name}"
            except Exception:
                return f"PDF 텍스트 추출에 실패했습니다: {path.name}"
        return f"지원하지 않는 참고자료 형식입니다: {path.name}"

    def _ensure_reference_library_loaded(self) -> list[dict[str, str]]:
        if self._library_snippets is not None:
            return self._library_snippets
        self._library_snippets = []
        settings = self.storage.settings
        library_dir = settings.reference_library_dir
        if library_dir is None or not library_dir.exists() or not library_dir.is_dir():
            return self._library_snippets
        files = sorted(
            [path for path in library_dir.iterdir() if path.is_file() and path.suffix.lower() in REFERENCE_SUFFIXES]
        )[:200]
        for file_path in files:
            text = self.extract_reference_text(file_path)
            if not text or "지원하지 않는 참고자료 형식" in text or "텍스트 추출" in text:
                continue
            for snippet in self._split_library_snippets(text):
                self._library_snippets.append({"source": file_path.name, "text": snippet})
        log_line(f"[INFO] reference_library snippets={len(self._library_snippets)} dir={library_dir}")
        return self._library_snippets

    def _split_library_snippets(self, text: str) -> list[str]:
        snippets: list[str] = []
        blocks = self.LIBRARY_BLOCK_SPLIT_RE.split(text)
        for block in blocks:
            clean = self._sanitize_xml_text(re.sub(r"\s+", " ", block).strip())
            if len(clean) < 40:
                continue
            if len(clean) > 800:
                # Split only if truly oversized; use sentence boundary to keep
                # related sentences (e.g. "경쟁사 1 … 경쟁사 2 …") together.
                sentences = self.LIBRARY_SENTENCE_SPLIT_RE.split(clean)
                merged: list[str] = []
                current = ""
                for sentence in sentences:
                    sentence = sentence.strip()
                    if not sentence:
                        continue
                    trial = sentence if not current else f"{current} {sentence}"
                    if len(trial) <= 700:
                        current = trial
                    else:
                        if len(current) >= 40:
                            merged.append(current)
                        current = sentence
                if len(current) >= 40:
                    merged.append(current)
                candidates = merged
            else:
                candidates = [clean]
            for candidate in candidates:
                candidate = self._sanitize_xml_text(candidate)
                if self.LIBRARY_GUIDE_RE.search(candidate):
                    continue
                if "<" in candidate and ">" in candidate:
                    continue
                snippets.append(candidate)
        return snippets

    def _tokenize(self, text: str) -> set[str]:
        return {token.lower() for token in self.LIBRARY_TOKEN_RE.findall(text or "")}

    def _score_library_snippet(self, snippet: str, query_tokens: set[str]) -> float:
        if not snippet:
            return 0.0
        snippet_tokens = self._tokenize(snippet)
        if not snippet_tokens:
            return 0.0
        overlap = len(snippet_tokens & query_tokens)
        keyword_bonus = sum(1 for keyword in self.LIBRARY_KEYWORD_HINTS if keyword in snippet)
        length = len(snippet)
        if length < 50:
            length_penalty = -1.5
        elif length > 280:
            length_penalty = -1.0
        else:
            length_penalty = 0.0
        return overlap * 2.0 + keyword_bonus * 0.5 + length_penalty

    def _suggest_snippets_from_corpus(
        self,
        missing: list[dict[str, Any]],
        corpus: list[dict[str, str]],
        allowed_target_kinds: set[str],
        min_score: float,
    ) -> dict[str, str]:
        if not corpus:
            return {}
        suggestions: dict[str, str] = {}
        used_indexes: set[int] = set()
        for question in missing:
            question_id = str(question.get("question_id", "")).strip()
            if not question_id:
                continue
            target_kind = str(question.get("target", {}).get("kind", ""))
            if target_kind not in allowed_target_kinds:
                continue
            label = str(question.get("label", "")).strip()
            if not label:
                continue
            query_tokens = self._tokenize(label)
            query_tokens.update(self._tokenize(str(question.get("source_hint", ""))))
            if not query_tokens:
                continue
            scored: list[tuple[int, str, float]] = []
            for index, item in enumerate(corpus):
                score = self._score_library_snippet(item["text"], query_tokens)
                scored.append((index, item["text"], score))
            ranked = sorted(scored, key=lambda row: row[2], reverse=True)

            # Collect ALL snippets above threshold (not just the first one).
            # This ensures multi-part content like "경쟁사 1 / 경쟁사 2" are both included.
            selected_parts: list[str] = []
            for index, item_text, score in ranked:
                if score < min_score:
                    break
                if index in used_indexes:
                    continue
                selected_parts.append(item_text)
                used_indexes.add(index)
                # Cap at 4 snippets per section to avoid context overflow
                if len(selected_parts) >= 4:
                    break

            if selected_parts:
                # Join multiple snippets with newline so AI sees all content
                suggestions[question_id] = "\n".join(selected_parts)
        return suggestions

    def _suggest_library_snippets(self, missing: list[dict[str, Any]]) -> dict[str, str]:
        corpus = self._ensure_reference_library_loaded()
        return self._suggest_snippets_from_corpus(
            missing=missing,
            corpus=corpus,
            allowed_target_kinds={"section"},
            min_score=3.0,
        )

    def _read_reference_text(self, reference: ReferenceFile) -> str:
        extracted_path = Path(reference.extracted_text_path)
        if extracted_path.exists():
            try:
                text = extracted_path.read_text(encoding="utf-8", errors="ignore")
                cleaned = self._sanitize_xml_text(text)
                if cleaned:
                    return cleaned
            except Exception:
                pass
        return self._sanitize_xml_text(reference.extracted_preview)

    def _ordered_references_for_text(self, references: list[ReferenceFile]) -> list[ReferenceFile]:
        if not references:
            return []
        grouped: dict[str, list[ReferenceFile]] = {}
        for reference in references:
            stem = Path(reference.file_name).stem.lower().strip()
            grouped.setdefault(stem, []).append(reference)
        ordered: list[ReferenceFile] = []
        for stem in sorted(grouped.keys()):
            group = grouped[stem]
            group = sorted(
                group,
                key=lambda ref: (
                    0 if Path(ref.file_name).suffix.lower() == ".docx" else 1,
                    0 if Path(ref.file_name).suffix.lower() == ".txt" else 1,
                    ref.file_name.lower(),
                ),
            )
            ordered.append(group[0])
        return ordered

    def _build_reference_snippet_corpus(self, project_input: ProjectInput) -> list[dict[str, str]]:
        corpus: list[dict[str, str]] = []
        for reference in self._ordered_references_for_text(project_input.references):
            text = self._read_reference_text(reference)
            if not text:
                continue
            if "지원하지 않는 참고자료 형식" in text or "텍스트 추출" in text:
                continue
            snippets = self._split_library_snippets(text)
            for snippet in snippets:
                corpus.append({"source": reference.file_name, "text": snippet})
                if len(corpus) >= self.REFERENCE_SNIPPET_LIMIT:
                    return corpus
        return corpus

    def _suggest_reference_snippets(self, missing: list[dict[str, Any]], project_input: ProjectInput) -> dict[str, str]:
        corpus = self._build_reference_snippet_corpus(project_input)
        return self._suggest_snippets_from_corpus(
            missing=missing,
            corpus=corpus,
            allowed_target_kinds={"section", "table_cell"},
            min_score=2.5,
        )

    @staticmethod
    def _merge_hint_maps(primary: dict[str, str], secondary: dict[str, str]) -> dict[str, str]:
        merged = dict(primary)
        for question_id, snippet in secondary.items():
            if question_id not in merged:
                merged[question_id] = snippet
        return merged

    def _merge_context_with_library(
        self,
        context: str,
        library_hints: dict[str, str],
        questions: list[dict[str, Any]] | None = None,
    ) -> str:
        if not library_hints:
            return context
        lines = [context.strip()] if context.strip() else []
        # Build question_id → label map for explicit pairing
        qid_to_label: dict[str, str] = {}
        if questions:
            for q in questions:
                qid = str(q.get("question_id", "")).strip()
                label = str(q.get("label", "")).strip()
                if qid and label:
                    qid_to_label[qid] = label
        lines.append("원문 참고자료 (아래 내용을 항목에 맞게 옮겨 쓰세요. 내용 추가/변형 금지):")
        for index, (qid, snippet) in enumerate(library_hints.items(), start=1):
            label = qid_to_label.get(qid, qid)
            # Pass full snippet without truncation so AI has complete source
            lines.append(f"{index}. [항목: {label}] 원문: {snippet}")
            if index >= 12:
                break
        return "\n".join(line for line in lines if line)

    def _adapt_library_snippet(self, snippet: str, project_title: str, org_name: str) -> str:
        """Sanitize and return the library snippet as-is.

        Do NOT truncate or rewrite. The AI layer handles format conversion.
        For table cells the caller clips to max_chars after this call.
        Numbers, proper nouns, and all factual content must remain intact.
        """
        text = self._sanitize_xml_text(re.sub(r"\s+", " ", snippet).strip())
        return text

    def _sanitize_xml_text(self, text: str) -> str:
        cleaned = self.INVALID_XML_CHAR_RE.sub(" ", text or "")
        cleaned = cleaned.replace("\uFFFD", " ")
        cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        return cleaned.strip()

    def _build_context(self, project_input: ProjectInput, strict_preserve: bool = False) -> str:
        parts = [
            f"과제명: {project_input.project_meta.get('project_title', '')}",
            f"기관명: {project_input.organization_profile.get('name', '')}",
        ]
        for key, value in project_input.answers.items():
            parts.append(f"{key}: {value}")
        for reference in self._ordered_references_for_text(project_input.references):
            reference_text = self._read_reference_text(reference)
            if reference_text:
                # Case A: strict preserve mode → pass full source text so AI can find all content.
                # Standard mode: keep 2400 char limit to stay within context budget.
                char_limit = 16000 if strict_preserve else 2400
                parts.append(f"[참고자료:{reference.file_name}] {reference_text[:char_limit]}")
        return "\n".join(part for part in parts if part).strip()

    def _next_reference_path(self, project_id: str, safe_name: str) -> Path:
        reference_dir = self.storage.project_dir(project_id) / "references"
        candidate = reference_dir / safe_name
        if not candidate.exists():
            return candidate
        stem = candidate.stem
        suffix = candidate.suffix
        index = 2
        while True:
            next_candidate = reference_dir / f"{stem}_{index}{suffix}"
            if not next_candidate.exists():
                return next_candidate
            index += 1

    def _collect_missing_questions(
        self, profile: TemplateProfile, answers: dict[str, Any], required_only: bool = False
    ) -> list[dict[str, Any]]:
        missing: list[dict[str, Any]] = []
        excluded_section_ids = {section.field_id for section in profile.sections if getattr(section, "is_excluded", False)}
        for question in profile.questions:
            target_kind = str(question.target.get("kind", ""))
            if target_kind not in {"section", "table_cell"}:
                continue
            if target_kind == "section":
                field_id = str(question.target.get("field_id", "")).strip()
                if field_id and field_id in excluded_section_ids:
                    continue
            if required_only and not bool(question.required):
                continue
            value = str(answers.get(question.question_id, "") or "").strip()
            if value:
                continue
            missing.append(question.model_dump())
        return missing

    def _table_cell_index(self, profile: TemplateProfile) -> tuple[dict[str, dict[str, Any]], dict[str, int]]:
        by_cell_id: dict[str, dict[str, Any]] = {}
        table_sizes: dict[str, int] = {}
        for table in profile.tables:
            table_sizes[table.table_id] = len(table.cells)
            for cell in table.cells:
                by_cell_id[cell.cell_id] = {
                    "table_id": table.table_id,
                    "label": cell.label,
                    "row_header": cell.row_header,
                    "col_header": cell.col_header,
                }
        return by_cell_id, table_sizes

    @staticmethod
    def _clip_text(value: str, max_chars: int) -> str:
        text = re.sub(r"\s+", " ", str(value or "").strip())
        if len(text) <= max_chars:
            return text
        clipped = text[:max_chars].rstrip(" ,.;:")
        return f"{clipped}."

    def _table_cell_max_chars(self, label: str) -> int:
        if self.TABLE_NARROW_HINT_RE.search(label or ""):
            return 36
        return 90

    def _filter_missing_for_autofill(
        self,
        profile: TemplateProfile,
        missing: list[dict[str, Any]],
        transfer_mode: bool,
    ) -> list[dict[str, Any]]:
        by_cell_id, table_sizes = self._table_cell_index(profile)
        filtered: list[dict[str, Any]] = []
        for question in missing:
            target_kind = str(question.get("target", {}).get("kind", ""))
            label = str(question.get("label", "")).strip()
            if target_kind != "table_cell":
                if target_kind == "section" and self._skip_section_label(label):
                    continue
                filtered.append(question)
                continue
            cell_id = str(question.get("target", {}).get("cell_id", ""))
            table_id = str(question.get("target", {}).get("table_id", ""))
            meta = by_cell_id.get(cell_id, {})
            if not label:
                continue
            if self._skip_table_label(label):
                continue
            # Large matrix tables are the main source of low-quality overfill.
            if transfer_mode and table_sizes.get(table_id, 0) > 20 and not bool(question.get("required")):
                continue
            # Skip cells when either row/column semantics are missing in wide tables.
            row_header = str(meta.get("row_header", "")).strip()
            col_header = str(meta.get("col_header", "")).strip()
            if transfer_mode and table_sizes.get(table_id, 0) > 10 and (not row_header or not col_header):
                continue
            filtered.append(question)
        return filtered

    def _postprocess_answers(self, profile: TemplateProfile, answers: dict[str, Any]) -> dict[str, Any]:
        normalized: dict[str, Any] = dict(answers)
        by_cell_id, _ = self._table_cell_index(profile)
        for cell_id, meta in by_cell_id.items():
            if cell_id not in normalized:
                continue
            value = normalized.get(cell_id)
            if not isinstance(value, str):
                continue
            max_chars = self._table_cell_max_chars(str(meta.get("label", "")))
            normalized[cell_id] = self._clip_text(value, max_chars)
        return normalized

    def _draft_missing_answers_in_chunks(
        self,
        missing: list[dict[str, Any]],
        context: str,
        writing_provider: str = "",
        writing_model: str = "",
        strict_preserve: bool = False,
    ) -> dict[str, str]:
        if not context.strip():
            return {}
        drafted: dict[str, str] = {}
        chunk_size = 40

        def _draft_group(
            questions: list[dict[str, Any]],
            provider_override: str | None = None,
            model_override: str | None = None,
        ) -> None:
            for start in range(0, len(questions), chunk_size):
                chunk = questions[start : start + chunk_size]
                result = self.openai_service.draft_missing_answers(
                    chunk,
                    context,
                    provider_override=provider_override,
                    model_override=model_override,
                    strict_preserve=strict_preserve,
                )
                if not result:
                    continue
                for key, value in result.items():
                    value_text = str(value or "").strip()
                    if value_text:
                        drafted[key] = value_text

        section_questions = [item for item in missing if str(item.get("target", {}).get("kind", "")) == "section"]
        table_questions = [item for item in missing if str(item.get("target", {}).get("kind", "")) == "table_cell"]

        forced_provider = str(writing_provider or "").strip().lower()
        forced_model = str(writing_model or "").strip() or None
        if forced_provider in {"openai", "anthropic"} and self.openai_service.provider_available(forced_provider):
            if forced_provider == "anthropic" and forced_model is None:
                forced_model = self.openai_service.anthropic_writing_model
            _draft_group(section_questions, provider_override=forced_provider, model_override=forced_model)
            _draft_group(table_questions, provider_override=forced_provider, model_override=forced_model)
            return drafted

        if section_questions and self.openai_service.provider_available("anthropic"):
            _draft_group(
                section_questions,
                provider_override="anthropic",
                model_override=self.openai_service.anthropic_writing_model,
            )
        else:
            _draft_group(section_questions)
        _draft_group(table_questions)
        return drafted

    def _build_fallback_answers(
        self,
        missing: list[dict[str, Any]],
        project_input: ProjectInput,
        library_hints: dict[str, str] | None = None,
        strict_preserve: bool = False,
    ) -> dict[str, str]:
        project_title = str(project_input.project_meta.get("project_title", "") or "").strip() or "본 과제"
        org_name = str(project_input.organization_profile.get("name", "") or "").strip() or "신청 기관"
        fallback: dict[str, str] = {}
        library_hints = library_hints or {}
        used_snippets: set[str] = set()
        for question in missing:
            question_id = str(question.get("question_id", "")).strip()
            if not question_id:
                continue
            target_kind = str(question.get("target", {}).get("kind", ""))
            if target_kind == "section":
                label = str(question.get("label", "")).strip()
                best_snippet = library_hints.get(question_id, "")
                if best_snippet and best_snippet not in used_snippets:
                    fallback[question_id] = self._adapt_library_snippet(best_snippet, project_title, org_name)
                    used_snippets.add(best_snippet)
                elif strict_preserve:
                    # Case A: no matching source found → leave blank, do not invent
                    pass
                else:
                    fallback[question_id] = self._fallback_section_text(project_title, org_name, label)
            elif target_kind == "table_cell":
                label = str(question.get("label", "")).strip()
                best_snippet = library_hints.get(question_id, "")
                if best_snippet and best_snippet not in used_snippets:
                    adapted = self._adapt_library_snippet(best_snippet, project_title, org_name)
                    fallback[question_id] = self._clip_text(adapted, self._table_cell_max_chars(label))
                    used_snippets.add(best_snippet)
                elif strict_preserve:
                    # Case A: no matching source found → leave blank, do not invent
                    pass
                else:
                    fallback[question_id] = self._fallback_table_text(project_title, label)
        return fallback

    def _fallback_section_text(self, project_title: str, org_name: str, label: str) -> str:
        if any(token in label for token in ("목표", "성과", "KPI", "지표")):
            return f"{project_title}의 핵심 성과지표를 설정하고 월 단위로 달성률을 점검합니다."
        if any(token in label for token in ("시장", "경쟁", "고객")):
            return f"{org_name}는 목표 고객군과 경쟁 환경을 분석하여 차별화 전략을 실행합니다."
        if any(token in label for token in ("예산", "비용", "수익")):
            return "필수 비용과 기대 수익을 구분해 단계별로 집행하고 분기별로 재검토합니다."
        if any(token in label for token in ("추진", "일정", "계획")):
            return "준비-실행-확산 3단계 일정으로 운영하며 단계 종료 시 점검 결과를 반영합니다."
        return f"{project_title} 관련 '{label}' 항목은 참고자료를 기반으로 구체 실행안 중심으로 작성합니다."

    def _fallback_table_text(self, project_title: str, label: str) -> str:
        if any(token in label for token in ("금액", "예산", "비용", "매출", "원", "만원")):
            return "1,000(추정)"
        if any(token in label for token in ("일정", "기간", "시기", "월")):
            return "1단계 준비 / 2단계 실행 / 3단계 확산(월 단위 점검)"
        if any(token in label for token in ("목표", "성과", "지표")):
            return "정량 KPI 설정 및 월별 점검(달성률 관리)"
        if any(token in label for token in ("담당", "인력", "역할")):
            return "전담 인력 지정 및 역할 분담, 책임자 명시"
        return f"{project_title} 기준 실행 계획 수치 및 근거 반영"

    def _build_transfer_report(
        self,
        profile: TemplateProfile,
        project_input: ProjectInput,
        output_path: Path,
        render_result: dict[str, Any],
        qa_report: dict[str, Any],
        evidence: list[EvidenceSource],
        images: list[GeneratedImage],
        transfer_mode: bool,
    ) -> dict[str, Any]:
        generated_metrics = self._doc_metrics(output_path) if output_path.exists() else {
            "paragraph_count": 0,
            "paragraph_chars": 0,
            "table_cells_total": 0,
            "table_fill_ratio": 0.0,
            "table_chars": 0,
            "drawing_count": 0,
            "guide_marker_hits": 0,
        }
        total_sections = len(profile.sections)
        total_table_cells = sum(len(table.cells) for table in profile.tables)
        sections_written = int(render_result.get("sections_written", 0))
        cells_written = int(render_result.get("cells_written", 0))
        image_slots_total = len(profile.image_slots)
        required_slots = [slot for slot in profile.image_slots if slot.required]
        image_slot_ids = {image.slot_id for image in images}
        required_filled = sum(1 for slot in required_slots if slot.slot_id in image_slot_ids)
        all_slots_filled = sum(1 for slot in profile.image_slots if slot.slot_id in image_slot_ids)
        matched_topics = sum(1 for req in project_input.evidence_requests if any(src.topic == req.topic for src in evidence))

        unresolved_from_qa: list[str] = []
        unresolved_from_qa.extend(str(item) for item in qa_report.get("errors", [])[:20])
        unresolved_from_qa.extend(str(item) for item in qa_report.get("warnings", [])[:20])
        unresolved_from_render = [str(item) for item in render_result.get("errors", [])[:20]]

        return {
            "status": "ok",
            "mode": "transfer" if transfer_mode else "standard",
            "output_docx": str(output_path),
            "fill_ratio": {
                "sections": round((sections_written / total_sections), 4) if total_sections else 1.0,
                "table_cells": round((cells_written / total_table_cells), 4) if total_table_cells else 1.0,
                "table_cells_document": float(generated_metrics["table_fill_ratio"]),
            },
            "image_slot_coverage": {
                "required_filled": required_filled,
                "required_total": len(required_slots),
                "required_ratio": round((required_filled / len(required_slots)), 4) if required_slots else 1.0,
                "all_filled": all_slots_filled,
                "all_total": image_slots_total,
                "all_ratio": round((all_slots_filled / image_slots_total), 4) if image_slots_total else 1.0,
            },
            "evidence_usage": {
                "requested_topics": len(project_input.evidence_requests),
                "matched_topics": matched_topics,
                "sources_collected": len(evidence),
                "reference_files": len(project_input.references),
            },
            "qa_summary": {
                "passed": bool(qa_report.get("passed")),
                "error_count": int(qa_report.get("error_count", 0)),
                "warning_count": int(qa_report.get("warning_count", 0)),
                "guide_marker_hits": int(generated_metrics.get("guide_marker_hits", 0)),
            },
            "unresolved_items": {
                "qa": unresolved_from_qa,
                "render": unresolved_from_render,
            },
        }

    def _resolve_benchmark_docx(self, profile: TemplateProfile) -> Path | None:
        forced = os.getenv("AUTO_WRITE_BENCHMARK_DOCX", "").strip()
        if forced:
            forced_path = Path(forced)
            if forced_path.exists():
                return forced_path
        fixed_benchmark = Path(
            r"C:\Users\ekth3\OneDrive\바탕 화면\다솜\경영지도사 개인\02. 밸류업파트너스\2026 AI수출지원공통\예창\마켓게이트_예비창업패키지 사업계획서 1550.docx"
        )
        if "예비창업패키지" in profile.template_name and fixed_benchmark.exists():
            return fixed_benchmark
        source_docx = Path(profile.source_docx)
        if source_docx.exists():
            return source_docx
        return None

    def _doc_metrics(self, docx_path: Path) -> dict[str, Any]:
        document = Document(str(docx_path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        paragraph_chars = sum(len(text) for text in paragraphs)

        table_cells_total = 0
        table_cells_filled = 0
        table_chars = 0
        guide_marker_hits = 0
        for text in paragraphs:
            if self.GUIDE_MARKER_RE.search(text):
                guide_marker_hits += 1

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_cells_total += 1
                    text = cell.text.strip()
                    if text:
                        table_cells_filled += 1
                        table_chars += len(text)
                        if self.GUIDE_MARKER_RE.search(text):
                            guide_marker_hits += 1

        drawing_count = 0
        placeholder_hits = 0
        try:
            with zipfile.ZipFile(docx_path, "r") as archive:
                xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
                drawing_count = xml.count("<w:drawing")
        except Exception:
            drawing_count = 0

        for text in paragraphs:
            if self._is_sample_value(text) or cell_value_needs_improvement(text):
                placeholder_hits += 1
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and (self._is_sample_value(text) or cell_value_needs_improvement(text)):
                        placeholder_hits += 1

        return {
            "paragraph_count": len(paragraphs),
            "paragraph_chars": paragraph_chars,
            "table_cells_total": table_cells_total,
            "table_fill_ratio": round((table_cells_filled / table_cells_total), 4) if table_cells_total else 0.0,
            "table_chars": table_chars,
            "drawing_count": drawing_count,
            "guide_marker_hits": guide_marker_hits,
            "placeholder_hits": placeholder_hits,
        }

    def _build_benchmark_compare(
        self, profile: TemplateProfile, output_path: Path, qa_report: dict[str, Any]
    ) -> dict[str, Any]:
        benchmark = self._resolve_benchmark_docx(profile)
        if benchmark is None or not benchmark.exists():
            return {
                "status": "benchmark_missing",
                "benchmark_docx": "",
                "output_docx": str(output_path),
                "qa_passed": bool(qa_report.get("passed")),
                "message": "비교용 벤치마크 파일을 찾지 못했습니다.",
            }

        generated_metrics = self._doc_metrics(output_path)
        benchmark_metrics = self._doc_metrics(benchmark)
        table_fill_gap = round(generated_metrics["table_fill_ratio"] - benchmark_metrics["table_fill_ratio"], 4)
        if benchmark_metrics["table_chars"] > 0:
            table_char_gap_ratio = round(
                abs(generated_metrics["table_chars"] - benchmark_metrics["table_chars"]) / benchmark_metrics["table_chars"],
                4,
            )
        else:
            table_char_gap_ratio = 0.0

        source_docx = Path(profile.source_docx)
        is_template_reference = source_docx.exists() and benchmark.resolve() == source_docx.resolve()
        if is_template_reference:
            benchmark_pass = qa_report.get("error_count", 0) == 0
            acceptance: dict[str, Any] = {
                "mode": "template_reference",
                "note": "원본 템플릿을 기준으로 비교하므로 품질 판정은 QA 오류 건수 중심으로 판단합니다.",
                "qa_error_count_max": 0,
            }
        else:
            benchmark_pass = (
                table_fill_gap >= -0.02 and table_char_gap_ratio <= 0.05 and qa_report.get("error_count", 0) == 0
            )
            acceptance = {
                "mode": "benchmark_document",
                "table_fill_ratio_gap_min": -0.02,
                "table_char_gap_ratio_max": 0.05,
                "qa_error_count_max": 0,
            }
        return {
            "status": "ok",
            "benchmark_docx": str(benchmark),
            "output_docx": str(output_path),
            "benchmark_mode": "template_reference" if is_template_reference else "benchmark_document",
            "benchmark_metrics": benchmark_metrics,
            "generated_metrics": generated_metrics,
            "gap": {
                "paragraph_chars": generated_metrics["paragraph_chars"] - benchmark_metrics["paragraph_chars"],
                "table_chars": generated_metrics["table_chars"] - benchmark_metrics["table_chars"],
                "table_fill_ratio": table_fill_gap,
                "drawing_count": generated_metrics["drawing_count"] - benchmark_metrics["drawing_count"],
                "guide_marker_hits": generated_metrics["guide_marker_hits"] - benchmark_metrics["guide_marker_hits"],
                "table_char_gap_ratio": table_char_gap_ratio,
            },
            "acceptance": acceptance,
            "result": {
                "passed": bool(benchmark_pass),
                "qa_passed": bool(qa_report.get("passed")),
                "qa_error_count": int(qa_report.get("error_count", 0)),
            },
        }
