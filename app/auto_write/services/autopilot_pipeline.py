"""autopilot_pipeline.py — 문서 품질 '수정' 전 단계를 무인 연속 실행하는 통합 파이프라인.

진단만 하던 단계들을 **실제 수정**으로 잇고 한 번에 돌린다:

  1) 백업 + 서식 수정 + 점수/게이트  — 기존 ``DocumentQualityOrchestrator.run``
     (안내문구 삭제·글머리표/표 공백·빈 단락·핵심문장 강조 + 100점 채점 + 85점 게이트)
  2) 이미지 실제 적용             — ``image_apply.apply_images``
     (그림 위치에 NotebookLM 슬라이드 생성 프롬프트 블록 삽입; 제안은 Claude, 폴백 키워드)
  3) PSST 보강                   — ``psst_fill.apply_psst_scaffold``
     (누락/미흡 영역에 작성 뼈대+가이드)
  4) 실사용 수용검사 게이트(R8)   — ``usage_acceptance.run_acceptance``
     (심사위원 관점 하드페일 검사. fail 결함이 있으면 출력 파일명을
      ``_DRAFT`` 로 강제해 '제출본' 이름으로 내보내지 않는다)
  5) 잔존 빈칸 스캔 + 통합 리포트

설계 원칙
---------
- **원본 절대 보존**: 1단계에서 원본을 ``results/backup/<ts>/`` 에 백업한다.
  이후 단계는 항상 새 출력 경로로 저장(in==out 이면 ValueError).
- **점수/게이트는 1단계(서식 수정) 기준**으로 판정한다. 이미지/PSST 보강은
  사람이 채워야 할 '시각화 자리'와 '작성 To-Do' 라서 점수를 부풀리지 않는다(날조 0).
- 빈칸 채움(submittable_filler)은 외부 plan 데이터가 필요하므로 무인 파이프라인의
  필수 단계로 넣지 않고, 잔존 빈칸을 **스캔해 보고**만 한다(있으면 To-Do 로 안내).
"""

from __future__ import annotations

import re
import shutil
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from docx import Document

from ..config import ensure_directories, get_settings
from .document_quality_orchestrator import DocumentQualityOrchestrator
from .image_apply import (
    ImageApplyReport, apply_images, extract_notebooklm_prompts, strip_notebooklm_blocks,
)
from .psst_fill import PSSTFillReport, apply_psst_scaffold
from .usage_acceptance import (
    AcceptanceConfig, SEV_FAIL, backup_existing_output, force_draft_name, run_acceptance,
)

# 잔존 빈칸(placeholder) 보수적 탐지 패턴
_RESIDUAL_RE = re.compile(
    r"(_{3,}|\[\s*\]|\(\s*작성\s*\)|※\s*작성|기재\s*바랍|예\s*:\s*\)|【\s*】)"
)


@dataclass
class AutopilotReport:
    input_docx: str
    output_docx: str = ""
    backup_dir: str = ""
    report_md: str = ""
    # 1단계(품질)
    doc_type: str = ""
    score_total: float = 0.0
    grade: str = ""
    passed: bool = False
    iterations: int = 0
    ops_summary: str = ""
    # 2단계(이미지)
    prompts_inserted: int = 0
    # 3단계(PSST)
    psst_overall_ratio: float = 0.0
    psst_areas_scaffolded: int = 0
    psst_items_added: int = 0
    psst_scaffolded_areas: list[str] = field(default_factory=list)
    # 4단계(수용검사 게이트 — R8)
    acceptance_submittable: bool = False
    acceptance_verdict: str = ""
    acceptance_fail_defects: int = 0
    acceptance_failed_checks: list[str] = field(default_factory=list)
    acceptance_error: str = ""
    draft_marked: bool = False
    draft_mark_error: str = ""
    format_mismatch: str = ""  # 산출 형식 게이트(ACC-5) — 요구 형식과 다르면 사유 기록
    overwrite_backup: str = ""  # 재실행 보호(PIPE-2) — 기존 산출물 백업 경로
    prompt_md: str = ""        # 제출 정리(US-6) — 보존된 슬라이드 프롬프트 md 경로
    strip_removed: int = 0     # 제출 정리(US-6) — 제거한 작업용 블록 단락 수
    # 5단계(잔존)
    residual_placeholders: list[str] = field(default_factory=list)
    manual_todo: list[str] = field(default_factory=list)

    def as_dict(self) -> dict[str, Any]:
        return {
            "input_docx": self.input_docx,
            "output_docx": self.output_docx,
            "backup_dir": self.backup_dir,
            "report_md": self.report_md,
            "doc_type": self.doc_type,
            "score_total": round(self.score_total, 1),
            "grade": self.grade,
            "passed": self.passed,
            "iterations": self.iterations,
            "ops_summary": self.ops_summary,
            "prompts_inserted": self.prompts_inserted,
            "psst_overall_ratio": round(self.psst_overall_ratio, 3),
            "psst_areas_scaffolded": self.psst_areas_scaffolded,
            "psst_items_added": self.psst_items_added,
            "psst_scaffolded_areas": self.psst_scaffolded_areas,
            "acceptance_submittable": self.acceptance_submittable,
            "acceptance_verdict": self.acceptance_verdict,
            "acceptance_fail_defects": self.acceptance_fail_defects,
            "acceptance_failed_checks": self.acceptance_failed_checks,
            "acceptance_error": self.acceptance_error,
            "draft_marked": self.draft_marked,
            "draft_mark_error": self.draft_mark_error,
            "format_mismatch": self.format_mismatch,
            "overwrite_backup": self.overwrite_backup,
            "prompt_md": self.prompt_md,
            "strip_removed": self.strip_removed,
            "residual_placeholders": self.residual_placeholders,
            "manual_todo": self.manual_todo,
        }


def _scan_residual(docx_path: str, *, limit: int = 20) -> list[str]:
    """결과 DOCX 에서 잔존 빈칸 표식을 보수적으로 스캔한다(보고용)."""
    found: list[str] = []
    try:
        doc = Document(docx_path)
    except Exception:
        return found
    for p in doc.paragraphs:
        t = p.text.strip()
        if t and _RESIDUAL_RE.search(t):
            found.append(t[:60])
            if len(found) >= limit:
                return found
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and _RESIDUAL_RE.search(t):
                    found.append(t[:60])
                    if len(found) >= limit:
                        return found
    return found


def _make_openai_service() -> Optional[Any]:
    """가용한 OpenAIService 를 만든다(키 없거나 실패하면 None)."""
    try:
        from .openai_client import OpenAIService

        settings = get_settings()
        svc = OpenAIService(settings)
        return svc if getattr(svc, "available", False) else None
    except Exception:
        return None


def _make_orchestrator(results_root: Path, *, use_ai: bool = False) -> DocumentQualityOrchestrator:
    openai_service = _make_openai_service() if use_ai else None
    return DocumentQualityOrchestrator(results_root, openai_service=openai_service)


def run_autopilot(
    input_docx: str,
    output_docx: Optional[str] = None,
    *,
    emphasize: bool = True,
    underline: bool = False,
    remove_guides: bool = True,
    normalize_fonts: bool = False,
    max_images: int = 8,
    placeholder_only: bool = False,
    psst_scaffold: bool = True,
    acceptance_gate: bool = True,
    blind_review: bool = False,
    required_format: Optional[str] = None,
    submit_clean: bool = False,
    max_pages: Optional[int] = None,
    ai_section_max: Optional[int] = None,
    write_report: bool = True,
) -> AutopilotReport:
    """문서 품질 수정 전 단계를 무인 연속 실행한다.

    Args:
        input_docx: 입력 DOCX 경로(원본, 자동 백업됨).
        output_docx: 최종 출력 경로. 미지정 시 results/ 아래 자동 명명.
        emphasize/underline/remove_guides/normalize_fonts: 서식 수정 옵션(1단계).
        max_images: 이미지 적용 최대 개수(2단계).
        placeholder_only: True 면 차트 생성 없이 자리표시만 삽입(2단계, 가장 안전).
        psst_scaffold: True 면 PSST 누락/미흡 영역에 작성 가이드 삽입(3단계).
        acceptance_gate: True 면 실사용 수용검사(usage_acceptance)를 실행하고,
            fail 결함이 있으면 출력 파일명을 ``_DRAFT`` 로 강제한다(4단계, R8).
        blind_review: 블라인드 공고 모드 — ○○○ 마스킹을 허용하고 실명 잔존을
            fail 로 검출한다(R10). 기본 False.
        required_format: 공고 요구 산출 형식(예: "hwp"). 최종 산출 확장자가 다르면
            제출명을 차단(_DRAFT)하고 변환 안내를 남긴다(ACC-5 — 판정은 이
            파이프라인 레벨에서만, run_acceptance 내부 아님). 기본 None=무검사.
        submit_clean: True 면 게이트 직전에 NotebookLM 프롬프트를 md 로 보존한 뒤
            작업용 블록을 제거한다(US-6) — '기본 산출이 항상 _DRAFT' 구조 해소.
            기본 False = 기존 동작(블록 유지, 작업용 중간본).
        write_report: True 면 통합 리포트(md/json) 생성.

    Returns:
        AutopilotReport — 전 단계 통합 결과.
    """
    settings = get_settings()
    ensure_directories(settings)
    results_root = Path(settings.results_root)
    results_root.mkdir(parents=True, exist_ok=True)

    in_path = Path(input_docx)
    if not in_path.exists():
        raise FileNotFoundError(f"입력 DOCX 가 없습니다: {input_docx}")

    stem = in_path.stem
    if output_docx:
        final_path = Path(output_docx)
    else:
        final_path = results_root / f"{stem}_autopilot.docx"
    if in_path.resolve() == final_path.resolve():
        raise ValueError("출력이 입력과 같습니다. 원본 덮어쓰기는 금지입니다.")

    report = AutopilotReport(input_docx=str(in_path), output_docx=str(final_path))

    # 재실행 보호(PIPE-2): 같은 산출명이 이미 있으면(이전 실행 결과·사용자 수정본)
    # 덮어쓰기 전에 타임스탬프 백업으로 보존한다. _DRAFT 변형 이름도 함께 본다.
    report.overwrite_backup = backup_existing_output(final_path)

    # --- 1단계: 백업 + 서식 수정 + 점수/게이트 ---
    tmp_quality = results_root / f"{stem}_ap1_quality.docx"
    orch = _make_orchestrator(results_root)
    qresult = orch.run(
        str(in_path),
        str(tmp_quality),
        emphasize=emphasize,
        underline=underline,
        remove_guides=remove_guides,
        normalize_fonts=normalize_fonts,
        write_report=False,
    )
    report.backup_dir = str(qresult.backup_dir)
    report.doc_type = f"{qresult.doc_type.type_label} ({qresult.doc_type.confidence:.0%})"
    report.score_total = qresult.score.total
    report.grade = qresult.score.grade
    report.passed = qresult.passed
    report.iterations = qresult.iterations
    o = qresult.ops
    report.ops_summary = (
        f"안내문구-{o.guide_paragraphs_removed} 글머리표-{o.bullet_spacing_fixed} "
        f"표셀-{o.table_cells_cleaned} 빈단락-{o.empty_paragraphs_removed} "
        f"강조-{o.paragraphs_emphasized}"
    )
    stage_in = Path(qresult.output_docx)

    # --- 2단계: 이미지 실제 적용(NotebookLM 슬라이드 프롬프트 삽입) ---
    tmp_img = results_root / f"{stem}_ap2_img.docx"
    img: ImageApplyReport = apply_images(
        str(stage_in), str(tmp_img),
        max_items=max_images, placeholder_only=placeholder_only,
        openai_service=_make_openai_service(),
    )
    report.prompts_inserted = img.prompts_inserted
    stage_in = tmp_img

    # --- 3단계: PSST 보강 ---
    if psst_scaffold:
        psst: PSSTFillReport = apply_psst_scaffold(str(stage_in), str(final_path))
        report.psst_overall_ratio = psst.overall_ratio
        report.psst_areas_scaffolded = psst.areas_scaffolded
        report.psst_items_added = psst.items_added
        report.psst_scaffolded_areas = list(psst.scaffolded_areas)
    else:
        # PSST 보강 생략 시 2단계 결과를 최종으로 복사
        final_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(str(stage_in), str(final_path))

    # --- 3.5단계(옵션): 제출 정리(--submit-clean, US-6) ---
    #     프롬프트를 md 로 보존(손실 0)한 뒤 작업용 블록을 제거하고 나서 게이트를
    #     통과시킨다 — 블록 삽입 때문에 '기본 산출이 항상 _DRAFT' 가 되는 구조 해소.
    if submit_clean:
        prompts = extract_notebooklm_prompts(str(final_path))
        if prompts:
            md_path = results_root / f"{stem}_슬라이드프롬프트.md"
            backup_existing_output(md_path)
            md_path.write_text("\n\n---\n\n".join(prompts), encoding="utf-8")
            report.prompt_md = str(md_path)
        tmp_clean = results_root / f"{stem}_ap3_clean.docx"
        strip_rep = strip_notebooklm_blocks(str(final_path), str(tmp_clean))
        report.strip_removed = strip_rep.paragraphs_removed
        # Path.replace 는 드라이브가 다르면 실패(WinError 17) — 복사로 대체
        shutil.copyfile(str(tmp_clean), str(final_path))

    # --- 4단계: 실사용 수용검사 게이트(R8) — fail 결함이 있으면 DRAFT 마킹 ---
    #     게이트 자신이 죽어도 통과로 취급하지 않고(fail-closed: 판정 불가 = 제출 금지),
    #     검사 예외로 리포트·백업 정보가 유실되지 않게 보호한다(submit 게이트와 동일 정책).
    if acceptance_gate:
        acc = None
        try:
            acc = run_acceptance(str(final_path), AcceptanceConfig(
                blind_review=blind_review, max_pages=max_pages, ai_section_max=ai_section_max))
        except Exception as exc:
            report.acceptance_error = f"{type(exc).__name__}: {exc}"
        if acc is not None:
            report.acceptance_submittable = acc.submittable
            report.acceptance_verdict = "제출가능" if acc.submittable else "제출불가(DRAFT)"
            report.acceptance_fail_defects = acc.fail_defects
            report.acceptance_failed_checks = [
                f"{r.label}: {r.detail}"
                for r in acc.results if r.severity == SEV_FAIL and not r.passed
            ]
        if acc is None or not acc.submittable:
            # '제출본' 이름으로 내보내지 않는다. 출력의 _DRAFT 이름이 입력 원본과
            # 겹치면 _DRAFT2 로 대체 마킹하고(원본 보존·침묵 스킵 금지),
            # rename 실패(파일 잠금 등)도 기록해 사용자에게 알린다.
            new_path, mark_error = force_draft_name(final_path, avoid=in_path)
            if mark_error:
                report.draft_mark_error = mark_error
            else:
                final_path = new_path
                report.output_docx = str(final_path)
                report.draft_marked = True

    # --- 4.5단계: 산출 형식 게이트(ACC-5) — 요구 형식(hwp 등)과 다르면 제출명 차단 ---
    #     판정은 파이프라인 레벨(최종 파일 확장자)에서만 한다 — run_acceptance 내부에
    #     넣으면 변환 전 DOCX 가 영구 fail 이 된다(설계 결정, ralplan v2).
    if required_format and final_path.suffix.lstrip(".").lower() != required_format.lstrip(".").lower():
        report.format_mismatch = (
            f"요구 산출형식 .{required_format.lstrip('.')} ↔ 실제 {final_path.suffix} — "
            f"scripts\\docx2hwp.py(대화형 PowerShell)로 변환 후 제출"
        )
        if not final_path.stem.endswith(("_DRAFT", "_DRAFT2")):
            new_path, mark_error = force_draft_name(final_path, avoid=in_path)
            if mark_error:
                report.draft_mark_error = report.draft_mark_error or mark_error
            else:
                final_path = new_path
                report.output_docx = str(final_path)
                report.draft_marked = True

    # --- 5단계: 잔존 빈칸 스캔 + To-Do ---
    report.residual_placeholders = _scan_residual(str(final_path))
    report.manual_todo = _build_todo(report)

    if write_report:
        report.report_md = _write_report(results_root, stem, report)

    return report


def _build_todo(report: AutopilotReport) -> list[str]:
    todo: list[str] = []
    if not report.passed:
        todo.append(
            f"품질점수 {report.score_total:.1f}점(게이트 미달) — 보완 후 재실행 권장."
        )
    if report.acceptance_verdict and not report.acceptance_submittable:
        todo.append(
            f"수용검사 {report.acceptance_verdict} (fail {report.acceptance_fail_defects}건) — "
            f"아래 결함을 해결해야 제출 가능합니다."
        )
        todo.extend(f"  · {c}" for c in report.acceptance_failed_checks)
    if report.acceptance_error:
        todo.append(
            f"수용검사 실행 실패({report.acceptance_error}) — 판정 불가. "
            f"self_diagnose 로 수동 진단 전 제출 금지(_DRAFT 표시)."
        )
    if report.draft_mark_error:
        todo.append(
            f"_DRAFT 마킹 실패({report.draft_mark_error}) — 출력 파일명이 제출 이름 "
            f"그대로이니 직접 변경 전 제출 금지."
        )
    if report.format_mismatch:
        todo.append(f"산출 형식 불일치 — {report.format_mismatch}")
    if report.prompts_inserted:
        todo.append(
            f"NotebookLM 슬라이드 프롬프트 {report.prompts_inserted}곳 — "
            f"각 프롬프트를 NotebookLM 에 붙여넣어 슬라이드를 만들고, 안내 블록은 삭제하세요."
        )
    for area in report.psst_scaffolded_areas:
        todo.append(f"PSST 작성 보강: {area}")
    if report.residual_placeholders:
        todo.append(
            f"잔존 빈칸 {len(report.residual_placeholders)}곳 — 직접 채우거나 submittable_filler 로 채움."
        )
    return todo


def _write_report(results_root: Path, stem: str, report: AutopilotReport) -> str:
    md_path = results_root / f"{stem}_autopilot_report.md"
    lines: list[str] = []
    lines.append(f"# 오토파일럿 최종 리포트 — {stem}")
    lines.append("")
    lines.append(f"- 입력: `{report.input_docx}`")
    lines.append(f"- 출력: `{report.output_docx}`")
    lines.append(f"- 백업: `{report.backup_dir}`")
    lines.append("")
    lines.append("## 1) 서식 수정 + 품질점수")
    lines.append(f"- 문서 유형: {report.doc_type}")
    gate = "통과" if report.passed else "미달"
    lines.append(
        f"- 품질점수: **{report.score_total:.1f}/100** ({report.grade}) | 게이트 {gate} "
        f"(반복 {report.iterations}회)"
    )
    lines.append(f"- 후처리: {report.ops_summary}")
    lines.append("")
    lines.append("## 2) 이미지 적용 (NotebookLM 슬라이드 프롬프트)")
    lines.append(f"- 슬라이드 프롬프트 삽입: {report.prompts_inserted}건")
    lines.append("- 각 프롬프트를 NotebookLM 슬라이드 생성에 붙여넣어 사용하세요.")
    lines.append("")
    lines.append("## 3) PSST 보강")
    lines.append(f"- 전체 충족률: {report.psst_overall_ratio*100:.0f}%")
    lines.append(
        f"- 보강 영역: {report.psst_areas_scaffolded}개 / 추가 항목: {report.psst_items_added}개"
    )
    if report.psst_scaffolded_areas:
        lines.append(f"- 대상: {', '.join(report.psst_scaffolded_areas)}")
    lines.append("")
    lines.append("## 4) 실사용 수용검사 (제출 가능성 게이트)")
    if report.acceptance_verdict:
        lines.append(
            f"- 판정: **{report.acceptance_verdict}** (fail 결함 {report.acceptance_fail_defects}건)"
        )
        for c in report.acceptance_failed_checks:
            lines.append(f"  - {c}")
        if report.draft_marked:
            lines.append("- fail 결함이 있어 출력 파일명에 `_DRAFT` 를 붙였습니다 — 결함 해결 전에는 제출하지 마세요.")
    else:
        lines.append("- (게이트 생략됨 — acceptance_gate=False)")
    lines.append("")
    lines.append("## 5) 수동 보완 To-Do")
    if report.manual_todo:
        for t in report.manual_todo:
            lines.append(f"- [ ] {t}")
    else:
        lines.append("- 없음 (제출 준비 완료)")
    lines.append("")
    md_path.write_text("\n".join(lines), encoding="utf-8")
    return str(md_path)
