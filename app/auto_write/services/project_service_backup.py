from __future__ import annotations

import json
import logging
import os
import re
import zipfile
from pathlib import Path
from typing import Any

from docx import Document

from ..analysis.docx_template import analyze_template, build_doc_summary, sanitize_template_profile
from ..models import (
    ArtifactBundle,
    EvidenceRequest,
    EvidenceSource,
    GeneratedImage,
    ImageRequest,
    ProjectInput,
    ReferenceFile,
    TemplateProfile,
)
from ..storage import Storage
from ..utils import log_line, read_json, sanitize_user_filename, unique_lines, write_json
from .evidence_service import EvidenceService
from .image_service import ImageService
from .openai_client import OpenAIService
from .qa_service import QAService
from .render_service import RenderService


class ProjectService:
    GUIDE_MARKER_RE = re.compile(r"(※|<[^>]+>|기재|작성요령|작성방법|예시|OOO|○○○)")
    INVALID_XML_CHAR_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")
    LIBRARY_SENTENCE_SPLIT_RE = re.compile(r"(?<=[\.\!\?])\s+")
    LIBRARY_BLOCK_SPLIT_RE = re.compile(r"(?:\r?\n){2,}")
    LIBRARY_GUIDE_RE = re.compile(r"(※|<[^>]+>|기재|작성요령|작성방법|예시|샘플|양식|붙임)")
    LIBRARY_TOKEN_RE = re.compile(r"[가-힣A-Za-z0-9]{2,}")
    LIBRARY_KEYWORD_HINTS = ("문제", "해결", "목표", "시장", "고객", "전략", "실행", "성과", "매출", "비용", "로드맵")
    REFERENCE_SNIPPET_LIMIT = 160
    TABLE_AUTOFILL_EXCLUDE_RE = re.compile(
        r"(열\s*\d+|행\s*\d+|협약해지|제재|환수|유의사항|집행기준|증빙|계좌이체|서류|동의|개인정보)",
        re.IGNORECASE,
    )
    TABLE_NARROW_HINT_RE = re.compile(r"(월|주차|진행률|금액|비중|건수|수량|단가|열\s*\d+|행\s*\d+)")

    def __init__(
        self,
        storage: Storage,
        openai_service: OpenAIService,
        evidence_service: EvidenceService,
        image_service: ImageService,
        render_service: RenderService,
        qa_service: QAService,
    ):
        self.storage = storage
        self.openai_service = openai_service
        self.evidence_service = evidence_service
        self.image_service = image_service
        self.render_service = render_service
        self.qa_service = qa_service
        self._library_snippets: list[dict[str, str]] | None = None
        logging.getLogger("pypdf").setLevel(logging.ERROR)

    def analyze_uploaded_template(self, file_name: str, content: bytes) -> TemplateProfile:
        template_id, output_path = self.storage.create_template_space(file_name)
        output_path.write_bytes(content)
        profile = analyze_template(output_path)
        profile.template_id = template_id
        profile.source_docx = str(output_path)
        if self.openai_service.available:
            refined = self.openai_service.refine_template_profile(profile.model_dump(), build_doc_summary(output_path))
            if refined:
                try:
                    merged = TemplateProfile.model_validate(refined)
                    merged.template_id = template_id
                    merged.source_docx = str(output_path)
                    merged.template_name = file_name
                    profile = merged
                except Exception:
                    pass
        profile = sanitize_template_profile(profile)
        self.storage.save_template_profile(profile)
        return profile

    def finalize_template(self, template_id: str, raw_json: str) -> TemplateProfile:
        try:
            data = json.loads(raw_json)
        except json.JSONDecodeError as exc:
            raise ValueError("템플릿 설정 JSON 형식이 올바르지 않습니다.") from exc
        profile = TemplateProfile.model_validate(data)
        profile.template_id = template_id
        folder = self.storage.template_dir(template_id)
        docx_files = sorted(folder.glob("*.docx"))
        if not docx_files:
            raise ValueError("원본 DOCX 파일을 찾지 못했습니다. 템플릿을 다시 업로드해 주세요.")
        source_docx = docx_files[0]
        profile.source_docx = str(source_docx)
        profile = sanitize_template_profile(profile)
        self.storage.save_template_profile(profile)
        return profile

    def create_project(self, template_id: str, project_name: str) -> str:
        project_id, _ = self.storage.create_project_space(template_id, project_name or "새 프로젝트")
        profile = self.storage.load_template_profile(template_id)
        blank = ProjectInput(template_id=template_id)
        self.storage.save_project_input(project_id, blank)
        write_json(self.storage.project_dir(project_id) / "template_snapshot.json", profile.model_dump())
        return project_id

    def save_project_form(
        self,
        project_id: str,
        answers: dict[str, Any],
        project_title: str,
        organization_name: str,
        evidence_topics: str,
        reference_files: list[tuple[str, bytes]],
    ) -> ProjectInput:
        profile = self.load_profile_for_project(project_id)
        references: list[ReferenceFile] = []
        for file_name, content in reference_files:
            if not file_name:
                continue
            safe_name = sanitize_user_filename(file_name)
            path = self._next_reference_path(project_id, safe_name)
            path.write_bytes(content)
            preview = self.extract_reference_text(path)
            preview_path = path.with_suffix(path.suffix + ".txt")
            preview_path.write_text(preview, encoding="utf-8")
            references.append(
                ReferenceFile(
                    file_name=path.name,
                    saved_path=str(path),
                    extracted_text_path=str(preview_path),
                    extracted_preview=preview[:400],
                )
            )

        evidence_requests = [EvidenceRequest(topic=line) for line in unique_lines(evidence_topics)]
        filtered_answers = {key: value for key, value in answers.items() if str(value).strip()}
        image_requests = []
        for slot in profile.image_slots:
            note_key = f"image_note_{slot.slot_id}"
            image_requests.append(
                ImageRequest(
                    slot_id=slot.slot_id,
                    label=slot.label,
                    prompt=str(filtered_answers.get(note_key, "")),
                    source="generated",
                )
            )
        project_input = ProjectInput(
            template_id=profile.template_id,
            project_meta={"project_title": project_title},
            organization_profile={"name": organization_name},
            answers=filtered_answers,
            references=references,
            evidence_requests=evidence_requests,
            image_requests=image_requests,
        )
        self.storage.save_project_input(project_id, project_input)
        return project_input

    def generate(self, project_id: str) -> ArtifactBundle:
        profile = self.load_profile_for_project(project_id)
        project_input = self.storage.load_project_input(project_id)
        context = self._build_context(project_input)
        transfer_mode = len(project_input.references) >= 2
        template_completed = self._template_looks_completed(profile)
        missing = self._collect_missing_questions(profile, project_input.answers, required_only=False)
        missing = self._filter_missing_for_autofill(profile, missing, transfer_mode)
        if missing and not template_completed:
            reference_hints = self._suggest_reference_snippets(missing, project_input)
            # When project references are uploaded, prioritize them and avoid unrelated library contamination.
            has_library_context = False if project_input.references else bool(self._ensure_reference_library_loaded())
            has_user_context = (
                bool(project_input.references)
                or len(project_input.answers) >= 2
                or bool(reference_hints)
                or has_library_context
            )
            library_hints = (
                self._suggest_library_snippets(missing)
                if has_user_context and not project_input.references
                else {}
            )
            combined_hints = self._merge_hint_maps(reference_hints, library_hints)
            full_context = self._merge_context_with_library(context, combined_hints)
            drafted = self._draft_missing_answers_in_chunks(missing, full_context) if has_user_context else {}
            if drafted:
                project_input.answers.update({k: v for k, v in drafted.items() if str(v).strip()})
            remaining = self._collect_missing_questions(profile, project_input.answers, required_only=False)
            remaining = self._filter_missing_for_autofill(profile, remaining, transfer_mode)
            if remaining:
                fallback = self._build_fallback_answers(remaining, project_input, combined_hints)
                project_input.answers.update(fallback)
            project_input.answers = self._postprocess_answers(profile, project_input.answers)
            project_input.answers = {
                key: self._sanitize_xml_text(value) if isinstance(value, str) else value
                for key, value in project_input.answers.items()
            }
            self.storage.save_project_input(project_id, project_input)

        evidence = self.evidence_service.search(project_input.evidence_requests)
        images = self.image_service.build_images(
            profile.image_slots,
            project_input.answers,
            evidence,
            self.storage.project_dir(project_id) / "generated_assets",
        )
        output_path = self.storage.project_dir(project_id) / "output" / "output.docx"
        render_result = self.render_service.render(profile, project_input, images, output_path)
        qa_report = self.qa_service.build_report(profile, project_input, render_result, images, evidence)

        sources_path = self.storage.project_dir(project_id) / "output" / "sources.json"
        qa_path = self.storage.project_dir(project_id) / "output" / "qa_report.json"
        benchmark_path = self.storage.project_dir(project_id) / "output" / "benchmark_compare.json"
        transfer_path = self.storage.project_dir(project_id) / "output" / "transfer_report.json"
        write_json(sources_path, [source.model_dump() for source in evidence])
        write_json(qa_path, qa_report)
        write_json(benchmark_path, self._build_benchmark_compare(profile, output_path, qa_report))
        write_json(
            transfer_path,
            self._build_transfer_report(
                profile=profile,
                project_input=project_input,
                output_path=output_path,
                render_result=render_result,
                qa_report=qa_report,
                evidence=evidence,
                images=images,
                transfer_mode=transfer_mode,
            ),
        )
        log_line(f"[DONE] project={project_id} docx={output_path.name} errors={qa_report['error_count']}")
        return ArtifactBundle(
            output_docx=str(output_path),
            qa_report=str(qa_path),
            sources=str(sources_path),
            benchmark_compare=str(benchmark_path),
            transfer_report=str(transfer_path),
            generated_assets=[image.path for image in images],
        )

    def _template_looks_completed(self, profile: TemplateProfile) -> bool:
        source_docx = Path(profile.source_docx)
        if not source_docx.exists():
            return False
        try:
            metrics = self._doc_metrics(source_docx)
        except Exception:
            return False
        if metrics["paragraph_count"] >= 40 and metrics["paragraph_chars"] >= 1800 and metrics["table_fill_ratio"] >= 0.9:
            return True
        return False

    def load_profile_for_project(self, project_id: str) -> TemplateProfile:
        snapshot = self.storage.project_dir(project_id) / "template_snapshot.json"
        if snapshot.exists():
            profile = TemplateProfile.model_validate(read_json(snapshot))
            return sanitize_template_profile(profile)
        meta = read_json(self.storage.project_dir(project_id) / "project_meta.json")
        return sanitize_template_profile(self.storage.load_template_profile(meta["template_id"]))

    def normalize_profile(self, profile: TemplateProfile) -> TemplateProfile:
        return sanitize_template_profile(profile)

    def visible_questions(self, profile: TemplateProfile) -> list[dict[str, Any]]:
        visible: list[dict[str, Any]] = []
        for question in profile.questions:
            if question.question_id in {"project_title", "organization_name"}:
                continue
            if question.required:
                visible.append(question.model_dump())
                continue
            target_kind = str(question.target.get("kind", ""))
            if target_kind == "image_slot_note":
                continue
            if target_kind == "section" and len(visible) < 12:
                visible.append(question.model_dump())
        return visible

    def extract_reference_text(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"}:
            return path.read_text(encoding="utf-8", errors="ignore")
        if suffix == ".json":
            return json.dumps(read_json(path), ensure_ascii=False, indent=2)
        if suffix == ".docx":
            document = Document(str(path))
            return "\n".join(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
        if suffix == ".pdf":
            try:
                from pypdf import PdfReader

                reader = PdfReader(str(path))
                pages: list[str] = []
                for page in reader.pages:
                    text = (page.extract_text() or "").strip()
                    if text:
                        pages.append(text)
                joined = self._sanitize_xml_text("\n".join(pages).strip())
                return joined or f"PDF에서 텍스트를 추출하지 못했습니다: {path.name}"
            except Exception:
                return f"PDF 텍스트 추출에 실패했습니다: {path.name}"
        return f"지원하지 않는 참고자료 형식입니다: {path.name}"

    def _ensure_reference_library_loaded(self) -> list[dict[str, str]]:
        if self._library_snippets is not None:
            return self._library_snippets
        self._library_snippets = []
        settings = self.storage.settings
        library_dir = settings.reference_library_dir
        if library_dir is None or not library_dir.exists() or not library_dir.is_dir():
            return self._library_snippets
        files = sorted(
            [path for path in library_dir.iterdir() if path.is_file() and path.suffix.lower() in {".pdf", ".docx", ".txt", ".md"}]
        )[:200]
        for file_path in files:
            text = self.extract_reference_text(file_path)
            if not text or "지원하지 않는 참고자료 형식" in text or "텍스트 추출" in text:
                continue
            for snippet in self._split_library_snippets(text):
                self._library_snippets.append({"source": file_path.name, "text": snippet})
        log_line(f"[INFO] reference_library snippets={len(self._library_snippets)} dir={library_dir}")
        return self._library_snippets

    def _split_library_snippets(self, text: str) -> list[str]:
        snippets: list[str] = []
        blocks = self.LIBRARY_BLOCK_SPLIT_RE.split(text)
        for block in blocks:
            clean = self._sanitize_xml_text(re.sub(r"\s+", " ", block).strip())
            if len(clean) < 40:
                continue
            if len(clean) > 260:
                sentences = self.LIBRARY_SENTENCE_SPLIT_RE.split(clean)
                merged: list[str] = []
                current = ""
                for sentence in sentences:
                    sentence = sentence.strip()
                    if not sentence:
                        continue
                    trial = sentence if not current else f"{current} {sentence}"
                    if len(trial) <= 240:
                        current = trial
                    else:
                        if len(current) >= 40:
                            merged.append(current)
                        current = sentence
                if len(current) >= 40:
                    merged.append(current)
                candidates = merged
            else:
                candidates = [clean]
            for candidate in candidates:
                candidate = self._sanitize_xml_text(candidate)
                if self.LIBRARY_GUIDE_RE.search(candidate):
                    continue
                if "<" in candidate and ">" in candidate:
                    continue
                if len(candidate) > 320:
                    continue
                snippets.append(candidate)
        return snippets

    def _tokenize(self, text: str) -> set[str]:
        return {token.lower() for token in self.LIBRARY_TOKEN_RE.findall(text or "")}

    def _score_library_snippet(self, snippet: str, query_tokens: set[str]) -> float:
        if not snippet:
            return 0.0
        snippet_tokens = self._tokenize(snippet)
        if not snippet_tokens:
            return 0.0
        overlap = len(snippet_tokens & query_tokens)
        keyword_bonus = sum(1 for keyword in self.LIBRARY_KEYWORD_HINTS if keyword in snippet)
        length = len(snippet)
        if length < 50:
            length_penalty = -1.5
        elif length > 280:
            length_penalty = -1.0
        else:
            length_penalty = 0.0
        return overlap * 2.0 + keyword_bonus * 0.5 + length_penalty

    def _suggest_snippets_from_corpus(
        self,
        missing: list[dict[str, Any]],
        corpus: list[dict[str, str]],
        allowed_target_kinds: set[str],
        min_score: float,
    ) -> dict[str, str]:
        if not corpus:
            return {}
        suggestions: dict[str, str] = {}
        used_indexes: set[int] = set()
        for question in missing:
            question_id = str(question.get("question_id", "")).strip()
            if not question_id:
                continue
            target_kind = str(question.get("target", {}).get("kind", ""))
            if target_kind not in allowed_target_kinds:
                continue
            label = str(question.get("label", "")).strip()
            if not label:
                continue
            query_tokens = self._tokenize(label)
            query_tokens.update(self._tokenize(str(question.get("source_hint", ""))))
            if not query_tokens:
                continue
            ranked = sorted(
                enumerate(corpus),
                key=lambda pair: self._score_library_snippet(pair[1]["text"], query_tokens),
                reverse=True,
            )
            selected = ""
            for index, item in ranked:
                score = self._score_library_snippet(item["text"], query_tokens)
                if score < min_score:
                    break
                if index in used_indexes:
                    continue
                selected = item["text"]
                used_indexes.add(index)
                break
            if selected:
                suggestions[question_id] = selected
        return suggestions

    def _suggest_library_snippets(self, missing: list[dict[str, Any]]) -> dict[str, str]:
        corpus = self._ensure_reference_library_loaded()
        return self._suggest_snippets_from_corpus(
            missing=missing,
            corpus=corpus,
            allowed_target_kinds={"section"},
            min_score=3.0,
        )

    def _read_reference_text(self, reference: ReferenceFile) -> str:
        extracted_path = Path(reference.extracted_text_path)
        if extracted_path.exists():
            try:
                text = extracted_path.read_text(encoding="utf-8", errors="ignore")
                cleaned = self._sanitize_xml_text(text)
                if cleaned:
                    return cleaned
            except Exception:
                pass
        return self._sanitize_xml_text(reference.extracted_preview)

    def _ordered_references_for_text(self, references: list[ReferenceFile]) -> list[ReferenceFile]:
        if not references:
            return []
        grouped: dict[str, list[ReferenceFile]] = {}
        for reference in references:
            stem = Path(reference.file_name).stem.lower().strip()
            grouped.setdefault(stem, []).append(reference)
        ordered: list[ReferenceFile] = []
        for stem in sorted(grouped.keys()):
            group = grouped[stem]
            group = sorted(
                group,
                key=lambda ref: (
                    0 if Path(ref.file_name).suffix.lower() == ".docx" else 1,
                    0 if Path(ref.file_name).suffix.lower() == ".txt" else 1,
                    ref.file_name.lower(),
                ),
            )
            ordered.append(group[0])
        return ordered

    def _build_reference_snippet_corpus(self, project_input: ProjectInput) -> list[dict[str, str]]:
        corpus: list[dict[str, str]] = []
        for reference in self._ordered_references_for_text(project_input.references):
            text = self._read_reference_text(reference)
            if not text:
                continue
            if "지원하지 않는 참고자료 형식" in text or "텍스트 추출" in text:
                continue
            snippets = self._split_library_snippets(text)
            for snippet in snippets:
                corpus.append({"source": reference.file_name, "text": snippet})
                if len(corpus) >= self.REFERENCE_SNIPPET_LIMIT:
                    return corpus
        return corpus

    def _suggest_reference_snippets(self, missing: list[dict[str, Any]], project_input: ProjectInput) -> dict[str, str]:
        corpus = self._build_reference_snippet_corpus(project_input)
        return self._suggest_snippets_from_corpus(
            missing=missing,
            corpus=corpus,
            allowed_target_kinds={"section", "table_cell"},
            min_score=2.5,
        )

    @staticmethod
    def _merge_hint_maps(primary: dict[str, str], secondary: dict[str, str]) -> dict[str, str]:
        merged = dict(primary)
        for question_id, snippet in secondary.items():
            if question_id not in merged:
                merged[question_id] = snippet
        return merged

    def _merge_context_with_library(self, context: str, library_hints: dict[str, str]) -> str:
        if not library_hints:
            return context
        lines = [context.strip()] if context.strip() else []
        lines.append("참고 우수문장:")
        for index, snippet in enumerate(library_hints.values(), start=1):
            lines.append(f"{index}. {snippet[:280]}")
            if index >= 12:
                break
        return "\n".join(line for line in lines if line)

    def _adapt_library_snippet(self, snippet: str, project_title: str, org_name: str) -> str:
        text = self._sanitize_xml_text(re.sub(r"\s+", " ", snippet).strip())
        if len(text) > 260:
            text = text[:260].rstrip(" ,.;") + "."
        return self._sanitize_xml_text(text)

    def _sanitize_xml_text(self, text: str) -> str:
        cleaned = self.INVALID_XML_CHAR_RE.sub(" ", text or "")
        cleaned = cleaned.replace("\uFFFD", " ")
        cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        return cleaned.strip()

    def _build_context(self, project_input: ProjectInput) -> str:
        parts = [
            f"과제명: {project_input.project_meta.get('project_title', '')}",
            f"기관명: {project_input.organization_profile.get('name', '')}",
        ]
        for key, value in project_input.answers.items():
            parts.append(f"{key}: {value}")
        for reference in self._ordered_references_for_text(project_input.references):
            reference_text = self._read_reference_text(reference)
            if reference_text:
                parts.append(f"[참고자료:{reference.file_name}] {reference_text[:2400]}")
        return "\n".join(part for part in parts if part).strip()

    def _next_reference_path(self, project_id: str, safe_name: str) -> Path:
        reference_dir = self.storage.project_dir(project_id) / "references"
        candidate = reference_dir / safe_name
        if not candidate.exists():
            return candidate
        stem = candidate.stem
        suffix = candidate.suffix
        index = 2
        while True:
            next_candidate = reference_dir / f"{stem}_{index}{suffix}"
            if not next_candidate.exists():
                return next_candidate
            index += 1

    def _collect_missing_questions(
        self, profile: TemplateProfile, answers: dict[str, Any], required_only: bool = False
    ) -> list[dict[str, Any]]:
        missing: list[dict[str, Any]] = []
        excluded_section_ids = {section.field_id for section in profile.sections if getattr(section, "is_excluded", False)}
        for question in profile.questions:
            target_kind = str(question.target.get("kind", ""))
            if target_kind not in {"section", "table_cell"}:
                continue
            if target_kind == "section":
                field_id = str(question.target.get("field_id", "")).strip()
                if field_id and field_id in excluded_section_ids:
                    continue
            if required_only and not bool(question.required):
                continue
            value = str(answers.get(question.question_id, "") or "").strip()
            if value:
                continue
            missing.append(question.model_dump())
        return missing

    def _table_cell_index(self, profile: TemplateProfile) -> tuple[dict[str, dict[str, Any]], dict[str, int]]:
        by_cell_id: dict[str, dict[str, Any]] = {}
        table_sizes: dict[str, int] = {}
        for table in profile.tables:
            table_sizes[table.table_id] = len(table.cells)
            for cell in table.cells:
                by_cell_id[cell.cell_id] = {
                    "table_id": table.table_id,
                    "label": cell.label,
                    "row_header": cell.row_header,
                    "col_header": cell.col_header,
                }
        return by_cell_id, table_sizes

    @staticmethod
    def _clip_text(value: str, max_chars: int) -> str:
        text = re.sub(r"\s+", " ", str(value or "").strip())
        if len(text) <= max_chars:
            return text
        clipped = text[:max_chars].rstrip(" ,.;:")
        return f"{clipped}."

    def _table_cell_max_chars(self, label: str) -> int:
        if self.TABLE_NARROW_HINT_RE.search(label or ""):
            return 36
        return 90

    def _filter_missing_for_autofill(
        self,
        profile: TemplateProfile,
        missing: list[dict[str, Any]],
        transfer_mode: bool,
    ) -> list[dict[str, Any]]:
        by_cell_id, table_sizes = self._table_cell_index(profile)
        filtered: list[dict[str, Any]] = []
        for question in missing:
            target_kind = str(question.get("target", {}).get("kind", ""))
            if target_kind != "table_cell":
                filtered.append(question)
                continue
            cell_id = str(question.get("target", {}).get("cell_id", ""))
            table_id = str(question.get("target", {}).get("table_id", ""))
            label = str(question.get("label", "")).strip()
            meta = by_cell_id.get(cell_id, {})
            if not label:
                continue
            if self.TABLE_AUTOFILL_EXCLUDE_RE.search(label):
                continue
            # Large matrix tables are the main source of low-quality overfill.
            if transfer_mode and table_sizes.get(table_id, 0) > 20 and not bool(question.get("required")):
                continue
            # Skip cells when either row/column semantics are missing in wide tables.
            row_header = str(meta.get("row_header", "")).strip()
            col_header = str(meta.get("col_header", "")).strip()
            if transfer_mode and table_sizes.get(table_id, 0) > 10 and (not row_header or not col_header):
                continue
            filtered.append(question)
        return filtered

    def _postprocess_answers(self, profile: TemplateProfile, answers: dict[str, Any]) -> dict[str, Any]:
        normalized: dict[str, Any] = dict(answers)
        by_cell_id, _ = self._table_cell_index(profile)
        for cell_id, meta in by_cell_id.items():
            if cell_id not in normalized:
                continue
            value = normalized.get(cell_id)
            if not isinstance(value, str):
                continue
            max_chars = self._table_cell_max_chars(str(meta.get("label", "")))
            normalized[cell_id] = self._clip_text(value, max_chars)
        return normalized

    def _draft_missing_answers_in_chunks(self, missing: list[dict[str, Any]], context: str) -> dict[str, str]:
        if not context.strip():
            return {}
        drafted: dict[str, str] = {}
        chunk_size = 40
        for start in range(0, len(missing), chunk_size):
            chunk = missing[start : start + chunk_size]
            result = self.openai_service.draft_missing_answers(chunk, context)
            if not result:
                continue
            for key, value in result.items():
                value_text = str(value or "").strip()
                if value_text:
                    drafted[key] = value_text
        return drafted

    def _build_fallback_answers(
        self,
        missing: list[dict[str, Any]],
        project_input: ProjectInput,
        library_hints: dict[str, str] | None = None,
    ) -> dict[str, str]:
        project_title = str(project_input.project_meta.get("project_title", "") or "").strip() or "본 과제"
        org_name = str(project_input.organization_profile.get("name", "") or "").strip() or "신청 기관"
        fallback: dict[str, str] = {}
        library_hints = library_hints or {}
        used_snippets: set[str] = set()
        for question in missing:
            question_id = str(question.get("question_id", "")).strip()
            if not question_id:
                continue
            target_kind = str(question.get("target", {}).get("kind", ""))
            if target_kind == "section":
                label = str(question.get("label", "")).strip()
                best_snippet = library_hints.get(question_id, "")
                if best_snippet and best_snippet not in used_snippets:
                    fallback[question_id] = self._adapt_library_snippet(best_snippet, project_title, org_name)
                    used_snippets.add(best_snippet)
                else:
                    fallback[question_id] = self._fallback_section_text(project_title, org_name, label)
            elif target_kind == "table_cell":
                label = str(question.get("label", "")).strip()
                best_snippet = library_hints.get(question_id, "")
                if best_snippet and best_snippet not in used_snippets:
                    adapted = self._adapt_library_snippet(best_snippet, project_title, org_name)
                    fallback[question_id] = self._clip_text(adapted, self._table_cell_max_chars(label))
                    used_snippets.add(best_snippet)
                else:
                    fallback[question_id] = self._fallback_table_text(project_title, label)
        return fallback

    def _fallback_section_text(self, project_title: str, org_name: str, label: str) -> str:
        if any(token in label for token in ("목표", "성과", "KPI", "지표")):
            return f"{project_title}의 핵심 성과지표를 설정하고 월 단위로 달성률을 점검합니다."
        if any(token in label for token in ("시장", "경쟁", "고객")):
            return f"{org_name}는 목표 고객군과 경쟁 환경을 분석하여 차별화 전략을 실행합니다."
        if any(token in label for token in ("예산", "비용", "수익")):
            return "필수 비용과 기대 수익을 구분해 단계별로 집행하고 분기별로 재검토합니다."
        if any(token in label for token in ("추진", "일정", "계획")):
            return "준비-실행-확산 3단계 일정으로 운영하며 단계 종료 시 점검 결과를 반영합니다."
        return f"{project_title} 관련 '{label}' 항목은 참고자료를 기반으로 구체 실행안 중심으로 작성합니다."

    def _fallback_table_text(self, project_title: str, label: str) -> str:
        if any(token in label for token in ("금액", "예산", "비용", "매출", "원", "만원")):
            return "1,000(추정)"
        if any(token in label for token in ("일정", "기간", "시기", "월")):
            return "1단계 준비 / 2단계 실행 / 3단계 확산(월 단위 점검)"
        if any(token in label for token in ("목표", "성과", "지표")):
            return "정량 KPI 설정 및 월별 점검(달성률 관리)"
        if any(token in label for token in ("담당", "인력", "역할")):
            return "전담 인력 지정 및 역할 분담, 책임자 명시"
        return f"{project_title} 기준 실행 계획 수치 및 근거 반영"

    def _build_transfer_report(
        self,
        profile: TemplateProfile,
        project_input: ProjectInput,
        output_path: Path,
        render_result: dict[str, Any],
        qa_report: dict[str, Any],
        evidence: list[EvidenceSource],
        images: list[GeneratedImage],
        transfer_mode: bool,
    ) -> dict[str, Any]:
        generated_metrics = self._doc_metrics(output_path) if output_path.exists() else {
            "paragraph_count": 0,
            "paragraph_chars": 0,
            "table_cells_total": 0,
            "table_fill_ratio": 0.0,
            "table_chars": 0,
            "drawing_count": 0,
            "guide_marker_hits": 0,
        }
        total_sections = len(profile.sections)
        total_table_cells = sum(len(table.cells) for table in profile.tables)
        sections_written = int(render_result.get("sections_written", 0))
        cells_written = int(render_result.get("cells_written", 0))
        image_slots_total = len(profile.image_slots)
        required_slots = [slot for slot in profile.image_slots if slot.required]
        image_slot_ids = {image.slot_id for image in images}
        required_filled = sum(1 for slot in required_slots if slot.slot_id in image_slot_ids)
        all_slots_filled = sum(1 for slot in profile.image_slots if slot.slot_id in image_slot_ids)
        matched_topics = sum(1 for req in project_input.evidence_requests if any(src.topic == req.topic for src in evidence))

        unresolved_from_qa: list[str] = []
        unresolved_from_qa.extend(str(item) for item in qa_report.get("errors", [])[:20])
        unresolved_from_qa.extend(str(item) for item in qa_report.get("warnings", [])[:20])
        unresolved_from_render = [str(item) for item in render_result.get("errors", [])[:20]]

        return {
            "status": "ok",
            "mode": "transfer" if transfer_mode else "standard",
            "output_docx": str(output_path),
            "fill_ratio": {
                "sections": round((sections_written / total_sections), 4) if total_sections else 1.0,
                "table_cells": round((cells_written / total_table_cells), 4) if total_table_cells else 1.0,
                "table_cells_document": float(generated_metrics["table_fill_ratio"]),
            },
            "image_slot_coverage": {
                "required_filled": required_filled,
                "required_total": len(required_slots),
                "required_ratio": round((required_filled / len(required_slots)), 4) if required_slots else 1.0,
                "all_filled": all_slots_filled,
                "all_total": image_slots_total,
                "all_ratio": round((all_slots_filled / image_slots_total), 4) if image_slots_total else 1.0,
            },
            "evidence_usage": {
                "requested_topics": len(project_input.evidence_requests),
                "matched_topics": matched_topics,
                "sources_collected": len(evidence),
                "reference_files": len(project_input.references),
            },
            "qa_summary": {
                "passed": bool(qa_report.get("passed")),
                "error_count": int(qa_report.get("error_count", 0)),
                "warning_count": int(qa_report.get("warning_count", 0)),
                "guide_marker_hits": int(generated_metrics.get("guide_marker_hits", 0)),
            },
            "unresolved_items": {
                "qa": unresolved_from_qa,
                "render": unresolved_from_render,
            },
        }

    def _resolve_benchmark_docx(self, profile: TemplateProfile) -> Path | None:
        forced = os.getenv("AUTO_WRITE_BENCHMARK_DOCX", "").strip()
        if forced:
            forced_path = Path(forced)
            if forced_path.exists():
                return forced_path
        fixed_benchmark = Path(
            r"C:\Users\ekth3\OneDrive\바탕 화면\다솜\경영지도사 개인\02. 밸류업파트너스\2026 AI수출지원공통\예창\마켓게이트_예비창업패키지 사업계획서 1550.docx"
        )
        if "예비창업패키지" in profile.template_name and fixed_benchmark.exists():
            return fixed_benchmark
        source_docx = Path(profile.source_docx)
        if source_docx.exists():
            return source_docx
        return None

    def _doc_metrics(self, docx_path: Path) -> dict[str, Any]:
        document = Document(str(docx_path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        paragraph_chars = sum(len(text) for text in paragraphs)

        table_cells_total = 0
        table_cells_filled = 0
        table_chars = 0
        guide_marker_hits = 0
        for text in paragraphs:
            if self.GUIDE_MARKER_RE.search(text):
                guide_marker_hits += 1

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_cells_total += 1
                    text = cell.text.strip()
                    if text:
                        table_cells_filled += 1
                        table_chars += len(text)
                        if self.GUIDE_MARKER_RE.search(text):
                            guide_marker_hits += 1

        drawing_count = 0
        try:
            with zipfile.ZipFile(docx_path, "r") as archive:
                xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
                drawing_count = xml.count("<w:drawing")
        except Exception:
            drawing_count = 0

        return {
            "paragraph_count": len(paragraphs),
            "paragraph_chars": paragraph_chars,
            "table_cells_total": table_cells_total,
            "table_fill_ratio": round((table_cells_filled / table_cells_total), 4) if table_cells_total else 0.0,
            "table_chars": table_chars,
            "drawing_count": drawing_count,
            "guide_marker_hits": guide_marker_hits,
        }

    def _build_benchmark_compare(
        self, profile: TemplateProfile, output_path: Path, qa_report: dict[str, Any]
    ) -> dict[str, Any]:
        benchmark = self._resolve_benchmark_docx(profile)
        if benchmark is None or not benchmark.exists():
            return {
                "status": "benchmark_missing",
                "benchmark_docx": "",
                "output_docx": str(output_path),
                "qa_passed": bool(qa_report.get("passed")),
                "message": "비교용 벤치마크 파일을 찾지 못했습니다.",
            }

        generated_metrics = self._doc_metrics(output_path)
        benchmark_metrics = self._doc_metrics(benchmark)
        table_fill_gap = round(generated_metrics["table_fill_ratio"] - benchmark_metrics["table_fill_ratio"], 4)
        if benchmark_metrics["table_chars"] > 0:
            table_char_gap_ratio = round(
                abs(generated_metrics["table_chars"] - benchmark_metrics["table_chars"]) / benchmark_metrics["table_chars"],
                4,
            )
        else:
            table_char_gap_ratio = 0.0

        source_docx = Path(profile.source_docx)
        is_template_reference = source_docx.exists() and benchmark.resolve() == source_docx.resolve()
        if is_template_reference:
            benchmark_pass = qa_report.get("error_count", 0) == 0
            acceptance: dict[str, Any] = {
                "mode": "template_reference",
                "note": "원본 템플릿을 기준으로 비교하므로 품질 판정은 QA 오류 건수 중심으로 판단합니다.",
                "qa_error_count_max": 0,
            }
        else:
            benchmark_pass = (
                table_fill_gap >= -0.02 and table_char_gap_ratio <= 0.05 and qa_report.get("error_count", 0) == 0
            )
            acceptance = {
                "mode": "benchmark_document",
                "table_fill_ratio_gap_min": -0.02,
                "table_char_gap_ratio_max": 0.05,
                "qa_error_count_max": 0,
            }
        return {
            "status": "ok",
            "benchmark_docx": str(benchmark),
            "output_docx": str(output_path),
            "benchmark_mode": "template_reference" if is_template_reference else "benchmark_document",
            "benchmark_metrics": benchmark_metrics,
            "generated_metrics": generated_metrics,
            "gap": {
                "paragraph_chars": generated_metrics["paragraph_chars"] - benchmark_metrics["paragraph_chars"],
                "table_chars": generated_metrics["table_chars"] - benchmark_metrics["table_chars"],
                "table_fill_ratio": table_fill_gap,
                "drawing_count": generated_metrics["drawing_count"] - benchmark_metrics["drawing_count"],
                "guide_marker_hits": generated_metrics["guide_marker_hits"] - benchmark_metrics["guide_marker_hits"],
                "table_char_gap_ratio": table_char_gap_ratio,
            },
            "acceptance": acceptance,
            "result": {
                "passed": bool(benchmark_pass),
                "qa_passed": bool(qa_report.get("passed")),
                "qa_error_count": int(qa_report.get("error_count", 0)),
            },
        }
