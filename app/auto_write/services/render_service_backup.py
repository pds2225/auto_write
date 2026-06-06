from __future__ import annotations

import shutil
from pathlib import Path
from typing import Any

from docx import Document

from ..models import GeneratedImage, ProjectInput, TemplateProfile
from .docx_ops import insert_image_after_paragraph, insert_image_in_cell, insert_text_after_paragraph, set_cell_text


class RenderService:
    def _resolve_table_cell(self, document: Document, table_index: int, row_index: int, cell_index: int):
        if table_index < 0 or table_index >= len(document.tables):
            return None, f"표 index 범위를 벗어났습니다(table_index={table_index}, tables={len(document.tables)})"
        table = document.tables[table_index]
        if row_index < 0 or row_index >= len(table.rows):
            return None, f"행 index 범위를 벗어났습니다(table_index={table_index}, row={row_index}, rows={len(table.rows)})"
        row = table.rows[row_index]
        if cell_index < 0 or cell_index >= len(row.cells):
            return (
                None,
                f"열 index 범위를 벗어났습니다(table_index={table_index}, row={row_index}, cell={cell_index}, cells={len(row.cells)})",
            )
        return row.cells[cell_index], ""

    def render(
        self,
        profile: TemplateProfile,
        project_input: ProjectInput,
        images: list[GeneratedImage],
        output_path: Path,
    ) -> dict[str, Any]:
        source_docx = Path(profile.source_docx)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source_docx, output_path)
        document = Document(str(output_path))
        answers = project_input.answers
        errors: list[str] = []
        warnings: list[str] = []

        section_count = 0
        for section in profile.sections:
            value = str(answers.get(section.field_id, "") or "").strip()
            if not value:
                continue
            if insert_text_after_paragraph(document, section.anchor_text, value, allow_create_paragraph=False):
                section_count += 1
            else:
                warnings.append(f"문단 anchor를 찾지 못했습니다: {section.anchor_text}")

        cell_count = 0
        for table in profile.tables:
            table_index = table.table_index
            for cell in table.cells:
                value = str(answers.get(cell.cell_id, "") or "").strip()
                if not value:
                    continue
                target_cell, reason = self._resolve_table_cell(document, table_index, cell.row, cell.cell)
                if target_cell is None:
                    errors.append(f"표 셀 위치 오류: {table.label} r{cell.row} c{cell.cell} ({reason})")
                    continue
                set_cell_text(target_cell, value)
                cell_count += 1

        image_by_slot = {image.slot_id: image for image in images}
        image_count = 0
        for slot in profile.image_slots:
            generated = image_by_slot.get(slot.slot_id)
            if generated is None:
                continue
            image_path = Path(generated.path)
            if not image_path.exists():
                errors.append(f"이미지 파일을 찾지 못했습니다: {slot.label} ({image_path.name})")
                continue
            anchor_ref = slot.anchor_ref
            ok = False
            if slot.anchor_type == "table_cell":
                table_index = int(anchor_ref.get("table_index", 0))
                row_index = int(anchor_ref.get("row", 0))
                cell_index = int(anchor_ref.get("cell", 0))
                _, reason = self._resolve_table_cell(document, table_index, row_index, cell_index)
                if reason:
                    errors.append(f"이미지 표 위치 오류: {slot.label} ({reason})")
                    continue
                ok = insert_image_in_cell(
                    document,
                    table_index=table_index,
                    row=row_index,
                    cell_index=cell_index,
                    image_path=image_path,
                    width_cm=float(slot.size_hint.get("width_cm", 12.0)),
                )
            else:
                ok = insert_image_after_paragraph(
                    document,
                    anchor_text=str(anchor_ref.get("anchor_text", "")),
                    image_path=image_path,
                    width_cm=float(slot.size_hint.get("width_cm", 14.0)),
                    insert_offset=int(anchor_ref.get("insert_offset", 2)),
                    allow_create_paragraph=False,
                )
            if ok:
                image_count += 1
            else:
                errors.append(f"이미지를 넣지 못했습니다: {slot.label}")

        document.save(str(output_path))
        return {
            "output_path": str(output_path),
            "sections_written": section_count,
            "cells_written": cell_count,
            "images_written": image_count,
            "errors": errors,
            "warnings": warnings,
        }
