"""infographic_suggest.py

문서 내용을 분석해 **인포그래픽·도식·참고이미지 삽입 위치를 제안**하는 리포트를 만든다.
실제 이미지를 삽입하지는 않는다(삽입은 ``image_apply`` 가 담당한다). 본 모듈은
"어디에 / 어떤 형태로 / 어떤 슬라이드 프롬프트로" 넣으면 좋을지 **제안만** 한다.

두 가지 경로를 제공한다.
  - ``suggest_images``    : 키워드 → 시각화 유형 매핑(결정론, AI 미사용). 폴백 기본값.
  - ``suggest_images_ai`` : Claude 로 위치·유형·슬라이드 프롬프트를 생성한다.
    AI 키가 없거나 호출에 실패하면 자동으로 ``suggest_images`` 로 폴백한다.

각 제안에는 NotebookLM 의 슬라이드 생성 기능에 그대로 붙여넣을 ``slide_prompt`` 가 담긴다.
문서에 없는 수치는 프롬프트에 만들어 넣지 않는다(날조 0).
"""

from __future__ import annotations

import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from docx import Document

# (앵커 키워드들, 추천 시각화 유형, 캡션 템플릿, 생성 프롬프트 템플릿)
_SUGGESTION_RULES: list[tuple[tuple[str, ...], str, str, str]] = [
    (("시장규모", "TAM", "SAM", "SOM", "시장 전망", "성장률"),
     "막대/도넛 차트",
     "[그림] 목표 시장규모 및 성장 전망(TAM·SAM·SOM)",
     "TAM/SAM/SOM 3단계 시장규모를 보여주는 깔끔한 도넛 또는 동심원 인포그래픽, 한글 라벨, 정부지원사업 보고서 톤"),
    (("추진일정", "추진 일정", "로드맵", "마일스톤", "일정계획", "단계별"),
     "타임라인/간트",
     "[그림] 사업 추진 일정 로드맵",
     "분기별 마일스톤을 표시하는 수평 타임라인/간트 차트, 한글, 단정한 비즈니스 스타일"),
    (("팀구성", "팀 구성", "조직도", "조직 구성", "인력구성"),
     "조직도",
     "[그림] 팀 구성·역할 조직도",
     "대표/핵심팀/외부협력으로 구성된 조직도, 역할 라벨 포함, 한글, 심플 플랫 디자인"),
    (("비즈니스모델", "BM", "수익모델", "수익 구조", "밸류체인", "가치사슬"),
     "플로우/밸류체인 도식",
     "[그림] 비즈니스 모델·수익 구조 도식",
     "가치사슬과 수익 흐름을 화살표로 연결한 비즈니스 모델 다이어그램, 한글, 인포그래픽"),
    (("프로세스", "절차", "처리 과정", "동작 원리", "구조도", "아키텍처", "시스템 구성"),
     "플로우차트/구성도",
     "[그림] 핵심 기술·시스템 구성도",
     "시스템/기술 아키텍처 블록 다이어그램, 모듈 간 연결, 한글 라벨, 기술 보고서 스타일"),
    (("경쟁사", "경쟁력", "비교", "차별성", "포지셔닝"),
     "비교표/포지셔닝맵",
     "[그림] 경쟁사 대비 차별성·포지셔닝 맵",
     "2축 포지셔닝 맵 또는 경쟁 비교 인포그래픽, 자사 강조, 한글, 깔끔한 비즈니스 톤"),
    (("매출", "재무", "손익", "매출계획", "재무계획", "추정"),
     "추세 선/막대 그래프",
     "[그림] 연도별 매출·재무 추정",
     "연도별 매출/영업이익 추정을 보여주는 막대+선 복합 그래프, 한글, 보고서 스타일"),
]


@dataclass
class ImageSuggestion:
    anchor_text: str            # 제안 위치(가까운 단락 텍스트)
    visual_type: str            # 추천 시각화 유형
    caption: str                # 캡션(문서 삽입용)
    prompt: str                 # (하위호환) 이미지 생성 프롬프트
    keyword: str                # 트리거된 키워드(AI 제안이면 "AI")
    slide_prompt: str = ""      # NotebookLM 슬라이드 생성 기능에 붙여넣을 프롬프트

    def as_dict(self) -> dict[str, Any]:
        return {
            "anchor_text": self.anchor_text[:80],
            "visual_type": self.visual_type,
            "caption": self.caption,
            "prompt": self.prompt,
            "keyword": self.keyword,
            "slide_prompt": self.slide_prompt,
        }


@dataclass
class InfographicReport:
    suggestions: list[ImageSuggestion] = field(default_factory=list)
    existing_images: int = 0

    def as_dict(self) -> dict[str, Any]:
        return {
            "suggestion_count": len(self.suggestions),
            "existing_images": self.existing_images,
            "suggestions": [s.as_dict() for s in self.suggestions],
        }


def _count_existing_images(doc: Document) -> int:
    from docx.oxml.ns import qn
    return len(doc.element.body.findall(".//" + qn("w:drawing")))


def suggest_images(doc: Document, *, max_suggestions: int = 8) -> InfographicReport:
    """문서 단락을 훑어 도식 삽입 제안을 생성한다(중복 유형은 1회만)."""
    report = InfographicReport(existing_images=_count_existing_images(doc))
    used_types: set[str] = set()

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # 표 헤더 텍스트도 앵커 후보에 포함
    for table in doc.tables:
        if table.rows:
            header = " ".join(c.text.strip() for c in table.rows[0].cells if c.text.strip())
            if header:
                paragraphs.append(header)

    for text in paragraphs:
        if len(report.suggestions) >= max_suggestions:
            break
        for keywords, vtype, caption, prompt in _SUGGESTION_RULES:
            if vtype in used_types:
                continue
            hit = next((kw for kw in keywords if kw in text), None)
            if hit:
                report.suggestions.append(ImageSuggestion(
                    anchor_text=text,
                    visual_type=vtype,
                    caption=caption,
                    prompt=prompt,
                    keyword=hit,
                    slide_prompt=_keyword_slide_prompt(vtype, caption),
                ))
                used_types.add(vtype)
                break
    return report


def suggest_images_docx(path: str | Path, *, max_suggestions: int = 8) -> InfographicReport:
    return suggest_images(Document(str(Path(path))), max_suggestions=max_suggestions)


# --------------------------------------------------------------------------- 슬라이드 프롬프트
# 모든 슬라이드 프롬프트에 디폴트로 덧붙이는 디자인 규칙(흰 배경·블루/무채색·키워드 위주 등).
_SLIDE_DESIGN_GUIDE = (
    "[디자인] 흰 배경에 블루 계열 또는 무채색 텍스트로 단정하고 깔끔하게. "
    "의미 전달이 최우선이며 가독성·가시성을 중시한다. "
    "여러 항목은 세로가 아니라 가로로 나열한다. "
    "한 텍스트 상자는 3줄을 넘기지 않는다. "
    "상세 설명이 아니라 키워드 위주로(한 항목당 최대 1~2문장). "
    "색은 흰 배경·블루 계열·단순한 컬러 중심으로 절제한다. "
    "정보량이 많으면 표·비교표·프로세스·레이어 구조로 바꿔 정리한다."
)


def _append_design_guide(slide_prompt: str) -> str:
    """슬라이드 프롬프트에 디자인 규칙이 없으면 디폴트로 덧붙인다(중복 방지)."""
    if not slide_prompt or "[디자인]" in slide_prompt:
        return slide_prompt or _SLIDE_DESIGN_GUIDE
    sep = "" if slide_prompt.endswith((" ", ".", "。", "!", "?")) else " "
    return f"{slide_prompt}{sep}{_SLIDE_DESIGN_GUIDE}"


def _keyword_slide_prompt(visual_type: str, caption: str) -> str:
    """키워드 폴백용 NotebookLM 슬라이드 생성 프롬프트(결정론). 디자인 규칙 포함."""
    topic = caption.replace("[그림] ", "").strip() or visual_type
    body = (
        f"다음 내용을 사업계획서용 슬라이드 1장으로 만들어줘: '{topic}'. "
        f"표현 형식은 '{visual_type}'. "
        f"내가 제공한 문서 본문에 실제로 적힌 수치·항목만 사용하고, "
        f"문서에 없는 데이터는 추측하거나 만들어 넣지 마. "
        f"한국어, 정부지원사업 보고서 톤."
    )
    return _append_design_guide(body)


def _match_anchor(anchor: str, para_texts: list[str]) -> str:
    """AI 가 돌려준 anchor 를 문서에 실제로 있는 단락 텍스트로 보정한다(없으면 "")."""
    if not anchor:
        return ""
    for t in para_texts:
        if anchor in t or t in anchor:
            return t
    key = anchor[:12].strip()
    if key:
        for t in para_texts:
            if key in t:
                return t
    return ""


_AI_SYSTEM_PROMPT = (
    "당신은 한국 정부지원사업 사업계획서를 분석해, NotebookLM 의 슬라이드 생성 기능에 "
    "그대로 붙여넣을 '슬라이드 생성 프롬프트'를 만드는 전문가입니다.\n\n"
    "[규칙 — 위반 시 결과 폐기]\n"
    "1. 문서에서 시각자료(차트·도식·타임라인·조직도 등)가 들어가면 좋은 위치를 고른다.\n"
    "2. 각 위치마다 anchor / visual_type / caption / slide_prompt 를 만든다.\n"
    "3. anchor 는 반드시 입력 문서에 실제로 등장한 문구의 일부를 그대로 사용한다(창작 금지).\n"
    "4. slide_prompt 에는 문서에 실제로 있는 수치·항목만 쓰고, 없는 수치는 절대 지어내지 않는다.\n"
    "5. slide_prompt 는 한국어로, 상세 설명이 아니라 '키워드 위주'(항목당 최대 1~2문장)로 간결히 쓴다.\n"
    "6. 색상·레이아웃 등 디자인 규칙은 시스템이 자동으로 덧붙이므로, slide_prompt 에는 "
    "무엇을 담을지(내용)에만 집중한다.\n"
    "7. 서로 다른 유형 위주로 최대 {max} 개까지만 제안한다.\n\n"
    "반환 형식(JSON object):\n"
    '{"suggestions": [{"anchor": "문서 문구 일부", "visual_type": "유형", '
    '"caption": "[그림] 제목", "slide_prompt": "NotebookLM 에 붙여넣을 지시문"}]}'
)


def suggest_images_ai(
    doc: Document,
    *,
    openai_service: Optional[Any] = None,
    max_suggestions: int = 8,
) -> InfographicReport:
    """Claude 로 그림 위치·유형·슬라이드 프롬프트를 제안한다.

    ``openai_service`` 가 없거나(키 미연결) 호출/파싱에 실패하면 ``suggest_images``
    (키워드 규칙)로 자동 폴백한다 — 따라서 키가 없어도 항상 동작한다.
    """
    available = bool(openai_service is not None and getattr(openai_service, "available", False))
    if not available:
        return suggest_images(doc, max_suggestions=max_suggestions)

    para_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    doc_text = "\n".join(para_texts)[:6000]

    system = _AI_SYSTEM_PROMPT.replace("{max}", str(max_suggestions))
    user = json.dumps(
        {"max_suggestions": max_suggestions, "document": doc_text},
        ensure_ascii=False,
    )
    try:
        result = openai_service.complete_json(system, user, max_tokens=2048)
    except Exception:
        result = None

    if not isinstance(result, dict) or not isinstance(result.get("suggestions"), list):
        return suggest_images(doc, max_suggestions=max_suggestions)

    report = InfographicReport(existing_images=_count_existing_images(doc))
    for item in result["suggestions"]:
        if len(report.suggestions) >= max_suggestions:
            break
        if not isinstance(item, dict):
            continue
        vtype = str(item.get("visual_type", "")).strip()
        slide_prompt = str(item.get("slide_prompt", "")).strip()
        if not (vtype and slide_prompt):
            continue
        anchor = str(item.get("anchor", "")).strip()
        caption = str(item.get("caption", "")).strip() or f"[그림] {vtype}"
        slide_prompt = _append_design_guide(slide_prompt)  # 디자인 규칙 디폴트 부가
        matched = _match_anchor(anchor, para_texts)
        report.suggestions.append(ImageSuggestion(
            anchor_text=matched or anchor,
            visual_type=vtype,
            caption=caption,
            prompt=slide_prompt,
            keyword="AI",
            slide_prompt=slide_prompt,
        ))

    # AI 가 빈 제안을 돌려주면 키워드 폴백으로 대체
    if not report.suggestions:
        return suggest_images(doc, max_suggestions=max_suggestions)
    return report


def suggest_images_ai_docx(
    path: str | Path,
    *,
    openai_service: Optional[Any] = None,
    max_suggestions: int = 8,
) -> InfographicReport:
    return suggest_images_ai(
        Document(str(Path(path))),
        openai_service=openai_service,
        max_suggestions=max_suggestions,
    )
