from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document

from ..models import EvidenceSource, GeneratedImage, ProjectInput, TemplateProfile


class QAService:
    GUIDE_MARKER_RE = re.compile(r"(※|<[^>]+>|기재|작성요령|작성방법|예시|OOO|○○○)")
    CRITICAL_GUIDE_MARKER_RE = re.compile(
        r"(<[^>]*(기재|작성|입력|예시|제목|OOO|○○○|000)[^>]*>|OOO|○○○|000|작성요령|작성방법)"
    )
    NON_BUSINESS_LABEL_RE = re.compile(
        r"(^(?:[◦ㅇ■□●·•]+)$|담\s*당\s*자|사실과 다름이 없음을 확인|추천기관|신청기업|제출목록|증빙서류|평가대상에서 제외|별첨|붙임)",
        re.IGNORECASE,
    )

    @staticmethod
    def _msg(category: str, message: str) -> str:
        return f"[{category}] {message}"

    @classmethod
    def _is_non_business_label(cls, text: str) -> bool:
        return bool(cls.NON_BUSINESS_LABEL_RE.search(str(text or "").strip()))

    def _normalize_match_text(self, text: str) -> str:
        normalized = re.sub(r"\s+", " ", text or "").strip()
        normalized = normalized.replace("–", "-").replace("—", "-").replace("‑", "-")
        for _ in range(2):
            if len(normalized) < 16 or len(normalized) % 2 != 0:
                break
            half = len(normalized) // 2
            left = normalized[:half].strip()
            right = normalized[half:].strip()
            if left and left == right:
                normalized = left
                continue
            break
        return normalized

    def _match_anchor(self, candidate: str, anchor_text: str) -> bool:
        anchor = self._normalize_match_text(anchor_text)
        text = self._normalize_match_text(candidate)
        if not anchor or not text:
            return False
        if anchor in text:
            return True
        compact_anchor = anchor.replace(" ", "")
        compact_text = text.replace(" ", "")
        return bool(compact_anchor and compact_anchor in compact_text)

    @staticmethod
    def _paragraph_text(paragraph) -> str:
        parts = [
            str(node.text or "").strip()
            for node in paragraph.iter()
            if str(getattr(node, "tag", "")).endswith("}t")
        ]
        return " ".join(part for part in parts if part).strip()

    def _is_meaningful_text(self, text: str, anchor_text: str = "") -> bool:
        candidate = str(text or "").strip()
        if not candidate:
            return False
        if self._is_non_business_label(candidate):
            return False
        if self.GUIDE_MARKER_RE.search(candidate):
            return False
        if anchor_text and self._match_anchor(candidate, anchor_text):
            return False
        return True

    def _table_has_meaningful_text(self, table_element, anchor_text: str) -> bool:
        for cell in table_element.iter():
            if not str(getattr(cell, "tag", "")).endswith("}tc"):
                continue
            texts = [str(node.text or "").strip() for node in cell.iter() if str(getattr(node, "tag", "")).endswith("}t")]
            joined = re.sub(r"\s+", " ", " ".join(item for item in texts if item)).strip()
            if self._is_meaningful_text(joined, anchor_text):
                return True
        return False

    def _has_content_after_anchor(self, document: Document, anchor_text: str) -> bool:
        body_items = list(document.element.body)
        for index, element in enumerate(body_items):
            if not str(getattr(element, "tag", "")).endswith("}p"):
                continue
            current = self._paragraph_text(element)
            if not self._match_anchor(current, anchor_text):
                continue
            for next_index in range(index + 1, min(index + 10, len(body_items))):
                next_item = body_items[next_index]
                tag = str(getattr(next_item, "tag", ""))
                if tag.endswith("}p"):
                    if self._is_meaningful_text(self._paragraph_text(next_item), anchor_text):
                        return True
                    continue
                if tag.endswith("}tbl") and self._table_has_meaningful_text(next_item, anchor_text):
                    return True
        return False

    def _cell_has_meaningful_text(self, document: Document, table_index: int, row: int, cell: int) -> bool:
        if table_index < 0 or table_index >= len(document.tables):
            return False
        table = document.tables[table_index]
        if row < 0 or row >= len(table.rows):
            return False
        if cell < 0 or cell >= len(table.rows[row].cells):
            return False
        text = table.rows[row].cells[cell].text.strip()
        return self._is_meaningful_text(text)

    def _collect_guide_markers(self, document: Document, limit: int = 12) -> list[str]:
        found: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if self.GUIDE_MARKER_RE.search(text):
                found.append(text[:120])
            if len(found) >= limit:
                return found
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if not text:
                        continue
                    if self.GUIDE_MARKER_RE.search(text):
                        found.append(text[:120])
                    if len(found) >= limit:
                        return found
        return found

    def _iter_document_text_with_page(self, document: Document):
        page_no = 1
        for element in list(document.element.body):
            tag = str(getattr(element, "tag", ""))
            if tag.endswith("}p"):
                text = self._paragraph_text(element)
                if text:
                    yield page_no, text
                for node in element.iter():
                    node_tag = str(getattr(node, "tag", ""))
                    if node_tag.endswith("}lastRenderedPageBreak"):
                        page_no += 1
                        continue
                    if not node_tag.endswith("}br"):
                        continue
                    break_type = (
                        node.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type")
                        or node.get("w:type")
                        or node.get("type")
                    )
                    if str(break_type or "").lower() == "page":
                        page_no += 1
            elif tag.endswith("}tbl"):
                for cell in element.iter():
                    if not str(getattr(cell, "tag", "")).endswith("}tc"):
                        continue
                    texts = [str(node.text or "").strip() for node in cell.iter() if str(getattr(node, "tag", "")).endswith("}t")]
                    joined = re.sub(r"\s+", " ", " ".join(item for item in texts if item)).strip()
                    if joined:
                        yield page_no, joined

    def _collect_placeholder_pages(self, document: Document, limit: int = 8) -> list[int]:
        pages: list[int] = []
        seen: set[int] = set()
        for page_no, text in self._iter_document_text_with_page(document):
            if "○○○" not in text:
                continue
            if page_no in seen:
                continue
            seen.add(page_no)
            pages.append(page_no)
            if len(pages) >= limit:
                break
        return pages

    @staticmethod
    def _missing_field_message(question: Any, target_kind: str) -> str:
        label = str(getattr(question, "label", "") or "").strip() or "필수 입력"
        question_id = str(getattr(question, "question_id", "") or "").strip()
        target = getattr(question, "target", {}) or {}
        if target_kind == "project_meta":
            key = str(target.get("key", "")).strip() or "project_title"
            return f"❌ '{label}' 항목이 비어있습니다. project_input.json의 meta.{key}를 채워주세요."
        if target_kind == "organization_profile":
            key = str(target.get("key", "")).strip() or "name"
            return f"❌ '{label}' 항목이 비어있습니다. project_input.json의 organization.{key}를 채워주세요."
        if target_kind in {"section", "table_cell"} and question_id:
            return f"❌ '{label}' 항목이 비어있습니다. project_input.json의 answers.{question_id}를 채워주세요."
        if question_id:
            return f"❌ '{label}' 항목이 비어있습니다. project_input.json의 answers.{question_id}를 채워주세요."
        return f"❌ '{label}' 항목이 비어있습니다. 입력값을 확인해주세요."

    def build_report(
        self,
        profile: TemplateProfile,
        project_input: ProjectInput,
        render_result: dict[str, Any],
        images: list[GeneratedImage],
        evidence: list[EvidenceSource],
        preview_result: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        errors: list[str] = []
        warnings: list[str] = []
        answers = project_input.answers
        output_path = Path(render_result["output_path"])
        output_doc = Document(str(output_path)) if output_path.exists() else None
        section_by_field_id = {section.field_id: section for section in profile.sections}
        table_by_id = {table.table_id: table for table in profile.tables}

        for question in profile.questions:
            if not question.required:
                continue
            if self._is_non_business_label(question.label):
                continue
            target = question.target.get("kind")
            if target == "project_meta":
                value = project_input.project_meta.get(question.target.get("key", ""))
                if not str(value or "").strip():
                    errors.append(self._msg("필수입력", self._missing_field_message(question, str(target))))
            elif target == "organization_profile":
                value = project_input.organization_profile.get(question.target.get("key", ""))
                if not str(value or "").strip():
                    errors.append(self._msg("필수입력", self._missing_field_message(question, str(target))))
            elif target == "section":
                section_field_id = str(question.target.get("field_id", "") or question.question_id)
                section = section_by_field_id.get(section_field_id)
                if output_doc is not None and section is not None:
                    if not self._has_content_after_anchor(output_doc, section.anchor_text):
                        message = self._missing_field_message(question, str(target)).replace(
                            "비어있습니다.", "결과 문서에서 비어있습니다."
                        )
                        errors.append(self._msg("필수입력", message))
                else:
                    value = str(answers.get(section_field_id, answers.get(question.question_id, "")) or "").strip()
                    if not value:
                        errors.append(self._msg("필수입력", self._missing_field_message(question, str(target))))
            elif target == "table_cell":
                table = table_by_id.get(str(question.target.get("table_id", "")))
                cell_id = str(question.target.get("cell_id", ""))
                table_cell = next((cell for cell in table.cells if cell.cell_id == cell_id), None) if table else None
                if output_doc is not None and table is not None and table_cell is not None:
                    if not self._cell_has_meaningful_text(output_doc, table.table_index, table_cell.row, table_cell.cell):
                        message = self._missing_field_message(question, str(target)).replace(
                            "비어있습니다.", "결과 문서에서 비어있습니다."
                        )
                        errors.append(self._msg("필수입력", message))
                else:
                    value = str(answers.get(cell_id, answers.get(question.question_id, "")) or "").strip()
                    if not value:
                        errors.append(self._msg("필수입력", self._missing_field_message(question, str(target))))
            else:
                value = answers.get(question.question_id)
                if not str(value or "").strip():
                    errors.append(self._msg("필수입력", self._missing_field_message(question, str(target))))

        image_by_slot = {image.slot_id for image in images}
        for slot in profile.image_slots:
            if self._is_non_business_label(slot.label):
                continue
            if slot.required and slot.slot_id not in image_by_slot:
                errors.append(
                    self._msg("필수이미지", f"❌ '{slot.label}' 이미지가 생성되지 않았습니다. 이미지 설명 입력을 확인해주세요.")
                )

        if len(images) > 1:
            warnings.append(self._msg("이미지수", f"⚠️ 이미지가 {len(images)}장 입력되었습니다. 데이터바우처 규정상 1장만 삽입됩니다."))

        if render_result.get("errors"):
            for issue in render_result["errors"]:
                errors.append(self._msg("렌더링", f"❌ 문서 생성 중 문제가 발생했습니다. {issue}"))
        if render_result.get("warnings"):
            for issue in render_result["warnings"]:
                warnings.append(self._msg("렌더링", f"⚠️ 문서 생성 중 확인이 필요한 항목이 있습니다. {issue}"))

        for table in profile.tables:
            for cell in table.cells:
                value = str(answers.get(cell.cell_id, "") or "")
                if len(value) > 180:
                    warnings.append(
                        self._msg("표길이", f"⚠️ '{cell.label}' 항목이 너무 길어 표가 깨질 수 있습니다. 문장을 짧게 줄여주세요.")
                    )

        if output_doc is not None:
            guide_markers = self._collect_guide_markers(output_doc)
            for marker in guide_markers[:6]:
                if self.CRITICAL_GUIDE_MARKER_RE.search(marker):
                    errors.append(self._msg("가이드문구", f"❌ 결과 문서에 가이드 문구가 남아 있습니다: {marker}"))
                else:
                    warnings.append(self._msg("가이드문구", f"⚠️ 결과 문서에 가이드 문구가 남아 있습니다: {marker}"))
            placeholder_pages = self._collect_placeholder_pages(output_doc)
            for page_no in placeholder_pages:
                warnings.append(
                    self._msg("빈칸", f"⚠️ {page_no}번째 페이지에 빈 칸(○○○)이 남아있습니다. 실제 값으로 교체해주세요.")
                )

        if project_input.evidence_requests and not evidence:
            warnings.append(self._msg("출처", "⚠️ 검색 요청은 있었지만 저장된 출처가 없습니다. 검색 키워드를 확인해주세요."))

        for request in project_input.evidence_requests:
            if request.topic and not any(source.topic == request.topic for source in evidence):
                warnings.append(self._msg("출처", f"⚠️ '{request.topic}' 주제의 검색 결과를 찾지 못했습니다. 주제를 더 구체적으로 입력해주세요."))

        if not output_path.exists():
            errors.append(self._msg("산출물", "❌ 최종 DOCX 파일이 생성되지 않았습니다. 템플릿 파일과 입력값을 다시 확인해주세요."))
        preview_result = preview_result or {}
        preview_status = str(preview_result.get("status", "") or "").strip().lower()
        if preview_status == "failed":
            errors.append(self._msg("미리보기", "❌ DOCX 페이지 미리보기를 만들지 못했습니다. 문서 화면 검수가 빠졌습니다."))
        elif preview_status == "skipped":
            warnings.append(self._msg("미리보기", "⚠️ DOCX 페이지 미리보기를 만들지 못해 화면 검수를 건너뛰었습니다."))
        for issue in preview_result.get("warnings", []) or []:
            warnings.append(self._msg("미리보기", f"⚠️ {issue}"))
        for issue in preview_result.get("errors", []) or []:
            errors.append(self._msg("미리보기", f"❌ {issue}"))

        return {
            "passed": not errors,
            "error_count": len(errors),
            "warning_count": len(warnings),
            "errors": errors,
            "warnings": warnings,
            "summary": {
                "sections_written": render_result.get("sections_written", 0),
                "cells_written": render_result.get("cells_written", 0),
                "images_written": render_result.get("images_written", 0),
                "sources_collected": len(evidence),
                "preview_pages": int(preview_result.get("page_count", 0) or 0),
            },
        }
