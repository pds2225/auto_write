from __future__ import annotations

import json
import shutil
from pathlib import Path
from typing import Any

from docx import Document


class DocxEditService:
    """Editable text projection for DOCX without overwriting the generated source."""

    def load_blocks(self, docx_path: str | Path, limit: int = 240) -> list[dict[str, Any]]:
        path = Path(docx_path)
        if not path.is_file():
            return []
        doc = Document(str(path))
        rows: list[dict[str, Any]] = []
        for index, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                rows.append({"id": f"p:{index}", "kind": "본문", "location": f"문단 {index + 1}", "text": paragraph.text})
                if len(rows) >= limit:
                    return rows
        for table_index, table in enumerate(doc.tables):
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if not text:
                        continue
                    rows.append(
                        {
                            "id": f"t:{table_index}:r:{row_index}:c:{cell_index}",
                            "kind": "표",
                            "location": f"표 {table_index + 1} / 행 {row_index + 1} / 열 {cell_index + 1}",
                            "text": cell.text,
                        }
                    )
                    if len(rows) >= limit:
                        return rows
        return rows

    @staticmethod
    def _set_paragraph_text(paragraph, text: str) -> None:
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(text)

    @staticmethod
    def _set_cell_text(cell, text: str) -> None:
        if not cell.paragraphs:
            cell.text = text
            return
        first = cell.paragraphs[0]
        DocxEditService._set_paragraph_text(first, text)
        for paragraph in cell.paragraphs[1:]:
            DocxEditService._set_paragraph_text(paragraph, "")

    def apply_edits(
        self,
        source_docx: str | Path,
        edits: dict[str, str],
        output_docx: str | Path,
        lock_path: str | Path,
    ) -> dict[str, Any]:
        source = Path(source_docx)
        output = Path(output_docx)
        if source.resolve() == output.resolve():
            raise ValueError("원본 덮어쓰기는 금지입니다.")
        output.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, output)
        doc = Document(str(output))
        applied: dict[str, str] = {}

        for block_id, new_text in edits.items():
            text = str(new_text)
            try:
                if block_id.startswith("p:"):
                    index = int(block_id.split(":")[1])
                    paragraph = doc.paragraphs[index]
                    if paragraph.text != text:
                        self._set_paragraph_text(paragraph, text)
                        applied[block_id] = text
                elif block_id.startswith("t:"):
                    parts = block_id.split(":")
                    table_index = int(parts[1])
                    row_index = int(parts[3])
                    cell_index = int(parts[5])
                    cell = doc.tables[table_index].rows[row_index].cells[cell_index]
                    if cell.text != text:
                        self._set_cell_text(cell, text)
                        applied[block_id] = text
            except (IndexError, ValueError):
                continue

        doc.save(str(output))
        lock_file = Path(lock_path)
        existing = {}
        if lock_file.exists():
            try:
                existing = json.loads(lock_file.read_text(encoding="utf-8"))
            except Exception:
                existing = {}
        locks = existing.get("locks", {}) if isinstance(existing, dict) else {}
        locks.update(applied)
        lock_file.write_text(
            json.dumps({"source": str(source), "edited_output": str(output), "locks": locks}, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        return {"output": str(output), "applied_count": len(applied), "locks": locks}

    def load_locks(self, lock_path: str | Path) -> dict[str, str]:
        path = Path(lock_path)
        if not path.exists():
            return {}
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
            locks = data.get("locks", {})
            return locks if isinstance(locks, dict) else {}
        except Exception:
            return {}
