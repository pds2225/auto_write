"""doc_quality_score.py

후처리가 끝난 DOCX 의 **문서 품질 점수(100점 만점)** 를 결정론적으로 산정한다.

각 항목은 "후처리 후에도 남아있는 결함" 을 세어 만점에서 감점하는 방식이다.
AI 를 호출하지 않으며, 동일 입력에 항상 동일 점수를 반환한다.

배점(총 100)
-----------
1. 안내문구 제거            15
2. 글머리표 공백 정리        10
3. 문단·공백 정리           10
4. 글자크기·스타일 일관성    15
5. 표 내부 품질             10
6. 주요문장 강조 적정성      10
7. 문서 유형별 구조 적합성   15
8. PSST/보고서 구조 충족도   10
9. 이미지·도식 제안 적정성    5

품질 게이트: 90↑ 우수 / 85↑ 통과 / 70↑ 보완 필요 / 70 미만 실패
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

from .doc_quality_ops import (
    _BULLET_PREFIX_RE, _MULTI_SPACE_RE, _PURE_GUIDE_RE, _is_guide_text,
    _text_nodes, _element_has_drawing,
)
from .docx_ops import _iter_body_paragraphs, _paragraph_text
from .qa_service import QAService

# 플레이스홀더(OOO / ○○○ / 000) — 명백한 미작성 흔적 (general 안내 판정 기준)
# 이전에 쓰던 GUIDE_MARKER_RE("기재"·"예시"·"※" 단독)는 실제 내용에 과도하게 매치되어
# false positive가 많았으므로 더 보수적인 기준으로 교체한다.
_PLACEHOLDER_RE = re.compile(r"(?<!\w)(OOO|○○○|000)(?!\w)")

# 유형별 '있어야 할' 구조 키워드(구조 적합성 채점용)
_TYPE_STRUCTURE_KEYWORDS: dict[str, tuple[str, ...]] = {
    "business_plan": ("문제", "해결", "시장", "성장", "팀", "사업화"),
    "rnd_plan": ("목표", "기술", "성능", "방법", "일정", "사업화"),
    "pitch_deck": ("문제", "솔루션", "시장", "팀", "투자"),
    "consulting_report": ("현황", "진단", "개선", "실행", "기대효과"),
    "policy_fund_report": ("자금", "용도", "상환", "매출", "리스크"),
    "certification_report": ("요건", "충족", "보완", "서류"),
    "export_report": ("수출", "시장", "바이어", "전략"),
    "field_clinic_report": ("현황", "진단", "애로", "처방", "개선"),
    "generic_submission": (),
}


@dataclass
class ScoreItem:
    key: str
    label: str
    score: float
    max_score: float
    defects: int
    detail: str = ""

    def as_dict(self) -> dict[str, Any]:
        return {
            "key": self.key, "label": self.label,
            "score": round(self.score, 1), "max_score": self.max_score,
            "defects": self.defects, "detail": self.detail,
        }


@dataclass
class QualityScore:
    total: float
    grade: str
    passed: bool                  # 85점 이상
    items: list[ScoreItem] = field(default_factory=list)

    def as_dict(self) -> dict[str, Any]:
        return {
            "total": round(self.total, 1),
            "grade": self.grade,
            "passed": self.passed,
            "items": [i.as_dict() for i in self.items],
        }


# ---------------------------------------------------------------------------
# 잔존 결함 스캐너
# ---------------------------------------------------------------------------

def _iter_all_cell_texts(doc: Document):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield cell


def _scan_guide(doc: Document) -> tuple[int, int]:
    """(critical, general) 잔존 안내문구 — body 직계 단락만 검사.

    remove_guide_paragraphs 와 완전히 동일한 범위(body 직계)로 스캔한다.
    표 셀 내부는 (별도 표 안내 삭제기가 담당하므로) 이 채점에서는 제외한다.
    - critical : _is_guide_text 에 매치 (삭제기와 동일한 단일 기준 = _PURE_GUIDE_RE
                 또는 _GUIDE_EXTRA_RE; scan-range == delete-range 로 비대칭 방지)
    - general  : body 직계 단락에 OOO/○○○/000 플레이스홀더가 있는 경우
    """
    critical = general = 0
    for p in _iter_body_paragraphs(doc):
        t = _paragraph_text(p).strip()
        if not t:
            continue
        if _is_guide_text(t):
            critical += 1
        elif _PLACEHOLDER_RE.search(t):
            general += 1
    return critical, general


def _scan_bullet(doc: Document) -> int:
    defects = 0
    for p in doc.paragraphs:
        t = p.text
        if not t:
            continue
        if _BULLET_PREFIX_RE.search(t) or _MULTI_SPACE_RE.search(t):
            defects += 1
    return defects


def _scan_empty_groups(doc: Document) -> int:
    """본문 직계 요소를 순회하며 연속 빈 단락 그룹(2개 이상) 수를 센다.

    remove_empty_paragraphs 와 완전히 동일한 순회 방식:
    - body 직계 자식을 전부 순회한다.
    - 표(w:tbl)·섹션속성 등 비단락 요소는 연속 카운터를 리셋한다.
      (표로 분리된 빈단락이 연속으로 오인되는 false-positive 방지)
    - 완전 빈 단락 = 텍스트 없고 이미지/도형 없는 w:p.
    """
    body = doc.element.body
    groups = 0
    run = 0
    for el in body:
        if el.tag != qn("w:p"):
            if run >= 2:
                groups += 1
            run = 0
            continue
        is_empty = (not _paragraph_text(el)) and (not _element_has_drawing(el))
        if is_empty:
            run += 1
        else:
            if run >= 2:
                groups += 1
            run = 0
    if run >= 2:
        groups += 1
    return groups


def _scan_font_sizes(doc: Document) -> tuple[int, int]:
    """(본문 폰트 크기 종류 수, 이상치 수)."""
    sizes: dict[float, int] = {}
    outliers = 0
    for p in doc.paragraphs:
        for run in p.runs:
            sz = run.font.size
            if sz is None:
                continue
            pt = sz.pt
            sizes[pt] = sizes.get(pt, 0) + 1
            if pt < 8 or pt > 18:
                outliers += 1
    return len(sizes), outliers


def _scan_table_ws(doc: Document) -> int:
    """표 셀의 실제 w:t 노드에서 공백 결함 셀 수를 센다.

    cleanup_table_whitespace 와 동일한 기준으로 검사한다:
    - 첫 w:t 노드 앞 공백 / 마지막 w:t 노드 뒤 공백 / 노드 내부 다중 공백.
    이전에 cell.text(단락 전체 합침, \\n 포함)를 검사하던 방식은
    단락 구분 \\n 자체를 결함으로 오탐했으므로 노드 수준으로 수정한다.
    merged cell 중복 카운팅도 방지한다.
    """
    seen: set[int] = set()
    defects = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tc_id = id(tc)
                if tc_id in seen:
                    continue
                seen.add(tc_id)
                if _element_has_drawing(tc):
                    continue
                nodes = [n for n in _text_nodes(tc) if n.text]
                if not nodes:
                    continue
                has_defect = False
                if nodes[0].text != nodes[0].text.lstrip("  　\t"):
                    has_defect = True
                elif nodes[-1].text != nodes[-1].text.rstrip("  　\t"):
                    has_defect = True
                else:
                    for n in nodes:
                        if _MULTI_SPACE_RE.search(n.text):
                            has_defect = True
                            break
                if has_defect:
                    defects += 1
    return defects


def _count_bold_paragraphs(doc: Document) -> int:
    cnt = 0
    for p in doc.paragraphs:
        if any(r.bold for r in p.runs if (r.text or "").strip()):
            cnt += 1
    return cnt


def _count_nonempty_paragraphs(doc: Document) -> int:
    return sum(1 for p in doc.paragraphs if p.text.strip())


# ---------------------------------------------------------------------------
# 채점기
# ---------------------------------------------------------------------------

def score_document(
    doc: Document,
    *,
    doc_type: str = "generic_submission",
    type_confidence: float = 0.0,
    psst_ratio: float | None = None,
    image_suggestions: int = 0,
    existing_images: int = 0,
    empty_required_cells: int = 0,
) -> QualityScore:
    """후처리된 문서를 채점한다.

    Parameters 는 오케스트레이터가 분류/PSST/이미지제안 단계 결과를 주입한다.
    psst_ratio 가 None 이면 PSST 미적용 유형 → 8번 항목은 보고서 구조 키워드로 대체.
    empty_required_cells 는 (오케스트레이터가 양식 재분석으로 산출한) 미입력 필수셀 수로,
    현재는 **참고용(informational, 가중치 0)** 으로만 표기해 85점 게이트를 흔들지 않는다.
    """
    items: list[ScoreItem] = []
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for cell in _iter_all_cell_texts(doc):
        full_text += "\n" + cell.text

    # 1. 안내문구 제거 (15)
    crit, gen = _scan_guide(doc)
    s1 = max(0.0, 15.0 - crit * 5.0 - gen * 1.0)
    items.append(ScoreItem("guide_removal", "안내문구 제거", s1, 15,
                           crit + gen, f"critical={crit}, general={gen}"))

    # 2. 글머리표 공백 정리 (10)
    b = _scan_bullet(doc)
    s2 = max(0.0, 10.0 - b * 1.0)
    items.append(ScoreItem("bullet_spacing", "글머리표 공백 정리", s2, 10, b,
                           f"잔존 글머리표/다중공백 단락={b}"))

    # 3. 문단·공백 정리 (10)
    eg = _scan_empty_groups(doc)
    s3 = max(0.0, 10.0 - eg * 2.0)
    items.append(ScoreItem("paragraph_cleanup", "문단·공백 정리", s3, 10, eg,
                           f"연속 빈단락 그룹={eg}"))

    # 4. 글자크기·스타일 일관성 (15)
    kinds, outliers = _scan_font_sizes(doc)
    # 정부양식 문서는 제목·소제목·본문·표 등으로 6종 이하면 양호로 본다(이전 기준 4종은 과도).
    # 폰트 정규화(normalize_font_sizes)가 기본 off 이므로 이상치 계수만 적용한다.
    penalty4 = max(0, kinds - 6) * 1.0 + outliers * 2.0
    s4 = max(0.0, 15.0 - penalty4)
    items.append(ScoreItem("font_consistency", "글자크기·스타일 일관성", s4, 15,
                           max(0, kinds - 4) + outliers,
                           f"폰트 종류={kinds}, 이상치={outliers}"))

    # 5. 표 내부 품질 (10)
    tw = _scan_table_ws(doc)
    s5 = max(0.0, 10.0 - tw * 1.0)
    detail5 = f"공백 결함 셀={tw}"
    if empty_required_cells:
        # 참고용 표기만(가중치 0) — 게이트 점수에는 반영하지 않는다.
        detail5 += f" / 미입력 필수셀(참고)={empty_required_cells}"
    items.append(ScoreItem("table_quality", "표 내부 품질", s5, 10, tw, detail5))

    # 6. 주요문장 강조 적정성 (10)
    bold_p = _count_bold_paragraphs(doc)
    total_p = max(1, _count_nonempty_paragraphs(doc))
    ratio = bold_p / total_p
    if bold_p == 0:
        s6, d6 = 4.0, "강조 없음(핵심문장 미강조)"
    elif ratio > 0.35:
        s6, d6 = 5.0, f"과잉 강조(비율 {ratio:.0%})"
    else:
        s6, d6 = 10.0, f"강조 {bold_p}개(비율 {ratio:.0%}) 적정"
    items.append(ScoreItem("emphasis", "주요문장 강조 적정성", s6, 10,
                           0 if s6 == 10 else 1, d6))

    # 7. 문서 유형별 구조 적합성 (15)
    kws = _TYPE_STRUCTURE_KEYWORDS.get(doc_type, ())
    present = sum(1 for kw in kws if kw in full_text)
    struct_ratio = (present / len(kws)) if kws else 0.7
    s7 = round(15.0 * (0.4 * min(1.0, type_confidence) + 0.6 * struct_ratio), 1)
    items.append(ScoreItem("type_structure", "문서 유형별 구조 적합성", s7, 15,
                           len(kws) - present,
                           f"유형={doc_type}, 구조키워드 {present}/{len(kws) or '-'}, conf={type_confidence:.2f}"))

    # 8. PSST/보고서 구조 충족도 (10)
    if psst_ratio is not None:
        s8 = round(10.0 * psst_ratio, 1)
        d8 = f"PSST 충족 {psst_ratio:.0%}"
    else:
        # 보고서 구조: 현황/진단/개선/기대효과 등 일반 보고 구조 키워드
        report_kw = ("현황", "분석", "개선", "결론", "계획", "기대")
        rp = sum(1 for kw in report_kw if kw in full_text)
        s8 = round(10.0 * (rp / len(report_kw)), 1)
        d8 = f"보고서 구조 키워드 {rp}/{len(report_kw)}"
    items.append(ScoreItem("psst_structure", "PSST/보고서 구조 충족도", s8, 10, 0, d8))

    # 9. 이미지·도식 제안 적정성 (5)
    if image_suggestions > 0 or existing_images > 0:
        s9 = 5.0
        d9 = f"제안 {image_suggestions}건, 기존 이미지 {existing_images}장"
    else:
        s9 = 2.0
        d9 = "도식 제안 없음(시각화 여지 점검 필요)"
    items.append(ScoreItem("image_suggestion", "이미지·도식 제안 적정성", s9, 5,
                           0 if s9 == 5 else 1, d9))

    total = sum(i.score for i in items)
    grade = ("우수" if total >= 90 else "통과" if total >= 85
             else "보완 필요" if total >= 70 else "실패")
    return QualityScore(total=total, grade=grade, passed=total >= 85, items=items)
