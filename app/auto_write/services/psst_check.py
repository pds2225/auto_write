"""psst_check.py

사업계획서 유형 문서의 **PSST 구조 충실도**를 검사한다.

PSST = Problem(문제인식) / Solution(실현가능성) / Scale-up(성장전략) / Team(팀구성)

검사 방식
---------
- 섹션 존재 여부: ``project_service.ProjectService`` 의 PSST 정규식(``PSST_*_RE``)을
  재사용해 양식 섹션 헤더가 있는지 확인(중복 구현 금지).
- 내용 충실도: 각 영역의 핵심 하위 요소를 본문 텍스트에서 키워드로 탐지하고,
  영역별로 [누락 / 미흡 / 적정 / 우수] 4단계 등급을 매긴다.

본 모듈은 결정론적이며 AI 를 호출하지 않는다.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document

# 기존 PSST 섹션 정규식 재사용
from .project_service import ProjectService

# 영역별 핵심 하위 요소와 탐지 키워드 (사용자 정의 PSST 체크리스트)
_PSST_ITEMS: dict[str, list[tuple[str, tuple[str, ...]]]] = {
    "problem": [
        ("고객/시장 문제", ("고객", "시장", "니즈", "수요", "페인", "불편", "문제")),
        ("기존 대안 한계", ("기존", "대안", "한계", "기존방식", "현행", "종래")),
        ("문제 심각성", ("심각", "위험", "리스크", "비용", "손실", "확산")),
        ("수치 근거", ("%", "억", "만", "건", "명", "배", "규모", "통계")),
    ],
    "solution": [
        ("해결방안/핵심기능", ("해결", "솔루션", "핵심기능", "기능", "서비스", "제품")),
        ("차별성", ("차별", "경쟁력", "독창", "우위", "특허", "기술력")),
        ("구현 가능성", ("구현", "개발", "실현", "검증", "시제품", "TRL", "프로토타입")),
        ("고객 적용 시나리오", ("적용", "활용", "시나리오", "사용", "도입", "고객사")),
    ],
    "scale": [
        ("시장규모", ("시장규모", "TAM", "SAM", "SOM", "시장", "성장률")),
        ("수익모델", ("수익", "BM", "비즈니스모델", "과금", "라이선스", "구독")),
        ("판로/성장전략", ("판로", "유통", "마케팅", "성장전략", "확장", "진출")),
        ("KPI/매출계획", ("KPI", "매출", "목표", "계획", "로드맵", "마일스톤")),
    ],
    "team": [
        ("대표자 역량", ("대표", "경력", "역량", "전공", "이력", "창업자")),
        ("팀 구성", ("팀", "구성원", "인력", "조직", "직원", "멤버")),
        ("외부 협력", ("협력", "파트너", "제휴", "자문", "네트워크", "MOU")),
        ("수행 경험/실행력", ("경험", "수행", "실적", "성과", "추진", "실행")),
    ],
}

_AREA_LABELS = {
    "problem": "Problem(문제인식)",
    "solution": "Solution(실현가능성)",
    "scale": "Scale-up(성장전략)",
    "team": "Team(팀구성)",
}


@dataclass
class PSSTAreaResult:
    area: str
    label: str
    section_present: bool
    items_total: int
    items_found: int
    missing_items: list[str] = field(default_factory=list)
    grade: str = "누락"          # 누락 / 미흡 / 적정 / 우수

    def as_dict(self) -> dict[str, Any]:
        return {
            "area": self.area,
            "label": self.label,
            "section_present": self.section_present,
            "items_total": self.items_total,
            "items_found": self.items_found,
            "missing_items": self.missing_items,
            "grade": self.grade,
        }


@dataclass
class PSSTReport:
    applicable: bool
    areas: list[PSSTAreaResult] = field(default_factory=list)
    overall_ratio: float = 0.0   # 전체 충족 항목 비율
    summary: str = ""

    def as_dict(self) -> dict[str, Any]:
        return {
            "applicable": self.applicable,
            "overall_ratio": round(self.overall_ratio, 3),
            "summary": self.summary,
            "areas": [a.as_dict() for a in self.areas],
        }


def _grade(found: int, total: int) -> str:
    if total <= 0 or found == 0:
        return "누락"
    ratio = found / total
    if ratio >= 0.9:
        return "우수"
    if ratio >= 0.6:
        return "적정"
    return "미흡"


def _extract_text(doc: Document, *, limit: int = 30000) -> str:
    parts: list[str] = []
    total = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
            total += len(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
                    total += len(t)
        if total >= limit:
            break
    return "\n".join(parts)


def check_psst(doc: Document) -> PSSTReport:
    """Document 에 대해 PSST 4영역 충실도를 검사한다."""
    text = _extract_text(doc)

    # 섹션 헤더 존재 여부 (기존 정규식 재사용)
    section_flags = {
        "problem": bool(ProjectService.PSST_PROBLEM_RE.search(text)),
        "solution": bool(ProjectService.PSST_SOLUTION_RE.search(text)),
        "scale": bool(ProjectService.PSST_SCALE_RE.search(text)),
        "team": bool(ProjectService.PSST_TEAM_RE.search(text)),
    }

    areas: list[PSSTAreaResult] = []
    total_found = 0
    total_items = 0
    for area, items in _PSST_ITEMS.items():
        found = 0
        missing: list[str] = []
        for item_label, keywords in items:
            if any(kw.lower() in text.lower() for kw in keywords):
                found += 1
            else:
                missing.append(item_label)
        total_found += found
        total_items += len(items)
        areas.append(PSSTAreaResult(
            area=area,
            label=_AREA_LABELS[area],
            section_present=section_flags[area],
            items_total=len(items),
            items_found=found,
            missing_items=missing,
            grade=_grade(found, len(items)),
        ))

    overall = total_found / total_items if total_items else 0.0
    weak = [a.label for a in areas if a.grade in ("누락", "미흡")]
    summary = (
        f"PSST 전체 충족 {total_found}/{total_items} ({overall*100:.0f}%). "
        + ("보완 필요: " + ", ".join(weak) if weak else "전 영역 적정 이상.")
    )
    return PSSTReport(applicable=True, areas=areas, overall_ratio=overall, summary=summary)


def check_psst_docx(path: str | Path) -> PSSTReport:
    return check_psst(Document(str(Path(path))))
