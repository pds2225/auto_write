from __future__ import annotations

import re
import shutil
from pathlib import Path
from typing import Any

from docx import Document

from ..models import GeneratedImage, ProjectInput, TemplateProfile
from .docx_ops import insert_image_after_paragraph, insert_image_in_cell, insert_text_after_paragraph, logical_cells, set_cell_text


class RenderService:
    NON_RENDERABLE_SECTION_RE = re.compile(
        r"(^(?:[◦ㅇ■□●·•]+)$|담\s*당\s*자|사실과 다름이 없음을 확인|추천기관|신청기업|제출목록|증빙서류|평가대상에서 제외|별첨|붙임)",
        re.IGNORECASE,
    )
    NON_RENDERABLE_SLOT_RE = re.compile(
        r"(증빙서류|제출목록|동의서|서약서|확인서|추천기관|신청기업|담\s*당\s*자|평가대상에서 제외|별첨|붙임)",
        re.IGNORECASE,
    )

    @staticmethod
    def _parse_anchor_index(value: Any, default: int = -1) -> int:
        try:
            text = str(value).strip()
            if not text:
                return default
            return int(text)
        except (TypeError, ValueError):
            return default

    def _resolve_table_cell(self, document: Document, table_index: int, row_index: int, cell_index: int, *, logical: bool = False):
        if table_index < 0 or table_index >= len(document.tables):
            return None, f"표 index 범위를 벗어났습니다(table_index={table_index}, tables={len(document.tables)})"
        table = document.tables[table_index]
        if row_index < 0 or row_index >= len(table.rows):
            return None, f"행 index 범위를 벗어났습니다(table_index={table_index}, row={row_index}, rows={len(table.rows)})"
        row = table.rows[row_index]
        # 텍스트 채움(logical=True)은 '논리 셀'(서로 다른 w:tc)로 해석해야 가로 병합 표에서도
        # 분석이 기록한 좌표와 같은 자리에 들어간다. 이미지 슬롯은 grid 인덱스라 logical=False.
        cells = logical_cells(row) if logical else list(row.cells)
        if cell_index < 0 or cell_index >= len(cells):
            return (
                None,
                f"열 index 범위를 벗어났습니다(table_index={table_index}, row={row_index}, cell={cell_index}, cells={len(cells)})",
            )
        return cells[cell_index], ""

    @classmethod
    def _should_skip_section(cls, label: str) -> bool:
        return bool(cls.NON_RENDERABLE_SECTION_RE.search(str(label or "").strip()))

    @classmethod
    def _should_skip_image_slot(cls, slot_label: str) -> bool:
        return bool(cls.NON_RENDERABLE_SLOT_RE.search(str(slot_label or "").strip()))

    def render(
        self,
        profile: TemplateProfile,
        project_input: ProjectInput,
        images: list[GeneratedImage],
        output_path: Path,
        *,
        psst_only: bool = False,
        psst_field_ids: set[str] | None = None,
        core_table_ids: set[str] | None = None,
    ) -> dict[str, Any]:
        source_docx = Path(profile.source_docx)
        if not source_docx.is_file():
            raise ValueError(
                f"템플릿 원본 DOCX를 찾을 수 없습니다: {source_docx.name}. "
                "템플릿 DOCX를 다시 업로드한 뒤 프로젝트를 새로 만드세요."
            )
        output_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source_docx, output_path)
        document = Document(str(output_path))
        answers = project_input.answers
        errors: list[str] = []
        warnings: list[str] = []
        image_by_slot = {image.slot_id: image for image in images}
        image_slot_ids = {slot_id for slot_id, image in image_by_slot.items() if Path(image.path).exists()}
        table_cell_image_slots: dict[tuple[int, int, int], str] = {}
        for slot in profile.image_slots:
            if slot.anchor_type != "table_cell":
                continue
            try:
                key = (
                    int(slot.anchor_ref.get("table_index", -1)),
                    int(slot.anchor_ref.get("row", -1)),
                    int(slot.anchor_ref.get("cell", -1)),
                )
            except (TypeError, ValueError):
                continue
            if key[0] < 0 or key[1] < 0 or key[2] < 0:
                continue
            table_cell_image_slots[key] = slot.slot_id

        section_count = 0
        for section in profile.sections:
            # psst_only는 PSST 핵심 섹션만 렌더링하도록 제한하지만, 그런 섹션이 실제로
            # 감지됐을 때만 적용해야 한다. psst_field_ids가 비어 있으면 템플릿에 PSST
            # 제목이 없다는 뜻이므로, 사용자가 입력한 내용을 누락시키지 않도록 채워진
            # 모든 섹션을 렌더링한다(_restrict_autofill_targets의 동작과 일치).
            if psst_only and psst_field_ids and section.field_id not in psst_field_ids:
                continue
            value = str(answers.get(section.field_id, "") or "").strip()
            if not value:
                continue
            if self._should_skip_section(section.label) or self._should_skip_section(section.anchor_text):
                warnings.append(f"학습 규칙에 따라 비본문 섹션은 건너뛰었습니다: {section.label}")
                continue
            if insert_text_after_paragraph(document, section.anchor_text, value, allow_create_paragraph=True):
                section_count += 1
            else:
                warnings.append(f"문단 anchor를 찾지 못했습니다: {section.anchor_text}")

        cell_count = 0
        for table in profile.tables:
            # 섹션과 동일한 graceful degradation: 핵심 표가 감지되지 않으면
            # 채워진 표 셀을 억제하지 않는다.
            if psst_only and core_table_ids and table.table_id not in core_table_ids:
                continue
            table_index = table.table_index
            for cell in table.cells:
                slot_id = table_cell_image_slots.get((table_index, cell.row, cell.cell))
                # Skip placeholder text for cells that already have a real image file.
                if slot_id and slot_id in image_slot_ids:
                    continue
                value = str(answers.get(cell.cell_id, "") or "").strip()
                if not value:
                    continue
                target_cell, reason = self._resolve_table_cell(
                    document, table_index, cell.row, cell.cell, logical=True
                )
                if target_cell is None:
                    errors.append(f"표 셀 위치 오류: {table.label} r{cell.row} c{cell.cell} ({reason})")
                    continue
                set_cell_text(target_cell, value)
                cell_count += 1

        image_count = 0
        for slot in profile.image_slots:
            generated = image_by_slot.get(slot.slot_id)
            if generated is None:
                continue
            if self._should_skip_image_slot(slot.label):
                warnings.append(f"학습 규칙에 따라 비본문 이미지 슬롯은 건너뛰었습니다: {slot.label}")
                continue
            image_path = Path(generated.path)
            if not image_path.exists():
                errors.append(f"이미지 파일을 찾지 못했습니다: {slot.label} ({image_path.name})")
                continue
            anchor_ref = slot.anchor_ref
            ok = False
            if slot.anchor_type == "table_cell":
                table_index = self._parse_anchor_index(anchor_ref.get("table_index"))
                row_index = self._parse_anchor_index(anchor_ref.get("row"))
                cell_index = self._parse_anchor_index(anchor_ref.get("cell"))
                _, reason = self._resolve_table_cell(document, table_index, row_index, cell_index)
                if reason:
                    errors.append(f"이미지 표 위치 오류: {slot.label} ({reason})")
                    continue
                ok = insert_image_in_cell(
                    document,
                    table_index=table_index,
                    row=row_index,
                    cell_index=cell_index,
                    image_path=image_path,
                    width_cm=float(slot.size_hint.get("width_cm", 12.0)),
                )
            else:
                ok = insert_image_after_paragraph(
                    document,
                    anchor_text=str(anchor_ref.get("anchor_text", "")),
                    image_path=image_path,
                    width_cm=float(slot.size_hint.get("width_cm", 14.0)),
                    insert_offset=int(anchor_ref.get("insert_offset", 2)),
                    allow_create_paragraph=False,
                )
            if ok:
                image_count += 1
            else:
                errors.append(f"이미지를 넣지 못했습니다: {slot.label}")

        document.save(str(output_path))
        return {
            "output_path": str(output_path),
            "sections_written": section_count,
            "cells_written": cell_count,
            "images_written": image_count,
            "errors": errors,
            "warnings": warnings,
        }
