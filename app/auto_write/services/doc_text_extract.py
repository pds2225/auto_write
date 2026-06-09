"""doc_text_extract.py — 공고/양식 파일에서 본문 텍스트를 추출한다.

지원 형식
---------
- ``.txt`` / ``.md``      : 그대로 읽기(utf-8 → cp949 폴백)
- ``.docx``               : python-docx(문단 + 표 셀)
- ``.pdf``                : pypdf
- ``.hwp`` / ``.hwpx`` / ``.doc`` : ``document_ingest.ensure_template_docx`` 로 DOCX 변환 후 추출.
  변환 실패 시 HWP 미리보기 스트림(PrvText, olefile) 으로 폴백한다(일부만 추출될 수 있음).

추출 실패해도 예외를 던지지 않고 ``(텍스트, 안내노트)`` 를 반환한다 — 호출측이
안내노트를 사용자에게 보여주고 "한글/PDF에서 .docx·.txt 로 저장 후 재시도" 를 권할 수 있다.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

_TEXT_EXTS = {".txt", ".md"}
_CONVERT_EXTS = {".hwp", ".hwpx", ".doc"}


def _read_textfile(path: Path) -> str:
    for enc in ("utf-8", "cp949", "euc-kr"):
        try:
            return path.read_text(encoding=enc)
        except Exception:
            continue
    return path.read_bytes().decode("utf-8", errors="replace")


def _docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception:
        return ""
    try:
        reader = PdfReader(str(path))
        return "\n".join((page.extract_text() or "") for page in reader.pages).strip()
    except Exception:
        return ""


def _hwp_prvtext(path: Path) -> str:
    """HWP 의 미리보기 텍스트(PrvText, UTF-16LE)를 추출한다(본문 전체는 아님)."""
    try:
        import olefile
    except Exception:
        return ""
    try:
        if not olefile.isOleFile(str(path)):
            return ""
        ole = olefile.OleFileIO(str(path))
        try:
            if ole.exists("PrvText"):
                data = ole.openstream("PrvText").read()
                return data.decode("utf-16-le", errors="replace").strip()
        finally:
            ole.close()
    except Exception:
        return ""
    return ""


def extract_text(path: str | Path) -> tuple[str, list[str]]:
    """파일에서 본문 텍스트를 추출한다. (text, notes) 를 반환한다."""
    p = Path(path)
    notes: list[str] = []
    if not p.exists():
        return "", [f"파일이 없습니다: {p}"]

    ext = p.suffix.lower()
    if ext in _TEXT_EXTS:
        return _read_textfile(p), notes
    if ext == ".docx":
        try:
            return _docx_text(p), notes
        except Exception as exc:
            return "", [f"DOCX 읽기 실패: {exc}"]
    if ext == ".pdf":
        t = _pdf_text(p)
        if t.strip():
            return t, notes
        notes.append("PDF 텍스트 추출 실패(스캔본일 수 있음) — .docx/.txt 로 저장 후 재시도 권장.")
        return "", notes

    if ext in _CONVERT_EXTS:
        # HWP/HWPX/DOC → DOCX 변환 후 추출(기존 양식 변환기 재사용)
        try:
            from ..document_ingest import ensure_template_docx

            docx_path, conv_notes = ensure_template_docx(p)
            notes.extend(conv_notes)
            t = _docx_text(Path(docx_path))
            if t.strip():
                return t, notes
        except Exception as exc:
            notes.append(f"문서 변환 실패: {exc}")
        if ext == ".hwp":
            t = _hwp_prvtext(p)
            if t.strip():
                notes.append("HWP 미리보기 텍스트(PrvText)만 추출 — 본문 일부가 누락될 수 있음.")
                return t, notes
        notes.append("텍스트 추출 실패 — 한글/오피스에서 .docx 또는 .txt 로 저장 후 재시도 권장.")
        return "", notes

    # 알 수 없는 확장자 → 텍스트로 시도
    try:
        return _read_textfile(p), [f"알 수 없는 형식({ext}) — 텍스트로 읽음."]
    except Exception as exc:
        return "", [f"읽기 실패: {exc}"]
