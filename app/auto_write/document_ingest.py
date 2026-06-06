from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from lxml import etree

TEMPLATE_SUFFIXES = {".docx", ".hwpx", ".hwp"}
REFERENCE_SUFFIXES = {".docx", ".pdf", ".txt", ".md", ".json", ".hwpx", ".hwp"}

HWPX_SECTION_RE = re.compile(r"Contents/section(\d+)\.xml$", re.IGNORECASE)
HWP_PREVIEW_STREAM = ["PrvText"]
HWP_GUIDE_ROW_RE = re.compile(
    r"(※|안내문구|작성요령|작성방법|유의사항|주의사항|참고용|참고|예시|샘플|단위\s*:|삭제\s*후\s*제출|공란|해당\s*시|기재|서식|입력\s*예시|작성\s*예시)",
    re.IGNORECASE,
)
HWPX_TEXT_NS = {
    "hp": "http://www.hancom.co.kr/hwpml/2011/paragraph",
    "hs": "http://www.hancom.co.kr/hwpml/2011/section",
}


def is_supported_template_file(file_name: str) -> bool:
    return Path(str(file_name or "")).suffix.lower() in TEMPLATE_SUFFIXES


def is_supported_reference_file(file_name: str) -> bool:
    return Path(str(file_name or "")).suffix.lower() in REFERENCE_SUFFIXES


def template_upload_detail() -> str:
    return "DOCX, HWPX, HWP 파일만 업로드할 수 있습니다."


def template_accept_value() -> str:
    return ".docx,.hwpx,.hwp"


def reference_accept_value() -> str:
    return ".docx,.hwpx,.hwp,.pdf,.txt,.md,.json"


def ensure_template_docx(source_path: Path) -> tuple[Path, list[str]]:
    suffix = source_path.suffix.lower()
    if suffix == ".docx":
        return source_path, []

    converted_path = source_path.with_name(f"{source_path.stem}_converted.docx")
    if suffix == ".hwpx":
        _convert_hwpx_to_docx(source_path, converted_path)
        return converted_path, ["HWPX 양식을 분석용 DOCX로 변환한 뒤 기존 템플릿 엔진으로 분석했습니다."]
    if suffix == ".hwp":
        try:
            _convert_hwp_to_docx(source_path, converted_path)
            return converted_path, ["HWP 양식을 `unhwp`로 표 구조까지 복원한 뒤 기존 템플릿 엔진으로 분석했습니다."]
        except Exception:
            preview_text = extract_hwp_preview_text(source_path)
            if not preview_text.strip():
                raise ValueError("HWP 파일의 구조화된 내용을 읽지 못했습니다. HWPX 또는 DOCX로 저장한 뒤 다시 올려주세요.")
            _write_text_docx(preview_text, converted_path, title=source_path.stem)
            return converted_path, ["HWP 양식은 미리보기 텍스트로 분석용 DOCX를 만들었습니다. 표 복원은 제한됩니다."]
    raise ValueError(template_upload_detail())


def extract_additional_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".hwpx":
        text = extract_hwpx_preview_text(path)
        return text or extract_hwpx_text(path)
    if suffix == ".hwp":
        text = extract_hwp_markdown_text(path)
        return text or extract_hwp_preview_text(path)
    return ""


def extract_hwpx_preview_text(path: Path) -> str:
    try:
        with zipfile.ZipFile(path, "r") as archive:
            if "Preview/PrvText.txt" not in archive.namelist():
                return ""
            data = archive.read("Preview/PrvText.txt")
        return _normalize_preview_text(data.decode("utf-8", errors="ignore"))
    except Exception:
        return ""


def extract_hwpx_text(path: Path) -> str:
    paragraphs: list[str] = []
    try:
        with zipfile.ZipFile(path, "r") as archive:
            for section_name in _iter_hwpx_section_paths(archive):
                root = etree.fromstring(archive.read(section_name))
                for paragraph in root.findall("hp:p", HWPX_TEXT_NS):
                    text = _extract_hwpx_paragraph_text(paragraph)
                    if text:
                        paragraphs.append(text)
                    if any(_local_name(elem) == "pic" for elem in paragraph.iter()) and not text:
                        paragraphs.append("이미지")
        return "\n".join(paragraphs).strip()
    except Exception:
        return ""


def extract_hwp_preview_text(path: Path) -> str:
    olefile_mod = _load_olefile()
    if olefile_mod is None:
        return ""
    try:
        if not olefile_mod.isOleFile(str(path)):
            return ""
        ole = olefile_mod.OleFileIO(str(path))
        try:
            if not ole.exists(HWP_PREVIEW_STREAM):
                return ""
            data = ole.openstream(HWP_PREVIEW_STREAM).read()
        finally:
            ole.close()
        return _normalize_preview_text(data.decode("utf-16-le", errors="ignore"))
    except Exception:
        return ""


def extract_hwp_markdown_text(path: Path) -> str:
    unhwp_mod = _load_unhwp()
    if unhwp_mod is None:
        return ""
    try:
        markdown = unhwp_mod.to_markdown(str(path))
    except Exception:
        return ""
    if not isinstance(markdown, str):
        return ""
    return _normalize_hwp_markdown_text(markdown)


def _load_unhwp():
    try:
        import unhwp  # type: ignore

        return unhwp
    except Exception:
        return None


def _load_olefile():
    try:
        import olefile  # type: ignore

        return olefile
    except Exception:
        return None


def _iter_hwpx_section_paths(archive: zipfile.ZipFile) -> list[str]:
    entries: list[tuple[int, str]] = []
    for name in archive.namelist():
        match = HWPX_SECTION_RE.fullmatch(name)
        if match:
            entries.append((int(match.group(1)), name))
    return [name for _, name in sorted(entries)]


def _local_name(tag: Any) -> str:
    if not isinstance(tag, str):
        return ""
    if tag.startswith("{"):
        return tag.split("}", 1)[1]
    return tag


def _iter_direct_children(element, local_name: str):
    for child in element:
        if _local_name(getattr(child, "tag", "")) == local_name:
            yield child


def _element_has_ancestor(elem, names: set[str], stop) -> bool:
    current = elem.getparent()
    while current is not None and current is not stop:
        if _local_name(getattr(current, "tag", "")) in names:
            return True
        current = current.getparent()
    return False


def _extract_hwpx_paragraph_text(paragraph) -> str:
    parts: list[str] = []
    for elem in paragraph.iter():
        local_name = _local_name(getattr(elem, "tag", ""))
        if local_name == "lineBreak":
            parts.append("\n")
            continue
        if local_name != "t":
            continue
        if _element_has_ancestor(elem, {"tbl", "tc", "subList"}, paragraph):
            continue
        text = str(elem.text or "")
        if text:
            parts.append(text)
    return _normalize_joined_text(parts)


def _extract_hwpx_cell_text(cell) -> str:
    parts: list[str] = []
    for elem in cell.iter():
        local_name = _local_name(getattr(elem, "tag", ""))
        if local_name == "lineBreak":
            parts.append("\n")
            continue
        if local_name == "t" and elem.text:
            parts.append(str(elem.text))
    text = _normalize_joined_text(parts)
    if text:
        return text
    if any(_local_name(getattr(elem, "tag", "")) == "pic" for elem in cell.iter()):
        return "이미지"
    return ""


def _normalize_joined_text(parts: list[str]) -> str:
    joined = "".join(parts).replace("\r", "")
    joined = re.sub(r"\n{3,}", "\n\n", joined)
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in joined.split("\n")]
    compact = "\n".join(line for line in lines if line)
    return compact.strip()


def _normalize_preview_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in normalized.split("\n")]
    return "\n".join(line for line in lines if line).strip()


def _normalize_hwp_markdown_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    normalized = normalized.replace("<br />", "\n").replace("<br/>", "\n").replace("<br>", "\n")
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in normalized.split("\n")]
    return "\n".join(line for line in lines if line).strip()


def _extract_hwp_json(path: Path) -> dict[str, Any] | None:
    unhwp_mod = _load_unhwp()
    if unhwp_mod is None:
        return None
    try:
        with unhwp_mod.parse(str(path)) as result:
            raw_json = getattr(result, "json", "")
            if not raw_json:
                return None
            data = json.loads(raw_json) if isinstance(raw_json, str) else raw_json
            if isinstance(data, dict):
                return data
    except Exception:
        return None
    return None


def _convert_hwp_to_docx(source_path: Path, output_path: Path) -> None:
    data = _extract_hwp_json(source_path)
    if not data:
        raise ValueError("HWP 파일에서 구조화된 내용을 추출하지 못했습니다.")

    document = Document()
    wrote_anything = False
    for section in data.get("sections", []):
        if not isinstance(section, dict):
            continue
        for item in section.get("content", []):
            if _append_unhwp_item(document, item):
                wrote_anything = True

    if not wrote_anything:
        raise ValueError("HWP 파일에서 변환할 내용을 찾지 못했습니다.")

    document.save(output_path)


def _append_unhwp_item(document: Document, item: Any) -> bool:
    if not isinstance(item, dict):
        return False
    if "Paragraph" in item:
        return _append_unhwp_paragraph(document, item["Paragraph"])
    if "Table" in item:
        return _append_unhwp_table(document, item["Table"])
    if "Text" in item:
        text = _normalize_joined_text([str(item["Text"].get("text", ""))])
        if text:
            document.add_paragraph(text)
            return True
        return False
    if "Image" in item:
        document.add_paragraph("이미지")
        return True
    wrote_anything = False
    for value in item.values():
        if isinstance(value, list):
            for child in value:
                wrote_anything = _append_unhwp_item(document, child) or wrote_anything
        elif isinstance(value, dict):
            wrote_anything = _append_unhwp_item(document, value) or wrote_anything
    return wrote_anything


def _append_unhwp_paragraph(document: Document, paragraph: Any) -> bool:
    if not isinstance(paragraph, dict):
        return False
    parts: list[str] = []
    for item in paragraph.get("content", []):
        text = _extract_unhwp_inline_text(item)
        if text:
            parts.append(text)
    text = _normalize_joined_text(parts)
    if not text:
        return False
    heading_level = 0
    style = paragraph.get("style", {})
    if isinstance(style, dict):
        try:
            heading_level = int(style.get("heading_level") or 0)
        except Exception:
            heading_level = 0
    if heading_level > 0 and "\n" not in text:
        document.add_heading(text, level=min(heading_level, 9))
    else:
        document.add_paragraph(text)
    return True


def _append_unhwp_table(document: Document, table: Any) -> bool:
    if not isinstance(table, dict):
        return False
    segments = _split_unhwp_table_segments(table)
    if not segments:
        return False

    wrote_anything = False
    for kind, payload in segments:
        if kind == "note":
            note_text = str(payload or "").strip()
            if note_text:
                document.add_paragraph(note_text)
                wrote_anything = True
            continue
        if kind == "table" and _render_unhwp_table(document, payload):
            wrote_anything = True
    return wrote_anything


def _extract_unhwp_cell_text(cell: Any) -> str:
    if not isinstance(cell, dict):
        return ""
    parts: list[str] = []
    for item in cell.get("content", []):
        text = _extract_unhwp_block_text(item)
        if text:
            parts.append(text)
    return _normalize_joined_text(["\n".join(parts)])


def _extract_unhwp_block_text(node: Any) -> str:
    if isinstance(node, str):
        if node in {"LineBreak", "lineBreak", "HardBreak", "hardBreak"}:
            return "\n"
        return node
    if isinstance(node, list):
        parts: list[str] = []
        for child in node:
            text = _extract_unhwp_block_text(child)
            if text:
                parts.append(text)
        return _normalize_joined_text(parts)
    if not isinstance(node, dict):
        return ""
    if "Text" in node:
        return str(node["Text"].get("text", ""))
    if "Paragraph" in node:
        return _extract_unhwp_paragraph_text(node["Paragraph"])
    if "Table" in node:
        return _extract_unhwp_table_text(node["Table"])
    if "Image" in node:
        return "이미지"
    if "content" in node and isinstance(node["content"], list):
        parts: list[str] = []
        for child in node["content"]:
            text = _extract_unhwp_block_text(child)
            if text:
                parts.append(text)
        return "\n".join(parts)
    return ""


def _extract_unhwp_inline_text(node: Any) -> str:
    if isinstance(node, str):
        if node in {"LineBreak", "lineBreak", "HardBreak", "hardBreak"}:
            return "\n"
        return node
    if isinstance(node, list):
        parts: list[str] = []
        for child in node:
            text = _extract_unhwp_inline_text(child)
            if text:
                parts.append(text)
        return "".join(parts)
    if not isinstance(node, dict):
        return ""
    if "Text" in node:
        return str(node["Text"].get("text", ""))
    if "Image" in node:
        return "이미지"
    if "Paragraph" in node:
        return _extract_unhwp_paragraph_text(node["Paragraph"])
    if "Table" in node:
        return _extract_unhwp_table_text(node["Table"])
    if "content" in node and isinstance(node["content"], list):
        parts: list[str] = []
        for child in node["content"]:
            text = _extract_unhwp_inline_text(child)
            if text:
                parts.append(text)
        return "".join(parts)
    return ""


def _extract_unhwp_paragraph_text(paragraph: Any) -> str:
    if not isinstance(paragraph, dict):
        return ""
    parts: list[str] = []
    for item in paragraph.get("content", []):
        text = _extract_unhwp_inline_text(item)
        if text:
            parts.append(text)
    return _normalize_joined_text(parts)


def _extract_unhwp_table_text(table: Any) -> str:
    if not isinstance(table, dict):
        return ""
    rows: list[str] = []
    for row in table.get("rows", []):
        if not isinstance(row, dict):
            continue
        row_parts: list[str] = []
        for cell in row.get("cells", []):
            text = _extract_unhwp_cell_text(cell)
            if text:
                row_parts.append(text)
        if row_parts:
            rows.append("\t".join(row_parts))
    return "\n".join(rows)


def _split_unhwp_table_segments(table: Any) -> list[tuple[str, Any]]:
    if not isinstance(table, dict):
        return []
    rows = [row for row in table.get("rows", []) if isinstance(row, dict)]
    if not rows:
        return []
    max_col = max(
        (
            sum(max(1, int((cell or {}).get("colspan") or 1)) for cell in row.get("cells", []) if isinstance(cell, dict))
            for row in rows
        ),
        default=0,
    )
    segments: list[tuple[str, Any]] = []
    current_rows: list[dict[str, Any]] = []
    for row in rows:
        row_texts = [_extract_unhwp_cell_text(cell) for cell in row.get("cells", []) if isinstance(cell, dict)]
        row_text = _normalize_joined_text([text for text in row_texts if text])
        if _is_unhwp_guidance_row(row, row_text, max_col):
            if current_rows:
                segments.append(("table", current_rows))
                current_rows = []
            if row_text:
                segments.append(("note", row_text))
            continue
        current_rows.append(row)
    if current_rows:
        segments.append(("table", current_rows))
    return segments


def _is_unhwp_guidance_row(row: dict[str, Any], row_text: str, max_col: int) -> bool:
    text = re.sub(r"\s+", " ", row_text or "").strip()
    if not text:
        return False
    if HWP_GUIDE_ROW_RE.search(text):
        return True
    cells = row.get("cells", [])
    if len(cells) == 1 and isinstance(cells[0], dict):
        cell = cells[0]
        colspan = max(1, int(cell.get("colspan") or 1))
        rowspan = max(1, int(cell.get("rowspan") or 1))
        if colspan >= max(1, max_col) and rowspan == 1 and HWP_GUIDE_ROW_RE.search(text):
            return True
    return False


def _render_unhwp_table(document: Document, rows: list[dict[str, Any]]) -> bool:
    placements: list[dict[str, Any]] = []
    occupied: set[tuple[int, int]] = set()
    max_row = 0
    max_col = 0

    for row_index, row in enumerate(rows):
        cells = row.get("cells", [])
        current_col = 0
        is_header = bool(row.get("is_header"))
        for cell in cells:
            if not isinstance(cell, dict):
                continue
            while (row_index, current_col) in occupied:
                current_col += 1
            rowspan = max(1, int(cell.get("rowspan") or 1))
            colspan = max(1, int(cell.get("colspan") or 1))
            text = _extract_unhwp_cell_text(cell)
            placements.append(
                {
                    "row": row_index,
                    "col": current_col,
                    "rowspan": rowspan,
                    "colspan": colspan,
                    "text": text,
                    "header": is_header,
                }
            )
            for delta_row in range(rowspan):
                for delta_col in range(colspan):
                    occupied.add((row_index + delta_row, current_col + delta_col))
            max_row = max(max_row, row_index + rowspan)
            max_col = max(max_col, current_col + colspan)
            current_col += colspan

    if max_row <= 0 or max_col <= 0:
        return False

    doc_table = document.add_table(rows=max_row, cols=max_col)
    for cell_info in placements:
        row_index = int(cell_info["row"])
        col_index = int(cell_info["col"])
        text = str(cell_info["text"])
        _set_docx_cell_text(doc_table.cell(row_index, col_index), text, bold=bool(cell_info.get("header")))
    for cell_info in placements:
        row_index = int(cell_info["row"])
        col_index = int(cell_info["col"])
        rowspan = int(cell_info["rowspan"])
        colspan = int(cell_info["colspan"])
        if rowspan <= 1 and colspan <= 1:
            continue
        start = doc_table.cell(row_index, col_index)
        end = doc_table.cell(row_index + rowspan - 1, col_index + colspan - 1)
        if start is end:
            continue
        try:
            start.merge(end)
        except Exception:
            continue
    return True


def _set_docx_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold


def _convert_hwpx_to_docx(source_path: Path, output_path: Path) -> None:
    document = Document()
    try:
        with zipfile.ZipFile(source_path, "r") as archive:
            for section_name in _iter_hwpx_section_paths(archive):
                root = etree.fromstring(archive.read(section_name))
                for paragraph in root.findall("hp:p", HWPX_TEXT_NS):
                    paragraph_text = _extract_hwpx_paragraph_text(paragraph)
                    tables = [elem for elem in paragraph.iter() if _local_name(getattr(elem, "tag", "")) == "tbl"]
                    has_picture = any(_local_name(getattr(elem, "tag", "")) == "pic" for elem in paragraph.iter())
                    if paragraph_text:
                        document.add_paragraph(paragraph_text)
                    for table_elem in tables:
                        _append_hwpx_table(document, table_elem)
                    if not paragraph_text and not tables and has_picture:
                        document.add_paragraph("이미지")
    except Exception as exc:
        preview_text = extract_hwpx_preview_text(source_path)
        if preview_text.strip():
            _write_text_docx(preview_text, output_path, title=source_path.stem)
            return
        raise ValueError(f"HWPX 파일 변환에 실패했습니다: {exc}") from exc

    if not any(paragraph.text.strip() for paragraph in document.paragraphs) and not document.tables:
        preview_text = extract_hwpx_preview_text(source_path)
        if preview_text.strip():
            _write_text_docx(preview_text, output_path, title=source_path.stem)
            return
    document.save(output_path)


def _append_hwpx_table(document: Document, table_elem) -> None:
    cells: list[dict[str, int | str]] = []
    max_row = 0
    max_col = 0
    current_row_index = 0
    for row_elem in _iter_direct_children(table_elem, "tr"):
        current_col_index = 0
        for cell_elem in _iter_direct_children(row_elem, "tc"):
            cell_addr = next(_iter_direct_children(cell_elem, "cellAddr"), None)
            cell_span = next(_iter_direct_children(cell_elem, "cellSpan"), None)
            row_index = int(str(cell_addr.get("rowAddr", current_row_index))) if cell_addr is not None else current_row_index
            col_index = int(str(cell_addr.get("colAddr", current_col_index))) if cell_addr is not None else current_col_index
            row_span = int(str(cell_span.get("rowSpan", 1))) if cell_span is not None else 1
            col_span = int(str(cell_span.get("colSpan", 1))) if cell_span is not None else 1
            text = _extract_hwpx_cell_text(cell_elem)
            cells.append(
                {
                    "row": row_index,
                    "col": col_index,
                    "row_span": row_span,
                    "col_span": col_span,
                    "text": text,
                }
            )
            max_row = max(max_row, row_index + row_span)
            max_col = max(max_col, col_index + col_span)
            current_col_index = col_index + col_span
        current_row_index += 1

    if max_row <= 0 or max_col <= 0:
        return

    table = document.add_table(rows=max_row, cols=max_col)
    for cell_info in cells:
        row_index = int(cell_info["row"])
        col_index = int(cell_info["col"])
        table.cell(row_index, col_index).text = str(cell_info["text"])
    for cell_info in cells:
        row_index = int(cell_info["row"])
        col_index = int(cell_info["col"])
        row_span = int(cell_info["row_span"])
        col_span = int(cell_info["col_span"])
        if row_span <= 1 and col_span <= 1:
            continue
        start = table.cell(row_index, col_index)
        end = table.cell(row_index + row_span - 1, col_index + col_span - 1)
        if start is end:
            continue
        try:
            start.merge(end)
        except Exception:
            continue


def _write_text_docx(text: str, output_path: Path, title: str = "") -> None:
    document = Document()
    normalized = _normalize_preview_text(text)
    if title:
        document.add_paragraph(title)
    for line in normalized.split("\n"):
        clean = line.strip()
        if clean:
            document.add_paragraph(clean)
    document.save(output_path)
