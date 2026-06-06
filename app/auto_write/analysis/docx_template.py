from __future__ import annotations

import re
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from lxml import etree

from ..models import (
    ImageSlotProfile,
    QuestionProfile,
    SectionProfile,
    TableCellProfile,
    TableProfile,
    TemplateProfile,
)
from ..utils import short_id, slugify

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": WNS}


def _w(tag: str) -> str:
    return f"{{{WNS}}}{tag}"


SECTION_RE = re.compile(r"^((\d+[\.\)])|([IVX]+\.)|([가-하]\.)|([A-Z]\.)|([①-⑮]))\s*.+")
IMAGE_HINT_RE = re.compile(r"(이미지|사진|도식|그림|chart|graph|diagram)", re.IGNORECASE)
GUIDE_HINT_RE = re.compile(
    r"(필요 시|작성|기재|예시|단위:|주의|해당|만 해당|붙임|서식|입력|작성요령|작성방법|참고용)",
    re.IGNORECASE,
)
MARKER_ONLY_RE = re.compile(r"^[◦ㅇ■□●·•\-\u2013\u2014\s]+$")
ADMIN_SECTION_RE = re.compile(
    r"(개인정보|수집.?이용 동의|제3자 제공|동의서|서약서|확인서|체크리스트|제출서류|첨부서류|별지|별첨|붙임|서명|날인|귀하|주민등록|생년월일|휴대전화|이메일|전화번호|주소)",
    re.IGNORECASE,
)
LEGAL_SENTENCE_RE = re.compile(r"(개인정보보호법|제\d+조|동의를 거부할 수 있으나|선정 결과 통보일|필수사항입니다)")
PLACEHOLDER_HINT_RE = re.compile(r"(<[^>]+>|○○○|OOO|000)")
IMPROVABLE_CELL_VALUE_RE = re.compile(
    r"(○{2,}|O{2,}(?!penai)|0{2,}백만|00\.00|00백만|OO학|○○기업|"
    r"예정\(['\u2019]?0{2}|완료\(['\u2019]?0{2}|\.{3,}|…{2,}|^00$|에정\b|"
    r"S\s*/\s*W|S/W)",
    re.IGNORECASE,
)
REQUIRED_SECTION_RE = re.compile(r"^((\d+[\.\)])|([가-하]\.)|(□)|([①-⑮]))\s*.+")
REQUIRED_SECTION_KEYWORD_RE = re.compile(
    r"(개요|필요성|문제|목표|전략|추진|시장|수익|예산|인력|성과|기대효과|실행|사업화|활용)"
)
FORM_GUIDE_RE = re.compile(r"(※|<[^>]+>|기재|작성요령|작성방법|양식|붙임|별첨|단위:)")
QUESTION_EXCLUDE_RE = re.compile(
    r"(작성요령|작성방법|유의사항|기재요령|예시문|체크.?박스|동의함|동의하지 않음|년\s*월\s*일|서명|날인|귀하)",
    re.IGNORECASE,
)
NON_CONTENT_SECTION_RE = re.compile(
    r"(협약해지|제재|환수|집행 유의사항|상기\s*제재|위원회\s*심의|증빙서류|계좌이체|별첨|붙임)",
    re.IGNORECASE,
)
SIGNATURE_HINT_RE = re.compile(
    r"(사실과 다름이 없음을 확인|기업 추천합니다|추천기관|신청기업|담\s*당\s*자|직인|\(인\)|평가대상에서 제외|제출목록 안내)",
    re.IGNORECASE,
)
BUSINESS_CONTEXT_KEYWORD_RE = re.compile(
    r"(개요|아이템|문제|필요성|실현|해결|시장|경쟁|전략|계획|로드맵|자금|성과|인프라|역량|조직|팀|고객|사업화|기술|차별성|성장)"
)
GENERIC_TABLE_LABEL_RE = re.compile(
    r"^(구분|구\s*분|비목|순번|명\s*칭|아이템명|참가구분|기업명|추천기관명|협약기간|수상일|투자유치일)$"
)
GENERIC_IMAGE_LABEL_RE = re.compile(r"^(관련\s*)?(이미지|사진|도식|그림)$", re.IGNORECASE)
TABLE_LABEL_HINT_RE = re.compile(
    r"(신청 분야|신청 제외 대상|자가진단|과제내용|일반현황|기업 구성|재직 인력 고용현황|추가 인력 고용계획|"
    r"업무파트너|성과공유제|사업비|지식재산권|아이템|개요|현황|계획|협업 이력|투자유치|목표시장|"
    r"개발 동기|개발 방법|고용현황|수혜 여부|창업아이템|지원사업|사업계획)",
    re.IGNORECASE,
)
GENERIC_CELL_PART_RE = re.compile(r"^(순번|구분|내용|상세 내용|항목|기타|비고|해당 여부)$")
NOISE_CELL_PART_RE = re.compile(r"^(\.{2,}|…+|[①-⑮]|[0-9]+|[YyNn])$")
EXCLUDED_SECTION_KEYWORDS = [
    "동의서",
    "서약서",
    "확인서",
    "개인정보",
    "청렴",
    "보안",
    "서명란",
    "작성요령",
    "별첨",
    "붙임",
    "이탤릭",
    "파란색 안내",
]


def para_text(p: etree._Element) -> str:
    return "".join(r.text or "" for r in p.iter(_w("t"))).strip()


def _normalize_heading_text(text: str) -> str:
    compact = re.sub(r"\s+", " ", text or "").strip()
    candidate = compact
    for _ in range(2):
        if len(candidate) < 20 or len(candidate) % 2 != 0:
            break
        half = len(candidate) // 2
        left = candidate[:half].strip()
        right = candidate[half:].strip()
        if left and left == right:
            candidate = left
            continue
        break
    return candidate


def _looks_like_marker_only(text: str) -> bool:
    stripped = re.sub(r"\s+", " ", text or "").strip()
    if not stripped:
        return True
    return bool(MARKER_ONLY_RE.fullmatch(stripped))


def _is_signature_or_confirmation_text(text: str) -> bool:
    stripped = re.sub(r"\s+", " ", text or "").strip()
    if not stripped:
        return False
    return bool(SIGNATURE_HINT_RE.search(stripped))


def _is_generic_table_label(text: str) -> bool:
    stripped = re.sub(r"\s+", " ", text or "").strip()
    if not stripped:
        return False
    return bool(GENERIC_TABLE_LABEL_RE.fullmatch(stripped))


def _is_generic_image_label(text: str) -> bool:
    stripped = re.sub(r"\s+", " ", text or "").strip()
    if not stripped:
        return False
    return bool(GENERIC_IMAGE_LABEL_RE.fullmatch(stripped))


def _contextualize_label(base_label: str, context_heading: str, fallback: str) -> str:
    label = re.sub(r"\s+", " ", base_label or "").strip()
    context = re.sub(r"\s+", " ", context_heading or "").strip()
    if not label:
        return context or fallback
    if context and context != label and (
        _looks_like_marker_only(label)
        or _is_generic_table_label(label)
        or _is_generic_cell_fragment(label)
    ):
        return f"{context} / {label}"
    return label


def _clean_label_fragment(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", text or "").strip()
    if not cleaned:
        return ""
    cleaned = cleaned.replace("\r", " ").replace("\n", " ")
    cleaned = re.sub(r"^[◦□●■•·\-\*]+\s*", "", cleaned)
    cleaned = re.sub(r"^(?:\d+[\.\)]\s*)+", "", cleaned)
    cleaned = re.sub(r"^(?:[①-⑮]\s*)+", "", cleaned)
    cleaned = re.sub(r"\s{2,}", " ", cleaned).strip(" /")
    parts = [part.strip() for part in cleaned.split("/") if part.strip()]
    if len(parts) > 1:
        deduped: list[str] = []
        for part in parts:
            if part not in deduped:
                deduped.append(part)
        cleaned = " / ".join(deduped)
    return cleaned.strip()


def _is_noise_fragment(text: str) -> bool:
    cleaned = _clean_label_fragment(text)
    if not cleaned:
        return True
    if _looks_like_marker_only(cleaned):
        return True
    if NOISE_CELL_PART_RE.fullmatch(cleaned):
        return True
    return False


def _is_generic_cell_fragment(text: str) -> bool:
    cleaned = _clean_label_fragment(text)
    if not cleaned:
        return True
    if GENERIC_CELL_PART_RE.fullmatch(cleaned):
        return True
    if _is_noise_fragment(cleaned):
        return True
    return False


def _is_booleanish_fragment(text: str) -> bool:
    cleaned = _clean_label_fragment(text)
    return cleaned in {"Y", "N", "y", "n", "O", "X"}


def _is_numericish_fragment(text: str) -> bool:
    cleaned = _clean_label_fragment(text)
    if not cleaned:
        return False
    return bool(re.fullmatch(r"[0-9]+(?:[-.][0-9]+)?", cleaned) or re.fullmatch(r"[①-⑮]", cleaned))


def _score_table_label_candidate(text: str, context_heading: str) -> int:
    cleaned = _clean_label_fragment(text)
    if not cleaned:
        return -100
    score = len(cleaned)
    if SECTION_RE.match(cleaned) or REQUIRED_SECTION_RE.match(cleaned):
        score += 8
    if TABLE_LABEL_HINT_RE.search(cleaned):
        score += 6
    if cleaned == _clean_label_fragment(context_heading):
        score += 1
    if _is_generic_table_label(cleaned):
        score -= 6
    if _is_noise_fragment(cleaned):
        score -= 8
    if PLACEHOLDER_HINT_RE.search(cleaned) or re.search(r"(?:○|O{2,}|0{2,})", cleaned):
        score -= 5
    if cleaned.endswith("?") or cleaned.endswith("？") or re.search(r"(나요|인가요)\?", cleaned):
        score -= 8
    if "?" in cleaned or "？" in cleaned:
        score -= 12
        if len(cleaned) > 20:
            score -= 20
    if "," in cleaned or "，" in cleaned:
        score -= 4
    word_count = len([part for part in re.split(r"\s+", cleaned) if part])
    if word_count > 8:
        score -= 4
    elif word_count > 5:
        score -= 2
    if len(cleaned) <= 2:
        score -= 2
    if len(cleaned) > 48:
        score -= 2
    elif len(cleaned) > 30:
        score -= 3
    return score


def _pick_table_label(rows: list[list[str]], context_heading: str, fallback: str) -> str:
    cleaned_context = _clean_label_fragment(context_heading)
    for row in rows[:4]:
        row_candidates: list[str] = []
        for cell in row:
            cleaned = _clean_label_fragment(cell)
            if cleaned and not _is_noise_fragment(cleaned):
                row_candidates.append(cleaned)
        if not row_candidates:
            continue

        informative = [candidate for candidate in row_candidates if not _is_generic_cell_fragment(candidate)]
        base = informative[0] if informative else row_candidates[0]
        if _is_noise_fragment(base) or _is_numericish_fragment(base) or _is_booleanish_fragment(base):
            continue
        if cleaned_context:
            return _contextualize_label(base, cleaned_context, fallback)
        return base
    return cleaned_context or fallback


def _build_cell_label(table_label: str, row_header: str, col_header: str, fallback: str) -> str:
    parts: list[str] = []
    for raw_part in (row_header, col_header):
        cleaned = _clean_label_fragment(raw_part)
        if not cleaned or _is_noise_fragment(cleaned):
            continue
        parts.append(cleaned)

    deduped: list[str] = []
    for part in parts:
        if part not in deduped:
            deduped.append(part)

    informative = [part for part in deduped if not _is_generic_cell_fragment(part)]
    if len(informative) >= 2:
        return " / ".join(informative)
    if len(informative) == 1:
        only = informative[0]
        if _is_booleanish_fragment(only):
            label = _clean_label_fragment(table_label)
            if label:
                return f"{label} / {only}"
        if _is_numericish_fragment(only):
            label = _clean_label_fragment(table_label)
            return label or fallback
        return only
    if deduped:
        only = deduped[0]
        if _is_booleanish_fragment(only):
            label = _clean_label_fragment(table_label)
            if label:
                return f"{label} / {only}"
        if _is_numericish_fragment(only):
            label = _clean_label_fragment(table_label)
            return label or fallback
        return only
    label = _clean_label_fragment(table_label)
    return label or fallback


def _is_context_heading_candidate(text: str) -> bool:
    normalized = _normalize_heading_text(text)
    if not normalized:
        return False
    if _looks_like_marker_only(normalized):
        return False
    if _is_guide_text(normalized):
        return False
    if _is_administrative_text(normalized):
        return False
    if _is_signature_or_confirmation_text(normalized):
        return False
    if SECTION_RE.match(normalized) or REQUIRED_SECTION_RE.match(normalized):
        return True
    if BUSINESS_CONTEXT_KEYWORD_RE.search(normalized) and len(normalized) <= 80:
        return True
    return len(normalized) <= 40 and len(normalized.split()) <= 8


def _is_guide_text(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return True
    if _looks_like_marker_only(stripped):
        return True
    if stripped.startswith(("*", "-", "※")):
        return True
    if PLACEHOLDER_HINT_RE.search(stripped):
        return True
    return bool(GUIDE_HINT_RE.search(stripped))


def _is_required_section(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    if _is_guide_text(stripped):
        return False
    if _is_signature_or_confirmation_text(stripped):
        return False
    if REQUIRED_SECTION_RE.match(stripped):
        return True
    if REQUIRED_SECTION_KEYWORD_RE.search(stripped) and len(stripped) <= 80:
        return True
    return False


def _is_excluded_section(text: str) -> bool:
    stripped = re.sub(r"\s+", " ", text or "").strip().lower()
    if not stripped:
        return False
    if _looks_like_marker_only(stripped):
        return True
    if _is_signature_or_confirmation_text(stripped):
        return True
    return any(keyword.lower() in stripped for keyword in EXCLUDED_SECTION_KEYWORDS)


def _is_administrative_text(text: str) -> bool:
    stripped = re.sub(r"\s+", " ", text or "").strip()
    if not stripped:
        return True
    if _looks_like_marker_only(stripped):
        return True
    if _is_excluded_section(stripped):
        return True
    if ADMIN_SECTION_RE.search(stripped):
        return True
    if _is_signature_or_confirmation_text(stripped):
        return True
    if LEGAL_SENTENCE_RE.search(stripped) and len(stripped) >= 40:
        return True
    return False


def _is_business_question_label(text: str) -> bool:
    stripped = re.sub(r"\s+", " ", text or "").strip()
    if not stripped:
        return False
    if _looks_like_marker_only(stripped):
        return False
    if _is_guide_text(stripped):
        return False
    if _is_administrative_text(stripped):
        return False
    if _is_signature_or_confirmation_text(stripped):
        return False
    if NON_CONTENT_SECTION_RE.search(stripped):
        return False
    if FORM_GUIDE_RE.search(stripped):
        return False
    if QUESTION_EXCLUDE_RE.search(stripped):
        return False
    if re.match(r"^\(.*단위.*\)$", stripped):
        return False
    if re.match(r"^\(총\s*사업비.*\)$", stripped):
        return False
    if stripped.count("/") >= 2 and ("동의서" in stripped or "개인정보" in stripped):
        return False
    if len(stripped) >= 180:
        return False
    return True


def read_docx_xml(docx_path: Path) -> etree._Element:
    with zipfile.ZipFile(docx_path, "r") as archive:
        xml_bytes = archive.read("word/document.xml")
    return etree.fromstring(xml_bytes)


def build_doc_summary(docx_path: Path) -> dict[str, Any]:
    root = read_docx_xml(docx_path)
    body = root.find(_w("body"), NS)
    summary: dict[str, Any] = {"paragraphs": [], "tables": []}
    table_index = 0
    current_context_heading = ""
    for element in body:
        if element.tag == _w("p"):
            text = para_text(element)
            if text:
                summary["paragraphs"].append(text)
                normalized = _normalize_heading_text(text)
                if _is_context_heading_candidate(normalized):
                    current_context_heading = normalized
        elif element.tag == _w("tbl"):
            rows: list[list[str]] = []
            for row in element.findall(_w("tr"), NS):
                row_values: list[str] = []
                for cell in row.findall(_w("tc"), NS):
                    cell_value = " / ".join(
                        paragraph
                        for paragraph in (para_text(p) for p in cell.findall(_w("p"), NS))
                        if paragraph
                    )
                    row_values.append(cell_value)
                rows.append(row_values)
            summary["tables"].append(
                {
                    "table_index": table_index,
                    "rows": rows,
                    "context_heading": current_context_heading,
                }
            )
            table_index += 1
    return summary


def cell_value_needs_improvement(value: str) -> bool:
    stripped = re.sub(r"\s+", " ", str(value or "")).strip()
    if not stripped:
        return False
    if PLACEHOLDER_HINT_RE.search(stripped):
        return True
    if IMPROVABLE_CELL_VALUE_RE.search(stripped):
        return True
    if re.fullmatch(r"0{2,}", stripped):
        return True
    return False


def _is_major_section_anchor(text: str) -> bool:
    normalized = _normalize_heading_text(text)
    if not normalized or _looks_like_marker_only(normalized):
        return False
    if normalized.startswith("○") or normalized.startswith("•"):
        return False
    if SECTION_RE.match(normalized) or normalized.startswith("□"):
        return True
    if re.match(r"^[①-⑨]", normalized):
        return True
    if re.match(r"^\d+\.\s", normalized):
        return True
    return False


def detect_sections(paragraphs: list[str], *, strict_headings: bool = False) -> list[SectionProfile]:
    sections: list[SectionProfile] = []
    seen_anchors: set[str] = set()
    for index, text in enumerate(paragraphs):
        normalized = _normalize_heading_text(text)
        if normalized in seen_anchors:
            continue
        if _looks_like_marker_only(normalized):
            continue
        if _is_signature_or_confirmation_text(normalized):
            continue
        if strict_headings:
            is_heading = _is_major_section_anchor(normalized)
        else:
            is_heading = (
                SECTION_RE.match(normalized)
                or normalized.startswith("□")
                or normalized.endswith(":")
                or normalized.endswith("]")
                or (len(normalized) <= 48 and len(normalized.split()) <= 8)
            )
        if not is_heading:
            continue
        seen_anchors.add(normalized)
        field_id = f"section_{index + 1:03d}_{slugify(normalized, 'section')}"
        sections.append(
            SectionProfile(
                field_id=field_id,
                label=normalized,
                anchor_text=normalized,
                required=_is_required_section(normalized),
                is_excluded=_is_excluded_section(normalized),
            )
        )
    return sections


def _nearest_row_header(rows: list[list[str]], row_index: int, col_index: int) -> str:
    if row_index < len(rows):
        for left in range(col_index - 1, -1, -1):
            if left >= len(rows[row_index]):
                continue
            value = (rows[row_index][left] or "").strip()
            if value and not _is_noise_fragment(value):
                return value
    return ""


def _nearest_col_header(rows: list[list[str]], row_index: int, col_index: int) -> str:
    for up in range(row_index - 1, -1, -1):
        if col_index < len(rows[up]):
            value = (rows[up][col_index] or "").strip()
            if value and not _is_noise_fragment(value):
                return value
    return ""


def detect_tables(
    table_data: list[dict[str, Any]],
    *,
    include_improvable_filled: bool = False,
) -> list[TableProfile]:
    tables: list[TableProfile] = []
    for index, info in enumerate(table_data):
        rows = info["rows"]
        context_heading = str(info.get("context_heading", "") or "").strip()
        if not rows:
            continue
        row_count = len(rows)
        col_count = max((len(row) for row in rows), default=0)
        label = _pick_table_label(rows, context_heading, f"표 {index + 1}")

        cells: list[TableCellProfile] = []
        for row_index, row in enumerate(rows):
            for col_index in range(col_count):
                value = row[col_index].strip() if col_index < len(row) else ""
                if value:
                    if not (include_improvable_filled and cell_value_needs_improvement(value)):
                        continue
                row_header = _nearest_row_header(rows, row_index, col_index)
                col_header = _nearest_col_header(rows, row_index, col_index)
                if not row_header and not col_header:
                    continue
                cell_label = _build_cell_label(label, row_header, col_header, f"표 {index + 1}")
                if not cell_label.strip():
                    continue
                field_type = "textarea" if row_count <= 6 and col_count <= 4 else "text"
                cells.append(
                    TableCellProfile(
                        cell_id=f"table_{index + 1:02d}_r{row_index}_c{col_index}",
                        label=cell_label,
                        row=row_index,
                        cell=col_index,
                        row_header=row_header,
                        col_header=col_header,
                        field_type=field_type,
                        required=bool(value and cell_value_needs_improvement(value)),
                    )
                )
        anchors = [anchor for anchor in (context_heading, label) if anchor]
        tables.append(
            TableProfile(
                table_id=f"table_{index + 1:02d}",
                label=label,
                table_index=index,
                anchors=anchors,
                row_count=row_count,
                col_count=col_count,
                cells=cells,
            )
        )
    return tables


def detect_image_slots(document: Document, table_data: list[dict[str, Any]] | None = None) -> list[ImageSlotProfile]:
    image_slots: list[ImageSlotProfile] = []
    seen_labels: set[tuple[int, str]] = set()
    table_contexts: dict[int, str] = {}
    for index, info in enumerate(table_data or []):
        table_index = int(info.get("table_index", index))
        context_heading = str(info.get("context_heading", "") or "").strip()
        if not context_heading:
            rows = info.get("rows", [])
            for row in rows[:3]:
                context_heading = next((re.sub(r"\s+", " ", cell or "").strip() for cell in row if str(cell or "").strip()), "")
                if context_heading:
                    break
        table_contexts[table_index] = context_heading

    def extract_image_label(raw_text: str, context_heading: str) -> str:
        lines = [re.sub(r"\s+", " ", line).strip() for line in raw_text.splitlines() if line.strip()]
        for line in lines:
            if IMAGE_HINT_RE.search(line):
                candidate = line[:120]
                if context_heading and _is_generic_image_label(candidate):
                    return f"{context_heading} 이미지"
                return candidate
        compact = re.sub(r"\s+", " ", raw_text).strip()
        if compact:
            if context_heading and _is_generic_image_label(compact):
                return f"{context_heading} 이미지"
            return compact[:120]
        if context_heading:
            return f"{context_heading} 이미지"
        return "관련 이미지"

    def is_required_image_slot(label: str) -> bool:
        lowered = label.lower()
        optional_tokens = ("예시", "참고", "필요 시", "제목", "삽입", "설계도", "sample")
        if any(token in lowered for token in optional_tokens):
            return False
        if "<" in label and ">" in label:
            return False
        required_tokens = ("관련 이미지", "필수", "첨부")
        if any(token in lowered for token in required_tokens):
            return True
        return not _is_guide_text(label)

    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                text = cell.text.strip()
                if IMAGE_HINT_RE.search(text):
                    context_heading = table_contexts.get(table_index, "")
                    label = extract_image_label(text, context_heading)
                    key = (table_index, re.sub(r"\s+", " ", label).strip().lower())
                    if key in seen_labels:
                        continue
                    seen_labels.add(key)
                    image_slots.append(
                        ImageSlotProfile(
                            slot_id=f"img_table_{table_index}_{row_index}_{cell_index}",
                            label=label or f"표 {table_index + 1} 이미지",
                            required=is_required_image_slot(label),
                            anchor_type="table_cell",
                            anchor_ref={"table_index": table_index, "row": row_index, "cell": cell_index},
                        )
                    )
    return image_slots


def build_questions(
    sections: list[SectionProfile], tables: list[TableProfile], image_slots: list[ImageSlotProfile]
) -> list[QuestionProfile]:
    questions: list[QuestionProfile] = [
        QuestionProfile(
            question_id="project_title",
            label="과제명",
            field_type="text",
            required=True,
            source_hint="문서 제목 또는 사업명",
            target={"kind": "project_meta", "key": "project_title"},
        ),
        QuestionProfile(
            question_id="organization_name",
            label="기관명 또는 회사명",
            field_type="text",
            required=True,
            source_hint="신청 주체 이름",
            target={"kind": "organization_profile", "key": "name"},
        ),
    ]
    for section in sections:
        if section.is_excluded:
            continue
        questions.append(
            QuestionProfile(
                question_id=section.field_id,
                label=section.label,
                field_type="textarea",
                required=section.required,
                source_hint=f"문단 기준: {section.anchor_text}",
                target={"kind": "section", "field_id": section.field_id},
            )
        )
    for table in tables:
        for cell in table.cells:
            questions.append(
                QuestionProfile(
                    question_id=cell.cell_id,
                    label=cell.label,
                    field_type=cell.field_type,
                    required=cell.required,
                    source_hint=f"{table.label} 표의 빈 칸 채우기",
                    target={"kind": "table_cell", "table_id": table.table_id, "cell_id": cell.cell_id},
                )
            )
    for image_slot in image_slots:
        questions.append(
            QuestionProfile(
                question_id=f"image_note_{image_slot.slot_id}",
                label=f"{image_slot.label} 설명",
                field_type="textarea",
                required=False,
                source_hint="이미지 생성에 쓸 핵심 설명이나 강조 포인트",
                target={"kind": "image_slot_note", "slot_id": image_slot.slot_id},
            )
        )
    return questions


def analyze_template(docx_path: Path) -> TemplateProfile:
    template_id = short_id("tpl")
    summary = build_doc_summary(docx_path)
    document = Document(str(docx_path))
    is_completed_doc = _is_completed_document(summary)
    if is_completed_doc:
        sections = detect_sections(summary["paragraphs"], strict_headings=True)
        tables = detect_tables(summary["tables"], include_improvable_filled=True)
    else:
        sections = detect_sections(summary["paragraphs"])
        tables = detect_tables(summary["tables"])
    image_slots = detect_image_slots(document, summary["tables"])
    if is_completed_doc:
        sections = [
            section.model_copy(update={"required": _is_required_section(section.label), "source": "ai"})
            for section in sections
        ]
        tables = [
            table.model_copy(
                update={
                    "cells": [
                        cell.model_copy(update={"required": cell.required or cell_value_needs_improvement(cell.label)})
                        for cell in table.cells
                    ],
                }
            )
            for table in tables
        ]
        image_slots = [slot.model_copy(update={"required": False}) for slot in image_slots]
    if not image_slots:
        for section in sections[:3]:
            image_slots.append(
                ImageSlotProfile(
                    slot_id=f"img_after_{slugify(section.anchor_text, 'image')}",
                    label=f"{section.label} 설명 이미지",
                    required=False,
                    anchor_type="after_paragraph",
                    anchor_ref={"anchor_text": section.anchor_text, "insert_offset": 2},
                    source="suggested",
                )
            )
    questions = build_questions(sections, tables, image_slots)
    notes = [
        "자동 분석으로 문단 anchor, 표 빈 칸, 이미지 위치를 추정했습니다.",
        "애매한 항목은 템플릿 상세 화면에서 JSON으로 직접 수정할 수 있습니다.",
    ]
    profile = TemplateProfile(
        template_id=template_id,
        template_name=docx_path.name,
        source_docx=str(docx_path),
        sections=sections,
        tables=tables,
        image_slots=image_slots,
        questions=questions,
        analysis_notes=notes,
    )
    return sanitize_template_profile(profile)


def _is_completed_document(summary: dict[str, Any]) -> bool:
    paragraphs = [text for text in summary.get("paragraphs", []) if str(text).strip()]
    tables = summary.get("tables", [])
    guide_hits = sum(1 for text in paragraphs if FORM_GUIDE_RE.search(str(text)))
    paragraph_chars = sum(len(str(text).strip()) for text in paragraphs)
    total_cells = 0
    empty_cells = 0
    for table in tables:
        for row in table.get("rows", []):
            for cell in row:
                total_cells += 1
                if not str(cell).strip():
                    empty_cells += 1
    empty_ratio = (empty_cells / total_cells) if total_cells else 0.0
    long_paragraphs = sum(1 for text in paragraphs if len(str(text).strip()) >= 80)
    paragraph_count = len(paragraphs)
    guide_ratio = (guide_hits / paragraph_count) if paragraph_count else 0.0
    # Completed docs usually have substantial narrative and fewer template guide markers/blanks.
    if long_paragraphs >= 8 and empty_ratio <= 0.18 and guide_hits <= 12:
        return True
    if paragraph_count >= 40 and paragraph_chars >= 1800 and empty_ratio <= 0.12 and guide_ratio <= 0.25:
        return True
    return False


def sanitize_template_profile(profile: TemplateProfile) -> TemplateProfile:
    source_docx = Path(profile.source_docx)
    is_partial_source = False
    if source_docx.exists():
        try:
            is_partial_source = _is_completed_document(build_doc_summary(source_docx))
        except Exception:
            is_partial_source = False

    sections = [
        section
        for section in profile.sections
        if not section.is_excluded and _is_business_question_label(section.label)
    ]
    if is_partial_source:
        sections = [section for section in sections if _is_major_section_anchor(section.anchor_text)]
        try:
            summary = build_doc_summary(source_docx)
            enriched_tables = detect_tables(summary["tables"], include_improvable_filled=True)
            table_by_id = {table.table_id: table for table in enriched_tables}
            merged_tables: list[TableProfile] = []
            for table in profile.tables:
                extra = table_by_id.get(table.table_id)
                if extra is None:
                    merged_tables.append(table)
                    continue
                cell_by_id = {cell.cell_id: cell for cell in table.cells}
                for cell in extra.cells:
                    if cell.cell_id not in cell_by_id:
                        cell_by_id[cell.cell_id] = cell
                merged_tables.append(
                    table.model_copy(update={"cells": list(cell_by_id.values())})
                )
            profile = profile.model_copy(update={"tables": merged_tables})
        except Exception:
            pass
    tables: list[TableProfile] = []
    for table in profile.tables:
        if not _is_business_question_label(table.label):
            continue
        cells = [cell for cell in table.cells if _is_business_question_label(cell.label)]
        if not cells:
            continue
        tables.append(table.model_copy(update={"cells": cells}))
    image_slots = [slot for slot in profile.image_slots if _is_business_question_label(slot.label)]
    questions = build_questions(sections, tables, image_slots)
    notes = list(profile.analysis_notes)
    if len(questions) != len(profile.questions):
        notes.append("행정성 문구(동의서, 개인정보, 별지, 첨부서류 등)는 입력 폼에서 제외했습니다.")
    return profile.model_copy(
        update={
            "sections": sections,
            "tables": tables,
            "image_slots": image_slots,
            "questions": questions,
            "analysis_notes": notes,
        }
    )
