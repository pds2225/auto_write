"""초기창업패키지(AI 인재 실증형) 양식 분석·생성 빌드 스크립트 (임시, 검증 후 삭제).
실제 워크스페이스(D:\\auto_write\\workspace)를 대상으로 동작한다.
"""
from __future__ import annotations
import os
import sys
from datetime import date
from pathlib import Path

from auto_write.config import get_settings, ensure_directories
from auto_write.storage import Storage
from auto_write.services.openai_client import OpenAIService
from auto_write.services.evidence_service import EvidenceService
from auto_write.services.image_service import ImageService
from auto_write.services.render_service import RenderService
from auto_write.services.qa_service import QAService
from auto_write.services.project_service import ProjectService

FORM = Path(os.getenv("AUTO_WRITE_CHOCHANG_FORM", r"C:\Users\ekth3\Downloads\(초기) [별첨 1] 2026년도 초기창업패키지(AI 인재 실증형) 사업계획서 양식.docx"))


def make_service() -> ProjectService:
    settings = get_settings()
    ensure_directories(settings)
    storage = Storage(settings)
    oa = OpenAIService(settings)
    svc = ProjectService(storage, oa, EvidenceService(oa), ImageService(oa), RenderService(), QAService())
    return settings, storage, svc


def cmd_analyze() -> None:
    settings, storage, svc = make_service()
    print("workspace:", settings.workspace_root)
    print("results  :", settings.results_root)
    print("has_openai:", settings.has_openai, "| has_anthropic:", settings.has_anthropic, "| provider:", settings.ai_provider)
    print("FORM exists:", FORM.exists(), "|", FORM.name)
    profile = svc.analyze_uploaded_template(FORM.name, FORM.read_bytes())
    print("=" * 60)
    print("template_id :", profile.template_id)
    print("source_docx :", profile.source_docx)
    print("sections    :", len(profile.sections))
    for s in profile.sections:
        print("   SEC", s.field_id, "|", (s.anchor_text or "")[:55], "|", (s.label or "")[:40])
    print("tables      :", len(profile.tables))
    for t in profile.tables:
        print("   TBL", t.table_id, "idx", t.table_index, "r", t.row_count, "c", t.col_count,
              "cells", len(t.cells), "|", (t.label or "")[:45])
    print("image_slots :", len(profile.image_slots))
    for im in profile.image_slots:
        label = getattr(im, "label", "") or ""
        print("   IMG", im.slot_id, getattr(im, "anchor_type", ""), "|", str(label)[:45])
    print("questions   :", len(profile.questions))


MIRAE_BASE = Path(os.getenv("AUTO_WRITE_MIRAE_BASE", r"C:\Users\ekth3\OneDrive\바탕 화면\다솜\경영지도사 개인\02. 밸류업파트너스\2026 미래큐러스"))

USER_BRIEF = (
    "미래큐러스(Miraculous)는 'AI가 실패하거나 해킹되어도 독립 하드웨어 안전 도메인이 인간을 보호한다'는 "
    "가치를 구현하는 AI Safety & Security Infrastructure 연구기업이다.\n"
    "· AI·휴머노이드·드론·BCI 시장 급성장으로 물리 세계와 연결된 AI 안전 사고 위험이 현실화되고 있음\n"
    "· 기존 AI 시스템은 Software 단일 구조로, AI 오작동·해킹 시 안전 기능도 함께 마비됨\n"
    "· 4대 한계: 단일 시스템 구조 / Software 의존 / 독립 안전 도메인 부재 / 즉각 차단 구조 부재\n"
    "· 자동차의 ISO 26262처럼 AI·로봇 산업도 Functional Safety 규제 적용 가능성이 커지고 EU AI Act 등 규제가 강화되는 흐름\n"
    "· 미래큐러스는 등록특허 KR 10-2915951 기반 독립 하드웨어 안전 아키텍처로 First Mover 포지션 확보를 추진"
)

USER_NOTES = "\n\n".join([
    (
        "[실현 가능성] HSCC(Humanoid Safety Control Chip) 핵심 구조\n"
        "· 등록특허 KR 10-2915951 기반, 메인 AI 프로세서와 완전히 독립된 하드웨어 안전 제어 구조\n"
        "· FSM(유한상태기계) 기반 하드웨어 로직으로 CPU·OS를 우회, 80ms 이내 결정론적 응답 목표\n"
        "· NVIDIA 생태계(Omniverse·Isaac·CUDA·Jetson) 활용해 AI Safety 구조를 시뮬레이션·검증할 예정\n"
        "· 협약기간 내 Robot AI Safety 설계안, NVIDIA 연계 검증 보고서, 추가 특허 2~4건 출원을 목표"
    ),
    (
        "[성장전략] 핵심 비즈니스 모델: Fabless + Licensing + Platform\n"
        "· 밸류체인(추진 예정): 미래큐러스(원천 특허·Architecture)→가온칩스 등(SoC 설계)→삼성전자(파운드리)→글로벌 로봇 OEM\n"
        "· 수익 구조: 기술 라이선스(선급금·로열티·MG) + Safety SoC 공급 + Safety Platform 구독\n"
        "· 단계: Robot AI Safety→Financial Security→Mobile AI Safety→Neural Safety(가정 시나리오)\n"
        "· 협약기간 KPI(목표): AI 위험 시뮬레이션 30건 이상, 구조 설계안 4식, 추가 특허 출원 2~4건, 협력 검토 3건 이상"
    ),
    (
        "[팀 구성] 대표자 및 팀원 구성 계획\n"
        "· 대표: 원천 특허 개발 및 AI Safety 플랫폼 구조 연구 총괄(약학 전공, KR 10-2915951 등 특허 보유)\n"
        "· 부사장: 산업안전 총괄(한국산업안전공단 안전연구소장 경력)\n"
        "· 현재 재직(대표 제외) 2명 / 협약기간 내 AI Safety 연구 인력 1명 추가 채용 예정(AI 교육 수료 인재)\n"
        "· 외부 협력: 삼성전자·가온칩스 Safety SoC 협력 검토 중, NVIDIA 생태계 활용 추진 예정"
    ),
])


def _report_artifacts(storage, pid, art) -> None:
    import json as _json
    from docx import Document
    print("=" * 60)
    print("project_id  :", pid)
    print("output_docx :", art.output_docx)
    print("results_dir :", art.results_dir)
    print("results_docx:", art.results_docx)
    qa = _json.loads(Path(art.qa_report).read_text(encoding="utf-8"))
    print("QA passed   :", qa.get("passed"), "| errors:", qa.get("error_count"), "| warnings:", qa.get("warning_count"))
    for e in (qa.get("errors") or [])[:8]:
        print("   ERR", str(e)[:110])
    for w in (qa.get("warnings") or [])[:6]:
        print("   WRN", str(w)[:110])
    transfer = _json.loads(Path(art.transfer_report).read_text(encoding="utf-8"))
    print("transfer mode:", transfer.get("mode"), "| fill_ratio:", transfer.get("fill_ratio"))
    doc = Document(art.output_docx)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    print("output paragraphs:", len(paras), "| tables:", len(doc.tables))
    # placeholder leftovers
    leftover = [t for t in paras if "○○○" in t or "OOO" in t or "□□□" in t]
    print("placeholder leftovers:", len(leftover))


def cmd_generate() -> None:
    settings, storage, svc = make_service()
    print("workspace:", settings.workspace_root, "| provider:", settings.ai_provider)
    profile = svc.analyze_uploaded_template(FORM.name, FORM.read_bytes())
    tid = profile.template_id
    print("template_id:", tid)
    pid = svc.create_project(tid, "2026 초기창업패키지(AI 인재 실증형) - 미래큐러스")
    refs = []
    for fn in ["body_extracted.md", "body.docx.txt", "미래큐러스_사업계획서_완성본.docx.txt"]:
        p = MIRAE_BASE / fn
        if p.exists():
            refs.append((fn, p.read_bytes()))
    print("references attached:", [r[0] for r in refs])
    svc.save_project_form(
        project_id=pid,
        answers={"user_brief": USER_BRIEF, "user_notes": USER_NOTES},
        project_title="Robot AI Safety & Security Infrastructure 플랫폼",
        organization_name="미래큐러스(Miraculous)",
        evidence_topics="",
        reference_files=refs,
    )
    print("generating... (AI 호출 포함, 수 분 소요 가능)")
    art = svc.generate(pid)
    _report_artifacts(storage, pid, art)


def build_mirae_plan() -> dict:
    """미래큐러스 → 초기창업패키지(AI 인재 실증형) 양식 채움 plan."""
    item_name = "Robot AI Safety & Security Infrastructure 플랫폼"
    deliverable = "Robot AI Safety Architecture 설계안 / NVIDIA 연계 AI Safety 검증 보고서 / 추가 특허 2~4건 출원(PCT 포함)"
    return {
        "identity": {
            "기업명": "미래큐러스(Miraculous)",
            "개업연월일": "2025.03.17",
            "사업자 구분": "개인사업자",
            "대표자 유형": "단독",
            "사업자등록번호": "(사업자등록증 상 등록번호 기재)",
            "사업자 소재지": "서울특별시 강남구",
        },
        "overview": {
            "명칭": "Robot AI Safety & Security Infrastructure Platform",
            "범주": "AI Safety 플랫폼 / Functional Safety 반도체 / AI Security Architecture",
            "아이템개요": "AI가 실패·해킹되어도 독립 하드웨어 안전 도메인이 인간을 보호하는 AI Safety & Security Infrastructure. 등록특허 KR 10-2915951 기반 HSCC(Humanoid Safety Control Chip) 독립 안전 제어 구조.",
            "문제인식": "AI·휴머노이드·드론·BCI 확산으로 물리 세계와 연결된 AI 안전사고 위험이 현실화. 기존 Software 단일 구조는 오작동·해킹 시 안전기능도 함께 마비됨.",
            "실현가능성": "등록특허 기반 메인 AI와 독립된 FSM 하드웨어 안전제어(80ms 이내 결정론적 응답 목표). NVIDIA 생태계(Omniverse·Isaac·CUDA·Jetson)로 시뮬레이션·검증.",
            "성장전략": "Fabless + Licensing + Platform BM. Robot→Financial→Mobile→Neural 단계 확장. 기술 라이선스 + Safety SoC 공급 + Safety Platform 구독.",
            "팀구성": "대표(원천특허·AI Safety 총괄), 부사장(산업안전 총괄), 협약기간 내 AI Safety 연구 인력 1명 추가 채용 예정.",
            "AI인재활용": "NVIDIA AI 생태계 기반 AI Safety 시뮬레이션·검증에 AI 연구인력 1명 채용. AI 위험 시뮬레이션 30건 이상, 구조 설계안 4식 도출 목표.",
        },
        "row_rewrites": [
            # 표4(병합 요약표, idx=4): 창업아이템명/산출물/팀/AI채용
            {"table_index": 4, "row": 0, "cols": [None, item_name]},
            {"table_index": 4, "row": 1, "cols": [None, deliverable]},
            {"table_index": 4, "row": 3, "cols": [None, "1명", "-"]},  # 채용인원 정규직1/계약직-
            {"table_index": 4, "row": 14, "cols": ["1", "부사장", "산업안전 총괄", "한국산업안전공단 안전연구소장 경력", "완료"]},
            {"table_index": 4, "row": 15, "cols": ["2", "부장", "반도체 설계 협력", "팹리스·파운드리 실무 경험", "예정('26.下)"]},
            {"table_index": 4, "row": 19, "cols": ["1", "정규직", "-", "40시간", "AI Safety 연구·구조 분석", "'26.7월"]},
            {"table_index": 4, "row": 20, "cols": ["", "", "", "", "", ""]},  # 2번째 예시행 제거(채용 1명)
            # 표6(개요요약 이미지캡션은 replacements에서 처리)
            # 표7(추진일정, idx=6)
            {"table_index": 6, "row": 1, "cols": ["1", "AI 인재 채용", "2026.07~", "AI Safety 연구 인력 1명 채용"]},
            {"table_index": 6, "row": 2, "cols": ["2", "Robot AI Safety 구조 연구", "2026.07~08", "Functional Safety Architecture 고도화·AI 위험행동 분석"]},
            {"table_index": 6, "row": 3, "cols": ["3", "NVIDIA 생태계 검증", "2026.07~09", "Omniverse·CUDA·Isaac 시뮬레이션 30건 이상"]},
            {"table_index": 6, "row": 4, "cols": ["4", "Financial Security 구조 연구", "2026.08~10", "Hardware Security 구조 설계안 도출"]},
            {"table_index": 6, "row": 5, "cols": ["5", "추가 특허 출원", "협약기간 내", "AI Safety 확장 특허 2~4건 출원(PCT 포함)"]},
            # 표9(사업비, idx=8): 비목·집행계획만 실제값, 금액은 협약 후 확정(공란)
            # 비목 r3·r4는 세로 병합 셀이므로 동일 라벨로 통일(마지막 기록이 병합셀 전체 적용)
            {"table_index": 8, "row": 3, "cols": ["기계장치·SW 및 외주용역비", "GPU 연산환경·AI Safety 연구 소프트웨어", "", "", "", ""]},
            {"table_index": 8, "row": 4, "cols": ["기계장치·SW 및 외주용역비", "반도체 설계 검토·디지털 트윈 시뮬레이션 환경 구축", "", "", "", ""]},
            {"table_index": 8, "row": 5, "cols": ["특허취득비", "PCT 국제출원·미국 등 개별국 출원", "", "", "", ""]},
            {"table_index": 8, "row": 6, "cols": ["인건비", "AI Safety 연구 인력(대표자 제외)", "", "", "", ""]},
            {"table_index": 8, "row": 7, "cols": ["합 계", "(협약 후 사업비 산정·확정)", "", "", "", ""]},
            # 표10(전체 일정, idx=9)
            {"table_index": 9, "row": 1, "cols": ["1", "Mobile AI Safety 구조 연구", "2026.09~11", "AI 실행 통제 구조 연구"]},
            {"table_index": 9, "row": 2, "cols": ["2", "Neural Safety 구조 연구", "2026.09~11", "BCI 기반 Human Protection 구조 연구"]},
            {"table_index": 9, "row": 3, "cols": ["3", "최종 산출물 제출", "2026.11", "설계안·검증 보고서·사업계획서"]},
            {"table_index": 9, "row": 4, "cols": ["4", "글로벌 사업화 기반", "협약 이후", "PCT 국제출원·해외 특허 진입 기반 확보"]},
            # 표11(팀 구성, idx=10)
            {"table_index": 10, "row": 1, "cols": ["1", "부사장", "산업안전 총괄", "한국산업안전공단 안전연구소장 경력", "완료"]},
            {"table_index": 10, "row": 2, "cols": ["2", "부장", "반도체 설계 협력", "팹리스·파운드리 실무 경험", "예정('26.下)"]},
            {"table_index": 10, "row": 3, "cols": ["3", "AI 인재", "AI Safety 연구·구조 분석", "AI·반도체·SoC 전공 / Functional Safety 이해", "예정('26.7)"]},
            # 표12(협력 파트너, idx=11)
            {"table_index": 11, "row": 1, "cols": ["1", "삼성전자", "AI 반도체·파운드리", "Safety SoC 구조 협력·특허 실시 검토", "검토 중"]},
            {"table_index": 11, "row": 2, "cols": ["2", "가온칩스", "반도체 설계·디자인하우스", "Safety SoC 설계 협력", "검토 중"]},
            {"table_index": 11, "row": 3, "cols": ["3", "NVIDIA", "GPU·AI·Omniverse·Edge AI", "AI 위험 시뮬레이션·Safety 검증", "협약기간 중 추진 예정"]},
            # 표13(AI 인재 채용, idx=12)
            {"table_index": 12, "row": 1, "cols": ["1", "정규직", "-", "40시간", "AI Safety 연구·구조 분석", "예정('26.7)"]},
            {"table_index": 12, "row": 2, "cols": ["", "", "", "", "", ""]},
            # 표14(AI 인재 채용·핵심과제 추진 일정, idx=13, 3열) — 템플릿 예시 제거·실제 일정
            {"table_index": 13, "row": 1, "cols": ["AI 인재 채용", "'26.07", "AI Safety 연구 인력 1명 채용(근로계약·4대보험)"]},
            {"table_index": 13, "row": 2, "cols": ["핵심 개발과제 추진", "'26.07~", "Robot AI Safety 구조 연구·NVIDIA 생태계 검증 착수"]},
            {"table_index": 13, "row": 3, "cols": ["1단계 핵심 산출물 완성", "'26.11", "Robot AI Safety 설계안·NVIDIA 연계 검증 보고서"]},
            {"table_index": 13, "row": 4, "cols": ["2단계 구조 연구 확장", "'26.12~", "Financial·Mobile·Neural Safety 구조 연구"]},
            {"table_index": 13, "row": 5, "cols": ["추가 특허 출원", "협약기간 내", "AI Safety 확장 특허 2~4건(PCT 포함)"]},
            {"table_index": 13, "row": 6, "cols": ["글로벌 사업화 기반", "'27 상반기", "PCT 국제출원·해외 특허 진입 기반 확보"]},
            {"table_index": 13, "row": 7, "cols": ["최종 산출물 제출", "'27 이후", "설계안·검증 보고서·사업계획서 종합"]},
        ],
        "replacements": {
            "< 사진(이미지) 또는 설계도 제목 >": "[그림] AI Safety & Security Infrastructure 플랫폼 구조도 — 독립 하드웨어 안전 도메인 개념(협약 후 실제 이미지 삽입)",
        },
        "replacements_prefix": {
            "OO기술이 적용된": item_name,
            "모바일 어플리케이션(0개), 웹사이트(0개)": deliverable,
            "OO학 박사, OO학과 교수 재직": "한국산업안전공단 안전연구소장 경력",
            "OO학 학사, OO 관련 경력": "팹리스·파운드리 실무 경험",
            "AX 로드맵 수립": "AI Safety 연구·구조 분석",
            "성과측정(ROI 분석)": "AI 위험 시뮬레이션·검증",
            "성과측정(ROI분석": "AI 위험 시뮬레이션·검증",
        },
        # 5. AI 인재활용 계획(핵심 평가항목) — profile 미포함 본문이라 가이드 문구를 직접 교체
        "paragraph_fills": [
            {
                "anchor": "ㅇ AI 도입의 필요성에 대해 세부내용 작성",
                "lines": [
                    "ㅇ AI 도입의 필요성",
                    "· AI Safety·Security Infrastructure 구조는 방대한 위험 시나리오의 시뮬레이션·검증이 필수이며, NVIDIA GPU·AI 생태계 없이는 검증 속도·규모 확보가 어려움",
                    "· 등록특허 KR 10-2915951 기반 독립 하드웨어 안전 구조(HSCC)를 Omniverse·Isaac·CUDA 디지털 트윈으로 검증해 설계 신뢰성을 선제 확보",
                    "· AI 위험 행동 분석·Risk Score 연산 가속에 AI 전문 인력과 AI 인프라가 직접 활용됨",
                ],
            },
            {
                "anchor": "ㅇ AI 인재 활용 계획에 대한 상세 내용 작성(채용시기, 채용기간, 채용 인원및 업무 등)",
                "lines": [
                    "ㅇ AI 인재 활용 계획 (채용시기·기간·인원·업무)",
                    "· 채용시기: 2026.07월 / 채용유형: 정규직 1명 / 근로시간: 주 40시간",
                    "· 담당업무: AI Safety 연구·구조 분석, AI 위험 시뮬레이션(Omniverse) 30건 이상 수행·검증",
                    "· 요구역량: AI·반도체·SoC·로봇 관련 전공, Functional Safety 이해, AI 시뮬레이션 분석 가능자",
                    "· 활용 성과(목표): 구조 설계안 4식(Robot/Financial/Mobile/Neural), 추가 특허 2~4건 출원 기여",
                ],
            },
        ],
    }


def cmd_finalize(pid: str) -> None:
    import json as _json
    from docx import Document
    from auto_write.services.submittable_filler import SubmittableFiller
    settings, storage, svc = make_service()
    in_docx = settings.workspace_root / "projects" / pid / "output" / "output.docx"
    print("pid       :", pid)
    print("input docx:", in_docx, "| exists:", in_docx.exists())
    if not in_docx.exists():
        print("[중단] output.docx 없음 — 먼저 generate 필요")
        return
    try:
        pinput = storage.load_project_input(pid)
        print("org       :", pinput.organization_profile.get("name", ""))
        print("title     :", pinput.project_meta.get("project_title", ""))
    except Exception as exc:
        print("(project_input 로드 실패:", exc, ")")
    # 표 인덱스/헤더 검증
    doc = Document(str(in_docx))
    print("tables    :", len(doc.tables))
    for ti, tb in enumerate(doc.tables):
        r0 = " | ".join(c.text.strip().replace(chr(10), '/')[:16] for c in tb.rows[0].cells[:6]) if tb.rows else ""
        print(f"   T{ti} ({len(tb.rows)}x{len(tb.columns)}): {r0}")
    # 마감 실행
    plan = build_mirae_plan()
    out_name = f"미래큐러스_초기창업패키지_AI인재실증형_제출초안_{date.today().strftime('%Y%m%d')}.docx"
    out_docx = settings.results_root / out_name
    report = SubmittableFiller(plan).finalize(in_docx, out_docx)
    print("=" * 60)
    print("FINALIZE REPORT")
    print(_json.dumps({k: v for k, v in report.items() if k != "residual_remaining"}, ensure_ascii=False, indent=2))
    print("residual_remaining:", len(report.get("residual_remaining", [])))
    for r in report.get("residual_remaining", [])[:25]:
        print("   !", r)
    print("output    :", out_docx, "| exists:", out_docx.exists())


if __name__ == "__main__":
    cmd = sys.argv[1] if len(sys.argv) > 1 else "analyze"
    if cmd == "analyze":
        cmd_analyze()
    elif cmd == "generate":
        cmd_generate()
    elif cmd == "finalize":
        cmd_finalize(sys.argv[2])
    elif cmd == "inspect":
        from docx import Document
        docx_path = Path(sys.argv[2])
        doc = Document(str(docx_path))
        print("PARAGRAPHS (non-empty):")
        for i, p in enumerate(doc.paragraphs):
            t = p.text.strip()
            if t:
                print(f"  {i:03d}: {t[:130]}")
        print("TABLES:", len(doc.tables))
        for ti, tb in enumerate(doc.tables):
            print(f"-- Table {ti} ({len(tb.rows)}x{len(tb.columns)}) --")
            for r in tb.rows:
                cells = [c.text.strip().replace(chr(10), '/') for c in r.cells]
                print("   |", " | ".join(c[:22] for c in cells))
    elif cmd == "struct":
        # 양식 구조 정밀 덤프: 섹션 field_id + 표 cell_id + 원본 그리드 텍스트
        from docx import Document
        from auto_write.services.project_service import ProjectService as PS
        settings, storage, svc = make_service()
        profile = svc.analyze_uploaded_template(FORM.name, FORM.read_bytes())
        out = []
        def w(s=""):
            out.append(s)
        w("=" * 70)
        w(f"template_id: {profile.template_id} | sections={len(profile.sections)} tables={len(profile.tables)} images={len(profile.image_slots)} questions={len(profile.questions)}")
        w("\n### SECTIONS ###")
        for s in profile.sections:
            psst = []
            txt = f"{s.anchor_text} {s.label}"
            if PS.PSST_PROBLEM_RE.search(txt): psst.append("PROBLEM")
            if PS.PSST_SOLUTION_RE.search(txt): psst.append("SOLUTION")
            if PS.PSST_SCALE_RE.search(txt): psst.append("SCALE")
            if PS.PSST_TEAM_RE.search(txt): psst.append("TEAM")
            w(f"  [{s.field_id}] excl={s.is_excluded} {'/'.join(psst) or '-':8} | anchor={s.anchor_text[:48]!r} | label={s.label[:36]!r}")
        w("\n### TABLES ###")
        doc = Document(str(FORM))
        for t in profile.tables:
            core = bool(PS.CORE_TABLE_LABEL_RE.search(str(t.label or "")))
            w(f"\n-- {t.table_id} idx={t.table_index} {t.row_count}x{t.col_count} CORE={core} | label={t.label[:50]!r}")
            # 원본 그리드
            if 0 <= t.table_index < len(doc.tables):
                tb = doc.tables[t.table_index]
                for ri, r in enumerate(tb.rows):
                    cells = [c.text.strip().replace(chr(10), '/') for c in r.cells]
                    w(f"   GRID r{ri}: " + " | ".join(f"{c[:26]}" for c in cells))
            # profile cells
            for c in t.cells:
                w(f"   CELL {c.cell_id} r{c.row}c{c.cell} req={c.required} | label={c.label[:24]!r} rowH={c.row_header[:18]!r} colH={c.col_header[:18]!r}")
        w("\n### IMAGE SLOTS ###")
        for im in profile.image_slots:
            w(f"  [{im.slot_id}] {im.anchor_type} req={im.required} | label={str(im.label)[:40]!r} | ref={im.anchor_ref}")
        text = "\n".join(out)
        dump_path = Path(r"D:\auto_write\WORKS\form_struct_dump.txt")
        dump_path.parent.mkdir(parents=True, exist_ok=True)
        dump_path.write_text(text, encoding="utf-8")
        print(text)
        print("\n[written]", dump_path)
    elif cmd == "heads":
        # 양식 앞부분 표(일반현황 등 profile 미포함 표) 전체 그리드 덤프
        from docx import Document
        doc = Document(str(FORM))
        n = int(sys.argv[2]) if len(sys.argv) > 2 else 5
        print("total tables:", len(doc.tables))
        for ti, tb in enumerate(doc.tables[:n]):
            print(f"\n== RAW Table {ti} ({len(tb.rows)}x{len(tb.columns)}) ==")
            for ri, r in enumerate(tb.rows):
                cells = [c.text.strip().replace(chr(10), '/') for c in r.cells]
                # dedup 연속 병합 표시
                print(f"  r{ri}: " + " | ".join(c[:34] for c in cells))
    else:
        print("unknown cmd:", cmd)
