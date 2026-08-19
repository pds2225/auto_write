#!/usr/bin/env python3
"""Build the 2026 소셜벤처 리그 솔루션 제안서 (별첨1).

Official form = 공고 HWPX 별첨1 (지정 양식, 최대 5쪽).
Facts come from the 2026-08-11 서울 AI 허브 입주신청서 and the public 공고.
Unverified numbers are left as [확인필요]. Original notice is never overwritten.
"""
from __future__ import annotations

import os
import re
import sys
import zipfile
from copy import deepcopy
from pathlib import Path

from lxml import etree

APP = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(APP))

from core.docx.services.hwpx_fill import _q, _set_cell_text  # noqa: E402

_STANDALONE_RE = re.compile(rb"standalone\s*=\s*['\"](yes|no)['\"]")


def _detect_standalone(xml_bytes: bytes):
    m = _STANDALONE_RE.search(xml_bytes[:200])
    if not m:
        return None
    return m.group(1) == b"yes"

# --- 사실 근거 (날조 0) -------------------------------------------------
# 서울 AI 허브 입주신청서 v1.1 (2026-08-11, Drive) + 중기부 공고 제2026-482호
APPLICANT_NAME = "박다솜"
TEAM_NAME = "마켓게이트"
BIRTH_PLACEHOLDER = "[확인필요]"  # 이력 소스 생년월일이 1992.04.06 / 1992.04.15 로 불일치
TITLE = "수출 초보 중소기업의 AI 생산성 격차를 줄이는 수출실행 플랫폼 MarketGate"

# 공고 붙임1 과제④ 원문 요지 + AI 허브 사업계획서
BACKGROUND = (
    "본 제안은 공고 과제④ 「AI를 활용한 소상공인·중소기업의 생산성 및 혁신역량 제고」에 해당합니다. "
    "공고는 소상공인·중소기업이 인력난과 디지털 전환 지연으로 생산성 정체가 이어지는데도, "
    "개별 사업장이 스스로 혁신을 시도할 여력이 부족하다고 규정합니다. "
    "기존 지원 또한 개별 사업장 자금·컨설팅 직접지원에 머물러, "
    "다수 사업장의 공통 문제를 푸는 서비스·협업 모델은 지원 공백이라고 명시합니다.\n\n"
    "신청인 박다솜은 수출 컨설팅과 AI 서비스 기획·개발을 해 오며, "
    "같은 구조가 수출 현장에서도 반복되는 것을 확인했습니다. "
    "수출 데이터는 UN Comtrade·KITA 등 기관별로 파편화되어 있고, "
    "전담 인력이 없는 중소기업은 시장 선정 이후 바이어 발굴·컨택·후속관리·계약까지 "
    "실행을 이어가기 어렵습니다. "
    "이에 AI·공공데이터로 국가 추천부터 바이어 검증·컨택·성과 추적까지 연결하는 "
    "SaaS(MarketGate)로, 개별 컨설팅이 아니라 다수 기업이 같이 쓸 수 있는 생산성 인프라를 제안합니다."
)

PROBLEM = (
    "▶ 이 문제를 겪고 있는 사람들은 누구이고, 어느 정도 규모인가요?\n"
    "서비스를 받는 쪽은 해외진출을 준비하거나 기존 수출을 확대하려는 수출 초보 중소기업입니다. "
    "1차 현장 대상은 K-뷰티 수출(준비) 기업이며, 이후 다른 소비재로 확장하는 구조입니다. "
    "공급·일하는 쪽은 수출 컨설턴트·무역대행·협회 실무자로, "
    "건별 수작업 조사와 바이어 리스트 작성에 시간이 묶여 다수 기업을 동시에 돕지 못합니다. "
    "시장 규모 산정은 수출 중소기업 약 8만개사를 출발점으로 두었습니다"
    "(서울 AI 허브 입주 사업계획서, 2026.8.11.).\n\n"
    "▶ 이들은 어떤 상황에서 어떤 어려움을 겪고 있나요?\n"
    "① 수출데이터 파편화: 국가·품목 통계를 기관 사이트마다 따로 조회해야 합니다.\n"
    "② 수출국가 선정: 담당자 경험과 수작업 분석에 의존해, ‘왜 이 나라인가’를 설명하기 어렵습니다.\n"
    "③ 바이어 발굴·검증: 검색·리스트 확보에 시간이 들고, 적합·제재·인증 여부를 교차검증하기 어렵습니다.\n"
    "④ 컨택 이후 단절: Contacted → Replied → Meeting → RFQ → 계약으로 이어지는 실행 기록이 없어 "
    "성과를 관리하지 못합니다.\n"
    "공고 과제④가 지적한 ‘개별 사업장의 혁신 여력 부족’이, 수출 실행 구간에서 그대로 나타납니다.\n\n"
    "▶ 이 문제가 실제로 존재한다는 것을 어떻게 알게 되었나요?\n"
    "통계·제도 쪽은 공고 붙임1 과제④의 문제 현황·기존 지원제도 한계를 근거로 합니다. "
    "현장 쪽은 신청인이 수출 컨설팅을 수행하며 반복해서 본 패턴입니다. "
    "기업마다 같은 조사·리스트 작업을 처음부터 다시 했고, "
    "공공 정보는 있으나 다음 실행(컨택·후속·계약)으로 넘어가지 못했습니다. "
    "파일럿 수치(베타 기업 수, 전환율)는 아직 실증 전 계획이므로 본 제안서에 확정 실적으로 적지 않습니다."
)

EXISTING = (
    "▶ 현재 실제로 어떻게 대응하고 있나요? 왜 해결되지 않나요?\n"
    "가장 많이 쓰이는 방법은 세 갈래입니다.\n"
    "① 공공 수출지원(KOTRA·무역협회 등): 정보 제공과 지원 절차 중심입니다. "
    "데이터는 있으나 해석·추천·바이어 검증·파이프라인 관리까지 한 흐름으로 이어지지 않습니다.\n"
    "② 민간 바이어 발굴·컨설팅: 프로젝트형으로 가능하지만 건별 비용이 크고, "
    "끝나면 실행 기록이 기업 안에 남지 않습니다.\n"
    "③ 내부 수작업: 담당자가 포털을 검색하고 엑셀로 관리합니다. "
    "전담 인력이 없는 기업은 이 단계를 지속하지 못합니다.\n"
    "한계가 반복되는 이유는 공고가 적은 것과 같습니다. "
    "지원이 개별 사업장 자금·컨설팅에 머무르고, "
    "다수 사업장의 공통 문제를 푸는 상시 서비스 모델은 제도 바깥에 있습니다. "
    "비용(건별 고가), 제도(정보 제공형 공공서비스), 구조(실행 구간 단절)가 겹칩니다.\n\n"
    "▶ 본인 또는 팀이 직접 해봤거나 관여했던 활동은 무엇이고, 무엇이 부족했나요?\n"
    "신청인은 수출 컨설팅으로 기업별 시장조사·바이어 후보 정리·실행 자문을 수행했습니다. "
    "기업 한 곳에는 도움이 되었으나, 같은 조사 절차를 다음 기업에 거의 처음부터 다시 해야 했습니다. "
    "사람이 직접 하는 한 동시에 도울 수 있는 기업 수가 늘지 않고, "
    "컨택 이후 단계(회신·미팅·RFQ·계약)를 같은 화면에서 추적할 도구가 없었습니다. "
    "이 공백을 메우기 위해 MarketGate를 예비창업 아이템으로 설계했고, "
    "HS코드·공공데이터 기반 유망국 추천 및 바이어 매칭 특허를 출원했습니다"
    "(출원번호 10-2026-0026207). "
    "현재는 대표 1인 중심이며, 개발 협력 1인·전문자문 1인과 협업하고 있습니다. "
    "유료 전환율·미팅 성사율 등 운영 실적은 아직 없으며, "
    "K-뷰티 수출(준비) 중소기업 10개사 이상 1개월 무료 베타가 다음 검증 계획입니다."
)

SOLUTION = (
    "▶ 해결 방법과 작동 과정. 핵심 기술이 이 문제에 적합한 이유.\n"
    "제품은 MarketGate입니다. 수출 초보 중소기업(1차 K-뷰티)에게 "
    "국가 선정부터 바이어 발굴·검증, 컨택·Follow-up, 계약·수출성과 추적까지 "
    "한 흐름으로 제공합니다.\n"
    "① 국가 선정: HS코드·공공데이터를 묶어 유망 수출국을 추천합니다. "
    "파편화된 통계를 한 화면에서 보게 해 ‘왜 이 나라인가’를 설명할 근거를 만듭니다.\n"
    "② 바이어 발굴·검증: 품목·기업·거래 데이터로 후보를 찾고 교차검증합니다. "
    "검색 시간을 줄이고, 적합하지 않은 대상을 걸러 컨택 낭비를 줄입니다.\n"
    "③ 컨택·Follow-up: 이메일 발송과 후속 접촉을 기록·자동화합니다.\n"
    "④ 성과 파이프라인: Contacted → Replied → Qualified → Meeting → RFQ → 계약·선적까지 "
    "단계 전환을 남깁니다. 조회 수가 아니라 실행 성과를 관리합니다.\n"
    "핵심 기술은 HS코드·공공데이터 기반 추천·매칭이며, 위 특허 출원 범위와 같습니다. "
    "이 문제에 적합한 이유는, 과제④가 요구하는 것이 개별 컨설팅 대행이 아니라 "
    "다수 사업장이 같이 쓸 수 있는 AI 생산성 도구이기 때문입니다.\n\n"
    "▶ 기존 대응과 구체적으로 무엇이 다른가요?\n"
    "공공 서비스는 정보·절차, 민간 서비스는 부분 업무 대행에 가깝습니다. "
    "MarketGate는 상시 SaaS로 국가추천·바이어 검증·파이프라인·성과 추적을 연결합니다. "
    "신청인 본인의 기존 컨설팅과도 다릅니다. "
    "컨설팅은 사람-기업 1:1이고, 본 솔루션은 같은 절차를 소프트웨어로 재사용해 "
    "동시에 여러 기업이 실행하게 합니다.\n\n"
    "▶ 무엇으로 지속되나요? 누가 비용을 지불하나요?\n"
    "이용 기업이 비용을 지불합니다. "
    "수익은 저가 월 구독(Buyer CRM·파이프라인·트래킹)과, "
    "Qualified Lead·미팅·RFQ 등 실제 성과가 날 때 성과 크레딧을 차감하는 구조입니다. "
    "이후 관세·물류·법무 제휴와 기관형 B2B(수출바우처 등)로 확장하는 계획입니다. "
    "목표 숫자는 입주 사업계획서 기준이며 아직 실적이 아닙니다. "
    "1차년도 유료 약 40~50개·매출 0.4억원, 3차년도 연간 유료 250개사 이상·매출 6억원입니다. "
    "본 리그 서류평가 선정 시 임팩트 활동비(200만원)는 문제 구체화와 BM 고도화"
    "(베타 기업 모집·성과지표 정의)에 쓰겠습니다."
)

IMPACT = (
    "▶ 가장 먼저 일어나는 변화는 무엇이고, 누구에게 일어나나요?\n"
    "가장 먼저, K-뷰티 수출(준비) 중소기업 담당자의 국가 선정·바이어 후보 정리 시간이 줄어들고, "
    "컨택 이후 단계가 화면에 남습니다. "
    "‘조사는 했는데 다음 실행이 없다’는 단절이 줄어드는 것이 1차 변화입니다. "
    "공급 쪽(컨설턴트·유관 실무)은 같은 조사 작업을 기업마다 처음부터 반복하지 않게 됩니다. "
    "이 변화의 측정 지표는 조회 수가 아니라 "
    "Contacted → Replied → Meeting → RFQ 전환과, 베타 참여 기업 수입니다. "
    "베타는 10개사 이상 1개월 무료 실증이 계획이며, 본 제출 시점의 완료 실적은 아닙니다.\n\n"
    "▶ 사업이 성장하면 그 변화는 어디까지 커질 수 있나요? 근거는 무엇인가요?\n"
    "가까운 범위는 K-뷰티 수출 실행의 디지털화입니다. "
    "입주 사업계획서 기준으로 2차년도 유료 150~200개, 3차년도 250개사 이상·매출 6억원을 목표로 적었습니다. "
    "TAM 4,304억원 / SAM 807억원 / SOM 6억원은 같은 문서의 시장 가정이며, "
    "현재 매출은 예비창업으로 해당 없습니다. "
    "사회적 임팩트의 축은 매출 극대화보다, "
    "공고 과제④가 비운 ‘다수 사업장의 공통 문제를 푸는 서비스 모델’을 "
    "수출 실행 구간에 실제로 놓는 것입니다. "
    "기업 수가 늘수록 같은 AI·데이터 기반이 재사용되어 "
    "개별 컨설팅으로는 닿지 않는 기업의 수출 시도가 가능해집니다. "
    "확장 순서는 K-뷰티 검증 → 타 소비재 검토이며, "
    "전환율·재구매율·유지율을 확인한 뒤에 규모를 키우겠습니다."
)

TASK4_CELL = (
    "□1 지역사회 돌봄 공백 해소  "
    "□2 일자리 및 사회참여 확대  "
    "□3 중소사업장・취약계층의 에너지 비용 및 효율 개선  "
    "■4 AI를 활용한 소상공인・중소기업의 생산성 및 혁신역량 제고  "
    "□5 자원순환 촉진 및 폐기물 문제 해결"
    "※ 온라인 신청 화면에서 선택한 과제와 일치해야 합니다."
)


def _local(tag: str) -> str:
    return tag.split("}")[-1] if isinstance(tag, str) else ""


def _direct_tcs(tr):
    return [c for c in tr if _local(getattr(c, "tag", "")) == "tc"]


def _tbl_rows(tbl):
    return [r for r in tbl if _local(getattr(r, "tag", "")) == "tr"]


def set_cell_multiline(tc, text: str) -> bool:
    """Replace a cell's paragraphs with one hp:p per line, cloning the first p."""
    lines = text.replace("\r\n", "\n").split("\n")
    if not lines:
        return False
    sub = None
    for child in tc:
        if _local(getattr(child, "tag", "")) == "subList":
            sub = child
            break
    if sub is None:
        return _set_cell_text(tc, text.replace("\n", " "))
    paras = [p for p in sub if _local(getattr(p, "tag", "")) == "p"]
    if not paras:
        return _set_cell_text(tc, text.replace("\n", " "))
    template = paras[0]
    _set_cell_text(tc, lines[0])
    for extra in paras[1:]:
        sub.remove(extra)
    for line in lines[1:]:
        clone = deepcopy(template)
        ts = [el for el in clone.iter(_q("t"))]
        if ts:
            ts[0].text = line
            for t in ts[1:]:
                t.text = ""
        sub.append(clone)
    return True


def extract_annex1(src: Path, dst: Path) -> None:
    """Keep only 별첨1 (before 별첨2). Never overwrite src."""
    if src.resolve() == dst.resolve():
        raise ValueError("출력이 입력과 같습니다.")
    dst.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(src, "r") as zin:
        names = zin.namelist()
        data = {n: zin.read(n) for n in names}
    section_name = "Contents/section0.xml"
    root = etree.fromstring(data[section_name])
    kids = list(root)
    start = end = None
    for i, kid in enumerate(kids):
        text = "".join(kid.itertext())
        compact = re.sub(r"\s+", "", text)
        # 제출서류 표("1솔루션제안서*[별첨1]...")가 아니라 별첨1 제목 행만.
        if start is None and compact in {"별첨1솔루션제안서", "별첨1"}:
            start = i
        if start is not None and i > start and compact.startswith("별첨2"):
            end = i
            break
    if start is None or end is None:
        raise RuntimeError(f"별첨1 구간을 찾지 못했습니다 start={start} end={end}")
    for kid in kids[:start] + kids[end:]:
        root.remove(kid)
    standalone = _detect_standalone(data[section_name])
    data[section_name] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=standalone
    )
    _strip_unused_bindata(data)
    _write_hwpx(dst, data)


def _strip_unused_bindata(data: dict[str, bytes]) -> None:
    """Drop notice screenshots that 별첨1 no longer references."""
    for name in [n for n in data if n.startswith("BinData/")]:
        del data[name]
    hpf_name = "Contents/content.hpf"
    if hpf_name in data:
        xml = data[hpf_name].decode("utf-8")
        xml = re.sub(r'<opf:item id="image\d+"[^/]*/>', "", xml)
        data[hpf_name] = xml.encode("utf-8")
    if "Preview/PrvText.txt" in data:
        data["Preview/PrvText.txt"] = (
            "「모두의 창업:사회혁신 소셜벤처 리그」솔루션 제안서\n"
            "마켓게이트 / 박다솜 / 과제④\n"
        ).encode("utf-8")


def _write_hwpx(dst: Path, data: dict[str, bytes]) -> None:
    names = list(data)
    tmp = dst.with_suffix(dst.suffix + ".part")
    with zipfile.ZipFile(tmp, "w") as zout:
        if "mimetype" in data:
            zi = zipfile.ZipInfo("mimetype")
            zi.compress_type = zipfile.ZIP_STORED
            zout.writestr(zi, data["mimetype"])
        for name in names:
            if name == "mimetype":
                continue
            zout.writestr(name, data[name])
    os.replace(tmp, dst)


def fill_annex1(path: Path) -> dict[str, str]:
    with zipfile.ZipFile(path, "r") as zin:
        data = {n: zin.read(n) for n in zin.namelist()}
    section_name = "Contents/section0.xml"
    root = etree.fromstring(data[section_name])
    tbls = root.findall(".//{*}tbl")
    report: dict[str, str] = {"tables": str(len(tbls))}
    indexed = []
    for i, tbl in enumerate(tbls):
        t = "".join(tbl.itertext()).replace("\n", " ")
        indexed.append((i, t[:40]))
    report["index"] = repr(indexed)

    def find_tbl(pred):
        for i, tbl in enumerate(tbls):
            t = "".join(tbl.itertext())
            if pred(t):
                return i
        raise KeyError("table not found")

    i_name = find_tbl(lambda t: "생년월일" in t and "00.00.00" in t)
    i_sel = find_tbl(lambda t: "선택과제" in t and "제목" in t)
    i_bg = find_tbl(lambda t: "과제선택 배경" in t)
    i_body = find_tbl(lambda t: "기존대응 문제점" in t and "과제 해결방안" in t)

    name_row = _direct_tcs(_tbl_rows(tbls[i_name])[0])
    # 성명 | value | 생년월일 | 00.00.00
    _set_cell_text(name_row[1], APPLICANT_NAME)
    _set_cell_text(name_row[3], BIRTH_PLACEHOLDER)

    sel_rows = _tbl_rows(tbls[i_sel])
    _set_cell_text(_direct_tcs(sel_rows[0])[1], TASK4_CELL)
    _set_cell_text(_direct_tcs(sel_rows[1])[1], TITLE)

    bg_rows = _tbl_rows(tbls[i_bg])
    _set_cell_text(_direct_tcs(bg_rows[0])[1], "과제④를 선택한 이유")
    set_cell_multiline(_direct_tcs(bg_rows[1])[0], BACKGROUND + "\n\n" + PROBLEM)

    body_rows = _tbl_rows(tbls[i_body])
    set_cell_multiline(_direct_tcs(body_rows[1])[0], EXISTING)
    set_cell_multiline(_direct_tcs(body_rows[3])[0], SOLUTION)
    set_cell_multiline(_direct_tcs(body_rows[5])[0], IMPACT)

    standalone = _detect_standalone(data[section_name])
    data[section_name] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=standalone
    )
    _write_hwpx(path, data)
    report["ok"] = "1"
    report["name_tbl"] = str(i_name)
    return report


def build_docx(out: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
    from docx.oxml import OxmlElement

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)

    def set_run_font(run, size=10.5, bold=False, color=None, name="맑은 고딕"):
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Malgun Gothic"
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Malgun Gothic")
        rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
        rFonts.set(qn("w:eastAsia"), name)
        if color:
            run.font.color.rgb = color

    def shade(cell, hex_color: str) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    def set_cell(cell, text, *, size=10.5, bold=False, center=False, color=None, fill=None):
        if fill:
            shade(cell, fill)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        first = True
        for line in text.split("\n"):
            if first:
                run = p.add_run(line)
                first = False
            else:
                p2 = cell.add_paragraph()
                p2.paragraph_format.space_before = Pt(2)
                p2.paragraph_format.space_after = Pt(2)
                p2.paragraph_format.line_spacing = 1.15
                run = p2.add_run(line)
            set_run_font(run, size=size, bold=bold, color=color)

    NAVY = "1F4E79"
    WHITE = RGBColor(255, 255, 255)
    LABEL = "D6E3F0"

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("「모두의 창업:사회혁신 소셜벤처 리그」솔루션 제안서")
    set_run_font(r, size=16, bold=True, color=RGBColor(31, 78, 121))

    note = doc.add_paragraph()
    nr = note.add_run(
        "지정 양식(별첨1) 기준 작성 · 최대 5쪽 · 선택과제 ④ · 예비창업 마켓게이트 / 박다솜"
    )
    set_run_font(nr, size=9, color=RGBColor(89, 89, 89))
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    t0 = doc.add_table(rows=1, cols=4)
    t0.style = "Table Grid"
    row = t0.rows[0].cells
    set_cell(row[0], "성명", bold=True, center=True, color=WHITE, fill=NAVY)
    set_cell(row[1], APPLICANT_NAME, center=True)
    set_cell(row[2], "생년월일", bold=True, center=True, color=WHITE, fill=NAVY)
    set_cell(row[3], BIRTH_PLACEHOLDER + " (6자리)", center=True)

    t1 = doc.add_table(rows=2, cols=2)
    t1.style = "Table Grid"
    set_cell(t1.rows[0].cells[0], "선택과제", bold=True, center=True, color=WHITE, fill=NAVY)
    set_cell(
        t1.rows[0].cells[1],
        "□1 돌봄 공백  □2 일자리·사회참여  □3 에너지 효율\n"
        "■4 AI를 활용한 소상공인·중소기업의 생산성 및 혁신역량 제고\n"
        "□5 자원순환\n※ 온라인 신청 화면의 선택 과제와 일치해야 합니다.",
        size=10,
    )
    set_cell(t1.rows[1].cells[0], "제목", bold=True, center=True, color=WHITE, fill=NAVY)
    set_cell(t1.rows[1].cells[1], TITLE, bold=True, size=11)

    def section_table(label: str, prompt: str, body: str):
        tbl = doc.add_table(rows=2, cols=2)
        tbl.style = "Table Grid"
        set_cell(tbl.rows[0].cells[0], label, bold=True, center=True, color=WHITE, fill=NAVY)
        set_cell(tbl.rows[0].cells[1], prompt, size=10, fill=LABEL)
        merged = tbl.rows[1].cells[0].merge(tbl.rows[1].cells[1])
        set_cell(merged, body, size=10.5)
        # widen label col
        tbl.columns[0].width = Cm(3.4)
        tbl.columns[1].width = Cm(13.6)
        doc.add_paragraph()

    section_table("과제선택 배경", "과제를 선택하게 된 배경", BACKGROUND)
    section_table(
        "사회문제",
        "누구·규모 / 어떤 어려움 / 어떻게 알게 되었는지",
        PROBLEM,
    )
    section_table("기존대응 문제점", "이 문제는 지금까지 어떻게 다뤄져 왔나요?", EXISTING)
    section_table("과제 해결방안", "제안하는 제품·서비스는 이 문제를 어떤 방식으로 해결하나요?", SOLUTION)
    section_table("임팩트 성과", "이 시도가 성공하면 무엇이, 얼마나 달라지나요?", IMPACT)

    foot = doc.add_paragraph()
    fr = foot.add_run(
        "사실관계 출처: ① 중소벤처기업부 공고 제2026-482호 붙임1 과제④ ② 서울 AI 허브 입주신청서·사업계획서 "
        "(마켓게이트/박다솜, 2026.8.11.). AI 도구로 정리했으나 미확인 수치는 [확인필요]로 남겼습니다. "
        "지정 양식 5쪽을 넘는 내용은 평가에 반영되지 않습니다."
    )
    set_run_font(fr, size=8, color=RGBColor(89, 89, 89))

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))


def verify_hwpx(path: Path) -> list[str]:
    z = zipfile.ZipFile(path)
    root = etree.fromstring(z.read("Contents/section0.xml"))
    text = "".join(root.itertext())
    fails = []
    for needle in (
        "박다솜",
        "MarketGate",
        "■4",
        "10-2026-0026207",
        TITLE[:12],
        "과제④",
    ):
        if needle not in text:
            fails.append(f"missing {needle}")
    for banned in ("김지훈", "서울대학교 경영학과", "N=120", "성공률 82%"):
        if banned in text:
            fails.append(f"banned fact leaked: {banned}")
    if "별첨2" in text.replace(" ", ""):
        fails.append("별첨2 leaked into proposal file")
    if "모집공고" in text and "제2026" in text:
        fails.append("notice body leaked")
    return fails


def write_checklist(out: Path) -> None:
    out.write_text(
        """# 소셜벤처 리그 솔루션 제안서 — 제출 전 확인

마감: 2026-08-19(수) 18:00 정각 (ONE 플랫폼). 마감일 접속 지연 가능.
접수: https://www.kibo.or.kr/portal  → 솔루션제안서 양식 다운로드 → 업로드.
기 제출한 신청서는 수정 불가. 알림톡 수령해야 접수 완료.

## 이 파일이 뭔가
- 솔루션제안서_마켓게이트_박다솜_별첨1.hwpx : 공고 별첨1만 분리해 채운 지정 양식(다른 양식 제출 시 평가 제외).
- 솔루션제안서_마켓게이트_박다솜_별첨1.docx : 같은 내용 가독용. 업로드는 HWPX(또는 ONE에서 받은 원 양식에 옮겨 적기).

## 선택
- 과제④ AI를 활용한 소상공인·중소기업의 생산성 및 혁신역량 제고
- 예비창업 마켓게이트 / 대표 박다솜
- 온라인 신청 화면의 과제 선택과 반드시 일치

## 제출 전 본인이 채울 것
- [ ] 생년월일 6자리 (문서에는 [확인필요]. 소스 간 1992.04.06 / 1992.04.15 불일치)
- [ ] 공고일(2026.7.30) 기준 사업자 미등록인지 — 예비창업 자격. 등록돼 있으면 창업기업 서류(사업자등록증명) 추가
- [ ] ONE 플랫폼에서 받은 최신 별첨1 양식과 표 칸이 같은지 한 번 대조
- [ ] 한글에서 HWPX 열림·5쪽 이내 확인 (5쪽 초과분은 평가 제외)
- [ ] 동의서·서약서는 플랫폼에서 체크 (이번 산출물은 제안서만)
- [ ] 대표자 본인 계정으로 제출 (타인 신청 시 탈락)

## 넣지 않은 것 (날조 0)
- 없는 팀원 실명, 미완료 베타 실적, 출처 불명 시장 통계
- 특허 10-2026-0026207·매출 목표·조직(대표1+개발협력1+자문1)은 2026-08-11 서울 AI 허브 신청서 기준

문의: 기보 소셜벤처가치평가센터 02-3407-2925 / 2921 / 2932
""",
        encoding="utf-8",
    )


def main() -> int:
    notice = Path(os.environ.get("SVL_NOTICE_HWPX", "/tmp/svl/notice.hwpx"))
    out_dir = Path(os.environ.get("SVL_OUT_DIR", "/opt/cursor/artifacts"))
    out_dir.mkdir(parents=True, exist_ok=True)
    hwpx = out_dir / "솔루션제안서_마켓게이트_박다솜_별첨1.hwpx"
    docx = out_dir / "솔루션제안서_마켓게이트_박다솜_별첨1.docx"
    extract_annex1(notice, hwpx)
    rep = fill_annex1(hwpx)
    build_docx(docx)
    write_checklist(out_dir / "svl_proposal_submit_checklist.md")
    fails = verify_hwpx(hwpx)
    print("extract/fill", rep)
    print("docx", docx, "bytes", docx.stat().st_size)
    print("hwpx", hwpx, "bytes", hwpx.stat().st_size)
    if fails:
        print("VERIFY FAIL", fails)
        return 2
    print("VERIFY PASS")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
