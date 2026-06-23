"""test_cross_form_paragraph_fields.py — cross-form 자동전사의 **본문 단락형 빈칸**
("라벨 : ____") 채우기 회귀 테스트.

기존 자동전사는 표(table) 칸만 채웠다. 실제 정부·지자체 양식은 표 외에도
``○ 신청기관명 : ________`` / ``대표자 :`` 같은 본문 단락형 빈칸을 많이 쓴다.
이 테스트는 그 단락형 빈칸을 표와 동일한 보수적 규칙(high 자동 + 사용자 확정,
날조 0, 원본 미수정, 마스킹값 보존)으로 채우는지 검증한다.

검증 항목:
  ① 단락형 빈칸 타깃 탐지(blank 만, 이미 채워짐/문장/마스킹은 제외)
  ② 단락 빈칸 high 자동전사 + 선행 장식(○·번호) 보존
  ③ 소스에 없는 라벨은 단락도 미전사(날조 0)
  ④ 원본 source/target 파일 미수정
  ⑤ 마스킹 값(○○○)이 든 단락은 '이미 채워짐'으로 보존(덮어쓰기 금지)
  ⑥ 표·단락 혼합 양식에서 둘 다 전사
  ⑦ 단락 needs_confirm 후보를 --confirm 으로 확정 적용
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from auto_write.services.cross_form_autofill import (
    autofill_from_source,
    find_target_fields,
    match_fields,
    extract_source_fields,
)


def _source_table(path: Path) -> None:
    """소스: 라벨|값 2열 표(기업명/대표자/사업자등록번호/제품명)."""
    doc = Document()
    doc.add_heading("사업계획서", 0)
    t = doc.add_table(rows=4, cols=2)
    rows = [
        ("기업명", "밸류업(주)"),
        ("대표자", "홍길동"),
        ("사업자등록번호", "123-45-67890"),
        ("제품명", "스마트 수질센서"),
    ]
    for i, (lab, val) in enumerate(rows):
        t.rows[i].cells[0].text = lab
        t.rows[i].cells[1].text = val
    doc.save(str(path))


def _para_text_by_prefix(docx_path: Path, prefix_norm: str) -> str:
    """타깃 문서 본문 단락 중 정규화 라벨이 prefix_norm 으로 시작하는 단락 텍스트."""
    from auto_write.services.cross_form_autofill import _key

    doc = Document(str(docx_path))
    for p in doc.paragraphs:
        txt = p.text or ""
        if ":" in txt or "：" in txt:
            head = txt.replace("：", ":").split(":", 1)[0]
            if _key(head) == prefix_norm:
                return txt
    return ""


# --- ① 탐지 -------------------------------------------------------------------

def test_find_paragraph_blank_targets(tmp_path: Path) -> None:
    tgt = tmp_path / "b.docx"
    doc = Document()
    doc.add_paragraph("○ 기업명 : ________")     # 빈칸 → 타깃
    doc.add_paragraph("대표자 :")                  # 빈칸(콜론 뒤 없음) → 타깃
    doc.add_paragraph("주소 : 서울시 강남구")        # 이미 채워짐 → 타깃 아님
    doc.add_paragraph("참고 : 빈칸을 채우시오")       # 문장 → 타깃 아님
    doc.add_paragraph("연락처 : ○○○")             # 마스킹값 → 타깃 아님(보존)
    doc.save(str(tgt))

    targets = find_target_fields(tgt)
    para_targets = [t for t in targets if t.get("kind") == "paragraph"]
    norms = {t["normalized"] for t in para_targets}
    assert norms == {"기업명", "대표자"}
    # 단락 타깃은 para_index 를 가지고 table 좌표는 -1
    for t in para_targets:
        assert t["para_index"] >= 0
        assert t["table_index"] == -1


# --- ② high 자동전사 + 장식 보존 ----------------------------------------------

def test_autofill_paragraph_high_and_decoration(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _source_table(src)

    doc = Document()
    doc.add_paragraph("○ 기업명 : ________")
    doc.add_paragraph("1. 대표자 :")
    doc.save(str(tgt))

    rep = autofill_from_source(src, tgt, out)
    assert rep.ok
    assert rep.transcribed >= 2
    # 장식(○·1.)은 보존하고 값만 채운다.
    assert _para_text_by_prefix(out, "기업명") == "○ 기업명 : 밸류업(주)"
    assert _para_text_by_prefix(out, "대표자") == "1. 대표자 : 홍길동"


# --- ③ 날조 0(소스에 없는 라벨은 단락도 미전사) -------------------------------

def test_autofill_paragraph_no_fabrication(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _source_table(src)

    doc = Document()
    doc.add_paragraph("사업명 :")  # 소스에 '사업명' 없음
    doc.save(str(tgt))

    rep = autofill_from_source(src, tgt, out)
    # 값이 채워지지 않았다 → 단락 그대로(콜론 뒤 비어 있음)
    assert _para_text_by_prefix(out, "사업명").replace(" ", "") in ("사업명:",)
    # 전사 0이라 ok=False(보수적), 사업명은 미전사
    assert all(m.normalized != "사업명" or not m.value for m in rep.matches)


# --- ④ 원본 미수정 ------------------------------------------------------------

def test_autofill_paragraph_originals_untouched(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _source_table(src)
    doc = Document()
    doc.add_paragraph("기업명 : ____")
    doc.save(str(tgt))

    before = Path(tgt).read_bytes()
    autofill_from_source(src, tgt, out)
    after = Path(tgt).read_bytes()
    assert before == after  # 타깃 원본 바이트 동일


# --- ⑤ 마스킹값 보존(덮어쓰기 금지) -------------------------------------------

def test_autofill_paragraph_masked_value_preserved(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _source_table(src)
    doc = Document()
    doc.add_paragraph("대표자 : ○○○")  # 블라인드 마스킹된 이름 = 이미 채워짐
    doc.save(str(tgt))

    autofill_from_source(src, tgt, out)
    # ○○○ 은 빈칸이 아니므로 홍길동으로 덮어쓰지 않는다.
    assert _para_text_by_prefix(out, "대표자") == "대표자 : ○○○"


# --- ⑥ 표·단락 혼합 -----------------------------------------------------------

def test_autofill_table_and_paragraph_mixed(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _source_table(src)

    doc = Document()
    # 표 빈칸: 사업자번호(동의어 → 사업자등록번호)
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "사업자번호"
    t.rows[0].cells[1].text = ""
    # 단락 빈칸: 기업명
    doc.add_paragraph("○ 기업명 : ____")
    doc.save(str(tgt))

    rep = autofill_from_source(src, tgt, out)
    assert rep.ok
    # 표 셀 전사 확인
    doc_out = Document(str(out))
    cell_val = doc_out.tables[0].rows[0].cells[1].text.strip()
    assert cell_val == "123-45-67890"
    # 단락 전사 확인
    assert _para_text_by_prefix(out, "기업명") == "○ 기업명 : 밸류업(주)"


# --- ⑦ 단락 needs_confirm 후보 확정 적용 --------------------------------------

def test_autofill_paragraph_confirm_flow(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _source_table(src)  # 소스에 '제품명' 존재

    doc = Document()
    doc.add_paragraph("제품명칭 :")  # 제품명 ↔ 제품명칭 = 퍼지(자동전사 보류)
    doc.save(str(tgt))

    # 확정 없이는 보류
    rep0 = autofill_from_source(src, tgt, out)
    assert any(c["normalized"] == "제품명칭" for c in rep0.needs_confirm)
    assert _para_text_by_prefix(out, "제품명칭").replace(" ", "") == "제품명칭:"

    # 확정하면 채워진다(날조 0 — 소스 실값).
    out2 = tmp_path / "out2.docx"
    rep1 = autofill_from_source(
        src, tgt, out2, confirmations={"제품명칭": "제품명"})
    assert rep1.confirmed >= 1
    assert _para_text_by_prefix(out2, "제품명칭") == "제품명칭 : 스마트 수질센서"
