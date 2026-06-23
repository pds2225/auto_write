"""test_cross_form_multifield.py — cross-form 자동전사의 **한 줄 여러 칸**
(멀티필드 단락) 채우기 회귀 테스트.

실제 정부·지자체 신청서 표지는 한 줄에 여러 칸을 나란히 둔다:

    기업명 : ____________    대표자 : ____________
    연락처 : (생략)          이메일 : ____

기존 자동전사는 단락의 **첫 콜론 뒤 전체**를 한 값으로 봐서, 같은 줄에 다른 칸이
이어지면 "이미 채워짐"으로 오판해 **그 줄 전체를 놓쳤다**(두 칸 다 미전사).
이 테스트는 한 줄의 각 "라벨 : 빈칸" 칸을 개별 인식해, 빈칸만 골라 소스 실값으로
채우고(높은 신뢰 high + 사용자 확정), 이미 채워진/마스킹된 칸과 칸 사이 간격은
그대로 보존하는지(날조 0·원본 미수정) 검증한다.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from auto_write.services.cross_form_autofill import (
    autofill_from_source,
    extract_source_fields,
    find_target_fields,
)


def _source_table(path: Path) -> None:
    """소스: 라벨|값 2열 표(기업명/대표자/사업자등록번호/제품명)."""
    doc = Document()
    doc.add_heading("사업계획서", 0)
    rows = [
        ("기업명", "밸류업(주)"),
        ("대표자", "홍길동"),
        ("사업자등록번호", "123-45-67890"),
        ("제품명", "스마트 수질센서"),
    ]
    t = doc.add_table(rows=len(rows), cols=2)
    for i, (lab, val) in enumerate(rows):
        t.rows[i].cells[0].text = lab
        t.rows[i].cells[1].text = val
    doc.save(str(path))


def _para_line(docx_path: Path, contains: str) -> str:
    """타깃 문서 본문 단락 중 contains 를 포함하는 첫 단락 텍스트."""
    doc = Document(str(docx_path))
    for p in doc.paragraphs:
        if contains in (p.text or ""):
            return p.text
    return ""


# --- ① 한 줄 여러 칸 탐지 -----------------------------------------------------

def test_multifield_line_detected(tmp_path: Path) -> None:
    tgt = tmp_path / "b.docx"
    doc = Document()
    doc.add_paragraph("기업명 : ____________    대표자 : ____________")
    doc.save(str(tgt))

    targets = find_target_fields(tgt)
    para = [t for t in targets if t.get("kind") == "paragraph"]
    norms = {t["normalized"] for t in para}
    assert norms == {"기업명", "대표자"}      # 두 칸 모두 인식(전엔 0개)
    # 같은 단락의 두 칸 → 같은 para_index, 표 좌표는 -1
    assert len({t["para_index"] for t in para}) == 1
    assert all(t["table_index"] == -1 and t["value_cell"] == -1 for t in para)


# --- ② 한 줄 두 칸 자동전사(둘 다 high) + 간격 보존 ---------------------------

def test_multifield_both_high_filled(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _source_table(src)
    doc = Document()
    doc.add_paragraph("기업명 : ____________    대표자 : ____________")
    doc.save(str(tgt))

    rep = autofill_from_source(src, tgt, out)
    assert rep.ok
    assert rep.transcribed == 2
    # 두 값이 각자 자리에, 칸 사이 4칸 간격은 보존
    assert _para_line(out, "기업명") == "기업명 : 밸류업(주)    대표자 : 홍길동"


# --- ③ 한 칸은 빈칸·한 칸은 이미 채워짐 → 빈칸만 채우고 나머지 보존 ----------

def test_multifield_mixed_blank_and_filled(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _source_table(src)
    doc = Document()
    doc.add_paragraph("기업명 : ____    대표자 : 김삿갓")  # 대표자는 이미 채워짐
    doc.save(str(tgt))

    rep = autofill_from_source(src, tgt, out)
    assert rep.transcribed == 1                         # 빈칸 1개만
    # 기업명만 채우고, 이미 적힌 '김삿갓'은 홍길동으로 덮어쓰지 않는다.
    assert _para_line(out, "기업명") == "기업명 : 밸류업(주)    대표자 : 김삿갓"


# --- ④ 한 줄에 마스킹값(○○○)이 섞여도 보존 ----------------------------------

def test_multifield_masked_segment_preserved(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _source_table(src)
    doc = Document()
    doc.add_paragraph("기업명 : ____    연락처 : ○○○")  # 연락처는 블라인드 마스킹
    doc.save(str(tgt))

    autofill_from_source(src, tgt, out)
    # ○○○ 은 빈칸이 아니므로 보존, 기업명만 채움
    assert _para_line(out, "기업명") == "기업명 : 밸류업(주)    연락처 : ○○○"


# --- ⑤ 날조 0: 소스에 없는 칸은 빈칸 그대로 ----------------------------------

def test_multifield_no_fabrication(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _source_table(src)
    doc = Document()
    doc.add_paragraph("기업명 : ____    사업명 : ____")  # 소스에 사업명 없음
    doc.save(str(tgt))

    rep = autofill_from_source(src, tgt, out)
    assert rep.transcribed == 1
    # 기업명만 채워지고 사업명 칸은 빈칸(____)으로 보존(없는 값 지어내지 않음)
    assert _para_line(out, "기업명") == "기업명 : 밸류업(주)    사업명 : ____"


# --- ⑥ 탭(\t)으로 구분된 한 줄 두 칸 ------------------------------------------

def test_multifield_tab_separated(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _source_table(src)
    doc = Document()
    doc.add_paragraph("기업명 : ____\t대표자 : ____")
    doc.save(str(tgt))

    rep = autofill_from_source(src, tgt, out)
    assert rep.transcribed == 2
    assert _para_line(out, "기업명") == "기업명 : 밸류업(주)\t대표자 : 홍길동"


# --- ⑦ 소스의 한 줄 여러 칸 본문도 각 칸으로 추출 ----------------------------

def test_source_multifield_body_extracted(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    doc = Document()
    # 표가 아니라 본문 한 줄에 두 칸이 채워져 있는 소스
    doc.add_paragraph("기업명 : 밸류업(주)    대표자 : 홍길동")
    doc.save(str(src))

    fields = extract_source_fields(src)
    assert fields.get("기업명") == "밸류업(주)"
    assert fields.get("대표자") == "홍길동"        # 전엔 첫 칸 값에 오염되어 묶였음


# --- ⑧ 한 줄 두 칸 + needs_confirm 확정 적용 ----------------------------------

def test_multifield_confirm_flow(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _source_table(src)
    doc = Document()
    doc.add_paragraph("기업명 : ____    제품명칭 : ____")  # 제품명칭=퍼지(제품명)
    doc.save(str(tgt))

    # 확정 없이: 기업명만 high 자동, 제품명칭은 보류
    rep0 = autofill_from_source(src, tgt, out)
    assert rep0.transcribed == 1
    assert any(c["normalized"] == "제품명칭" for c in rep0.needs_confirm)

    # 확정하면 같은 줄의 둘째 칸도 채워진다(날조 0 — 소스 실값).
    out2 = tmp_path / "out2.docx"
    rep1 = autofill_from_source(src, tgt, out2, confirmations={"제품명칭": "제품명"})
    assert rep1.transcribed == 2
    assert rep1.confirmed >= 1
    assert _para_line(out2, "기업명") == "기업명 : 밸류업(주)    제품명칭 : 스마트 수질센서"


# --- ⑨ 원본(소스·타깃) 미수정 ------------------------------------------------

def test_multifield_originals_untouched(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _source_table(src)
    doc = Document()
    doc.add_paragraph("기업명 : ____    대표자 : ____")
    doc.save(str(tgt))

    src_before = src.read_bytes()
    tgt_before = tgt.read_bytes()
    autofill_from_source(src, tgt, out)
    assert src.read_bytes() == src_before
    assert tgt.read_bytes() == tgt_before
