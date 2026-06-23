"""test_cross_form_label_variants.py — 라벨 변형 recall 확대 회귀.

실제 정부지원사업 양식 라벨에 거의 항상 붙는 장식(글머리표 ○ ▶ · /
순번 1. ① 가. Ⅰ.)이 있어도 정확일치·동의어 매칭이 작동해 빈 칸이
자동전사되는지 검증한다. 동시에 장식 제거가 오매칭(사업명↛사업자명)을
새로 만들지 않는지(보수성 유지)도 함께 잠근다.

핵심 효과(before→after):
  - before: '○ 기업명', '1. 성명' 처럼 장식이 붙은 타깃 라벨은 fuzzy 로 강등돼
            자동전사가 전혀 안 됐다(빈 양식 그대로).
  - after:  장식을 벗긴 핵심 라벨로 정확일치·동의어 매칭이 되어 자동 전사된다.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from auto_write.services.cross_form_autofill import (
    _best_source_for_target,
    _key,
    autofill_from_source,
)


def _value_for(docx_path: Path, label: str) -> str:
    doc = Document(str(docx_path))
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if cells and cells[0].text.strip() == label:
                return cells[1].text.strip()
    return ""


# --- 단위: 정규화 키가 장식을 벗긴다 -----------------------------------------

def test_key_strips_bullet_prefix() -> None:
    assert _key("○ 기업명") == _key("기업명")
    assert _key("● 대표자") == "대표자"
    assert _key("▶ 연락처") == "연락처"
    assert _key("· 주소") == "주소"
    assert _key("※ 이메일") == "이메일"


def test_key_strips_number_prefix() -> None:
    assert _key("1. 기업명") == "기업명"
    assert _key("12) 대표자") == "대표자"
    assert _key("① 연락처") == "연락처"
    assert _key("가. 주소") == "주소"
    assert _key("Ⅰ. 사업명") == "사업명"


def test_key_preserves_label_without_decoration() -> None:
    # 순번처럼 보여도 구분자(./)) 가 없는 정상 라벨은 깎지 않는다.
    assert _key("1차년도매출") == "1차년도매출"
    assert _key("가산점항목") == "가산점항목"
    assert _key("기업명") == "기업명"


def test_key_all_decoration_keeps_original() -> None:
    # 전부 장식(글자 없음)이면 원본 유지(마스킹/잔여물 보호).
    assert _key("○○○") != ""
    assert _key("○○○") == _key("○○○")


# --- 단위: 매칭이 장식 라벨에서도 high 가 된다 --------------------------------

def test_match_bullet_label_exact_high() -> None:
    source = {"기업명": "밸류업(주)"}
    sel, conf, _cand = _best_source_for_target(_key("○ 기업명"), source)
    assert conf == "high"
    assert sel == "기업명"


def test_match_number_label_synonym_high() -> None:
    source = {"대표자": "홍길동"}
    sel, conf, _cand = _best_source_for_target(_key("1. 성명"), source)
    assert conf == "high"
    assert sel == "대표자"


# --- e2e: 장식 라벨 타깃에 전사 ----------------------------------------------

def test_e2e_bulleted_and_numbered_target_labels(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"

    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "기업명"
    t.rows[0].cells[1].text = "밸류업(주)"
    t.rows[1].cells[0].text = "대표자"
    t.rows[1].cells[1].text = "홍길동"
    doc.save(str(src))

    doc2 = Document()
    t2 = doc2.add_table(rows=2, cols=2)
    t2.rows[0].cells[0].text = "○ 기업명"        # 글머리표 + 정확일치
    t2.rows[0].cells[1].text = ""
    t2.rows[1].cells[0].text = "1. 성명"          # 순번 + 동의어(대표자)
    t2.rows[1].cells[1].text = ""
    doc2.save(str(tgt))

    report = autofill_from_source(src, tgt, out, use_ai=False)
    assert _value_for(out, "○ 기업명") == "밸류업(주)"
    assert _value_for(out, "1. 성명") == "홍길동"
    assert report.transcribed == 2


# --- e2e: 확장 동의어 클러스터(실무 라벨) -------------------------------------

def test_e2e_extended_cluster_employees(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"

    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "상시근로자수"
    t.rows[0].cells[1].text = "12명"
    doc.save(str(src))

    doc2 = Document()
    t2 = doc2.add_table(rows=1, cols=2)
    t2.rows[0].cells[0].text = "종업원수"  # 동의어(상시근로자수)
    t2.rows[0].cells[1].text = ""
    doc2.save(str(tgt))

    report = autofill_from_source(src, tgt, out, use_ai=False)
    assert _value_for(out, "종업원수") == "12명"
    assert report.transcribed == 1


def test_e2e_extended_cluster_fax(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"

    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "팩스번호"
    t.rows[0].cells[1].text = "02-123-4567"
    doc.save(str(src))

    doc2 = Document()
    t2 = doc2.add_table(rows=1, cols=2)
    t2.rows[0].cells[0].text = "팩스"  # 동의어(팩스번호)
    t2.rows[0].cells[1].text = ""
    doc2.save(str(tgt))

    report = autofill_from_source(src, tgt, out, use_ai=False)
    assert _value_for(out, "팩스") == "02-123-4567"
    assert report.transcribed == 1


# --- 적대 반례: 장식 제거가 오매칭을 새로 만들지 않는다 ----------------------

def test_decoration_strip_does_not_create_false_match(tmp_path: Path) -> None:
    """장식을 벗겨도 '사업명' ↛ '사업자명' 자동전사 금지(보수성 유지)."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"

    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "○ 사업명"
    t.rows[0].cells[1].text = "AI 인재실증 사업"
    doc.save(str(src))

    doc2 = Document()
    t2 = doc2.add_table(rows=1, cols=2)
    t2.rows[0].cells[0].text = "1. 사업자명"
    t2.rows[0].cells[1].text = ""
    doc2.save(str(tgt))

    report = autofill_from_source(src, tgt, out, use_ai=False)
    assert _value_for(out, "1. 사업자명") == ""
    assert report.transcribed == 0
    assert any(n["normalized"] == _key("사업자명") for n in report.needs_confirm)
