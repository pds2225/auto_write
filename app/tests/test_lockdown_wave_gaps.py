"""JSON gap 잠금 — L004·L014·L048·L049·L072·L105 (L050 생성은 BLOCKED).

기계화 4점: guard + test + coverage JSON + runtime wiring.
JSON L050(동일명 HWP+PDF 병행 생성)은 이 클라우드에 한글/LibreOffice 가 없어 gap 유지.
"""
from __future__ import annotations

from pathlib import Path
from unittest import mock

import pytest
from docx import Document
from pypdf import PdfReader, PdfWriter

from auto_write.document_ingest import _render_unhwp_table
from auto_write.models import ProjectInput
from auto_write.services.company_extract import extract_company_from_file
from auto_write.services.doc_text_extract import (
    extract_tax_invoice_buyer,
    looks_like_tax_invoice,
)
from auto_write.services.docx_ops import (
    GENERATED_TABLE_HEADER_FILL,
    add_generated_table,
    cell_shading_fill,
    style_generated_table,
)
from auto_write.services.document_quality_orchestrator import DocumentQualityOrchestrator
from auto_write.services.evaluation_service import EvaluationService
from auto_write.services.openai_client import OpenAIService
from auto_write.services.skill_frontmatter import (
    collect_skill_frontmatter_errors,
    parse_skill_frontmatter,
)
from auto_write.services.submission_gates import (
    announcement_tuple_stem,
    build_submit_layout_dir,
    is_dated_notice_folder,
    is_draft_artifact,
    is_submit_layout_path,
    merge_pdfs,
    missing_pdf_pair,
)
from auto_write.services.submission_orchestrator import SubmissionPipeline
from test_submission_pipeline import (
    _FakeProjectService,
    _FakeStorage,
    _profile,
    _settings,
)


_REPO = Path(__file__).resolve().parents[2]

_TAX_TEXT = """전자세금계산서
공급자
상호 (법인명)
주식회사 공급자알파
공급받는자
상호 (법인명)
홍길동상사
"""


def _blank_pdf(path: Path, pages: int = 1) -> Path:
    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=72, height=72)
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("wb") as fh:
        writer.write(fh)
    return path


def test_l004_second_corp_name_is_buyer_even_if_personlike():
    assert looks_like_tax_invoice(Path("전자세금계산서.pdf"), "")
    assert extract_tax_invoice_buyer(_TAX_TEXT) == "홍길동상사"
    assert extract_tax_invoice_buyer("상호 (법인명) 갑\n상호 (법인명) 을") == "을"


def test_l004_company_extract_overrides_supplier(tmp_path: Path):
    src = tmp_path / "전자세금계산서.txt"
    src.write_text(_TAX_TEXT, encoding="utf-8")
    fields, notes = extract_company_from_file(src)
    assert fields["기업명"]["value"] == "홍길동상사"
    assert any("L004" in n for n in notes) or fields["기업명"]["raw_label"].startswith("(법인명)#2")


def test_l014_add_generated_table_header_shaded_body_plain():
    doc = Document()
    table = add_generated_table(doc, [["구분", "값"], ["매출", "100"], ["비용", "40"]])
    header_fill = GENERATED_TABLE_HEADER_FILL.upper()
    assert cell_shading_fill(table.rows[0].cells[0]) == header_fill
    assert table.rows[0].cells[0].paragraphs[0].runs[0].bold is True
    assert cell_shading_fill(table.rows[1].cells[0]) == ""
    assert table.rows[1].cells[0].paragraphs[0].runs[0].bold is not True
    assert cell_shading_fill(table.rows[2].cells[1]) == ""


def test_l014_unhwp_render_wires_generated_header_style():
    doc = Document()
    rows = [
        {
            "is_header": True,
            "cells": [{"content": [{"Text": {"text": "항목"}}], "colspan": 1, "rowspan": 1}],
        },
        {
            "is_header": False,
            "cells": [{"content": [{"Text": {"text": "매출"}}], "colspan": 1, "rowspan": 1}],
        },
    ]
    assert _render_unhwp_table(doc, rows)
    table = doc.tables[0]
    assert cell_shading_fill(table.rows[0].cells[0]) == GENERATED_TABLE_HEADER_FILL.upper()
    assert cell_shading_fill(table.rows[1].cells[0]) == ""


def test_l014_style_generated_does_not_run_on_untouched_form_table():
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "라벨"
    t.cell(0, 0).paragraphs[0].runs[0].bold = True
    assert cell_shading_fill(t.rows[0].cells[0]) == ""
    style_generated_table(t)
    assert cell_shading_fill(t.rows[0].cells[0]) == GENERATED_TABLE_HEADER_FILL.upper()


def test_l048_announcement_tuple_and_merge_pdfs(tmp_path: Path):
    stem = announcement_tuple_stem(
        yyyymmdd="2026-08-31", notice_name="한경련 더하기", doc_kind="증빙합본"
    )
    assert stem == "20260831_한경련더하기_증빙합본"
    a = _blank_pdf(tmp_path / "a.pdf", 1)
    b = _blank_pdf(tmp_path / "b.pdf", 2)
    dest = tmp_path / f"{stem}.pdf"
    merge_pdfs([a, b], dest)
    assert dest.is_file()
    assert len(PdfReader(str(dest)).pages) == 3


def test_l048_pipeline_merges_evidence_pdfs(tmp_path: Path):
    notice = tmp_path / "20260831 테스트공고"
    notice.mkdir()
    p1 = _blank_pdf(notice / "사업자등록증.pdf")
    p2 = _blank_pdf(notice / "납세증명.pdf")
    settings = _settings(tmp_path)
    storage = _FakeStorage(tmp_path)
    oa = OpenAIService(settings)
    prof = _profile(tmp_path)
    ps = _FakeProjectService(storage, prof, oa)
    storage.save_project_input(
        "p1",
        ProjectInput(
            template_id="t1",
            organization_profile={"기업명": "테스트(주)"},
            project_meta={
                "notice_folder": str(notice),
                "evidence_pdfs": [str(p1), str(p2)],
            },
        ),
    )
    pipeline = SubmissionPipeline(ps, EvaluationService(oa), storage, settings)
    fake = mock.Mock(submittable=True, fail_defects=0, results=[])
    with mock.patch(
        "auto_write.services.submission_orchestrator.run_acceptance",
        return_value=fake,
    ):
        report = pipeline.run(
            "p1", announcement_text="", enable_images=False, enable_notebooklm=False
        )
    merged = Path(report["merged_pdf"])
    assert merged.is_file()
    assert "증빙합본" in merged.name
    assert len(PdfReader(str(merged)).pages) == 2
    assert is_submit_layout_path(report["final_docx"])


def test_l049_submit_layout_dir_dated_and_undated(tmp_path: Path):
    dated = tmp_path / "20260831 한경련더하기"
    dated.mkdir()
    layout = build_submit_layout_dir(dated)
    assert layout == dated / "제출"
    assert is_dated_notice_folder(dated)
    assert is_submit_layout_path(layout / "신청서.hwpx")
    undated = tmp_path / "작업폴더"
    undated.mkdir()
    plain = build_submit_layout_dir(undated)
    assert plain == undated / "제출"
    wrapped = build_submit_layout_dir(undated, today="20260831", wrap_undated=True)
    assert wrapped == undated / "20260831 작업폴더" / "제출"
    assert is_submit_layout_path(wrapped)


def test_l050_pair_check_skips_draft_and_flags_final(tmp_path: Path):
    final = tmp_path / "신청서.hwpx"
    final.write_bytes(b"PK")
    assert missing_pdf_pair(final) is True
    draft = tmp_path / "신청서_DRAFT.hwpx"
    draft.write_bytes(b"PK")
    assert is_draft_artifact(draft)
    assert missing_pdf_pair(draft) is False
    (tmp_path / "신청서.pdf").write_bytes(b"%PDF")
    assert missing_pdf_pair(final) is False


def test_l072_quality_orchestrator_rolls_back_when_worse(tmp_path: Path):
    src = tmp_path / "in.docx"
    doc = Document()
    doc.add_paragraph("KEEPME_UNIQUE")
    doc.add_paragraph("※ 작성요령: 여기에 기재하세요")
    doc.save(str(src))

    from auto_write.services.doc_quality_score import QualityScore

    scores = [90.0, 40.0, 40.0, 40.0]

    def fake_score(*_a, **_k):
        total = scores.pop(0) if scores else 40.0
        return QualityScore(total=total, grade="x", passed=total >= 85, items=[])

    orch = DocumentQualityOrchestrator(tmp_path / "results")
    with mock.patch(
        "core.docx.services.document_quality_orchestrator.score_document",
        fake_score,
    ):
        res = orch.run(src, write_report=False)
    assert res.heuristic_rolled_back is True
    assert res.score_before >= res.score.total
    out = Document(res.output_docx)
    texts = [p.text for p in out.paragraphs]
    assert "KEEPME_UNIQUE" in texts
    assert any("작성요령" in t for t in texts)


def test_l072_quality_orchestrator_keeps_improvement(tmp_path: Path):
    src = tmp_path / "in.docx"
    Document().save(str(src))
    from auto_write.services.doc_quality_score import QualityScore

    scores = [40.0, 88.0, 88.0]

    def fake_score(*_a, **_k):
        total = scores.pop(0) if scores else 88.0
        return QualityScore(total=total, grade="x", passed=total >= 85, items=[])

    orch = DocumentQualityOrchestrator(tmp_path / "results")
    with mock.patch(
        "core.docx.services.document_quality_orchestrator.score_document",
        fake_score,
    ):
        res = orch.run(src, write_report=False)
    assert res.heuristic_rolled_back is False
    assert res.score.total >= 85


def test_l105_broken_bracket_description_rejected():
    broken = "---\nname: x\ndescription: [한글] 훅이 깨진다\n---\nbody\n"
    with pytest.raises(ValueError, match="L105"):
        parse_skill_frontmatter(broken)


def test_l105_folded_scalar_ok_and_repo_skills_parse():
    ok = "---\nname: demo\ndescription: >-\n  요청 원문 훅\n---\n"
    data = parse_skill_frontmatter(ok)
    assert "요청 원문" in data["description"]
    errors = collect_skill_frontmatter_errors(_REPO)
    assert errors == [], errors
