"""test_orch_confirm_markers.py — DocumentQualityOrchestrator._count_confirm_markers 회귀.

세로 병합셀 과다집계(row 단위 dedup) + lxml proxy id 재사용 오판을,
usage_acceptance.dedup_cells 단일출처 재사용으로 막았다(재발방지 1-B).
본문+표 합산, 병합 중복은 1회만 집계됨을 고정한다.
"""
from __future__ import annotations

from docx import Document as Docx

from auto_write.services.document_quality_orchestrator import (
    DocumentQualityOrchestrator as Orch,
)


def test_vertical_merge_counted_once():
    """세로 3행 병합셀의 [확인필요]를 row 수만큼(3회) 세지 않고 1회만 센다."""
    doc = Docx()
    t = doc.add_table(rows=3, cols=2)
    merged = t.cell(0, 0).merge(t.cell(2, 0))  # 세로 3행 병합
    merged.text = "[확인필요]"
    t.cell(0, 1).text = "[확인필요]"  # 일반 셀 1개
    # row 단위 dedup 이면 병합셀을 3번 세어 4가 된다 — 단일출처 dedup 으로 2 여야 함.
    assert Orch._count_confirm_markers(doc) == 2


def test_body_and_table_summed():
    """본문 단락 + 표 셀의 [확인필요]를 합산한다."""
    doc = Docx()
    doc.add_paragraph("설명 [확인필요] 문장")
    t = doc.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "[확인필요]"
    assert Orch._count_confirm_markers(doc) == 2


def test_no_markers_zero():
    doc = Docx()
    doc.add_paragraph("정상 문장")
    assert Orch._count_confirm_markers(doc) == 0
