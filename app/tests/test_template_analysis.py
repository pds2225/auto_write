from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from auto_write.analysis.docx_template import analyze_template


class TemplateAnalysisTests(unittest.TestCase):
    def test_detects_sections_tables_and_images(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "sample.docx"
            doc = Document()
            doc.add_paragraph("1. 사업 개요")
            doc.add_paragraph("2. 추진 전략")
            doc.add_paragraph("* 필요 시 칸을 늘려 작성할 수 있음")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "구분"
            table.cell(0, 1).text = "내용"
            table.cell(1, 0).text = "핵심 서비스"
            table.cell(1, 1).text = ""
            image_table = doc.add_table(rows=2, cols=1)
            image_table.cell(0, 0).text = "관련 이미지"
            image_table.cell(1, 0).text = "< 사진(이미지) 또는 설계도 제목 >"
            doc.save(path)

            profile = analyze_template(path)

            self.assertGreaterEqual(len(profile.sections), 2)
            self.assertTrue(any(table.cells for table in profile.tables))
            self.assertTrue(any(slot.anchor_type == "table_cell" for slot in profile.image_slots))
            self.assertTrue(any(question.question_id == "project_title" for question in profile.questions))
            guide_sections = [section for section in profile.sections if "필요 시" in section.label]
            self.assertFalse(guide_sections)
            self.assertTrue(any("이미지" in slot.label for slot in profile.image_slots))
            optional_image_slots = [slot for slot in profile.image_slots if "설계도 제목" in slot.label]
            self.assertFalse(optional_image_slots)

    def test_filters_out_consent_and_privacy_sections(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "consent_sample.docx"
            doc = Document()
            doc.add_paragraph("1. 사업 개요")
            doc.add_paragraph("제품 소개 및 시장 진입 전략")
            doc.add_paragraph("※ 작성요령: 아래 예시를 참고하여 기재")
            doc.add_paragraph("[별지 제4호] 개인정보 수집·이용 동의서")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "개인정보 수집·이용 동의서"
            table.cell(0, 1).text = "개인정보 수집·이용 동의서"
            table.cell(1, 0).text = "동의 여부"
            table.cell(1, 1).text = ""
            doc.save(path)

            profile = analyze_template(path)

            self.assertTrue(any("사업 개요" in section.label for section in profile.sections))
            self.assertFalse(any("개인정보" in section.label for section in profile.sections))
            self.assertFalse(any("작성요령" in section.label for section in profile.sections))
            self.assertFalse(any("개인정보" in table.label for table in profile.tables))
            self.assertFalse(any("작성요령" in table.label for table in profile.tables))
            self.assertFalse(any("개인정보" in question.label for question in profile.questions))
            self.assertFalse(any("작성요령" in question.label for question in profile.questions))

    def test_contextualizes_generic_table_and_image_labels(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "contextual_sample.docx"
            doc = Document()
            doc.add_paragraph("창업아이템 개요(요약)")
            doc.add_paragraph("◦")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "구분"
            table.cell(0, 1).text = "내용"
            table.cell(1, 0).text = "핵심 기능"
            table.cell(1, 1).text = ""
            image_table = doc.add_table(rows=1, cols=1)
            image_table.cell(0, 0).text = "이미지"
            doc.save(path)

            profile = analyze_template(path)

            self.assertFalse(any(section.label == "◦" for section in profile.sections))
            self.assertTrue(any(table.label.startswith("창업아이템 개요(요약) / 구분") for table in profile.tables))
            self.assertTrue(any(slot.label.startswith("창업아이템 개요(요약)") for slot in profile.image_slots))

    def test_filters_out_signature_and_confirmation_sections(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "signature_sample.docx"
            doc = Document()
            doc.add_paragraph("1. 사업 개요")
            doc.add_paragraph("■ 신청기업")
            doc.add_paragraph("상기 내용은 사실과 다름이 없음을 확인하며, 기업 추천합니다.")
            doc.add_paragraph("담 당 자 : (인)")
            doc.save(path)

            profile = analyze_template(path)

            self.assertTrue(any("사업 개요" in section.label for section in profile.sections))
            self.assertFalse(any("신청기업" in section.label for section in profile.sections))
            self.assertFalse(any("사실과 다름이 없음을 확인" in section.label for section in profile.sections))
            self.assertFalse(any("담 당 자" in section.label for section in profile.sections))


if __name__ == "__main__":
    unittest.main()
