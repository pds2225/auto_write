"""test_cross_form_single_summary.py — 단일 A→B 전사 결과의 비개발자용 요약 회귀.

배경(닫지 못했던 루프)
---------------------
단일 전사 CLI(``cross_form_fill.py`` fill/legacy)는 성공해도 stdout 에 **원본 JSON**
(``report.as_dict()``)을 통째로 덤프했다. 비개발자 사용자는 그 JSON 을 뜯어보지 않으면
"무엇이 채워졌는지 / 무엇을 확인해야 하는지 / 무엇을 직접 채워야 하는지"를 알 수 없었고,
특히 needs_confirm 후보를 ``--confirm`` 으로 확정하려면 어떤 ``타깃=소스`` 를 넣어야 하는지
JSON 에서 직접 찾아야 했다(사실상 막다른 길).

이 기능(format_single_summary_korean)
------------------------------------
성공 경로 기본 출력을 **사람이 읽는 한국어 요약**으로 바꾼다:
  - 자동으로 채운 칸(값·동의어 소스 라벨 포함)
  - 자동 체크한 선택칸
  - 확인 필요 칸 + **그대로 붙여 다시 실행할 수 있는 ``--confirm "타깃=소스"`` 명령**
  - 완성본에 값이 없어 비워둔 칸(직접 채움 안내)
  - 다음 행동 한 줄
``--json`` 을 주면 종전처럼 기계용 원본 JSON 을 낸다(하위호환·오류 경로는 항상 JSON).

핵심 효과(before→after):
  - before: 성공해도 raw JSON 벽 → 비개발자가 해석 불가, 확정 명령을 손으로 조립.
  - after:  "채움/확인필요/빈칸 + 복붙 확정 명령"을 바로 보고 사람 확인 루프를 닫는다.

안전 불변(이 테스트가 잠그는 것):
  - 읽기 전용 — 요약은 report 만 읽고 값을 지어내지 않는다(채움 수 = 실제 transcribed).
  - cp949(윈도우 한글 콘솔 기본) 인코딩 안전 — 이모지 없이 안전 문자만 사용.
  - 확정 명령이 실제로 동작한다(round-trip) — 제안한 ``--confirm`` 을 그대로 먹이면 채워짐.
"""

from __future__ import annotations

import io
import json
import re
from pathlib import Path

from docx import Document

from auto_write.services.cross_form_autofill import (
    AutofillReport,
    Match,
    autofill_from_source,
    format_single_summary_korean,
)
import cross_form_fill


# --- 도우미 ------------------------------------------------------------------

def _table_doc(path: Path, pairs: list[tuple[str, str]]) -> None:
    doc = Document()
    t = doc.add_table(rows=len(pairs), cols=2)
    for ri, (label, value) in enumerate(pairs):
        t.rows[ri].cells[0].text = label
        t.rows[ri].cells[1].text = value
    doc.save(str(path))


def _value_for(docx_path: Path, label: str) -> str:
    doc = Document(str(docx_path))
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if cells and cells[0].text.strip() == label:
                return cells[1].text.strip()
    return ""


def _match(target: str, value: str, *, source: str = "", conf: str = "high") -> Match:
    from auto_write.services.cross_form_autofill import _key
    return Match(
        target_label=target, normalized=_key(target),
        source_label=source, value=value, confidence=conf,
        table_index=0, row=0, value_cell=1,
    )


# --- 단위: 각 구획이 정확히 렌더된다 -----------------------------------------

def test_summary_lists_filled_confirm_and_blank() -> None:
    """채운 칸·확인 필요(+확정 명령)·빈칸이 각 구획으로 나온다."""
    rep = AutofillReport(source="a.docx", target="b.docx", output="out.docx", ok=True)
    rep.matches = [_match("기업명", "미래큐러스")]
    rep.transcribed = 1
    rep.needs_confirm = [{"target_label": "제품 명칭", "normalized": "제품명칭",
                          "candidates": ["제품명"], "confidence": "fuzzy"}]
    rep.unmatched_targets = [{"target_label": "생년월일", "normalized": "생년월일"}]

    text = format_single_summary_korean(rep)
    assert "[자동으로 채운 칸] 1개" in text
    assert "기업명 ← 미래큐러스" in text
    assert "[확인 필요]" in text and "제품 명칭" in text
    # 복붙 가능한 확정 명령이 원본 타깃 라벨 그대로 제시된다.
    assert '--confirm "제품 명칭=제품명"' in text
    assert "[빈칸]" in text and "생년월일" in text
    assert "다음:" in text


def test_summary_synonym_shows_source_label() -> None:
    """동의어 매칭(성명←대표자)은 소스 라벨을 병기, 정확일치는 병기 안 함."""
    rep = AutofillReport(source="a", target="b", output="o", ok=True, transcribed=2)
    rep.matches = [
        _match("성명", "홍길동", source="대표자"),   # 동의어 → 소스 라벨 병기
        _match("기업명", "미래큐러스", source="기업명"),  # 정확일치 → 병기 없음
    ]
    text = format_single_summary_korean(rep)
    assert "성명 ← 홍길동  (소스 라벨: 대표자)" in text
    assert "기업명 ← 미래큐러스" in text
    # 정확일치 줄에는 소스 라벨 병기가 없어야 한다.
    kiup_line = [ln for ln in text.splitlines() if "기업명 ←" in ln][0]
    assert "소스 라벨" not in kiup_line


def test_summary_preserved_cells_note() -> None:
    """matches 수 > transcribed 면 '채우지 않았습니다' 주석이 붙는다(원인 불가지)."""
    rep = AutofillReport(source="a", target="b", output="o", ok=True, transcribed=1)
    rep.matches = [_match("기업명", "미래큐러스"), _match("대표자", "홍길동")]
    text = format_single_summary_korean(rep)
    assert "[자동으로 채운 칸] 1개" in text
    assert "1칸은 채우지 않았습니다" in text


def test_summary_skipped_note_is_cause_agnostic() -> None:
    """단락 필드가 위치 미확정/빈값으로 미기입돼 skipped 로 잡혀도, '이미 값이
    있어'라고 원인을 단정하지 않는다(비개발자 오해 방지 — 코드리뷰 MEDIUM#1).

    표-셀 미기입만 '이미 값이 있음'이고, 단락(paragraph)·셀단락 미기입은 위치
    미확정/빈값 등 원인이 달라 단정 문구는 거짓 설명이 된다.
    """
    rep = AutofillReport(source="a", target="b", output="o", ok=True, transcribed=0)
    pm = Match(target_label="비고", normalized="비고", source_label="비고",
               value="홍길동", confidence="high", table_index=-1, row=-1,
               value_cell=-1, kind="paragraph", para_index=0)
    rep.matches = [pm]
    text = format_single_summary_korean(rep)
    assert "1칸은 채우지 않았습니다" in text
    assert "이미 값이 있어 그대로 두었습니다" not in text


def test_summary_checkbox_section() -> None:
    """자동 체크한 선택칸이 라벨·소스값과 함께 나온다."""
    rep = AutofillReport(source="a", target="b", output="o", ok=True,
                         checkbox_checked=1)
    rep.checkbox_groups = [{
        "label": "사업자 형태", "normalized": "사업자형태",
        "source_value": "개인사업자", "checked_option_index": 0,
        "confidence": "high"}]
    text = format_single_summary_korean(rep)
    assert "[자동 체크한 선택칸] 1개" in text
    assert "사업자 형태" in text and "개인사업자" in text
    # checked_option_index < 0 (미체크)은 목록에 안 나온다.
    rep2 = AutofillReport(source="a", target="b", output="o", ok=True)
    rep2.checkbox_groups = [{"label": "X", "normalized": "x", "source_value": "",
                             "checked_option_index": -1, "confidence": "low"}]
    assert "[자동 체크한 선택칸]" not in format_single_summary_korean(rep2)


def test_summary_caps_long_lists() -> None:
    """확인 필요·빈칸이 12개 초과면 '…외 N개'로 접는다(벽 방지)."""
    rep = AutofillReport(source="a", target="b", output="o", ok=True)
    rep.needs_confirm = [
        {"target_label": f"항목{i}", "normalized": f"항목{i}",
         "candidates": [f"후보{i}"], "confidence": "fuzzy"} for i in range(15)
    ]
    rep.unmatched_targets = [
        {"target_label": f"빈칸{i}", "normalized": f"빈칸{i}"} for i in range(20)
    ]
    text = format_single_summary_korean(rep)
    assert "[확인 필요] 비슷하지만 확실치 않아 비워둠 15개" in text
    assert "…외 3개" in text     # 15 - 12
    assert "[빈칸] 완성본에 값이 없어 비워둔 칸 20개" in text
    assert "…외 8개" in text     # 20 - 12


def test_summary_not_ok_shows_draft_notice() -> None:
    """ok=False 면 '아직 제출본 아님' + 사유가 나온다."""
    rep = AutofillReport(source="a", target="b", output="o", ok=False,
                         notes=["전사 0건 — 소스 3필드/타깃 0빈칸"])
    text = format_single_summary_korean(rep)
    assert "[아직 제출본 아님]" in text
    assert "전사 0건" in text


def test_summary_empty_body_hint() -> None:
    """ok 인데 채움/확인/빈칸이 전무하면 안내 문구를 준다(빈 결과 방지)."""
    rep = AutofillReport(source="a", target="b", output="o", ok=True)
    text = format_single_summary_korean(rep)
    assert "채울 빈칸이 없거나 이미 모두 채워져" in text


def test_summary_is_cp949_safe() -> None:
    """윈도우 한글 콘솔(cp949) 기본 인코딩에서 크래시 없이 출력 가능해야 한다.

    (이모지 U+1F4C4 등 보충평면 문자를 쓰면 cp949 에서 UnicodeEncodeError.)
    """
    rep = AutofillReport(source="a", target="b", output="o", ok=True, transcribed=1)
    rep.matches = [_match("성명", "홍길동", source="대표자")]
    rep.needs_confirm = [{"target_label": "주소", "normalized": "주소",
                          "candidates": ["자택주소"], "confidence": "fuzzy"}]
    rep.unmatched_targets = [{"target_label": "매출액", "normalized": "매출액"}]
    text = format_single_summary_korean(rep)
    text.encode("cp949")   # 예외가 나면 실패


def test_summary_confirm_escapes_special_label() -> None:
    """타깃 라벨에 = 나 " 가 있으면 깨지는 인라인 --confirm 대신 --confirm-file 안내.

    인라인 ``--confirm "타깃=소스"`` 는 파서가 첫 = 에서 분리·셸 따옴표가 붕괴해
    엉뚱한 칸을 채운다(적대적 코드리뷰 MEDIUM#1). 단일·배치 공용 _confirm_hint 방어.
    """
    rep = AutofillReport(source="a", target="b", output="o", ok=True)
    rep.needs_confirm = [{"target_label": "매출(전년=100)", "normalized": "매출",
                          "candidates": ["매출액"], "confidence": "fuzzy"}]
    text = format_single_summary_korean(rep)
    assert '--confirm "매출(전년=100)=매출액"' not in text  # 깨지는 명령 미방출
    assert "--confirm-file" in text                          # 값 손상 없는 대안 안내
    text.encode("cp949")   # 대안 안내도 cp949 안전


# --- round-trip: 제안한 --confirm 이 실제로 그 칸을 채운다(핵심 증명) --------

def test_summary_confirm_command_actually_fills(tmp_path: Path) -> None:
    """요약이 제안한 ``--confirm "타깃=소스"`` 를 그대로 먹이면 그 칸이 채워진다.

    비개발자가 요약만 보고(JSON 해석 없이) 사람 확인 루프를 닫을 수 있음을 증명.
    """
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _table_doc(src, [("제품명", "스마트센서")])
    _table_doc(tgt, [("제품명칭", "")])

    rep = autofill_from_source(src, tgt, out, use_ai=False)
    assert rep.transcribed == 0            # 퍼지라 자동 전사 안 됨
    summary = format_single_summary_korean(rep)

    m = re.search(r'--confirm "(.+?)=(.+?)"', summary)
    assert m, f"요약에 --confirm 제안이 없음:\n{summary}"
    target_label, source_label = m.group(1), m.group(2)

    out2 = tmp_path / "out2.docx"
    rep2 = autofill_from_source(
        src, tgt, out2, use_ai=False,
        confirmations={target_label: source_label})
    assert rep2.transcribed == 1
    assert _value_for(out2, "제품명칭") == "스마트센서"


# --- CLI: 기본은 사람 요약, --json 은 원본 JSON ------------------------------

def _run_single_capsys(argv: list[str], capsys) -> tuple[int, str]:
    rc = cross_form_fill.main(argv)
    return rc, capsys.readouterr().out


def test_cli_single_default_prints_human_summary(tmp_path: Path, capsys) -> None:
    """fill 기본 출력은 사람 요약(JSON 덤프가 아님)."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _table_doc(src, [("기업명", "미래큐러스")])
    _table_doc(tgt, [("기업명", "")])

    rc, stdout = _run_single_capsys(
        ["fill", "--source", str(src), "--target", str(tgt), "-o", str(out)],
        capsys)
    assert rc == 0
    assert "완성본 → 빈 양식 전사 결과" in stdout
    assert "[자동으로 채운 칸]" in stdout
    assert not stdout.lstrip().startswith("{")   # raw JSON 아님


def test_cli_single_json_flag_prints_json(tmp_path: Path, capsys) -> None:
    """--json 이면 기계용 원본 JSON(하위호환)."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _table_doc(src, [("기업명", "미래큐러스")])
    _table_doc(tgt, [("기업명", "")])

    rc, stdout = _run_single_capsys(
        ["fill", "--source", str(src), "--target", str(tgt), "-o", str(out),
         "--json"], capsys)
    assert rc == 0
    parsed = json.loads(stdout)   # 유효한 JSON 이어야 함
    assert parsed["ok"] is True
    assert parsed["transcribed"] == 1


def test_cli_legacy_mode_also_human_summary(tmp_path: Path, capsys) -> None:
    """subcommand 없는 레거시 경로도 기본은 사람 요약."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _table_doc(src, [("기업명", "미래큐러스")])
    _table_doc(tgt, [("기업명", "")])

    rc, stdout = _run_single_capsys(
        ["--source", str(src), "--target", str(tgt), "-o", str(out)], capsys)
    assert rc == 0
    assert "[자동으로 채운 칸]" in stdout


def test_cli_error_path_still_json(tmp_path: Path, capsys) -> None:
    """오류 경로(비지원 입력)는 --json 없이도 JSON 오류 리포트(하위호환)."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.pdf"          # 비지원 확장자
    out = tmp_path / "out.docx"
    _table_doc(src, [("기업명", "미래큐러스")])
    tgt.write_bytes(b"%PDF-1.4\nfake\n")

    rc, stdout = _run_single_capsys(
        ["fill", "--source", str(src), "--target", str(tgt), "-o", str(out)],
        capsys)
    assert rc == 2
    parsed = json.loads(stdout)
    assert parsed["ok"] is False
