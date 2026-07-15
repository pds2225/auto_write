"""SFT 데이터 레이어 P1 테스트 — 사람 수정(human_approved) 캡처.

폼 재제출이 직전 AI 초안을 고친 것이면 (before=AI, after=사람) 페어를 feedback 으로 기록.
- edited / draft_rejected 구분
- 첫 divergence 게이트: 이미 사람이 고친(현재값≠AI반영본) 답변은 재기록 안 함
- 비-AI 필드(reflected 에 없는 qid)는 무시
"""

from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from auto_write.config import Settings, ensure_directories
from auto_write.services import learning_store
from auto_write.services.evidence_service import EvidenceService
from auto_write.services.image_service import ImageService
from auto_write.services.openai_client import OpenAIService
from auto_write.services.project_service import ProjectService
from auto_write.services.qa_service import QAService
from auto_write.services.render_service import RenderService
from auto_write.storage import Storage


def test_feedback_append_and_load_roundtrip(tmp_path: Path) -> None:
    rec = {"project_id": "p1", "qid": "q1", "source": "user",
           "action_type": "edited", "feedback": {"before": "A", "after": "B"}}
    learning_store.append_feedback(rec, root=tmp_path)
    loaded = learning_store.load_feedback(root=tmp_path)
    assert loaded == [rec]


def _build_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("□ 일반현황")
    doc.add_paragraph("1. 문제 인식 (Problem)_필요성")
    tb = doc.add_table(rows=2, cols=2)
    tb.cell(0, 0).text = "항목"; tb.cell(0, 1).text = "내용"
    tb.cell(1, 0).text = "기업명"; tb.cell(1, 1).text = "○○기업"
    doc.save(path)


class HumanApprovedCaptureTests(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp = tempfile.TemporaryDirectory()
        root = Path(self.tmp.name)
        self.settings = Settings(
            app_root=root / "app", workspace_root=root / "workspace",
            template_root=root / "workspace" / "templates",
            project_root=root / "workspace" / "projects", results_root=root / "results",
            static_root=root / "app" / "auto_write" / "static",
            template_view_root=root / "app" / "auto_write" / "templates",
            host="127.0.0.1", port=8765,
            openai_api_key="", openai_model="gpt-4.1-mini", openai_search_model="gpt-4.1-mini",
            openai_image_model="gpt-image-1", anthropic_api_key="",
            anthropic_model="claude-sonnet-4-20250514", anthropic_search_model="claude-sonnet-4-20250514",
        )
        ensure_directories(self.settings)
        self.storage = Storage(self.settings)
        oa = OpenAIService(self.settings)
        self.service = ProjectService(
            storage=self.storage, openai_service=oa,
            evidence_service=EvidenceService(oa), image_service=ImageService(oa),
            render_service=RenderService(), qa_service=QAService(),
        )
        # 스파이: feedback 을 실제 workspace 대신 리스트로 포착.
        self.events: list[dict] = []
        self._orig_append = learning_store.append_feedback
        learning_store.append_feedback = lambda rec, root=None: self.events.append(rec)  # type: ignore

        sample = Path(self.tmp.name) / "s.docx"
        _build_docx(sample)
        profile = self.service.analyze_uploaded_template("s.docx", sample.read_bytes())
        self.project_id = self.service.create_project(profile.template_id, "P1")
        self.qid = "section_ai_1"

    def tearDown(self) -> None:
        learning_store.append_feedback = self._orig_append  # type: ignore
        self.tmp.cleanup()

    def _save(self, answers: dict) -> None:
        base = {"user_brief": "b", "user_notes": "n"}
        base.update(answers)
        self.service.save_project_form(
            project_id=self.project_id, answers=base,
            project_title="T", organization_name="O", evidence_topics="",
            reference_files=[],
        )

    def _write_snapshot(self, reflected: dict) -> None:
        sft = self.storage.project_dir(self.project_id) / "sft"
        sft.mkdir(parents=True, exist_ok=True)
        (sft / "ai_draft_snapshot.json").write_text(
            json.dumps({"reflected": reflected}, ensure_ascii=False), encoding="utf-8"
        )

    def test_edited_answer_captured_as_pair(self) -> None:
        self._save({self.qid: "AI 초안값"})          # 이전 세대 저장값
        self._write_snapshot({self.qid: "AI 초안값"})  # P0 가 남긴 AI 반영본
        self.events.clear()
        self._save({self.qid: "사람이 고친 값"})        # 재제출 = 사람 수정
        edits = [e for e in self.events if e["qid"] == self.qid]
        self.assertEqual(len(edits), 1)
        self.assertEqual(edits[0]["action_type"], "edited")
        self.assertEqual(edits[0]["feedback"]["before"], "AI 초안값")
        self.assertEqual(edits[0]["feedback"]["after"], "사람이 고친 값")
        self.assertEqual(edits[0]["source"], "user")

    def test_emptied_answer_is_draft_rejected(self) -> None:
        self._save({self.qid: "AI 초안값"})
        self._write_snapshot({self.qid: "AI 초안값"})
        self.events.clear()
        self._save({self.qid: ""})  # 빈값 재제출 = AI 초안 거부
        edits = [e for e in self.events if e["qid"] == self.qid]
        self.assertEqual(len(edits), 1)
        self.assertEqual(edits[0]["action_type"], "draft_rejected")

    def test_already_human_edited_not_repaired(self) -> None:
        # 현재 저장값(사람 v1)이 AI 반영본과 다르면 첫 divergence 아님 → 재기록 안 함.
        self._save({self.qid: "사람 v1"})
        self._write_snapshot({self.qid: "원래 AI값"})  # 저장값(사람 v1) != 반영본
        self.events.clear()
        self._save({self.qid: "사람 v2"})
        self.assertEqual([e for e in self.events if e["qid"] == self.qid], [])

    def test_non_ai_field_ignored(self) -> None:
        self._save({self.qid: "AI 초안값", "manual_only": "손입력"})
        self._write_snapshot({self.qid: "AI 초안값"})  # manual_only 는 reflected 에 없음
        self.events.clear()
        self._save({self.qid: "AI 초안값", "manual_only": "손입력 수정"})
        self.assertEqual([e for e in self.events if e["qid"] == "manual_only"], [])


if __name__ == "__main__":
    unittest.main()
