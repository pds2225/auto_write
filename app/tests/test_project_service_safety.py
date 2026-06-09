from __future__ import annotations

import json
import os
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from auto_write.config import Settings, ensure_directories
from auto_write.models import ProjectInput, QuestionProfile, ReferenceFile, TableCellProfile, TableProfile, TemplateProfile
from auto_write.services.evidence_service import EvidenceService
from auto_write.services.image_service import ImageService
from auto_write.services.openai_client import OpenAIService
from auto_write.services.project_service import ProjectService
from auto_write.services.qa_service import QAService
from auto_write.services.render_service import RenderService
from auto_write.storage import Storage


def build_settings(root: Path, reference_library_dir: Path | None = None) -> Settings:
    app_root = root / "app"
    workspace_root = root / "workspace"
    return Settings(
        app_root=app_root,
        workspace_root=workspace_root,
        template_root=workspace_root / "templates",
        project_root=workspace_root / "projects",
        results_root=root / "results",
        static_root=app_root / "auto_write" / "static",
        template_view_root=app_root / "auto_write" / "templates",
        host="127.0.0.1",
        port=8765,
        openai_api_key="",
        openai_model="gpt-4.1-mini",
        openai_search_model="gpt-4.1-mini",
        openai_image_model="gpt-image-1",
        anthropic_api_key="",
        anthropic_model="claude-sonnet-4-20250514",
        anthropic_search_model="claude-sonnet-4-20250514",
        reference_library_dir=reference_library_dir,
    )


class ProjectServiceSafetyTests(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp_dir = tempfile.TemporaryDirectory()
        root = Path(self.tmp_dir.name)
        self.settings = build_settings(root)
        ensure_directories(self.settings)
        self.storage = Storage(self.settings)
        openai_service = OpenAIService(self.settings)
        self.service = ProjectService(
            storage=self.storage,
            openai_service=openai_service,
            evidence_service=EvidenceService(openai_service),
            image_service=ImageService(openai_service),
            render_service=RenderService(),
            qa_service=QAService(),
        )

    def tearDown(self) -> None:
        self.tmp_dir.cleanup()

    def _save_template_profile(self, template_id: str) -> TemplateProfile:
        profile = TemplateProfile(
            template_id=template_id,
            template_name="sample.docx",
            source_docx=str(self.storage.template_dir(template_id) / "sample.docx"),
        )
        self.storage.template_dir(template_id).mkdir(parents=True, exist_ok=True)
        self.storage.save_template_profile(profile)
        return profile

    def test_finalize_template_rejects_missing_docx(self):
        profile = self._save_template_profile("tpl_missing")
        raw_json = json.dumps(profile.model_dump(), ensure_ascii=False)
        with self.assertRaises(ValueError) as ctx:
            self.service.finalize_template("tpl_missing", raw_json)
        self.assertIn("원본 DOCX 파일", str(ctx.exception))

    def test_save_project_form_rejects_path_like_filename(self):
        self._save_template_profile("tpl_safe")
        project_id = self.service.create_project("tpl_safe", "테스트")
        with self.assertRaises(ValueError) as ctx:
            self.service.save_project_form(
                project_id=project_id,
                answers={},
                project_title="p",
                organization_name="o",
                evidence_topics="",
                reference_files=[("../evil.txt", b"bad")],
            )
        self.assertIn("경로 구분자", str(ctx.exception))

    def test_save_project_form_deduplicates_same_filename(self):
        self._save_template_profile("tpl_dup")
        project_id = self.service.create_project("tpl_dup", "중복 파일명 테스트")
        project_input = self.service.save_project_form(
            project_id=project_id,
            answers={},
            project_title="p",
            organization_name="o",
            evidence_topics="",
            reference_files=[("memo.txt", b"alpha"), ("memo.txt", b"beta")],
        )
        names = [ref.file_name for ref in project_input.references]
        self.assertEqual(len(names), 2)
        self.assertEqual(names[0], "memo.txt")
        self.assertEqual(names[1], "memo_2.txt")
        for ref in project_input.references:
            self.assertTrue(Path(ref.saved_path).exists())

    def test_save_project_form_preserves_existing_references_without_new_upload(self):
        self._save_template_profile("tpl_keep_refs")
        project_id = self.service.create_project("tpl_keep_refs", "기존 참고자료 유지 테스트")
        first_input = self.service.save_project_form(
            project_id=project_id,
            answers={"user_brief": "초기 입력"},
            project_title="초기 과제",
            organization_name="초기 기관",
            evidence_topics="",
            reference_files=[("initial.txt", b"seed reference")],
        )
        self.assertEqual(len(first_input.references), 1)
        first_ref = first_input.references[0]

        second_input = self.service.save_project_form(
            project_id=project_id,
            answers={"user_notes": "재실행"},
            project_title="재실행 과제",
            organization_name="재실행 기관",
            evidence_topics="",
            reference_files=[],
        )

        self.assertEqual(len(second_input.references), 1)
        self.assertEqual(second_input.references[0].file_name, first_ref.file_name)
        self.assertEqual(second_input.references[0].saved_path, first_ref.saved_path)
        self.assertTrue(Path(second_input.references[0].saved_path).exists())

    def test_save_project_form_preserves_non_submitted_answers_and_allows_explicit_clear(self):
        self._save_template_profile("tpl_keep_answers")
        project_id = self.service.create_project("tpl_keep_answers", "기존 답변 유지 테스트")
        first_input = self.service.save_project_form(
            project_id=project_id,
            answers={"hidden_field": "기존 값", "user_brief": "초기 입력"},
            project_title="초기 과제",
            organization_name="초기 기관",
            evidence_topics="",
            reference_files=[],
        )
        self.assertEqual(first_input.answers.get("hidden_field"), "기존 값")

        second_input = self.service.save_project_form(
            project_id=project_id,
            answers={"user_notes": "재실행 입력"},
            project_title="재실행 과제",
            organization_name="재실행 기관",
            evidence_topics="",
            reference_files=[],
        )
        self.assertEqual(second_input.answers.get("hidden_field"), "기존 값")
        self.assertEqual(second_input.answers.get("user_notes"), "재실행 입력")

        third_input = self.service.save_project_form(
            project_id=project_id,
            answers={"hidden_field": "   "},
            project_title="재실행 과제",
            organization_name="재실행 기관",
            evidence_topics="",
            reference_files=[],
        )
        self.assertNotIn("hidden_field", third_input.answers)

    def test_save_project_form_preserves_project_meta_when_blank_values_submitted(self):
        self._save_template_profile("tpl_keep_meta")
        project_id = self.service.create_project("tpl_keep_meta", "메타 유지 테스트")
        first_input = self.service.save_project_form(
            project_id=project_id,
            answers={"user_brief": "초기"},
            project_title="초기 과제명",
            organization_name="초기 기관명",
            evidence_topics="",
            reference_files=[],
        )
        self.assertEqual(first_input.project_meta.get("project_title"), "초기 과제명")
        self.assertEqual(first_input.organization_profile.get("name"), "초기 기관명")

        second_input = self.service.save_project_form(
            project_id=project_id,
            answers={"user_notes": "재실행"},
            project_title="   ",
            organization_name="",
            evidence_topics="",
            reference_files=[],
        )
        self.assertEqual(second_input.project_meta.get("project_title"), "초기 과제명")
        self.assertEqual(second_input.organization_profile.get("name"), "초기 기관명")

    def test_render_docx_previews_treats_existing_png_as_success_even_on_nonzero_exit(self):
        docx_path = Path(self.tmp_dir.name) / "sample.docx"
        Document().save(docx_path)
        preview_dir = Path(self.tmp_dir.name) / "preview"
        script_path = Path(self.tmp_dir.name) / "fake_render.py"
        script_path.write_text(
            "\n".join(
                [
                    "from pathlib import Path",
                    "import sys",
                    "out_dir = Path(sys.argv[sys.argv.index('--output_dir') + 1])",
                    "out_dir.mkdir(parents=True, exist_ok=True)",
                    "png_bytes = b'\\x89PNG\\r\\n\\x1a\\n' + b'0' * 32",
                    "(out_dir / 'page-1.png').write_bytes(png_bytes)",
                    "print('Pages rendered to ' + str(out_dir))",
                    "raise SystemExit(1)",
                ]
            ),
            encoding="utf-8",
        )
        with patch.dict(os.environ, {"AUTO_WRITE_RENDER_DOCX_SCRIPT": str(script_path)}):
            result = self.service._render_docx_previews(docx_path, preview_dir)

        self.assertEqual(result["status"], "passed")
        self.assertEqual(result["page_count"], 1)
        self.assertEqual(result["pages"], ["page-1.png"])
        self.assertTrue(any("종료 코드" in item for item in result["warnings"]))

    def test_save_project_form_stores_selected_writing_provider_and_model(self):
        self._save_template_profile("tpl_writing_model")
        project_id = self.service.create_project("tpl_writing_model", "모델 선택 테스트")

        first_input = self.service.save_project_form(
            project_id=project_id,
            answers={},
            project_title="초기 과제명",
            organization_name="초기 기관명",
            evidence_topics="",
            reference_files=[],
            writing_provider="anthropic",
            writing_model="claude-opus-4-7",
        )
        self.assertEqual(first_input.project_meta.get("writing_provider"), "anthropic")
        self.assertEqual(first_input.project_meta.get("writing_model"), "claude-opus-4-7")

        second_input = self.service.save_project_form(
            project_id=project_id,
            answers={},
            project_title="초기 과제명",
            organization_name="초기 기관명",
            evidence_topics="",
            reference_files=[],
            writing_provider="",
            writing_model="",
        )
        self.assertEqual(second_input.project_meta.get("writing_provider"), "anthropic")
        self.assertEqual(second_input.project_meta.get("writing_model"), "claude-opus-4-7")

    def test_save_project_form_deduplicates_existing_reference_entries_by_path(self):
        self._save_template_profile("tpl_ref_dedupe")
        project_id = self.service.create_project("tpl_ref_dedupe", "참고자료 중복 제거 테스트")
        initial = self.service.save_project_form(
            project_id=project_id,
            answers={},
            project_title="초기 과제",
            organization_name="초기 기관",
            evidence_topics="",
            reference_files=[("dup.txt", b"same ref")],
        )
        self.assertEqual(len(initial.references), 1)
        ref = initial.references[0]

        duplicated = ProjectInput(
            template_id=initial.template_id,
            project_meta=initial.project_meta,
            organization_profile=initial.organization_profile,
            answers=initial.answers,
            references=[ref, ReferenceFile.model_validate(ref.model_dump())],
            evidence_requests=initial.evidence_requests,
            image_requests=initial.image_requests,
        )
        self.storage.save_project_input(project_id, duplicated)

        saved = self.service.save_project_form(
            project_id=project_id,
            answers={},
            project_title="초기 과제",
            organization_name="초기 기관",
            evidence_topics="",
            reference_files=[],
        )
        self.assertEqual(len(saved.references), 1)
        self.assertEqual(saved.references[0].saved_path, ref.saved_path)

    def test_generate_inserts_section_when_anchor_has_no_blank_paragraph(self):
        sample_docx = Path(self.tmp_dir.name) / "heading_only_template.docx"
        doc = Document()
        doc.add_paragraph("1. 사업 개요")
        doc.save(sample_docx)

        profile = self.service.analyze_uploaded_template("heading_only_template.docx", sample_docx.read_bytes())
        project_id = self.service.create_project(profile.template_id, "문단 생성 테스트")
        self.service.save_project_form(
            project_id=project_id,
            answers={"section_001_1": "앵커 아래 빈 문단이 없어도 본문이 생성되어야 합니다."},
            project_title="문단 생성 테스트",
            organization_name="테스트 기관",
            evidence_topics="",
            reference_files=[],
        )

        artifacts = self.service.generate(project_id)
        qa_report = json.loads(Path(artifacts.qa_report).read_text(encoding="utf-8"))
        transfer = json.loads(Path(artifacts.transfer_report).read_text(encoding="utf-8"))
        output_text = "\n".join(paragraph.text for paragraph in Document(artifacts.output_docx).paragraphs)

        self.assertIn("앵커 아래 빈 문단이 없어도", output_text)
        self.assertEqual(qa_report.get("error_count"), 0)
        self.assertGreaterEqual(transfer.get("fill_ratio", {}).get("sections", 0), 1.0)

    def test_generate_fills_empty_fields_and_writes_benchmark_report(self):
        sample_docx = Path(self.tmp_dir.name) / "template_sample.docx"
        doc = Document()
        doc.add_paragraph("1. 사업 개요")
        doc.add_paragraph("2. 추진 계획")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "구분"
        table.cell(0, 1).text = "내용"
        table.cell(1, 0).text = "실행안"
        table.cell(1, 1).text = ""
        doc.save(sample_docx)

        profile = self.service.analyze_uploaded_template("template_sample.docx", sample_docx.read_bytes())
        project_id = self.service.create_project(profile.template_id, "자동채움 테스트")
        self.service.save_project_form(
            project_id=project_id,
            answers={},
            project_title="테스트 과제",
            organization_name="테스트 기관",
            evidence_topics="",
            reference_files=[],
        )

        artifacts = self.service.generate(project_id)
        generated_input = self.storage.load_project_input(project_id)
        self.assertGreater(len(generated_input.answers), 0)

        self.assertTrue(Path(artifacts.output_docx).exists())
        self.assertTrue(Path(artifacts.qa_report).exists())
        self.assertTrue(Path(artifacts.sources).exists())
        self.assertTrue(Path(artifacts.benchmark_compare).exists())
        self.assertTrue(Path(artifacts.transfer_report).exists())

        compare = json.loads(Path(artifacts.benchmark_compare).read_text(encoding="utf-8"))
        self.assertEqual(compare.get("status"), "ok")
        self.assertIn("generated_metrics", compare)
        self.assertIn("result", compare)
        transfer = json.loads(Path(artifacts.transfer_report).read_text(encoding="utf-8"))
        self.assertEqual(transfer.get("status"), "ok")
        self.assertIn("fill_ratio", transfer)
        self.assertIn("image_slot_coverage", transfer)
        self.assertIn("evidence_usage", transfer)
        self.assertIn("qa_summary", transfer)
        self.assertEqual(transfer.get("mode"), "standard")

    def test_library_snippet_used_for_section_fallback(self):
        library_dir = Path(self.tmp_dir.name) / "library"
        library_dir.mkdir(parents=True, exist_ok=True)
        (library_dir / "sample.txt").write_text(
            "시장 분석 결과를 바탕으로 핵심 고객군을 정의하고 차별화된 실행 전략을 단계적으로 수행합니다.",
            encoding="utf-8",
        )
        self.settings = build_settings(Path(self.tmp_dir.name), reference_library_dir=library_dir)
        ensure_directories(self.settings)
        self.storage = Storage(self.settings)
        openai_service = OpenAIService(self.settings)
        self.service = ProjectService(
            storage=self.storage,
            openai_service=openai_service,
            evidence_service=EvidenceService(openai_service),
            image_service=ImageService(openai_service),
            render_service=RenderService(),
            qa_service=QAService(),
        )

        sample_docx = Path(self.tmp_dir.name) / "market_sample.docx"
        doc = Document()
        doc.add_paragraph("1. 시장 분석")
        doc.add_paragraph("2. 추진 전략")
        doc.save(sample_docx)

        profile = self.service.analyze_uploaded_template("market_sample.docx", sample_docx.read_bytes())
        project_id = self.service.create_project(profile.template_id, "라이브러리 활용 테스트")
        self.service.save_project_form(
            project_id=project_id,
            answers={},
            project_title="시장 테스트",
            organization_name="테스트 기관",
            evidence_topics="",
            reference_files=[],
        )
        self.service.generate(project_id)
        generated_input = self.storage.load_project_input(project_id)

        section_values = [
            value
            for key, value in generated_input.answers.items()
            if key.startswith("section_")
        ]
        self.assertTrue(any("핵심 고객군" in str(text) for text in section_values))

    def test_collect_missing_questions_returns_all_empty_when_required_only_false(self):
        profile = TemplateProfile(
            template_id="tpl_q",
            template_name="q.docx",
            source_docx="q.docx",
            questions=[
                QuestionProfile(
                    question_id="section_required",
                    label="필수 문단",
                    required=True,
                    target={"kind": "section", "field_id": "section_required"},
                ),
                QuestionProfile(
                    question_id="table_optional",
                    label="선택 표 셀",
                    required=False,
                    target={"kind": "table_cell", "table_id": "table_01", "cell_id": "table_optional"},
                ),
            ],
        )
        answers = {"section_required": "입력됨"}

        required_only = self.service._collect_missing_questions(profile, answers, required_only=True)
        all_missing = self.service._collect_missing_questions(profile, answers, required_only=False)

        self.assertEqual(len(required_only), 0)
        self.assertEqual([item["question_id"] for item in all_missing], ["table_optional"])

    def test_draft_missing_answers_routes_sections_to_anthropic_only(self):
        calls: list[dict[str, object]] = []

        def fake_draft_missing_answers(
            questions: list[dict[str, object]],
            context: str,
            provider_override: str | None = None,
            model_override: str | None = None,
            strict_preserve: bool = False,
        ) -> dict[str, str]:
            calls.append(
                {
                    "provider_override": provider_override,
                    "model_override": model_override,
                    "question_ids": [str(question.get("question_id", "")) for question in questions],
                }
            )
            return {str(question.get("question_id", "")): "생성 결과" for question in questions}

        self.service.openai_service.draft_missing_answers = fake_draft_missing_answers  # type: ignore[method-assign]
        self.service.openai_service.provider_available = lambda provider: provider == "anthropic"  # type: ignore[method-assign]
        self.service.openai_service._anthropic_writing_model = "claude-opus-4-1-20250805"

        missing = [
            {
                "question_id": "section_001",
                "label": "1. 사업 개요",
                "target": {"kind": "section", "field_id": "section_001"},
            },
            {
                "question_id": "table_01_r1_c1",
                "label": "사업비",
                "target": {"kind": "table_cell", "table_id": "table_01", "cell_id": "table_01_r1_c1"},
            },
        ]

        drafted = self.service._draft_missing_answers_in_chunks(missing, "테스트 문맥")

        self.assertEqual(drafted["section_001"], "생성 결과")
        self.assertEqual(drafted["table_01_r1_c1"], "생성 결과")
        self.assertEqual(len(calls), 2)
        self.assertEqual(calls[0]["provider_override"], "anthropic")
        self.assertEqual(calls[0]["question_ids"], ["section_001"])
        self.assertEqual(calls[1]["provider_override"], None)
        self.assertEqual(calls[1]["question_ids"], ["table_01_r1_c1"])

    def test_draft_missing_answers_honors_selected_provider_and_model_for_all_chunks(self):
        calls: list[dict[str, object]] = []

        def fake_draft_missing_answers(
            questions: list[dict[str, object]],
            context: str,
            provider_override: str | None = None,
            model_override: str | None = None,
            strict_preserve: bool = False,
        ) -> dict[str, str]:
            calls.append(
                {
                    "provider_override": provider_override,
                    "model_override": model_override,
                    "question_ids": [str(question.get("question_id", "")) for question in questions],
                }
            )
            return {str(question.get("question_id", "")): "생성 결과" for question in questions}

        self.service.openai_service.draft_missing_answers = fake_draft_missing_answers  # type: ignore[method-assign]
        self.service.openai_service.provider_available = lambda provider: provider == "anthropic"  # type: ignore[method-assign]

        missing = [
            {
                "question_id": "section_001",
                "label": "1. 사업 개요",
                "target": {"kind": "section", "field_id": "section_001"},
            },
            {
                "question_id": "table_01_r1_c1",
                "label": "사업비",
                "target": {"kind": "table_cell", "table_id": "table_01", "cell_id": "table_01_r1_c1"},
            },
        ]

        drafted = self.service._draft_missing_answers_in_chunks(
            missing,
            "테스트 문맥",
            writing_provider="anthropic",
            writing_model="claude-opus-4-7",
        )

        self.assertEqual(drafted["section_001"], "생성 결과")
        self.assertEqual(drafted["table_01_r1_c1"], "생성 결과")
        self.assertEqual(len(calls), 2)
        self.assertEqual(calls[0]["provider_override"], "anthropic")
        self.assertEqual(calls[1]["provider_override"], "anthropic")
        self.assertEqual(calls[0]["model_override"], "claude-opus-4-7")
        self.assertEqual(calls[1]["model_override"], "claude-opus-4-7")

    def test_transfer_mode_enabled_when_multiple_references_exist(self):
        sample_docx = Path(self.tmp_dir.name) / "transfer_template.docx"
        doc = Document()
        doc.add_paragraph("1. 사업 개요")
        doc.add_paragraph("2. 시장 분석")
        doc.save(sample_docx)

        profile = self.service.analyze_uploaded_template("transfer_template.docx", sample_docx.read_bytes())
        project_id = self.service.create_project(profile.template_id, "전환 모드 테스트")
        self.service.save_project_form(
            project_id=project_id,
            answers={},
            project_title="테스트 과제",
            organization_name="테스트 기관",
            evidence_topics="",
            reference_files=[
                ("source_a.txt", "시장 분석과 고객 문제 정의".encode("utf-8")),
                ("source_b.txt", "실행 전략과 일정 수립".encode("utf-8")),
            ],
        )

        artifacts = self.service.generate(project_id)
        transfer = json.loads(Path(artifacts.transfer_report).read_text(encoding="utf-8"))
        self.assertEqual(transfer.get("mode"), "transfer")
        self.assertEqual(transfer.get("evidence_usage", {}).get("reference_files"), 2)

    def test_filter_missing_skips_large_matrix_table_cells_in_transfer_mode(self):
        cells = [
            TableCellProfile(
                cell_id=f"table_01_r{i}_c1",
                label=f"구분 / 항목{i}",
                row=i,
                cell=1,
                required=False,
                row_header="구분",
                col_header=f"항목{i}",
            )
            for i in range(25)
        ]
        profile = TemplateProfile(
            template_id="tpl_matrix",
            template_name="matrix.docx",
            source_docx="matrix.docx",
            tables=[
                TableProfile(
                    table_id="table_01",
                    label="구분",
                    table_index=0,
                    row_count=10,
                    col_count=5,
                    cells=cells,
                )
            ],
            questions=[
                QuestionProfile(
                    question_id=cell.cell_id,
                    label=cell.label,
                    required=False,
                    target={"kind": "table_cell", "table_id": "table_01", "cell_id": cell.cell_id},
                )
                for cell in cells
            ],
        )

        missing = self.service._collect_missing_questions(profile, answers={}, required_only=False)
        filtered = self.service._filter_missing_for_autofill(profile, missing, transfer_mode=True)
        self.assertEqual(len(filtered), 0)

    def test_filter_missing_skips_signature_sections_and_sample_table_cells(self):
        profile = TemplateProfile(
            template_id="tpl_filtering",
            template_name="filtering.docx",
            source_docx="filtering.docx",
            tables=[
                TableProfile(
                    table_id="table_01",
                    label="사업 개요 / 핵심 전략",
                    table_index=0,
                    row_count=2,
                    col_count=2,
                    cells=[
                        TableCellProfile(
                            cell_id="table_bad",
                            label="참가구분 / OO.OO.OO / 법인기업",
                            row=1,
                            cell=1,
                            row_header="OO.OO.OO",
                            col_header="법인기업",
                        ),
                        TableCellProfile(
                            cell_id="table_good",
                            label="사업 개요 / 핵심 전략 / 실행 방안",
                            row=1,
                            cell=0,
                            row_header="핵심 전략",
                            col_header="실행 방안",
                        ),
                    ],
                )
            ],
            questions=[
                QuestionProfile(
                    question_id="section_bad",
                    label="담 당 자 : (인)",
                    target={"kind": "section", "field_id": "section_bad"},
                ),
                QuestionProfile(
                    question_id="section_good",
                    label="1. 사업 개요",
                    target={"kind": "section", "field_id": "section_good"},
                ),
                QuestionProfile(
                    question_id="table_bad",
                    label="참가구분 / OO.OO.OO / 법인기업",
                    target={"kind": "table_cell", "table_id": "table_01", "cell_id": "table_bad"},
                ),
                QuestionProfile(
                    question_id="table_good",
                    label="사업 개요 / 핵심 전략 / 실행 방안",
                    target={"kind": "table_cell", "table_id": "table_01", "cell_id": "table_good"},
                ),
            ],
        )

        filtered = self.service._filter_missing_for_autofill(profile, profile.model_dump()["questions"], transfer_mode=False)
        filtered_ids = {item["question_id"] for item in filtered}

        self.assertIn("section_good", filtered_ids)
        self.assertIn("table_good", filtered_ids)
        self.assertNotIn("section_bad", filtered_ids)
        self.assertNotIn("table_bad", filtered_ids)

    # ---- STEP2: 빈셀 미입력(item 5) ----
    def test_fallback_table_text_no_number_fabrication(self):
        """금액/예산 셀은 임의 수치('1,000(추정)') 대신 [확인필요] 로 표시(P2)."""
        self.assertEqual(self.service._fallback_table_text("과제", "사업비 금액"), "[확인필요]")
        self.assertEqual(self.service._fallback_table_text("과제", "예산 비용"), "[확인필요]")
        self.assertNotIn("1,000", self.service._fallback_table_text("과제", "매출 원"))
        # 산문 분기(일정/목표/담당)는 보존 — 과대치환 방지(P1)
        self.assertNotEqual(self.service._fallback_table_text("과제", "추진 일정"), "[확인필요]")
        self.assertNotEqual(self.service._fallback_table_text("과제", "담당 인력"), "[확인필요]")

    def test_filter_missing_surfaces_required_dropped_cell(self):
        """와이드 표에서 행/열 의미가 없어 드롭되는 '필수' 셀은 dropped_required 로 수집된다."""
        cells = []
        questions = []
        # 필수·col_header 누락 셀(드롭 대상이지만 surface 돼야 함)
        cells.append(TableCellProfile(cell_id="cell_req", label="사업비 총액",
                                      row=1, cell=1, required=True,
                                      row_header="사업비", col_header=""))
        questions.append(QuestionProfile(question_id="cell_req", label="사업비 총액", required=True,
                                         target={"kind": "table_cell", "table_id": "table_big", "cell_id": "cell_req"}))
        # 나머지 11개는 정상 헤더 보유(표 크기 > 10 확보)
        for i in range(11):
            cid = f"cell_ok_{i}"
            cells.append(TableCellProfile(cell_id=cid, label=f"항목{i} / 값",
                                          row=i + 2, cell=1, required=False,
                                          row_header=f"행{i}", col_header="값"))
            questions.append(QuestionProfile(question_id=cid, label=f"항목{i} / 값", required=False,
                                             target={"kind": "table_cell", "table_id": "table_big", "cell_id": cid}))
        profile = TemplateProfile(
            template_id="tpl_big", template_name="big.docx", source_docx="big.docx",
            tables=[TableProfile(table_id="table_big", label="사업비", table_index=0,
                                 row_count=12, col_count=1, cells=cells)],
            questions=questions,
        )
        q_dicts = profile.model_dump()["questions"]
        dropped: list[dict] = []
        filtered = self.service._filter_missing_for_autofill(profile, q_dicts, transfer_mode=True, dropped_required=dropped)
        filtered_ids = {q["question_id"] for q in filtered}
        dropped_ids = {q.get("question_id") for q in dropped}
        self.assertNotIn("cell_req", filtered_ids)      # 자동작성에서는 제외
        self.assertIn("cell_req", dropped_ids)           # 그러나 surface 대상으로 수집됨


if __name__ == "__main__":
    unittest.main()
