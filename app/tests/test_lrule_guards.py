# test_lrule_guards.py — mechanized callable wiring
from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from auto_write.domains.domain_classifier import Domain
from auto_write.domains.pipeline_gate import run_to_final
from auto_write.services.lrule_enforcer import LRuleEnforcer, enforce_lrules, rule_code
from auto_write.services.lrule_guards import build_lrule_guards


def _bp_doc(path: Path, extra: str = "") -> None:
    doc = Document()
    doc.add_heading("사업계획서", level=1)
    doc.add_paragraph("1. 문제 인식 (Problem)")
    doc.add_paragraph("시장에서 기존 솔루션의 한계를 발견하고 새로운 접근이 필요합니다.")
    if extra:
        doc.add_paragraph(extra)
    doc.save(str(path))


def test_build_lrule_guards_covers_every_mechanized(tmp_path):
    path = tmp_path / "bp.docx"
    _bp_doc(path)
    guards = build_lrule_guards(path)
    enforcer = LRuleEnforcer()
    mech = [l for l in enforcer._lessons if l.get("category") == "mechanized"]
    coverage = json.loads(
        (Path(__file__).resolve().parent / "lessons_coverage.json").read_text(
            encoding="utf-8"
        )
    )
    assert len(mech) == coverage["counts"]["mechanized"]
    for lesson in mech:
        assert lesson["id"] in guards, lesson["id"]
        assert rule_code(lesson["id"]) in guards
        assert "passed" in guards[lesson["id"]]


def test_l009_fails_on_unresolved_marker(tmp_path):
    path = tmp_path / "bp.docx"
    _bp_doc(path, extra="매출은 [확인필요] 입니다.")
    guards = build_lrule_guards(path)
    assert guards["L009"]["passed"] is False
    report = enforce_lrules(
        domain=Domain.BUSINESS_PLAN, artifact_path=path, guards=guards
    )
    l009 = next(r for r in report.rules if rule_code(r["id"]) == "L009")
    assert l009["status"] == "FAIL"


def test_run_to_final_auto_guards_zero_unverifiable(tmp_path):
    path = tmp_path / "bp.docx"
    _bp_doc(path)
    gate = run_to_final(
        path, explicit_domain="business_plan", apply_draft_name=False
    )
    assert gate.lrule_report is not None
    assert gate.lrule_report.summary.get("unverifiable", 0) == 0
    mech = [
        r for r in gate.lrule_report.rules
        if r["phase"] == "mechanized" and r["applicable"]
    ]
    assert mech
    assert all(r["status"] in {"PASS", "FAIL"} for r in mech)
    assert not gate.finalizer.submittable  # judgment/gap still REVIEW_REQUIRED


def test_submit_hwpx_stays_r9_only():
    """HWPX 제출 경로는 LRule/run_to_final 을 타지 않는다 (R9 수용검사 KEEP)."""
    src = (
        Path(__file__).resolve().parents[1] / "auto_write" / "services" / "hwpx_submit.py"
    ).read_text(encoding="utf-8")
    lowered = src.lower()
    assert "lrule" not in lowered
    assert "run_to_final" not in src
    assert "build_lrule_guards" not in src


def test_lookup_guard_by_short_code():
    report = enforce_lrules(
        domain=Domain.BUSINESS_PLAN,
        guards={"L009": {"passed": True, "evidence": "short-code"}},
    )
    l009 = next(r for r in report.rules if rule_code(r["id"]) == "L009")
    assert l009["status"] == "PASS"
    assert l009["evidence"] == "short-code"
