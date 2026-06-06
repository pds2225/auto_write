from __future__ import annotations

import json
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document

from auto_write.analysis.docx_template import analyze_template
from auto_write.config import Settings, ensure_directories
from auto_write.document_ingest import ensure_template_docx, extract_additional_text
from auto_write.services.evidence_service import EvidenceService
from auto_write.services.image_service import ImageService
from auto_write.services.openai_client import OpenAIService
from auto_write.services.project_service import ProjectService
from auto_write.services.qa_service import QAService
from auto_write.services.render_service import RenderService
from auto_write.storage import Storage


def build_settings(root: Path) -> Settings:
    app_root = root / "app"
    workspace_root = root / "workspace"
    return Settings(
        app_root=app_root,
        workspace_root=workspace_root,
        template_root=workspace_root / "templates",
        project_root=workspace_root / "projects",
        results_root=root / "results",
        static_root=app_root / "auto_write" / "static",
        template_view_root=app_root / "auto_write" / "templates",
        host="127.0.0.1",
        port=8765,
        openai_api_key="",
        openai_model="gpt-4.1-mini",
        openai_search_model="gpt-4.1-mini",
        openai_image_model="gpt-image-1",
        anthropic_api_key="",
        anthropic_model="claude-sonnet-4-20250514",
        anthropic_search_model="claude-sonnet-4-20250514",
    )


def write_sample_hwpx(path: Path) -> None:
    section_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section" xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p id="1">
    <hp:run charPrIDRef="0"><hp:t>1. 사업 개요</hp:t></hp:run>
    <hp:linesegarray/>
  </hp:p>
  <hp:p id="2">
    <hp:run charPrIDRef="0">
      <hp:tbl id="1" rowCnt="2" colCnt="2">
        <hp:tr>
          <hp:tc>
            <hp:subList><hp:p><hp:run><hp:t>구분</hp:t></hp:run></hp:p></hp:subList>
            <hp:cellAddr rowAddr="0" colAddr="0"/>
            <hp:cellSpan rowSpan="1" colSpan="1"/>
          </hp:tc>
          <hp:tc>
            <hp:subList><hp:p><hp:run><hp:t>내용</hp:t></hp:run></hp:p></hp:subList>
            <hp:cellAddr rowAddr="0" colAddr="1"/>
            <hp:cellSpan rowSpan="1" colSpan="1"/>
          </hp:tc>
        </hp:tr>
        <hp:tr>
          <hp:tc>
            <hp:subList><hp:p><hp:run><hp:t>핵심 전략</hp:t></hp:run></hp:p></hp:subList>
            <hp:cellAddr rowAddr="1" colAddr="0"/>
            <hp:cellSpan rowSpan="1" colSpan="1"/>
          </hp:tc>
          <hp:tc>
            <hp:subList><hp:p><hp:run/></hp:p></hp:subList>
            <hp:cellAddr rowAddr="1" colAddr="1"/>
            <hp:cellSpan rowSpan="1" colSpan="1"/>
          </hp:tc>
        </hp:tr>
      </hp:tbl>
    </hp:run>
    <hp:linesegarray/>
  </hp:p>
  <hp:p id="3">
    <hp:run charPrIDRef="0"><hp:pic id="img1"/></hp:run>
    <hp:linesegarray/>
  </hp:p>
</hs:sec>
"""
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("mimetype", "application/hwpx")
        archive.writestr("Contents/section0.xml", section_xml)
        archive.writestr("Preview/PrvText.txt", "1. 사업 개요\n핵심 전략")


def _make_hwp_style() -> dict[str, object]:
    return {
        "heading_level": 0,
        "alignment": "Left",
        "list_style": None,
        "indent_level": 0,
        "line_spacing": None,
        "space_before": None,
        "space_after": None,
    }


def _make_hwp_text(text: str) -> dict[str, object]:
    return {
        "Text": {
            "text": text,
            "style": {
                "bold": False,
                "italic": False,
                "underline": False,
                "strikethrough": False,
                "superscript": False,
                "subscript": False,
                "font_name": "맑은 고딕",
                "font_size": 12.0,
                "color": "#000000",
                "background_color": None,
            },
        }
    }


def _make_hwp_paragraph(text: str, heading_level: int = 0) -> dict[str, object]:
    style = _make_hwp_style()
    style["heading_level"] = heading_level
    return {
        "Paragraph": {
            "style": style,
            "content": [_make_hwp_text(text)],
        }
    }


def _make_hwp_cell(text: str, rowspan: int = 1, colspan: int = 1) -> dict[str, object]:
    return {
        "content": [_make_hwp_paragraph(text)],
        "rowspan": rowspan,
        "colspan": colspan,
        "alignment": "Left",
        "vertical_alignment": "Top",
        "background_color": None,
    }


def _make_hwp_json() -> str:
    data = {
        "metadata": {
            "title": "",
            "author": "kised",
            "subject": "",
            "keywords": [],
            "created": "2026-01-01T00:00:00Z",
            "modified": "2026-01-01T00:00:00Z",
            "creator_app": None,
            "format_version": "5.1.1.0",
            "is_distribution": False,
        },
        "sections": [
            {
                "index": 0,
                "content": [
                    _make_hwp_paragraph("사업 개요", heading_level=2),
                    {
                        "Table": {
                            "rows": [
                                {
                                    "cells": [
                                        _make_hwp_cell("신청 분야", colspan=2),
                                    ],
                                    "is_header": True,
                                },
                                {
                                    "cells": [
                                        _make_hwp_cell("전략 분야"),
                                        _make_hwp_cell("주관기관"),
                                    ],
                                    "is_header": True,
                                },
                                {
                                    "cells": [
                                        _make_hwp_cell("AI"),
                                        _make_hwp_cell("인공지능산업융합사업단"),
                                    ],
                                    "is_header": False,
                                },
                            ],
                            "column_widths": [1800, 3200],
                            "has_header": True,
                        }
                    },
                ],
                "header": None,
                "footer": None,
            }
        ],
        "styles": {
            "char_styles": [],
            "para_styles": [],
            "named_styles": [],
            "named_style_defs": [],
        },
        "resources": {},
    }
    return json.dumps(data, ensure_ascii=False)


def _make_hwp_guidance_json() -> str:
    data = json.loads(_make_hwp_json())
    data["sections"][0]["content"] = [
        _make_hwp_paragraph("신청현황", heading_level=2),
        {
            "Table": {
                "rows": [
                    {
                        "cells": [
                            _make_hwp_cell("※ 표는 삭제 후 제출", colspan=2),
                        ],
                        "is_header": True,
                    },
                    {
                        "cells": [
                            _make_hwp_cell("신청 분야", colspan=2),
                        ],
                        "is_header": True,
                    },
                    {
                        "cells": [
                            _make_hwp_cell("전략 분야"),
                            _make_hwp_cell("주관기관"),
                        ],
                        "is_header": True,
                    },
                    {
                        "cells": [
                            _make_hwp_cell("AI"),
                            _make_hwp_cell("인공지능산업융합사업단"),
                        ],
                        "is_header": False,
                    },
                ],
                "column_widths": [1800, 3200],
                "has_header": True,
            }
        },
    ]
    return json.dumps(data, ensure_ascii=False)


class _FakeUnhwpResult:
    def __init__(self, json_text: str):
        self.json = json_text

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc, tb):
        return False


class _FakeUnhwpModule:
    def __init__(self, json_text: str, markdown_text: str):
        self._json_text = json_text
        self._markdown_text = markdown_text

    def parse(self, _path):
        return _FakeUnhwpResult(self._json_text)

    def to_markdown(self, _path):
        return self._markdown_text


class _FakeOle:
    def __init__(self, data: bytes):
        self._data = data

    def exists(self, candidate):
        return candidate == ["PrvText"]

    def openstream(self, candidate):
        class _Stream:
            def __init__(self, payload: bytes):
                self._payload = payload

            def read(self):
                return self._payload

        return _Stream(self._data)

    def close(self):
        return None


class _FakeOleModule:
    def __init__(self, data: bytes):
        self._data = data

    @staticmethod
    def isOleFile(_path: str) -> bool:
        return True

    def OleFileIO(self, _path: str):
        return _FakeOle(self._data)


class DocumentIngestTests(unittest.TestCase):
    def test_ensure_template_docx_separates_hwp_guidance_rows(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            hwp_path = Path(tmp_dir) / "sample.hwp"
            hwp_path.write_bytes(b"fake")
            fake_unhwp = _FakeUnhwpModule(
                _make_hwp_guidance_json(),
                "신청현황\n\n※ 표는 삭제 후 제출\n\n신청 분야\n전략 분야\n주관기관\nAI\n인공지능산업융합사업단",
            )

            with patch("auto_write.document_ingest._load_unhwp", return_value=fake_unhwp):
                converted_path, notes = ensure_template_docx(hwp_path)

            self.assertTrue(converted_path.exists())
            self.assertTrue(any("표 구조" in note for note in notes))

            loaded_doc = Document(str(converted_path))
            paragraph_texts = [paragraph.text.strip() for paragraph in loaded_doc.paragraphs if paragraph.text.strip()]
            self.assertIn("신청현황", paragraph_texts)
            self.assertIn("※ 표는 삭제 후 제출", paragraph_texts)
            self.assertEqual(loaded_doc.tables[0].cell(0, 0).text.strip(), "신청 분야")
            self.assertEqual(loaded_doc.tables[0].cell(1, 0).text.strip(), "전략 분야")
            self.assertFalse(any("Left" in cell.text or "Justify" in cell.text for row in loaded_doc.tables[0].rows for cell in row.cells))

            profile = analyze_template(converted_path)
            self.assertTrue(any(table.label == "신청 분야" for table in profile.tables))

    def test_ensure_template_docx_converts_hwp_with_tables(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            hwp_path = Path(tmp_dir) / "sample.hwp"
            hwp_path.write_bytes(b"fake")
            fake_unhwp = _FakeUnhwpModule(
                _make_hwp_json(),
                "사업 개요\n\n| 신청 분야 | |\n| --- | --- |\n| AI | 인공지능산업융합사업단 |",
            )

            with patch("auto_write.document_ingest._load_unhwp", return_value=fake_unhwp):
                converted_path, notes = ensure_template_docx(hwp_path)

            self.assertTrue(converted_path.exists())
            self.assertTrue(converted_path.name.endswith("_converted.docx"))
            self.assertTrue(any("unhwp" in note for note in notes))

            loaded_doc = Document(str(converted_path))
            self.assertTrue(any("사업 개요" in paragraph.text for paragraph in loaded_doc.paragraphs))
            self.assertGreaterEqual(len(loaded_doc.tables), 1)
            self.assertIn("신청 분야", loaded_doc.tables[0].cell(0, 0).text)
            self.assertIn("인공지능산업융합사업단", loaded_doc.tables[0].cell(2, 1).text)

    def test_ensure_template_docx_converts_hwpx(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            hwpx_path = Path(tmp_dir) / "sample.hwpx"
            write_sample_hwpx(hwpx_path)

            converted_path, notes = ensure_template_docx(hwpx_path)

            self.assertTrue(converted_path.exists())
            self.assertTrue(converted_path.name.endswith("_converted.docx"))
            self.assertTrue(any("HWPX" in note for note in notes))

            profile = analyze_template(converted_path)
            self.assertTrue(any("사업 개요" in section.label for section in profile.sections))
            self.assertTrue(any(table.cells for table in profile.tables))

    def test_extract_additional_text_reads_hwpx_preview(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            hwpx_path = Path(tmp_dir) / "preview.hwpx"
            write_sample_hwpx(hwpx_path)

            extracted = extract_additional_text(hwpx_path)

            self.assertIn("1. 사업 개요", extracted)
            self.assertIn("핵심 전략", extracted)

    def test_extract_additional_text_prefers_hwp_markdown(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            hwp_path = Path(tmp_dir) / "sample.hwp"
            hwp_path.write_bytes(b"fake")
            fake_unhwp = _FakeUnhwpModule(
                _make_hwp_json(),
                "사업 개요\n\n| 신청 분야 | |\n| --- | --- |\n| AI | 인공지능산업융합사업단 |",
            )

            with patch("auto_write.document_ingest._load_unhwp", return_value=fake_unhwp):
                extracted = extract_additional_text(hwp_path)

            self.assertIn("사업 개요", extracted)
            self.assertIn("| 신청 분야 | |", extracted)
            self.assertIn("인공지능산업융합사업단", extracted)

    def test_extract_additional_text_reads_hwp_preview_stream(self):
        preview_text = "2026년 모집 공고\n핵심 지원 내용".encode("utf-16-le")
        fake_ole = _FakeOleModule(preview_text)
        with tempfile.TemporaryDirectory() as tmp_dir:
            hwp_path = Path(tmp_dir) / "sample.hwp"
            hwp_path.write_bytes(b"fake")
            with patch("auto_write.document_ingest._load_olefile", return_value=fake_ole):
                extracted = extract_additional_text(hwp_path)

        self.assertIn("2026년 모집 공고", extracted)
        self.assertIn("핵심 지원 내용", extracted)

    def test_project_service_accepts_hwpx_template_upload(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            root = Path(tmp_dir)
            settings = build_settings(root)
            ensure_directories(settings)
            storage = Storage(settings)
            openai_service = OpenAIService(settings)
            service = ProjectService(
                storage=storage,
                openai_service=openai_service,
                evidence_service=EvidenceService(openai_service),
                image_service=ImageService(openai_service),
                render_service=RenderService(),
                qa_service=QAService(),
            )
            hwpx_path = root / "service_sample.hwpx"
            write_sample_hwpx(hwpx_path)

            profile = service.analyze_uploaded_template(hwpx_path.name, hwpx_path.read_bytes())

            self.assertEqual(profile.template_name, hwpx_path.name)
            self.assertTrue(profile.source_docx.endswith("_converted.docx"))
            self.assertTrue(Path(profile.source_docx).exists())
            self.assertTrue(any("HWPX" in note for note in profile.analysis_notes))
            loaded_doc = Document(profile.source_docx)
            self.assertTrue(any("1. 사업 개요" in paragraph.text for paragraph in loaded_doc.paragraphs))


if __name__ == "__main__":
    unittest.main()
