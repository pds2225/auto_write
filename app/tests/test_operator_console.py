from __future__ import annotations

import json
import shutil
import subprocess
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from auto_write.operator_main import app
from auto_write.services.docx_edit_service import DocxEditService
from auto_write.services.git_sync_service import GitSyncError, GitSyncService
from auto_write.services.lrule_console_service import LRuleConsoleService
from auto_write.services.system_map_service import SystemMapService
from auto_write.services.workflow_monitor import WorkflowMonitor


REPO_ROOT = Path(__file__).resolve().parents[2]
RULE_REL = Path("app/tests/lessons_coverage.json")


def test_operator_console_smoke():
    client = TestClient(app)
    response = client.get("/console")
    assert response.status_code == 200
    assert "문서 작업" in response.text
    assert "L 규칙" in response.text
    assert "GitHub" in response.text


def test_operator_lrules_exposes_entire_canonical_registry():
    service = LRuleConsoleService(REPO_ROOT)
    data = service.load()
    rules = service.list_rules()
    assert len(rules) == len(data["lessons"])
    assert len(rules) == data["counts"]["total"]
    assert rules[0]["_code"].startswith("L")


def test_architecture_map_checks_real_files():
    lrules = LRuleConsoleService(REPO_ROOT)
    service = SystemMapService(REPO_ROOT, lrules)
    overview = service.overview()
    keys = {node["key"] for node in overview["nodes"]}
    assert {"web", "router", "lrule", "project", "converter"}.issubset(keys)
    router = next(node for node in overview["nodes"] if node["key"] == "router")
    assert router["path"] == "app/auto_write/domains/domain_router.py"
    assert router["status"] == "NORMAL"


def test_docx_browser_edit_never_overwrites_source_and_locks(tmp_path):
    source = tmp_path / "source.docx"
    doc = Document()
    doc.add_paragraph("원문")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "표 원문"
    doc.save(source)

    output = tmp_path / "edited.docx"
    locks = tmp_path / "user_locks.json"
    service = DocxEditService()
    report = service.apply_edits(
        source,
        {"p:0": "사용자 수정", "t:0:r:0:c:0": "표 수정"},
        output,
        locks,
    )

    assert report["applied_count"] == 2
    assert source.read_bytes() != output.read_bytes()
    assert Document(source).paragraphs[0].text == "원문"
    assert Document(output).paragraphs[0].text == "사용자 수정"
    assert json.loads(locks.read_text(encoding="utf-8"))["locks"]["p:0"] == "사용자 수정"

    second = service.apply_edits(
        Path(report["output"]),
        {"p:0": "사용자 2차 수정"},
        output,
        locks,
    )
    assert Path(second["output"]).name == "output_user_edited_v2.docx"
    assert Document(output).paragraphs[0].text == "사용자 수정"
    assert Document(second["output"]).paragraphs[0].text == "사용자 2차 수정"


def test_docx_editor_deduplicates_merged_table_cells(tmp_path):
    source = tmp_path / "merged.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.text = "병합 셀"
    doc.save(source)

    blocks = DocxEditService().load_blocks(source)
    table_blocks = [block for block in blocks if block["kind"] == "표"]
    assert len(table_blocks) == 1
    assert table_blocks[0]["text"] == "병합 셀"


def test_workflow_monitor_records_actual_wrapped_step(tmp_path):
    monitor = WorkflowMonitor(tmp_path / "runs.json")
    run_id = monitor.start_run("write", "문서 작성")
    with monitor.step(run_id, "route", "업무 분류", "DomainRouter"):
        value = 1 + 1
        assert value == 2
    monitor.finish_run(run_id, {"result": "ok"})

    run = monitor.get_run(run_id)
    assert run is not None
    assert run["status"] == "SUCCESS"
    assert run["steps"][0]["status"] == "SUCCESS"
    assert run["steps"][0]["service"] == "DomainRouter"


def test_git_sync_source_forbids_force_and_base_direct_push():
    source = (REPO_ROOT / "app/auto_write/services/git_sync_service.py").read_text(encoding="utf-8")
    assert '"--force"' not in source
    assert '"--force-with-lease"' not in source
    assert "AUTO_WRITE_GIT_WRITE_MODE" not in source
    assert 'write_mode = "direct"' not in source


def _git(cwd: Path, *args: str) -> str:
    proc = subprocess.run(
        ["git", *args],
        cwd=cwd,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    if proc.returncode != 0:
        raise AssertionError((proc.stderr or proc.stdout).strip())
    return (proc.stdout or "").strip()


def _rule_payload(summary: str) -> str:
    return json.dumps(
        {
            "counts": {"mechanized": 1, "gap": 0, "judgment": 0, "total": 1},
            "lessons": [
                {
                    "id": "L001 | 테스트 규칙",
                    "summary": summary,
                    "mechanizable": "yes",
                    "category": "mechanized",
                    "guard_ref": "app/example.py; app/tests/test_example.py",
                    "gap_desc": "",
                    "impact": "high",
                    "domain": "all",
                }
            ],
        },
        ensure_ascii=False,
        indent=2,
    ) + "\n"


def _write_rule(repo: Path, summary: str) -> None:
    path = repo / RULE_REL
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(_rule_payload(summary), encoding="utf-8")


def _setup_git_remote(tmp_path: Path) -> tuple[Path, Path, Path]:
    remote = tmp_path / "remote.git"
    seed = tmp_path / "seed"
    web = tmp_path / "web"
    peer = tmp_path / "peer"

    subprocess.run(["git", "init", "--bare", str(remote)], check=True, capture_output=True)
    subprocess.run(["git", "clone", str(remote), str(seed)], check=True, capture_output=True)
    _git(seed, "config", "user.email", "test@example.com")
    _git(seed, "config", "user.name", "Auto Write Test")
    _write_rule(seed, "v1")
    _git(seed, "add", str(RULE_REL).replace("\\", "/"))
    _git(seed, "commit", "-m", "seed rule")
    _git(seed, "branch", "-M", "master")
    _git(seed, "push", "-u", "origin", "master")

    subprocess.run(["git", "clone", "-b", "master", str(remote), str(web)], check=True, capture_output=True)
    subprocess.run(["git", "clone", "-b", "master", str(remote), str(peer)], check=True, capture_output=True)
    for repo in (web, peer):
        _git(repo, "config", "user.email", "test@example.com")
        _git(repo, "config", "user.name", "Auto Write Test")
    return remote, web, peer


def _disable_gh(monkeypatch) -> None:
    original_which = shutil.which
    monkeypatch.setattr(shutil, "which", lambda name: None if name == "gh" else original_which(name))


@pytest.mark.skipif(shutil.which("git") is None, reason="git executable required")
def test_git_remote_to_web_and_stale_base_guard(tmp_path):
    _, web, peer = _setup_git_remote(tmp_path)
    service = GitSyncService(web)

    assert service.snapshot(fetch=True).status == "SYNCED"

    (peer / "remote_change.txt").write_text("remote v2\n", encoding="utf-8")
    _git(peer, "add", "remote_change.txt")
    _git(peer, "commit", "-m", "remote update")
    _git(peer, "push", "origin", "master")
    synced = service.sync_from_remote()
    assert synced.status == "SYNCED"
    assert (web / "remote_change.txt").read_text(encoding="utf-8") == "remote v2\n"

    stale_base = service.base_remote_sha()
    (peer / "remote_change_2.txt").write_text("remote v3\n", encoding="utf-8")
    _git(peer, "add", "remote_change_2.txt")
    _git(peer, "commit", "-m", "remote race")
    _git(peer, "push", "origin", "master")
    with pytest.raises(GitSyncError, match="REMOTE_CHANGED"):
        service.assert_write_base(stale_base)


@pytest.mark.skipif(shutil.which("git") is None, reason="git executable required")
def test_git_web_rule_edit_uses_feature_branch_and_returns_after_merge(tmp_path, monkeypatch):
    _, web, _ = _setup_git_remote(tmp_path)
    _disable_gh(monkeypatch)
    service = GitSyncService(web)

    expected = service.base_remote_sha()
    service.assert_write_base(expected)
    _write_rule(web, "v2-web")
    result = service.commit_and_push(
        [str(RULE_REL).replace("\\", "/")],
        message="web rule update",
        expected_base_remote_sha=expected,
    )

    assert result["branch"].startswith("web/lrule-")
    assert service.current_branch() == result["branch"]
    assert service.remote_sha(result["branch"]) == service.local_sha()
    assert service.base_remote_sha() != service.local_sha()

    # Simulate a normal merge of the review branch into master.
    _git(web, "push", "origin", f"{result['branch']}:master")
    synced = service.sync_from_remote()
    assert service.current_branch() == "master"
    assert synced.status == "SYNCED"
    assert json.loads((web / RULE_REL).read_text(encoding="utf-8"))["lessons"][0]["summary"] == "v2-web"


@pytest.mark.skipif(shutil.which("git") is None, reason="git executable required")
def test_git_feature_branch_reuses_branch_and_syncs_new_master(tmp_path, monkeypatch):
    _, web, peer = _setup_git_remote(tmp_path)
    _disable_gh(monkeypatch)
    service = GitSyncService(web)

    base = service.base_remote_sha()
    service.assert_write_base(base)
    _write_rule(web, "v2")
    first = service.commit_and_push(
        [str(RULE_REL).replace("\\", "/")],
        message="rule v2",
        expected_base_remote_sha=base,
    )

    (peer / "master_change.txt").write_text("advanced\n", encoding="utf-8")
    _git(peer, "add", "master_change.txt")
    _git(peer, "commit", "-m", "advance master")
    _git(peer, "push", "origin", "master")

    synced = service.sync_from_remote()
    assert service.current_branch() == first["branch"]
    assert synced.status == "SYNCED"
    assert (web / "master_change.txt").read_text(encoding="utf-8") == "advanced\n"
    assert service.remote_sha(first["branch"]) == service.local_sha()
    assert subprocess.run(
        ["git", "merge-base", "--is-ancestor", "origin/master", "HEAD"],
        cwd=web,
        capture_output=True,
    ).returncode == 0

    base2 = service.base_remote_sha()
    service.assert_write_base(base2)
    _write_rule(web, "v3")
    second = service.commit_and_push(
        [str(RULE_REL).replace("\\", "/")],
        message="rule v3",
        expected_base_remote_sha=base2,
    )
    assert second["branch"] == first["branch"]


@pytest.mark.skipif(shutil.which("git") is None, reason="git executable required")
def test_git_squash_equivalent_rule_content_returns_to_master(tmp_path, monkeypatch):
    _, web, peer = _setup_git_remote(tmp_path)
    _disable_gh(monkeypatch)
    service = GitSyncService(web)

    base = service.base_remote_sha()
    service.assert_write_base(base)
    _write_rule(web, "squash-value")
    feature = service.commit_and_push(
        [str(RULE_REL).replace("\\", "/")],
        message="feature rule",
        expected_base_remote_sha=base,
    )
    feature_sha = service.local_sha()

    # Create an independent master commit with the same canonical registry value,
    # equivalent to a squash merge where the feature tip is not an ancestor.
    _write_rule(peer, "squash-value")
    _git(peer, "add", str(RULE_REL).replace("\\", "/"))
    _git(peer, "commit", "-m", "squash merged rule")
    _git(peer, "push", "origin", "master")
    _git(web, "fetch", "origin")
    assert subprocess.run(
        ["git", "merge-base", "--is-ancestor", feature_sha, "origin/master"],
        cwd=web,
        capture_output=True,
    ).returncode != 0

    synced = service.sync_from_remote()
    assert service.current_branch() == "master"
    assert synced.status == "SYNCED"
    assert feature["branch"].startswith("web/lrule-")


@pytest.mark.skipif(shutil.which("git") is None, reason="git executable required")
def test_git_failed_feature_push_rolls_back_worktree(tmp_path, monkeypatch):
    remote, web, _ = _setup_git_remote(tmp_path)
    _disable_gh(monkeypatch)
    service = GitSyncService(web)

    hook = remote / "hooks" / "pre-receive"
    hook.write_text("#!/bin/sh\necho rejected >&2\nexit 1\n", encoding="utf-8")
    hook.chmod(0o755)

    expected = service.base_remote_sha()
    service.assert_write_base(expected)
    _write_rule(web, "must-rollback")

    with pytest.raises(GitSyncError):
        service.commit_and_push(
            [str(RULE_REL).replace("\\", "/")],
            message="rejected update",
            expected_base_remote_sha=expected,
        )

    assert service.current_branch() == "master"
    assert _git(web, "status", "--porcelain") == ""
    assert json.loads((web / RULE_REL).read_text(encoding="utf-8"))["lessons"][0]["summary"] == "v1"


@pytest.mark.skipif(shutil.which("git") is None, reason="git executable required")
def test_git_precommit_remote_change_restores_managed_registry(tmp_path, monkeypatch):
    _, web, peer = _setup_git_remote(tmp_path)
    _disable_gh(monkeypatch)
    service = GitSyncService(web)

    base = service.base_remote_sha()
    service.assert_write_base(base)
    _write_rule(web, "v2")
    first = service.commit_and_push(
        [str(RULE_REL).replace("\\", "/")],
        message="first branch edit",
        expected_base_remote_sha=base,
    )

    stale_base = service.base_remote_sha()
    service.assert_write_base(stale_base)
    _write_rule(web, "pending-v3")

    (peer / "master_change.txt").write_text("remote advanced\n", encoding="utf-8")
    _git(peer, "add", "master_change.txt")
    _git(peer, "commit", "-m", "advance master")
    _git(peer, "push", "origin", "master")

    with pytest.raises(GitSyncError, match="REMOTE_CHANGED"):
        service.commit_and_push(
            [str(RULE_REL).replace("\\", "/")],
            message="must abort",
            expected_base_remote_sha=stale_base,
        )

    assert service.current_branch() == first["branch"]
    assert _git(web, "status", "--porcelain") == ""
    assert json.loads((web / RULE_REL).read_text(encoding="utf-8"))["lessons"][0]["summary"] == "v2"


@pytest.mark.skipif(shutil.which("git") is None, reason="git executable required")
def test_git_unrelated_dirty_abort_preserves_pending_rule_edit_for_retry(tmp_path, monkeypatch):
    _, web, _ = _setup_git_remote(tmp_path)
    _disable_gh(monkeypatch)
    service = GitSyncService(web)

    base = service.base_remote_sha()
    service.assert_write_base(base)
    _write_rule(web, "pending")
    (web / "scratch.txt").write_text("unrelated\n", encoding="utf-8")

    with pytest.raises(GitSyncError, match="규칙 파일 외"):
        service.commit_and_push(
            [str(RULE_REL).replace("\\", "/")],
            message="blocked",
            expected_base_remote_sha=base,
        )
    assert json.loads((web / RULE_REL).read_text(encoding="utf-8"))["lessons"][0]["summary"] == "pending"

    (web / "scratch.txt").unlink()
    service.assert_write_base(base)
    result = service.commit_and_push(
        [str(RULE_REL).replace("\\", "/")],
        message="retry succeeds",
        expected_base_remote_sha=base,
    )
    assert result["branch"].startswith("web/lrule-")


@pytest.mark.skipif(shutil.which("git") is None, reason="git executable required")
def test_rule_history_tracks_field_edits_not_only_id_occurrence_changes(tmp_path):
    _, web, _ = _setup_git_remote(tmp_path)
    service = GitSyncService(web)

    _write_rule(web, "v2")
    _git(web, "add", str(RULE_REL).replace("\\", "/"))
    _git(web, "commit", "-m", "change L001 summary to v2")
    (web / "other.txt").write_text("not a rule\n", encoding="utf-8")
    _git(web, "add", "other.txt")
    _git(web, "commit", "-m", "unrelated commit")
    _write_rule(web, "v3")
    _git(web, "add", str(RULE_REL).replace("\\", "/"))
    _git(web, "commit", "-m", "change L001 summary to v3")

    history = service.rule_history("L001", str(RULE_REL).replace("\\", "/"), limit=10)
    subjects = [item["subject"] for item in history]
    assert subjects[0] == "change L001 summary to v3"
    assert "change L001 summary to v2" in subjects
    assert "seed rule" in subjects
    assert "unrelated commit" not in subjects
