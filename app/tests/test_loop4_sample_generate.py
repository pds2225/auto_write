from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from auto_write.config import Settings, ensure_directories
from auto_write.services.evidence_service import EvidenceService
from auto_write.services.image_service import ImageService
from auto_write.services.openai_client import OpenAIService
from auto_write.services.project_service import ProjectService
from auto_write.services.qa_service import QAService
from auto_write.services.render_service import RenderService
from auto_write.storage import Storage


def build_settings(root: Path) -> Settings:
    app_root = root / "app"
    workspace_root = root / "workspace"
    return Settings(
        app_root=app_root,
        workspace_root=workspace_root,
        template_root=workspace_root / "templates",
        project_root=workspace_root / "projects",
        results_root=root / "results",
        static_root=app_root / "auto_write" / "static",
        template_view_root=app_root / "auto_write" / "templates",
        host="127.0.0.1",
        port=8765,
        openai_api_key="",
        openai_model="gpt-4.1-mini",
        openai_search_model="gpt-4.1-mini",
        openai_image_model="gpt-image-1",
        anthropic_api_key="",
        anthropic_model="claude-sonnet-4-20250514",
        anthropic_search_model="claude-sonnet-4-20250514",
    )

SAMPLE_BRIEF = "LOOP4 샘플 사업 개요: AI 안전 제어 스타트업"
SAMPLE_NOTES = "LOOP4 해결 파트\n\nLOOP4 성장 파트\n\nLOOP4 팀 파트"


def _build_psst_sample_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("□ 일반현황")
    doc.add_paragraph("1. 문제 인식 (Problem)_창업 아이템의 필요성")
    doc.add_paragraph("2. 실현 가능성 (Solution)_창업 아이템의 개발 계획")
    doc.add_paragraph("3. 성장전략 (Scale-up)_사업화 추진 전략")
    doc.add_paragraph("4. 팀 구성 (Team)_대표자 및 팀원 구성 계획")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "항목"
    table.cell(0, 1).text = "내용"
    table.cell(1, 0).text = "기업명"
    table.cell(1, 1).text = "○○기업"
    doc.save(path)


class Loop4SampleGenerateTests(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp_dir = tempfile.TemporaryDirectory()
        root = Path(self.tmp_dir.name)
        self.settings = build_settings(root)
        ensure_directories(self.settings)
        self.storage = Storage(self.settings)
        openai_service = OpenAIService(self.settings)
        self.service = ProjectService(
            storage=self.storage,
            openai_service=openai_service,
            evidence_service=EvidenceService(openai_service),
            image_service=ImageService(openai_service),
            render_service=RenderService(),
            qa_service=QAService(),
        )

    def tearDown(self) -> None:
        self.tmp_dir.cleanup()

    def test_loop4_sample_generate_writes_results_and_psst_text(self) -> None:
        sample_docx = Path(self.tmp_dir.name) / "psst_sample.docx"
        _build_psst_sample_docx(sample_docx)
        profile = self.service.analyze_uploaded_template("psst_sample.docx", sample_docx.read_bytes())
        project_id = self.service.create_project(profile.template_id, "LOOP4 검증")

        self.service.save_project_form(
            project_id=project_id,
            answers={
                "user_brief": SAMPLE_BRIEF,
                "user_notes": SAMPLE_NOTES,
            },
            project_title="LOOP4 테스트 과제",
            organization_name="LOOP4 테스트 기업",
            evidence_topics="",
            reference_files=[],
            improve_partial=True,
            psst_only=True,
            disable_images=True,
        )

        artifacts = self.service.generate(project_id)
        self.assertTrue(Path(artifacts.output_docx).exists())
        self.assertTrue(Path(artifacts.results_dir).exists())
        self.assertTrue(Path(artifacts.hwp_paste).exists())
        self.assertTrue(Path(artifacts.results_docx).exists())

        hwp_text = Path(artifacts.hwp_paste).read_text(encoding="utf-8")
        self.assertIn(SAMPLE_BRIEF, hwp_text)
        self.assertIn("LOOP4 해결 파트", hwp_text)
        self.assertIn("LOOP4 성장 파트", hwp_text)
        self.assertIn("LOOP4 팀 파트", hwp_text)

        doc_text = "\n".join(paragraph.text for paragraph in Document(artifacts.output_docx).paragraphs)
        self.assertIn(SAMPLE_BRIEF, doc_text)
        self.assertIn("LOOP4 해결 파트", doc_text)

        transfer = json.loads(Path(artifacts.transfer_report).read_text(encoding="utf-8"))
        image_cov = transfer.get("image_slot_coverage", {})
        self.assertEqual(image_cov.get("all_filled", 0), 0)

        fill_map = json.loads(Path(artifacts.fill_map).read_text(encoding="utf-8"))
        self.assertTrue(fill_map.get("user_brief_mapped_to"))
        self.assertTrue(any(item.get("psst") == "problem" for item in fill_map.get("sections", [])))

        project_input = self.storage.load_project_input(project_id)
        psst = self.service._find_psst_field_ids(profile)
        self.assertEqual(
            str(project_input.answers.get(psst["problem"], "")),
            SAMPLE_BRIEF,
        )
        self.assertEqual(
            str(project_input.answers.get(psst["solution"], "")),
            "LOOP4 해결 파트",
        )


if __name__ == "__main__":
    unittest.main()
