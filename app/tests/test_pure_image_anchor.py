"""test_pure_image_anchor.py — image_apply._find_anchor 2-패스 앵커 매칭 회귀.

프롬프트/이미지 블록을 '정위치'에 삽입하기 위한 앵커 탐색 계약을 in-memory Document 로
고정한다(python-docx 로컬 객체만, COM/네트워크 없음). 과거 회귀:
- 단일 패스 시절 짧은 부분문자열 본문이 표 셀 정방향 앵커보다 먼저 잡혀 오삽입 → forward-first.
- 역포함(2차)에 임계(길이>4 & 앵커의 25% 이상)를 둬 단순 키워드 오매칭 차단.

Returns 계약: (paragraph, top_table). 표 안 앵커면 top_table 이 그 최상위 표(삽입은 표 뒤).
"""

from __future__ import annotations

from docx import Document

from auto_write.services.image_apply import _find_anchor


def test_forward_match_in_body_returns_paragraph_no_table():
    doc = Document()
    doc.add_paragraph("① 사업 개요 및 추진 배경")
    para, table = _find_anchor(doc, "① 사업 개요")
    assert para is not None
    assert "사업 개요" in para.text
    assert table is None          # 본문 직계 → 표 컨텍스트 없음


def test_forward_match_in_table_cell_returns_top_table():
    doc = Document()
    doc.add_paragraph("무관 본문")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "신청기업 개요"
    para, table = _find_anchor(doc, "신청기업 개요")
    assert para is not None and para.text == "신청기업 개요"
    # 표 안 앵커 → 그 셀이 아니라 최상위 표를 함께 반환(삽입은 표 뒤).
    # python-docx 는 doc.tables 마다 새 프록시를 만들어 is 동일성이 깨지므로
    # 안정적인 하위 oxml 요소(_tbl)로 동일 표인지 비교한다.
    assert table is not None and table._tbl is tbl._tbl


def test_forward_beats_reverse_prefers_exact_cell_anchor():
    # 짧은 본문 부분문자열이 있어도, 정방향(표 셀 완전포함)이 우선되어야 한다.
    doc = Document()
    doc.add_paragraph("추진")   # 앵커의 짧은 부분문자열(역포함 후보)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.rows[0].cells[0].text = "추진 전략 및 성장 계획"
    para, table = _find_anchor(doc, "추진 전략 및 성장 계획")
    # 1차 정방향에서 표 셀이 잡혀야 함(본문 "추진"이 먼저 잡히면 오삽입).
    assert para.text == "추진 전략 및 성장 계획"
    assert table is not None and table._tbl is tbl._tbl


def test_reverse_match_when_paragraph_is_substring_of_anchor():
    doc = Document()
    doc.add_paragraph("사업화 추진 전략")   # 앵커의 부분(길이 9 > 4, 25% 이상)
    para, table = _find_anchor(doc, "제3장 사업화 추진 전략 세부 계획")
    assert para is not None
    assert para.text == "사업화 추진 전략"
    assert table is None


def test_reverse_rejects_too_short_paragraph():
    doc = Document()
    doc.add_paragraph("팀")   # 길이 1 (>4 아님) → 역포함 거부
    para, table = _find_anchor(doc, "팀 구성 및 보유 역량")
    assert (para, table) == (None, None)


def test_reverse_rejects_below_ratio_threshold():
    doc = Document()
    doc.add_paragraph("핵심요약칸")   # 길이 5(>4)지만 앵커 대비 25% 미만
    long_anchor = "핵심요약칸" + "0123456789" * 3   # 총 35자 → 25%=8.75 > 5
    para, table = _find_anchor(doc, long_anchor)
    assert (para, table) == (None, None)


def test_not_found_returns_none_pair():
    doc = Document()
    doc.add_paragraph("전혀 다른 본문 내용")
    assert _find_anchor(doc, "존재하지 않는 앵커 문구") == (None, None)


def test_blank_anchor_returns_none_pair():
    doc = Document()
    doc.add_paragraph("본문")
    assert _find_anchor(doc, "   ") == (None, None)
    assert _find_anchor(doc, "") == (None, None)
