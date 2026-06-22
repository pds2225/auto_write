"""사업계획서 작성 규칙 1단계(저위험) 테스트.

Phase 1: 신규 규칙은 모두 옵트인이며, 켜지 않으면 현행 동작과 동등하다.
- quality_rules 프리셋 플래그 / resolve_ruleset 매핑·override·검증
- run_all(rules=None) 및 run_all(rules=프리셋) == 레거시(서식 count 동일, Phase 1 no-op)
- ⑦ check_unverified_claims: 기본 off, 활성 시 '확정 단정'만 warn, 게이트(submittable) 불변

Phase 2(②③⑤ 배선)에서 run_all 의 rules 동등성 단언은 의도적으로 갱신될 예정이다.
"""

from __future__ import annotations

import pytest
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from auto_write.services.quality_rules import (
    BizplanRulesConfig, PRESETS, resolve_ruleset,
)
from auto_write.services.doc_quality_ops import (
    run_all, normalize_colored_text_to_black, unify_paragraph_formatting,
)
from auto_write.services.usage_acceptance import (
    AcceptanceConfig, SEV_WARN, run_acceptance, check_unverified_claims,
)


# --- 프리셋 / resolve_ruleset ----------------------------------------------

def test_no_preset_enables_score_empty_required_by_default():
    # ⑤ 채점 연동은 보수적으로 어떤 프리셋도 기본 on 하지 않는다(사용자 결정 2026-06-22).
    enabled = [n for n in PRESETS if PRESETS[n].score_empty_required]
    assert enabled == []


def test_presets_flag_values():
    biz = PRESETS["bizplan"]
    assert biz.color_to_black is True
    assert biz.body_font_pt == 10.0            # ③ 본문 10pt
    assert biz.score_empty_required is False   # ⑤ 보수적 off(사용자 결정)

    rep = PRESETS["report"]
    assert rep.body_font_pt is None
    assert rep.score_empty_required is False

    mini = PRESETS["minimal"]
    assert mini.color_to_black is False
    assert mini.blank_undecided is False

    off = PRESETS["off"]                        # 전부 off — 레거시와 논리 동등
    assert off.color_to_black is False
    assert off.body_font_pt is None
    assert off.suggest_notebooklm is False
    assert off.blank_undecided is False
    assert off.flag_unverified_claims is False
    assert off.score_empty_required is False


def test_default_config_is_conservative():
    # 명시 지정 없이 만든 기본 설정은 보수적(고위험·동작변경 규칙 off).
    d = BizplanRulesConfig()
    assert d.score_empty_required is False
    assert d.enforce_confirm_marker is False
    assert d.flag_unverified_claims is False
    assert d.body_font_pt == 10.0


def test_resolve_ruleset_mapping_and_override():
    assert resolve_ruleset("business_plan") is PRESETS["bizplan"]
    assert resolve_ruleset("pitch_deck") is PRESETS["bizplan"]
    assert resolve_ruleset("rnd_plan") is PRESETS["report"]    # 기타/미상 → report
    assert resolve_ruleset(None) is PRESETS["report"]
    # override(CLI --ruleset)가 doc_type 보다 우선
    assert resolve_ruleset("business_plan", override="minimal") is PRESETS["minimal"]
    assert resolve_ruleset(override="off") is PRESETS["off"]


def test_resolve_ruleset_unknown_override_raises():
    with pytest.raises(ValueError):
        resolve_ruleset(override="does-not-exist")


# --- 하위호환: run_all(rules=...) == 레거시 (Phase 1 no-op) -----------------

def _build_doc() -> Document:
    """후처리가 실제로 일을 하도록 빈 단락 + 파란 런을 포함한 문서."""
    doc = Document()
    doc.add_paragraph("정상 본문 문장입니다.")
    doc.add_paragraph("")  # 빈 단락 → empty_paragraphs_removed 유발
    p = doc.add_paragraph("실적 ")
    run = p.add_run("핵심 성과 수치")  # 안내문구 패턴 비매칭(삭제 안 됨)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # 유색 → colored_runs_normalized
    return doc


def test_run_all_rules_none_equals_legacy():
    legacy = run_all(_build_doc())
    rules_none = run_all(_build_doc(), rules=None)
    assert legacy.as_dict() == rules_none.as_dict()


def test_run_all_bizplan_applies_color_and_target_pt():
    # bizplan: color_to_black=True(색→검정 on) + body_font_pt=10(본문 10pt 통일).
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("본문 문장 텍스트")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    rep = run_all(doc, rules=PRESETS["bizplan"])
    assert rep.colored_runs_normalized >= 1            # ② 색→검정 적용
    assert _run_sz(p.runs[0]) == "20"                  # ③ 본문 10pt(half-point 20)


def test_run_all_off_preset_disables_color_normalization():
    # off: color_to_black=False → 색→검정을 끈다(사용자 결정 — 프리셋이 제어).
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("본문 문장 텍스트").font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    rep = run_all(doc, rules=PRESETS["off"])
    assert rep.colored_runs_normalized == 0            # 색 변환 안 함
    color = p.runs[0]._element.find(qn("w:rPr")).find(qn("w:color"))
    assert color.get(qn("w:val")).lower() == "0000ff"  # 파랑 보존


def test_run_all_minimal_preset_disables_color_normalization():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("본문 문장 텍스트").font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    rep = run_all(doc, rules=PRESETS["minimal"])
    assert rep.colored_runs_normalized == 0            # minimal 도 색 변환 끔


def test_run_all_rules_none_keeps_default_color_normalization():
    # 대조군: rules 미지정(현행)은 기본대로 색→검정 on(하위호환 — 항상 켜짐).
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("본문 문장 텍스트").font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    rep = run_all(doc, rules=None)
    assert rep.colored_runs_normalized >= 1


def test_run_all_actually_does_work():
    """위 동등성이 '0==0' 공허가 아님을 보증 — 파란 런이 실제로 검정화된다."""
    rep = run_all(_build_doc())
    assert rep.colored_runs_normalized >= 1


def test_run_all_legacy_args_unaffected_by_rules_none():
    """오케스트레이터가 쓰는 인자 조합에서도 rules=None 이 결과를 바꾸지 않음."""
    kw = dict(remove_guides=True, emphasize=True, underline=False, normalize_fonts=False)
    legacy = run_all(_build_doc(), **kw)
    rules_none = run_all(_build_doc(), rules=None, **kw)
    assert legacy.as_dict() == rules_none.as_dict()


# --- ⑦ check_unverified_claims: warn·기본 off·게이트 불변 -------------------

def _doc_with(*texts: str) -> Document:
    doc = Document()
    for t in texts:
        doc.add_paragraph(t)
    return doc


def test_unverified_claims_off_by_default():
    doc = _doc_with("삼성전자와 협업(확정)")
    r_none = check_unverified_claims(doc, None)              # config 미지정 → 비활성
    assert r_none.severity == SEV_WARN and r_none.defects == 0
    r_default = check_unverified_claims(doc, AcceptanceConfig())  # 기본 flag off → 비활성
    assert r_default.defects == 0


def test_unverified_claims_active_flags_confirmed_only():
    cfg = AcceptanceConfig(flag_unverified_claims=True)
    doc = _doc_with("삼성전자와 협업(확정)", "B사와 협력 검토 중")
    r = check_unverified_claims(doc, cfg)
    assert r.severity == SEV_WARN
    assert r.defects == 1   # '협업(확정)'만 — '검토 중'은 헷지 표현으로 제외


def test_unverified_claims_excludes_evidence_marked():
    cfg = AcceptanceConfig(flag_unverified_claims=True)
    doc = _doc_with("삼성전자와 협업(확정) [산출근거: 계약서 2026.01]")
    assert check_unverified_claims(doc, cfg).defects == 0   # 근거 표기 → 제외


def test_unverified_claims_warn_does_not_change_gate(tmp_path):
    """warn 추가는 AcceptanceReport.submittable 에 영향 없음(SEV_FAIL 만 평가)."""
    doc = _doc_with("삼성전자와 협업(확정)")
    p = tmp_path / "claim.docx"
    doc.save(str(p))
    rep_off = run_acceptance(p, AcceptanceConfig(flag_unverified_claims=False))
    rep_on = run_acceptance(p, AcceptanceConfig(flag_unverified_claims=True))
    assert rep_off.submittable == rep_on.submittable   # 게이트(submittable) 불변
    assert rep_on.warn_defects >= 1                     # 경고는 잡힌다
    assert "unverified_claims" in {r.check_id for r in rep_on.results}  # 검사 등록 확인


# --- ② 색→검정 정규화는 형광펜·음영(highlight/shd)을 보존해야 함 (강조 증발 방지) ----

def _doc_with_colored_highlighted_run():
    doc = Document()
    p = doc.add_paragraph("핵심 ")
    run = p.add_run("성과 강조 문장")
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)     # 파란 글자색
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW    # 노란 형광펜(w:highlight)
    rpr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "FFFF00")
    rpr.append(shd)                                     # 음영(w:shd)
    return doc, run


def test_color_normalize_preserves_highlight_and_shd():
    """색 정규화는 글자색만 검정화하고 형광펜·음영은 보존해야 한다.

    검출기 check_residual_colored_runs 는 색만 보고 highlight 는 안 보므로, 색변환이
    highlight/shd 까지 지우면 게이트가 못 잡는 '강조 증발' 회귀가 된다(계획 §2-D-②).
    """
    doc, run = _doc_with_colored_highlighted_run()
    n = normalize_colored_text_to_black(doc)
    assert n >= 1                                       # 파란 글자색 → 검정
    rpr = run._element.find(qn("w:rPr"))
    color = rpr.find(qn("w:color"))
    assert color is not None and color.get(qn("w:val")) == "000000"
    assert rpr.find(qn("w:highlight")) is not None, "형광펜(highlight)이 보존돼야 함"
    assert rpr.find(qn("w:shd")) is not None, "음영(shd)이 보존돼야 함"
    assert normalize_colored_text_to_black(doc) == 0    # 멱등 — 2회차 0건


# --- ③ unify_paragraph_formatting target_pt 본문 통일 (지배값과 상호배타·명시 분기) ----

def _run_sz(run):
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        return None
    sz = rpr.find(qn("w:sz"))
    return sz.get(qn("w:val")) if sz is not None else None


def test_target_pt_sets_body_size_overriding_dominant():
    # 본문: 9pt·9pt·12pt → 지배값 9pt(half 18)이지만 target_pt=10 이면 모두 20.
    doc = Document()
    p = doc.add_paragraph()
    for pt in (9, 9, 12):
        p.add_run("문장 ").font.size = Pt(pt)
    n = unify_paragraph_formatting(doc, target_pt=10.0)
    assert n == 1
    assert {_run_sz(r) for r in p.runs} == {"20"}      # target 10pt=half 20, 지배값 18 미적용
    assert unify_paragraph_formatting(doc, target_pt=10.0) == 0   # 멱등


def test_target_pt_imposes_size_on_theme_inherited_body():
    # 명시 크기 없는 본문(테마 상속)에도 target_pt 모드는 크기를 '부여'한다.
    doc = Document()
    p = doc.add_paragraph("크기 미지정 본문")
    assert _run_sz(p.runs[0]) is None
    n = unify_paragraph_formatting(doc, target_pt=10.0)
    assert n == 1
    assert _run_sz(p.runs[0]) == "20"


def test_dominant_mode_preserves_theme_inherited_body():
    # 대조군: target_pt=None(지배값 모드)은 명시 크기 없는 본문을 보존(부여 안 함).
    doc = Document()
    p = doc.add_paragraph("크기 미지정 본문")
    assert unify_paragraph_formatting(doc, target_pt=None) == 0
    assert _run_sz(p.runs[0]) is None                  # 날조 금지 — 크기 미부여


def test_target_pt_skips_headings_and_table_cells():
    doc = Document()
    h = doc.add_heading("제목", level=1)
    h.runs[0].font.size = Pt(15)                       # 제목 크기 명시
    t = doc.add_table(rows=1, cols=1)
    cell_run = t.cell(0, 0).paragraphs[0].add_run("셀 내용")
    cell_run.font.size = Pt(9)
    body = doc.add_paragraph()
    body.add_run("본문").font.size = Pt(9)
    unify_paragraph_formatting(doc, target_pt=10.0)
    assert _run_sz(h.runs[0]) == "30"                  # 제목 미변경(15pt)
    assert _run_sz(cell_run) == "18"                   # 표 셀 미변경(target·지배 둘 다 미적용)
    assert _run_sz(body.runs[0]) == "20"               # 본문만 target 10pt


# --- 오케스트레이터 --ruleset 노출 (opt-in, 기본 None=현행) -------------------

def test_orchestrator_ruleset_controls_color_normalization(tmp_path):
    from auto_write.services.document_quality_orchestrator import DocumentQualityOrchestrator

    def _blue_docx(name):
        d = Document()
        d.add_paragraph().add_run("본문 문장 텍스트").font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
        path = tmp_path / name
        d.save(str(path))
        return path

    orch = DocumentQualityOrchestrator(tmp_path, openai_service=None)
    # ruleset 미지정(현행) → 색→검정 on(하위호환)
    res_default = orch.run(_blue_docx("in_default.docx"), tmp_path / "out_default.docx",
                           ruleset=None, write_report=False)
    assert res_default.ops.colored_runs_normalized >= 1
    # off 프리셋 → 색→검정 끔(프리셋이 run_all 까지 제어)
    res_off = orch.run(_blue_docx("in_off.docx"), tmp_path / "out_off.docx",
                       ruleset="off", write_report=False)
    assert res_off.ops.colored_runs_normalized == 0
