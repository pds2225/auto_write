"""검출↔교정 정합 강제 메타 테스트 (티어1-C).

배경: '지웠는데 검출됨' / '검출만 되고 교정 안 됨' 류의 정합 불일치가 R11(유색
텍스트 검정화) 한 기능에서만 5회 재발했다. 근본 원인은 **검출기와 교정기가 순회
범위·값 로직을 공유하지 않으면**(각자 병렬로 트리를 걷는다) 한쪽만 새 범위를
얻거나 잃어 어긋나는 것이다.

이 파일은 개별 기능 회귀(test_usage_acceptance.py 의 body/header/hyperlink/table
개별 케이스)와 겹치지 않는 **메타(정합 강제) 각도**만 담는다:

  (1) 안내 정합 — usage_acceptance 의 fail 검사와 acceptance_remediation._REMEDIES
      사이의 대응을, test_acceptance_remediation.py 가 이미 강제하는 '정방향(모든
      check_id 에 안내)' 외의 **역방향(고아 안내 금지)·config 의존 fail 집합(strict
      승격)·모든 AUTO fail 명령의 완전 치환** 각도로 보강한다.

  (2) 검출↔교정 거울 정합 — 유색 텍스트를 본문·표셀·머리글·바닥글·하이퍼링크 각
      범위에 **정확히 1개씩** 심어, check_residual_colored_runs(검출)와
      normalize_colored_text_to_black(교정)이 **같은 순회 범위를 정확한 개수로**
      보는지 범위별로 고정한다(기존 개별 테스트는 `>= 1`, 여기선 `== 1` 로 잠근다).
      한 범위라도 한쪽만 보면 정직하게 실패한다.

프로덕션 코드는 건드리지 않는다(읽기 전용 메타 테스트).
"""

from __future__ import annotations

import pytest
from docx import Document
from docx.oxml import parse_xml
from docx.shared import RGBColor

from auto_write.services import acceptance_remediation as rem
from auto_write.services import usage_acceptance as ua
from auto_write.services.usage_acceptance import (
    AcceptanceConfig,
    SEV_FAIL,
    check_residual_colored_runs,
)
from auto_write.services.doc_quality_ops import normalize_colored_text_to_black


# ===========================================================================
# (1) 안내(remediation) 정합 — 정방향은 test_acceptance_remediation 이 이미 강제.
#     여기서는 역방향·config 의존·명령 완전치환의 보강 각도만.
# ===========================================================================

def _all_check_ids() -> set[str]:
    """실제 검사(_ALL_CHECKS)가 내는 check_id 전체 — 단일 출처."""
    doc = Document()
    doc.add_paragraph("본문 한 줄")
    cfg = AcceptanceConfig()
    return {check(doc, cfg).check_id for check in ua._ALL_CHECKS}


def _fail_check_ids(tmp_path, config: AcceptanceConfig) -> set[str]:
    """주어진 config 에서 severity 가 fail 인 check_id 집합.

    run_acceptance 를 거친다 — strict_acceptance 승격은 개별 check 함수가 아니라
    run_acceptance 가 사후에 적용하므로, 실제 게이트가 보는 fail 집합을 얻으려면
    반드시 이 경로로 판정해야 한다(개별 호출은 승격 전 native severity 라 어긋남).
    """
    doc = Document()
    doc.add_paragraph("본문 한 줄")
    p = tmp_path / "판정용.docx"
    doc.save(str(p))
    report = ua.run_acceptance(p, config)
    return {r.check_id for r in report.results if r.severity == SEV_FAIL}


def test_no_orphan_remedies() -> None:
    """역방향 정합: _REMEDIES 의 모든 키가 실재하는 check_id 여야 한다.

    검사가 이름 변경·삭제됐는데 안내만 남거나(고아), 안내 키에 오타가 나면
    '지웠는데 안내에 남음' 류의 드리프트가 생긴다 — 여기서 강제 실패시킨다.
    (정방향 '모든 검사에 안내'는 test_acceptance_remediation 이 담당.)
    """
    check_ids = _all_check_ids()
    orphan = set(rem._REMEDIES) - check_ids
    assert not orphan, f"실재하지 않는 check_id 에 걸린 고아 안내: {sorted(orphan)}"


def test_strict_mode_promoted_fails_all_have_remedies(tmp_path) -> None:
    """config 의존 fail 집합 정합: strict_acceptance 로 warn→fail 승격된 검사도
    전부 안내(AUTO/HUMAN/MANUAL)를 가져야 한다.

    기본 config 의 fail 집합만 안내가 있고 승격된 fail 이 누락되면, 공고가 해당
    항목을 필수로 요구(strict)할 때 '제출불가인데 다음 행동 안내 없음'이 된다.
    """
    strict = AcceptanceConfig(strict_acceptance=True)
    fail_ids = _fail_check_ids(tmp_path, strict)

    # 승격 대상(_PROMOTABLE_WARN_IDS)이 실제로 fail 로 편입됐는지(배선 확인)
    for pid in ua._PROMOTABLE_WARN_IDS:
        assert pid in fail_ids, f"{pid} 가 strict 모드에서 fail 로 승격되지 않음(배선 끊김)"

    # 그 확장된 fail 집합 전부에 전용 안내가 있어야 한다
    for cid in fail_ids:
        r = rem.remedy_for(cid)
        assert r is not rem._DEFAULT, f"fail 검사 {cid} 에 전용 안내 없음 — _REMEDIES 추가 필요"
        assert r.kind in (rem.KIND_AUTO, rem.KIND_HUMAN, rem.KIND_MANUAL)


def test_every_auto_fail_command_is_fully_substituted(tmp_path) -> None:
    """모든 AUTO 계열 fail 검사의 안내 명령이 실제 문서 경로로 완전 치환돼야 한다.

    (기존 test_build_substitutes_doc_path_in_command 는 self_inserted_blocks 한
    건만 확인.) 여기서는 AUTO 로 분류된 모든 fail 검사에 대해 build_remediation 이
    {doc} 토큰을 남기지 않고 문서명을 담는지 메타로 강제한다 — 새 AUTO 안내가
    토큰 치환을 빠뜨리면(사용자에게 '{doc}' 그대로 노출) 실패한다.
    """
    doc_name = "제출본 v2.docx"
    fail_ids = _fail_check_ids(tmp_path, AcceptanceConfig())
    auto_fail = [cid for cid in fail_ids if rem.remedy_for(cid).kind == rem.KIND_AUTO]
    assert auto_fail, "AUTO 로 자동해결되는 fail 검사가 최소 1개는 있어야(가드)"

    for cid in auto_fail:
        result = ua.CheckResult(cid, "라벨", SEV_FAIL, 1)
        items = rem.build_remediation([result], doc_name)
        cmd = items[0]["command"]
        assert cmd, f"{cid} 는 AUTO 인데 명령이 비어 있음"
        assert rem.DOC_TOKEN not in cmd, f"{cid} 명령에 미치환 토큰 {rem.DOC_TOKEN} 잔존: {cmd}"
        assert doc_name in cmd, f"{cid} 명령에 실제 문서명이 안 들어감: {cmd}"


# ===========================================================================
# (2) 검출↔교정 거울 정합 — 유색 텍스트를 각 범위에 정확히 1개씩 심어
#     검출/교정이 '같은 범위를 같은 개수로' 보는지 범위별로 고정.
# ===========================================================================

_BLUE = RGBColor(0x00, 0x00, 0xFF)   # 파란 안내문구 색(검정도 보존색도 아님 → 결함)

# 하이퍼링크로 감싼 파란 run(직계가 아닌 .//w:r 경로) — 검출·교정 모두 이 범위를
# 봐야 한다(과거: 직계 run 만 봐서 하이퍼링크형 안내문구가 통째로 우회).
_HYPERLINK_BLUE_XML = (
    '<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:r><w:rPr><w:color w:val="0000FF"/></w:rPr><w:t>여기서 상세보기</w:t></w:r>'
    '</w:hyperlink>'
)


def _build_body() -> Document:
    """본문 직계 단락에 파란 run 1개."""
    d = Document()
    d.add_paragraph().add_run("파란 본문 안내문구")
    d.paragraphs[-1].runs[0].font.color.rgb = _BLUE
    return d


def _build_table_cell() -> Document:
    """표 셀 안 단락에 파란 run 1개(정부양식 핵심 위치)."""
    d = Document()
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "라벨"                      # 무색
    run = t.cell(0, 1).paragraphs[0].add_run("파란 표셀 안내문구")
    run.font.color.rgb = _BLUE
    return d


def _build_header() -> Document:
    """머리글 단락에 파란 run 1개(ACC-9 본문 밖 표시 영역)."""
    d = Document()
    hdr = d.sections[0].header
    hdr.is_linked_to_previous = False
    run = hdr.paragraphs[0].add_run("파란 머리글 안내문구")
    run.font.color.rgb = _BLUE
    return d


def _build_footer() -> Document:
    """바닥글 단락에 파란 run 1개 — 기존 테스트가 안 덮은 범위(머리글만 있었음)."""
    d = Document()
    ftr = d.sections[0].footer
    ftr.is_linked_to_previous = False
    run = ftr.paragraphs[0].add_run("파란 바닥글 안내문구")
    run.font.color.rgb = _BLUE
    return d


def _build_hyperlink() -> Document:
    """본문 단락 안 <w:hyperlink> 로 감싼 파란 run 1개(.//w:r 범위)."""
    d = Document()
    p = d.add_paragraph("안내: ")               # 앞 run 은 무색
    p._p.append(parse_xml(_HYPERLINK_BLUE_XML))
    return d


_RANGE_BUILDERS = [
    ("본문", _build_body),
    ("표셀", _build_table_cell),
    ("머리글", _build_header),
    ("바닥글", _build_footer),
    ("하이퍼링크", _build_hyperlink),
]


@pytest.mark.parametrize("where,builder", _RANGE_BUILDERS, ids=[w for w, _ in _RANGE_BUILDERS])
def test_detect_correct_mirror_parity_per_range(where, builder) -> None:
    """범위별 거울 정합: 각 범위에 유색 run 1개 → 검출 1 · 교정 1 · 재검출 0.

    - 검출만 되고 교정 안 됨  → 교정=0 또는 재검출>0 로 실패.
    - 교정만 되고 검출 안 됨  → 검출=0 으로 실패.
    - 한쪽 순회 범위가 이 범위를 빠뜨림 → 위 둘 중 하나로 정직하게 실패.
    """
    d = builder()

    # 검출(단락당 1건)·교정(런당 1건)이 정확히 1로 일치해야 한다(범위당 유색 run 1개).
    assert check_residual_colored_runs(d).defects == 1, f"[{where}] 검출이 1이 아님"
    assert normalize_colored_text_to_black(d) == 1, f"[{where}] 교정이 1이 아님"
    # 교정 후 재검출 0(거울 정합의 핵심) + 멱등(두 번째 교정 0)
    assert check_residual_colored_runs(d).defects == 0, f"[{where}] 교정 후에도 검출됨(정합 깨짐)"
    assert normalize_colored_text_to_black(d) == 0, f"[{where}] 멱등 위반(재교정 발생)"


def _build_all_ranges() -> Document:
    """본문·표셀·머리글·바닥글·하이퍼링크 5개 유색 run 을 한 문서에 동시에."""
    d = Document()
    # 본문
    d.add_paragraph().add_run("파란 본문 안내문구").font.color.rgb = _BLUE
    # 하이퍼링크(본문 내)
    p = d.add_paragraph("안내: ")
    p._p.append(parse_xml(_HYPERLINK_BLUE_XML))
    # 표 셀
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "라벨"
    t.cell(0, 1).paragraphs[0].add_run("파란 표셀 안내문구").font.color.rgb = _BLUE
    # 머리글·바닥글
    hdr = d.sections[0].header
    hdr.is_linked_to_previous = False
    hdr.paragraphs[0].add_run("파란 머리글 안내문구").font.color.rgb = _BLUE
    ftr = d.sections[0].footer
    ftr.is_linked_to_previous = False
    ftr.paragraphs[0].add_run("파란 바닥글 안내문구").font.color.rgb = _BLUE
    return d


def test_detect_correct_mirror_parity_all_ranges_at_once() -> None:
    """모든 범위를 한 문서에 동시에 심어 검출 개수 == 교정 개수 == 5 를 강제한다.

    한 범위라도 검출·교정 중 한쪽에서만 보이면 총합이 5 에서 벌어져 실패한다
    (per-range 테스트가 어느 범위인지 짚어주고, 이 테스트는 누락 자체를 막는다).
    """
    d = _build_all_ranges()

    detected = check_residual_colored_runs(d)
    assert detected.defects == 5, f"5개 범위 유색을 모두 검출해야 함(실제 {detected.defects})"

    corrected = normalize_colored_text_to_black(d)
    assert corrected == 5, f"5개 범위 유색을 모두 교정해야 함(실제 {corrected})"

    assert check_residual_colored_runs(d).defects == 0, "교정 후 전 범위 0"

    # 검출 표본이 4개 구조 범위(본문/표/머리글/바닥글)를 실제로 훑었는지 교차 확인
    # (하이퍼링크는 본문 단락 안이라 where='본문' 으로 집계됨).
    wheres = {s.split("]")[0].lstrip("[") for s in detected.samples}
    assert {"본문", "표", "머리글", "바닥글"} <= wheres, f"검출이 놓친 구조 범위 있음: {wheres}"
