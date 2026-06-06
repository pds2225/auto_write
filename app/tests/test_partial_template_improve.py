from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from auto_write.analysis.docx_template import analyze_template, cell_value_needs_improvement, sanitize_template_profile


class PartialTemplateImproveTests(unittest.TestCase):
    def test_cell_value_needs_improvement_detects_placeholders(self):
        self.assertTrue(cell_value_needs_improvement("00백만원"))
        self.assertTrue(cell_value_needs_improvement("○○기업"))
        self.assertFalse(cell_value_needs_improvement("미래큐러스"))

    def test_completed_partial_doc_keeps_major_sections_only(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "partial.docx"
            doc = Document()
            doc.add_paragraph("□ 일반현황")
            doc.add_paragraph("○ 휴머노이드 로봇 및 AI 산업의 급속한 성장")
            doc.add_paragraph("1. 문제 인식 (Problem)_창업 아이템의 필요성")
            for _ in range(45):
                doc.add_paragraph("본문 내용을 채운 문단입니다. " * 4)
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "총 사업비"
            table.cell(0, 1).text = "금액"
            table.cell(1, 0).text = "정부지원"
            table.cell(1, 1).text = "00백만원"
            doc.save(path)

            profile = analyze_template(path)

            self.assertTrue(any(section.label.startswith("□") for section in profile.sections))
            self.assertFalse(any(section.anchor_text.startswith("○") for section in profile.sections))
            self.assertTrue(
                any(
                    cell_value_needs_improvement(cell.label) or cell.required
                    for table in profile.tables
                    for cell in table.cells
                )
            )

    def test_sanitize_filters_bullet_sections_for_saved_partial_profile(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "partial.docx"
            doc = Document()
            doc.add_paragraph("□ 일반현황")
            doc.add_paragraph("○ 본문 불릿")
            doc.add_paragraph("1. 문제 인식")
            for _ in range(45):
                doc.add_paragraph("작성된 본문입니다. " * 20)
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "항목"
            table.cell(0, 1).text = "값"
            table.cell(1, 0).text = "예산"
            table.cell(1, 1).text = "00백만원"
            doc.save(path)

            profile = analyze_template(path)
            profile = sanitize_template_profile(profile)
            self.assertFalse(any(section.anchor_text.startswith("○") for section in profile.sections))
            self.assertTrue(
                any(
                    cell.cell_id.endswith("_r1_c1")
                    for table in profile.tables
                    for cell in table.cells
                )
            )
