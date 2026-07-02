"""test_gate_faildraft_invariant.py — fail-open 게이트 우회 방지 불변식 테스트.

배경 (과거 최다 재발 결함, 6건+)
--------------------------------
제출불가(fail 등급 결함) 문서가 ``_DRAFT`` 마킹을 잃고 깨끗한 '제출용' 이름으로
빠져나가는 **fail-open** 누수가 진입점마다 따로 발견·수정됐다. 그러나 회귀 테스트가
인스턴스별로 흩어져 있어 새 경로가 생길 때마다 같은 계열 버그가 재발했다.

이 파일은 진입점 전체에 걸쳐 하나의 불변식(property)을 고정한다:

  ┌────────────────────────────────────────────────────────────────────┐
  │ fail 등급 결함이 있는 문서는, 어떤 진입점을 거쳐도                    │
  │   (1) '제출용'(비-DRAFT) 이름의 산출물을 만들지 않고                  │
  │   (2) 판정을 '제출가능'으로 보고하지 않는다.                          │
  └────────────────────────────────────────────────────────────────────┘

대상 진입점 (실사용 4경로):
  - ``self_diagnose.main``                     — 읽기 전용 진단 CLI(산출물 없음 →
                                                  게이트는 exit 2 로 고정)
  - ``autopilot_pipeline.run_autopilot``       — 문서 품질 오토파일럿
  - ``submission_orchestrator.SubmissionPipeline`` — ``auto_write/submit.py`` 엔진
  - ``bizplan_autopilot.run_bizplan_autopilot``    — 사업계획서 완성 오케스트레이터

fail 결함은 3종으로 파라미터화한다 — '특정 결함 하나'가 아니라 결함 클래스
전반에서 게이트가 닫히는지 본다:
  · marker      : ``[확인필요]`` 등 미해결 작성 마커(unresolved_markers)
  · unchecked   : ``택 1`` 행의 미체크 선택란 □(unchecked_choices)
  · empty_label : 필수 라벨(명 칭) 옆 칸 공란(empty_label_fields)

프로덕션 코드는 건드리지 않는다 — 테스트만. 기존 방어가 있으면 통과하고, 구멍이
있으면 그 케이스가 '정직하게 실패'한다(거짓 통과 금지).
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path
from unittest import mock

import pytest
from docx import Document

# app/ 를 import 기준에 올린다(test_submission_pipeline.py 와 동일 관례).
APP_DIR = Path(__file__).resolve().parent.parent
if str(APP_DIR) not in sys.path:
    sys.path.insert(0, str(APP_DIR))

from auto_write.config import Settings
from auto_write.models import ImageSlotProfile, ProjectInput, TemplateProfile
from auto_write.services.evaluation_service import EvaluationService
from auto_write.services.evidence_service import EvidenceService
from auto_write.services.image_service import ImageService
from auto_write.services.openai_client import OpenAIService
from auto_write.services.render_service import RenderService
from auto_write.services.submission_orchestrator import SubmissionPipeline

# fail 결함 3종 — 각각 usage_acceptance 의 서로 다른 fail 검사를 발동시킨다.
_FAIL_KINDS = ("marker", "unchecked", "empty_label")


def _make_fail_docx(path: Path, kind: str) -> None:
    """지정한 종류의 fail 등급 결함 1개를 심은 DOCX 를 만든다.

    세 결함 모두 서식 청소(품질 오토파일럿) 단계를 통과해 살아남는 것으로 확인됨
    (probe 검증) — 그래서 게이트가 마지막까지 결함을 포착해야 한다.
    """
    doc = Document()
    doc.add_paragraph("개요: 게이트 fail-open 불변식 검증용 문서.")
    if kind == "marker":
        doc.add_paragraph("사업비 [확인필요] 원")
    elif kind == "unchecked":
        t = doc.add_table(rows=1, cols=2)
        t.cell(0, 0).text = "지원 유형(택 1)"
        t.cell(0, 1).text = "□ 제조  □ 지식서비스"     # 체크 표시 없음
    elif kind == "empty_label":
        t = doc.add_table(rows=1, cols=2)
        t.cell(0, 0).text = "명 칭"
        t.cell(0, 1).text = ""                         # 필수 라벨 옆 칸 공란
    else:  # pragma: no cover - 방어
        raise ValueError(f"알 수 없는 결함 종류: {kind}")
    doc.save(str(path))


def _is_draft_name(path: str | Path) -> bool:
    return Path(path).stem.endswith(("_DRAFT", "_DRAFT2"))


# ---------------------------------------------------------------------------
# 1) self_diagnose — 읽기 전용 진단. 산출물이 없으므로 게이트는 종료코드(2)로 고정.
# ---------------------------------------------------------------------------

@pytest.mark.parametrize("kind", _FAIL_KINDS)
def test_self_diagnose_fail_doc_returns_reject_code(kind: str, tmp_path: Path) -> None:
    """fail 결함 문서 → self_diagnose 는 '제출가능(0)'을 절대 반환하지 않고
    '제출불가(2)'로 고정한다. 또한 읽기 전용이라 원본을 수정하지 않는다."""
    import self_diagnose

    src = tmp_path / f"diag_{kind}.docx"
    _make_fail_docx(src, kind)
    before = src.read_bytes()

    rc = self_diagnose.main([str(src)])

    assert rc == 2, f"{kind}: 제출불가여야 하는데 exit {rc}"
    assert src.read_bytes() == before, "진단은 읽기 전용 — 원본이 수정됨"


# ---------------------------------------------------------------------------
# 2) autopilot_pipeline.run_autopilot — 최종 산출명에 반드시 _DRAFT.
# ---------------------------------------------------------------------------

@pytest.mark.parametrize("kind", _FAIL_KINDS)
def test_autopilot_fail_doc_forces_draft(kind: str, tmp_path: Path) -> None:
    """fail 결함 문서 → run_autopilot 은 요청한 '제출용' 이름으로 내보내지 않고
    출력명에 _DRAFT 를 강제한다(비-DRAFT 제출명 산출물 0개)."""
    from auto_write.services.autopilot_pipeline import run_autopilot

    src = tmp_path / f"ap_{kind}.docx"
    out = tmp_path / f"ap_{kind}_out.docx"      # 사용자가 요청한 '제출용' 이름
    _make_fail_docx(src, kind)

    # 자기삽입 블록(NotebookLM·PSST 가이드)을 넣지 않아 '주입 결함' 하나만 남긴다 —
    # 게이트가 그 순수 결함을 포착하는지 격리 검증(max_images=0·psst 끔).
    report = run_autopilot(
        str(src), str(out), max_images=0, psst_scaffold=False, write_report=False
    )

    assert report.acceptance_submittable is False, f"{kind}: 제출가능으로 오판"
    assert report.draft_marked is True, f"{kind}: _DRAFT 마킹 누락"
    assert _is_draft_name(report.output_docx), f"{kind}: 산출명 비-DRAFT ({report.output_docx})"
    assert Path(report.output_docx).exists()
    # 사용자가 요청한 깨끗한 '제출용' 이름으로는 파일이 존재하면 안 된다.
    assert not out.exists(), f"{kind}: 제출용 이름으로 파일 유출({out.name})"


# ---------------------------------------------------------------------------
# 3) bizplan_autopilot.run_bizplan_autopilot — 최종 복사본까지 _DRAFT 전파.
# ---------------------------------------------------------------------------

@pytest.mark.parametrize("kind", _FAIL_KINDS)
def test_bizplan_fail_doc_forces_draft(kind: str, tmp_path: Path) -> None:
    """fail 결함 문서 → run_bizplan_autopilot 은 최종 복사본 이름에도 _DRAFT 를
    강제한다(shutil.copyfile 로 깨끗한 이름으로 세탁되던 과거 누수 차단)."""
    from auto_write.services.bizplan_autopilot import run_bizplan_autopilot

    src = tmp_path / f"bp_{kind}.docx"
    out = tmp_path / f"bp_{kind}_out.docx"
    _make_fail_docx(src, kind)

    report = run_bizplan_autopilot(
        str(src), str(out), use_ai=False, write_report=False
    )

    assert report.acceptance_submittable is False, f"{kind}: 제출가능으로 오판"
    assert report.draft_marked is True, f"{kind}: _DRAFT 마킹 누락"
    assert _is_draft_name(report.output_docx), f"{kind}: 산출명 비-DRAFT ({report.output_docx})"
    assert Path(report.output_docx).exists()
    assert not out.exists(), f"{kind}: 제출용 이름으로 파일 유출({out.name})"


# ---------------------------------------------------------------------------
# 4) SubmissionPipeline (auto_write/submit.py 엔진) — 격리 results_root 에서
#    '제출초안_*' 비-DRAFT 산출물 0개 불변식.
# ---------------------------------------------------------------------------

def _settings(tmp: Path) -> Settings:
    return Settings(
        app_root=tmp, workspace_root=tmp, template_root=tmp, project_root=tmp,
        results_root=tmp / "results", static_root=tmp, template_view_root=tmp,
        host="127.0.0.1", port=8765,
        openai_api_key="", openai_model="m", openai_search_model="m", openai_image_model="gpt-image-1",
        anthropic_api_key="", anthropic_model="m", anthropic_search_model="m",
        gemini_api_key="",
    )


def _profile(tmp: Path) -> TemplateProfile:
    slot = ImageSlotProfile(
        slot_id="img1", label="추진 체계 인포그래픽", required=True,
        anchor_type="table_cell",
        anchor_ref={"table_index": 0, "row": 0, "cell": 0},
        source="template",
    )
    return TemplateProfile(
        template_id="t1", template_name="s.docx",
        source_docx=str(tmp / "s.docx"), image_slots=[slot],
    )


class _FakeStorage:
    def __init__(self, base: Path):
        self._base = Path(base)
        self._inputs: dict = {}

    def project_dir(self, pid):
        d = self._base / "projects" / pid
        d.mkdir(parents=True, exist_ok=True)
        return d

    def load_project_input(self, pid):
        return self._inputs[pid]

    def save_project_input(self, pid, pi):
        self._inputs[pid] = pi


class _FakeProjectService:
    """generate() 가 지정한 fail 결함을 심은 output.docx 를 만든다."""

    def __init__(self, storage, profile, oa, fail_kind: str):
        self.storage = storage
        self._profile = profile
        self.openai_service = oa
        self.image_service = ImageService(oa)
        self.render_service = RenderService()
        self.evidence_service = EvidenceService(oa)
        self._fail_kind = fail_kind

    def generate(self, pid):
        out = self.storage.project_dir(pid) / "output" / "output.docx"
        out.parent.mkdir(parents=True, exist_ok=True)
        _make_fail_docx(out, self._fail_kind)
        return None

    def load_profile_for_project(self, pid):
        return self._profile


@pytest.mark.parametrize("kind", _FAIL_KINDS)
def test_submission_pipeline_fail_doc_forces_draft(kind: str, tmp_path: Path) -> None:
    """fail 결함 문서 → SubmissionPipeline 최종본이 _DRAFT 이고, results_root 최상위에
    비-DRAFT '제출초안_*.docx' 산출물이 하나도 남지 않는다(중간본까지 마킹 전파)."""
    settings = _settings(tmp_path)
    storage = _FakeStorage(tmp_path)
    oa = OpenAIService(settings)
    prof = _profile(tmp_path)
    ps = _FakeProjectService(storage, prof, oa, kind)
    storage.save_project_input(
        "p1",
        ProjectInput(template_id="t1", organization_profile={"기업명": "테스트(주)"},
                     project_meta={}),
    )
    pipeline = SubmissionPipeline(ps, EvaluationService(oa), storage, settings)
    report = pipeline.run(
        "p1", announcement_text="", enable_images=False, enable_notebooklm=False,
    )

    assert report["acceptance"]["submittable"] is False, f"{kind}: 제출가능으로 오판"
    assert _is_draft_name(report["final_docx"]), f"{kind}: 최종본 비-DRAFT ({report['final_docx']})"
    assert Path(report["final_docx"]).exists()
    # results_root 최상위에 '제출초안_*' 비-DRAFT(제출용) 산출물이 남으면 안 된다.
    leftovers = [
        p.name for p in Path(settings.results_root).glob("제출초안_*.docx")
        if not _is_draft_name(p)
    ]
    assert leftovers == [], f"{kind}: 비-DRAFT 제출 산출물 잔존 {leftovers}"


# ---------------------------------------------------------------------------
# 5) rename 실패(파일 잠금) — fail-closed. 마킹이 물리적으로 실패해도 시스템은
#    '제출가능'으로 보고하지 않는다(오류를 기록하고, 자동 체인은 exit 3 으로 막는다).
# ---------------------------------------------------------------------------

def test_autopilot_rename_lock_is_fail_closed(tmp_path: Path, monkeypatch) -> None:
    """_DRAFT rename 이 파일 잠금으로 실패해도(force_draft_name 이 오류 반환),
    run_autopilot 은 침묵 통과하지 않는다 — draft_mark_error 를 기록하고 판정을
    '제출가능'으로 올리지 않는다(fail-closed).

    ※ rename 자체가 실패하면 깨끗한 이름 파일이 물리적으로 남는 것은 불가피하다
      (rename 이 안 됐으므로). 이 케이스의 fail-closed 계약은 '보고가 성공으로
      둔갑하지 않는 것' — draft_mark_error 기록 + acceptance_submittable=False +
      manual_todo 경고이며, 자동 체인 차단은 아래 CLI --strict(exit 3) 가 담당한다.
    """
    from auto_write.services import autopilot_pipeline

    def _locked(path, avoid=None):  # 실제 시그니처(path, *, avoid) 모사 — 잠금 실패
        return path, "PermissionError: locked"

    monkeypatch.setattr(autopilot_pipeline, "force_draft_name", _locked)

    src = tmp_path / "lock_in.docx"
    out = tmp_path / "lock_out.docx"
    _make_fail_docx(src, "marker")
    report = autopilot_pipeline.run_autopilot(
        str(src), str(out), max_images=0, psst_scaffold=False, write_report=False
    )

    assert report.acceptance_submittable is False       # 결함은 여전히 fail
    assert report.draft_marked is False                 # rename 이 실패했으므로 마킹 X
    assert report.draft_mark_error, "잠금 실패가 draft_mark_error 로 기록되어야 함"
    assert any("_DRAFT 마킹 실패" in t for t in report.manual_todo)


def test_autopilot_cli_strict_exit3_on_rename_lock(tmp_path: Path, monkeypatch) -> None:
    """자동 체인 차단: rename 잠금 실패 시 --strict 종료코드가 3(검사불능/판정불가)
    으로, 절대 0(성공)으로 새지 않는다 — fail-open 자동화 유출 차단."""
    import auto_write_autopilot as cli
    from auto_write.services import autopilot_pipeline

    def _locked(path, avoid=None):
        return path, "PermissionError: locked"

    monkeypatch.setattr(autopilot_pipeline, "force_draft_name", _locked)

    src = tmp_path / "lock_cli.docx"
    _make_fail_docx(src, "marker")
    rc = cli.main([
        str(src), "-o", str(tmp_path / "lock_cli_out.docx"),
        "--strict", "--no-psst", "--max-images", "0", "--no-report",
    ])
    assert rc == 3, f"잠금 실패인데 exit {rc} — fail-closed(3) 여야 함"


if __name__ == "__main__":  # pragma: no cover
    import unittest
    sys.argv = [sys.argv[0]]
    raise SystemExit(pytest.main([str(Path(__file__))]))
