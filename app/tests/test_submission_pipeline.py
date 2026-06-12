import sys
import tempfile
import unittest
from pathlib import Path
from unittest import mock

from docx import Document

APP_DIR = Path(__file__).resolve().parent.parent
if str(APP_DIR) not in sys.path:
    sys.path.insert(0, str(APP_DIR))

from auto_write.config import Settings
from auto_write.models import GeneratedImage, ImageSlotProfile, ProjectInput, TemplateProfile
from auto_write.services.evaluation_service import EvaluationService
from auto_write.services.evidence_service import EvidenceService
from auto_write.services.image_service import ImageService
from auto_write.services.openai_client import OpenAIService
from auto_write.services.plan_builder import build_fill_plan
from auto_write.services.render_service import RenderService
from auto_write.services.submission_orchestrator import SubmissionPipeline


def _settings(tmp: Path) -> Settings:
    return Settings(
        app_root=tmp, workspace_root=tmp, template_root=tmp, project_root=tmp,
        results_root=tmp / "results", static_root=tmp, template_view_root=tmp,
        host="127.0.0.1", port=8765,
        openai_api_key="", openai_model="m", openai_search_model="m", openai_image_model="gpt-image-1",
        anthropic_api_key="", anthropic_model="m", anthropic_search_model="m",
        gemini_api_key="",
    )


def _profile(tmp: Path) -> TemplateProfile:
    slot = ImageSlotProfile(
        slot_id="img1", label="추진 체계 인포그래픽", required=True,
        anchor_type="table_cell",
        anchor_ref={"table_index": 0, "row": 0, "cell": 0},
        source="template",
    )
    return TemplateProfile(
        template_id="t1", template_name="s.docx",
        source_docx=str(tmp / "s.docx"), image_slots=[slot],
    )


class _FakeStorage:
    def __init__(self, base: Path):
        self._base = Path(base)
        self._inputs = {}

    def project_dir(self, pid):
        d = self._base / "projects" / pid
        d.mkdir(parents=True, exist_ok=True)
        return d

    def load_project_input(self, pid):
        return self._inputs[pid]

    def save_project_input(self, pid, pi):
        self._inputs[pid] = pi


class _FakeProjectService:
    def __init__(self, storage, profile, oa):
        self.storage = storage
        self._profile = profile
        self.openai_service = oa
        self.image_service = ImageService(oa)
        self.render_service = RenderService()
        self.evidence_service = EvidenceService(oa)

    def generate(self, pid):
        out = self.storage.project_dir(pid) / "output" / "output.docx"
        out.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.add_paragraph("1. 문제 인식")
        doc.add_paragraph("시장의 문제와 해결 방안을 서술한 본문 내용입니다.")
        t = doc.add_table(rows=1, cols=1)
        t.rows[0].cells[0].text = "< 사진(이미지) 또는 설계도 제목 >"
        doc.save(str(out))
        return None

    def load_profile_for_project(self, pid):
        return self._profile


class SubmissionPipelineTest(unittest.TestCase):
    def test_plan_builder_org_and_external(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            pi = ProjectInput(
                template_id="t1",
                organization_profile={"기업명": "테스트(주)", "대표자": "홍길동", "빈값": ""},
                project_meta={"overview": {"명칭": "테스트 사업"}},
            )
            prof = _profile(tmp_path)
            ext_dir = tmp_path / "ext"
            ext_dir.mkdir()
            (ext_dir / "fill_plan.json").write_text(
                '{"row_rewrites":[{"table_index":0,"row":0,"cols":["A"]}]}', encoding="utf-8"
            )
            plan = build_fill_plan(prof, pi, external_plan_dir=ext_dir)
            self.assertEqual(plan["identity"]["기업명"], "테스트(주)")
            self.assertNotIn("빈값", plan["identity"])
            self.assertEqual(plan["overview"]["명칭"], "테스트 사업")
            self.assertEqual(len(plan["row_rewrites"]), 1)

    def test_insert_images_into_docx(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            docx_path = tmp_path / "d.docx"
            doc = Document()
            t = doc.add_table(rows=1, cols=1)
            t.rows[0].cells[0].text = "placeholder"
            doc.save(str(docx_path))
            from PIL import Image
            png = tmp_path / "i.png"
            Image.new("RGB", (200, 120), "#2563EB").save(str(png))
            slot = ImageSlotProfile(
                slot_id="img1", label="도식", required=True,
                anchor_type="table_cell",
                anchor_ref={"table_index": 0, "row": 0, "cell": 0},
                source="template",
            )
            prof = TemplateProfile(
                template_id="t1", template_name="s.docx",
                source_docx=str(docx_path), image_slots=[slot],
            )
            img = GeneratedImage(slot_id="img1", label="도식", path=str(png), source="generated")
            rep = RenderService().insert_images_into_docx(prof, [img], docx_path)
            self.assertEqual(rep["images_written"], 1)
            out = Document(str(docx_path))
            self.assertIn("w:drawing", out.tables[0].rows[0].cells[0]._tc.xml)

    def test_pipeline_end_to_end_no_ai(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            settings = _settings(tmp_path)
            storage = _FakeStorage(tmp_path)
            oa = OpenAIService(settings)
            prof = _profile(tmp_path)
            ps = _FakeProjectService(storage, prof, oa)
            pi = ProjectInput(
                template_id="t1",
                organization_profile={"기업명": "테스트(주)"},
                project_meta={},
            )
            storage.save_project_input("p1", pi)
            pipeline = SubmissionPipeline(ps, EvaluationService(oa), storage, settings)
            report = pipeline.run("p1", announcement_text="", enable_images=True)
            self.assertIn("generate", report["steps"])
            self.assertIn("finalize", report["steps"])
            self.assertIn("images", report["steps"])
            self.assertTrue(Path(report["submit_docx"]).exists())
            self.assertTrue(Path(report["final_docx"]).exists())
            self.assertTrue((storage.project_dir("p1") / "output" / "output.docx").exists())
            self.assertGreaterEqual(report["images"]["inserted"], 1)

    def test_pipeline_inserts_notebooklm_prompts(self):
        """버그② 회귀: submit 파이프라인이 NotebookLM 슬라이드 프롬프트를
        최종본에 삽입해야 한다(이전엔 image_service(PNG)만 쓰고 미연결이라 안 나옴)."""
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            settings = _settings(tmp_path)
            storage = _FakeStorage(tmp_path)
            oa = OpenAIService(settings)
            prof = _profile(tmp_path)

            class _PSWithKeyword(_FakeProjectService):
                def generate(self, pid):
                    out = self.storage.project_dir(pid) / "output" / "output.docx"
                    out.parent.mkdir(parents=True, exist_ok=True)
                    doc = Document()
                    doc.add_paragraph("나. 추진일정 로드맵 — 단계별 마일스톤.")  # 간트 트리거
                    doc.add_paragraph("맺음말 본문 단락.")
                    doc.save(str(out))
                    return None

            ps = _PSWithKeyword(storage, prof, oa)
            pi = ProjectInput(
                template_id="t1",
                organization_profile={"기업명": "테스트(주)"},
                project_meta={},
            )
            storage.save_project_input("p1", pi)
            pipeline = SubmissionPipeline(ps, EvaluationService(oa), storage, settings)
            report = pipeline.run(
                "p1", announcement_text="",
                enable_images=False, enable_notebooklm=True,
            )
            self.assertIn("notebooklm", report["steps"])
            self.assertGreaterEqual(report["notebooklm"]["prompts_inserted"], 1)
            text = "\n".join(p.text for p in Document(report["final_docx"]).paragraphs)
            self.assertIn("NotebookLM", text)

    def test_pipeline_no_notebooklm_flag_skips(self):
        """--no-notebooklm 에 해당하는 enable_notebooklm=False 면 단계가 빠져야 한다."""
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            settings = _settings(tmp_path)
            storage = _FakeStorage(tmp_path)
            oa = OpenAIService(settings)
            prof = _profile(tmp_path)
            ps = _FakeProjectService(storage, prof, oa)
            pi = ProjectInput(
                template_id="t1",
                organization_profile={"기업명": "테스트(주)"},
                project_meta={},
            )
            storage.save_project_input("p1", pi)
            pipeline = SubmissionPipeline(ps, EvaluationService(oa), storage, settings)
            report = pipeline.run(
                "p1", announcement_text="",
                enable_images=False, enable_notebooklm=False,
            )
            self.assertNotIn("notebooklm", report["steps"])

    def test_pipeline_acceptance_gate_drafts_notebooklm_output(self):
        """R7: NotebookLM 작업용 블록이 든 최종본은 '제출' 이름으로 내보내지 않고
        _DRAFT 를 강제해야 한다(수용검사 fail → 파일명 차단)."""
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            settings = _settings(tmp_path)
            storage = _FakeStorage(tmp_path)
            oa = OpenAIService(settings)
            prof = _profile(tmp_path)

            class _PSWithKeyword(_FakeProjectService):
                def generate(self, pid):
                    out = self.storage.project_dir(pid) / "output" / "output.docx"
                    out.parent.mkdir(parents=True, exist_ok=True)
                    doc = Document()
                    doc.add_paragraph("나. 추진일정 로드맵 — 단계별 마일스톤.")  # 간트 트리거
                    doc.save(str(out))
                    return None

            ps = _PSWithKeyword(storage, prof, oa)
            pi = ProjectInput(
                template_id="t1",
                organization_profile={"기업명": "테스트(주)"},
                project_meta={},
            )
            storage.save_project_input("p1", pi)
            pipeline = SubmissionPipeline(ps, EvaluationService(oa), storage, settings)
            report = pipeline.run(
                "p1", announcement_text="",
                enable_images=False, enable_notebooklm=True,
            )
            self.assertIn("acceptance", report["steps"])
            self.assertFalse(report["acceptance"]["submittable"])
            self.assertTrue(report["final_docx"].endswith("_DRAFT.docx"))
            self.assertTrue(Path(report["final_docx"]).exists())
            self.assertTrue(any("제출 금지" in n for n in report["needs_input"]))

    def test_pipeline_acceptance_gate_disabled(self):
        """acceptance_gate=False 면 게이트 단계 없이 기존 동작을 유지한다."""
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            settings = _settings(tmp_path)
            storage = _FakeStorage(tmp_path)
            oa = OpenAIService(settings)
            prof = _profile(tmp_path)
            ps = _FakeProjectService(storage, prof, oa)
            pi = ProjectInput(
                template_id="t1",
                organization_profile={"기업명": "테스트(주)"},
                project_meta={},
            )
            storage.save_project_input("p1", pi)
            pipeline = SubmissionPipeline(ps, EvaluationService(oa), storage, settings)
            report = pipeline.run(
                "p1", announcement_text="",
                enable_images=False, enable_notebooklm=False,
                acceptance_gate=False,
            )
            self.assertNotIn("acceptance", report["steps"])
            self.assertNotIn("acceptance", report)
            self.assertFalse(report["final_docx"].endswith("_DRAFT.docx"))

    def test_pipeline_gate_fail_closed_on_acceptance_error(self):
        """R9: 게이트 자신이 죽어도 통과로 취급하지 않는다(fail-closed) —
        acceptance_error 기록 + needs_input 경고 + _DRAFT 강제."""
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            settings = _settings(tmp_path)
            storage = _FakeStorage(tmp_path)
            oa = OpenAIService(settings)
            prof = _profile(tmp_path)
            ps = _FakeProjectService(storage, prof, oa)
            pi = ProjectInput(
                template_id="t1",
                organization_profile={"기업명": "테스트(주)"},
                project_meta={},
            )
            storage.save_project_input("p1", pi)
            pipeline = SubmissionPipeline(ps, EvaluationService(oa), storage, settings)
            with mock.patch(
                "auto_write.services.submission_orchestrator.run_acceptance",
                side_effect=RuntimeError("boom"),
            ):
                report = pipeline.run(
                    "p1", announcement_text="",
                    enable_images=False, enable_notebooklm=False,
                )
            self.assertIn("acceptance_error", report)
            self.assertNotIn("acceptance", report["steps"])
            self.assertTrue(report["final_docx"].endswith("_DRAFT.docx"))
            self.assertTrue(Path(report["final_docx"]).exists())
            self.assertTrue(any("제출 금지" in n for n in report["needs_input"]))

    def test_pipeline_gate_updates_report_paths_after_draft(self):
        """R9: fail 로 최종본이 _DRAFT 로 rename 되면, 같은 파일을 가리키던
        submit_docx/quality_docx 리포트 경로도 갱신된다(댕글링 금지)."""
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            settings = _settings(tmp_path)
            storage = _FakeStorage(tmp_path)
            oa = OpenAIService(settings)
            prof = _profile(tmp_path)
            ps = _FakeProjectService(storage, prof, oa)
            pi = ProjectInput(
                template_id="t1",
                organization_profile={"기업명": "테스트(주)"},
                project_meta={},
            )
            storage.save_project_input("p1", pi)
            pipeline = SubmissionPipeline(ps, EvaluationService(oa), storage, settings)
            fake = mock.Mock(submittable=False, fail_defects=2, results=[])
            with mock.patch(
                "auto_write.services.submission_orchestrator.run_acceptance",
                return_value=fake,
            ):
                report = pipeline.run(
                    "p1", announcement_text="",
                    enable_images=False, enable_notebooklm=False,
                )
            self.assertTrue(report["final_docx"].endswith("_DRAFT.docx"))
            self.assertTrue(report["acceptance"]["draft_marked"])
            for key in ("submit_docx", "quality_docx"):
                if key in report:
                    self.assertTrue(
                        Path(report[key]).exists(),
                        f"{key} 가 존재하지 않는 경로를 가리킴: {report[key]}",
                    )

    def test_pipeline_gate_drafts_intermediate_artifacts(self):
        """R9 잔여 #9: 게이트 fail 시 최종본만이 아니라 중간 산출물(제출초안_*·_품질)도
        _DRAFT 마킹되어 results 에 '제출' 이름(비 DRAFT) 파일이 남지 않는다."""
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            settings = _settings(tmp_path)
            storage = _FakeStorage(tmp_path)
            oa = OpenAIService(settings)
            prof = _profile(tmp_path)

            class _PSWithKeyword(_FakeProjectService):
                def generate(self, pid):
                    out = self.storage.project_dir(pid) / "output" / "output.docx"
                    out.parent.mkdir(parents=True, exist_ok=True)
                    doc = Document()
                    doc.add_paragraph("나. 추진일정 로드맵 — 단계별 마일스톤.")  # 간트 트리거
                    doc.save(str(out))
                    return None

            ps = _PSWithKeyword(storage, prof, oa)
            pi = ProjectInput(
                template_id="t1",
                organization_profile={"기업명": "테스트(주)"},
                project_meta={},
            )
            storage.save_project_input("p1", pi)
            pipeline = SubmissionPipeline(ps, EvaluationService(oa), storage, settings)
            report = pipeline.run(
                "p1", announcement_text="",
                enable_images=False, enable_notebooklm=True,   # NotebookLM 블록 → 게이트 fail
            )
            self.assertFalse(report["acceptance"]["submittable"])
            # 리포트 경로가 _DRAFT 로 바뀐 실파일을 가리킨다
            for key in ("submit_docx", "quality_docx"):
                if key in report:
                    p = Path(report[key])
                    self.assertTrue(
                        p.stem.endswith(("_DRAFT", "_DRAFT2")), f"{key} 미마킹: {p.name}"
                    )
                    self.assertTrue(p.exists(), f"{key} 경로 없음: {p}")
            # results 최상위에 '제출' 이름(비 DRAFT) 산출물 잔존 0
            leftovers = [
                p.name for p in Path(settings.results_root).glob("제출초안_*.docx")
                if not p.stem.endswith(("_DRAFT", "_DRAFT2"))
            ]
            self.assertEqual(leftovers, [])

    def test_pipeline_gate_pass_keeps_submit_names(self):
        """R9 잔여 #9 반대 사양: 게이트 PASS 면 어떤 산출물도 _DRAFT 로 바뀌지 않는다
        (마킹은 'fail/판정불가 시만' — 조건이 망가지면 이 테스트가 잡는다)."""
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            settings = _settings(tmp_path)
            storage = _FakeStorage(tmp_path)
            oa = OpenAIService(settings)
            prof = _profile(tmp_path)
            ps = _FakeProjectService(storage, prof, oa)
            pi = ProjectInput(
                template_id="t1",
                organization_profile={"기업명": "테스트(주)"},
                project_meta={},
            )
            storage.save_project_input("p1", pi)
            pipeline = SubmissionPipeline(ps, EvaluationService(oa), storage, settings)
            fake = mock.Mock(submittable=True, fail_defects=0, results=[])
            with mock.patch(
                "auto_write.services.submission_orchestrator.run_acceptance",
                return_value=fake,
            ):
                report = pipeline.run(
                    "p1", announcement_text="",
                    enable_images=False, enable_notebooklm=False,
                )
            self.assertTrue(report["acceptance"]["submittable"])
            self.assertFalse(report["acceptance"]["draft_marked"])
            self.assertFalse(report["final_docx"].endswith("_DRAFT.docx"))
            for key in ("submit_docx", "quality_docx"):
                if key in report:
                    p = Path(report[key])
                    self.assertFalse(
                        p.stem.endswith(("_DRAFT", "_DRAFT2")), f"{key} 오마킹: {p.name}"
                    )
                    self.assertTrue(p.exists(), f"{key} 경로 없음: {p}")

    def test_pipeline_notebooklm_partial_output_loses_submit_name(self):
        """R9 잔여 보강: apply_images 가 파일 생성 '후' 죽으면 부분 생성본이
        제출 이름의 고아로 남지 않는다(즉시 _DRAFT 박탈 + 게이트 artifacts 등재)."""
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            settings = _settings(tmp_path)
            storage = _FakeStorage(tmp_path)
            oa = OpenAIService(settings)
            prof = _profile(tmp_path)
            ps = _FakeProjectService(storage, prof, oa)
            pi = ProjectInput(
                template_id="t1",
                organization_profile={"기업명": "테스트(주)"},
                project_meta={},
            )
            storage.save_project_input("p1", pi)
            pipeline = SubmissionPipeline(ps, EvaluationService(oa), storage, settings)

            def _boom(in_docx, out_docx, **kw):
                Path(out_docx).write_bytes(Path(in_docx).read_bytes())  # 파일 생성 후 사망
                raise RuntimeError("nlm crashed")

            with mock.patch(
                "auto_write.services.image_apply.apply_images", side_effect=_boom
            ):
                report = pipeline.run(
                    "p1", announcement_text="",
                    enable_images=False, enable_notebooklm=True,
                )
            self.assertIn("notebooklm_error", report)
            leftovers = [
                p.name for p in Path(settings.results_root).glob("제출초안_*노트북LM*.docx")
                if not p.stem.endswith(("_DRAFT", "_DRAFT2"))
            ]
            self.assertEqual(leftovers, [], "부분 생성본이 제출 이름으로 잔존")


if __name__ == "__main__":
    unittest.main()
