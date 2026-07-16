"""test_pure_psst_fill.py — psst_fill 안전 가드·직렬화·스캐폴딩 안전망.

PSST 작성 뼈대 삽입기의 핵심 안전 계약을 고정한다:
  - 원본 덮어쓰기 금지: in_docx == out_docx 면 ValueError(파일 접근 전 즉시 실패).
  - 없는 입력은 FileNotFoundError.
  - PSSTFillReport.as_dict 는 비율을 3자리로 반올림(순수).
  - 약한(누락/미흡) 영역이 있으면 문서 '끝'에 보강 가이드를 추가하되 원본은 미수정.

python-docx 로컬 객체만 사용(COM/네트워크 없음). 야간 순수함수 안전망(2026-07-13).
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from auto_write.services.psst_fill import (
    PSSTFillReport,
    apply_psst_scaffold,
    _DEFAULT_TARGET_GRADES,
)
from auto_write.services.usage_acceptance import SCAFFOLD_HEADING


# --- PSSTFillReport.as_dict: 순수 직렬화(비율 반올림) -----------------------

def test_report_as_dict_rounds_ratio_to_three_places():
    rep = PSSTFillReport(overall_ratio=0.123456)
    assert rep.as_dict()["overall_ratio"] == 0.123


def test_report_as_dict_target_grades_is_list_of_defaults():
    d = PSSTFillReport().as_dict()
    assert d["target_grades"] == list(_DEFAULT_TARGET_GRADES)
    assert d["areas_scaffolded"] == 0
    assert d["items_added"] == 0
    assert d["scaffolded_areas"] == []


# --- 안전 가드: 원본 보호 ---------------------------------------------------

def test_apply_scaffold_rejects_same_in_out(tmp_path: Path):
    same = str(tmp_path / "원본.docx")
    with pytest.raises(ValueError):
        apply_psst_scaffold(same, same)


def test_apply_scaffold_missing_input_raises(tmp_path: Path):
    with pytest.raises(FileNotFoundError):
        apply_psst_scaffold(str(tmp_path / "없음.docx"), str(tmp_path / "out.docx"))


# --- 실제 스캐폴딩: 약한 영역 보강 + 원본 미수정 ---------------------------

def test_apply_scaffold_adds_guide_and_preserves_original(tmp_path: Path):
    src = tmp_path / "빈약한계획서.docx"
    doc = Document()
    doc.add_paragraph("이 문서에는 PSST 섹션이 없습니다.")
    doc.save(str(src))
    before = src.read_bytes()

    out = tmp_path / "보강본.docx"
    rep = apply_psst_scaffold(str(src), str(out))

    # PSST 섹션이 없으니 최소 1개 영역은 보강 대상(누락/미흡)이 되어야 한다.
    assert rep.areas_scaffolded >= 1
    assert rep.items_added >= 1
    assert rep.scaffolded_areas
    assert rep.output_docx == str(out)

    # 결과 문서 끝에 표준 보강 헤딩이 실제로 들어간다(검출기와 공유하는 시그니처).
    out_text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert SCAFFOLD_HEADING in out_text

    # 원본은 바이트 단위로 그대로(미수정) — 결과는 별도 파일.
    assert src.read_bytes() == before
    assert out.exists() and out != src
