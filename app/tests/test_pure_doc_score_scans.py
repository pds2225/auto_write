"""test_pure_doc_score_scans.py — doc_quality_score 잔존결함 스캐너 회귀.

채점(score_document)의 감점 근거가 되는 결정론 스캐너들을 in-memory Document 로
직접 단위 검증한다(AI·COM·네트워크 없음, python-docx 로컬 객체만).

고정하는 계약:
- _scan_font_sizes: 본문 폰트 크기 종류 수 + 이상치(<8pt·>18pt) 수, size=None 은 무시.
- _count_nonempty_paragraphs / _count_bold_paragraphs: 비어있지 않은/굵은 단락 수.
- _scan_bullet: 줄머리 글머리표+2칸 이상 공백 또는 내부 2칸 이상 공백 결함 수.
- _scan_empty_groups: 연속 빈 단락(2개 이상) 그룹 수(표는 연속 카운터 리셋).
- _scan_guide: body 직계 안내문구(critical) / OOO·○○○ 플레이스홀더(general).
"""

from __future__ import annotations

from docx import Document
from docx.shared import Pt

from auto_write.services.doc_quality_score import (
    _count_bold_paragraphs,
    _count_nonempty_paragraphs,
    _scan_bullet,
    _scan_empty_groups,
    _scan_font_sizes,
    _scan_guide,
)


# --------------------------------------------------------------------------
# _scan_font_sizes
# --------------------------------------------------------------------------

def test_font_sizes_counts_distinct_and_no_outliers():
    doc = Document()
    for sz in (10, 10, 12):
        p = doc.add_paragraph()
        r = p.add_run("본문")
        r.font.size = Pt(sz)
    distinct, outliers = _scan_font_sizes(doc)
    assert distinct == 2      # {10, 12}
    assert outliers == 0


def test_font_sizes_flags_too_small_and_too_large():
    doc = Document()
    for sz in (6, 20, 11):     # 6<8, 20>18 → 이상치 2
        p = doc.add_paragraph()
        p.add_run("x").font.size = Pt(sz)
    distinct, outliers = _scan_font_sizes(doc)
    assert distinct == 3
    assert outliers == 2


def test_font_sizes_boundaries_8_and_18_not_outliers():
    doc = Document()
    for sz in (8, 18):
        p = doc.add_paragraph()
        p.add_run("경계").font.size = Pt(sz)
    _, outliers = _scan_font_sizes(doc)
    assert outliers == 0       # 8·18 은 경계 안(이상치 아님)


def test_font_sizes_ignores_none_size():
    doc = Document()
    doc.add_paragraph("크기 미지정")  # run.font.size is None
    distinct, outliers = _scan_font_sizes(doc)
    assert distinct == 0
    assert outliers == 0


# --------------------------------------------------------------------------
# _count_nonempty_paragraphs / _count_bold_paragraphs
# --------------------------------------------------------------------------

def test_count_nonempty_paragraphs():
    doc = Document()
    doc.add_paragraph("있음")
    doc.add_paragraph("   ")     # 공백뿐 → 미포함
    doc.add_paragraph("또 있음")
    assert _count_nonempty_paragraphs(doc) == 2


def test_count_bold_paragraphs_requires_nonempty_bold_run():
    doc = Document()
    p1 = doc.add_paragraph()
    p1.add_run("강조").bold = True          # 굵은 + 내용 → 카운트
    p2 = doc.add_paragraph()
    r = p2.add_run("")                       # 굵지만 내용 없음 → 제외
    r.bold = True
    doc.add_paragraph("보통 문장")          # 굵지 않음 → 제외
    assert _count_bold_paragraphs(doc) == 1


# --------------------------------------------------------------------------
# _scan_bullet
# --------------------------------------------------------------------------

def test_scan_bullet_flags_bullet_prefix_and_inner_multispace():
    doc = Document()
    doc.add_paragraph("•  글머리표 뒤 두 칸")   # 글머리표 + 2칸 → 결함
    doc.add_paragraph("문장 안에  두 칸 공백")   # 내부 2칸 → 결함
    doc.add_paragraph("정상 한 칸 문장")          # 결함 아님
    assert _scan_bullet(doc) == 2


def test_scan_bullet_clean_document_zero():
    doc = Document()
    doc.add_paragraph("깔끔한 문장 하나")
    doc.add_paragraph("또 다른 정상 문장")
    assert _scan_bullet(doc) == 0


# --------------------------------------------------------------------------
# _scan_empty_groups
# --------------------------------------------------------------------------

def test_empty_groups_counts_consecutive_run():
    doc = Document()
    doc.add_paragraph("A")
    doc.add_paragraph("")
    doc.add_paragraph("")   # 연속 2개 → 한 그룹
    doc.add_paragraph("B")
    assert _scan_empty_groups(doc) == 1


def test_empty_groups_single_empties_not_counted():
    doc = Document()
    doc.add_paragraph("A")
    doc.add_paragraph("")   # 단독
    doc.add_paragraph("B")
    doc.add_paragraph("")   # 끝 단독
    assert _scan_empty_groups(doc) == 0


def test_empty_groups_table_resets_run():
    doc = Document()
    doc.add_paragraph("")   # 표 앞 단독 빈 단락
    doc.add_table(rows=1, cols=1)   # 표 → 연속 카운터 리셋
    doc.add_paragraph("")   # 표 뒤 단독 빈 단락
    # 표가 사이를 끊으므로 두 빈 단락이 '연속 2개'로 오인되지 않는다.
    assert _scan_empty_groups(doc) == 0


# --------------------------------------------------------------------------
# _scan_guide
# --------------------------------------------------------------------------

def test_scan_guide_clean_document_zero():
    doc = Document()
    doc.add_paragraph("우리 회사는 2020년에 설립되었습니다.")
    assert _scan_guide(doc) == (0, 0)


def test_scan_guide_detects_critical_guide_phrase():
    doc = Document()
    doc.add_paragraph("작성요령: 아래에 사업 내용을 기술")
    critical, general = _scan_guide(doc)
    assert critical == 1
    assert general == 0


def test_scan_guide_detects_placeholder_as_general():
    doc = Document()
    doc.add_paragraph("성명: ○○○")   # 플레이스홀더 → general
    critical, general = _scan_guide(doc)
    assert critical == 0
    assert general == 1
