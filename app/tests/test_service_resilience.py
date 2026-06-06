from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
import httpx

from auto_write.config import Settings, ensure_directories
from auto_write.models import (
    GeneratedImage,
    ImageSlotProfile,
    ProjectInput,
    QuestionProfile,
    SectionProfile,
    TableCellProfile,
    TableProfile,
    TemplateProfile,
)
from auto_write.services.image_service import ImageService
from auto_write.services.openai_client import OpenAIService
from auto_write.services.qa_service import QAService
from auto_write.services.render_service import RenderService


class _FailingResponses:
    def create(self, **kwargs):
        raise RuntimeError("forced responses failure")


class _FailingImages:
    def generate(self, **kwargs):
        raise RuntimeError("forced image failure")


class _FailingClient:
    def __init__(self):
        self.responses = _FailingResponses()
        self.images = _FailingImages()


def build_settings(root: Path) -> Settings:
    app_root = root / "app"
    workspace_root = root / "workspace"
    return Settings(
        app_root=app_root,
        workspace_root=workspace_root,
        template_root=workspace_root / "templates",
        project_root=workspace_root / "projects",
        results_root=root / "results",
        static_root=app_root / "auto_write" / "static",
        template_view_root=app_root / "auto_write" / "templates",
        host="127.0.0.1",
        port=8765,
        openai_api_key="",
        openai_model="gpt-4.1-mini",
        openai_search_model="gpt-4.1-mini",
        openai_image_model="gpt-image-1",
        anthropic_api_key="",
        anthropic_model="claude-sonnet-4-20250514",
        anthropic_search_model="claude-sonnet-4-20250514",
    )


class ServiceResilienceTests(unittest.TestCase):
    def test_service_uses_anthropic_provider_when_only_anthropic_key_exists(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            root = Path(tmp_dir)
            settings = Settings(
                app_root=root / "app",
                workspace_root=root / "workspace",
                template_root=root / "workspace" / "templates",
                project_root=root / "workspace" / "projects",
                results_root=root / "results",
                static_root=root / "app" / "auto_write" / "static",
                template_view_root=root / "app" / "auto_write" / "templates",
                host="127.0.0.1",
                port=8765,
                openai_api_key="",
                openai_model="gpt-4.1-mini",
                openai_search_model="gpt-4.1-mini",
                openai_image_model="gpt-image-1",
                anthropic_api_key="test-key",
                anthropic_model="claude-sonnet-4-20250514",
                anthropic_search_model="claude-sonnet-4-20250514",
            )
            service = OpenAIService(settings)
            self.assertTrue(service.available)
            self.assertEqual(service.provider, "anthropic")
            self.assertIn("Claude", service.status_text)

    def test_openai_service_returns_safe_defaults_on_client_error(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            settings = build_settings(Path(tmp_dir))
            ensure_directories(settings)
            service = OpenAIService(settings)
            service._client = _FailingClient()

            self.assertIsNone(service.complete_json("s", "u"))
            self.assertEqual(service.web_search_sources("시장 분석"), [])
            self.assertIn("설명", service.propose_image_prompt("설명 이미지", "context"))
            target = Path(tmp_dir) / "img.png"
            self.assertFalse(service.generate_image_file("prompt", target))

    def test_normalize_sources_payload_filters_invalid_and_duplicate_urls(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            settings = build_settings(Path(tmp_dir))
            ensure_directories(settings)
            service = OpenAIService(settings)
            payload = {
                "sources": [
                    {"title": "A", "url": "https://example.com/a", "organization": "OrgA", "summary": "one"},
                    {"title": "A-dup", "url": "https://example.com/a", "organization": "OrgA2", "summary": "dup"},
                    {"title": "", "url": "https://example.com/b", "organization": "OrgB", "summary": "empty title"},
                    {"title": "No URL", "url": "", "organization": "OrgC", "summary": "empty url"},
                    {"title": "C", "url": "https://example.com/c", "organization": "OrgC", "summary": "three"},
                ]
            }
            normalized = service._normalize_sources_payload(payload)
            self.assertEqual(len(normalized), 2)
            self.assertEqual(normalized[0]["url"], "https://example.com/a")
            self.assertEqual(normalized[1]["url"], "https://example.com/c")

    @patch("auto_write.services.openai_client.httpx.post")
    def test_anthropic_auth_error_disables_runtime(self, mock_post):
        with tempfile.TemporaryDirectory() as tmp_dir:
            root = Path(tmp_dir)
            settings = Settings(
                app_root=root / "app",
                workspace_root=root / "workspace",
                template_root=root / "workspace" / "templates",
                project_root=root / "workspace" / "projects",
                results_root=root / "results",
                static_root=root / "app" / "auto_write" / "static",
                template_view_root=root / "app" / "auto_write" / "templates",
                host="127.0.0.1",
                port=8765,
                openai_api_key="",
                openai_model="gpt-4.1-mini",
                openai_search_model="gpt-4.1-mini",
                openai_image_model="gpt-image-1",
                anthropic_api_key="test-invalid",
                anthropic_model="claude-sonnet-4-20250514",
                anthropic_search_model="claude-sonnet-4-20250514",
            )
            service = OpenAIService(settings)
            self.assertTrue(service.available)

            request = httpx.Request("POST", "https://api.anthropic.com/v1/messages")
            response = httpx.Response(401, request=request)
            mock_post.side_effect = httpx.HTTPStatusError("401 Unauthorized", request=request, response=response)

            result = service.complete_json("Return JSON only", "test")
            self.assertIsNone(result)
            self.assertFalse(service.available)
            self.assertIn("인증 실패", service.status_text)

    def test_render_and_qa_report_table_index_error_clearly(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            source_docx = Path(tmp_dir) / "sample.docx"
            doc = Document()
            doc.add_paragraph("1. 사업 개요")
            table = doc.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "제목"
            doc.save(source_docx)

            profile = TemplateProfile(
                template_id="tpl_x",
                template_name="sample.docx",
                source_docx=str(source_docx),
                tables=[
                    TableProfile(
                        table_id="table_01",
                        label="테스트 표",
                        table_index=0,
                        row_count=1,
                        col_count=1,
                        cells=[
                            TableCellProfile(
                                cell_id="table_01_r2_c0",
                                label="잘못된 셀",
                                row=2,
                                cell=0,
                                required=True,
                            )
                        ],
                    )
                ],
            )
            project_input = ProjectInput(
                template_id="tpl_x",
                answers={"table_01_r2_c0": "값"},
            )

            render_service = RenderService()
            output_path = Path(tmp_dir) / "output.docx"
            render_result = render_service.render(profile, project_input, [], output_path)
            self.assertTrue(any("행 index" in item for item in render_result["errors"]))

            qa_service = QAService()
            report = qa_service.build_report(profile, project_input, render_result, [], [])
            self.assertTrue(any(item.startswith("[렌더링]") for item in report["errors"]))

    def test_render_handles_invalid_image_anchor_index_without_crash(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            source_docx = Path(tmp_dir) / "img_anchor.docx"
            doc = Document()
            doc.add_paragraph("1. 이미지 영역")
            table = doc.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "이미지"
            doc.save(source_docx)

            image_path = Path(tmp_dir) / "dummy.png"
            image_path.write_bytes(b"not-a-real-png")

            profile = TemplateProfile(
                template_id="tpl_img_anchor",
                template_name="img_anchor.docx",
                source_docx=str(source_docx),
                image_slots=[
                    ImageSlotProfile(
                        slot_id="slot_bad",
                        label="잘못된 이미지 슬롯",
                        required=True,
                        anchor_type="table_cell",
                        anchor_ref={"table_index": "x", "row": "0", "cell": "0"},
                    )
                ],
            )
            project_input = ProjectInput(template_id="tpl_img_anchor", answers={})

            render_service = RenderService()
            output_path = Path(tmp_dir) / "output.docx"
            render_result = render_service.render(
                profile,
                project_input,
                [GeneratedImage(slot_id="slot_bad", label="잘못된 이미지 슬롯", path=str(image_path), source="generated")],
                output_path,
            )

            self.assertTrue(Path(output_path).exists())
            self.assertTrue(any("이미지 표 위치 오류" in item for item in render_result["errors"]))

    def test_render_skips_non_business_marker_section(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            source_docx = Path(tmp_dir) / "marker.docx"
            doc = Document()
            doc.add_paragraph("◦")
            doc.save(source_docx)

            profile = TemplateProfile(
                template_id="tpl_marker",
                template_name="marker.docx",
                source_docx=str(source_docx),
                sections=[
                    SectionProfile(
                        field_id="section_bad",
                        label="◦",
                        anchor_text="◦",
                    )
                ],
            )
            project_input = ProjectInput(template_id="tpl_marker", answers={"section_bad": "작성되면 안 됨"})

            render_service = RenderService()
            output_path = Path(tmp_dir) / "output.docx"
            render_result = render_service.render(profile, project_input, [], output_path)
            output_doc = Document(str(output_path))

            self.assertEqual(render_result["sections_written"], 0)
            self.assertTrue(any("비본문 섹션" in item for item in render_result["warnings"]))
            self.assertFalse(any("작성되면 안 됨" in paragraph.text for paragraph in output_doc.paragraphs))

    def test_qa_anchor_match_handles_duplicated_heading_text(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            doc = Document()
            doc.add_paragraph("2. 실현 가능성 (Solution)_창업 아이템의 개발 계획2. 실현 가능성 (Solution)_창업 아이템의 개발 계획")
            doc.add_paragraph("핵심 기능 구현과 실증 계획을 협약기간 내 완료합니다.")
            path = Path(tmp_dir) / "dup.docx"
            doc.save(path)

            qa_service = QAService()
            loaded = Document(str(path))
            self.assertTrue(
                qa_service._has_content_after_anchor(loaded, "2. 실현 가능성 (Solution)_창업 아이템의 개발 계획")
            )

    def test_qa_prioritizes_output_doc_content_for_required_section(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = Path(tmp_dir) / "output.docx"
            doc = Document()
            doc.add_paragraph("1. 사업 개요")
            doc.add_paragraph("시장 문제를 해결하기 위한 핵심 전략과 실행 계획")
            doc.save(output_path)

            profile = TemplateProfile(
                template_id="tpl_section",
                template_name="section.docx",
                source_docx=str(output_path),
                sections=[
                    SectionProfile(
                        field_id="section_001",
                        label="1. 사업 개요",
                        anchor_text="1. 사업 개요",
                        required=True,
                    )
                ],
                questions=[
                    QuestionProfile(
                        question_id="section_001",
                        label="1. 사업 개요",
                        required=True,
                        target={"kind": "section", "field_id": "section_001"},
                    )
                ],
            )
            project_input = ProjectInput(template_id="tpl_section", answers={})
            render_result = {
                "output_path": str(output_path),
                "sections_written": 1,
                "cells_written": 0,
                "images_written": 0,
                "errors": [],
                "warnings": [],
            }
            qa_service = QAService()
            report = qa_service.build_report(profile, project_input, render_result, [], [])
            self.assertEqual(report["error_count"], 0)

    def test_qa_ignores_required_signature_style_questions(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = Path(tmp_dir) / "output.docx"
            doc = Document()
            doc.add_paragraph("담 당 자 : (인)")
            doc.save(output_path)

            profile = TemplateProfile(
                template_id="tpl_signature",
                template_name="signature.docx",
                source_docx=str(output_path),
                sections=[
                    SectionProfile(
                        field_id="section_signature",
                        label="담 당 자 : (인)",
                        anchor_text="담 당 자 : (인)",
                        required=True,
                    )
                ],
                questions=[
                    QuestionProfile(
                        question_id="section_signature",
                        label="담 당 자 : (인)",
                        required=True,
                        target={"kind": "section", "field_id": "section_signature"},
                    )
                ],
            )
            project_input = ProjectInput(template_id="tpl_signature", answers={})
            render_result = {
                "output_path": str(output_path),
                "sections_written": 0,
                "cells_written": 0,
                "images_written": 0,
                "errors": [],
                "warnings": [],
            }

            qa_service = QAService()
            report = qa_service.build_report(profile, project_input, render_result, [], [])
            self.assertEqual(report["error_count"], 0)

    def test_image_service_skips_non_business_suggested_slot_without_note(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            settings = build_settings(Path(tmp_dir))
            ensure_directories(settings)
            openai_service = OpenAIService(settings)
            image_service = ImageService(openai_service)
            images = image_service.build_images(
                [
                    ImageSlotProfile(
                        slot_id="img_skip",
                        label="평가대상에서 제외될 수 있음 설명 이미지",
                        required=False,
                        anchor_type="after_paragraph",
                        anchor_ref={"anchor_text": "평가대상에서 제외될 수 있음"},
                        source="suggested",
                    )
                ],
                answers={},
                evidence=[],
                output_dir=Path(tmp_dir) / "generated",
            )

            self.assertEqual(images, [])

    def test_qa_uses_target_field_id_when_question_id_differs(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = Path(tmp_dir) / "output.docx"
            doc = Document()
            doc.add_paragraph("1. 사업 개요")
            doc.add_paragraph("핵심 실행 전략과 일정")
            doc.save(output_path)

            profile = TemplateProfile(
                template_id="tpl_section_alias",
                template_name="section_alias.docx",
                source_docx=str(output_path),
                sections=[
                    SectionProfile(
                        field_id="section_real",
                        label="1. 사업 개요",
                        anchor_text="1. 사업 개요",
                        required=True,
                    )
                ],
                questions=[
                    QuestionProfile(
                        question_id="legacy_question_id",
                        label="1. 사업 개요",
                        required=True,
                        target={"kind": "section", "field_id": "section_real"},
                    )
                ],
            )
            project_input = ProjectInput(template_id="tpl_section_alias", answers={})
            render_result = {
                "output_path": str(output_path),
                "sections_written": 1,
                "cells_written": 0,
                "images_written": 0,
                "errors": [],
                "warnings": [],
            }
            qa_service = QAService()
            report = qa_service.build_report(profile, project_input, render_result, [], [])
            self.assertEqual(report["error_count"], 0)


if __name__ == "__main__":
    unittest.main()
