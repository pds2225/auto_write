from __future__ import annotations

from io import BytesIO

from docx import Document

from auto_write.operator_main import _prefill_template_from_references


def _docx_bytes(build) -> bytes:
    doc = Document()
    build(doc)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_unified_write_route_prefills_new_form_from_existing_plan():
    source = _docx_bytes(
        lambda doc: (
            (lambda table: (
                setattr(table.cell(0, 0), "text", "대표자"),
                setattr(table.cell(0, 1), "text", "홍길동"),
            ))(doc.add_table(rows=1, cols=2))
        )
    )
    target = _docx_bytes(
        lambda doc: (
            (lambda table: (
                setattr(table.cell(0, 0), "text", "대표자"),
                setattr(table.cell(0, 1), "text", ""),
            ))(doc.add_table(rows=1, cols=2))
        )
    )

    name, filled_bytes, stats = _prefill_template_from_references(
        "신규_사업계획서_양식.docx",
        target,
        [("기존_완성_사업계획서.docx", source)],
    )

    assert name == "신규_사업계획서_양식.docx"
    assert stats["mode"] == "cross_form_then_bizplan"
    assert stats["transcribed"] >= 1
    assert stats["applied_sources"] == ["기존_완성_사업계획서.docx"]

    filled = Document(BytesIO(filled_bytes))
    assert filled.tables[0].cell(0, 1).text == "홍길동"


def test_unified_write_route_keeps_new_generation_when_no_existing_document_reference():
    target = _docx_bytes(
        lambda doc: (
            (lambda table: (
                setattr(table.cell(0, 0), "text", "대표자"),
                setattr(table.cell(0, 1), "text", ""),
            ))(doc.add_table(rows=1, cols=2))
        )
    )

    _, same_bytes, stats = _prefill_template_from_references(
        "신규_사업계획서_양식.docx",
        target,
        [("지원사업_공고.pdf", b"not-a-real-pdf")],
    )

    assert stats["mode"] == "bizplan"
    assert stats["attempted"] == 0
    assert same_bytes == target
