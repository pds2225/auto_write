"""Wave B/C/D 교훈 가드 — 제출·이력서·COM kill spy.

계획 ID 와 JSON ID 가 다른 항목은 JSON 요약을 기준으로만 mechanized 표시한다.
L154–L156 은 lessons.md 에 없으면 coverage JSON 에 넣지 않고 코드+테스트만.
"""
from __future__ import annotations

import sys
import types
from pathlib import Path
from unittest import mock

import pytest
from docx import Document

from auto_write.document_ingest import ensure_template_docx
from auto_write.models import ProjectInput
from auto_write.services.doc_quality_ops import bold_bullet_paren_labels, run_all
from auto_write.services.evaluation_service import EvaluationService
from auto_write.services.hwpx_fill import fill_hwpx
from auto_write.services.hwpx_pic_insert import REF_IMAGE_FRAME_MM as PIC_FRAME
from auto_write.services.hwpx_resume_supplement import supplement_hwpx_from_resume
from auto_write.services.openai_client import OpenAIService
from auto_write.services.resume_extract import parse_profile_text
from auto_write.services.submission_gates import (
    AnnouncementFormError,
    PORTFOLIO_MARKER,
    PHOTO_MARKER,
    REF_IMAGE_FRAME_MM,
    SampleSectionError,
    dropped_other_history,
    estimate_page_count,
    infer_hangul_required,
    is_announcement_form_path,
    missing_resume_skeleton_sections,
    page_count_increased,
    photo_slot_ok,
    portfolio_ok,
    reference_image_frame_hwpunit,
    require_sample_ok,
    resume_layout_warnings,
    safe_body_accent,
    slash_combo_headers,
    submit_folder_contamination,
    work_suffix_hits,
)
from auto_write.services.submission_orchestrator import SubmissionPipeline
from auto_write.services.usage_acceptance import run_acceptance
from core.docx.services import hwp_docx_convert
from core.docx.services.hwp_docx_convert import kill_hangul_processes
from test_submission_pipeline import (
    _FakeProjectService,
    _FakeStorage,
    _profile,
    _settings,
)


# ---------------------------------------------------------------------------
# L059 / L048 파일명·폴더
# ---------------------------------------------------------------------------


def test_l059_work_suffix_hits_denylist_not_draft():
    assert "_converted" in work_suffix_hits("양식_converted.docx")
    assert "_노트북LM" in work_suffix_hits("제출초안_p1_노트북LM.docx")
    assert work_suffix_hits("신청서_박다솜_DRAFT.hwpx") == []
    assert work_suffix_hits("신청서_박다솜.docx") == []


def test_l059_acceptance_warns_on_converted_filename(tmp_path: Path):
    p = tmp_path / "초안_converted.docx"
    Document().save(str(p))
    report = run_acceptance(p)
    hits = [r for r in report.results if r.check_id == "work_suffix_filename"]
    assert hits and hits[0].defects >= 1
    assert hits[0].severity == "warn"


def test_l048_submit_folder_contamination_mix_names():
    dirty = submit_folder_contamination(
        ["원본.docx", "신청서_박다솜.hwpx", "모집공고.hwpx", "a_converted.docx"]
    )
    assert "원본.docx" in dirty
    assert "모집공고.hwpx" in dirty
    assert "a_converted.docx" in dirty
    assert "신청서_박다솜.hwpx" not in dirty


def test_l059_pipeline_flags_notebooklm_suffix(tmp_path: Path):
    settings = _settings(tmp_path)
    storage = _FakeStorage(tmp_path)
    oa = OpenAIService(settings)
    prof = _profile(tmp_path)
    ps = _FakeProjectService(storage, prof, oa)
    storage.save_project_input(
        "p1", ProjectInput(template_id="t1", organization_profile={"기업명": "테스트(주)"}, project_meta={})
    )
    pipeline = SubmissionPipeline(ps, EvaluationService(oa), storage, settings)
    fake = mock.Mock(submittable=True, fail_defects=0, results=[])
    with mock.patch(
        "auto_write.services.submission_orchestrator.run_acceptance",
        return_value=fake,
    ):
        report = pipeline.run(
            "p1", announcement_text="",
            enable_images=False, enable_notebooklm=True,
        )
    assert report.get("work_suffix"), report
    assert any("L059" in n for n in report["needs_input"])
    assert report.get("submit_mix")


# ---------------------------------------------------------------------------
# L037 (= 계획 L049) 공고를 양식으로 채우지 않음
# ---------------------------------------------------------------------------


def test_l037_announcement_path_detection():
    assert is_announcement_form_path("모집공고.hwpx")
    assert is_announcement_form_path("안내.pdf")
    assert not is_announcement_form_path("전문상담위원_참여신청서.hwpx")
    assert not is_announcement_form_path("공고결과보고서.docx")  # 단독 '공고' 오탐 금지


def test_l037_fill_hwpx_rejects_announcement_name(tmp_path: Path):
    src = tmp_path / "모집공고.hwpx"
    src.write_bytes(b"PK\x03\x04 dummy")
    with pytest.raises(AnnouncementFormError, match="L049|공고"):
        fill_hwpx(src, tmp_path / "out.hwpx", identity={"상호": "테스트"})


def test_l037_ensure_template_rejects_pdf(tmp_path: Path):
    pdf = tmp_path / "모집요강.pdf"
    pdf.write_bytes(b"%PDF-1.4")
    with pytest.raises(AnnouncementFormError):
        ensure_template_docx(pdf)


# ---------------------------------------------------------------------------
# L050 계획: 한글 전용 → DOCX 산출은 형식 게이트 _DRAFT
# (JSON L050 은 HWP+PDF 쌍이라 gap 유지)
# ---------------------------------------------------------------------------


def test_infer_hangul_required_patterns():
    assert infer_hangul_required("제출은 한글 전용입니다")
    assert infer_hangul_required("HWPX로 제출")
    assert not infer_hangul_required("")
    assert not infer_hangul_required("한글과 컴퓨터를 사용해도 됩니다")


def test_l050_hangul_required_announcement_drafts_docx(tmp_path: Path):
    settings = _settings(tmp_path)
    storage = _FakeStorage(tmp_path)
    oa = OpenAIService(settings)
    prof = _profile(tmp_path)
    ps = _FakeProjectService(storage, prof, oa)
    storage.save_project_input(
        "p1", ProjectInput(template_id="t1", organization_profile={"기업명": "테스트(주)"}, project_meta={})
    )
    pipeline = SubmissionPipeline(ps, EvaluationService(oa), storage, settings)
    fake = mock.Mock(submittable=True, fail_defects=0, results=[])
    with mock.patch(
        "auto_write.services.submission_orchestrator.run_acceptance",
        return_value=fake,
    ):
        report = pipeline.run(
            "p1",
            announcement_text="본 사업 제출 파일은 한글 전용입니다.",
            enable_images=False,
            enable_notebooklm=False,
        )
    assert report.get("format_mismatch")
    assert Path(report["final_docx"]).name.endswith("_DRAFT.docx")


def test_l040_pipeline_required_doc_drafts(tmp_path: Path):
    settings = _settings(tmp_path)
    storage = _FakeStorage(tmp_path)
    oa = OpenAIService(settings)
    prof = _profile(tmp_path)
    ps = _FakeProjectService(storage, prof, oa)
    storage.save_project_input(
        "p1", ProjectInput(template_id="t1", organization_profile={"기업명": "테스트(주)"}, project_meta={})
    )
    pipeline = SubmissionPipeline(ps, EvaluationService(oa), storage, settings)
    report = pipeline.run(
        "p1",
        announcement_text="",
        enable_images=False,
        enable_notebooklm=False,
        required_documents=("청렴서약서",),
    )
    acc = report.get("acceptance") or {}
    assert acc.get("submittable") is False
    assert Path(report["final_docx"]).name.endswith("_DRAFT.docx")


# ---------------------------------------------------------------------------
# L095 페이지 기준선
# ---------------------------------------------------------------------------


def test_l095_estimate_page_count_and_fill_report(tmp_path: Path):
    assert estimate_page_count(tmp_path / "missing.txt") == 1
    assert page_count_increased(1, 2)
    assert not page_count_increased(2, 2)
    from test_hwpx_fill import _make_hwpx

    src = tmp_path / "form.hwpx"
    _make_hwpx(src)
    out = tmp_path / "out.hwpx"
    rep = fill_hwpx(src, out, identity={"상호": "테스트법인"})
    assert rep.pages_before >= 1
    assert rep.pages_after >= 1
    assert "pages_before" in rep.as_dict()


# ---------------------------------------------------------------------------
# L080 글머리 직후 괄호 라벨 볼드
# ---------------------------------------------------------------------------


def test_l080_bold_bullet_paren_label_not_mid_sentence():
    doc = Document()
    doc.add_paragraph("ㅇ (문제인식) 본문과 (중간) 괄호")
    doc.add_paragraph("문장 중간만 (참고) 있다")
    n = bold_bullet_paren_labels(doc)
    assert n == 1
    first = [(r.text, bool(r.bold)) for r in doc.paragraphs[0].runs]
    assert any("(문제인식)" in (t or "") and b for t, b in first)
    mid = "".join(r.text or "" for r in doc.paragraphs[1].runs)
    assert "(참고)" in mid
    assert not any(bool(r.bold) and "(참고)" in (r.text or "") for r in doc.paragraphs[1].runs)


def test_l080_run_all_wires_bold_paren():
    doc = Document()
    doc.add_paragraph("ㅇ (문제인식) 본문")
    report = run_all(doc, emphasize=False)
    assert report.bullet_paren_labels_bolded >= 1
    assert report.as_dict()["bullet_paren_labels_bolded"] >= 1


# ---------------------------------------------------------------------------
# Wave C 이력서
# ---------------------------------------------------------------------------


def test_l038_company_count_and_total_row_kept():
    text = (
        "수행기간 | 프로젝트명 | 수행내용 | 발주처 | 업체수\n"
        "2023.01~2023.12 | A사업 | 컨설팅 | 공단 | 12\n"
        "합계 | 12 |  |  |\n"
    )
    p = parse_profile_text(text, source="fx.hwp")
    assert p.projects[0].company_count == "12"
    assert p.projects[0].is_total is False
    totals = [x for x in p.projects if x.is_total]
    assert totals and totals[0].company_count == "12"


def test_l060_lecture_kind_still_parsed():
    text = (
        "수행기간 | 주최기관명 | 강의 주제 | 회차/시간 | 구분\n"
        "2024.12.15 | 초이비즈니스그룹 | 기초 | 1회 2시간 | 민간\n"
    )
    p = parse_profile_text(text, source="fx.hwp")
    assert p.lectures[0].kind == "민간"
    assert p.lectures[0].count == "1회 2시간"


def test_l039_l043_l044_l061_resume_layout_helpers():
    assert slash_combo_headers("학력 / 경력 통합")
    assert not slash_combo_headers("학력\n경력사항\n강의")
    blob = "성명 박다솜\n경력사항\n주최기관명\n프로젝트명\n기타 이력"
    assert missing_resume_skeleton_sections(blob) == []
    assert missing_resume_skeleton_sections("성명만") == ["경력", "강의", "수행"]
    assert dropped_other_history("기타 이력 있음", "인적만")
    assert not dropped_other_history("기타", "기타 유지")
    assert not portfolio_ok("텍스트만")
    assert portfolio_ok(f"끝\n{PORTFOLIO_MARKER}")
    assert portfolio_ok("끝", has_image=True)
    assert not photo_slot_ok("텍스트만")
    assert photo_slot_ok(PHOTO_MARKER)
    warns = resume_layout_warnings("성명만")
    assert any(w.startswith("L039") for w in warns)
    assert any(w.startswith("L044") for w in warns)
    assert any(w.startswith("L061") for w in warns)


def test_l061_confirm_output_plan_still_required():
    """계획 L061(출력 형식 확인)은 기존 validate_output_plan. JSON L061은 사진칸."""
    from auto_write.services.cross_form_output_policy import OutputPolicyError, validate_output_plan
    from test_cross_form_output_policy import _plan

    with pytest.raises(OutputPolicyError, match="confirm-output-plan"):
        validate_output_plan(_plan(["docx"], "docx-crossform", confirmed=False))


# ---------------------------------------------------------------------------
# L154–L156 코드만 (coverage JSON 금지)
# ---------------------------------------------------------------------------


def test_l154_sample_ok_required_before_full_document():
    require_sample_ok(sample_ok=True, full_document=True)
    require_sample_ok(sample_ok=False, full_document=False)
    with pytest.raises(SampleSectionError, match="L154"):
        require_sample_ok(sample_ok=False, full_document=True)
    with pytest.raises(SampleSectionError):
        supplement_hwpx_from_resume(
            "missing.hwpx", "out.hwpx", sample_ok=False, full_document=True
        )


def test_l155_l156_accent_and_reference_frame():
    assert safe_body_accent("#0000FF") == "2E74B5"
    assert safe_body_accent("0000FF") == "2E74B5"
    assert safe_body_accent("FF0000") == "FF0000"
    assert REF_IMAGE_FRAME_MM == (170.0, 55.0)
    assert PIC_FRAME == REF_IMAGE_FRAME_MM
    w, h = reference_image_frame_hwpunit()
    assert w > 0 and h > 0
    assert w > h


# ---------------------------------------------------------------------------
# Wave D L003 COM kill spy
# ---------------------------------------------------------------------------


def test_l003_kill_hangul_noop_off_windows():
    if sys.platform == "win32":
        pytest.skip("Windows 실 taskkill 은 이 클라우드에서 돌리지 않음")
    assert kill_hangul_processes() == []


def test_l003_dispatch_calls_kill_before_com(monkeypatch):
    calls: list[str] = []
    monkeypatch.setattr(
        hwp_docx_convert, "kill_hangul_processes", lambda: calls.append("kill") or ["Hwp.exe"]
    )

    class _Client:
        @staticmethod
        def Dispatch(progid):
            calls.append(f"dispatch:{progid}")
            return object()

    win32 = types.SimpleNamespace(client=_Client)
    monkeypatch.setitem(sys.modules, "win32com", win32)
    monkeypatch.setitem(sys.modules, "win32com.client", _Client)
    hwp_docx_convert._dispatch_hwp(skip_com_guard=True)
    assert calls[0] == "kill"
    assert calls[1] == "dispatch:HWPFrame.HwpObject"
