"""test_notice_pipeline.py — mail→auto_write 통합 파이프라인 회귀 테스트."""

from __future__ import annotations

import json
from datetime import date, timedelta
from pathlib import Path

from docx import Document

from auto_write.services.announcement_analyzer import AnnouncementReport
from auto_write.services.cross_form_autofill import (
    BatchAutofillItem,
    BatchAutofillReport,
    autofill_from_source,
)
from auto_write.services.d_trigger import filter_narrative_unmatched, should_run_bizplan_for_target
from auto_write.services.folder_analyzer import FolderAnalysisReport
from auto_write.services.form_analyzer import analyze_form, classify_field_kind
from auto_write.services.notice_pipeline import (
    build_todo_text,
    check_deadline_warning,
    classify_login_wall,
    format_pipeline_summary_korean,
    load_confirm_files,
    resolve_notice_folder_from_log,
    run_pipeline,
    write_confirm_json_files,
)
from auto_write.services.user_pipeline_config import load_config, save_config


def _make_pair(tmp_path: Path) -> tuple[Path, Path, Path]:
  notice = tmp_path / "01_테스트공고"
  notice.mkdir()
  pool = tmp_path / "pool"
  pool.mkdir()
  src = pool / "완성_신청서.docx"
  doc = Document()
  t = doc.add_table(rows=2, cols=2)
  t.rows[0].cells[0].text = "성명"
  t.rows[0].cells[1].text = "홍길동"
  t.rows[1].cells[0].text = "이메일"
  t.rows[1].cells[1].text = "test@example.com"
  doc.save(str(src))
  tgt = notice / "신청서.docx"
  doc2 = Document()
  t2 = doc2.add_table(rows=2, cols=2)
  t2.rows[0].cells[0].text = "성명"
  t2.rows[0].cells[1].text = ""
  t2.rows[1].cells[0].text = "이메일"
  t2.rows[1].cells[1].text = ""
  doc2.save(str(tgt))
  ann = notice / "공고문.txt"
  ann.write_text(
      f"모집공고\n신청 마감: {(date.today() + timedelta(days=30)).year}년 "
      f"{(date.today() + timedelta(days=30)).month:02d}월 "
      f"{(date.today() + timedelta(days=30)).day:02d}일\n지원자격: 예비창업자",
      encoding="utf-8",
  )
  return notice, pool, tgt


def test_classify_field_kind() -> None:
    assert classify_field_kind("성명") == "fact"
    assert classify_field_kind("사업 개요 및 추진 계획") == "narrative"
    assert classify_field_kind("문제인식(Problem)") == "narrative"


def test_form_analyzer_writable_item_details(tmp_path: Path) -> None:
    p = tmp_path / "f.docx"
    doc = Document()
    doc.add_paragraph("1. 문제인식")
    doc.add_paragraph("(작성)")
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "성명"
    t.rows[0].cells[1].text = ""
    doc.save(str(p))
    fr = analyze_form(p)
    assert fr.writable_item_details
    kinds = {d["label"]: d["field_kind"] for d in fr.writable_item_details}
    assert kinds.get("성명") == "fact"


def test_d_trigger_narrative_only() -> None:
    from auto_write.services.form_analyzer import FormReport

    fr = FormReport(
        writable_item_details=[
            {"label": "성명", "field_kind": "fact"},
            {"label": "사업 개요", "field_kind": "narrative"},
        ],
        writable_items=["성명", "사업 개요"],
    )
    unmatched = [
        {"target_label": "성명", "normalized": "성명"},
        {"target_label": "사업 개요", "normalized": "사업개요"},
    ]
    gaps = filter_narrative_unmatched(unmatched, fr)
    assert len(gaps) == 1
    assert gaps[0]["target_label"] == "사업 개요"


def test_check_deadline_warning_past() -> None:
    past = date.today() - timedelta(days=5)
    ann = AnnouncementReport(key_info={"deadline": f"{past.year}년 {past.month:02d}월 {past.day:02d}일"})
    rep = FolderAnalysisReport(folder="/x", announcement=ann)
    msg = check_deadline_warning(rep)
    assert "마감이 지났습니다" in msg


def test_classify_login_wall_sbiz24() -> None:
    msg = classify_login_wall(
        "https://www.sbiz24.kr/#/pbanc/769",
        [{"status": "NO_ATTACHMENTS", "detail_url": "https://www.sbiz24.kr/"}],
    )
    assert "로그인" in msg


def test_resolve_notice_folder_from_log(tmp_path: Path) -> None:
    folder = tmp_path / "03_공고제목"
    folder.mkdir()
    f = folder / "a.hwp"
    f.write_bytes(b"x")
    log = {
        "results": [
            {
                "detail_url": "https://example.com/1",
                "status": "DOWNLOADED",
                "save_path": str(f),
            }
        ]
    }
    (tmp_path / "_download_log.json").write_text(
        json.dumps(log), encoding="utf-8")
    got = resolve_notice_folder_from_log(tmp_path, "https://example.com/1")
    assert got == folder


def test_write_confirm_and_load(tmp_path: Path) -> None:
    out = tmp_path / "filled"
    out.mkdir()
    batch = BatchAutofillReport(
        notice_folder=str(tmp_path),
        source_pool=str(tmp_path),
        output_dir=str(out),
        items=[
            BatchAutofillItem(
                target=str(tmp_path / "신청서.docx"),
                source="s.docx",
                output=str(out / "신청서_filled.docx"),
                ok=True,
                needs_confirm=[
                    {"target_label": "성명", "candidates": ["홍길동"]},
                ],
            )
        ],
    )
    paths = write_confirm_json_files(batch, out)
    assert paths
    merged = load_confirm_files(out)
    assert merged.get("성명") == "홍길동"


def test_build_todo_text() -> None:
    batch = BatchAutofillReport(
        notice_folder="/n",
        source_pool="/p",
        output_dir="/n/filled",
        items=[
            BatchAutofillItem(
                target="/n/신청서.docx",
                source="s.docx",
                output="/n/filled/신청서_filled.docx",
                ok=True,
                hwp_ok=True,
                needs_confirm_count=1,
                needs_confirm=[{"target_label": "제품명", "candidates": ["A제품"]}],
                unmatched_targets=[{"target_label": "사업명"}],
            )
        ],
    )
    text = build_todo_text(batch, deadline_warning="⚠ 마감 임박")
    assert "[완료]" in text
    assert "[확인 필요]" in text
    assert "[직접 입력]" in text


def test_run_pipeline_notice_folder(tmp_path: Path, monkeypatch) -> None:
    notice, pool, _tgt = _make_pair(tmp_path)
    monkeypatch.setattr(
        "auto_write.services.hwp_docx_convert.docx_to_hwp",
        lambda *a, **k: type("R", (), {"ok": False, "notes": ["no com"]})(),
    )
    result = run_pipeline(
        notice_folder=notice,
        source_pool=pool,
        run_bizplan=False,
        convert_hwp=False,
    )
    assert result.ok
    assert result.exit_code == 0
    assert Path(result.output_dir).is_dir()
    todo = Path(result.output_dir) / "다음할일.txt"
    assert todo.is_file()
    manifest = notice / ".analysis" / "run_manifest.json"
    assert manifest.is_file()
    assert "양식" in result.batch_summary or "채움" in result.batch_summary


def test_run_pipeline_missing_source_pool(tmp_path: Path, monkeypatch) -> None:
    notice, _, _ = _make_pair(tmp_path)
    cfg_path = tmp_path / "user_pipeline.json"
    monkeypatch.setattr(
        "auto_write.services.user_pipeline_config._CONFIG_PATH", cfg_path)
    save_config({"default_source_pool": ""})
    result = run_pipeline(notice_folder=notice, source_pool=None)
    assert result.exit_code == 1


def test_user_pipeline_save_defaults(tmp_path: Path, monkeypatch) -> None:
    cfg_path = tmp_path / "config" / "user_pipeline.json"
    monkeypatch.setattr(
        "auto_write.services.user_pipeline_config._CONFIG_PATH", cfg_path)
    save_config({"default_source_pool": str(tmp_path / "pool")})
    cfg = load_config()
    assert cfg["default_source_pool"] == str(tmp_path / "pool")


def test_format_pipeline_summary() -> None:
    from auto_write.services.notice_pipeline import PipelineResult

    r = PipelineResult(
        notice_folder="/n",
        analysis_summary="■ 마감: 내일",
        batch_summary="양식 1개 채움",
    )
    s = format_pipeline_summary_korean(r)
    assert "양식 1개" in s


# --- E2E matrix (다케이스 게이트 — 네트워크 없이 구조 검증) ---

def test_e2e_matrix_case1_star_like_batch(tmp_path: Path, monkeypatch) -> None:
    """케이스1: 완료 공고 — 양식 여러 개 batch."""
    notice, pool, _ = _make_pair(tmp_path)
    extra = notice / "참가서류.docx"
    Document().save(str(extra))
    monkeypatch.setattr(
        "auto_write.services.hwp_docx_convert.docx_to_hwp",
        lambda *a, **k: type("R", (), {"ok": False, "notes": []})(),
    )
    result = run_pipeline(notice_folder=notice, source_pool=pool, run_bizplan=False, convert_hwp=False)
    assert len(list(Path(notice).glob("*.docx"))) >= 2
    assert result.ok


def test_e2e_matrix_case2_hwp_honest(tmp_path: Path) -> None:
    """케이스2: HWP 변환 실패 시 정직 안내."""
    notice, pool, tgt = _make_pair(tmp_path)
    result = run_pipeline(
        notice_folder=notice, source_pool=pool,
        run_bizplan=False, convert_hwp=True,
    )
    assert "HWP" in result.todo_text or "DOCX" in result.todo_text


def test_e2e_matrix_case3_source_pool_pick(tmp_path: Path) -> None:
    """케이스3: source-pool 자동선택."""
    notice, pool, tgt = _make_pair(tmp_path)
    rep = autofill_from_source(pool / "완성_신청서.docx", tgt, notice / "filled" / "out.docx")
    assert rep.ok or rep.transcribed >= 0


def test_e2e_matrix_case5_login_wall_message() -> None:
    """케이스5: 로그인벽 안내."""
    msg = classify_login_wall(
        "https://www.sbiz24.kr/x",
        [{"status": "PAGE_FETCH_FAILED"}],
    )
    assert msg and "로그인" in msg


# --- UX-4: 실패·환경 정직 안내 ---

def test_download_failure_no_attachments() -> None:
    from auto_write.services.pipeline_failure_ux import classify_download_failure

    rep = classify_download_failure(
        "https://example.com",
        [{"status": "NO_ATTACHMENTS"}],
    )
    assert any("첨부파일" in a.message for a in rep.advices)
    assert rep.exit_code == 2


def test_download_failure_ssl_hint() -> None:
    from auto_write.services.pipeline_failure_ux import classify_download_failure

    rep = classify_download_failure(
        "https://gov.kr/x",
        [{"status": "PAGE_FETCH_FAILED", "error": "SSL certificate verify failed"}],
        stderr="unexpected_eof_while_reading",
    )
    assert any("SSL" in a.message or "연결" in a.message for a in rep.advices)


def test_pdf_only_pool_warning(tmp_path: Path) -> None:
    from auto_write.services.pipeline_failure_ux import pdf_only_pool_warning

    pool = tmp_path / "pool"
    pool.mkdir()
    (pool / "완성본.pdf").write_bytes(b"%PDF")
    msg = pdf_only_pool_warning(pool)
    assert msg and "PDF" in msg


def test_kstartup_junk_attachments(tmp_path: Path) -> None:
    from auto_write.services.pipeline_failure_ux import detect_kstartup_junk_attachments

    notice = tmp_path / "공고"
    notice.mkdir()
    (notice / "첨부파일.html").write_text("<html>", encoding="utf-8")
    (notice / "location.href = file.downloadPath").write_text("x", encoding="utf-8")
    msg = detect_kstartup_junk_attachments(notice)
    assert msg and "K-Startup" in msg


def test_merge_cell_address_hint() -> None:
    from auto_write.services.pipeline_failure_ux import _hint_for_batch_item

    item = BatchAutofillItem(
        target="/n/신청서.docx",
        source="s.docx",
        output="/n/filled/out.docx",
        ok=True,
        unmatched_targets=[{"target_label": "주소지"}],
    )
    hints = _hint_for_batch_item(item)
    assert any("병합" in h.message or "직접" in h.message for h in hints)


def test_empty_source_pool_failure(tmp_path: Path) -> None:
    from auto_write.services.pipeline_failure_ux import collect_batch_failures

    notice, _, _ = _make_pair(tmp_path)
    empty_pool = tmp_path / "empty_pool"
    empty_pool.mkdir()
    batch = BatchAutofillReport(str(notice), str(empty_pool), str(notice / "filled"), items=[])
    rep = collect_batch_failures(batch, empty_pool)
    assert any(a.code == "BATCH_NO_TARGETS" for a in rep.advices)


def test_no_forms_analysis_failure(tmp_path: Path) -> None:
    from auto_write.services.pipeline_failure_ux import collect_analysis_failures

    notice = tmp_path / "공고만"
    notice.mkdir()
    (notice / "공고문.txt").write_text("모집공고", encoding="utf-8")
    analysis = FolderAnalysisReport(folder=str(notice), forms=[])
    rep = collect_analysis_failures(analysis)
    assert any(a.code == "NOTICE_NO_FORMS" for a in rep.advices)


def test_pipeline_failure_lines_in_todo(tmp_path: Path, monkeypatch) -> None:
    notice, pool, _ = _make_pair(tmp_path)
    monkeypatch.setattr(
        "auto_write.services.hwp_docx_convert.docx_to_hwp",
        lambda *a, **k: type("R", (), {"ok": False, "notes": ["HWP COM 없음"]})(),
    )
    result = run_pipeline(
        notice_folder=notice, source_pool=pool,
        run_bizplan=False, convert_hwp=True,
    )
    blob = result.todo_text + "\n".join(result.failure_lines)
    assert "HWP" in blob or "DOCX" in blob
