# -*- coding: utf-8 -*-
"""quality_ratchet — 품질 기준선(ratchet) 게이트 판정·집계 테스트.

mail accuracy_baseline 이식(2026-07-17). 게이트 의미론(시딩·전진·하락 차단·
골든셋 변경 재설정)을 고정한다. 파일 IO 없는 순수 함수 위주 + 실제 DOCX
채점 통합 1건(결정론 경로).
"""
from __future__ import annotations

import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from auto_write.services.quality_ratchet import (  # noqa: E402
    DIMENSIONS, TREND_HEADER, aggregate_dimensions, build_summary,
    build_trend_row, gate,
)

_NOW = "2026-07-17T00:00:00Z"

_ALL_KEYS = [k for keys in DIMENSIONS.values() for k in keys]


def _items(score_each: float = 5.0) -> list[dict]:
    return [{"key": k, "score": score_each, "max_score": 10, "defects": 0}
            for k in _ALL_KEYS]


def _summary(*, n_docs=1, avg=90.0, doc_set=None, tests=None, run="t"):
    docs = [{"name": n, "total": avg, "passed": True,
             "dims": {"formatting": 50.0, "placement": 36.0, "image": 4.0}}
            for n in (doc_set or [f"d{i}.docx" for i in range(n_docs)])]
    return build_summary(docs, tests, run)


# ── 집계 ──────────────────────────────────────────────────────────────

def test_aggregate_dimensions_sums_by_axis():
    dims = aggregate_dimensions(_items(5.0))
    assert dims == {"formatting": 25.0, "placement": 15.0, "image": 5.0}


def test_aggregate_dimensions_unknown_key_raises():
    items = _items() + [{"key": "new_item", "score": 1.0}]
    with pytest.raises(KeyError):
        aggregate_dimensions(items)


def test_build_summary_averages_and_empty():
    s = _summary(doc_set=["a.docx", "b.docx"])
    assert s["n_docs"] == 2 and s["avg_total"] == 90.0
    assert s["dims_avg"]["formatting"] == 50.0
    empty = build_summary([], {"passed": 10, "failed": 0}, "t")
    assert empty["n_docs"] == 0 and empty["avg_total"] is None


def test_trend_row_matches_header():
    row = build_trend_row(_summary(tests={"passed": 3, "failed": 0}))
    assert len(row) == len(TREND_HEADER)


# ── 게이트 ────────────────────────────────────────────────────────────

def test_gate_seeds_without_baseline():
    r = gate(_summary(tests={"passed": 100, "failed": 0}), None, _NOW)
    assert r.exit_code == 0 and r.status.startswith("SEED")
    assert r.baseline_out["tests_passed"] == 100
    assert r.baseline_out["avg_total"] == 90.0


def test_gate_fails_on_test_failure():
    base = gate(_summary(tests={"passed": 100, "failed": 0}), None, _NOW).baseline_out
    r = gate(_summary(tests={"passed": 100, "failed": 2}), base, _NOW)
    assert r.exit_code == 2 and "tests_failed" in r.status


def test_gate_fails_on_passed_drop():
    base = gate(_summary(tests={"passed": 100, "failed": 0}), None, _NOW).baseline_out
    r = gate(_summary(tests={"passed": 90, "failed": 0}), base, _NOW)
    assert r.exit_code == 2 and "tests_passed" in r.status


def test_gate_fails_on_score_drop_same_set():
    base = gate(_summary(avg=90.0, doc_set=["a.docx"]), None, _NOW).baseline_out
    r = gate(_summary(avg=80.0, doc_set=["a.docx"]), base, _NOW)
    assert r.exit_code == 2 and "avg_total" in r.status
    # 실패 시 기준선은 후퇴하지 않는다
    assert r.baseline_out["avg_total"] == 90.0


def test_gate_advances_baseline_on_improvement():
    base = gate(_summary(tests={"passed": 100, "failed": 0}), None, _NOW).baseline_out
    r = gate(_summary(tests={"passed": 120, "failed": 0}), base, _NOW)
    assert r.exit_code == 0 and "갱신" in r.status
    assert r.baseline_out["tests_passed"] == 120


def test_gate_set_change_rebaselines_without_fail():
    base = gate(_summary(avg=90.0, doc_set=["a.docx"]), None, _NOW).baseline_out
    r = gate(_summary(avg=70.0, doc_set=["b.docx"]), base, _NOW)  # 셋 교체+점수↓
    assert r.exit_code == 0 and "재설정" in r.status
    assert r.baseline_out["doc_set"] == ["b.docx"]
    assert r.baseline_out["avg_total"] == 70.0


def test_gate_skip_tests_keeps_previous_test_baseline():
    base = gate(_summary(tests={"passed": 100, "failed": 0}), None, _NOW).baseline_out
    r = gate(_summary(tests=None), base, _NOW)  # 이번 실행은 테스트 생략
    assert r.exit_code == 0
    assert r.baseline_out["tests_passed"] == 100


def test_gate_seed_flag_overwrites_baseline():
    base = gate(_summary(tests={"passed": 100, "failed": 0}), None, _NOW).baseline_out
    r = gate(_summary(tests={"passed": 50, "failed": 0}), base, _NOW, seed=True)
    assert r.exit_code == 0 and r.status.startswith("SEED")
    assert r.baseline_out["tests_passed"] == 50


# ── 통합(결정론 채점 경로) ────────────────────────────────────────────

def test_measure_docx_real_document(tmp_path):
    from docx import Document as NewDocument

    from quality_ratchet import measure_docx

    doc = NewDocument()
    doc.add_heading("사업 개요", level=1)
    doc.add_paragraph("문제 인식과 해결 방안, 시장 분석 및 성장 전략을 기술한다.")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "항목"
    t.cell(0, 1).text = "내용"
    p = tmp_path / "sample.docx"
    doc.save(str(p))

    r = measure_docx(p)
    assert set(r["dims"]) == set(DIMENSIONS)
    assert 0 <= r["total"] <= 100
    # 3축 합 == 총점 (9항목 전부가 정확히 한 축에 속함)
    assert abs(sum(r["dims"].values()) - r["total"]) < 0.2
    # 동일 입력 → 동일 점수 (결정론)
    assert measure_docx(p)["total"] == r["total"]
