"""test_cross_form_vertical.py — 세로형(라벨 위 / 값 아래) 표 카드 전사 회귀 테스트.

정부양식 표지·개요 박스는 라벨을 한 줄에 나열하고(성명·생년월일·연락처·이메일)
그 **아래 줄**에 값을 적는 '정보 카드' 구조가 흔하다. 기존 cross-form 자동채우기는
가로형(라벨|값 옆칸)·본문/셀 인라인만 채워 이 세로형 카드는 **한 칸도 못 채웠다**.
이 테스트는 세로형 카드 탐지·전사가 실제로 동작하고, 로스터/데이터 행/병합 어긋남에는
오작동하지 않으며(오매칭<빈칸·날조0·원본 미수정), 기존 가로형은 그대로임을 증명한다.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from auto_write.services.cross_form_autofill import (
    autofill_from_source,
    find_target_fields,
)


# --- 픽스처 헬퍼 -------------------------------------------------------------

def _make_person_source(path: Path) -> None:
    """소스: 라벨|값 2열 표(대표자=성명 동의어 포함)."""
    doc = Document()
    doc.add_heading("사업계획서", 0)
    t = doc.add_table(rows=4, cols=2)
    pairs = [
        ("대표자", "홍길동"),          # 성명 ← 대표자 (동의어)
        ("생년월일", "1990.05.15"),    # 정상 날짜(플레이스홀더 아님)
        ("연락처", "010-1234-5678"),
        ("이메일", "hong@test.com"),
    ]
    for i, (lab, val) in enumerate(pairs):
        t.rows[i].cells[0].text = lab
        t.rows[i].cells[1].text = val
    doc.save(str(path))


def _make_vertical_card_target(path: Path) -> None:
    """타깃: 세로형 카드 — 헤더 행(라벨 4개) + 아래 행(전부 빈칸)."""
    doc = Document()
    doc.add_heading("지원신청서", 0)
    t = doc.add_table(rows=2, cols=4)
    for c, lab in enumerate(["성명", "생년월일", "연락처", "이메일"]):
        t.rows[0].cells[c].text = lab
        t.rows[1].cells[c].text = ""  # 값 대기 칸
    doc.save(str(path))


def _below(path: Path, col: int, *, ti: int = 0, row: int = 1) -> str:
    doc = Document(str(path))
    return doc.tables[ti].rows[row].cells[col].text.strip()


# --- 탐지 ------------------------------------------------------------------

def test_vertical_card_targets_point_to_below_row(tmp_path: Path) -> None:
    """세로형 카드에서 각 헤더 라벨이 '아래 행'을 값칸으로 하는 타깃이 된다."""
    tgt = tmp_path / "card.docx"
    _make_vertical_card_target(tgt)
    targets = find_target_fields(tgt)
    by_label = {t["normalized"]: t for t in targets}
    # 4개 라벨 전부 타깃으로 잡혀야 한다.
    for lab in ("성명", "생년월일", "연락처", "이메일"):
        assert lab in by_label, f"{lab} 세로형 타깃 미탐지"
        t = by_label[lab]
        assert t["kind"] == "table"
        assert t["row"] == 1          # 값칸은 헤더 '아래' 행
        assert t["table_index"] == 0


# --- 전사(end-to-end) --------------------------------------------------------

def test_vertical_card_filled_end_to_end(tmp_path: Path) -> None:
    """전에는 0칸, 이제는 세로형 카드 아래 행에 값이 실제로 채워진다(동의어 포함)."""
    src = tmp_path / "src.docx"
    tgt = tmp_path / "tgt.docx"
    out = tmp_path / "out.docx"
    _make_person_source(src)
    _make_vertical_card_target(tgt)

    rep = autofill_from_source(src, tgt, out)

    assert rep.ok
    assert rep.transcribed == 4
    assert _below(out, 0) == "홍길동"          # 성명 ← 대표자(동의어)
    assert _below(out, 1) == "1990.05.15"      # 생년월일
    assert _below(out, 2) == "010-1234-5678"   # 연락처
    assert _below(out, 3) == "hong@test.com"   # 이메일


def test_vertical_fabrication_zero(tmp_path: Path) -> None:
    """소스에 값이 없는 헤더 칸은 빈칸으로 남는다(날조 0)."""
    src = tmp_path / "src.docx"
    tgt = tmp_path / "tgt.docx"
    out = tmp_path / "out.docx"

    # 소스에 '생년월일' 없음 → 그 칸은 채우지 않아야 한다.
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "대표자"
    t.rows[0].cells[1].text = "홍길동"
    t.rows[1].cells[0].text = "연락처"
    t.rows[1].cells[1].text = "010-1234-5678"
    doc.save(str(src))
    _make_vertical_card_target(tgt)

    rep = autofill_from_source(src, tgt, out)
    assert _below(out, 0) == "홍길동"
    assert _below(out, 1) == ""            # 생년월일 = 소스 없음 → 빈칸(날조 0)
    assert _below(out, 2) == "010-1234-5678"
    assert _below(out, 3) == ""            # 이메일 = 소스 없음 → 빈칸


# --- 안전 가드: 로스터/데이터 행/병합 ------------------------------------------

def test_vertical_roster_not_mass_filled(tmp_path: Path) -> None:
    """헤더 + 여러 빈 행(명부/로스터)은 세로 채움 금지 — 첫 행에 일괄기입하지 않는다."""
    src = tmp_path / "src.docx"
    tgt = tmp_path / "tgt.docx"
    out = tmp_path / "out.docx"
    _make_person_source(src)

    doc = Document()
    t = doc.add_table(rows=4, cols=4)  # 헤더 + 빈 행 3개 = 다중 레코드 표
    for c, lab in enumerate(["성명", "생년월일", "연락처", "이메일"]):
        t.rows[0].cells[c].text = lab
    for r in range(1, 4):
        for c in range(4):
            t.rows[r].cells[c].text = ""
    doc.save(str(tgt))

    # 세로형 타깃이 하나도 생기지 않아야 한다(로스터 보호).
    targets = find_target_fields(tgt)
    assert all(t["row"] != 1 for t in targets), "로스터 첫 행이 세로 타깃으로 오인됨"

    rep = autofill_from_source(src, tgt, out)
    # 어떤 데이터 행에도 값이 새로 박히면 안 된다(로스터 일괄기입 방지).
    for r in range(1, 4):
        for c in range(4):
            assert _below(out, c, row=r) == "", f"로스터 행{r} 열{c} 오기입"


def test_vertical_data_row_preserved(tmp_path: Path) -> None:
    """헤더 아래 행에 이미 값이 있으면 세로 카드로 보지 않고 실값을 보존한다."""
    src = tmp_path / "src.docx"
    tgt = tmp_path / "tgt.docx"
    out = tmp_path / "out.docx"
    _make_person_source(src)

    doc = Document()
    t = doc.add_table(rows=2, cols=4)
    for c, lab in enumerate(["성명", "생년월일", "연락처", "이메일"]):
        t.rows[0].cells[c].text = lab
    existing = ["김철수", "1985.01.01", "010-0000-0000", "kim@x.com"]
    for c, val in enumerate(existing):
        t.rows[1].cells[c].text = val
    doc.save(str(tgt))

    rep = autofill_from_source(src, tgt, out)
    for c, val in enumerate(existing):
        assert _below(out, c) == val, f"열{c} 실값이 덮어써짐(덮어쓰기 금지 위반)"


def test_vertical_merge_mismatch_skipped(tmp_path: Path) -> None:
    """위·아래 논리셀 수가 병합으로 어긋나면 세로 정렬 불가 → 제외(크래시 없음)."""
    src = tmp_path / "src.docx"
    tgt = tmp_path / "tgt.docx"
    out = tmp_path / "out.docx"
    _make_person_source(src)

    doc = Document()
    t = doc.add_table(rows=2, cols=3)
    for c, lab in enumerate(["성명", "생년월일", "연락처"]):
        t.rows[0].cells[c].text = lab
    for c in range(3):
        t.rows[1].cells[c].text = ""
    # 아래 행의 두 칸을 병합 → 논리셀 수 3(위) != 2(아래)
    t.rows[1].cells[0].merge(t.rows[1].cells[1])
    doc.save(str(tgt))

    targets = find_target_fields(tgt)          # 크래시 없이
    assert all(t["row"] != 1 for t in targets), "병합 어긋남인데 세로 타깃 생성됨"
    rep = autofill_from_source(src, tgt, out)  # 예외 없이 완료
    assert rep is not None


# --- 회귀: 기존 가로형 유지 --------------------------------------------------

def test_horizontal_card_still_works_and_no_double(tmp_path: Path) -> None:
    """가로형(라벨|값 옆칸) 카드는 그대로 1회만 채워진다(세로 로직이 방해하지 않음)."""
    src = tmp_path / "src.docx"
    tgt = tmp_path / "tgt.docx"
    out = tmp_path / "out.docx"
    _make_person_source(src)

    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "성명"
    t.rows[0].cells[1].text = ""
    t.rows[1].cells[0].text = "연락처"
    t.rows[1].cells[1].text = ""
    doc.save(str(tgt))

    rep = autofill_from_source(src, tgt, out)
    assert rep.transcribed == 2               # 성명·연락처 각 1회(중복 전사 없음)
    doc_out = Document(str(out))
    assert doc_out.tables[0].rows[0].cells[1].text.strip() == "홍길동"
    assert doc_out.tables[0].rows[1].cells[1].text.strip() == "010-1234-5678"


def test_originals_untouched(tmp_path: Path) -> None:
    """소스·타깃 원본 파일은 바이트 그대로(원본 미수정 불변)."""
    src = tmp_path / "src.docx"
    tgt = tmp_path / "tgt.docx"
    out = tmp_path / "out.docx"
    _make_person_source(src)
    _make_vertical_card_target(tgt)
    src_bytes = src.read_bytes()
    tgt_bytes = tgt.read_bytes()

    autofill_from_source(src, tgt, out)

    assert src.read_bytes() == src_bytes
    assert tgt.read_bytes() == tgt_bytes
