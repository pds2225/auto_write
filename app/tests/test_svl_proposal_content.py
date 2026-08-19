"""소셜벤처 리그 제안서 본문 불변: 과제④, 날조 팀원 금지, 지정 양식 마커."""
from __future__ import annotations

import importlib.util
from pathlib import Path

TOOL = Path(__file__).resolve().parents[1] / "tools" / "build_svl_proposal.py"


def _load():
    spec = importlib.util.spec_from_file_location("build_svl_proposal", TOOL)
    mod = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(mod)
    return mod


def test_task4_checked_and_others_unchecked():
    mod = _load()
    assert "■4" in mod.TASK4_CELL
    assert mod.TASK4_CELL.count("■") == 1
    for n in ("□1", "□2", "□3", "□5"):
        assert n in mod.TASK4_CELL


def test_no_fabricated_team_or_unverified_stats():
    mod = _load()
    blob = "\n".join(
        [mod.BACKGROUND, mod.PROBLEM, mod.EXISTING, mod.SOLUTION, mod.IMPACT]
    )
    for banned in ("김지훈", "서울대학교 경영학과", "N=120", "성공률 82%"):
        assert banned not in blob
    assert "10-2026-0026207" in blob
    assert "과제④" in blob
    assert "예비창업" in blob or "대표 1인" in blob


def test_docx_builds_and_keeps_facts(tmp_path):
    mod = _load()
    out = tmp_path / "svl.docx"
    mod.build_docx(out)
    assert out.is_file() and out.stat().st_size > 1000
    from docx import Document

    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    assert "박다솜" in text
    assert "MarketGate" in text
    assert "10-2026-0026207" in text
    assert "김지훈" not in text
