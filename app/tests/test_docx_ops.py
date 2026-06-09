from __future__ import annotations

import unittest

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from auto_write.services.docx_ops import set_cell_text


class DocxOpsTests(unittest.TestCase):
    def test_set_cell_text_clears_run_highlight_from_placeholder_text(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run("※ 작성요령")
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        set_cell_text(cell, "완료 문장")

        highlights = cell._tc.findall(".//" + qn("w:highlight"))
        self.assertEqual(len(highlights), 0)
        self.assertIn("완료 문장", cell.text)

    def test_set_cell_text_removes_placeholder_cell_shading(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        tc = cell._tc
        tc_pr = tc.find(qn("w:tcPr"))
        if tc_pr is None:
            tc_pr = OxmlElement("w:tcPr")
            tc.insert(0, tc_pr)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "FFFF00")
        tc_pr.append(shading)

        set_cell_text(cell, "채워진 값")

        tc_pr_after = cell._tc.find(qn("w:tcPr"))
        self.assertIsNotNone(tc_pr_after)
        self.assertIsNone(tc_pr_after.find(qn("w:shd")))
        self.assertIn("채워진 값", cell.text)

    # ---- STEP4: 텍스트 맞는 자리(가로 병합셀 좌표 정합, item 4) ----
    def test_logical_cells_dedups_horizontal_merge(self):
        from auto_write.services.docx_ops import logical_cells
        doc = Document()
        table = doc.add_table(rows=1, cols=3)
        table.cell(0, 0).merge(table.cell(0, 1))
        row = table.rows[0]
        self.assertEqual(len(row.cells), 3)              # grid 는 병합셀을 중복 노출
        logical = logical_cells(row)
        self.assertEqual(len(logical), 2)                # 논리 셀은 2개
        self.assertIs(logical[1]._tc, row.cells[2]._tc)  # 논리 1 = 병합 안 된 세 번째 grid 셀

    def test_resolve_table_cell_text_uses_logical_index(self):
        from auto_write.services.render_service import RenderService
        rs = RenderService()
        doc = Document()
        table = doc.add_table(rows=1, cols=3)
        table.cell(0, 0).merge(table.cell(0, 1))
        row = table.rows[0]
        # 텍스트 경로(logical=True): 논리 인덱스 1 = 세 번째 grid 셀(분석이 기록한 올바른 자리)
        cell_logical, err1 = rs._resolve_table_cell(doc, 0, 0, 1, logical=True)
        self.assertEqual(err1, "")
        self.assertIs(cell_logical._tc, row.cells[2]._tc)
        # 이미지 경로(기본 grid): 인덱스 1 = 병합 셀(다른 자리)
        cell_grid, err2 = rs._resolve_table_cell(doc, 0, 0, 1)
        self.assertEqual(err2, "")
        self.assertIs(cell_grid._tc, row.cells[1]._tc)
        # 병합 표에서 두 해석이 달라짐(이전엔 텍스트가 grid 로 해석돼 엉뚱한 셀로 갔던 지점)
        self.assertIsNot(cell_logical._tc, cell_grid._tc)

    def test_resolve_table_cell_logical_out_of_range_reports_error(self):
        from auto_write.services.render_service import RenderService
        rs = RenderService()
        doc = Document()
        table = doc.add_table(rows=1, cols=3)
        table.cell(0, 0).merge(table.cell(0, 1))         # 논리 셀 2개
        # 논리 인덱스 2 는 범위 밖 → 조용히 덮어쓰지 않고 오류로 보고
        cell, reason = rs._resolve_table_cell(doc, 0, 0, 2, logical=True)
        self.assertIsNone(cell)
        self.assertIn("열 index", reason)


if __name__ == "__main__":
    unittest.main()

