"""SFT 데이터 레이어 P0 테스트 — AI 호출 trace 저장 + 훅 fail-safe + provenance E2E.

- generation_store: record_ai_call append/load 왕복, blob dedup, 빈 본문, resolve.
- openai_client 훅: 성공 시 trace 기록(purpose/project_id/attempt 태깅),
  로깅 실패가 AI 호출을 절대 깨지 않음(fail-safe, 적대검증 HIGH).
- project_service E2E: generate() 가 무키(폴백) 경로에서도 크래시 없이
  input_before_generation / answers_provenance / ai_draft_snapshot 를 남긴다.
"""

from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from auto_write.config import Settings, ensure_directories
from auto_write.services import generation_store, learning_store
from auto_write.services.evidence_service import EvidenceService
from auto_write.services.image_service import ImageService
from auto_write.services.openai_client import OpenAIService
from auto_write.services.project_service import ProjectService
from auto_write.services.qa_service import QAService
from auto_write.services.render_service import RenderService
from auto_write.storage import Storage


# --- generation_store 단위 테스트 (root 격리) ---------------------------------

def test_record_ai_call_append_and_load_roundtrip(tmp_path: Path) -> None:
    rec = generation_store.record_ai_call(
        provider="anthropic",
        model="claude-x",
        system_prompt="sys",
        user_prompt="질문+컨텍스트",
        raw_response='{"q1": "답"}',
        purpose="draft_answers",
        project_id="prj_1",
        attempt=1,
        duration_ms=12,
        root=tmp_path,
    )
    traces = learning_store.load_generation_traces(root=tmp_path)
    assert len(traces) == 1
    assert traces[0]["trace_id"] == rec["trace_id"]
    assert traces[0]["purpose"] == "draft_answers"
    assert traces[0]["project_id"] == "prj_1"
    assert traces[0]["empty"] is False
    bodies = generation_store.resolve_trace_bodies(traces[0], root=tmp_path)
    assert bodies["user_prompt"] == "질문+컨텍스트"
    assert bodies["raw_response"] == '{"q1": "답"}'
    assert bodies["system_prompt"] == "sys"


def test_blob_dedup_same_text_one_file(tmp_path: Path) -> None:
    d1 = generation_store.store_blob("동일한 큰 컨텍스트", root=tmp_path)
    d2 = generation_store.store_blob("동일한 큰 컨텍스트", root=tmp_path)
    assert d1 == d2 and d1 != ""
    blob_files = list((tmp_path / "gen_blobs").glob("*.txt"))
    assert len(blob_files) == 1


def test_empty_response_marks_empty_and_blank_ref(tmp_path: Path) -> None:
    rec = generation_store.record_ai_call(
        provider="anthropic", model="m", system_prompt="s",
        user_prompt="u", raw_response="", root=tmp_path,
    )
    assert rec["empty"] is True
    assert rec["response_ref"] == ""


# --- openai_client 훅: 성공 기록 ---------------------------------------------

def _settings(tmp_path: Path, *, anthropic_key: str = "sk-test") -> Settings:
    return Settings(
        app_root=tmp_path / "app",
        workspace_root=tmp_path / "workspace",
        template_root=tmp_path / "workspace" / "templates",
        project_root=tmp_path / "workspace" / "projects",
        results_root=tmp_path / "results",
        static_root=tmp_path / "static",
        template_view_root=tmp_path / "views",
        host="127.0.0.1",
        port=8765,
        openai_api_key="",
        openai_model="gpt-4.1-mini",
        openai_search_model="gpt-4.1-mini",
        openai_image_model="gpt-image-1",
        anthropic_api_key=anthropic_key,
        anthropic_model="claude-sonnet-4-20250514",
        anthropic_search_model="claude-sonnet-4-20250514",
    )


def test_complete_text_records_trace_on_success(tmp_path: Path, monkeypatch) -> None:
    svc = OpenAIService(_settings(tmp_path))
    monkeypatch.setattr(svc, "_complete_text_anthropic", lambda *a, **k: '{"q1":"답"}')
    captured: list[dict] = []
    monkeypatch.setattr(
        generation_store, "record_ai_call",
        lambda **kw: captured.append(kw) or {"trace_id": "t"},
    )
    out = svc._complete_text(
        "sys", "usr", provider_override="anthropic",
        log_meta={"project_id": "prj_9", "purpose": "draft_answers", "attempt": 1},
    )
    assert out == '{"q1":"답"}'
    assert len(captured) == 1
    assert captured[0]["purpose"] == "draft_answers"
    assert captured[0]["project_id"] == "prj_9"
    assert captured[0]["provider"] == "anthropic"
    assert captured[0]["raw_response"] == '{"q1":"답"}'


def test_logging_failure_never_breaks_ai_call(tmp_path: Path, monkeypatch) -> None:
    """적대검증 HIGH: 로깅 예외가 AI 호출 실패로 승격되면 안 된다."""
    svc = OpenAIService(_settings(tmp_path))
    monkeypatch.setattr(svc, "_complete_text_anthropic", lambda *a, **k: "정상응답")

    def _boom(**kw):
        raise OSError("디스크 풀 시뮬레이션")

    monkeypatch.setattr(generation_store, "record_ai_call", _boom)
    out = svc._complete_text("sys", "usr", provider_override="anthropic")
    assert out == "정상응답"  # 예외를 삼키고 응답을 그대로 반환


def test_draft_missing_answers_tags_retry_attempt(tmp_path: Path, monkeypatch) -> None:
    """anthropic 작성모델 실패 후 기본모델 재시도 시 attempt=2 로 태깅된다."""
    svc = OpenAIService(_settings(tmp_path))
    calls: list[dict] = []
    # complete_json 실패(=None)를 강제해 재시도 경로로 유도 → 하지만 trace 는 남는다.
    monkeypatch.setattr(svc, "_complete_text_anthropic", lambda *a, **k: "")  # 빈 응답 → 파싱 실패
    monkeypatch.setattr(
        generation_store, "record_ai_call",
        lambda **kw: calls.append(kw) or {"trace_id": "t"},
    )
    svc.draft_missing_answers(
        [{"question_id": "q1"}], "컨텍스트",
        provider_override="anthropic",
        model_override="claude-writing-x",
        log_meta={"project_id": "prj_r"},
    )
    attempts = sorted(c["attempt"] for c in calls)
    assert attempts == [1, 2]
    assert all(c["purpose"] == "draft_answers" for c in calls)
    assert all(c["project_id"] == "prj_r" for c in calls)


# --- project_service E2E: 무키 폴백 경로에서도 스냅샷·provenance 생성 ----------

def _build_psst_sample_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("□ 일반현황")
    doc.add_paragraph("1. 문제 인식 (Problem)_창업 아이템의 필요성")
    doc.add_paragraph("2. 실현 가능성 (Solution)_창업 아이템의 개발 계획")
    doc.add_paragraph("3. 성장전략 (Scale-up)_사업화 추진 전략")
    doc.add_paragraph("4. 팀 구성 (Team)_대표자 및 팀원 구성 계획")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "항목"
    table.cell(0, 1).text = "내용"
    table.cell(1, 0).text = "기업명"
    table.cell(1, 1).text = "○○기업"
    doc.save(path)


class GenerateProvenanceE2ETests(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp = tempfile.TemporaryDirectory()
        root = Path(self.tmp.name)
        self.settings = Settings(
            app_root=root / "app",
            workspace_root=root / "workspace",
            template_root=root / "workspace" / "templates",
            project_root=root / "workspace" / "projects",
            results_root=root / "results",
            static_root=root / "app" / "auto_write" / "static",
            template_view_root=root / "app" / "auto_write" / "templates",
            host="127.0.0.1",
            port=8765,
            openai_api_key="",  # 무키 → AI 폴백 경로
            openai_model="gpt-4.1-mini",
            openai_search_model="gpt-4.1-mini",
            openai_image_model="gpt-image-1",
            anthropic_api_key="",
            anthropic_model="claude-sonnet-4-20250514",
            anthropic_search_model="claude-sonnet-4-20250514",
        )
        ensure_directories(self.settings)
        self.storage = Storage(self.settings)
        openai_service = OpenAIService(self.settings)
        self.service = ProjectService(
            storage=self.storage,
            openai_service=openai_service,
            evidence_service=EvidenceService(openai_service),
            image_service=ImageService(openai_service),
            render_service=RenderService(),
            qa_service=QAService(),
        )

    def tearDown(self) -> None:
        self.tmp.cleanup()

    def test_generate_writes_sft_snapshots_without_crash(self) -> None:
        sample = Path(self.tmp.name) / "psst_sample.docx"
        _build_psst_sample_docx(sample)
        profile = self.service.analyze_uploaded_template("psst_sample.docx", sample.read_bytes())
        project_id = self.service.create_project(profile.template_id, "P0 검증")
        self.service.save_project_form(
            project_id=project_id,
            answers={"user_brief": "P0 개요: AI 안전 제어", "user_notes": "해결\n\n성장\n\n팀"},
            project_title="P0 테스트 과제",
            organization_name="P0 테스트 기업",
            evidence_topics="",
            reference_files=[],
            improve_partial=True,
            psst_only=True,
            disable_images=True,
        )
        # 무키 폴백 경로 — 크래시 없이 완주해야 한다.
        artifacts = self.service.generate(project_id)
        self.assertTrue(Path(artifacts.output_docx).exists())

        sft_dir = self.storage.project_dir(project_id) / "sft"
        snap = sft_dir / "input_before_generation.json"
        prov = sft_dir / "answers_provenance.json"
        draft = sft_dir / "ai_draft_snapshot.json"
        self.assertTrue(snap.exists(), "입력 스냅샷 미생성")
        self.assertTrue(prov.exists(), "provenance 미생성")
        self.assertTrue(draft.exists(), "ai_draft_snapshot 미생성")

        prov_data = json.loads(prov.read_text(encoding="utf-8"))["provenance"]
        # 무키 → AI source 는 없어야 하고, 사용자/폴백 등 비-AI source 로 채워진다(오라벨 방지).
        sources = {v["source"] for v in prov_data.values()}
        self.assertNotIn("ai", sources)
        self.assertTrue(sources, "provenance 가 비어있음")
        # 사용자 폼 입력(user_brief)은 source=="user" 로 정확히 라벨링돼야 한다.
        self.assertEqual(prov_data.get("user_brief", {}).get("source"), "user")
        # 입력 스냅샷은 AI 변형 전 answers 를 담는다(user_brief 포함).
        snap_data = json.loads(snap.read_text(encoding="utf-8"))
        self.assertIn("user_brief", snap_data.get("answers", {}))


if __name__ == "__main__":
    unittest.main()
