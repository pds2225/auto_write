"""test_conversion_fidelity.py — 변환 일치도 측정 하네스 테스트.

comparator(compare_docx_structure)는 COM 무관하게 합성 DOCX 픽스처로 단위검증한다:
동일=100, 표/단락 인위 누락=감점+lost_items. roundtrip 은 COM 게이트만(monkeypatch)
검증한다 — COM 미가용 시 ok=False + 안내 + 예외 전파 X.
"""

from __future__ import annotations

import io
import struct
import zlib
from pathlib import Path

import pytest
from docx import Document

from auto_write.services import conversion_fidelity as mod
from auto_write.services.conversion_fidelity import (
    FidelityReport,
    compare_docx_structure,
    measure_roundtrip_fidelity,
)


# --- 픽스처 도우미 ------------------------------------------------------------

def _png_1x1() -> bytes:
    """유효한 1x1 PNG 바이트(python-docx add_picture 가 헤더를 읽음)."""
    def chunk(tag: bytes, data: bytes) -> bytes:
        body = tag + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    raw = b"\x00\xff\xff\xff"  # 1 scanline: filter byte + RGB
    idat = zlib.compress(raw)
    return sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def _make_doc(path: Path, *, paragraphs: int = 2, table: bool = True, image: bool = False) -> None:
    doc = Document()
    for i in range(paragraphs):
        doc.add_paragraph(f"본문 단락 {i} 핵심 전략")
    if table:
        t = doc.add_table(rows=2, cols=2)
        t.cell(0, 0).text = "항목"
        t.cell(0, 1).text = "수치"
        t.cell(1, 0).text = "매출"
        t.cell(1, 1).text = "100"
    if image:
        doc.add_picture(io.BytesIO(_png_1x1()))
    doc.save(str(path))


# --- comparator: 동일 문서 ----------------------------------------------------

def test_identical_docs_score_100(tmp_path: Path) -> None:
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    _make_doc(a)
    _make_doc(b)  # 동일 내용 2부
    r = compare_docx_structure(str(a), str(b))
    assert isinstance(r, FidelityReport)
    assert r.ok is True
    assert r.overall_score == 100.0
    assert all(v == 100.0 for v in r.metrics.values())
    assert r.lost_items == []


# --- comparator: 표 인위 제거 -------------------------------------------------

def test_missing_table_penalizes(tmp_path: Path) -> None:
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    _make_doc(a, table=True)
    _make_doc(b, table=False)  # b 에서 표 제거
    r = compare_docx_structure(str(a), str(b))
    assert r.metrics["tables"] < 100.0
    assert r.metrics["cells"] < 100.0
    assert r.overall_score < 100.0
    assert r.lost_items  # 비어있지 않음
    assert any("표" in item for item in r.lost_items)


# --- comparator: 단락 누락 ----------------------------------------------------

def test_missing_paragraphs_penalizes(tmp_path: Path) -> None:
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    _make_doc(a, paragraphs=4, table=False)
    _make_doc(b, paragraphs=1, table=False)  # 단락 누락
    r = compare_docx_structure(str(a), str(b))
    assert r.metrics["paragraphs"] < 100.0
    assert r.overall_score < 100.0
    assert any("단락" in item for item in r.lost_items)


# --- comparator: 이미지 카운트 ------------------------------------------------

def test_image_count_metric(tmp_path: Path) -> None:
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    _make_doc(a, table=False, image=True)
    _make_doc(b, table=False, image=False)  # 이미지 없음
    r = compare_docx_structure(str(a), str(b))
    assert r.counts["a"]["images"] == 1
    assert r.counts["b"]["images"] == 0
    assert r.metrics["images"] < 100.0
    assert any("이미지" in item for item in r.lost_items)


# --- comparator: 구조 한계 고지 -----------------------------------------------

def test_report_notes_structural_only(tmp_path: Path) -> None:
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    _make_doc(a)
    _make_doc(b)
    r = compare_docx_structure(str(a), str(b))
    assert any("구조" in n and "시각" in n for n in r.notes)  # 구조≠시각 명시


# --- roundtrip: COM 게이트 ----------------------------------------------------

def test_roundtrip_com_unavailable_gates_off(tmp_path: Path, monkeypatch) -> None:
    src = tmp_path / "plan.docx"
    _make_doc(src)
    monkeypatch.setattr(
        "auto_write.services.hwp_docx_convert.hancom_com_available", lambda: False)
    r = measure_roundtrip_fidelity(str(src), use_com=True)
    assert r.ok is False               # 측정 불가
    assert r.method == "roundtrip"
    assert any("COM" in n and "측정 불가" in n for n in r.notes)  # 안내
    assert r.metrics == {}             # 측정 생략


def test_roundtrip_no_com_flag_gates_off(tmp_path: Path) -> None:
    """--no-com(use_com=False)도 게이트 — COM 등록 여부와 무관히 측정 생략."""
    src = tmp_path / "plan.docx"
    _make_doc(src)
    r = measure_roundtrip_fidelity(str(src), use_com=False)
    assert r.ok is False
    assert any("COM" in n for n in r.notes)


def test_roundtrip_missing_input_raises(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        measure_roundtrip_fidelity(str(tmp_path / "없음.docx"))


# --- as_dict 계약 -------------------------------------------------------------

def test_as_dict_shape(tmp_path: Path) -> None:
    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    _make_doc(a)
    _make_doc(b)
    d = compare_docx_structure(str(a), str(b)).as_dict()
    for key in ("ok", "overall_score", "metrics", "counts", "lost_items", "method", "notes"):
        assert key in d
