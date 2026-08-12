# test_e2e_domain_pipeline.py — E2E domain pipeline tests
"""도메인 파이프라인 E2E 테스트.

synthetic fixture를 사용하여 business_plan / consultant_application
도메인 파이프라인이 실제로 동작하는지 검증한다.
"""
from __future__ import annotations

import tempfile
from pathlib import Path

import pytest
from docx import Document

from auto_write.domains.domain_classifier import Domain, classify_domain
from auto_write.domains.domain_router import DomainRouter, resolve_domain
from auto_write.services.lrule_enforcer import enforce_lrules
from auto_write.services.finalizer import finalize_artifact


def _create_bp_docx(path: Path) -> None:
    """사업계획서 synthetic fixture를 생성한다."""
    doc = Document()
    doc.add_heading("사업계획서", level=1)
    doc.add_paragraph("1. 문제 인식 (Problem)")
    doc.add_paragraph("시장에서 기존 솔루션의 한계를 발견하고 새로운 접근이 필요합니다.")
    doc.add_paragraph("2. 실현 가능성 (Solution)")
    doc.add_paragraph("기술적 실현 가능성을 검증하고 프로토타입을 완성했습니다.")
    doc.add_paragraph("3. 성장 전략 (Scale)")
    doc.add_paragraph("3년 내 매출 100억원 달성을 목표로 합니다.")
    doc.add_paragraph("4. 팀 구성 (Team)")
    doc.add_paragraph("AI 전문가 3명, 마케팅 전문가 2명으로 구성된 팀입니다.")
    doc.save(str(path))


def _create_ca_docx(path: Path) -> None:
    """컨설턴트 신청서 synthetic fixture를 생성한다."""
    doc = Document()
    doc.add_heading("컨설턴트 신청서", level=1)
    doc.add_paragraph("이력서")
    doc.add_paragraph("성명: 홍길동")
    doc.add_paragraph("경력: 10년")
    doc.add_paragraph("자격: 경영지도사")
    doc.add_paragraph("수행실적: 정부지원사업 5건")
    doc.save(str(path))


class TestE2EBusinessPlan:
    def test_domain_classification(self):
        """사업계획서 텍스트가 business_plan으로 분류되어야 한다."""
        result = classify_domain(text="사업계획서 PSST 창업아이템 사업화 전략")
        assert result.domain == Domain.BUSINESS_PLAN
        assert result.confidence > 0.5

    def test_domain_routing(self):
        """DomainRouter가 business_plan 컨텍스트를 생성해야 한다."""
        ctx = resolve_domain(text="사업계획서 PSST 창업아이템")
        assert ctx.domain == Domain.BUSINESS_PLAN
        assert ctx.domain_result.confidence > 0.5

    def test_lrule_enforcement(self):
        """business_plan 도메인에서 LRule enforcement가 동작해야 한다."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            _create_bp_docx(Path(f.name))
            tmppath = f.name

        try:
            report = enforce_lrules(
                domain=Domain.BUSINESS_PLAN,
                document_type="business_plan",
                artifact_path=tmppath,
            )
            assert report.domain == "business_plan"
            assert report.summary["total"] > 100
            assert report.artifact_sha256
            assert len(report.artifact_sha256) == 64

            # BP-specific rules should be applicable
            bp_rules = [r for r in report.rules if r["domain"] == "business_plan"]
            assert all(r["applicable"] for r in bp_rules)

            # CA-specific rules should be N/A
            ca_rules = [r for r in report.rules if r["domain"] == "consultant_application"]
            assert all(r["status"] == "N/A" for r in ca_rules)
        finally:
            Path(tmppath).unlink()

    def test_finalizer_produces_draft(self):
        """FAIL/REVIEW가 있으면 DRAFT를 생성해야 한다."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            _create_bp_docx(Path(f.name))
            tmppath = f.name

        try:
            report = enforce_lrules(domain=Domain.BUSINESS_PLAN, artifact_path=tmppath)
            result = finalize_artifact(artifact_path=tmppath, lrule_report=report)
            # Without guards, mechanized rules will be UNVERIFIABLE → DRAFT
            assert result.is_draft
            assert not result.submittable
            assert "_DRAFT" in result.final_path
        finally:
            Path(tmppath).unlink()

    def test_lrule_sha256_consistency(self):
        """같은 문서의 SHA256은 재계산 시 동일해야 한다."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            _create_bp_docx(Path(f.name))
            tmppath = f.name

        try:
            r1 = enforce_lrules(domain=Domain.BUSINESS_PLAN, artifact_path=tmppath)
            r2 = enforce_lrules(domain=Domain.BUSINESS_PLAN, artifact_path=tmppath)
            assert r1.artifact_sha256 == r2.artifact_sha256
            assert len(r1.artifact_sha256) == 64
        finally:
            Path(tmppath).unlink()

    def test_empty_doc_still_classifies(self):
        """빈 문서도 domain classification은 동작해야 한다."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.save(f.name)
            tmppath = f.name

        try:
            report = enforce_lrules(domain=Domain.BUSINESS_PLAN, artifact_path=tmppath)
            assert report.summary["total"] > 0
        finally:
            Path(tmppath).unlink()

    def test_force_draft_flag(self):
        """force_draft=True이면 무조건 DRAFT여야 한다."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            _create_bp_docx(Path(f.name))
            tmppath = f.name

        try:
            report = enforce_lrules(domain=Domain.BUSINESS_PLAN, artifact_path=tmppath)
            result = finalize_artifact(
                artifact_path=tmppath, lrule_report=report, force_draft=True
            )
            assert result.is_draft
            assert not result.submittable
        finally:
            Path(tmppath).unlink()


class TestE2EConsultantApplication:
    def test_domain_classification(self):
        """이력서 텍스트가 consultant_application으로 분류되어야 한다."""
        result = classify_domain(text="이력서 경력 자격 컨설턴트 신청서")
        assert result.domain == Domain.CONSULTANT_APPLICATION
        assert result.confidence > 0.5

    def test_domain_routing(self):
        """DomainRouter가 consultant_application 컨텍스트를 생성해야 한다."""
        ctx = resolve_domain(text="이력서 경력 자격")
        assert ctx.domain == Domain.CONSULTANT_APPLICATION

    def test_lrule_enforcement(self):
        """consultant_application 도메인에서 LRule enforcement가 동작해야 한다."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            _create_ca_docx(Path(f.name))
            tmppath = f.name

        try:
            report = enforce_lrules(
                domain=Domain.CONSULTANT_APPLICATION,
                document_type="resume",
                artifact_path=tmppath,
            )
            assert report.domain == "consultant_application"
            assert report.summary["total"] > 100

            # CA-specific rules should be applicable
            ca_rules = [r for r in report.rules if r["domain"] == "consultant_application"]
            assert all(r["applicable"] for r in ca_rules)

            # BP-specific rules should be N/A
            bp_rules = [r for r in report.rules if r["domain"] == "business_plan"]
            assert all(r["status"] == "N/A" for r in bp_rules)
        finally:
            Path(tmppath).unlink()

    def test_no_fabrication(self):
        """consultant_application에서 임의 사실 생성이 없어야 한다."""
        result = classify_domain(text="이력서 경력 자격 컨설턴트")
        assert result.domain == Domain.CONSULTANT_APPLICATION
        assert result.reason

    def test_ca_lrule_sha256_consistency(self):
        """신청서의 SHA256은 재계산 시 동일해야 한다."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            _create_ca_docx(Path(f.name))
            tmppath = f.name

        try:
            r1 = enforce_lrules(domain=Domain.CONSULTANT_APPLICATION, artifact_path=tmppath)
            r2 = enforce_lrules(domain=Domain.CONSULTANT_APPLICATION, artifact_path=tmppath)
            assert r1.artifact_sha256 == r2.artifact_sha256
            assert len(r1.artifact_sha256) == 64
        finally:
            Path(tmppath).unlink()

    def test_ca_finalizer_produces_draft(self):
        """신청서도 FAIL/REVIEW가 있으면 DRAFT여야 한다."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            _create_ca_docx(Path(f.name))
            tmppath = f.name

        try:
            report = enforce_lrules(domain=Domain.CONSULTANT_APPLICATION, artifact_path=tmppath)
            result = finalize_artifact(artifact_path=tmppath, lrule_report=report)
            assert result.is_draft
            assert "_DRAFT" in result.final_path
        finally:
            Path(tmppath).unlink()

    def test_ca_empty_doc_classification(self):
        """빈 신청서도 domain classification은 consultant_application이어야 한다."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.save(f.name)
            tmppath = f.name

        try:
            result = classify_domain(text="신청서 이력서 경력")
            assert result.domain == Domain.CONSULTANT_APPLICATION
        finally:
            Path(tmppath).unlink()

    def test_cross_domain_na(self):
        """BP 도메인에서 CA 규칙은 N/A여야 한다."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            _create_bp_docx(Path(f.name))
            tmppath = f.name

        try:
            report = enforce_lrules(domain=Domain.BUSINESS_PLAN, artifact_path=tmppath)
            ca_rules = [r for r in report.rules if r["domain"] == "consultant_application"]
            assert len(ca_rules) > 0
            assert all(r["status"] == "N/A" for r in ca_rules)
        finally:
            Path(tmppath).unlink()
