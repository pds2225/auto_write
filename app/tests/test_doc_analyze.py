"""test_doc_analyze.py — 공고/양식 분석 + 텍스트 추출 회귀 테스트(AI 비의존).

AI 키에 의존하지 않는 결정론 경로만 검증한다:
  - doc_text_extract: txt/docx 추출
  - announcement_analyzer: 휴리스틱(자격·금액·마감·제출서류) 추출
  - form_analyzer: 양식 작성 항목·PSST 구조 요약
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


class _NoAI:
    """available=False 더미 — 휴리스틱/규칙기반 경로 강제."""

    available = False

    def complete_json(self, *a, **k):
        return None

    def parse_announcement(self, *a, **k):
        return []


def _make_form(path: Path) -> None:
    doc = Document()
    doc.add_heading("사업계획서 양식", 0)
    for s in ["1. 문제인식(Problem)", "2. 실현가능성(Solution)",
              "3. 성장전략(Scale-up)", "4. 팀구성(Team)"]:
        doc.add_heading(s, level=1)
        doc.add_paragraph("(작성)")
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "항목"
    t.rows[0].cells[1].text = "내용"
    doc.save(str(path))


def test_extract_text_txt(tmp_path: Path) -> None:
    from auto_write.services.doc_text_extract import extract_text

    p = tmp_path / "a.txt"
    p.write_text("공고 본문\n평가기준 30점", encoding="utf-8")
    text, notes = extract_text(p)
    assert "평가기준" in text
    assert notes == []


def test_extract_text_docx(tmp_path: Path) -> None:
    from auto_write.services.doc_text_extract import extract_text

    p = tmp_path / "f.docx"
    _make_form(p)
    text, _ = extract_text(p)
    assert "문제인식" in text


def test_form_analyzer_field_kind_in_report(tmp_path: Path) -> None:
    from auto_write.services.form_analyzer import analyze_form

    p = tmp_path / "plan.docx"
    doc = Document()
    doc.add_heading("1. 문제인식(Problem)", level=1)
    doc.add_paragraph("(작성)")
    doc.save(str(p))
    fr = analyze_form(p)
    assert any(d.get("field_kind") == "narrative" for d in fr.writable_item_details)


def test_extract_missing_file(tmp_path: Path) -> None:
    from auto_write.services.doc_text_extract import extract_text

    text, notes = extract_text(tmp_path / "nope.docx")
    assert text == ""
    assert notes  # 안내 노트 존재


def test_announcement_heuristic(tmp_path: Path) -> None:
    from auto_write.services.announcement_analyzer import analyze_announcement

    p = tmp_path / "ann.txt"
    p.write_text(
        "지원대상: 예비창업자\n"
        "지원금액: 최대 1억원 이내\n"
        "신청 접수기간: 2026.07.01 ~ 2026.07.31\n"
        "제출서류: 사업계획서, 신분증 사본\n",
        encoding="utf-8",
    )
    r = analyze_announcement(p, openai_service=_NoAI())
    assert r.ai_used is False
    ki = r.key_info
    assert ki.get("funding_amount")          # 금액 추출
    assert ki.get("deadline")                # 마감 추출
    assert ki.get("required_documents")      # 제출서류 추출


def test_announcement_text_mode(tmp_path: Path) -> None:
    from auto_write.services.announcement_analyzer import analyze_announcement

    r = analyze_announcement(
        "지원금액: 최대 5천만원\n신청마감: 2026.08.15\n",
        is_text=True,
        openai_service=_NoAI(),
    )
    assert r.text_chars > 0
    assert r.key_info.get("deadline")


def test_form_analyzer(tmp_path: Path) -> None:
    from auto_write.services.form_analyzer import analyze_form

    p = tmp_path / "form.docx"
    _make_form(p)
    r = analyze_form(p)
    assert r.question_count >= 1
    # PSST 4영역 라벨이 모두 존재
    assert all(r.psst_present.values())
    assert r.as_dict()["psst_missing"] == []


def _make_announcement_txt(path: Path) -> None:
    path.write_text(
        "지원대상: 예비창업자\n"
        "신청 접수기간: 2026.07.01 ~ 2026.07.31\n"
        "제출서류: 사업계획서, 참가신청서\n",
        encoding="utf-8",
    )


def test_folder_classify_announcement_and_forms(tmp_path: Path) -> None:
    from auto_write.services.folder_analyzer import classify_folder_files

    _make_announcement_txt(tmp_path / "모집공고.txt")
    _make_form(tmp_path / "참가신청서.docx")
    (tmp_path / "포스터.jpg").write_bytes(b"x")

    ann, forms, other = classify_folder_files(tmp_path)
    assert any(p.name == "모집공고.txt" for p in ann)
    assert any(p.name == "참가신청서.docx" for p in forms)


def test_analyze_folder_korean_summary(tmp_path: Path) -> None:
    from auto_write.services.folder_analyzer import (
        analyze_folder,
        format_folder_summary_korean,
    )

    _make_announcement_txt(tmp_path / "공고문.txt")
    _make_form(tmp_path / "신청서.docx")

    r = analyze_folder(tmp_path, openai_service=_NoAI(), save_json=True)
    text = format_folder_summary_korean(r)
    assert "마감" in text
    assert "양식" in text
    assert (tmp_path / ".analysis" / "folder_analysis.json").exists()


def test_analyze_docs_folder_cli(tmp_path: Path) -> None:
    import os
    import subprocess

    app_dir = Path(__file__).resolve().parents[1]
    _make_announcement_txt(tmp_path / "공고문.txt")
    _make_form(tmp_path / "양식.docx")
    env = dict(os.environ)
    env["PYTHONIOENCODING"] = "utf-8"
    env["PYTHONPATH"] = str(app_dir)
    res = subprocess.run(
        ["py", "-3.11", str(app_dir / "analyze_docs.py"), "folder", str(tmp_path),
         "--no-ai", "--no-save-json"],
        capture_output=True, text=True, encoding="utf-8", cwd=str(app_dir), env=env,
    )
    assert res.returncode == 0
    assert "공고 폴더" in res.stdout
    assert not res.stdout.strip().startswith("{")
