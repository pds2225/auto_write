"""learn_run.py CLI 테스트 — 종료코드 계약, 저장 산출물, 원본 불변, 저장 실패에도 판정 유지.

종료코드 계약(self_diagnose 와 동일): 0=제출가능 / 1=입력오류 / 2=제출불가 / 3=검사불능(미검사 포함)
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

APP_DIR = Path(__file__).resolve().parent.parent
if str(APP_DIR) not in sys.path:
    sys.path.insert(0, str(APP_DIR))

import learn_run
from auto_write.services import learning_store as store


def _make_fail_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("사업비 [확인필요] 원")
    doc.save(str(path))


def test_defective_docx_exit2_with_runs_and_report(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setattr(learn_run, "LEARNING_ROOT", tmp_path / "learning")
    src = tmp_path / "제출본.docx"
    _make_fail_docx(src)

    rc = learn_run.main(["--final-file", str(src)])

    assert rc == 2
    runs = store.load_recent_runs(root=tmp_path / "learning")
    assert len(runs) == 1
    assert runs[0]["verdict"] == "제출불가"
    assert (tmp_path / "learning" / "learning_report.md").exists()
    defects = store.load_defects(root=tmp_path / "learning")
    assert any(d["check_id"] == "unresolved_markers" for d in defects)


def test_missing_file_exit1(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setattr(learn_run, "LEARNING_ROOT", tmp_path / "learning")
    rc = learn_run.main(["--final-file", str(tmp_path / "없는파일.docx")])
    assert rc == 1
    assert store.load_recent_runs(root=tmp_path / "learning") == []


def test_hwpx_without_acceptance_exit3_never_0_or_2(tmp_path: Path, monkeypatch) -> None:
    """§9 F2/F3: .hwpx 는 run_acceptance 를 호출하지 않는다 — 미검사→exit 3(0/2 절대 아님)."""
    monkeypatch.setattr(learn_run, "LEARNING_ROOT", tmp_path / "learning")
    src = tmp_path / "제출본.hwpx"
    src.write_bytes(b"fake hwpx bytes - content irrelevant, acceptance is never called")

    rc = learn_run.main(["--final-file", str(src)])

    assert rc == 3
    assert rc not in (0, 2)
    runs = store.load_recent_runs(root=tmp_path / "learning")
    assert len(runs) == 1
    assert runs[0]["verdict"] == "미검사"
    # HWPX 는 결함 분류를 하지 않는다(§9 F3) — defects.jsonl 에 기록이 없어야 한다.
    assert store.load_defects(root=tmp_path / "learning") == []


def test_original_document_bytes_unchanged(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setattr(learn_run, "LEARNING_ROOT", tmp_path / "learning")
    src = tmp_path / "제출본.docx"
    _make_fail_docx(src)
    before = src.read_bytes()

    learn_run.main(["--final-file", str(src)])

    assert src.read_bytes() == before


def test_store_write_failure_still_returns_verdict_code(tmp_path: Path, monkeypatch) -> None:
    """§9 M3: runs.jsonl 저장이 실패(권한 등)해도 이미 확정된 판정 종료코드는 오염되지 않는다."""
    monkeypatch.setattr(learn_run, "LEARNING_ROOT", tmp_path / "learning")
    src = tmp_path / "제출본.docx"
    _make_fail_docx(src)

    def _boom(*args, **kwargs):
        raise OSError("permission denied (simulated)")

    monkeypatch.setattr(learn_run, "append_run", _boom)

    rc = learn_run.main(["--final-file", str(src)])

    assert rc == 2  # 문서 결함 판정(제출불가)은 저장 실패와 무관하게 그대로 유지된다


def test_submittable_docx_exit0(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setattr(learn_run, "LEARNING_ROOT", tmp_path / "learning")
    src = tmp_path / "깨끗한문서.docx"
    doc = Document()
    doc.add_paragraph("결함 없는 평범한 본문입니다.")
    doc.save(str(src))

    rc = learn_run.main(["--final-file", str(src)])

    assert rc == 0
    runs = store.load_recent_runs(root=tmp_path / "learning")
    assert runs[0]["verdict"] == "제출가능"
