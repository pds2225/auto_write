"""수용검사 결함별 '다음 행동' 안내(acceptance_remediation) + self_diagnose 통합 검증.

기능: 진단이 결함 개수만 내던 것을, 결함마다 "자동 명령으로 해결 / 사람이 값 입력 /
한글에서 수동 수정" 중 무엇을 어떻게 해야 제출 가능한지 안내하도록 만든 것을 증명한다.
"""

from __future__ import annotations

from docx import Document

from auto_write.services import acceptance_remediation as rem
from auto_write.services import usage_acceptance as ua
from auto_write.services.usage_acceptance import CheckResult, SEV_FAIL, SEV_WARN


def _run_check_ids() -> set[str]:
    """실제 검사(_ALL_CHECKS)를 빈 문서에 돌려 나오는 check_id 전체(단일 출처)."""
    doc = Document()
    doc.add_paragraph("본문 한 줄")
    cfg = ua.AcceptanceConfig()
    return {check(doc, cfg).check_id for check in ua._ALL_CHECKS}


def test_every_check_id_has_specific_remedy() -> None:
    """실제 검사 전부에 전용 안내가 있어야 한다(새 검사 추가 시 여기서 강제 실패)."""
    ids = _run_check_ids()
    assert ids  # sanity: 검사가 하나라도 있어야
    for cid in ids:
        r = rem.remedy_for(cid)
        assert r is not rem._DEFAULT, f"{cid} 에 전용 안내가 없음 — _REMEDIES 에 추가 필요"
        assert r.kind in (rem.KIND_AUTO, rem.KIND_HUMAN, rem.KIND_MANUAL)
        assert r.action.strip()
        if r.kind == rem.KIND_AUTO:
            assert rem.DOC_TOKEN in r.command, f"{cid} 자동 안내에 실행 명령이 없음"
        else:
            assert r.command == "", f"{cid} 는 자동이 아닌데 명령이 붙어 있음"


def test_unknown_check_id_falls_back_to_default() -> None:
    r = rem.remedy_for("존재하지_않는_검사")
    assert r is rem._DEFAULT
    assert r.kind == rem.KIND_MANUAL


def test_build_excludes_passed_and_orders_fail_first() -> None:
    results = [
        CheckResult("font_size_spread", "글자크기 분산", SEV_WARN, 2),
        CheckResult("self_inserted_blocks", "자기삽입 블록", SEV_FAIL, 1),
        CheckResult("empty_table_rows", "빈 표 행", SEV_WARN, 0),   # passed → 제외
        CheckResult("unresolved_markers", "미해결 마커", SEV_FAIL, 3),
    ]
    items = rem.build_remediation(results, "제출본.docx")
    ids = [it["check_id"] for it in items]

    assert "empty_table_rows" not in ids           # 통과한 검사는 남은 일에서 빠진다
    assert ids[:2] == ["self_inserted_blocks", "unresolved_markers"]  # fail 먼저(원순서 유지)
    assert ids[2] == "font_size_spread"            # warn 은 뒤


def test_build_substitutes_doc_path_in_command() -> None:
    results = [CheckResult("self_inserted_blocks", "자기삽입 블록", SEV_FAIL, 1)]
    items = rem.build_remediation(results, "내 문서 v2.docx")
    cmd = items[0]["command"]
    assert "내 문서 v2.docx" in cmd
    assert rem.DOC_TOKEN not in cmd
    assert cmd.startswith("python strip_notebooklm.py")


def test_human_kind_items_have_no_command() -> None:
    results = [CheckResult("unresolved_markers", "미해결 마커", SEV_FAIL, 2)]
    items = rem.build_remediation(results, "x.docx")
    assert items[0]["kind"] == rem.KIND_HUMAN
    assert items[0]["command"] == ""


def test_unique_commands_dedups_in_first_seen_order() -> None:
    # residual_colored_runs 와 template_placeholders 는 같은 autopilot 명령을 공유
    results = [
        CheckResult("self_inserted_blocks", "블록", SEV_FAIL, 1),
        CheckResult("residual_colored_runs", "유색", SEV_FAIL, 1),
        CheckResult("template_placeholders", "자리표시", SEV_FAIL, 1),
    ]
    items = rem.build_remediation(results, "a.docx")
    cmds = rem.unique_commands(items)
    assert len(cmds) == 2                          # autopilot 명령은 한 번만
    assert cmds[0].startswith("python strip_notebooklm.py")
    assert any("auto_write_autopilot.py" in c for c in cmds)


def test_format_text_all_pass_returns_ready_line() -> None:
    lines = rem.format_remediation_text([])
    assert len(lines) == 1
    assert "바로 제출 가능" in lines[0]


def test_format_text_groups_by_kind_and_lists_commands() -> None:
    results = [
        CheckResult("self_inserted_blocks", "자기삽입 블록", SEV_FAIL, 2),
        CheckResult("unresolved_markers", "미해결 마커", SEV_FAIL, 5),
        CheckResult("font_size_spread", "글자크기 분산", SEV_WARN, 3),
    ]
    items = rem.build_remediation(results, "제출본.docx")
    text = "\n".join(rem.format_remediation_text(items))

    assert "제출까지 남은 일" in text
    assert "자동으로 해결" in text and "사람이 직접 입력" in text
    assert "자기삽입 블록(2)" in text
    assert "미해결 마커(5)" in text
    assert "글자크기 분산(3) (경고)" in text        # warn 은 (경고) 꼬리표
    assert 'python strip_notebooklm.py "제출본.docx"' in text  # 실행 명령 그대로 복붙 가능


# --- self_diagnose CLI 통합: 결함 있는 문서 진단 시 '다음 행동'이 출력·JSON 에 실린다 ---

def _defective_docx(path) -> None:
    doc = Document()
    doc.add_paragraph("[확인필요]")                       # unresolved_markers (사람 입력)
    doc.add_paragraph("이 블록은 삭제하세요")               # self_inserted_blocks (자동 제거)
    doc.save(str(path))


def test_self_diagnose_prints_and_saves_remediation(tmp_path, capsys) -> None:
    import self_diagnose as sd

    docx = tmp_path / "결함문서.docx"
    _defective_docx(docx)
    out_json = tmp_path / "결과.json"

    rc = sd.main([str(docx), "--json", str(out_json)])
    captured = capsys.readouterr().out

    assert rc == 2                                        # 제출불가(fail 존재)
    assert "제출까지 남은 일" in captured
    # 자동 해결 명령이 실제 문서 경로로 안내된다
    assert "strip_notebooklm.py" in captured
    # 사람이 입력해야 하는 항목도 구체 행동으로 안내된다
    assert "지어내면 안 됩니다" in captured

    import json
    data = json.loads(out_json.read_text(encoding="utf-8"))
    assert "remediation" in data
    kinds = {it["kind"] for it in data["remediation"]}
    assert rem.KIND_AUTO in kinds and rem.KIND_HUMAN in kinds
    # 각 항목이 check_id·action 을 기계가 읽을 수 있게 담는다
    assert all(it["check_id"] and it["action"] for it in data["remediation"])


def test_self_diagnose_clean_doc_reports_no_remaining_work(tmp_path, capsys) -> None:
    import self_diagnose as sd

    docx = tmp_path / "깨끗한문서.docx"
    doc = Document()
    doc.add_paragraph("정상 본문입니다. 결함이 없습니다.")
    doc.save(str(docx))

    rc = sd.main([str(docx)])
    captured = capsys.readouterr().out

    assert rc == 0
    assert "바로 제출 가능" in captured


# --- 구체 결함 위치(punch-list) + 진행도 요약 + 체크리스트 파일 저장 -------------

def test_build_remediation_carries_concrete_items() -> None:
    """안내가 '이 문서의 실제 결함 위치'(samples)를 그대로 담아야 한다."""
    results = [
        CheckResult("empty_label_fields", "필수 라벨 옆 칸 공란", SEV_FAIL, 2,
                    ["[표] '대표자명' 옆 칸 공란", "[표] '기업명' 옆 칸 공란"]),
    ]
    items = rem.build_remediation(results, "B.docx")
    assert items[0]["items"] == ["[표] '대표자명' 옆 칸 공란", "[표] '기업명' 옆 칸 공란"]


def test_remediation_summary_counts_by_kind() -> None:
    results = [
        CheckResult("self_inserted_blocks", "블록", SEV_FAIL, 2, ["[본문] 삭제블록"]),   # auto
        CheckResult("residual_colored_runs", "유색", SEV_FAIL, 3, ["[본문] 파란글씨"]),  # auto (같은 명령 계열)
        CheckResult("unresolved_markers", "마커", SEV_FAIL, 4),                          # human 4칸
        CheckResult("empty_label_fields", "라벨공란", SEV_FAIL, 1, ["[표] '명칭' 옆"]),  # human 1칸
        CheckResult("empty_table_rows", "빈행", SEV_WARN, 5),                            # manual (warn)
    ]
    s = rem.remediation_summary(rem.build_remediation(results, "B.docx"))
    assert s["auto_commands"] == 2            # strip_notebooklm + autopilot (서로 다른 명령)
    assert s["human_fields"] == 5             # 마커 4 + 라벨공란 1
    assert s["manual_items"] == 5             # 빈행 5
    assert s["blocking_defects"] == 10        # fail 결함만(2+3+4+1), warn 제외
    assert s["remaining_checks"] == 5


def test_summary_line_empty_and_populated() -> None:
    assert "바로 제출 가능" in rem.summary_line(rem.remediation_summary([]))
    items = rem.build_remediation(
        [CheckResult("unresolved_markers", "마커", SEV_FAIL, 3)], "x.docx")
    line = rem.summary_line(rem.remediation_summary(items))
    assert "제출까지 남은 일 요약" in line
    assert "사람이 채울 칸 3개" in line


def test_format_text_lists_concrete_items() -> None:
    results = [
        CheckResult("empty_label_fields", "필수 라벨 옆 칸 공란", SEV_FAIL, 2,
                    ["[표] '대표자명' 옆 칸 공란", "[표] '기업명' 옆 칸 공란"]),
    ]
    text = "\n".join(rem.format_remediation_text(rem.build_remediation(results, "B.docx")))
    assert "'대표자명' 옆 칸 공란" in text          # 일반 예시가 아니라 실제 결함 위치
    assert "'기업명' 옆 칸 공란" in text
    assert "제출까지 남은 일 요약" in text           # 진행도 한 줄


def test_format_text_truncates_many_items_with_more_note() -> None:
    many = [f"[표] '항목{i}' 옆 칸 공란" for i in range(5)]
    results = [CheckResult("empty_label_fields", "라벨공란", SEV_FAIL, 5, many)]
    text = "\n".join(rem.format_remediation_text(rem.build_remediation(results, "B.docx")))
    assert "… 외 2개" in text                        # 3개만 보이고 나머지는 --checklist 안내


def test_build_checklist_markdown_structure() -> None:
    results = [
        CheckResult("self_inserted_blocks", "자기삽입 블록", SEV_FAIL, 1, ["[본문] 삭제블록"]),
        CheckResult("empty_label_fields", "필수 라벨 옆 칸 공란", SEV_FAIL, 2,
                    ["[표] '대표자명' 옆 칸 공란", "[표] '기업명' 옆 칸 공란"]),
    ]
    md = rem.build_checklist_markdown(rem.build_remediation(results, "제출본.docx"),
                                      "제출본.docx", submittable=False)
    assert md.startswith("# 제출 준비 체크리스트 — 제출본.docx")
    assert "- [ ]" in md                              # 체크박스 to-do
    assert "제출불가(DRAFT)" in md
    assert "'대표자명' 옆 칸 공란" in md               # 구체 항목 전부 나열
    assert "'기업명' 옆 칸 공란" in md
    assert 'python strip_notebooklm.py "제출본.docx"' in md   # 자동 명령 복붙 블록
    assert "지어내지 마세요" in md                     # 날조0 안내


def test_build_checklist_markdown_clean_doc() -> None:
    md = rem.build_checklist_markdown([], "깨끗.docx", submittable=True)
    assert "제출 준비 완료" in md
    assert "- [ ]" not in md                          # 남은 일 없으면 체크박스 없음


def _defective_docx_with_table(path) -> None:
    """마커·자기삽입 블록 + 표 라벨 공란(구체 결함 위치가 잡히는) 문서."""
    doc = Document()
    doc.add_paragraph("[확인필요]")                    # unresolved_markers
    doc.add_paragraph("이 블록은 삭제하세요")            # self_inserted_blocks
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "대표자명"                # empty_label_fields → 구체 샘플
    t.rows[0].cells[1].text = ""                        # 옆 칸 공란
    doc.save(str(path))


def test_self_diagnose_writes_checklist_file_with_concrete_items(tmp_path, capsys) -> None:
    import self_diagnose as sd

    docx = tmp_path / "결함표문서.docx"
    _defective_docx_with_table(docx)
    chk = tmp_path / "제출준비.md"
    out_json = tmp_path / "결과.json"

    rc = sd.main([str(docx), "--checklist", str(chk), "--json", str(out_json)])
    captured = capsys.readouterr().out

    assert rc == 2
    # 화면에 진행도 요약과 구체 결함 위치가 뜬다
    assert "제출까지 남은 일 요약" in captured
    assert "'대표자명' 옆 칸 공란" in captured
    assert "제출 준비 체크리스트 저장" in captured

    # 체크리스트 파일이 실제로 생기고 체크박스·구체 항목을 담는다
    assert chk.exists()
    md = chk.read_text(encoding="utf-8")
    assert "- [ ]" in md
    assert "'대표자명' 옆 칸 공란" in md
    assert 'python strip_notebooklm.py' in md

    # JSON 에도 진행도 요약이 실려 기계가 소비할 수 있다
    import json
    data = json.loads(out_json.read_text(encoding="utf-8"))
    assert data["remediation_summary"]["human_fields"] >= 1
    assert any(it.get("items") for it in data["remediation"])


def test_self_diagnose_checklist_refuses_to_overwrite_source(tmp_path, capsys) -> None:
    """체크리스트 경로가 원본과 같으면 저장을 거부해 원본을 보호한다(날조/훼손 0)."""
    import self_diagnose as sd

    docx = tmp_path / "원본.docx"
    _defective_docx_with_table(docx)
    before = docx.read_bytes()

    rc = sd.main([str(docx), "--checklist", str(docx)])
    err = capsys.readouterr().err

    assert rc == 2
    assert "원본 보호" in err
    assert docx.read_bytes() == before               # 원본 바이트 그대로(미수정)
    # 원본이 여전히 정상 DOCX 인지(마크다운으로 덮이지 않았는지) 확인
    Document(str(docx))
