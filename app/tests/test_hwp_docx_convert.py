"""test_hwp_docx_convert.py — HWP/HWPX ↔ DOCX 양방향 변환 서비스 테스트.

한글(Hancom) COM 은 CI/백그라운드에서 못 띄우므로 가짜 객체로 검증하고,
HWP(unhwp)·PrvText 경로는 document_ingest 의 로더를 patch 해 검증한다.
HWPX 는 실제 zip 픽스처로 end-to-end 구조 변환을 검증한다.
"""

from __future__ import annotations

import json
import zipfile
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

from auto_write.services import hwp_docx_convert as mod
from auto_write.services.hwp_docx_convert import convert, docx_to_hwp, hwp_to_docx


# --- 픽스처 도우미 ------------------------------------------------------------

def _make_hwpx(path: Path) -> None:
    section_xml = """<?xml version="1.0" encoding="UTF-8"?>
<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section"
        xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph">
  <hp:p id="1"><hp:run><hp:t>사업 개요</hp:t></hp:run></hp:p>
  <hp:p id="2"><hp:run><hp:t>핵심 전략 본문</hp:t></hp:run></hp:p>
</hs:sec>"""
    with zipfile.ZipFile(path, "w") as z:
        z.writestr("mimetype", "application/hwpx")
        z.writestr("Contents/section0.xml", section_xml)
        z.writestr("Preview/PrvText.txt", "사업 개요\n핵심 전략 본문")


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("변환 대상 본문")
    doc.save(str(path))


def _unhwp_payload() -> dict:
    return {
        "sections": [{
            "content": [
                {"Paragraph": {"style": {"heading_level": 0},
                               "content": [{"Text": {"text": "유니크문단 본문"}}]}},
                {"Table": {"rows": [
                    {"is_header": True, "cells": [
                        {"content": [{"Text": {"text": "항목"}}], "colspan": 1, "rowspan": 1},
                        {"content": [{"Text": {"text": "수치"}}], "colspan": 1, "rowspan": 1},
                    ]},
                    {"is_header": False, "cells": [
                        {"content": [{"Text": {"text": "매출"}}], "colspan": 1, "rowspan": 1},
                        {"content": [{"Text": {"text": "100"}}], "colspan": 1, "rowspan": 1},
                    ]},
                ]}},
            ],
        }],
    }


class _FakeUnhwpResult:
    def __init__(self, payload: dict) -> None:
        self.json = json.dumps(payload)

    def __enter__(self):
        return self

    def __exit__(self, *args):
        return False


class _FakeUnhwp:
    def parse(self, path: str) -> _FakeUnhwpResult:
        return _FakeUnhwpResult(_unhwp_payload())


class _FakeOleStream:
    def __init__(self, data: bytes) -> None:
        self._data = data

    def read(self) -> bytes:
        return self._data


class _FakeOle:
    def __init__(self, data: bytes) -> None:
        self._data = data

    def exists(self, name) -> bool:
        return True

    def openstream(self, name) -> _FakeOleStream:
        return _FakeOleStream(self._data)

    def close(self) -> None:
        pass


class _FakeOlefileMod:
    def __init__(self, text: str) -> None:
        self._data = text.encode("utf-16-le")

    def isOleFile(self, path) -> bool:
        return True

    def OleFileIO(self, path) -> _FakeOle:
        return _FakeOle(self._data)


class _FakeHwpCom:
    """한글 COM 흉내 — SaveAs 가 실제로 파일을 만들어야 존재 검사를 통과한다."""

    def __init__(self) -> None:
        self.saved: list[tuple[str, str]] = []

    def RegisterModule(self, *a):
        return True

    def SetMessageBoxMode(self, *a):
        return 0

    def Open(self, path, fmt, opts):
        return True

    def SaveAs(self, path, fmt, opts):
        Path(path).write_bytes(b"FAKE-HWP-BINARY")
        self.saved.append((path, fmt))
        return True

    def Clear(self, *a):
        pass

    def Quit(self):
        pass


def _all_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


# --- 안전장치 -----------------------------------------------------------------

def test_out_equals_in_blocked(tmp_path: Path) -> None:
    src = tmp_path / "a.hwpx"
    _make_hwpx(src)
    with pytest.raises(ValueError):
        hwp_to_docx(str(src), str(src))


def test_missing_input_raises(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        hwp_to_docx(str(tmp_path / "없음.hwp"))


def test_convert_rejects_unknown_ext(tmp_path: Path) -> None:
    src = tmp_path / "a.pdf"
    src.write_bytes(b"%PDF")
    with pytest.raises(ValueError):
        convert(str(src))


# --- HWP/HWPX → DOCX ----------------------------------------------------------

def test_hwpx_to_docx_structure(tmp_path: Path) -> None:
    src = tmp_path / "form.hwpx"
    out = tmp_path / "form.docx"
    _make_hwpx(src)
    r = hwp_to_docx(str(src), str(out), use_com=False)
    assert r.ok and r.method == "hwpx_xml"
    assert "핵심 전략 본문" in _all_text(out)


def test_hwp_to_docx_unhwp_structure(tmp_path: Path) -> None:
    src = tmp_path / "form.hwp"
    out = tmp_path / "form.docx"
    src.write_bytes(b"HWP-DUMMY")
    with patch("auto_write.document_ingest._load_unhwp", return_value=_FakeUnhwp()):
        r = hwp_to_docx(str(src), str(out), use_com=False)
    assert r.ok and r.method == "unhwp"
    text = _all_text(out)
    assert "유니크문단 본문" in text      # 문단 보존
    assert "매출" in text and "100" in text   # 표 보존


def test_hwp_to_docx_prvtext_fallback(tmp_path: Path) -> None:
    """unhwp 미가용이면 미리보기 텍스트로라도 DOCX 를 만든다."""
    src = tmp_path / "form.hwp"
    out = tmp_path / "form.docx"
    src.write_bytes(b"HWP-DUMMY")
    fake_ole = _FakeOlefileMod("미리보기 본문 텍스트")
    with patch("auto_write.document_ingest._load_unhwp", return_value=None), \
         patch("auto_write.document_ingest._load_olefile", return_value=fake_ole):
        r = hwp_to_docx(str(src), str(out), use_com=False)
    assert r.ok and r.method == "prvtext"
    assert "미리보기 본문 텍스트" in _all_text(out)
    assert any("누락" in n for n in r.notes)   # 한계 고지


def test_hwp_to_docx_com_failure_falls_back(tmp_path: Path, monkeypatch) -> None:
    """COM 등록은 됐지만 세션에서 못 뜨는 경우(백그라운드) 구조 변환으로 폴백."""
    src = tmp_path / "form.hwpx"
    out = tmp_path / "form.docx"
    _make_hwpx(src)
    monkeypatch.setattr(mod, "hancom_com_available", lambda: True)

    def _boom():
        raise RuntimeError("COM 서버를 띄울 수 없음")

    monkeypatch.setattr(mod, "_dispatch_hwp", _boom)
    r = hwp_to_docx(str(src), str(out), use_com=True)
    assert r.ok and r.method == "hwpx_xml"
    assert any("COM 변환 실패" in n for n in r.notes)


def test_hwp_to_docx_com_success(tmp_path: Path, monkeypatch) -> None:
    src = tmp_path / "form.hwp"
    out = tmp_path / "form.docx"
    src.write_bytes(b"HWP-DUMMY")
    fake = _FakeHwpCom()
    monkeypatch.setattr(mod, "hancom_com_available", lambda: True)
    monkeypatch.setattr(mod, "_dispatch_hwp", lambda: fake)
    r = hwp_to_docx(str(src), str(out), use_com=True)
    assert r.ok and r.method == "hancom_com"
    assert fake.saved and fake.saved[0][1] == "DOCX"


# --- DOCX → HWP ----------------------------------------------------------------

def test_docx_to_hwp_com_unavailable_reports_failure(tmp_path: Path, monkeypatch) -> None:
    src = tmp_path / "plan.docx"
    _make_docx(src)
    monkeypatch.setattr(mod, "hancom_com_available", lambda: False)
    r = docx_to_hwp(str(src))
    assert r.ok is False
    assert any("한글" in n for n in r.notes)   # 사람이 할 일 안내


def test_docx_to_hwp_with_fake_com(tmp_path: Path, monkeypatch) -> None:
    src = tmp_path / "plan.docx"
    out = tmp_path / "plan.hwp"
    _make_docx(src)
    fake = _FakeHwpCom()
    monkeypatch.setattr(mod, "hancom_com_available", lambda: True)
    monkeypatch.setattr(mod, "_dispatch_hwp", lambda: fake)
    r = docx_to_hwp(str(src), str(out))
    assert r.ok and r.method == "hancom_com"
    assert out.exists()
    assert fake.saved[0][1] == "HWP"


def test_docx_to_hwpx_uses_hwpx_format(tmp_path: Path, monkeypatch) -> None:
    src = tmp_path / "plan.docx"
    out = tmp_path / "plan.hwpx"
    _make_docx(src)
    fake = _FakeHwpCom()
    monkeypatch.setattr(mod, "hancom_com_available", lambda: True)
    monkeypatch.setattr(mod, "_dispatch_hwp", lambda: fake)
    r = docx_to_hwp(str(src), str(out))
    assert r.ok
    assert fake.saved[0][1] in ("HWPX", "HWPML2X")


# --- 방향 자동 인식 + CLI ------------------------------------------------------

def test_convert_auto_direction(tmp_path: Path, monkeypatch) -> None:
    hwpx = tmp_path / "form.hwpx"
    _make_hwpx(hwpx)
    r1 = convert(str(hwpx), str(tmp_path / "o1.docx"), use_com=False)
    assert r1.direction == "hwp->docx" and r1.ok

    docx = tmp_path / "plan.docx"
    _make_docx(docx)
    monkeypatch.setattr(mod, "hancom_com_available", lambda: False)
    r2 = convert(str(docx), str(tmp_path / "o2.hwp"))
    assert r2.direction == "docx->hwp" and r2.ok is False


def test_cli_main(tmp_path: Path) -> None:
    import hwp_docx

    src = tmp_path / "form.hwpx"
    out = tmp_path / "form_cli.docx"
    _make_hwpx(src)
    assert hwp_docx.main([str(src), "-o", str(out), "--no-com"]) == 0
    assert out.exists()
    assert hwp_docx.main([str(tmp_path / "없는파일.hwp")]) == 2
