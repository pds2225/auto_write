"""test_auto_write_apply.py — image_apply / psst_fill / autopilot_pipeline 회귀 테스트.

새로 추가한 '실제 수정' 모듈들이:
  - 원본을 덮어쓰지 않고(out==in 가드)
  - 표 실측치가 있으면 차트를, 없으면 자리표시를 삽입하며
  - PSST 미흡/누락 영역에 작성 가이드를 추가하고
  - autopilot 이 전 단계를 무인 연속 실행하는지
를 검증한다. (숫자 날조가 없어야 함 — placeholder 폴백 동작 포함)
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from auto_write.services.image_apply import apply_images
from auto_write.services.psst_fill import apply_psst_scaffold


def _make_doc(path: Path, *, with_table: bool = True) -> None:
    doc = Document()
    doc.add_heading("사업계획서", 0)
    doc.add_paragraph("가. 문제인식: 고객 시장 니즈와 기존 대안의 한계로 비용 손실.")
    doc.add_paragraph("나. 추진일정 로드맵 — 단계별 마일스톤.")          # gantt 트리거
    doc.add_paragraph("다. 목표 시장규모 TAM SAM SOM 성장률 전망.")       # 막대/도넛 트리거
    if with_table:
        t = doc.add_table(rows=2, cols=3)
        t.rows[0].cells[0].text = "2024년"
        t.rows[0].cells[1].text = "2025년"
        t.rows[0].cells[2].text = "2026년"
        t.rows[1].cells[0].text = "100"
        t.rows[1].cells[1].text = "200"
        t.rows[1].cells[2].text = "350"
    doc.save(str(path))


def test_apply_images_in_equals_out_blocked(tmp_path: Path) -> None:
    src = tmp_path / "in.docx"
    _make_doc(src)
    with pytest.raises(ValueError):
        apply_images(str(src), str(src))


def test_apply_images_inserts_notebooklm_prompt(tmp_path: Path) -> None:
    src = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    _make_doc(src, with_table=True)
    report = apply_images(str(src), str(out))  # openai_service=None → 키워드 폴백
    assert out.exists()
    # 원본은 그대로(문단 수 변화 없음)
    assert len(Document(str(src)).paragraphs) < len(Document(str(out)).paragraphs)
    # 그림 위치마다 NotebookLM 슬라이드 프롬프트 블록이 삽입되어야 한다
    assert report.prompts_inserted >= 1
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "NotebookLM" in text


def test_apply_images_placeholder_only_still_inserts_prompt(tmp_path: Path) -> None:
    src = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    _make_doc(src, with_table=True)
    # placeholder_only 는 하위호환용(동작에 영향 없음) — 항상 프롬프트 블록 삽입
    report = apply_images(str(src), str(out), placeholder_only=True)
    assert report.prompts_inserted >= 1


def test_apply_images_no_table_still_inserts_prompt(tmp_path: Path) -> None:
    src = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    _make_doc(src, with_table=False)
    report = apply_images(str(src), str(out))
    # 표가 없어도 키워드 매칭 위치에 슬라이드 프롬프트가 들어간다 (숫자 날조 없음)
    assert report.prompts_inserted >= 1


def test_apply_images_table_anchor_inserts_after_table_not_end(tmp_path: Path) -> None:
    """버그① 회귀: 키워드가 '표 헤더'에만 있는(표 기반 양식) 경우에도
    NotebookLM 프롬프트가 문서 끝에 덤프되지 않고 해당 표 바로 뒤에 들어가야 한다."""
    src = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph("개요: 본 사업계획서 본문(키워드 없음).")
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "추진 일정"          # 로드맵/간트 트리거(표 헤더)
    table.rows[0].cells[1].text = "마일스톤"
    table.rows[0].cells[2].text = "담당"
    table.rows[1].cells[0].text = "1분기"
    doc.add_paragraph("맺음말: 마지막 본문 단락.")       # 표보다 뒤에 있는 본문
    doc.save(str(src))

    report = apply_images(str(src), str(out))            # openai_service=None → 키워드 폴백
    assert report.prompts_inserted >= 1
    assert report.anchors_missing == 0                   # 표 앵커를 찾았어야 함

    # 본문 순서상: 표(tbl) < NotebookLM 프롬프트 < 맺음말  (끝에 덤프 아님)
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph as _P
    out_doc = Document(str(out))
    seq = []
    for child in out_doc.element.body:
        if child.tag == qn("w:tbl"):
            seq.append(("tbl", ""))
        elif child.tag == qn("w:p"):
            seq.append(("p", _P(child, out_doc).text))
    idx_tbl = next(i for i, s in enumerate(seq) if s[0] == "tbl")
    idx_prompt = next(i for i, s in enumerate(seq) if "NotebookLM" in s[1])
    idx_end = next(i for i, s in enumerate(seq) if "맺음말" in s[1])
    assert idx_tbl < idx_prompt < idx_end


def test_submittable_filler_paragraph_fill_in_table_cell(tmp_path: Path) -> None:
    """버그①b 회귀: 채울 본문 앵커가 '표 셀 안'에 있어도 누락 없이 채워야 한다
    (이전엔 doc.paragraphs 만 봐서 표 셀 앵커를 '본문 앵커 미발견'으로 건너뜀)."""
    from auto_write.services.submittable_filler import SubmittableFiller

    src = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph("머리말")
    t = doc.add_table(rows=1, cols=1)
    t.rows[0].cells[0].text = "5. AI 인재활용 계획 세부내용 작성"   # 표 셀 안 앵커
    doc.save(str(src))

    plan = {"paragraph_fills": [
        {"anchor": "5. AI 인재활용 계획 세부내용 작성",
         "lines": ["실제 인재활용 계획 내용입니다.", "하위 항목 1"]}
    ]}
    report = SubmittableFiller(plan).finalize(src, out)
    assert report["paragraphs_filled"] == 1
    assert not any("앵커 미발견" in n for n in report["notes"])
    cell_text = "\n".join(
        c.text for tb in Document(str(out)).tables for r in tb.rows for c in r.cells
    )
    assert "실제 인재활용 계획 내용입니다." in cell_text


def test_psst_scaffold_adds_guidance(tmp_path: Path) -> None:
    src = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    _make_doc(src)  # team/scale 영역이 비어 미흡/누락
    report = apply_psst_scaffold(str(src), str(out))
    assert out.exists()
    assert report.areas_scaffolded >= 1
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "작성 보강 가이드" in text


def test_psst_scaffold_in_equals_out_blocked(tmp_path: Path) -> None:
    src = tmp_path / "in.docx"
    _make_doc(src)
    with pytest.raises(ValueError):
        apply_psst_scaffold(str(src), str(src))


def test_autopilot_end_to_end(tmp_path: Path) -> None:
    from auto_write.services.autopilot_pipeline import run_autopilot

    src = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    _make_doc(src, with_table=True)
    report = run_autopilot(str(src), str(out), write_report=False)
    assert out.exists()
    assert report.backup_dir  # 원본 백업이 생성되어야 함
    assert report.score_total > 0
    # 그림 위치에 NotebookLM 프롬프트가 최소 1개, PSST 보강이 일어났을 것
    assert report.prompts_inserted >= 1
    assert report.psst_areas_scaffolded >= 1


def test_autopilot_in_equals_out_blocked(tmp_path: Path) -> None:
    from auto_write.services.autopilot_pipeline import run_autopilot

    src = tmp_path / "in.docx"
    _make_doc(src)
    with pytest.raises(ValueError):
        run_autopilot(str(src), str(src))


# --- bizplan 생성·완성 오케스트레이터 (AI 비의존 결정론 경로) ---

def test_ai_writer_skips_without_key(tmp_path: Path) -> None:
    from auto_write.services.bizplan_ai_writer import ai_write_areas

    src = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    _make_doc(src)
    r = ai_write_areas(str(src), str(out), openai_service=None)
    assert r.skipped is True          # AI 키 없으면 본문 작성 생략(안전)
    assert out.exists()
    assert r.areas_written == 0


def test_bizplan_no_ai_completes(tmp_path: Path) -> None:
    from auto_write.services.bizplan_autopilot import run_bizplan_autopilot

    src = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    _make_doc(src, with_table=True)
    r = run_bizplan_autopilot(str(src), str(out), use_ai=False, write_report=False)
    assert out.exists()
    assert r.loops_run == 1           # 공고 없음 → 1회 완성
    assert r.ai_used is False
    assert r.backup_dir               # 원본 백업 생성


def test_bizplan_in_equals_out_blocked(tmp_path: Path) -> None:
    from auto_write.services.bizplan_autopilot import run_bizplan_autopilot

    src = tmp_path / "in.docx"
    _make_doc(src)
    with pytest.raises(ValueError):
        run_bizplan_autopilot(str(src), str(src), use_ai=False)


def test_anchor_forward_match_wins_over_reverse_substring(tmp_path: Path) -> None:
    """역포함 오매칭 회귀: 앞쪽 본문에 앵커 부분문자열(4~7자)이 있어도
    뒤쪽 표의 정방향 매칭 앵커가 선택되어야 한다(1차 패스 우선)."""
    from auto_write.services.image_apply import _find_anchor

    doc = Document()
    # 앞쪽 본문: "사업화" — 길이 3이라 역포함 임계(>=8) 미달, 4~7자 경계 테스트를 위해
    # 7자("추진계획")도 추가해 임계 정확성 검증
    doc.add_paragraph("사업화")          # 3자 — 역포함 임계(>=8) 미달
    doc.add_paragraph("추진계획")        # 4자 — 구 코드(>=4)에서 오매칭, 신 코드(>=8)에서 통과
    doc.add_paragraph("사업화추진")      # 5자 — 구 코드 오매칭, 신 코드 통과
    # 뒤쪽 표: 정방향 매칭 앵커("사업화추진계획 로드맵")
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "사업화추진계획 로드맵 단계별 마일스톤"
    doc.add_paragraph("맺음말 단락.")

    # anchor_text 앞 40자 = "사업화추진계획 로드맵 단계별 마일스톤"
    # 정방향: key("사업화추진계획 로드맵 단계별 마일스톤") in cell_text → 표 셀 매칭
    # 역포함 후보: "추진계획"(4자, >=8 불만족), "사업화추진"(5자, >=8 불만족) → 모두 탈락
    anchor = "사업화추진계획 로드맵 단계별 마일스톤"
    para, table_found = _find_anchor(doc, anchor)

    assert para is not None, "앵커를 찾지 못함 — 표 셀 정방향 매칭이 동작해야 함"
    assert table_found is not None, "표 안 단락이 아님 — 뒤쪽 표 셀이 선택되어야 함"
    assert anchor in para.text, "선택된 단락 텍스트가 앵커 키를 포함해야 함"


# --- autopilot_pipeline 결함 수정 회귀 테스트 ---

def test_residual_re_matches_confirmation_needed() -> None:
    """(B) _RESIDUAL_RE 가 [확인필요]와 [산출근거]를 정확히 매칭해야 한다."""
    from auto_write.services.autopilot_pipeline import _RESIDUAL_RE

    assert _RESIDUAL_RE.search("[확인필요]"), "[확인필요] 미매칭"
    assert _RESIDUAL_RE.search("[산출근거]"), "[산출근거] 미매칭"
    # 기존 패턴도 유지되어야 함
    assert _RESIDUAL_RE.search("___"), "밑줄 패턴 미매칭"
    assert _RESIDUAL_RE.search("(작성)"), "작성 패턴 미매칭"


def test_autopilot_report_fields_and_write_report(tmp_path: Path) -> None:
    """(A)+(B) AutopilotReport 에 신규 필드와 리포트 md 표기를 검증한다.

    run_autopilot E2E 대신 _write_report/_build_todo 를 직접 호출해
    메모리 압박 없이 결정론적으로 검증한다.
    """
    from auto_write.services.autopilot_pipeline import AutopilotReport, _build_todo, _write_report

    report = AutopilotReport(
        input_docx=str(tmp_path / "in.docx"),
        output_docx=str(tmp_path / "out.docx"),
        backup_dir=str(tmp_path / "backup"),
        doc_type="사업계획서 (90%)",
        score_total=92.0,
        grade="우수",
        passed=True,
        iterations=1,
        ops_summary="안내문구-2 글머리표-3 표셀-1 빈단락-4 강조-5",
        prompts_inserted=2,
        psst_overall_ratio=0.75,
        psst_areas_scaffolded=1,
        psst_items_added=3,
        psst_scaffolded_areas=["팀구성"],
        residual_placeholders=["[확인필요] 예산", "[산출근거] 매출"],
        residual_total=5,  # 총 5건(샘플은 2건)
        final_score_total=84.0,
        final_passed=False,
    )

    # AutopilotReport 에 신규 필드가 있어야 함
    assert hasattr(report, "final_score_total"), "final_score_total 필드 없음"
    assert hasattr(report, "final_passed"), "final_passed 필드 없음"
    assert hasattr(report, "residual_total"), "residual_total 필드 없음"

    # as_dict 에도 포함되어야 함
    d = report.as_dict()
    assert "final_score_total" in d
    assert "final_passed" in d
    assert "residual_total" in d
    assert d["residual_total"] == 5
    assert d["final_score_total"] == 84.0
    assert d["final_passed"] is False

    # _build_todo: 최종점수 미달 → To-Do 항목 추가되어야 함
    todo = _build_todo(report)
    assert any("최종본 재채점" in t for t in todo), "최종본 재채점 미달 To-Do 없음"
    assert any("총 5곳" in t for t in todo), "잔존 총건수 To-Do 없음"

    # _write_report: md 에 '최종본 재채점' 줄과 잔존 총건수 포함
    report.manual_todo = todo
    md_path = _write_report(tmp_path, "test_stem", report)
    md_text = Path(md_path).read_text(encoding="utf-8")
    assert "최종본 재채점" in md_text, "리포트에 최종본 재채점 줄 없음"
    assert "잔존 빈칸/[확인필요] 총" in md_text, "리포트에 잔존 총건수 표기 없음"
