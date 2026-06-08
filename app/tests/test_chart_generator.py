"""데이터 차트·도식 생성 모듈 테스트.

대상:
  - auto_write.services.chart_generator (bar/line/gantt/org 차트 PNG 생성)
  - auto_write.services.chart_insert     (PNG 를 DOCX 에 삽입)

검증 항목:
  1) 각 차트 함수가 유효 입력에서 PNG 파일을 실제로 생성하는지(존재 + 크기>0)
  2) 빈/잘못된 입력에서 None 을 반환하는지(예외 없이)
  3) insert_image_after_anchor 가 out_docx 를 만들고, in==out 이면 ValueError 인지

실행: (app 디렉토리 기준)  python -m pytest tests/test_chart_generator.py -q
"""
from __future__ import annotations

import tempfile
from pathlib import Path

import pytest
from docx import Document

from auto_write.services import chart_generator as cg
from auto_write.services.chart_insert import insert_image_after_anchor


# --------------------------------------------------------------------------- helpers
def _png_ok(path) -> bool:
    return path is not None and Path(path).exists() and Path(path).stat().st_size > 0


# --------------------------------------------------------------------------- 1: 막대
def test_bar_chart_creates_png(tmp_path: Path):
    out = tmp_path / "bar.png"
    res = cg.bar_chart(str(out), "연도별 매출", ["2023", "2024", "2025"], [10, 25, 40], "억원")
    assert res == str(out)
    assert _png_ok(res)


def test_bar_chart_invalid_returns_none(tmp_path: Path):
    out = str(tmp_path / "b.png")
    assert cg.bar_chart(out, "t", [], []) is None
    assert cg.bar_chart(out, "t", ["a", "b"], [1]) is None  # 길이 불일치
    assert cg.bar_chart(out, "t", ["a"], ["x"]) is None  # 숫자 아님
    assert cg.bar_chart(out, "t", [], [1, 2]) is None


# --------------------------------------------------------------------------- 2: 꺾은선
def test_line_chart_creates_png(tmp_path: Path):
    out = tmp_path / "line.png"
    res = cg.line_chart(
        str(out), "성장 추이", ["1Q", "2Q", "3Q"],
        {"매출": [5, 10, 18], "이익": [1, 3, 6]}, "억원",
    )
    assert res == str(out)
    assert _png_ok(res)


def test_line_chart_invalid_returns_none(tmp_path: Path):
    out = str(tmp_path / "l.png")
    assert cg.line_chart(out, "t", [], {"a": [1]}) is None
    assert cg.line_chart(out, "t", ["x", "y"], {}) is None
    assert cg.line_chart(out, "t", ["x", "y"], {"a": [1]}) is None  # 길이 불일치
    assert cg.line_chart(out, "t", ["x"], {"a": ["nope"]}) is None  # 숫자 아님


# --------------------------------------------------------------------------- 3: 간트
def test_gantt_chart_creates_png(tmp_path: Path):
    out = tmp_path / "gantt.png"
    tasks = [
        {"name": "기획", "start": 0, "end": 2},
        {"name": "개발", "start": 2, "end": 6},
        {"name": "검증", "start": 5, "end": 8},
    ]
    res = cg.gantt_chart(str(out), "추진 일정", tasks)
    assert res == str(out)
    assert _png_ok(res)


def test_gantt_chart_invalid_returns_none(tmp_path: Path):
    out = str(tmp_path / "g.png")
    assert cg.gantt_chart(out, "t", []) is None
    assert cg.gantt_chart(out, "t", "not a list") is None  # type: ignore[arg-type]
    # 유효 항목이 하나도 없는 경우(end<=start, 키 누락)
    assert cg.gantt_chart(out, "t", [{"name": "x", "start": 3, "end": 1}]) is None
    assert cg.gantt_chart(out, "t", [{"name": "x", "start": 0}]) is None


# --------------------------------------------------------------------------- 4: 조직도
def test_org_chart_creates_png(tmp_path: Path):
    out = tmp_path / "org.png"
    nodes = [
        {"id": "ceo", "label": "대표", "parent": None},
        {"id": "dev", "label": "개발팀", "parent": "ceo"},
        {"id": "biz", "label": "사업팀", "parent": "ceo"},
        {"id": "fe", "label": "프론트", "parent": "dev"},
    ]
    res = cg.org_chart(str(out), "팀 조직도", nodes)
    assert res == str(out)
    assert _png_ok(res)


def test_org_chart_invalid_returns_none(tmp_path: Path):
    out = str(tmp_path / "o.png")
    assert cg.org_chart(out, "t", []) is None
    # 루트 0개(모두 부모 보유) → None
    assert cg.org_chart(out, "t", [
        {"id": "a", "label": "A", "parent": "b"},
        {"id": "b", "label": "B", "parent": "a"},
    ]) is None
    # 존재하지 않는 부모 참조 → None
    assert cg.org_chart(out, "t", [
        {"id": "a", "label": "A", "parent": None},
        {"id": "c", "label": "C", "parent": "ghost"},
    ]) is None


# --------------------------------------------------------------------------- 5: 삽입
def test_insert_image_after_anchor_creates_out(tmp_path: Path):
    # 이미지 준비
    png = tmp_path / "chart.png"
    assert _png_ok(cg.bar_chart(str(png), "x", ["a", "b"], [1, 2]))
    # 원본 DOCX 준비
    in_docx = tmp_path / "in.docx"
    d = Document()
    d.add_paragraph("3. 추진 일정 (Schedule)")
    d.add_paragraph("뒷 내용")
    d.save(str(in_docx))

    out_docx = tmp_path / "out.docx"
    found = insert_image_after_anchor(
        str(in_docx), str(out_docx), "추진 일정", str(png),
        caption="[그림] 추진 일정", width_inches=6.0,
    )
    assert found is True
    assert out_docx.exists()
    # 원본은 손대지 않음
    assert in_docx.exists()


def test_insert_anchor_not_found_returns_false(tmp_path: Path):
    png = tmp_path / "c2.png"
    assert _png_ok(cg.bar_chart(str(png), "x", ["a"], [1]))
    in_docx = tmp_path / "in2.docx"
    d = Document()
    d.add_paragraph("아무 내용")
    d.save(str(in_docx))
    out_docx = tmp_path / "out2.docx"
    found = insert_image_after_anchor(str(in_docx), str(out_docx), "없는앵커", str(png))
    assert found is False
    assert out_docx.exists()


def test_insert_same_path_raises(tmp_path: Path):
    png = tmp_path / "c3.png"
    assert _png_ok(cg.bar_chart(str(png), "x", ["a"], [1]))
    in_docx = tmp_path / "same.docx"
    Document().save(str(in_docx))
    with pytest.raises(ValueError):
        insert_image_after_anchor(str(in_docx), str(in_docx), "x", str(png))


# --------------------------------------------------------------------------- 6: 정리
def test_tempfile_cleanup():
    # tempfile 로 만든 PNG 가 컨텍스트 종료 후 정리되는지(누수 없음) 확인
    with tempfile.TemporaryDirectory() as td:
        out = Path(td) / "t.png"
        assert _png_ok(cg.bar_chart(str(out), "x", ["a", "b"], [3, 4]))
        assert out.exists()
    assert not out.exists()
