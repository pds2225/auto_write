"""test_cross_form_autofill.py — cross-form value autofill(v1) 회귀 테스트.

검증 항목:
  ① 정확일치 라벨("기업명") 전사
  ② 동의어 전사("대표자" → "성명")  ← 핵심
  ③ 동의어 전사("사업자등록번호" → "사업자번호")
  ④ 소스에 없는 라벨("사업명")은 미매칭 → 빈칸 유지(날조 0)
  ⑤ 원본 source·target 파일 미수정
  ⑥ out==source / out==target → ValueError
  ⑦ 전사 값이 정확히 소스 값과 동일(날조 없음), 무관 라벨엔 값 안 들어감(보수적)

적대 반례(오매칭 차단):
  ⑧ '사업명' ↛ '사업자명' 자동전사 금지(퍼지 → needs_confirm)
  ⑨ '주소' ↛ '주소지정' 자동전사 금지
  ⑩ 클러스터 충돌(기업명+회사명 → 상호) → needs_confirm·전사 0
  ⑪ 마스킹 값 '○○○'·'OOO-OO-OOOOO' 전사 후 보존(훼손 없음)
  ⑫ 4컬럼 label|value|label|value 행 정확 전사
  ⑬ 추출 오염 없음(값 셀이 라벨로 안 들어감)
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from auto_write.services.cross_form_autofill import (
    autofill_from_source,
    extract_source_fields,
    find_checkbox_targets,
    find_target_fields,
    match_checkbox_groups,
    match_fields,
)


def _make_source(path: Path) -> None:
    """소스: 라벨|값 2열 표."""
    doc = Document()
    doc.add_heading("사업계획서", 0)
    t = doc.add_table(rows=3, cols=2)
    t.rows[0].cells[0].text = "기업명"
    t.rows[0].cells[1].text = "밸류업(주)"
    t.rows[1].cells[0].text = "대표자"
    t.rows[1].cells[1].text = "홍길동"
    t.rows[2].cells[0].text = "사업자등록번호"
    t.rows[2].cells[1].text = "123-45-67890"
    doc.save(str(path))


def _make_target(path: Path) -> None:
    """타깃: 라벨|<빈칸> 2열 표 (값칸 비어 있음)."""
    doc = Document()
    doc.add_heading("지원신청서", 0)
    t = doc.add_table(rows=4, cols=2)
    labels = ["기업명", "성명", "사업자번호", "사업명"]
    for i, lab in enumerate(labels):
        t.rows[i].cells[0].text = lab
        t.rows[i].cells[1].text = ""  # 빈 값칸
    doc.save(str(path))


def _value_for(docx_path: Path, label: str) -> str:
    """타깃 문서에서 라벨 행의 값칸 텍스트를 읽는다."""
    doc = Document(str(docx_path))
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if cells and cells[0].text.strip() == label:
                return cells[1].text.strip()
    return ""


# --- 단위: 추출/탐지/매칭 -----------------------------------------------------

def test_extract_source_fields(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    _make_source(src)
    fields = extract_source_fields(src)
    assert fields["기업명"] == "밸류업(주)"
    assert fields["대표자"] == "홍길동"
    assert fields["사업자등록번호"] == "123-45-67890"


def test_find_target_fields_only_blank(tmp_path: Path) -> None:
    tgt = tmp_path / "b.docx"
    _make_target(tgt)
    targets = find_target_fields(tgt)
    norms = {t["normalized"] for t in targets}
    assert {"기업명", "성명", "사업자번호", "사업명"} <= norms
    # 값셀 인덱스는 라벨 다음 칸(1)
    assert all(t["value_cell"] == 1 for t in targets)


def test_match_synonym_daepyoja_to_seongmyeong(tmp_path: Path) -> None:
    """동의어 매칭 핵심: 소스 '대표자' → 타깃 '성명'."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    _make_source(src)
    _make_target(tgt)
    source = extract_source_fields(src)
    targets = find_target_fields(tgt)
    matches = match_fields(source, targets)
    by_norm = {m.normalized: m for m in matches}
    assert by_norm["성명"].source_label == "대표자"
    assert by_norm["성명"].value == "홍길동"
    assert by_norm["성명"].confidence == "high"
    # 소스에 없는 사업명은 매칭 안 됨
    assert by_norm["사업명"].source_label == ""
    assert by_norm["사업명"].value == ""


# --- end-to-end: 전사 --------------------------------------------------------

def test_autofill_exact_synonym_and_no_fabrication(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _make_source(src)
    _make_target(tgt)

    report = autofill_from_source(src, tgt, out, use_ai=False)
    assert report.ok
    assert out.exists()

    # ① 정확일치 전사
    assert _value_for(out, "기업명") == "밸류업(주)"
    # ② 동의어 전사(핵심): 대표자 → 성명
    assert _value_for(out, "성명") == "홍길동"
    # ③ 동의어 전사: 사업자등록번호 → 사업자번호
    assert _value_for(out, "사업자번호") == "123-45-67890"
    # ④ 소스에 없는 사업명은 미매칭 → 빈칸 유지(날조 0)
    assert _value_for(out, "사업명") == ""

    # 전사 건수 = 3 (기업명/성명/사업자번호)
    assert report.transcribed == 3
    assert {m.normalized for m in report.matches} == {"기업명", "성명", "사업자번호"}
    # 사업명은 unmatched 로 보고
    assert any(u["normalized"] == "사업명" for u in report.unmatched_targets)


def test_autofill_preserves_originals(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _make_source(src)
    _make_target(tgt)
    src_before = src.read_bytes()
    tgt_before = tgt.read_bytes()

    autofill_from_source(src, tgt, out, use_ai=False)

    assert src.read_bytes() == src_before  # 소스 미수정
    assert tgt.read_bytes() == tgt_before  # 타깃 미수정
    # 타깃 값칸은 여전히 비어 있어야 함(원본 보존)
    assert _value_for(tgt, "성명") == ""


def test_autofill_out_equals_source_or_target_raises(tmp_path: Path) -> None:
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    _make_source(src)
    _make_target(tgt)
    with pytest.raises(ValueError):
        autofill_from_source(src, tgt, src, use_ai=False)
    with pytest.raises(ValueError):
        autofill_from_source(src, tgt, tgt, use_ai=False)


def test_autofill_conservative_no_unrelated_fill(tmp_path: Path) -> None:
    """보수적: 무관한 라벨('주소' 등 소스에 없는 항목)엔 값이 들어가지 않는다."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _make_source(src)
    # 타깃에 소스와 무관한 라벨을 추가
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "기업명"
    t.rows[0].cells[1].text = ""
    t.rows[1].cells[0].text = "추진목표"   # 소스에 없음, 동의어도 아님
    t.rows[1].cells[1].text = ""
    doc.save(str(tgt))

    report = autofill_from_source(src, tgt, out, use_ai=False)
    assert _value_for(out, "기업명") == "밸류업(주)"
    assert _value_for(out, "추진목표") == ""
    assert report.transcribed == 1


# --- 적대 반례: 오매칭 차단 --------------------------------------------------

def test_사업명_does_not_autofill_사업자명(tmp_path: Path) -> None:
    """접미사 다른 합성어: 소스 '사업명'이 타깃 '사업자명'으로 자동전사되면 안 된다.

    퍼지로 후보 제안(needs_confirm)은 가능하나, 전사·confident 0 이어야 한다.
    """
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "사업명"
    t.rows[0].cells[1].text = "AI 인재실증 사업"
    doc.save(str(src))
    doc2 = Document()
    t2 = doc2.add_table(rows=1, cols=2)
    t2.rows[0].cells[0].text = "사업자명"
    t2.rows[0].cells[1].text = ""
    doc2.save(str(tgt))

    report = autofill_from_source(src, tgt, out, use_ai=False)
    # 자동전사 안 됨 → 빈칸 유지
    assert _value_for(out, "사업자명") == ""
    assert report.transcribed == 0
    assert all(m.normalized != "사업자명" for m in report.matches)
    # 퍼지 후보로 needs_confirm 에 잡혀야 함(제안만)
    assert any(n["normalized"] == "사업자명" for n in report.needs_confirm)


def test_주소_does_not_autofill_주소지정(tmp_path: Path) -> None:
    """소스 '주소'가 타깃 '주소지정'으로 자동전사되면 안 된다(퍼지 → needs_confirm)."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "주소"
    t.rows[0].cells[1].text = "서울시 강남구"
    doc.save(str(src))
    doc2 = Document()
    t2 = doc2.add_table(rows=1, cols=2)
    t2.rows[0].cells[0].text = "주소지정"
    t2.rows[0].cells[1].text = ""
    doc2.save(str(tgt))

    report = autofill_from_source(src, tgt, out, use_ai=False)
    assert _value_for(out, "주소지정") == ""
    assert report.transcribed == 0
    assert any(n["normalized"] == "주소지정" for n in report.needs_confirm)


def test_cluster_conflict_goes_to_needs_confirm(tmp_path: Path) -> None:
    """동의어 클러스터 충돌: 소스에 '기업명'+'회사명' 둘 다 → 타깃 '상호'는 전사 보류.

    한 타깃에 같은 클러스터 소스가 복수면 conflict → 전사 0, needs_confirm.
    """
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "기업명"
    t.rows[0].cells[1].text = "밸류업(주)"
    t.rows[1].cells[0].text = "회사명"
    t.rows[1].cells[1].text = "밸류업 주식회사"
    doc.save(str(src))
    doc2 = Document()
    t2 = doc2.add_table(rows=1, cols=2)
    t2.rows[0].cells[0].text = "상호"   # 같은 클러스터(기업명/회사명/상호)
    t2.rows[0].cells[1].text = ""
    doc2.save(str(tgt))

    report = autofill_from_source(src, tgt, out, use_ai=False)
    assert _value_for(out, "상호") == ""
    assert report.transcribed == 0
    conf = next(n for n in report.needs_confirm if n["normalized"] == "상호")
    assert conf["confidence"] == "conflict"
    assert set(conf["candidates"]) == {"기업명", "회사명"}


def test_masked_value_preserved_after_transcribe(tmp_path: Path) -> None:
    """전사값 보존: 마스킹 이름 '○○○'·'OOO-OO-OOOOO' 가 지워지거나 망가지지 않는다."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "대표자"
    t.rows[0].cells[1].text = "○○○"
    t.rows[1].cells[0].text = "사업자등록번호"
    t.rows[1].cells[1].text = "OOO-OO-OOOOO"
    doc.save(str(src))
    doc2 = Document()
    t2 = doc2.add_table(rows=2, cols=2)
    t2.rows[0].cells[0].text = "성명"          # 동의어(대표자)
    t2.rows[0].cells[1].text = ""
    t2.rows[1].cells[0].text = "사업자등록번호"  # 정확일치
    t2.rows[1].cells[1].text = ""
    doc2.save(str(tgt))

    report = autofill_from_source(src, tgt, out, use_ai=False)
    # 마스킹 값이 원형 그대로 보존(blind-review 규칙 위반 없음)
    assert _value_for(out, "성명") == "○○○"
    assert _value_for(out, "사업자등록번호") == "OOO-OO-OOOOO"
    assert report.transcribed == 2


def test_four_column_label_value_row(tmp_path: Path) -> None:
    """4컬럼 label|value|label|value 행을 정확히 전사한다(짝수 인덱스만 라벨)."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    t = doc.add_table(rows=1, cols=4)
    t.rows[0].cells[0].text = "기업명"
    t.rows[0].cells[1].text = "밸류업(주)"
    t.rows[0].cells[2].text = "대표자"
    t.rows[0].cells[3].text = "홍길동"
    doc.save(str(src))
    fields = extract_source_fields(src)
    # 값 셀(밸류업/홍길동)이 라벨로 오염되지 않아야 함
    assert fields["기업명"] == "밸류업(주)"
    assert fields["대표자"] == "홍길동"
    assert "밸류업" not in fields
    assert "홍길동" not in fields

    doc2 = Document()
    t2 = doc2.add_table(rows=1, cols=4)
    t2.rows[0].cells[0].text = "기업명"
    t2.rows[0].cells[1].text = ""
    t2.rows[0].cells[2].text = "성명"   # 동의어(대표자)
    t2.rows[0].cells[3].text = ""
    doc2.save(str(tgt))

    report = autofill_from_source(src, tgt, out, use_ai=False)
    doc_out = Document(str(out))
    row = doc_out.tables[0].rows[0]
    assert row.cells[1].text.strip() == "밸류업(주)"
    assert row.cells[3].text.strip() == "홍길동"
    assert report.transcribed == 2


def test_extract_no_value_cell_pollution(tmp_path: Path) -> None:
    """추출 오염 없음: 값 셀이 라벨 키로 들어가지 않는다(짝수 인덱스만 라벨)."""
    src = tmp_path / "a.docx"
    doc = Document()
    t = doc.add_table(rows=1, cols=3)
    t.rows[0].cells[0].text = "항목"
    t.rows[0].cells[1].text = "연락처"
    t.rows[0].cells[2].text = "010-1234-5678"
    doc.save(str(src))
    fields = extract_source_fields(src)
    # 짝수 인덱스(0,2)만 라벨: (0,1)=항목→연락처. (2,3) 없음.
    # 값 셀 '연락처'(idx1)·'010-…'(idx2의 값 없음)이 라벨로 등록되면 안 됨.
    assert fields.get("항목") == "연락처"
    assert "연락처" not in fields  # 값 셀이 라벨로 오염 안 됨
    assert "010-1234-5678" not in fields


# ============================================================================
# 신규 회귀 테스트 (버그 헌트 C1/H1/H2/H3/H6/H7/M3/M5/M6 잠금)
# TDD: 수정 전 RED → 그룹 A→B→C 수정 후 GREEN.
# ============================================================================

import subprocess  # noqa: E402

CLI = Path(__file__).resolve().parents[1] / "cross_form_fill.py"
APP_DIR = Path(__file__).resolve().parents[1]


def _run_cli(source: str, target: str, out: str) -> "subprocess.CompletedProcess[str]":
    """cross_form_fill.py 를 subprocess 로 실행(CLI 계약 검증용)."""
    import os
    env = dict(os.environ)
    env["PYTHONIOENCODING"] = "utf-8"
    env["PYTHONPATH"] = str(APP_DIR)
    return subprocess.run(
        ["py", "-3.11", str(CLI),
         "--source", source, "--target", target, "-o", out],
        capture_output=True, text=True, encoding="utf-8", cwd=str(APP_DIR), env=env,
    )


# --- 그룹 A: C1/H1/M3 — 라벨이 값으로 둔갑하지 않음 ---------------------------

def test_column_major_budget_table_no_label_as_value(tmp_path: Path) -> None:
    """컬럼-메이저/행렬 예산표: 헤더행(라벨들)+데이터행(숫자) 구조에서
    extract 결과에 라벨이 '값'으로 들어가면 안 된다(C1/M3).

    행0=[국고보조금, 자기부담금, 총사업비] (모두 라벨)
    행1=[30,000,000, 10,000,000, 40,000,000] (모두 숫자값)
    짝수=라벨 위치단정이면 행0에서 국고보조금→자기부담금(라벨→라벨)으로 페어링됨.
    """
    src = tmp_path / "a.docx"
    doc = Document()
    t = doc.add_table(rows=2, cols=3)
    t.rows[0].cells[0].text = "국고보조금"
    t.rows[0].cells[1].text = "자기부담금"
    t.rows[0].cells[2].text = "총사업비"
    t.rows[1].cells[0].text = "30,000,000"
    t.rows[1].cells[1].text = "10,000,000"
    t.rows[1].cells[2].text = "40,000,000"
    doc.save(str(src))

    fields = extract_source_fields(src)
    # 라벨이 다른 라벨/숫자를 '값'으로 갖지 않아야 한다.
    assert fields.get("국고보조금") != "자기부담금"
    assert fields.get("총사업비") != "국고보조금"
    # 라벨 텍스트가 어떤 키의 값으로도 새어들면 안 됨.
    label_keys = {"국고보조금", "자기부담금", "총사업비"}
    for v in fields.values():
        assert v not in label_keys, f"라벨 '{v}' 이 값으로 추출됨"


def test_seed_labels_not_high_transcribed(tmp_path: Path) -> None:
    """실파일류 시드: 국고보조금/자기부담금/총사업비 칸에 라벨텍스트가
    high(자동전사)로 들어가면 안 된다(C1).

    소스 행렬 예산표 + 타깃에 동일 라벨 빈칸 → 라벨→값 전사가 일어나면 high 오전사.
    """
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    t = doc.add_table(rows=2, cols=3)
    t.rows[0].cells[0].text = "국고보조금"
    t.rows[0].cells[1].text = "자기부담금"
    t.rows[0].cells[2].text = "총사업비"
    t.rows[1].cells[0].text = "30,000,000"
    t.rows[1].cells[1].text = "10,000,000"
    t.rows[1].cells[2].text = "40,000,000"
    doc.save(str(src))

    doc2 = Document()
    t2 = doc2.add_table(rows=3, cols=2)
    for i, lab in enumerate(["국고보조금", "자기부담금", "총사업비"]):
        t2.rows[i].cells[0].text = lab
        t2.rows[i].cells[1].text = ""
    doc2.save(str(tgt))

    report = autofill_from_source(src, tgt, out, use_ai=False)
    # 라벨 텍스트가 high 로 전사되면 안 됨.
    for lab in ("국고보조금", "자기부담금", "총사업비"):
        v = _value_for(out, lab)
        assert v not in ("국고보조금", "자기부담금", "총사업비", "인증평가", "현금"), \
            f"'{lab}' 칸에 라벨텍스트 '{v}' 가 high 전사됨"


# --- 그룹 B: H2 중복 타깃 라벨 / H3 괄호 토큰 충돌 ---------------------------

def test_duplicate_target_label_only_one_high(tmp_path: Path) -> None:
    """동일 라벨 빈칸 2개 타깃 → 한 칸만 자동전사, 나머지는 needs_confirm(H2)."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "기업명"
    t.rows[0].cells[1].text = "밸류업(주)"
    doc.save(str(src))

    doc2 = Document()
    t2 = doc2.add_table(rows=2, cols=2)
    t2.rows[0].cells[0].text = "기업명"
    t2.rows[0].cells[1].text = ""
    t2.rows[1].cells[0].text = "기업명"   # 동일 라벨 중복 빈칸
    t2.rows[1].cells[1].text = ""
    doc2.save(str(tgt))

    report = autofill_from_source(src, tgt, out, use_ai=False)
    # 동일값이 두 칸 모두에 복제되면 안 됨 → 정확히 1건만 전사.
    assert report.transcribed == 1, f"중복 타깃에 {report.transcribed}건 전사(복제)"
    # 두 번째 기업명은 needs_confirm 로 강등되어야 함.
    assert any(n["normalized"] == "기업명" for n in report.needs_confirm)


def test_bracket_token_mismatch_not_high(tmp_path: Path) -> None:
    """금액(국고) 소스 + 금액(자부담) 타깃 → 괄호 토큰이 달라 high 전사 안 됨(H3)."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "금액(국고)"
    t.rows[0].cells[1].text = "60,000"
    doc.save(str(src))

    doc2 = Document()
    t2 = doc2.add_table(rows=1, cols=2)
    t2.rows[0].cells[0].text = "금액(자부담)"
    t2.rows[0].cells[1].text = ""
    doc2.save(str(tgt))

    report = autofill_from_source(src, tgt, out, use_ai=False)
    assert _value_for(out, "금액(자부담)") == "", "괄호 토큰 다른데 high 전사됨"
    assert report.transcribed == 0


# --- 이름(성명)류 필드 값-타입 가드 (실측 회귀: 역할서술 오전사) ----------------

def _name_target(norm: str = "대표자명") -> list[dict]:
    return [{
        "orig_label": norm, "normalized": norm, "kind": "table",
        "table_index": 0, "row": 0, "value_cell": 1, "para_index": -1,
    }]


def test_name_field_rejects_role_description_value() -> None:
    """이름(성명)류 타깃에 역할서술 값(콤마·'및'·'총괄')이 오면 high 금지·needs_confirm."""
    source = {"대표자": "기술개발, 특허전략 및 사업화 총괄"}  # 역할서술(이름 아님)
    m = match_fields(source, _name_target("대표자명"))[0]
    assert m.confidence != "high"           # 자동전사 금지(오매칭<빈칸)
    assert m.source_label == "" and m.value == ""
    assert "대표자" in m.candidates         # needs_confirm 후보로 노출(투명)


def test_name_field_accepts_real_name_value() -> None:
    """이름 모양 값(홍길동)은 정상 high 전사 — 회귀 보호."""
    m = match_fields({"대표자": "홍길동"}, _name_target("대표자명"))[0]
    assert m.confidence == "high" and m.value == "홍길동"


def test_name_field_accepts_masked_name_value() -> None:
    """마스킹된 이름(○○○)도 이름 모양으로 보고 전사 허용(보존)."""
    m = match_fields({"대표자": "○○○"}, _name_target("성명"))[0]
    assert m.confidence == "high" and m.value == "○○○"


def test_name_field_rejects_misc_separator_role_value() -> None:
    """ㆍ·세미콜론 등 드문 구분자로 나열된 역할서술도 이름칸 high 금지(G1)."""
    m = match_fields({"대표자": "전략기획ㆍ사업운영"}, _name_target("성명"))[0]
    assert m.confidence != "high"


def test_non_name_field_value_shape_unaffected() -> None:
    """비이름 필드(연락처)는 값에 콤마가 있어도 가드 영향 없음 — 기존 동작 보존."""
    source = {"연락처": "02-1234-5678, 내선 7"}
    targets = [{
        "orig_label": "연락처", "normalized": "연락처", "kind": "table",
        "table_index": 0, "row": 0, "value_cell": 1, "para_index": -1,
    }]
    m = match_fields(source, targets)[0]
    assert m.confidence == "high" and m.value.startswith("02-1234")


def test_role_description_paragraph_not_autofilled_to_name(tmp_path: Path) -> None:
    """실측 회귀(미래큐러스 A): 본문 '대표자 : <역할서술>' 이 타깃 '대표자명'에 자동전사 안 됨."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    da = Document()
    da.add_paragraph("⑤ 사업 수행 체계")
    da.add_paragraph("대표자 : 기술개발, 특허전략 및 사업화 총괄")
    da.save(str(src))
    db = Document()
    t = db.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "대표자명"
    t.rows[0].cells[1].text = ""
    db.save(str(tgt))

    report = autofill_from_source(src, tgt, out, use_ai=False)
    assert _value_for(out, "대표자명") == "", "역할서술이 이름칸에 자동전사됨"
    assert report.transcribed == 0
    assert any(nc["normalized"] == "대표자명" for nc in report.needs_confirm)


# --- 예시 플레이스홀더 빈칸승격 + 가짜타깃 필터 (recall, 실측) -------------------

from auto_write.services.cross_form_autofill import (  # noqa: E402
    _is_obvious_placeholder, _is_noise_label,
)


def test_is_obvious_placeholder_vectors() -> None:
    """적대검증 지정 벡터: 명백한 예시만 True, 그럴듯한 실값·영문 O단어는 False."""
    promote = ["2000.00.00.", "2000. 00. 00.", "0000.00.00", "2000년 00월00일",
               "000억원", "000명", "00건", "000-00-00000"]
    # 월·일 둘 다 0 일 때만 더미날짜 — 한쪽만 0/버전문자열은 보존(적대리뷰 #2)
    preserve = ["", "경기도 성남시", "서울시 강남구", "2025.03.17", "2025년 03월17일",
                "2020.10.05", "327-29-01754", "5억원", "100억원", "2,000명",
                "정규직 1명", "GOOGLE", "BOOK", "SOHO 창업", "O2O 플랫폼",
                "㈜OOO컴퍼니", "홍길동",
                "2000.01.00", "2000.00.15", "2020.0.5", "펌웨어 1234.0.1"]
    for v in promote:
        assert _is_obvious_placeholder(v) is True, f"승격 실패: {v!r}"
    for v in preserve:
        assert _is_obvious_placeholder(v) is False, f"실값 오판(덮어쓰기 위험): {v!r}"


def test_find_target_promotes_placeholder_cell(tmp_path: Path) -> None:
    """예시 플레이스홀더('2000.00.00.')가 든 칸도 채울 타깃으로 승격된다."""
    tgt = tmp_path / "b.docx"
    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "창업일"
    t.cell(0, 1).text = "2000.00.00."   # 예시 플레이스홀더
    doc.save(str(tgt))
    targets = find_target_fields(tgt)
    assert any(t["normalized"] == "창업일" for t in targets), "PH 칸이 타깃 승격 안 됨"


def test_find_target_preserves_real_value_cell(tmp_path: Path) -> None:
    """0 없는 그럴듯한 실값('경기도 성남시')은 타깃 아님(덮어쓰기 금지)."""
    tgt = tmp_path / "b.docx"
    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "사업장 소재지"
    t.cell(0, 1).text = "경기도 성남시"   # 그럴듯한 실값
    doc.save(str(tgt))
    targets = find_target_fields(tgt)
    assert not any(t["normalized"] == "사업장소재지" for t in targets), "실값 칸이 타깃됨"


def test_noise_label_filtered_but_real_label_kept() -> None:
    """표 번호·예시토큰은 가짜 라벨로 드롭, 동의어 클러스터 라벨은 보호(R7)."""
    assert _is_noise_label("2", "2") is True            # 표 번호
    assert _is_noise_label("3", "3") is True
    assert _is_noise_label("...", "") is True           # 생략기호
    assert _is_noise_label("000 대표", "000대표") is True  # 예시 인명/직책
    assert _is_noise_label("OOO", "ooo") is True
    # 진짜 라벨은 절대 드롭 안 함
    assert _is_noise_label("기업명", "기업명") is False
    assert _is_noise_label("대표자명", "대표자명") is False
    assert _is_noise_label("창업일", "창업일") is False
    assert _is_noise_label("협업희망 수요기업명", "협업희망수요기업명") is False


def test_noise_label_keeps_quantity_hint_labels() -> None:
    """예시힌트(00명·000억원)가 붙은 진짜 수량필드 라벨은 드롭하지 않는다(적대리뷰 #1, recall 보호)."""
    from auto_write.services.cross_form_autofill import _key
    for lab in ["고용 계획(00명)", "매출 목표(000억원)", "참여연구원 00명", "보유 특허 00건"]:
        assert _is_noise_label(lab, _key(lab)) is False, f"수량 라벨 오드롭: {lab!r}"


def test_placeholder_promotion_fills_from_source_e2e(tmp_path: Path) -> None:
    """실측 회귀: 소스 개업연월일 → 타깃 '창업일' PH칸 자동채움(동의어). 실값칸은 보존."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    da = Document()
    ta = da.add_table(rows=1, cols=2)
    ta.cell(0, 0).text = "개업연월일"
    ta.cell(0, 1).text = "2025.03.17"
    da.save(str(src))
    db = Document()
    tb = db.add_table(rows=2, cols=2)
    tb.cell(0, 0).text = "창업일"
    tb.cell(0, 1).text = "2000.00.00."        # PH → 채워져야 함
    tb.cell(1, 0).text = "사업장 소재지"
    tb.cell(1, 1).text = "경기도 성남시"        # 실값 → 보존
    db.save(str(tgt))

    report = autofill_from_source(src, tgt, out, use_ai=False)
    assert _value_for(out, "창업일") == "2025.03.17", "PH 칸이 소스값으로 안 채워짐"
    assert _value_for(out, "사업장 소재지") == "경기도 성남시", "실값이 덮어써짐!"
    assert report.transcribed == 1


def test_fill_double_gate_never_overwrites_real_value(tmp_path: Path) -> None:
    """2중 게이트: 어떤 경로로든 이미 실값이 든 칸은 set_cell_text 로 덮어쓰지 않는다."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    da = Document()
    ta = da.add_table(rows=1, cols=2)
    ta.cell(0, 0).text = "기업명"
    ta.cell(0, 1).text = "미래큐러스"
    da.save(str(src))
    db = Document()
    tb = db.add_table(rows=1, cols=2)
    tb.cell(0, 0).text = "기업명"
    tb.cell(0, 1).text = "㈜오토라이트"   # 이미 실값 있음 → 절대 덮어쓰기 금지
    db.save(str(tgt))

    report = autofill_from_source(src, tgt, out, use_ai=False)
    assert _value_for(out, "기업명") == "㈜오토라이트", "실값이 덮어써짐(불변원칙 위반)"
    assert report.transcribed == 0


# --- 그룹 C: H7 비지원 입력 / M6 미존재 출력 폴더 ---------------------------

def test_cli_unsupported_input_exit2_json(tmp_path: Path) -> None:
    """PDF/비지원 입력 → exit 2 + JSON 리포트(traceback/exit1 아님)(H7/M5)."""
    src = tmp_path / "a.docx"
    _make_source(src)
    pdf = tmp_path / "bad.pdf"
    pdf.write_bytes(b"%PDF-1.4\n%fake pdf content\n")
    out = tmp_path / "out.docx"

    res = _run_cli(str(src), str(pdf), str(out))
    assert res.returncode == 2, f"exit={res.returncode}, stderr={res.stderr[:400]}"
    # stdout 은 JSON 리포트여야 한다(raw traceback 금지).
    assert "Traceback" not in res.stdout
    import json as _json
    parsed = _json.loads(res.stdout)
    assert parsed["ok"] is False


def test_output_nested_dir_autocreated(tmp_path: Path) -> None:
    """미존재 중첩 출력 폴더 → 자동 생성 후 정상 저장(M6)."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "nope" / "deep" / "out.docx"
    _make_source(src)
    _make_target(tgt)

    report = autofill_from_source(src, tgt, out, use_ai=False)
    assert out.exists(), "중첩 출력 폴더가 자동 생성되지 않음"
    assert report.ok


# === 체크박스(선택칸) 자동 체크 — 사업자형태 □개인/□법인 ====================
# 실측(미래큐러스 A → 오토라이트 B): B 사업자형태 행 = [사업자 형태 | □ 개인 | □ 법인],
# A 사업자구분 = "개인사업자" → □ 개인을 ■ 개인으로 체크해야 함(현재 미구현).

def _make_checkbox_target(
    path: Path,
    label: str = "사업자 형태",
    options: tuple[str, ...] = ("□ 개인", "□ 법인"),
    trailing: tuple[str, str] | None = ("사업장 소재지", "경기도 성남시"),
) -> None:
    """타깃: [라벨 | □옵션1 | □옵션2 | (다른라벨 | 값)] 한 행 표."""
    doc = Document()
    ncols = 1 + len(options) + (2 if trailing else 0)
    t = doc.add_table(rows=1, cols=ncols)
    cells = t.rows[0].cells
    cells[0].text = label
    for i, opt in enumerate(options):
        cells[1 + i].text = opt
    if trailing:
        cells[1 + len(options)].text = trailing[0]
        cells[1 + len(options) + 1].text = trailing[1]
    doc.save(str(path))


def _make_bizform_source(path: Path, value: str = "개인사업자") -> None:
    """소스: 사업자 구분 = <value> (사업자형태 동의어)."""
    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "사업자 구분"
    t.rows[0].cells[1].text = value
    doc.save(str(path))


def _cell_text(docx_path: Path, ti: int, ri: int, ci: int) -> str:
    doc = Document(str(docx_path))
    return (doc.tables[ti].rows[ri].cells[ci].text or "").strip()


# --- 단위: 탐지 --------------------------------------------------------------

def test_find_checkbox_targets_detects_group(tmp_path: Path) -> None:
    """라벨 다음 연속 '□옵션' 셀들을 한 그룹으로 탐지(라벨·옵션 정규화)."""
    tgt = tmp_path / "b.docx"
    _make_checkbox_target(tgt)
    groups = find_checkbox_targets(tgt)
    assert len(groups) == 1
    g = groups[0]
    assert g["normalized"] == "사업자형태"
    opt_norms = [o["normalized"] for o in g["options"]]
    assert opt_norms == ["개인", "법인"]
    # 다음 라벨('사업장 소재지')은 옵션에 포함되지 않는다(그룹 경계).
    assert all("소재지" not in o["normalized"] for o in g["options"])


def test_find_checkbox_ignores_lone_box(tmp_path: Path) -> None:
    """옵션 텍스트 없는 단독 '□' 셀(매트릭스 체크칸)은 옵션으로 보지 않는다."""
    tgt = tmp_path / "b.docx"
    # [구분 | □ | □ | □] — 옵션명 없는 순수 박스 → 그룹 없음
    _make_checkbox_target(
        tgt, label="구분", options=("□", "□", "□"), trailing=None)
    groups = find_checkbox_targets(tgt)
    assert groups == []


def test_find_checkbox_no_box_no_group(tmp_path: Path) -> None:
    """체크박스 기호가 전혀 없으면 그룹 0(일반 라벨-값 표는 무관)."""
    tgt = tmp_path / "b.docx"
    _make_target(tgt)  # 일반 라벨|빈칸 표
    assert find_checkbox_targets(tgt) == []


# --- 단위: 매칭 --------------------------------------------------------------

def test_checkbox_match_개인사업자_checks_개인(tmp_path: Path) -> None:
    """소스 '개인사업자' → 옵션 '개인'만 매칭(법인 아님)."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    _make_bizform_source(src, "개인사업자")
    _make_checkbox_target(tgt)
    fields = extract_source_fields(src)
    groups = find_checkbox_targets(tgt)
    decided = match_checkbox_groups(fields, groups)
    assert len(decided) == 1
    d = decided[0]
    assert d["confidence"] == "high"
    assert d["checked_option_index"] == 0          # □ 개인
    assert d["source_value"] == "개인사업자"


def test_checkbox_match_법인사업자_checks_법인(tmp_path: Path) -> None:
    """역방향 안전: 소스 '법인사업자' → 옵션 '법인'만 매칭."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    _make_bizform_source(src, "법인사업자")
    _make_checkbox_target(tgt)
    decided = match_checkbox_groups(
        extract_source_fields(src), find_checkbox_targets(tgt))
    assert decided[0]["checked_option_index"] == 1  # □ 법인


def test_checkbox_no_source_no_check(tmp_path: Path) -> None:
    """소스에 사업자형태 정보가 없으면 어떤 옵션도 체크하지 않는다(날조 0)."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    _make_source(src)            # 기업명/대표자/사업자등록번호만(사업자형태 없음)
    _make_checkbox_target(tgt)
    decided = match_checkbox_groups(
        extract_source_fields(src), find_checkbox_targets(tgt))
    # 매칭 보류: checked_option_index 는 -1 (자동 체크 금지)
    assert decided[0]["checked_option_index"] == -1
    assert decided[0]["confidence"] != "high"


def test_checkbox_ambiguous_no_check(tmp_path: Path) -> None:
    """소스 값이 두 옵션 모두 포함하면(개인 및 법인) 보류(오매칭<빈칸)."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    _make_bizform_source(src, "개인 및 법인")
    _make_checkbox_target(tgt)
    decided = match_checkbox_groups(
        extract_source_fields(src), find_checkbox_targets(tgt))
    assert decided[0]["checked_option_index"] == -1


# --- E2E: autofill 기입 ------------------------------------------------------

def test_autofill_checkbox_e2e_marks_box(tmp_path: Path) -> None:
    """E2E: □ 개인 → ■ 개인, □ 법인 보존, 옵션 글자 보존, transcribed 카운트."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _make_bizform_source(src, "개인사업자")
    _make_checkbox_target(tgt)
    report = autofill_from_source(src, tgt, out)
    assert _cell_text(out, 0, 0, 1) == "■ 개인"      # 체크됨
    assert _cell_text(out, 0, 0, 2) == "□ 법인"      # 비매칭 옵션 보존
    assert report.checkbox_checked == 1


def test_autofill_checkbox_idempotent(tmp_path: Path) -> None:
    """이미 ■ 개인이면 그대로(멱등), 중복 체크·훼손 없음."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _make_bizform_source(src, "개인사업자")
    _make_checkbox_target(tgt, options=("■ 개인", "□ 법인"))
    report = autofill_from_source(src, tgt, out)
    assert _cell_text(out, 0, 0, 1) == "■ 개인"
    assert _cell_text(out, 0, 0, 2) == "□ 법인"
    assert report.checkbox_checked == 0              # 이미 체크 → 새 변경 0


def test_autofill_checkbox_disabled_flag(tmp_path: Path) -> None:
    """enable_checkbox=False 면 체크박스를 건드리지 않는다."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _make_bizform_source(src, "개인사업자")
    _make_checkbox_target(tgt)
    report = autofill_from_source(src, tgt, out, enable_checkbox=False)
    assert _cell_text(out, 0, 0, 1) == "□ 개인"      # 무변경
    assert report.checkbox_checked == 0


def test_autofill_checkbox_preserves_originals(tmp_path: Path) -> None:
    """원본 source·target 미수정(체크박스 경로도 불변)."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _make_bizform_source(src, "개인사업자")
    _make_checkbox_target(tgt)
    autofill_from_source(src, tgt, out)
    assert _cell_text(tgt, 0, 0, 1) == "□ 개인"      # 타깃 원본 보존
    assert _cell_text(src, 0, 0, 1) == "개인사업자"  # 소스 원본 보존


# --- 적대검증 FIX_FIRST 회귀: 부분문자열 오체크 차단 -------------------------

def test_checkbox_substring_no_false_check(tmp_path: Path) -> None:
    """옵션명이 무관한 값의 부분문자열이어도 오체크 금지(적대검증 HIGH #1).

    '개인정보보호'·'법인영업 중심'은 옵션 '개인'/'법인'의 상위문자열이지만
    의미가 다르다 → 정확일치(환원 후) 아니므로 보류(-1)여야 한다.
    """
    tgt = tmp_path / "b.docx"
    _make_checkbox_target(tgt)
    groups = find_checkbox_targets(tgt)
    for bad in ("개인정보보호", "법인영업 중심의 사업 운영"):
        decided = match_checkbox_groups({"사업자형태": bad}, groups)
        assert decided[0]["checked_option_index"] == -1, f"{bad!r} 오체크"
        assert decided[0]["confidence"] != "high"


def test_checkbox_short_value_no_false_check(tmp_path: Path) -> None:
    """짧은 소스 값이 긴 옵션의 부분문자열이어도 오체크 금지(적대검증 HIGH #2).

    '소'(값)는 '중소기업'의 일부지만 사업자 구분이 아니다 → 보류(-1).
    """
    tgt = tmp_path / "b.docx"
    _make_checkbox_target(tgt, label="기업 구분",
                          options=("□ 중소기업", "□ 대기업"), trailing=None)
    groups = find_checkbox_targets(tgt)
    decided = match_checkbox_groups({"기업구분": "소"}, groups)
    assert decided[0]["checked_option_index"] == -1


def test_checkbox_placeholder_value_no_check(tmp_path: Path) -> None:
    """예시 플레이스홀더 값('00법인'/'000개인')은 체크 금지(적대검증 HIGH #3·날조 0)."""
    tgt = tmp_path / "b.docx"
    _make_checkbox_target(tgt)
    groups = find_checkbox_targets(tgt)
    for ph in ("00법인", "000개인", "○○○"):
        decided = match_checkbox_groups({"사업자형태": ph}, groups)
        assert decided[0]["checked_option_index"] == -1, f"{ph!r} 오체크"


def test_checkbox_real_value_still_checks(tmp_path: Path) -> None:
    """보수화 후에도 정상 값('개인사업자'/'법인')은 그대로 체크(양성 보존)."""
    tgt = tmp_path / "b.docx"
    _make_checkbox_target(tgt)
    groups = find_checkbox_targets(tgt)
    assert match_checkbox_groups({"사업자형태": "개인사업자"}, groups)[0]["checked_option_index"] == 0
    assert match_checkbox_groups({"사업자형태": "주식회사"}, groups)[0]["checked_option_index"] == 1
    assert match_checkbox_groups({"사업자형태": "법인"}, groups)[0]["checked_option_index"] == 1


def test_check_option_cell_hyperlink_wrapped_box(tmp_path: Path) -> None:
    """<w:hyperlink>로 감싸인 run 안의 □ 도 체크된다(적대검증 LOW #4·조용한 no-op 방지)."""
    from docx.oxml import OxmlElement
    from auto_write.services.cross_form_autofill import _check_option_cell
    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "사업자 형태"
    cell = t.rows[0].cells[1]
    para = cell.paragraphs[0]
    hl = OxmlElement("w:hyperlink")
    run = OxmlElement("w:r")
    wt = OxmlElement("w:t")
    wt.text = "□ 개인"
    run.append(wt)
    hl.append(run)
    para._p.append(hl)
    assert _check_option_cell(cell) is True
    assert "■ 개인" in cell.text


# ============================================================================
# 배치 모드: 공고 폴더 일괄 채우기 + HWP (F = #1 + #3)
# ============================================================================

from auto_write.services.cross_form_autofill import (  # noqa: E402
    BatchAutofillReport,
    batch_autofill_from_pool,
    discover_form_targets,
    format_batch_summary_korean,
    is_skipped_non_form,
    pick_source_from_pool,
    try_convert_filled_docx_to_hwp,
)

BATCH_CLI = Path(__file__).resolve().parents[1] / "cross_form_fill.py"


def _run_batch_cli(
    notice: str, pool: str, out_subdir: str = "filled", extra: list[str] | None = None,
) -> "subprocess.CompletedProcess[str]":
    import os
    env = dict(os.environ)
    env["PYTHONIOENCODING"] = "utf-8"
    env["PYTHONPATH"] = str(APP_DIR)
    cmd = [
        "py", "-3.11", str(BATCH_CLI), "batch",
        "--notice-folder", notice,
        "--source-pool", pool,
        "-o", out_subdir,
    ]
    if extra:
        cmd.extend(extra)
    return subprocess.run(
        cmd, capture_output=True, text=True, encoding="utf-8",
        cwd=str(APP_DIR), env=env,
    )


def test_is_skipped_non_form_announcement() -> None:
    assert is_skipped_non_form(Path("2026_모집공고.hwp")) is True
    assert is_skipped_non_form(Path("신청서.hwp")) is False


def test_discover_form_targets_skips_announcement(tmp_path: Path) -> None:
    (tmp_path / "공고문.docx").write_bytes(b"x")
    tgt = tmp_path / "참가신청서.docx"
    _make_target(tgt)
    found = discover_form_targets(tmp_path)
    assert len(found) == 1
    assert found[0].name == "참가신청서.docx"


def test_pick_source_from_pool_keyword_and_mtime(tmp_path: Path) -> None:
    pool = tmp_path / "pool"
    pool.mkdir()
    old = pool / "old_메모.docx"
    new = pool / "new_메모.docx"
    best = pool / "STAR_사업계획서.docx"
    for p in (old, new, best):
        _make_source(p)
    # mtime: best oldest but keyword + dry-run wins
    import os
    import time
    os.utime(old, (time.time() - 100, time.time() - 100))
    os.utime(new, (time.time(), time.time()))
    picked = pick_source_from_pool(pool, Path("b.docx"))
    assert picked is not None
    assert picked.name == "STAR_사업계획서.docx"


def test_pick_source_dry_run_per_target(tmp_path: Path) -> None:
    """dry-run: 필드가 맞는 소스를 타깃마다 고른다(최신 mtime·파일명만으로는 오선택 방지)."""
    from auto_write.services.cross_form_autofill import dry_run_transcribe_score

    pool = tmp_path / "pool"
    pool.mkdir()
    good = pool / "z_last_완성.docx"
    bad = pool / "a_first_신청서.docx"
    _make_source(good)
    # bad: 기업명만 있는 빈약 소스
    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "비고"
    t.rows[0].cells[1].text = "메모"
    doc.save(str(bad))

    tgt = tmp_path / "b.docx"
    _make_target(tgt)

    import os
    import time
    os.utime(bad, (time.time(), time.time()))
    os.utime(good, (time.time() - 3600, time.time() - 3600))

    assert dry_run_transcribe_score(good, tgt) >= 3
    assert dry_run_transcribe_score(bad, tgt) == 0
    picked = pick_source_from_pool(pool, tgt)
    assert picked is not None
    assert picked.name == "z_last_완성.docx"


def test_batch_autofill_fills_multiple_targets(tmp_path: Path) -> None:
    pool = tmp_path / "pool"
    notice = tmp_path / "notice"
    pool.mkdir()
    notice.mkdir()
    _make_source(pool / "완성_사업계획서.docx")

    t1 = notice / "신청서1.docx"
    t2 = notice / "신청서2.docx"
    _make_target(t1)
    _make_target(t2)

    report = batch_autofill_from_pool(notice, pool, "filled", convert_hwp=False)
    assert report.ok_count == 2
    out_dir = notice / "filled"
    assert (out_dir / "신청서1_filled.docx").exists()
    assert (out_dir / "신청서2_filled.docx").exists()
    assert report.items[0].transcribed >= 1


def test_format_batch_summary_korean() -> None:
    from auto_write.services.cross_form_autofill import BatchAutofillItem
    report = BatchAutofillReport(
        notice_folder="/n", source_pool="/p", output_dir="/n/filled",
        items=[
            BatchAutofillItem(
                target="a.docx", source="s.docx", output="/n/filled/a_filled.docx",
                ok=True, transcribed=3, needs_confirm_count=2, hwp_ok=True,
            ),
        ],
    )
    text = format_batch_summary_korean(report)
    assert "양식 1개 채움" in text
    assert "2칸 확인필요" in text
    assert "HWP 1개 생성" in text


def test_try_convert_hwp_skip_when_com_fails(tmp_path: Path, monkeypatch) -> None:
    docx = tmp_path / "out.docx"
    _make_source(docx)

    from auto_write.services.hwp_docx_convert import ConvertReport
    import auto_write.services.hwp_docx_convert as hdc

    def _fail(*_a, **_k):
        return ConvertReport(direction="docx->hwp", ok=False, notes=["COM 없음"])

    monkeypatch.setattr(hdc, "docx_to_hwp", _fail)

    hwp, notes = try_convert_filled_docx_to_hwp(docx)
    assert hwp is None
    assert notes


def test_batch_cli_korean_stdout_not_json_only(tmp_path: Path) -> None:
    pool = tmp_path / "pool"
    notice = tmp_path / "notice"
    pool.mkdir()
    notice.mkdir()
    _make_source(pool / "완성.docx")
    tgt = notice / "양식.docx"
    _make_target(tgt)

    res = _run_batch_cli(str(notice), str(pool), extra=["--no-hwp"])
    assert res.returncode == 0
    assert "양식" in res.stdout
    assert "채움" in res.stdout
    # 기본은 JSON only 아님
    assert not res.stdout.strip().startswith("{")


def test_batch_cli_hwp_skip_message(tmp_path: Path, monkeypatch) -> None:
    pool = tmp_path / "pool"
    notice = tmp_path / "notice"
    pool.mkdir()
    notice.mkdir()
    _make_source(pool / "완성.docx")
    tgt = notice / "양식.docx"
    _make_target(tgt)

    import auto_write.services.hwp_docx_convert as hdc
    from auto_write.services.hwp_docx_convert import ConvertReport

    def _fail(*_a, **_k):
        return ConvertReport(direction="docx->hwp", ok=False, notes=["COM 없음"])

    monkeypatch.setattr(hdc, "docx_to_hwp", _fail)

    res = _run_batch_cli(str(notice), str(pool))
    assert res.returncode == 0
    assert "DOCX" in res.stdout or "HWP" in res.stdout
    assert (notice / "filled" / "양식_filled.docx").exists()


# --- 표 셀 안 인라인 빈칸("라벨 : ______") 전사 (recall 확대) ------------------
# 정부양식 표지/개요 박스는 한 셀 안에 "성명 : ___   생년월일 : ___" 처럼 인라인 필드를
# 담는다. python-docx 의 doc.paragraphs 는 표 셀 단락을 포함하지 않아 기존엔 이 빈칸이
# 전혀 탐지·전사되지 않았다(전에는 못 채우던 칸). kind="cell_paragraph" 로 채운다.

def _cell_text(docx_path: Path, ti: int, ri: int, ci: int) -> str:
    """타깃 문서의 특정 표 셀 텍스트를 읽는다(셀 인라인 전사 검증용)."""
    doc = Document(str(docx_path))
    return doc.tables[ti].cell(ri, ci).text


def _make_inline_cell_target(path: Path) -> None:
    """타깃: 한 셀 안에 인라인 '라벨 : 밑줄빈칸' 이 여러 개인 개요 박스(1x1 표)."""
    doc = Document()
    doc.add_heading("지원신청서", 0)
    t = doc.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "성명 : ______    생년월일 : ______"
    doc.save(str(path))


def test_find_target_fields_cell_inline(tmp_path: Path) -> None:
    """표 셀 안 인라인 '라벨 : 빈칸' 이 cell_paragraph 타깃으로 탐지된다."""
    tgt = tmp_path / "b.docx"
    _make_inline_cell_target(tgt)
    targets = find_target_fields(tgt)
    by_norm = {t["normalized"]: t for t in targets}
    assert "성명" in by_norm
    assert "생년월일" in by_norm
    assert by_norm["성명"]["kind"] == "cell_paragraph"
    # 좌표: 0번 표, 0행, 0번 논리셀, 그 셀 0번 단락
    assert by_norm["성명"]["table_index"] == 0
    assert by_norm["성명"]["row"] == 0
    assert by_norm["성명"]["value_cell"] == 0
    assert by_norm["성명"]["cell_para_index"] == 0


def test_autofill_cell_inline_fills(tmp_path: Path) -> None:
    """소스 값이 표 셀 안 인라인 빈칸에 전사된다(동의어 성명←대표자 포함)."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"

    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "대표자"
    t.cell(0, 1).text = "홍길동"
    t.cell(1, 0).text = "생년월일"
    t.cell(1, 1).text = "1990-01-01"
    doc.save(str(src))

    _make_inline_cell_target(tgt)

    report = autofill_from_source(src, tgt, out, use_ai=False)
    assert report.ok
    filled = _cell_text(out, 0, 0, 0)
    assert "홍길동" in filled           # 성명 ← 대표자(동의어)
    assert "1990-01-01" in filled       # 생년월일 정확일치
    # 라벨·콜론은 보존
    assert "성명" in filled and "생년월일" in filled
    # 밑줄 빈칸은 사라짐(값으로 대체)
    assert "______" not in filled
    assert report.transcribed == 2


def test_autofill_cell_inline_no_fabrication(tmp_path: Path) -> None:
    """소스에 없는 인라인 라벨은 빈칸으로 남는다(날조 0)."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"

    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "대표자"
    t.cell(0, 1).text = "홍길동"
    doc.save(str(src))

    _make_inline_cell_target(tgt)  # 성명·생년월일 (생년월일은 소스에 없음)

    report = autofill_from_source(src, tgt, out, use_ai=False)
    filled = _cell_text(out, 0, 0, 0)
    assert "홍길동" in filled                 # 성명 채움
    # 생년월일은 소스에 없음 → 밑줄 빈칸 유지(지어내지 않음)
    seg = filled.split("생년월일")[1]
    assert "____" in seg or "______" in seg


def test_autofill_cell_inline_preserves_filled_value(tmp_path: Path) -> None:
    """이미 값이 든 인라인 필드는 덮어쓰지 않는다(마스킹/실값 보존)."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"

    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "대표자"
    t.cell(0, 1).text = "홍길동"
    doc.save(str(src))

    dt = Document()
    tt = dt.add_table(rows=1, cols=1)
    # 성명은 이미 채워짐(마스킹) → 보존, 생년월일만 빈칸
    tt.cell(0, 0).text = "성명 : ○○○    생년월일 : ______"
    dt.save(str(tgt))

    autofill_from_source(src, tgt, out, use_ai=False)
    filled = _cell_text(out, 0, 0, 0)
    assert "○○○" in filled          # 이미 채워진 마스킹 성명 보존(덮어쓰기 금지)
    assert "홍길동" not in filled     # 실값이 있는 칸은 안 건드림


def test_autofill_cell_inline_original_untouched(tmp_path: Path) -> None:
    """셀 인라인 전사 후에도 원본 소스·타깃 바이트가 그대로다(원본 미수정)."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"

    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "대표자"
    t.cell(0, 1).text = "홍길동"
    doc.save(str(src))
    _make_inline_cell_target(tgt)

    src_before = src.read_bytes()
    tgt_before = tgt.read_bytes()
    autofill_from_source(src, tgt, out, use_ai=False)
    assert src.read_bytes() == src_before
    assert tgt.read_bytes() == tgt_before
