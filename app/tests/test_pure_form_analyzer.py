"""test_pure_form_analyzer.py — form_analyzer 순수 파생·분기 안전망.

양식 분석 리포트(FormReport)의 순수 파생(psst_missing 계산·직렬화 스키마)과
analyze_form 의 '파일 없음' 조기반환 분기를 외부 의존(COM/네트워크) 없이 고정한다.
psst_missing 은 psst_present(True/False 맵)에서 '없는 영역만' 사람이 읽는 라벨로
뽑아내는 파생값이라, 완성본에 있던 값이 새 양식에서 빠졌는지 판단하는 근거가 된다.
야간 순수함수 안전망(2026-07-13).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from auto_write.services.form_analyzer import (
    FormReport,
    _PSST_LABELS,
    analyze_form,
)


# --- FormReport.as_dict: psst_missing 파생(순수, I/O 없음) ------------------

def test_as_dict_derives_psst_missing_in_area_order():
    # 일부만 존재 → 없는 영역만 라벨로 나열(문제·성장은 있고, 실현·팀은 없음).
    rep = FormReport(psst_present={
        "problem": True,
        "solution": False,
        "scale": True,
        "team": False,
    })
    assert rep.as_dict()["psst_missing"] == [
        _PSST_LABELS["solution"],
        _PSST_LABELS["team"],
    ]


def test_as_dict_all_present_has_no_missing():
    rep = FormReport(psst_present={k: True for k in _PSST_LABELS})
    assert rep.as_dict()["psst_missing"] == []


def test_as_dict_all_absent_lists_every_label():
    rep = FormReport(psst_present={k: False for k in _PSST_LABELS})
    assert rep.as_dict()["psst_missing"] == list(_PSST_LABELS.values())


def test_as_dict_empty_present_map_yields_empty_missing():
    # psst 판정 자체를 못 한 경우(빈 맵) — KeyError 없이 빈 목록.
    assert FormReport().as_dict()["psst_missing"] == []


def test_as_dict_exposes_full_schema_keys():
    rep = FormReport(template_name="양식.docx", table_count=2, question_count=5)
    d = rep.as_dict()
    for key in (
        "template_name", "source_docx", "section_count", "table_count",
        "image_slot_count", "question_count", "required_question_count",
        "required_cell_count", "psst_present", "psst_missing",
        "writable_items", "analysis_notes",
    ):
        assert key in d
    assert d["template_name"] == "양식.docx"
    assert d["table_count"] == 2
    assert d["question_count"] == 5


# --- analyze_form: '파일 없음' 조기반환(변환·분석 진입 전, 원본 미접근) -------

def test_analyze_form_missing_file_returns_note_and_empty_report(tmp_path: Path):
    p = tmp_path / "존재하지않는양식.docx"
    rep = analyze_form(p)
    assert rep.template_name == p.name
    assert any("파일이 없습니다" in n for n in rep.analysis_notes)
    # 조기반환이라 어떤 구조 카운트도 잡히지 않는다(0) + PSST 판정 없음.
    assert rep.section_count == 0
    assert rep.question_count == 0
    assert rep.table_count == 0
    assert rep.psst_present == {}
    assert rep.as_dict()["psst_missing"] == []


# --- analyze_form: 일부 PSST 영역만 있는 실제 양식 → psst_missing 도출 -------

def _make_partial_form(path: Path, headings: list[str]) -> None:
    doc = Document()
    doc.add_heading("사업계획서 양식", 0)
    for h in headings:
        doc.add_heading(h, level=1)
        doc.add_paragraph("(작성)")
    doc.save(str(path))


def test_analyze_form_partial_psst_marks_absent_areas(tmp_path: Path):
    # 문제인식·팀구성만 있는 양식 → 실현가능성·성장전략은 빠진 것으로 잡혀야 한다.
    p = tmp_path / "부분양식.docx"
    _make_partial_form(p, ["1. 문제인식(Problem)", "2. 팀구성(Team)"])
    rep = analyze_form(p)
    assert rep.psst_present.get("problem") is True
    assert rep.psst_present.get("team") is True
    assert rep.psst_present.get("solution") is False
    assert rep.psst_present.get("scale") is False
    missing = rep.as_dict()["psst_missing"]
    assert _PSST_LABELS["solution"] in missing
    assert _PSST_LABELS["scale"] in missing
    assert _PSST_LABELS["problem"] not in missing
    assert _PSST_LABELS["team"] not in missing
