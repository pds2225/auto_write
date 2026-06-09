"""문서 품질 개선 하네스 회귀 테스트.

사용자 정의 필수 테스트(Phase 13):
  1) 글머리표 공백 정리  2) 표 내부 공백 정리  3) 안내문구 삭제 규칙
  4) 문서 유형 분류      5) PSST 검사          6) 이미지 제안 리포트
  7) 품질점수 산정       8) 백업 생성          9) 기존 기능 비훼손(import)

실행: (app 디렉토리 기준)  python -m pytest tests/test_document_quality_harness.py -q
"""
from __future__ import annotations

from pathlib import Path

from docx import Document

from auto_write.services import doc_quality_ops as dq
from auto_write.services.document_type_classifier import classify_text
from auto_write.services.psst_check import check_psst
from auto_write.services.infographic_suggest import suggest_images
from auto_write.services.doc_quality_score import score_document
from auto_write.services.document_quality_orchestrator import DocumentQualityOrchestrator


# --------------------------------------------------------------------------- helpers
def _plan_doc() -> Document:
    d = Document()
    d.add_paragraph("1. 문제인식 (Problem)")
    d.add_paragraph("○   고객 시장 문제: 기존 대안의 한계와 심각성, 비용 30% 증가")
    d.add_paragraph("※ 작성요령: 여기에 기재하세요")
    d.add_paragraph("")
    d.add_paragraph("")
    d.add_paragraph("2. 실현가능성 (Solution)")
    d.add_paragraph("ㅇ 핵심기능과 차별성을 구현하여 고객사 적용 시나리오 검증")
    d.add_paragraph("3. 성장전략 (Scale-up)")
    d.add_paragraph("· 시장규모 TAM 5000억, 수익모델 구독, 판로 확대, 매출 100억 KPI")
    d.add_paragraph("4. 팀구성 (Team)")
    d.add_paragraph("- 대표 경력과 팀 구성, 외부 협력 파트너, 수행 경험 보유")
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "  구분  "
    t.cell(0, 1).text = "내용   많은    공백"
    return d


# --------------------------------------------------------------------------- 1
def test_bullet_spacing_normalization():
    d = Document()
    d.add_paragraph("○    항목 하나")
    d.add_paragraph("ㅇ  항목   둘")
    fixed = dq.normalize_bullet_spacing(d)
    assert fixed >= 2
    # 다중 공백이 사라졌는지
    assert "    " not in d.paragraphs[0].text


# --------------------------------------------------------------------------- 2
def test_table_whitespace_cleanup():
    d = Document()
    t = d.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "  앞뒤  공백   많음  "
    cleaned = dq.cleanup_table_whitespace(d)
    assert cleaned == 1
    assert t.cell(0, 0).text == "앞뒤 공백 많음"


# --------------------------------------------------------------------------- 3
def test_guide_paragraph_removal():
    d = Document()
    d.add_paragraph("※ 작성요령: 여기에 기재")
    d.add_paragraph("실제 사업 내용입니다.")
    removed = dq.remove_guide_paragraphs(d)
    assert removed == 1
    remaining = [p.text for p in d.paragraphs if p.text.strip()]
    assert "실제 사업 내용입니다." in remaining


def test_table_guide_row_removal():
    """표 셀에 박힌 양식 안내문구 삭제: 안내전용표는 통째, 혼합표는 안내행만, 데이터표는 보존."""
    d = Document()
    # 1) 안내 전용 표 → 표 통째 삭제 대상
    t1 = d.add_table(rows=1, cols=1)
    t1.rows[0].cells[0].text = "※ 정부지원사업비는 최대 2억원 한도 이내로 작성"
    # 2) 데이터 표 → 보존돼야 함
    t2 = d.add_table(rows=2, cols=3)
    t2.rows[0].cells[0].text = "비목"
    t2.rows[0].cells[1].text = "집행 계획"
    t2.rows[0].cells[2].text = "금액"
    t2.rows[1].cells[0].text = "인건비"
    t2.rows[1].cells[1].text = "AI 전문인력 인건비"
    t2.rows[1].cells[2].text = "10,000,000"
    # 3) 혼합 표 → 안내 행만 삭제, 데이터 행 보존
    t3 = d.add_table(rows=2, cols=2)
    t3.rows[0].cells[0].text = "작성요령 : 예시를 삭제하고 작성"
    t3.rows[1].cells[0].text = "1"
    t3.rows[1].cells[1].text = "실제 데이터 행"

    before_tables = len(d.tables)
    removed = dq.remove_table_guide_rows(d)

    # 안내 전용 표(t1)는 통째 삭제 → 표 1개 감소
    assert len(d.tables) == before_tables - 1
    texts = [c.text for tb in d.tables for r in tb.rows for c in r.cells]
    # 데이터 표 보존
    assert "인건비" in texts and "10,000,000" in texts
    # 혼합 표: 데이터 행 보존 + 안내 행 제거
    assert "실제 데이터 행" in texts
    assert not any("작성요령" in x for x in texts)
    # 안내 전용 표 + 혼합표 안내행 = 최소 2건 제거
    assert removed >= 2


def test_empty_paragraph_removal():
    d = Document()
    d.add_paragraph("내용")
    d.add_paragraph("")
    d.add_paragraph("")
    d.add_paragraph("")
    d.add_paragraph("끝")
    removed = dq.remove_empty_paragraphs(d)
    assert removed >= 1


# --------------------------------------------------------------------------- 4
def test_document_type_classification():
    text = "사업계획서\n1. 문제인식 Problem\n창업아이템 PSST 성장전략 사업화"
    r = classify_text(text, filename="사업계획서.docx")
    assert r.type_code == "business_plan"
    assert r.confidence > 0.5

    rnd = classify_text("연구개발계획서 기술개발목표 TRL 성능지표 실험방법", filename="rnd.docx")
    assert rnd.type_code == "rnd_plan"

    generic = classify_text("안녕하세요 오늘 날씨가 좋습니다", filename="memo.docx")
    assert generic.type_code == "generic_submission"


def test_result_report_classification():
    """사업 결과·성과·정산 보고서 유형 인식 + 사업계획서 오분류 회귀 가드."""
    text = (
        "2024년 창업지원사업 결과보고서\n"
        "1. 사업 추진실적 및 목표 달성도\n"
        "2. 주요 성과보고 및 성과지표\n"
        "3. 사업비 집행실적 및 정산 내역(사업비정산)\n"
    )
    r = classify_text(text, filename="결과보고서.docx")
    assert r.type_code == "result_report"
    assert r.confidence > 0.5

    # 사업계획서가 결과보고서로 오분류되지 않는지(회귀 가드)
    plan = classify_text(
        "사업계획서\n문제인식 실현가능성 성장전략 팀구성 창업아이템 PSST",
        filename="plan.docx",
    )
    assert plan.type_code == "business_plan"


# --------------------------------------------------------------------------- 5
def test_psst_check():
    report = check_psst(_plan_doc())
    assert report.applicable
    assert len(report.areas) == 4
    assert report.overall_ratio > 0.5
    # 모든 영역 섹션 헤더 인식
    assert all(a.section_present for a in report.areas)


# --------------------------------------------------------------------------- 6
def test_infographic_suggestion():
    report = suggest_images(_plan_doc())
    assert len(report.suggestions) >= 1
    # as_dict 의 suggestion_count 가 실제 제안 수와 일치하는지(직렬화 정합성)
    assert report.as_dict()["suggestion_count"] == len(report.suggestions)
    # 시각화 유형 중복 없음
    vtypes = [s.visual_type for s in report.suggestions]
    assert len(vtypes) == len(set(vtypes))


# --------------------------------------------------------------------------- 7
def test_quality_scoring():
    d = _plan_doc()
    dq.run_all(d)
    score = score_document(
        d, doc_type="business_plan", type_confidence=0.9,
        psst_ratio=1.0, image_suggestions=3, existing_images=0,
    )
    assert 0 <= score.total <= 100
    assert len(score.items) == 9
    assert sum(i.max_score for i in score.items) == 100


# --------------------------------------------------------------------------- 8
def test_backup_and_full_run(tmp_path: Path):
    src = tmp_path / "sample.docx"
    _plan_doc().save(str(src))
    orch = DocumentQualityOrchestrator(tmp_path / "results")
    res = orch.run(src)

    # 백업 생성
    assert Path(res.backup_dir).exists()
    assert list(Path(res.backup_dir).glob("*.docx"))
    # 원본 비훼손 + 출력 분리
    assert Path(res.output_docx).exists()
    assert Path(res.output_docx).resolve() != src.resolve()
    # 리포트 생성
    assert Path(res.report_md).exists()
    assert Path(res.report_json).exists()
    # 점수/유형
    assert res.doc_type.type_code == "business_plan"
    assert 0 <= res.score.total <= 100

    # 롤백
    target = tmp_path / "restored.docx"
    assert DocumentQualityOrchestrator.rollback(res.backup_dir, target)
    assert target.exists()


def test_output_never_overwrites_input(tmp_path: Path):
    src = tmp_path / "in.docx"
    _plan_doc().save(str(src))
    orch = DocumentQualityOrchestrator(tmp_path / "results")
    import pytest
    with pytest.raises(ValueError):
        orch.run(src, src)  # 동일 경로 → 거부


# --------------------------------------------------------------------------- 9
def test_existing_modules_still_import():
    # 기존 핵심 서비스가 여전히 import 되는지(하네스 추가로 깨지지 않았는지)
    from auto_write.services import docx_ops, qa_service, project_service, render_service  # noqa
    from auto_write.services import evaluation_service, submittable_filler  # noqa
    assert hasattr(qa_service.QAService, "build_report")
    assert hasattr(project_service.ProjectService, "PSST_PROBLEM_RE")


# --------------------------------------------------------------------------- STEP1: 안내문구 삭제 강화 (item 3)
def test_guide_extra_imperative_removal():
    """글머리표 접두 + 명령형 어미 안내문구를 삭제하고 실제 본문은 보존한다."""
    d = Document()
    d.add_paragraph("○ 해당 내용을 구체적으로 기재하시오")   # 글머리표 접두 + 하시오
    d.add_paragraph("사업 목표를 작성하세요")                  # 명령형 어미(접두 없음)
    d.add_paragraph("핵심 인력 정보를 입력 바랍니다")          # 바랍니다
    d.add_paragraph("실제 사업 내용입니다.")                   # 본문 — 보존돼야 함
    removed = dq.remove_guide_paragraphs(d)
    remaining = [p.text for p in d.paragraphs if p.text.strip()]
    assert removed == 3
    assert remaining == ["실제 사업 내용입니다."]


def test_guide_no_false_delete():
    """선언형 '기재한/작성한다' 와 정량 본문은 절대 삭제하지 않는다(과삭제 트립와이어)."""
    d = Document()
    d.add_paragraph("본 사업은 매출 목표를 기재한 바와 같이 달성한다")  # 선언형(기재한)
    d.add_paragraph("매출 10억원 달성 및 고용 20명 창출")              # 정량 본문
    d.add_paragraph("사업계획을 작성하였으며 검토를 마쳤다")            # 선언형(작성하였)
    before = [p.text for p in d.paragraphs]
    removed = dq.remove_guide_paragraphs(d)
    after = [p.text for p in d.paragraphs if p.text.strip()]
    assert removed == 0
    assert after == before


def test_runall_wires_table_guides():
    """run_all 이 표 안내문구도 제거하고(배선 확인) 데이터 표는 보존한다."""
    d = Document()
    d.add_paragraph("실제 본문")
    # 안내 전용 표 → 통째 제거
    t1 = d.add_table(rows=1, cols=1)
    t1.rows[0].cells[0].text = "※ 정부지원사업비는 최대 2억원 한도 이내로 작성"
    # 데이터 표 → 보존
    t2 = d.add_table(rows=2, cols=2)
    t2.rows[0].cells[0].text = "비목"
    t2.rows[0].cells[1].text = "금액"
    t2.rows[1].cells[0].text = "인건비"
    t2.rows[1].cells[1].text = "10,000,000"
    before_tables = len(d.tables)
    report = dq.run_all(d)
    assert report.table_guide_rows_removed >= 1          # run_all 에서 실제 동작
    assert len(d.tables) == before_tables - 1            # 안내 전용 표 제거
    texts = [c.text for tb in d.tables for r in tb.rows for c in r.cells]
    assert "인건비" in texts and "10,000,000" in texts    # 데이터 표 보존


def test_table_guide_idempotent_in_runall():
    """run_all 을 두 번 돌려도(보완 루프 모사) 2회차엔 표 안내 삭제가 0(멱등)."""
    d = Document()
    d.add_paragraph("실제 본문 단락입니다")
    t = d.add_table(rows=1, cols=1)
    t.rows[0].cells[0].text = "작성요령 : 예시를 삭제하고 작성하시오"
    r1 = dq.run_all(d)
    r2 = dq.run_all(d)
    assert r1.table_guide_rows_removed >= 1
    assert r2.table_guide_rows_removed == 0
    assert r2.guide_paragraphs_removed == 0


def test_is_guide_text_single_source():
    """_is_guide_text 단일 기준 — 삭제기·채점기가 동일 판정을 공유함을 확인."""
    from auto_write.services.doc_quality_score import _scan_guide
    assert dq._is_guide_text("○ 내용을 기재하시오") is True
    assert dq._is_guide_text("※ 작성요령: 예시") is True
    assert dq._is_guide_text("매출 10억원을 달성하였다") is False
    # 채점기(_scan_guide)도 동일 기준으로 critical 카운트
    d = Document()
    d.add_paragraph("사업 내용을 구체적으로 작성하시기 바랍니다")  # 명령형 안내
    d.add_paragraph("실제 매출 10억원 달성")
    crit, _gen = _scan_guide(d)
    assert crit == 1


# --------------------------------------------------------------------------- STEP2: 빈셀 미입력 (item 5)
def test_score_empty_required_cells_informational():
    """미입력 필수셀 수는 참고용(informational) — item5 detail 에만 표기되고 게이트 점수 불변."""
    d = _plan_doc()
    dq.run_all(d)
    base = score_document(d, doc_type="business_plan", type_confidence=0.9,
                          psst_ratio=1.0, image_suggestions=3, existing_images=0)
    info = score_document(d, doc_type="business_plan", type_confidence=0.9,
                          psst_ratio=1.0, image_suggestions=3, existing_images=0,
                          empty_required_cells=3)
    # 총점·항목 점수 불변(게이트 안정성 P4)
    assert info.total == base.total
    item5 = next(i for i in info.items if i.key == "table_quality")
    assert item5.score == next(i for i in base.items if i.key == "table_quality").score
    # 참고 표기는 detail 에 노출
    assert "미입력 필수셀(참고)=3" in item5.detail


def test_orchestrator_counts_confirm_markers():
    """오케스트레이터가 본문/표의 '[확인필요]' 마커(미입력 필수칸)를 정확히 센다."""
    d = Document()
    d.add_paragraph("정상 본문")
    d.add_paragraph("핵심 수치: [확인필요]")
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "사업비"
    t.cell(0, 1).text = "[확인필요]"
    assert DocumentQualityOrchestrator._count_confirm_markers(d) == 2


# --------------------------------------------------------------------------- STEP3: 단락별 서식 통일 (item 2)
def test_unify_paragraph_formatting_dominant():
    """단락 내 크기가 다른 런을 '지배값(실재 다수값)'으로 통일한다(날조 없음)."""
    from docx.shared import Pt
    d = Document()
    p = d.add_paragraph()
    p.add_run("첫째 ").font.size = Pt(10)
    p.add_run("둘째 ").font.size = Pt(14)   # 소수 → 지배값으로 교정
    p.add_run("셋째").font.size = Pt(10)
    n = dq.unify_paragraph_formatting(d)
    assert n == 1
    assert {r.font.size.pt for r in p.runs} == {10.0}   # 다수값 10pt 로 통일


def test_unify_preserves_emphasis():
    """크기·글꼴만 통일하고 강조(bold)는 보존한다."""
    from docx.shared import Pt
    d = Document()
    p = d.add_paragraph()
    r1 = p.add_run("강조")
    r1.font.size = Pt(10)
    r1.bold = True
    p.add_run("일반").font.size = Pt(14)
    dq.unify_paragraph_formatting(d)
    assert p.runs[0].bold is True                       # 강조 보존
    assert {r.font.size.pt for r in p.runs} == {10.0}


def test_unify_skips_theme_inherited():
    """명시 크기/글꼴이 전혀 없는 단락은 테마 상속 보존 위해 건드리지 않는다."""
    d = Document()
    p = d.add_paragraph()
    p.add_run("명시 서식 없음 하나")
    p.add_run("명시 서식 없음 둘")
    n = dq.unify_paragraph_formatting(d)
    assert n == 0
    for r in p.runs:
        assert r.font.size is None                       # 강제 크기 주입 안 함


def test_unify_idempotent():
    """한 번 통일하면 재실행 시 0(멱등)."""
    from docx.shared import Pt
    d = Document()
    p = d.add_paragraph()
    p.add_run("a").font.size = Pt(10)
    p.add_run("b").font.size = Pt(14)
    n1 = dq.unify_paragraph_formatting(d)
    n2 = dq.unify_paragraph_formatting(d)
    assert n1 == 1
    assert n2 == 0


def test_runall_unify_live_default():
    """run_all 기본 호출(인자 없음)에서 unify 가 live 로 동작한다(죽은코드 아님)."""
    from docx.shared import Pt
    d = Document()
    p = d.add_paragraph()
    p.add_run("실행 본문 ").font.size = Pt(11)
    p.add_run("내용").font.size = Pt(15)
    report = dq.run_all(d)
    assert report.paragraphs_unified >= 1
