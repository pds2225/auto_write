"""test_crossform_cell_iteration.py — 재발 클래스 D(순회 범위 불완전) 회귀 고정.

배경
----
과거 반복 재발한 결함: ``doc.tables``/``doc.paragraphs`` 가 **표 셀 안 단락·중첩표**
(표 셀 '안'에 또 들어간 표)를 포함하지 않아, 소스추출·타깃탐지·전사·체크박스가 표 셀
안 필드를 매번 놓쳤다(2026-06-10 앵커 → 06-18 앵커 우선순위 → 07-02 셀 인라인).
cross_form_autofill 의 표 순회를 단일 헬퍼 ``_iter_all_tables`` (중첩표 재귀)로
단일화해 구조적으로 막았다.

이 파일은 그 불변식을 **빨간불로 고정**한다 — 미래에 어떤 소비자가 다시 최상위
``doc.tables`` 로 좁아지면(중첩표를 놓치면) 아래 테스트가 실패한다.

검증 불변식:
  ① ``_iter_all_tables`` 는 중첩표를 전위 순서로 평탄화한다.
  ② 중첩 없는 문서에선 ``_iter_all_tables(doc) == list(doc.tables)`` (회귀 0 보존).
  ③ 중첩표 셀의 '라벨|빈칸'이 ``find_target_fields`` 에 잡힌다.
  ④ 중첩표 셀 '안'의 인라인 빈칸('라벨 : ____')이 잡힌다.
  ⑤ 소스 A 의 중첩표 값이 ``extract_source_fields`` 로 추출된다.
  ⑥ 중첩표 칸이 end-to-end 로 실제 전사된다(값 보존·날조 0).
  ⑦ 2단계 깊이 중첩표까지 도달한다.
  ⑧ 중첩표 안 체크박스(□옵션) 그룹이 탐지된다.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from auto_write.services.cross_form_autofill import (
    _iter_all_tables,
    autofill_from_source,
    extract_source_fields,
    find_checkbox_targets,
    find_target_fields,
)


# --- 헬퍼: 중첩표 픽스처 --------------------------------------------------------

def _nested_target(path: Path) -> None:
    """타깃 B: 최상위 표 1개, 그 셀 안에 '라벨|빈칸' 중첩표를 둔다.

    최상위 표만 보는 순회는 이 중첩표를 통째로 놓친다(재발 클래스 D).
    """
    doc = Document()
    doc.add_heading("지원신청서", 0)
    outer = doc.add_table(rows=1, cols=1)
    cell = outer.rows[0].cells[0]
    cell.paragraphs[0].text = "신청기업 정보"
    inner = cell.add_table(rows=3, cols=2)
    inner.rows[0].cells[0].text = "기업명"
    inner.rows[0].cells[1].text = ""            # 채울 빈칸
    inner.rows[1].cells[0].text = "대표자"
    inner.rows[1].cells[1].text = ""
    # 중첩표 셀 '안' 인라인 빈칸(한 셀에 라벨:____)
    inner.rows[2].cells[0].text = "연락처 : ______"
    doc.save(str(path))


def _nested_source(path: Path) -> None:
    """소스 A: 중첩표 안에 라벨→값 을 둔다(소스 순회도 중첩을 봐야 한다)."""
    doc = Document()
    doc.add_heading("사업계획서", 0)
    outer = doc.add_table(rows=1, cols=1)
    cell = outer.rows[0].cells[0]
    inner = cell.add_table(rows=3, cols=2)
    inner.rows[0].cells[0].text = "기업명"
    inner.rows[0].cells[1].text = "밸류업(주)"
    inner.rows[1].cells[0].text = "대표자"
    inner.rows[1].cells[1].text = "홍길동"
    inner.rows[2].cells[0].text = "연락처"
    inner.rows[2].cells[1].text = "010-1234-5678"
    doc.save(str(path))


def _read_nested_pair(docx_path: Path, ti_outer: int = 0) -> dict[str, str]:
    """저장된 문서에서 최상위 표 → 그 셀의 첫 중첩표 → {라벨: 값}."""
    doc = Document(str(docx_path))
    inner = doc.tables[ti_outer].rows[0].cells[0].tables[0]
    out: dict[str, str] = {}
    for row in inner.rows:
        cells = row.cells
        if len(cells) >= 2 and cells[0].text.strip():
            out[cells[0].text.strip()] = cells[1].text.strip()
    return out


# --- ① / ② 헬퍼 자체 불변식 ---------------------------------------------------

# 주의: python-docx 는 접근할 때마다 새 Table 프록시를 만든다(lxml proxy 비-동일성).
# 그래서 표 동일성은 프록시 객체가 아니라 **밑단 XML 요소(_element)** 로 비교한다.

def test_iter_all_tables_flattens_nested_preorder(tmp_path: Path) -> None:
    """중첩표를 전위 순서(부모 → 셀 안 중첩 → 다음 부모)로 평탄화한다."""
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)  # 최상위 T0
    outer.rows[0].cells[0].add_table(rows=1, cols=1)  # T0 안의 N0
    doc.add_table(rows=1, cols=1)  # 최상위 T1
    els = [t._element for t in _iter_all_tables(doc)]
    # 최상위 2 + 중첩 1 = 3
    assert len(els) == 3, f"중첩표를 놓침: {len(els)} != 3"
    # 전위: T0, N0, T1 — N0(중첩)가 T1(다음 최상위)보다 먼저
    assert els[0] is doc.tables[0]._element
    assert els[2] is doc.tables[1]._element
    assert els[1] is doc.tables[0].rows[0].cells[0].tables[0]._element


def test_iter_all_tables_equals_doc_tables_when_no_nesting() -> None:
    """중첩이 없으면 밑단 표 요소 순서가 doc.tables 와 완전히 동일(기존 동작 보존·회귀 0)."""
    doc = Document()
    doc.add_table(rows=1, cols=2)
    doc.add_table(rows=2, cols=2)
    assert [t._element for t in _iter_all_tables(doc)] == \
        [t._element for t in doc.tables]


def test_iter_all_tables_reaches_two_levels_deep() -> None:
    """2단계 깊이 중첩표까지 도달한다(⑦)."""
    doc = Document()
    t0 = doc.add_table(rows=1, cols=1)
    t1 = t0.rows[0].cells[0].add_table(rows=1, cols=1)
    t2 = t1.rows[0].cells[0].add_table(rows=1, cols=1)
    els = [t._element for t in _iter_all_tables(doc)]
    assert t2._element in els, "2단계 깊이 중첩표에 도달하지 못함"
    assert len(els) == 3


# --- ③ / ④ 타깃 탐지가 중첩표를 포함 ------------------------------------------

def test_find_target_fields_detects_nested_table_cells(tmp_path: Path) -> None:
    """중첩표 셀의 '라벨|빈칸'이 타깃으로 탐지된다(③)."""
    tgt = tmp_path / "target.docx"
    _nested_target(tgt)
    targets = find_target_fields(tgt)
    norms = {t["normalized"] for t in targets}
    assert "기업명" in norms, "중첩표 '기업명' 빈칸 미탐지(순회 범위 불완전)"
    assert "대표자" in norms, "중첩표 '대표자' 빈칸 미탐지"


def test_find_target_fields_detects_nested_cell_inline_blank(tmp_path: Path) -> None:
    """중첩표 셀 '안' 인라인 빈칸('연락처 : ____')이 cell_paragraph 로 탐지된다(④)."""
    tgt = tmp_path / "target.docx"
    _nested_target(tgt)
    targets = find_target_fields(tgt)
    inline = [t for t in targets if t.get("kind") == "cell_paragraph"
              and t["normalized"] == "연락처"]
    assert inline, "중첩표 셀 인라인 '연락처 : ____' 미탐지"


# --- ⑤ 소스 추출이 중첩표를 포함 ----------------------------------------------

def test_extract_source_reads_nested_table_values(tmp_path: Path) -> None:
    """소스 A 중첩표 안의 라벨→값 이 추출된다(⑤)."""
    src = tmp_path / "source.docx"
    _nested_source(src)
    fields = extract_source_fields(src)
    assert fields.get("기업명") == "밸류업(주)", "중첩표 소스 값 미추출"
    assert fields.get("대표자") == "홍길동"
    assert fields.get("연락처") == "010-1234-5678"


# --- ⑥ end-to-end 전사(중첩표 채움·값 보존·날조 0) ----------------------------

def test_autofill_transcribes_into_nested_table(tmp_path: Path) -> None:
    """중첩표 칸이 실제로 채워지고, 값이 소스와 정확히 같다(⑥, 날조 0)."""
    src = tmp_path / "source.docx"
    tgt = tmp_path / "target.docx"
    out = tmp_path / "out.docx"
    _nested_source(src)
    _nested_target(tgt)

    rep = autofill_from_source(src, tgt, out)
    assert rep.ok is True
    assert rep.transcribed >= 3, f"중첩표 전사 부족: {rep.transcribed}"

    filled = _read_nested_pair(out)
    assert filled["기업명"] == "밸류업(주)"
    assert filled["대표자"] == "홍길동"

    # 중첩표 셀 '안' 인라인 빈칸('연락처 : ____')도 채워진다(값은 라벨 셀에 인라인 기입).
    doc = Document(str(out))
    inner = doc.tables[0].rows[0].cells[0].tables[0]
    inline_text = inner.rows[2].cells[0].text
    assert "010-1234-5678" in inline_text, \
        f"중첩표 셀 인라인 빈칸 미채움: {inline_text!r}"


def test_autofill_nested_does_not_modify_originals(tmp_path: Path) -> None:
    """중첩표 경로에서도 원본 A·B 는 변형되지 않는다(불변)."""
    src = tmp_path / "source.docx"
    tgt = tmp_path / "target.docx"
    out = tmp_path / "out.docx"
    _nested_source(src)
    _nested_target(tgt)
    src_before = src.read_bytes()
    tgt_before = tgt.read_bytes()

    autofill_from_source(src, tgt, out)

    assert src.read_bytes() == src_before, "소스 원본이 변형됨"
    assert tgt.read_bytes() == tgt_before, "타깃 원본이 변형됨"


# --- ⑧ 중첩표 안 체크박스 그룹 탐지 -------------------------------------------

def test_find_checkbox_targets_detects_nested_group(tmp_path: Path) -> None:
    """중첩표 안 '라벨 | □개인 | □법인' 선택칸 그룹이 탐지된다(⑧)."""
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    inner = outer.rows[0].cells[0].add_table(rows=1, cols=3)
    inner.rows[0].cells[0].text = "사업자형태"
    inner.rows[0].cells[1].text = "□ 개인"
    inner.rows[0].cells[2].text = "□ 법인"
    path = tmp_path / "cb.docx"
    doc.save(str(path))

    groups = find_checkbox_targets(path)
    assert any(g["normalized"] == "사업자형태" and len(g["options"]) == 2
               for g in groups), "중첩표 안 체크박스 그룹 미탐지"
