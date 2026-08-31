"""test_hangul_default_output.py — 자동생성·채움 기본 산출은 한글(HWPX).

사용자 원문: 「지 지금 자동으로 문서 만들면 워드로 만들어지냐 한글로 만들어지냐
한글로 만들어 줘야 돼」

DOCX 는 사용자가 명시할 때만. 이진 .hwp 는 COM 없으면 HWPX XML.
L014 생성표 서식은 이 경로에서 호출하지 않는다. L050 PDF 는 여기서 mechanized 하지 않는다.
"""

from __future__ import annotations

import inspect
import zipfile
from pathlib import Path

import pytest
from docx import Document

from auto_write.services.cross_form_output_policy import (
    OutputPolicyError,
    validate_output_plan,
)
from auto_write.services.hangul_default import (
    DEFAULT_AUTO_CREATE_EXT,
    default_auto_create_path,
    default_fill_output,
    default_user_facing_suffix,
    emit_hangul_file,
    is_explicit_docx,
    is_hangul_default_combo,
    write_text_hwpx,
)
from auto_write_hub import main as hub_main
from test_cross_form_output_policy import _plan

APP = Path(__file__).resolve().parents[1]
REPO = APP.parent


def _docx_with_text(path: Path, text: str = "사업 개요 본문") -> Path:
    doc = Document()
    doc.add_paragraph(text)
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "항목"
    table.rows[0].cells[1].text = "값"
    table.rows[1].cells[0].text = "기업명"
    table.rows[1].cells[1].text = "도보네비게이션(주)"
    doc.save(str(path))
    return path


def test_default_suffix_is_hwpx_not_docx() -> None:
    assert DEFAULT_AUTO_CREATE_EXT == ".hwpx"
    assert default_user_facing_suffix() == ".hwpx"
    assert not is_explicit_docx()


def test_explicit_docx_only_when_user_asks() -> None:
    assert is_explicit_docx(requested_format="docx")
    assert is_explicit_docx(output_path="out.docx")
    assert default_user_facing_suffix(requested_format="docx") == ".docx"
    assert default_user_facing_suffix(output_path=Path("a.docx")) == ".docx"
    assert default_user_facing_suffix(requested_format="hwpx") == ".hwpx"
    assert default_user_facing_suffix(requested_format="hwp") == ".hwp"


def test_default_auto_create_and_fill_paths(tmp_path: Path) -> None:
    p = default_auto_create_path("초안", tmp_path, kind="bizplan")
    assert p.suffix == ".hwpx"
    assert p.name == "초안_bizplan.hwpx"
    explicit = default_auto_create_path(
        "초안", tmp_path, kind="bizplan", output_path=tmp_path / "out.docx"
    )
    assert explicit.suffix == ".docx"
    src = tmp_path / "양식.hwpx"
    assert default_fill_output(src).name == "양식_제출.hwpx"
    assert default_fill_output(src, tmp_path / "지정.hwpx").name == "지정.hwpx"


def test_hangul_default_combo() -> None:
    assert is_hangul_default_combo(None, None)
    assert is_hangul_default_combo(["hwpx"], "rhwp-hwpx-fill")
    assert not is_hangul_default_combo(["docx"], "rhwp-hwpx-fill")
    assert not is_hangul_default_combo(["hwpx"], "docx-crossform")


def test_policy_hangul_default_needs_no_confirm() -> None:
    validate_output_plan(_plan(["hwpx"], "rhwp-hwpx-fill", confirmed=False))


def test_policy_docx_still_needs_confirm() -> None:
    with pytest.raises(OutputPolicyError, match="confirm-output-plan"):
        validate_output_plan(_plan(["docx"], "docx-crossform", confirmed=False))


def test_hub_fill_hangul_default_skips_confirm(tmp_path: Path, capsys) -> None:
    code = hub_main(["fill", "--notice-folder", str(tmp_path)])
    err = capsys.readouterr().err
    assert "confirm-output-plan 필수" not in err
    assert code != 0  # 양식 없음 등으로 실패할 수 있음. 승인 거절이 아님.


def test_hub_fill_docx_without_confirm_still_blocked(tmp_path: Path, capsys) -> None:
    code = hub_main(["fill", "--notice-folder", str(tmp_path), "--output", "docx"])
    err = capsys.readouterr().err
    assert code == 2
    assert "confirm-output-plan" in err or "DOCX" in err


def test_write_text_hwpx_mimetype_first(tmp_path: Path) -> None:
    out = tmp_path / "n.hwpx"
    write_text_hwpx(out, ["한글 본문"])
    assert out.is_file()
    with zipfile.ZipFile(out) as z:
        infos = z.infolist()
        assert infos[0].filename == "mimetype"
        assert infos[0].compress_type == zipfile.ZIP_STORED
        assert z.read("mimetype") == b"application/hwp+zip"
        body = z.read("Contents/section0.xml").decode("utf-8")
        assert "한글 본문" in body
        assert "linesegarray" not in body


def test_emit_docx_to_hwpx_xml_wrap(tmp_path: Path) -> None:
    src = _docx_with_text(tmp_path / "in.docx", "자동생성 본문")
    dst = tmp_path / "out.hwpx"
    rep = emit_hangul_file(src, dst)
    assert rep.ok
    assert Path(rep.output).suffix == ".hwpx"
    assert rep.method == "xml_wrap"
    with zipfile.ZipFile(rep.output) as z:
        body = z.read("Contents/section0.xml").decode("utf-8")
    assert "자동생성 본문" in body
    assert "도보네비게이션" in body
    assert src.exists()  # 원본 미수정


def test_emit_refuses_overwrite(tmp_path: Path) -> None:
    src = _docx_with_text(tmp_path / "same.docx")
    with pytest.raises(ValueError, match="덮어쓰기"):
        emit_hangul_file(src, src)


def test_emit_hwp_without_com_falls_back_to_hwpx(tmp_path: Path, monkeypatch) -> None:
    src = _docx_with_text(tmp_path / "in.docx")
    dst = tmp_path / "out.hwp"
    monkeypatch.setattr(
        "core.docx.services.hwp_docx_convert.hancom_com_available", lambda: False
    )
    rep = emit_hangul_file(src, dst)
    assert rep.ok
    assert Path(rep.output).suffix == ".hwpx"
    assert "Windows" in " ".join(rep.notes) or "COM" in " ".join(rep.notes)
    assert not dst.exists()  # 가짜 .hwp 를 만들지 않는다


def test_l014_styling_not_used_in_hangul_default() -> None:
    from auto_write.services import hangul_default as hd

    src = inspect.getsource(hd)
    assert "style_generated_table" not in src
    assert "add_generated_table" not in src


def test_l050_not_claimed_mechanized_here() -> None:
    from auto_write.services import hangul_default as hd

    src = inspect.getsource(hd)
    assert "try_generate_sibling_pdf" not in src


def test_bizplan_default_output_is_hwpx(tmp_path: Path, monkeypatch) -> None:
    from auto_write.config import Settings
    from auto_write.services.bizplan_autopilot import run_bizplan_autopilot

    settings = Settings(
        app_root=tmp_path,
        workspace_root=tmp_path,
        template_root=tmp_path,
        project_root=tmp_path,
        results_root=tmp_path / "results",
        static_root=tmp_path,
        template_view_root=tmp_path,
        host="127.0.0.1",
        port=8765,
        openai_api_key="",
        openai_model="m",
        openai_search_model="m",
        openai_image_model="gpt-image-1",
        anthropic_api_key="",
        anthropic_model="m",
        anthropic_search_model="m",
        gemini_api_key="",
    )
    monkeypatch.setattr("core.docx.services.bizplan_autopilot.get_settings", lambda: settings)
    monkeypatch.setattr("core.docx.services.bizplan_autopilot.ensure_directories", lambda s: None)
    src = _docx_with_text(tmp_path / "초안.docx", "초안 본문")
    report = run_bizplan_autopilot(str(src), None, use_ai=False, write_report=False)
    out = Path(report.output_docx)
    assert out.suffix.lower() == ".hwpx"
    assert out.is_file()
    assert src.suffix == ".docx"  # 원본은 워드 그대로


def test_bizplan_explicit_docx_stays_docx(tmp_path: Path) -> None:
    from auto_write.services.bizplan_autopilot import run_bizplan_autopilot

    src = _docx_with_text(tmp_path / "초안.docx")
    out = tmp_path / "지정.docx"
    report = run_bizplan_autopilot(str(src), str(out), use_ai=False, write_report=False)
    # 수용검사 fail 이면 _DRAFT.docx 가 정상
    assert Path(report.output_docx).suffix.lower() == ".docx"


def test_hwpx_submit_cli_default_suffix_is_hwpx() -> None:
    src = Path("양식.hwpx")
    assert default_fill_output(src).name.endswith(".hwpx")
    assert ".docx" not in default_fill_output(src).name


def test_skill_hooks_include_hangul_request() -> None:
    needle = "한글로 만들어 줘야 돼"
    for rel in (
        ".claude/skills/bizdoc-hub/SKILL.md",
        ".claude/skills/bizplan-orchestrator/SKILL.md",
        ".claude/skills/cross-form-submission/SKILL.md",
    ):
        text = (REPO / rel).read_text(encoding="utf-8")
        fm = text.split("---", 2)[1]
        assert needle in fm, f"{rel} description 훅에 요청 원문 없음"


def test_submission_pipeline_user_facing_is_hwpx(tmp_path: Path) -> None:
    """제출 파이프라인 작업본은 DOCX, 사용자 산출은 .hwpx."""
    from auto_write.models import ProjectInput
    from auto_write.services.evaluation_service import EvaluationService
    from auto_write.services.openai_client import OpenAIService
    from auto_write.services.submission_orchestrator import SubmissionPipeline
    from test_submission_pipeline import _FakeProjectService, _FakeStorage, _profile, _settings

    settings = _settings(tmp_path)
    storage = _FakeStorage(tmp_path)
    oa = OpenAIService(settings)
    ps = _FakeProjectService(storage, _profile(tmp_path), oa)
    storage.save_project_input(
        "p1",
        ProjectInput(template_id="t1", organization_profile={"기업명": "테스트(주)"}, project_meta={}),
    )
    report = SubmissionPipeline(ps, EvaluationService(oa), storage, settings).run(
        "p1", announcement_text="", enable_images=False, enable_notebooklm=False
    )
    hangul = Path(report["final_hangul"])
    assert hangul.suffix == ".hwpx"
    assert hangul.is_file()
    assert Path(report["final"]).suffix == ".hwpx"
    assert Path(report["final_docx"]).suffix == ".docx"
    assert Path(report["final_docx"]).is_file()

