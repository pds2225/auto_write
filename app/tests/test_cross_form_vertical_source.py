"""test_cross_form_vertical_source.py — 완성본 A의 세로형(라벨 위 / 값 아래)
정보 카드에서 값을 '추출'하는 회귀 테스트.

기존 cross-form 은 채우는 양식 B 쪽에서만 세로형 카드를 탐지했다(find_target_fields).
소스 A 쪽 추출(extract_source_fields)은 가로형(짝수 인덱스 라벨|값)·본문만 읽어,
A 가 정부양식 표지의 세로형 정보 카드(성명·생년월일·연락처를 윗줄에 나열하고
아랫줄에 값을 적는 구조)면 그 값들이 **통째로 유실**돼 B 로 전혀 전사되지 못했다.

이 테스트는 세로형 카드 값 추출이 실제로 동작하고(전에는 0개), 로스터/혼합 헤더/
병합 어긋남/예산 행렬에는 오작동하지 않으며(오추출<빈칸·날조0·원본 미수정),
end-to-end 로 A(세로 카드) → B(가로 빈 양식) 전사가 실제로 완성됨을 증명한다.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from auto_write.services.cross_form_autofill import (
    autofill_from_source,
    extract_source_fields,
)


# --- 픽스처 헬퍼 -------------------------------------------------------------

def _make_vertical_source(path: Path, *, header: list[str], values: list[str]) -> None:
    """소스 A: 세로형 카드 — 헤더 행(라벨) + 아래 행(값)."""
    doc = Document()
    doc.add_heading("사업계획서", 0)
    t = doc.add_table(rows=2, cols=len(header))
    for c, lab in enumerate(header):
        t.rows[0].cells[c].text = lab
    for c, val in enumerate(values):
        t.rows[1].cells[c].text = val
    doc.save(str(path))


def _make_blank_horizontal_target(path: Path, labels: list[str]) -> None:
    """타깃 B: 라벨|<빈칸> 2열 가로 양식."""
    doc = Document()
    doc.add_heading("지원신청서", 0)
    t = doc.add_table(rows=len(labels), cols=2)
    for i, lab in enumerate(labels):
        t.rows[i].cells[0].text = lab
        t.rows[i].cells[1].text = ""
    doc.save(str(path))


def _value_for(docx_path: Path, label: str) -> str:
    doc = Document(str(docx_path))
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if cells and cells[0].text.strip() == label:
                return cells[1].text.strip()
    return ""


# --- 추출(단위) --------------------------------------------------------------

def test_vertical_source_values_extracted(tmp_path: Path) -> None:
    """전에는 0개, 이제는 세로형 카드의 아래 행 값이 라벨별로 추출된다."""
    src = tmp_path / "src.docx"
    _make_vertical_source(
        src,
        header=["성명", "생년월일", "연락처", "이메일"],
        values=["홍길동", "1990.05.15", "010-1234-5678", "hong@test.com"],
    )
    fields = extract_source_fields(src)
    assert fields.get("성명") == "홍길동"
    assert fields.get("생년월일") == "1990.05.15"
    assert fields.get("연락처") == "010-1234-5678"
    assert fields.get("이메일") == "hong@test.com"


def test_vertical_source_synonym_header(tmp_path: Path) -> None:
    """세로 헤더가 동의어(대표자)여도 추출된다 → 매칭이 성명 타깃과 이어진다."""
    src = tmp_path / "src.docx"
    _make_vertical_source(
        src,
        header=["기업명", "대표자", "사업자등록번호"],
        values=["밸류업(주)", "홍길동", "123-45-67890"],
    )
    fields = extract_source_fields(src)
    assert fields.get("기업명") == "밸류업(주)"
    assert fields.get("대표자") == "홍길동"          # 성명 클러스터 대표
    assert fields.get("사업자등록번호") == "123-45-67890"


# --- end-to-end: A(세로 카드) → B(가로 빈 양식) 전사 --------------------------

def test_vertical_source_to_horizontal_target_e2e(tmp_path: Path) -> None:
    """핵심 win: 세로 카드 A 값이 실제로 가로 양식 B 에 전사된다(전에는 불가능)."""
    src = tmp_path / "src.docx"
    tgt = tmp_path / "tgt.docx"
    out = tmp_path / "out.docx"
    _make_vertical_source(
        src,
        header=["기업명", "대표자", "연락처", "이메일"],
        values=["밸류업(주)", "홍길동", "010-1234-5678", "hong@test.com"],
    )
    _make_blank_horizontal_target(tgt, ["기업명", "성명", "연락처", "이메일"])

    rep = autofill_from_source(src, tgt, out)

    assert rep.ok
    assert rep.transcribed == 4
    assert _value_for(out, "기업명") == "밸류업(주)"
    assert _value_for(out, "성명") == "홍길동"        # 성명 ← 대표자(동의어)
    assert _value_for(out, "연락처") == "010-1234-5678"
    assert _value_for(out, "이메일") == "hong@test.com"


def test_vertical_source_fabrication_zero(tmp_path: Path) -> None:
    """소스 세로 카드에 값이 없는(빈) 칸은 추출하지 않는다(날조 0)."""
    src = tmp_path / "src.docx"
    _make_vertical_source(
        src,
        header=["성명", "생년월일", "연락처"],
        values=["홍길동", "", "010-1234-5678"],   # 생년월일 값 없음
    )
    fields = extract_source_fields(src)
    assert fields.get("성명") == "홍길동"
    assert "생년월일" not in fields                  # 값 없음 → 미추출(날조 0)
    assert fields.get("연락처") == "010-1234-5678"


# --- 안전 가드 ---------------------------------------------------------------

def test_vertical_source_roster_not_extracted(tmp_path: Path) -> None:
    """헤더 + 여러 값 행(팀원 명부)은 세로 추출하지 않는다 — 첫 레코드를 대표로 오인 금지."""
    src = tmp_path / "src.docx"
    doc = Document()
    t = doc.add_table(rows=3, cols=4)  # 헤더 + 값 행 2개 = 다중 레코드
    for c, lab in enumerate(["성명", "생년월일", "연락처", "이메일"]):
        t.rows[0].cells[c].text = lab
    for c, val in enumerate(["홍길동", "1990.05.15", "010-1111-1111", "a@x.com"]):
        t.rows[1].cells[c].text = val
    for c, val in enumerate(["김철수", "1985.01.01", "010-2222-2222", "b@x.com"]):
        t.rows[2].cells[c].text = val
    doc.save(str(src))

    fields = extract_source_fields(src)
    # 로스터 → 세로 추출 금지: 성명/이메일이 명부 값으로 새지 않는다.
    assert "성명" not in fields
    assert "이메일" not in fields


def test_horizontal_labelvalue_not_misread_as_vertical(tmp_path: Path) -> None:
    """가로 라벨|값 행(col1=값)은 세로 헤더로 오인되지 않는다 — 아래 행을 읽지 않는다."""
    src = tmp_path / "src.docx"
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "기업명"
    t.rows[0].cells[1].text = "밸류업(주)"     # col1 = 값(라벨 아님)
    t.rows[1].cells[0].text = "대표자"
    t.rows[1].cells[1].text = "홍길동"
    doc.save(str(src))

    fields = extract_source_fields(src)
    # 가로 추출 그대로: 기업명→밸류업, 대표자→홍길동
    assert fields.get("기업명") == "밸류업(주)"
    assert fields.get("대표자") == "홍길동"
    # 세로 오독으로 '기업명→대표자'(아래 행) 같은 오염이 없어야 한다.
    assert fields.get("기업명") != "대표자"


def test_budget_matrix_header_not_vertically_extracted(tmp_path: Path) -> None:
    """예산 행렬 헤더(일부만 클러스터 라벨)는 세로 추출 대상이 아니다(오추출 방지)."""
    src = tmp_path / "src.docx"
    doc = Document()
    t = doc.add_table(rows=2, cols=3)
    for c, lab in enumerate(["국고보조금", "자기부담금", "총사업비"]):
        t.rows[0].cells[c].text = lab
    for c, val in enumerate(["30,000,000", "10,000,000", "40,000,000"]):
        t.rows[1].cells[c].text = val
    doc.save(str(src))

    fields = extract_source_fields(src)
    # '전부 인식 라벨' 조건 탈락(국고보조금·자기부담금 미등록) → 총사업비에 숫자 미기입.
    assert fields.get("총사업비") != "40,000,000"


def test_vertical_source_merge_mismatch_skipped(tmp_path: Path) -> None:
    """위·아래 논리셀 수가 병합으로 어긋나면 세로 추출을 건너뛴다(크래시 없음)."""
    src = tmp_path / "src.docx"
    doc = Document()
    t = doc.add_table(rows=2, cols=3)
    for c, lab in enumerate(["성명", "생년월일", "연락처"]):
        t.rows[0].cells[c].text = lab
    for c, val in enumerate(["홍길동", "1990.05.15", "010-1234-5678"]):
        t.rows[1].cells[c].text = val
    # 아래 행 두 칸 병합 → 논리셀 수 3(위) != 2(아래)
    t.rows[1].cells[0].merge(t.rows[1].cells[1])
    doc.save(str(src))

    fields = extract_source_fields(src)      # 크래시 없이
    # 정렬 불가 → 세로 추출 안 함(연락처는 병합에 안 걸린 열이지만 len 불일치로 전체 스킵)
    assert "성명" not in fields


def test_horizontal_row_with_label_word_value_not_misread(tmp_path: Path) -> None:
    """가로 라벨|값 행의 값이 라벨 어휘(예: '대표')여도 세로로 오인되지 않는다.

    [직위|대표] / [성명|홍길동] — '대표'는 '직위'의 값이지만 라벨 클러스터에도 속한다.
    세로로 오인하면 '대표→홍길동' 오추출이 생긴다(리뷰 Finding#1). 아래 행 라벨
    가드가 이를 막아 가로 추출의 정답('성명→홍길동')만 남긴다.
    """
    src = tmp_path / "src.docx"
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "직위"
    t.rows[0].cells[1].text = "대표"       # 직위의 값(라벨 어휘와 겹침)
    t.rows[1].cells[0].text = "성명"
    t.rows[1].cells[1].text = "홍길동"
    doc.save(str(src))

    fields = extract_source_fields(src)
    assert fields.get("성명") == "홍길동"      # 가로 추출 정답 보존
    assert "대표" not in fields                # 세로 오인 오추출 없음
    # '대표'가 성명 클러스터로 새어 성명 매칭을 conflict 로 퇴화시키지 않아야 한다.
    assert fields.get("성명") != "대표"


def test_wide_all_label_row_with_synonym_dup_not_vertical(tmp_path: Path) -> None:
    """4열 전부-라벨 행이라도 동의어 중복(부서·소속)이면 세로 오인하지 않는다.

    [직위|대표|부서|소속] — '대표'는 직위의 값, '소속'은 부서의 값이지만 둘 다 라벨
    어휘다(부서·소속은 같은 클러스터). 세로로 오인하면 가로 정답 '직위→대표'를 아래
    행 junk 로 덮어쓴다(리뷰 MEDIUM 잔여). 동의어 중복 배제 가드가 이를 막는다.
    """
    src = tmp_path / "src.docx"
    doc = Document()
    t = doc.add_table(rows=2, cols=4)
    for c, txt in enumerate(["직위", "대표", "부서", "소속"]):
        t.rows[0].cells[c].text = txt
    for c, txt in enumerate(["사원", "홍길동", "기술팀", "개발본부"]):
        t.rows[1].cells[c].text = txt
    doc.save(str(src))

    fields = extract_source_fields(src)
    assert fields.get("직위") == "대표"       # 가로 추출 정답 보존(직위=위치=대표)
    assert fields.get("직위") != "사원"        # 세로 오인 junk 없음
    assert fields.get("대표") != "홍길동"       # '대표→값' 오추출 없음


def test_vertical_incomplete_roster_not_extracted(tmp_path: Path) -> None:
    """2번째 레코드 일부 칸이 비어도 로스터로 보고 세로 추출하지 않는다(리뷰 Finding#2)."""
    src = tmp_path / "src.docx"
    doc = Document()
    t = doc.add_table(rows=3, cols=4)
    for c, lab in enumerate(["성명", "생년월일", "연락처", "이메일"]):
        t.rows[0].cells[c].text = lab
    for c, val in enumerate(["홍길동", "1990.05.15", "010-1111-1111", "a@x.com"]):
        t.rows[1].cells[c].text = val
    # 2번째 레코드: 이메일 칸만 비어 있음(현실 흔함) → all() 가드였다면 우회됐을 케이스
    for c, val in enumerate(["김철수", "1985.01.01", "010-2222-2222", ""]):
        t.rows[2].cells[c].text = val
    doc.save(str(src))

    fields = extract_source_fields(src)
    assert "성명" not in fields                # 첫 레코드가 대표값으로 새지 않음
    assert fields.get("성명") != "홍길동"


def test_vertical_source_originals_untouched(tmp_path: Path) -> None:
    """추출은 읽기 전용 — 소스 원본 바이트가 그대로다."""
    src = tmp_path / "src.docx"
    _make_vertical_source(
        src,
        header=["성명", "연락처"],
        values=["홍길동", "010-1234-5678"],
    )
    before = src.read_bytes()
    extract_source_fields(src)
    assert src.read_bytes() == before
