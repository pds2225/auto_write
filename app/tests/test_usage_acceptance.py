"""usage_acceptance 수용검사 엔진 테스트.

2026-06-10 실문서 진단에서 사람이 발견한 결함 7종이
자동으로 검출되는지 검증한다(오답노트 회귀).
"""

from __future__ import annotations

import pytest
from docx import Document

from auto_write.services.usage_acceptance import (
    AcceptanceConfig, run_acceptance,
    check_unresolved_markers, check_self_inserted_blocks,
    check_template_placeholders, check_unchecked_choices,
    check_empty_label_fields, check_font_name_mixing,
    check_empty_table_rows, check_recruit_date_conflict,
    check_residual_colored_runs,
    check_paren_choices, check_empty_label_fields_ext,
    check_empty_image_slots, check_page_overflow,
)


def _doc() -> Document:
    return Document()


def test_unresolved_markers_detected():
    d = _doc()
    d.add_paragraph("정부지원사업비는 [확인필요] 원으로 한다")
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 1).text = "[확인필요]"
    r = check_unresolved_markers(d)
    assert r.defects == 2 and r.severity == "fail"


def test_self_inserted_blocks_detected():
    d = _doc()
    d.add_paragraph("📊 [NotebookLM 슬라이드 생성용 프롬프트] · 유형: 타임라인")
    d.add_paragraph("(슬라이드 생성 후 이 블록은 삭제하세요)")
    r = check_self_inserted_blocks(d)
    assert r.defects == 2


def test_template_placeholders_detected():
    d = _doc()
    d.add_paragraph("< 사진(이미지) 또는 설계도 제목 >")
    d.add_paragraph("대표자 OOO 는 다음과 같이")
    r = check_template_placeholders(d)
    assert r.defects == 2


def test_unchecked_choices_detected_and_checked_passes():
    d = _doc()
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "지원 분야(택 1)"
    t.cell(0, 1).text = "□ 제조  □ 지식서비스"
    t.cell(1, 0).text = "지방우대 해당여부"
    t.cell(1, 1).text = "■ 비해당지역"
    r = check_unchecked_choices(d)
    assert r.defects == 1  # 체크된 행은 통과


def test_empty_label_fields_detected():
    d = _doc()
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "명 칭"
    t.cell(0, 1).text = ""
    r = check_empty_label_fields(d)
    assert r.defects == 1


def test_font_name_mixing_detected():
    d = _doc()
    for name in ("맑은 고딕", "나눔명조", "함초롬바탕", "굴림", "함초롬돋움", "HY헤드라인M"):
        run = d.add_paragraph().add_run("텍스트")
        run.font.name = name
    r = check_font_name_mixing(d)
    assert r.defects >= 1  # 허용 4종 초과


def test_empty_table_rows_and_date_conflict():
    d = _doc()
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "내용"
    # 2행은 완전 공란
    assert check_empty_table_rows(d).defects == 1
    d.add_paragraph("AI 인재 채용시기 ’26.7월")
    d.add_paragraph("채용은 ’26.10 협약 후 진행")
    assert check_recruit_date_conflict(d).defects == 1


def test_clean_document_is_submittable(tmp_path):
    d = _doc()
    d.add_paragraph("본 사업은 휴머노이드 안전제어 칩을 개발한다.")
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "명 칭"
    t.cell(0, 1).text = "Angel AI"
    p = tmp_path / "clean.docx"
    d.save(str(p))
    rep = run_acceptance(p)
    assert rep.submittable, [c.as_dict() for c in rep.results if not c.passed]


def test_defective_document_blocked(tmp_path):
    d = _doc()
    d.add_paragraph("사업비 [확인필요] 원")
    d.add_paragraph("(슬라이드 생성 후 이 블록은 삭제하세요)")
    p = tmp_path / "bad.docx"
    d.save(str(p))
    rep = run_acceptance(p)
    assert not rep.submittable
    assert rep.fail_defects >= 2


# --- US-1: 순회 범위 확장(ACC-9) + AcceptanceConfig ---------------------------

def test_header_and_footer_markers_detected(tmp_path):
    d = _doc()
    d.add_paragraph("본문은 정상")
    d.sections[0].header.paragraphs[0].text = "[확인필요] 머리글에 숨은 마커"
    d.sections[0].footer.paragraphs[0].text = "대표자 OOO"
    p = tmp_path / "hf.docx"
    d.save(str(p))
    rep = run_acceptance(p)
    by_id = {r.check_id: r for r in rep.results}
    assert by_id["unresolved_markers"].defects >= 1
    assert by_id["template_placeholders"].defects >= 1
    assert not rep.submittable


_TEXTBOX_PICT_XML = (
    '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:v="urn:schemas-microsoft-com:vml">'
    '<v:shape><v:textbox><w:txbxContent>'
    '<w:p><w:r><w:t>[확인필요] 텍스트박스 안 마커</w:t></w:r></w:p>'
    '</w:txbxContent></v:textbox></v:shape></w:pict>'
)


def test_textbox_marker_detected(tmp_path):
    from docx.oxml import parse_xml
    d = _doc()
    run = d.add_paragraph("그림 영역:").add_run("")
    run._r.append(parse_xml(_TEXTBOX_PICT_XML))
    p = tmp_path / "tb.docx"
    d.save(str(p))
    r = check_unresolved_markers(Document(str(p)))
    assert r.defects == 1


# --- US-3c: warn 선도입(괄호 선택란·라벨 확장·빈 그림칸·분량) -------------------

def test_paren_choices_warn():
    """( ) 괄호형 선택란 미선택은 warn, (V) 채움은 통과 (ACC-10)."""
    d = _doc()
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "지원 분야(택 1)"
    t.cell(0, 1).text = "( ) 제조  ( ) 지식서비스"
    r = check_paren_choices(d)
    assert r.severity == "warn" and r.defects == 1
    t.cell(0, 1).text = "(V) 제조  ( ) 지식서비스"
    assert check_paren_choices(d).defects == 0


def test_v_mark_counts_as_checked():
    """□ 행에 'V' 표기로 체크해도 미선택 오탐하지 않는다(ACC-10 오탐 방향 수정)."""
    d = _doc()
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "지방우대 해당여부"
    t.cell(0, 1).text = "□ 해당  V 비해당"
    assert check_unchecked_choices(d).defects == 0


def test_empty_label_fields_ext_warn():
    """라벨 변형('기 업 명')·확장 라벨(사업자등록번호) 공란은 warn (ACC-11)."""
    d = _doc()
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "기 업 명"
    t.cell(0, 1).text = ""
    t.cell(1, 0).text = "사업자등록번호"
    t.cell(1, 1).text = ""
    r = check_empty_label_fields_ext(d)
    assert r.severity == "warn" and r.defects == 2
    # 기존 fail 검사 담당분(정확일치 라벨)은 ext 에서 제외 — 이중 집계 방지
    d2 = _doc()
    t2 = d2.add_table(rows=1, cols=2)
    t2.cell(0, 0).text = "명 칭"
    t2.cell(0, 1).text = ""
    assert check_empty_label_fields_ext(d2).defects == 0
    assert check_empty_label_fields(d2).defects == 1


def test_empty_image_slots_warn():
    """'사진' 라벨 옆 칸에 텍스트도 그림도 없으면 warn (ACC-12)."""
    d = _doc()
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "대표 사진"
    t.cell(0, 1).text = ""
    t.cell(1, 0).text = "회사 로고"
    t.cell(1, 1).text = "첨부함"
    r = check_empty_image_slots(d)
    assert r.severity == "warn" and r.defects == 1


def test_page_overflow_config_gated():
    """분량 검사는 config 지정 시에만 — 근사 추정 warn (ACC-4)."""
    d = _doc()
    for _ in range(40):
        d.add_paragraph("가" * 100)  # 약 4,000자 ≈ 3p
    assert check_page_overflow(d).defects == 0  # 기본: 비활성
    r = check_page_overflow(d, AcceptanceConfig(max_pages=1))
    assert r.severity == "warn" and r.defects == 1
    assert check_page_overflow(d, AcceptanceConfig(max_pages=10)).defects == 0


# --- US-3a: 색상·스캐폴드·날짜·스타일폰트 + R8 재정의 --------------------------

def test_residual_colored_runs_detected():
    """파란 잔존 텍스트는 fail, 검정·색 미지정은 통과 (ACC-3)."""
    from docx.shared import RGBColor
    d = _doc()
    run = d.add_paragraph().add_run("※ 본 항목은 작성 후 파란 안내를 지우고 제출")
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    r = check_residual_colored_runs(d)
    assert r.severity == "fail" and r.defects == 1

    d2 = _doc()
    rr = d2.add_paragraph().add_run("검정 본문")
    rr.font.color.rgb = RGBColor(0, 0, 0)
    d2.add_paragraph("색 미지정 본문")
    assert check_residual_colored_runs(d2).defects == 0


def test_colored_self_block_not_double_counted():
    """자기삽입 블록의 유색 run 은 색상 검사에서 제외(이중 집계 방지)."""
    from docx.shared import RGBColor
    d = _doc()
    p = d.add_paragraph()
    run = p.add_run("(슬라이드 생성 후 이 블록은 삭제하세요)")
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    assert check_residual_colored_runs(d).defects == 0
    assert check_self_inserted_blocks(d).defects == 1


def test_psst_scaffold_blocks_detected():
    """psst_fill 이 넣는 스캐폴드 문구를 자기삽입 블록으로 검출 (ACC-6)."""
    from auto_write.services import usage_acceptance as ua
    d = _doc()
    d.add_paragraph(ua.SCAFFOLD_HEADING)
    d.add_paragraph(ua.SCAFFOLD_DELETE_NOTICE)
    d.add_paragraph(f"  □ 시장분석: {ua.SCAFFOLD_ITEM_SUFFIX}")
    r = check_self_inserted_blocks(d)
    assert r.defects == 3


def test_scaffold_phrases_are_detectable():
    """쓰기 모듈이 import 하는 표준 문구 전부가 검출 패턴에 걸린다(패턴 공유 보증)."""
    from auto_write.services import usage_acceptance as ua
    from auto_write.services import psst_fill, bizplan_ai_writer
    # 쓰기 모듈이 같은 상수를 쓰는지(동일 객체) + 패턴 매치 여부
    assert psst_fill.SCAFFOLD_HEADING is ua.SCAFFOLD_HEADING
    assert bizplan_ai_writer.AI_SECTION_DELETE_NOTICE is ua.AI_SECTION_DELETE_NOTICE
    for s in (ua.SCAFFOLD_HEADING, ua.SCAFFOLD_DELETE_NOTICE,
              f"□ x: {ua.SCAFFOLD_ITEM_SUFFIX}",
              ua.AI_SECTION_HEADING, ua.AI_SECTION_DELETE_NOTICE):
        assert ua.SELF_BLOCK_RE.search(s), s


def test_recruit_date_conflict_korean_format():
    """'2026년 3월' 한국식 표기와 ’26.7 표기의 모순을 검출 (ACC-13, warn)."""
    d = _doc()
    d.add_paragraph("AI 인재 채용 시기는 2026년 3월 예정")
    d.add_paragraph("채용은 ’26.7월 협약 후 진행")
    assert check_recruit_date_conflict(d).defects == 1


def test_style_inherited_font_mixing_detected():
    """run 직접 지정 없이 스타일로만 6종 혼용해도 검출 (ACC-7)."""
    from docx.enum.style import WD_STYLE_TYPE
    d = _doc()
    for i, name in enumerate(("맑은 고딕", "나눔명조", "함초롬바탕", "굴림", "함초롬돋움", "HY헤드라인M")):
        st = d.styles.add_style(f"TestStyle{i}", WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = name
        p = d.add_paragraph("스타일 상속 폰트 텍스트")
        p.style = st
    r = check_font_name_mixing(d)
    assert r.defects >= 1, r.as_dict()


def test_requirement_status_all_is_static():
    """R8(__all__)은 문서 상태가 아니라 원장의 정적 상태를 따른다 (LEDGER-1)."""
    import self_diagnose as sd
    req = {"check_ids": ["__all__"], "상태": "달성"}
    assert sd._requirement_status(req, {"unresolved_markers"}, False) == "달성"
    assert sd._requirement_status(req, set(), True) == "달성"


# --- US-2: 블라인드 마스킹 방향 역전(ACC-1·ACC-2, R10) -------------------------

def _blind_cfg() -> AcceptanceConfig:
    return AcceptanceConfig(blind_review=True)


def test_blind_review_allows_masking(tmp_path):
    """블라인드 모드: 올바른 ○○○ 마스킹 문서는 제출가능 (ACC-1 오탐 해소)."""
    d = _doc()
    d.add_paragraph("대표자 성명: ○○○ (블라인드 마스킹 완료)")
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "직장명"
    t.cell(0, 1).text = "○○○"
    p = tmp_path / "masked.docx"
    d.save(str(p))
    rep = run_acceptance(p, _blind_cfg())
    by_id = {r.check_id: r for r in rep.results}
    assert by_id["template_placeholders"].defects == 0
    assert by_id["masking_violation"].defects == 0
    assert rep.submittable, [c.as_dict() for c in rep.results if not c.passed]


def test_blind_review_detects_real_names(tmp_path):
    """블라인드 모드: 실명 잔존(진짜 위반)을 fail 로 검출 (ACC-2 미탐 해소)."""
    d = _doc()
    d.add_paragraph("대표자 성명: 홍길동 / 직장명: 삼성전자 / 대학명: 서울대학교")
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "성명"
    t.cell(0, 1).text = "김철수"
    p = tmp_path / "names.docx"
    d.save(str(p))
    rep = run_acceptance(p, _blind_cfg())
    mv = next(r for r in rep.results if r.check_id == "masking_violation")
    assert mv.severity == "fail" and mv.defects >= 4
    assert not rep.submittable


def test_default_masking_behavior_unchanged(tmp_path):
    """기본(비블라인드): ○○○는 여전히 자리표시 fail, 실명 검사는 비활성."""
    d = _doc()
    d.add_paragraph("대표자 성명: ○○○")
    d.add_paragraph("성명: 홍길동")
    p = tmp_path / "default.docx"
    d.save(str(p))
    rep = run_acceptance(p)
    by_id = {r.check_id: r for r in rep.results}
    assert by_id["template_placeholders"].defects >= 1  # ○○○ = 빈 자리표시(현행 유지)
    assert by_id["masking_violation"].defects == 0      # 비활성(오탐 0 원칙)


def test_blind_review_still_catches_other_placeholders(tmp_path):
    """블라인드 모드에서도 ○○○ 외 자리표시(<사진(이미지)>)는 그대로 fail."""
    d = _doc()
    d.add_paragraph("< 사진(이미지) 또는 설계도 제목 >")
    p = tmp_path / "ph.docx"
    d.save(str(p))
    rep = run_acceptance(p, _blind_cfg())
    ph = next(r for r in rep.results if r.check_id == "template_placeholders")
    assert ph.defects >= 1


# --- US-3b: 폰트 ascii/eastAsia 이중집계 오탐 수정(ACC-8) ----------------------

def test_font_pairs_not_double_counted():
    """정상 한·영 페어 3쌍 — 슬롯 분리 집계로 오탐(fail) 없어야 한다."""
    from docx.oxml.ns import qn
    d = _doc()
    pairs = (("Arial", "맑은 고딕"), ("Times New Roman", "바탕"), ("Calibri", "돋움"))
    for ascii_name, ea in pairs:
        run = d.add_paragraph().add_run("혼합 본문 텍스트")
        run.font.name = ascii_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), ea)
    r = check_font_name_mixing(d)
    assert r.defects == 0, r.as_dict()


def test_font_allowed_kinds_configurable():
    """AcceptanceConfig.allowed_fonts 로 허용 종수를 조정할 수 있다."""
    d = _doc()
    for name in ("Arial", "Times New Roman", "Calibri"):
        d.add_paragraph().add_run("텍스트").font.name = name
    assert check_font_name_mixing(d).defects == 0  # 기본 4종 허용
    r = check_font_name_mixing(d, AcceptanceConfig(allowed_fonts=1))
    assert r.defects == 2


def test_acceptance_config_default_is_noop(tmp_path):
    d = _doc()
    d.add_paragraph("본 사업은 정상 문서다.")
    p = tmp_path / "cfg.docx"
    d.save(str(p))
    base = run_acceptance(p)
    with_cfg = run_acceptance(p, AcceptanceConfig())
    assert base.submittable is True and with_cfg.submittable is True
    assert base.fail_defects == with_cfg.fail_defects == 0
    assert [r.check_id for r in base.results] == [r.check_id for r in with_cfg.results]


# --- US-4: 재실행 덮어쓰기 보호(PIPE-2) ----------------------------------------

def test_backup_existing_output(tmp_path):
    from pathlib import Path
    from auto_write.services.usage_acceptance import backup_existing_output
    assert backup_existing_output(tmp_path / "none.docx") == ""
    f = tmp_path / "a.docx"
    f.write_text("이전 산출물", encoding="utf-8")
    bak = backup_existing_output(f)
    assert bak and Path(bak).exists() and not f.exists()
    assert Path(bak).read_text(encoding="utf-8") == "이전 산출물"


def test_force_draft_name_preserves_existing_draft(tmp_path):
    """기존 _DRAFT(사용자 수정본일 수 있음)를 Path.replace 로 파괴하지 않는다."""
    from auto_write.services.usage_acceptance import force_draft_name
    old_draft = tmp_path / "out_DRAFT.docx"
    old_draft.write_text("사용자 수정본", encoding="utf-8")
    cur = tmp_path / "out.docx"
    cur.write_text("새 산출물", encoding="utf-8")
    new_path, err = force_draft_name(cur)
    assert err == "" and new_path.name == "out_DRAFT.docx"
    assert new_path.read_text(encoding="utf-8") == "새 산출물"
    baks = list(tmp_path.glob("out_DRAFT_prev*.docx"))
    assert len(baks) == 1 and baks[0].read_text(encoding="utf-8") == "사용자 수정본"


# --- US-5: self_diagnose 인코딩·종료코드 계약(ENC-1/2) --------------------------

def test_self_diagnose_cli_exit_codes(tmp_path):
    """ENC-1/2: cp949 캡처 환경에서 크래시 없이 완주(exit 2 + JSON 생성), 없는 파일 exit 1."""
    import os
    import subprocess
    import sys as _sys
    from pathlib import Path as _Path

    app_dir = _Path(__file__).resolve().parent.parent
    env = dict(os.environ)
    env.pop("PYTHONUTF8", None)
    env.pop("PYTHONIOENCODING", None)
    env["PYTHONPATH"] = str(app_dir)

    bad = tmp_path / "bad.docx"
    d = _doc()
    d.add_paragraph("사업비 [확인필요] 원 — 검사용(em-dash 포함)")
    d.save(str(bad))
    out_json = tmp_path / "diag.json"
    r = subprocess.run(
        [_sys.executable, str(app_dir / "self_diagnose.py"), str(bad), "--json", str(out_json)],
        capture_output=True, env=env, cwd=str(app_dir),
    )
    assert r.returncode == 2, r.stderr.decode("utf-8", "replace")
    assert out_json.exists()

    r1 = subprocess.run(
        [_sys.executable, str(app_dir / "self_diagnose.py"), str(tmp_path / "없음.docx")],
        capture_output=True, env=env, cwd=str(app_dir),
    )
    assert r1.returncode == 1


def test_self_diagnose_exit3_on_checker_crash(tmp_path, monkeypatch):
    """검사기 예외 = 판정 불가 → exit 3 (문서 결함 2 와 구분, fail-closed)."""
    import self_diagnose as sd

    src = tmp_path / "x.docx"
    _doc().save(str(src))

    def _boom(*_a, **_k):
        raise RuntimeError("checker down")

    monkeypatch.setattr(sd, "run_acceptance", _boom)
    assert sd.main([str(src)]) == 3


# --- US-7: 원장 판정 로직(LEDG-5/6) --------------------------------------------

def test_requirement_status_rules():
    """check_ids 공백=정적 상태 복창 / 매핑 검사 fail=부분·미달성 (LEDG-6)."""
    import self_diagnose as sd

    assert sd._requirement_status({"check_ids": [], "상태": "부분달성"}, set(), True) == "부분달성"
    req = {"check_ids": ["unresolved_markers", "self_inserted_blocks"]}
    assert sd._requirement_status(req, {"unresolved_markers"}, False) == "부분달성"
    assert sd._requirement_status(req, set(), True) == "달성"
    assert sd._requirement_status(req, {"unresolved_markers", "self_inserted_blocks"}, False) == "미달성"


# --- R11: 유색 텍스트 검정 정규화(교정) — ACC-3 역연산 회귀 ----------------------

def test_normalize_colored_text_flips_acc3():
    """파란/회색 유색 본문을 검정으로 정규화하면 ACC-3 결함이 0이 되어야 한다."""
    from docx.shared import RGBColor
    from auto_write.services.doc_quality_ops import normalize_colored_text_to_black

    d = _doc()
    p = d.add_paragraph()
    r = p.add_run("파란 안내문구")
    r.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    g = d.add_paragraph()
    gr = g.add_run("회색 가이드")
    gr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    d.add_paragraph().add_run("검정 본문(상속)")  # 색 미지정 — 보존 대상

    # before: ACC-3 가 2건 검출
    assert check_residual_colored_runs(d).defects == 2

    n = normalize_colored_text_to_black(d)
    assert n == 2, "유색 런 2개를 검정으로 바꿔야 함"

    # after: 0건 + 멱등(두 번째는 0)
    assert check_residual_colored_runs(d).defects == 0
    assert normalize_colored_text_to_black(d) == 0
    # 텍스트 무손실
    assert d.paragraphs[0].runs[0].text == "파란 안내문구"
    assert d.paragraphs[2].runs[0].text == "검정 본문(상속)"


def test_normalize_colored_text_preserves_white():
    """보존색(흰색 계열)은 검정으로 바꾸지 않는다(표 머리글 흰 글씨 등)."""
    from docx.shared import RGBColor
    from auto_write.services.doc_quality_ops import normalize_colored_text_to_black

    d = _doc()
    w = d.add_paragraph()
    wr = w.add_run("흰색 머리글")
    wr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    assert normalize_colored_text_to_black(d) == 0  # 보존색은 건드리지 않음


def test_color_preserve_set_in_sync_with_detection():
    """교정(docx_ops._PRESERVE_COLORS)과 검출(usage_acceptance._COLOR_PRESERVE) 보존색 동기 가드."""
    from auto_write.services.docx_ops import _PRESERVE_COLORS as REMEDIATE
    from auto_write.services.usage_acceptance import _COLOR_PRESERVE as DETECT
    assert REMEDIATE == DETECT, "보존색 집합이 검출/교정 간 어긋나면 안 됨"


def test_run_all_normalizes_colors():
    """run_all 이 유색 텍스트 정규화를 포함해 ACC-3 를 0으로 만든다(파이프라인 배선)."""
    from docx.shared import RGBColor
    from auto_write.services.doc_quality_ops import run_all

    d = _doc()
    pr = d.add_paragraph().add_run("파란 안내문구 본문입니다")
    pr.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    rep = run_all(d)
    assert rep.colored_runs_normalized >= 1
    assert check_residual_colored_runs(d).defects == 0
