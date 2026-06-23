"""test_cross_form_confirm.py — needs_confirm 후보의 사용자 확정/적용 흐름 회귀.

배경(닫지 못했던 루프)
---------------------
cross-form autofill 은 high(정확일치/동의어 단일후보)만 자동 전사하고,
퍼지·충돌 후보는 ``needs_confirm`` 으로 **보고만** 했다. 비개발자 사용자는
"이건 사람이 보세요"라는 후보 목록을 받고도 그것을 실제 문서에 적용할 방법이
없어 손으로 DOCX 를 편집해야 했다(막다른 길).

이 기능(confirmations)
---------------------
사용자가 후보를 검토해 ``{타깃 라벨: 소스 라벨}`` 확정을 넘기면, high 자동전사에
더해 **확정된 칸도 전사**해 루프를 닫는다.

핵심 효과(before→after):
  - before: 퍼지/충돌 후보는 보고만 되고 자동 채움 불가 → 사용자가 직접 편집.
  - after:  ``--confirm "제품명칭=제품명"`` 한 줄로 그 칸이 소스 실값으로 채워진다.

안전 불변(이 테스트가 잠그는 것):
  - 날조 0 — 확정이 가리키는 소스 라벨이 없거나 값이 비면 채우지 않고 보류·기록.
  - 원본 미수정 — out==source/target 거부는 그대로.
  - 보수성 — 확정하지 않은 퍼지/충돌은 여전히 자동 채우지 않는다.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from auto_write.services.cross_form_autofill import autofill_from_source
import cross_form_fill


def _table_doc(path: Path, pairs: list[tuple[str, str]]) -> None:
    """(라벨, 값) 쌍을 2열 표로 저장. 값이 ''면 빈칸."""
    doc = Document()
    t = doc.add_table(rows=len(pairs), cols=2)
    for ri, (label, value) in enumerate(pairs):
        t.rows[ri].cells[0].text = label
        t.rows[ri].cells[1].text = value
    doc.save(str(path))


def _value_for(docx_path: Path, label: str) -> str:
    doc = Document(str(docx_path))
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if cells and cells[0].text.strip() == label:
                return cells[1].text.strip()
    return ""


# --- 퍼지 후보 확정 적용 ------------------------------------------------------

def test_confirm_applies_fuzzy_candidate(tmp_path: Path) -> None:
    """제품명(소스) ↔ 제품명칭(타깃) 은 클러스터 밖 퍼지 → 확정하면 전사된다."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _table_doc(src, [("제품명", "스마트센서")])
    _table_doc(tgt, [("제품명칭", "")])

    # 확정 없으면: 퍼지라 자동 전사 안 됨, needs_confirm 에만 뜸
    base = autofill_from_source(src, tgt, out, use_ai=False)
    assert base.transcribed == 0
    assert any(n["normalized"] == "제품명칭" for n in base.needs_confirm)

    # 확정 적용: 제품명칭 칸을 소스 '제품명' 값으로 채운다
    out2 = tmp_path / "out2.docx"
    rep = autofill_from_source(
        src, tgt, out2, use_ai=False, confirmations={"제품명칭": "제품명"})
    assert _value_for(out2, "제품명칭") == "스마트센서"
    assert rep.transcribed == 1
    assert rep.confirmed == 1
    # 확정된 칸은 더 이상 needs_confirm 에 남지 않는다(해소됨)
    assert all(n["normalized"] != "제품명칭" for n in rep.needs_confirm)


# --- 충돌 후보 확정(여러 동의어 중 하나 선택) --------------------------------

def test_confirm_resolves_conflict_choosing_one(tmp_path: Path) -> None:
    """연락처 클러스터에 소스 후보 2개(전화번호·휴대폰) → conflict.
    사용자가 휴대폰을 고르면 그 값으로 채운다."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _table_doc(src, [("전화번호", "02-111-2222"), ("휴대폰", "010-3333-4444")])
    _table_doc(tgt, [("연락처", "")])

    base = autofill_from_source(src, tgt, out, use_ai=False)
    assert base.transcribed == 0
    nc = [n for n in base.needs_confirm if n["normalized"] == "연락처"]
    assert nc and nc[0]["confidence"] == "conflict"

    out2 = tmp_path / "out2.docx"
    rep = autofill_from_source(
        src, tgt, out2, use_ai=False, confirmations={"연락처": "휴대폰"})
    assert _value_for(out2, "연락처") == "010-3333-4444"
    assert rep.confirmed == 1


# --- 날조 0: 존재하지 않는 소스 라벨 확정은 무시 ------------------------------

def test_confirm_unknown_source_is_skipped(tmp_path: Path) -> None:
    """확정이 가리키는 소스 라벨이 소스에 없으면 채우지 않고 노트로 알린다(날조 0)."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _table_doc(src, [("제품명", "스마트센서")])
    _table_doc(tgt, [("제품명칭", "")])

    rep = autofill_from_source(
        src, tgt, out, use_ai=False, confirmations={"제품명칭": "존재하지않는라벨"})
    assert _value_for(out, "제품명칭") == ""   # 채우지 않음
    assert rep.confirmed == 0
    assert any("존재하지않는라벨" in n or "확인" in n for n in rep.notes)


# --- high 자동전사는 확정과 무관하게 그대로 ----------------------------------

def test_confirm_does_not_disturb_high(tmp_path: Path) -> None:
    """high(정확일치) 칸은 확정 인자가 있어도 그대로 자동 전사된다."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _table_doc(src, [("기업명", "밸류업(주)"), ("제품명", "스마트센서")])
    _table_doc(tgt, [("기업명", ""), ("제품명칭", "")])

    rep = autofill_from_source(
        src, tgt, out, use_ai=False, confirmations={"제품명칭": "제품명"})
    assert _value_for(out, "기업명") == "밸류업(주)"      # high 자동
    assert _value_for(out, "제품명칭") == "스마트센서"    # 확정
    assert rep.transcribed == 2
    assert rep.confirmed == 1


# --- 사용자 입력 정규화(장식·공백 허용) --------------------------------------

def test_confirm_normalizes_user_input(tmp_path: Path) -> None:
    """확정 라벨에 장식(○)·공백이 붙어 있어도 정규화해 매칭한다."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _table_doc(src, [("제품명", "스마트센서")])
    _table_doc(tgt, [("○ 제품명칭", "")])

    rep = autofill_from_source(
        src, tgt, out, use_ai=False,
        confirmations={" ○ 제품명칭 ": "  제품명 "})
    assert _value_for(out, "○ 제품명칭") == "스마트센서"
    assert rep.confirmed == 1


# --- 확정하지 않은 다른 퍼지는 여전히 보류(보수성) ---------------------------

def test_confirm_only_named_target(tmp_path: Path) -> None:
    """확정한 타깃만 채우고, 확정하지 않은 다른 퍼지 후보는 보류한다."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _table_doc(src, [("제품명", "센서"), ("서비스명", "구독형")])
    _table_doc(tgt, [("제품명칭", ""), ("서비스명칭", "")])

    rep = autofill_from_source(
        src, tgt, out, use_ai=False, confirmations={"제품명칭": "제품명"})
    assert _value_for(out, "제품명칭") == "센서"      # 확정됨
    assert _value_for(out, "서비스명칭") == ""        # 보류
    assert rep.confirmed == 1
    assert any(n["normalized"] == "서비스명칭" for n in rep.needs_confirm)


# --- CLI: --confirm 플래그 -----------------------------------------------------

def test_cli_confirm_flag(tmp_path: Path, capsys) -> None:
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _table_doc(src, [("제품명", "스마트센서")])
    _table_doc(tgt, [("제품명칭", "")])

    rc = cross_form_fill.main([
        "--source", str(src), "--target", str(tgt), "-o", str(out),
        "--confirm", "제품명칭=제품명",
    ])
    assert rc == 0
    assert _value_for(out, "제품명칭") == "스마트센서"


def test_cli_confirm_malformed_errors(tmp_path: Path) -> None:
    """= 없는 --confirm 은 raw traceback 대신 JSON 오류 + exit 2."""
    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    _table_doc(src, [("제품명", "스마트센서")])
    _table_doc(tgt, [("제품명칭", "")])

    rc = cross_form_fill.main([
        "--source", str(src), "--target", str(tgt), "-o", str(out),
        "--confirm", "잘못된형식",
    ])
    assert rc == 2


# --- CLI: --confirm-file (JSON) -----------------------------------------------

def test_cli_confirm_file(tmp_path: Path) -> None:
    import json

    src = tmp_path / "a.docx"
    tgt = tmp_path / "b.docx"
    out = tmp_path / "out.docx"
    cfile = tmp_path / "confirm.json"
    _table_doc(src, [("제품명", "스마트센서")])
    _table_doc(tgt, [("제품명칭", "")])
    cfile.write_text(json.dumps({"제품명칭": "제품명"}, ensure_ascii=False),
                     encoding="utf-8")

    rc = cross_form_fill.main([
        "--source", str(src), "--target", str(tgt), "-o", str(out),
        "--confirm-file", str(cfile),
    ])
    assert rc == 0
    assert _value_for(out, "제품명칭") == "스마트센서"
